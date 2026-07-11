from copy import deepcopy
import difflib

from docx import Document
from docx.oxml.ns import qn


def _set_run_font(run, font_name: str):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str):
    style.font.name = font_name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_font_everywhere(doc: Document, font_name: str):
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            _set_style_font(doc.styles[style_name], font_name)
        except Exception:
            pass

    for p in doc.paragraphs:
        for r in p.runs:
            _set_run_font(r, font_name)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_run_font(r, font_name)


def _norm_text(s: str) -> str:
    if not s:
        return ""
    return " ".join(str(s).split()).strip()


def _table_row_sigs_from_block(block):
    raw = getattr(block, "raw", None)
    if raw is not None and getattr(raw, "rows", None) is not None:
        sigs = []
        for row in raw.rows:
            cells = [_norm_text(cell.text) for cell in row.cells]
            sigs.append("\t".join(cells).rstrip())
        return sigs

    text = getattr(block, "text", "") or ""
    lines = [ln.rstrip() for ln in str(text).splitlines() if ln.strip()]
    return lines


def _block_sig(block):
    block_type = getattr(block, "block_type", "") or ""
    if block_type == "table":
        return ("table", "\n".join(_table_row_sigs_from_block(block)))
    return ("para", _norm_text(getattr(block, "text", "") or ""))


def _diff_table_rows(old_block, new_block):
    old_rows = _table_row_sigs_from_block(old_block)
    new_rows = _table_row_sigs_from_block(new_block)
    sm = difflib.SequenceMatcher(a=old_rows, b=new_rows, autojunk=False)

    old_keep = set()
    new_keep = set()
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        if tag in ("replace", "delete"):
            old_keep.update(range(i1, i2))
        if tag in ("replace", "insert"):
            new_keep.update(range(j1, j2))

    # 表格整体被判定为“变更”但行级 diff 找不到时：退化为保留全部
    if not old_keep and old_rows:
        old_keep = set(range(len(old_rows)))
    if not new_keep and new_rows:
        new_keep = set(range(len(new_rows)))

    # 默认把表头带上，便于阅读（仍然只输出变更数据行）
    if old_rows and len(old_keep) < len(old_rows):
        old_keep.add(0)
    if new_rows and len(new_keep) < len(new_rows):
        new_keep.add(0)

    return (
        {"kind": "table_slice", "raw": getattr(old_block, "raw", None), "keep_rows": sorted(old_keep)},
        {"kind": "table_slice", "raw": getattr(new_block, "raw", None), "keep_rows": sorted(new_keep)},
    )


def _diff_segment_items(old_seg, new_seg):
    old_blocks = list(getattr(old_seg, "blocks", []) or [])
    new_blocks = list(getattr(new_seg, "blocks", []) or [])

    a = [_block_sig(b) for b in old_blocks]
    b = [_block_sig(b) for b in new_blocks]

    sm = difflib.SequenceMatcher(a=a, b=b, autojunk=False)
    old_items = []
    new_items = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue

        if tag == "replace":
            if (
                i2 - i1 == 1
                and j2 - j1 == 1
                and getattr(old_blocks[i1], "block_type", None) == "table"
                and getattr(new_blocks[j1], "block_type", None) == "table"
            ):
                old_slice, new_slice = _diff_table_rows(old_blocks[i1], new_blocks[j1])
                old_items.append(old_slice)
                new_items.append(new_slice)
            else:
                old_items.extend(old_blocks[i1:i2])
                new_items.extend(new_blocks[j1:j2])

        elif tag == "delete":
            old_items.extend(old_blocks[i1:i2])

        elif tag == "insert":
            new_items.extend(new_blocks[j1:j2])

    return old_items, new_items


def _insert_text_paragraph_after(doc: Document, anchor, text: str):
    p = doc.add_paragraph(text)
    anchor.addnext(p._p)
    return p._p


def _insert_table_xml_after(anchor, raw_table, keep_rows=None):
    new_tbl = deepcopy(raw_table._tbl)

    if keep_rows is not None:
        keep = set(keep_rows)
        trs = list(new_tbl.xpath("./*[local-name()='tr']"))
        for idx, tr in enumerate(trs):
            if idx not in keep:
                new_tbl.remove(tr)

    anchor.addnext(new_tbl)
    return new_tbl


def _insert_item_after(doc: Document, anchor, item):
    if isinstance(item, dict) and item.get("kind") == "table_slice":
        raw = item.get("raw")
        if raw is None:
            return _insert_text_paragraph_after(doc, anchor, "(表格缺失)")
        return _insert_table_xml_after(anchor, raw, keep_rows=item.get("keep_rows"))

    block = item
    raw = getattr(block, "raw", None)

    if raw is None:
        return _insert_text_paragraph_after(doc, anchor, getattr(block, "text", "") or "")

    if getattr(raw, "_p", None) is not None:
        new_p = deepcopy(raw._p)
        anchor.addnext(new_p)
        return new_p

    if getattr(raw, "_tbl", None) is not None:
        return _insert_table_xml_after(anchor, raw, keep_rows=None)

    return _insert_text_paragraph_after(doc, anchor, getattr(block, "text", "") or "")


def _segment_first_text(seg):
    if seg is None:
        return ""
    for b in getattr(seg, "blocks", []) or []:
        t = (getattr(b, "text", "") or "").strip()
        if t:
            return t
    return ""


def _item_has_text(item) -> bool:
    if isinstance(item, dict) and item.get("kind") == "table_slice":
        raw = item.get("raw")
        keep_rows = item.get("keep_rows") or []
        return bool(raw is not None and keep_rows)
    t = (getattr(item, "text", "") or "").strip()
    return bool(t)


def _items_has_any_text(items) -> bool:
    return any(_item_has_text(x) for x in (items or []))


def _append_delete_content(doc: Document, old_seg, old_items):
    doc.add_paragraph("删除：")
    anchor = doc.paragraphs[-1]._p
    if old_items:
        for item in old_items:
            anchor = _insert_item_after(doc, anchor, item)
    elif old_seg is not None:
        blocks = list(getattr(old_seg, "blocks", []) or [])
        if blocks:
            for b in blocks:
                anchor = _insert_item_after(doc, anchor, b)
        else:
            _insert_text_paragraph_after(doc, anchor, "(无内容)")
    else:
        _insert_text_paragraph_after(doc, anchor, "(无内容)")


def render_change_order(changes, output_path):
    doc = Document()
    doc.add_heading("软件文档更改说明书", 0)

    for i, ch in enumerate(changes, 1):
        old_seg = ch.get("old")
        new_seg = ch.get("new")

        if old_seg is not None and new_seg is not None:
            old_items, new_items = _diff_segment_items(old_seg, new_seg)
        else:
            old_items = list(getattr(old_seg, "blocks", []) or []) if old_seg is not None else []
            new_items = list(getattr(new_seg, "blocks", []) or []) if new_seg is not None else []

        effective_type = ch.get("type", "修改")
        if effective_type == "修改" and _items_has_any_text(old_items) and not _items_has_any_text(new_items):
            effective_type = "删除"

        key = ch.get("key", "")
        seg = ch.get("seg", "")
        key_display = key.split(" > ")[-1] if key else ""
        seg_display = "逻辑流程图" if seg == "_MAIN" else seg
        doc.add_heading(f"（问题{i}，{effective_type}）{key_display} - {seg_display}", level=2)

        if effective_type == "删除":
            _append_delete_content(doc, old_seg, old_items)
            doc.add_paragraph("-" * 40)
            continue

        p_old = doc.add_paragraph("更改前：")
        anchor = p_old._p
        if old_seg is None:
            anchor = _insert_text_paragraph_after(doc, anchor, "(本版本前不存在该小节，为新增内容)")
        elif not old_items:
            anchor = _insert_text_paragraph_after(doc, anchor, "(无)")
        else:
            for item in old_items:
                anchor = _insert_item_after(doc, anchor, item)

        p_new = doc.add_paragraph("更改后：")
        anchor = p_new._p
        if new_seg is None:
            anchor = _insert_text_paragraph_after(doc, anchor, "(该小节已在新版本中删除)")
        elif not new_items:
            anchor = _insert_text_paragraph_after(doc, anchor, "(无)")
        else:
            for item in new_items:
                anchor = _insert_item_after(doc, anchor, item)

        doc.add_paragraph("-" * 40)

    _apply_font_everywhere(doc, "宋体")
    doc.save(output_path)
