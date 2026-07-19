from docx import Document
from docx.oxml.ns import qn


def _set_run_font(run, font_name: str):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str):
    style.font.name = font_name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_font_everywhere(doc: Document, font_name: str):
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            _set_style_font(doc.styles[style_name], font_name)
        except Exception:
            pass

    for p in doc.paragraphs:
        for r in p.runs:
            _set_run_font(r, font_name)


def _add_multiline_block(doc: Document, text: str):
    if not text:
        doc.add_paragraph("(无)")
        return
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)


def render_code_change_order(
    changes,
    output_path: str,
    problem_start: int = 1,
    tickets=None,
):
    from tickets.tickets import Ticket, format_problem_heading

    doc = Document()
    doc.add_heading("软件代码更改说明书", 0)

    current_key = None
    start = max(1, int(problem_start or 1))
    ticket_map = tickets or {}

    for offset, ch in enumerate(changes or [], 0):
        i = int(ch.get("problem_index") or (start + offset))
        key = ch.get("key", "未命名文件")
        seg = ch.get("seg", "全局区域")
        ctype = ch.get("type", "修改")

        if key != current_key:
            doc.add_heading(f"{key} -", level=2)
            current_key = key

        ticket_no = (ch.get("ticket_no") or "").strip()
        ticket_title = (ch.get("ticket_title") or "").strip()
        if not ticket_no and not ticket_title:
            t = ticket_map.get(i)
            if isinstance(t, Ticket):
                ticket_no = t.display_no()
                ticket_title = t.display_title()
            elif isinstance(t, dict):
                ticket_no = str(t.get("ticket_no") or t.get("问题单编号") or "").strip()
                ticket_title = str(t.get("title") or t.get("问题") or "").strip()

        if seg == "全局区域":
            seg_display = "全局区域"
            key_display = ""
        elif seg == "头文件":
            seg_display = "头文件"
            key_display = ""
        else:
            seg_display = f"{seg} 函数中"
            key_display = ""

        # 代码单标题习惯：（问题i，类型，单号）seg
        if ticket_no:
            head = f"（问题{i}，{ctype}，{ticket_no}）{seg_display}"
        else:
            head = f"（问题{i}，{ctype}）{seg_display}"
        # 兼容无 type 的旧样式时仍带单号
        if not ctype:
            head = format_problem_heading(i, "修改", key_display, seg_display, ticket_no=ticket_no)

        doc.add_heading(head, level=3)
        if ticket_title or ticket_no:
            parts = []
            if ticket_no:
                parts.append(f"问题单编号：{ticket_no}")
            if ticket_title:
                parts.append(f"问题：{ticket_title}")
            doc.add_paragraph("；".join(parts))

        doc.add_paragraph("更改前：")
        if ctype == "新增":
            doc.add_paragraph("(本版本前不存在该文件/内容，为新增内容)")
        else:
            _add_multiline_block(doc, ch.get("old_text", ""))

        doc.add_paragraph("更改后：")
        if ctype == "删除":
            doc.add_paragraph("(该文件/内容已在新版本中删除)")
        elif ctype == "新增" and seg not in {"全局区域", "头文件"}:
            doc.add_paragraph("(新增函数，省略函数体)")
        else:
            _add_multiline_block(doc, ch.get("new_text", ""))

        doc.add_paragraph("")

    _apply_font_everywhere(doc, "宋体")
    doc.save(output_path)
