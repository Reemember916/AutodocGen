import re
import difflib
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

#📌 2. 统一抽取文本（段落 + 文本框 + wps:txbx）
def extract_all_texts(p):
    """
    统一抽取文本（普通段落 + 文本框 + wps 文本框）
    """
    texts = []

    # 主体文字
    runs = p.xpath(".//w:t")
    txt = "".join(r.text or "" for r in runs).strip()
    if txt:
        texts.append(txt)

    # 文本框
    tbx_paths = [
        ".//*[local-name()='txbxContent']//*[local-name()='p']",
        ".//*[local-name()='txbx']//*[local-name()='txbxContent']//*[local-name()='p']"
    ]

    for path in tbx_paths:
        for p2 in p.xpath(path):
            runs = p2.xpath(".//w:t")
            txt2 = "".join(r.text or "" for r in runs).strip()
            if txt2:
                texts.append(txt2)

    return texts
#📌 3. 将整个 docx 转换为统一 LogicalParagraph 列表
class LogicalParagraph:
    def __init__(self, text, style, num_id, ilvl, is_heading, heading_level, raw):
        self.text = text
        self.style = style
        self.num_id = num_id
        self.ilvl = ilvl
        self.is_heading = is_heading
        self.heading_level = heading_level
        self.raw = raw

def iter_paragraphs_and_tables(parent):
    """
    同时遍历段落/表格，保持顺序
    """
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent._tbl

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

#📌 4. 遍历段落和表格
def detect_heading(para):
    style = para.style.name.lower() if para.style else ""
    for lv in range(1, 5):
        if style in {f"heading {lv}", f"heading{lv}", f"标题{lv}", f"标题 {lv}"}:
            return True, lv
    raw_style = para.style.name.strip() if para.style and para.style.name else ""
    m_custom = re.match(r"^\d+[_-]([1-4])$", raw_style)
    if m_custom:
        return True, int(m_custom.group(1))
    # outlineLvl（兼容处理：避免直接访问 CT_PPr 动态属性导致 AttributeError）
    ppr = para._p.pPr
    if ppr is not None:
        try:
            nodes = ppr.xpath("./*[local-name()='outlineLvl']")
            if nodes:
                val = nodes[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if val is not None:
                    return True, int(val) + 1
        except Exception:
            pass
    return False, None


def get_numbering(para):
    """
    返回 (numId, ilvl)
    """
    pPr = para._p.pPr
    if pPr is None:
        return None, None
    if pPr.numPr is None:
        return None, None
    numId = pPr.numPr.numId.val if pPr.numPr.numId is not None else None
    ilvl = pPr.numPr.ilvl.val if pPr.numPr.ilvl is not None else None
    return numId, ilvl

#📌 5. 构建整份文档的结构模型
def parse_logical_paragraphs(docx_path):
    doc = Document(docx_path)
    items = []

    for blk in iter_paragraphs_and_tables(doc):
        if isinstance(blk, Paragraph):
            texts = extract_all_texts(blk._p)
            text = texts[0] if texts else ""
            is_h, h_lv = detect_heading(blk)
            num_id, ilvl = get_numbering(blk)

            items.append(
                LogicalParagraph(
                    text=text,
                    style=blk.style.name if blk.style else "",
                    num_id=num_id,
                    ilvl=ilvl,
                    is_heading=is_h,
                    heading_level=h_lv,
                    raw=blk
                )
            )
        else:
            # 表格特殊处理
            items.append(("_TABLE_", blk))

    return items

# 📌 6. 按 H4 + 小节（a~e）分段
SUB_RE = re.compile(r'^\s*([a-e])\s*[).、．)]')

def parse_module_structure(items):
    result = {}
    ctx = ["", "", ""]

    cur_key = None
    pending_key = None
    pending_context = ""
    pending_title = ""
    cur_seg = "_MAIN"

    modules = {}

    def ensure_current_module():
        nonlocal cur_key, pending_key, pending_context, pending_title, cur_seg
        if cur_key or not pending_key:
            return cur_key
        cur_key = pending_key
        modules[cur_key] = {"context": pending_context, "title": pending_title, "segments": {}}
        pending_key = None
        pending_context = ""
        pending_title = ""
        cur_seg = "_MAIN"
        return cur_key

    for obj in items:
        if isinstance(obj, LogicalParagraph):
            lp = obj
            text = lp.text

            # 处理 H1~H3
            if lp.is_heading and lp.heading_level in (1, 2, 3):
                ctx[lp.heading_level - 1] = text
                for i in range(lp.heading_level, 3):
                    ctx[i] = ""
                cur_key = None
                pending_key = None
                pending_context = ""
                pending_title = ""
                cur_seg = "_MAIN"
                if lp.heading_level == 3:
                    pending_context = " > ".join(filter(None, ctx[:2]))
                    pending_title = text or "未命名 H3"
                    pending_key = " > ".join(filter(None, ctx[:2] + [pending_title]))
                continue

            # 处理 H4
            if lp.is_heading and lp.heading_level == 4:
                ctx_str = " > ".join(filter(None, ctx))
                title = text or "未命名 H4"
                cur_key = f"{ctx_str} > {title}"
                modules[cur_key] = {"context": ctx_str, "title": title, "segments": {}}
                pending_key = None
                pending_context = ""
                pending_title = ""
                cur_seg = "_MAIN"
                continue

            # 普通段落
            if ensure_current_module():
                m = SUB_RE.match(text)
                seg = m.group(1) if m else "_MAIN"
                if m:
                    cur_seg = seg
                else:
                    seg = cur_seg

                if seg not in modules[cur_key]["segments"]:
                    modules[cur_key]["segments"][seg] = {"blocks": [], "lines": []}

                modules[cur_key]["segments"][seg]["blocks"].append(lp)
                modules[cur_key]["segments"][seg]["lines"].append(text)

        else:
            # 表格
            if ensure_current_module():
                seg = cur_seg
                if seg not in modules[cur_key]["segments"]:
                    modules[cur_key]["segments"][seg] = {"blocks": [], "lines": []}
                modules[cur_key]["segments"][seg]["blocks"].append(obj)

    return modules

def copy_block(doc, blk):
    """
    blk 可以是 LogicalParagraph 或表格。
    文本框会展开为普通段落。
    """

    if isinstance(blk, LogicalParagraph):
        texts = extract_all_texts(blk.raw._p)
        if not texts:
            doc.add_paragraph("")
        else:
            for t in texts:
                p = doc.add_paragraph(t)
                if blk.style:
                    p.style = blk.style
    else:
        tbl = blk[1]
        doc.element.body.append(deepcopy(tbl._tbl))

def generate_change_order(
    old_path,
    new_path,
    output_path="更改单输出.docx",
    max_blocks_each_side=50,
    use_ellipsis=False,
    equal_ratio_threshold=0.01,
    page_break_between_changes=False
):
    print("解析旧版文档...")
    old_items = parse_logical_paragraphs(old_path)
    old_modules = parse_module_structure(old_items)

    print("解析新版文档...")
    new_items = parse_logical_paragraphs(new_path)
    new_modules = parse_module_structure(new_items)

    print("构建更改单...")

    # 创建输出文档
    doc = Document()

    # ========== 标题 ==========
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("软件文档更改说明书")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    run.bold = True

    # 差异统计
    change_list = []
    change_count = 0

    # 所有 H4 keys
    all_h4_keys = list(dict.fromkeys(list(old_modules.keys()) + list(new_modules.keys())))

    for h4_key in all_h4_keys:
        old_mod = old_modules.get(h4_key)
        new_mod = new_modules.get(h4_key)

        old_segments = old_mod["segments"] if old_mod else {}
        new_segments = new_mod["segments"] if new_mod else {}

        seg_keys = sorted(set(old_segments.keys()) | set(new_segments.keys()))

        ctx = (new_mod or old_mod).get("context", "")
        h4_title = (new_mod or old_mod).get("title", "未命名 H4")

        for seg in seg_keys:
            old_seg = old_segments.get(seg)
            new_seg = new_segments.get(seg)

            old_text = "\n".join(old_seg["lines"]) if old_seg else ""
            new_text = "\n".join(new_seg["lines"]) if new_seg else ""

            # ================= 判定差异 ===================
            if old_seg and new_seg:
                if old_text == new_text:
                    continue
                r = difflib.SequenceMatcher(None, old_text, new_text, autojunk=False).ratio()
                diff_ratio = 1 - r
                if diff_ratio <= equal_ratio_threshold:
                    continue
                change_type = "修改"

            elif old_seg and not new_seg:
                change_type = "删除"
            elif new_seg and not old_seg:
                change_type = "新增"
            else:
                continue

            change_count += 1

            # ------------ 记录差异表 ------------
            pretty_seg = f"{seg}) 小节" if seg != "_MAIN" else "正文"
            change_list.append({
                "id": change_count,
                "type": change_type,
                "h4": h4_title,
                "context": ctx,
                "seg": pretty_seg
            })

            # ========== 输出更改单正文 ==========
            if change_count > 1:
                doc.add_paragraph("")

            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(12)
            hr = head.add_run(f"（问题{change_count}，{change_type}） {ctx} > {h4_title} 中的 {pretty_seg}")
            hr.bold = True
            hr.font.name = "黑体"
            hr._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

            # ----------------- 更改前 -----------------
            p_old = doc.add_paragraph()
            r_old = p_old.add_run("更改前：")
            r_old.bold = True
            r_old.font.name = "黑体"

            if old_seg:
                for blk in old_seg["blocks"]:
                    copy_block(doc, blk)
            else:
                doc.add_paragraph("(本版本前不存在该小节，为新增内容)")

            # ----------------- 更改后 -----------------
            p_new = doc.add_paragraph()
            r_new = p_new.add_run("更改后：")
            r_new.bold = True
            r_new.font.name = "黑体"

            if new_seg:
                for blk in new_seg["blocks"]:
                    copy_block(doc, blk)
            else:
                doc.add_paragraph("(该小节已在新版本中删除)")

            doc.add_paragraph("-" * 40)

            if page_break_between_changes:
                doc.add_page_break()

    # ================================ 差异总表插入到文首 ================================
    table = doc.add_table(rows=1, cols=4)
    hdrs = ["序号", "更改类型", "标题（H4）", "所属章节 / 小节"]

    for i, h in enumerate(hdrs):
        table.rows[0].cells[i].text = h

    for item in change_list:
        row = table.add_row().cells
        row[0].text = str(item["id"])
        row[1].text = item["type"]
        row[2].text = item["h4"]
        row[3].text = f"{item['context']} > {item['h4']} > {item['seg']}"

    # 页面顶端插入标题
    doc.paragraphs[1].insert_paragraph_before("更改总览（自动生成）")

    doc.save(output_path)
    print("更改单已生成：", output_path)

if __name__ == "__main__":
    generate_change_order(
        old_path="作动器控制器控制管理软件设计说明(V1.01).docx",
        new_path="作动器控制器控制管理软件设计说明(V1.02).docx",
        output_path="更改单输出.docx",
        max_blocks_each_side=200,
        use_ellipsis=False,
        equal_ratio_threshold=0.01,
        page_break_between_changes=False
    )
