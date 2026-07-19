"""Phase-1 section matching / segment / noise regression tests."""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from canonical.normalize import _detect_sub_id, build_ast
from diff.block_diff import diff_segments
from diff.collect_changes import (
    _extract_doc_id,
    _normalize_title,
    build_match_report,
    build_section_pairs,
    collect_changes,
)
from model.ast import Block, DocumentAST, Section, Segment


def _section(level, title, key, body: str, seg_id: str = "_MAIN") -> Section:
    sec = Section(level=level, title=title, key=key, segments={})
    sec.segments[seg_id] = Segment(
        seg_id=seg_id,
        blocks=[
            Block(
                text=body,
                block_type="para",
                source="body",
                raw=None,
                path=(key, seg_id, 0),
            )
        ],
    )
    return sec


def _ast(sections) -> DocumentAST:
    return DocumentAST(sections=list(sections))


def _make_doc_with_h4(path: Path, h4_title: str, body: str, extra_paras=None):
    doc = Document()
    doc.add_heading("一级标题", level=1)
    doc.add_heading("二级标题", level=2)
    doc.add_heading("三级标题", level=3)
    doc.add_heading(h4_title, level=4)
    doc.add_paragraph(body)
    for p in extra_paras or []:
        doc.add_paragraph(p)
    doc.save(path)


class DocIdExtractionTests(unittest.TestCase):
    def test_slash_style(self):
        self.assertEqual(
            "D/R_SDD01_001_003",
            _extract_doc_id("IFBITStateUpdate（D/R_SDD01_001_003）"),
        )

    def test_dash_style(self):
        self.assertEqual("SDD-001-003", _extract_doc_id("状态机（SDD-001-003）"))

    def test_underscore_style(self):
        self.assertEqual("REQ_12_3", _extract_doc_id("接口处理（REQ_12_3）"))

    def test_plain_text_no_id(self):
        self.assertIsNone(_extract_doc_id("普通中文标题"))


class SegmentIdTests(unittest.TestCase):
    def test_letter_beyond_e(self):
        self.assertEqual("f", _detect_sub_id("f) 后续步骤"))
        self.assertEqual("z", _detect_sub_id("z、收尾"))

    def test_fullwidth_letter(self):
        self.assertEqual("f", _detect_sub_id("ｆ）全角"))

    def test_numbered(self):
        self.assertEqual("1", _detect_sub_id("1) 第一步"))
        self.assertEqual("12", _detect_sub_id("12）第十二步"))
        self.assertEqual("3", _detect_sub_id("（3）括号编号"))


class WhitespaceNoiseTests(unittest.TestCase):
    def test_fullwidth_space_not_a_change(self):
        old = Segment(
            "_MAIN",
            [Block("hello world", "para", "body", None, ("k", "_MAIN", 0))],
        )
        new = Segment(
            "_MAIN",
            [Block("hello　world", "para", "body", None, ("k", "_MAIN", 0))],
        )
        self.assertIsNone(diff_segments(old, new))


class DuplicateDocIdTests(unittest.TestCase):
    def test_duplicate_ids_disambiguated_by_title(self):
        old = _ast(
            [
                _section(
                    4,
                    "Alpha（SDD-100-001）",
                    "H1 > Alpha（SDD-100-001）",
                    "alpha unique body content here",
                ),
                _section(
                    4,
                    "Beta（SDD-100-001）",
                    "H1 > Beta（SDD-100-001）",
                    "beta unique body content here",
                ),
            ]
        )
        new = _ast(
            [
                _section(
                    4,
                    "AlphaRenamed（SDD-100-001）",
                    "H1 > AlphaRenamed（SDD-100-001）",
                    "alpha unique body content here",
                ),
                _section(
                    4,
                    "BetaRenamed（SDD-100-001）",
                    "H1 > BetaRenamed（SDD-100-001）",
                    "beta unique body content here",
                ),
            ]
        )
        pairs = build_section_pairs(old.sections, new.sections)
        methods = [m for _o, _n, m, _s in pairs]
        self.assertTrue(all(m in {"uid_title", "fuzzy"} for m in methods if m != "none"))
        self.assertEqual(2, sum(1 for o, n, m, _s in pairs if o and n))
        changes = collect_changes(old, new)
        # 正文相同 → 不应出现整章新增/删除
        self.assertFalse(any(c["type"] in ("新增", "删除") and c["seg"] != "章节标题" for c in changes))


class FuzzyThresholdTests(unittest.TestCase):
    def test_high_threshold_prevents_weak_pair(self):
        old = _ast([_section(4, "章甲", "H1 > 章甲", "完全不同的甲侧正文内容甲甲甲")])
        new = _ast([_section(4, "章乙", "H1 > 章乙", "完全不同的乙侧正文内容乙乙乙")])
        strict = build_section_pairs(old.sections, new.sections, fuzzy_min_score=0.95)
        # 严格阈值下应拆成 none（删 + 增）
        self.assertTrue(any(m == "none" for _o, _n, m, _s in strict))
        self.assertEqual(2, sum(1 for _o, _n, m, _s in strict if m == "none"))
        # 极低阈值仍可能配成 fuzzy；至少保证 API 可用
        loose = build_section_pairs(old.sections, new.sections, fuzzy_min_score=0.01)
        self.assertGreaterEqual(len(loose), 1)


class MatchReportTests(unittest.TestCase):
    def test_report_contains_methods_and_unmatched(self):
        old = _ast(
            [
                _section(4, "保留", "H1 > 保留", "keep body"),
                _section(4, "将删", "H1 > 将删", "delete only unique text"),
            ]
        )
        new = _ast([_section(4, "保留", "H1 > 保留", "keep body")])
        report = build_match_report(old, new)
        self.assertIn("method_counts", report)
        self.assertIn("pairs", report)
        self.assertTrue(report["unmatched_old"])
        self.assertEqual(1, len(report["unmatched_old"]))


class FootnoteAndDashIdDocxTests(unittest.TestCase):
    def test_dash_id_and_footnote_via_docx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp = Path(tmp_dir)
            old_path = tmp / "old.docx"
            new_path = tmp / "new.docx"
            _make_doc_with_h4(old_path, "控制逻辑（SDD-010-002）1", "步骤内容甲")
            _make_doc_with_h4(new_path, "控制逻辑（SDD-010-002）2", "步骤内容乙")

            old_ast = build_ast(str(old_path))
            new_ast = build_ast(str(new_path))
            self.assertEqual("SDD-010-002", _extract_doc_id(old_ast.sections[0].title))

            pairs = build_section_pairs(old_ast.sections, new_ast.sections)
            self.assertEqual(1, len(pairs))
            self.assertIn(pairs[0][2], {"uid", "key", "fuzzy"})
            changes = collect_changes(old_ast, new_ast)
            types = {c["type"] for c in changes}
            self.assertIn("修改", types)
            self.assertNotIn("新增", types)
            self.assertNotIn("删除", types)

    def test_segment_f_and_numbered_in_docx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp = Path(tmp_dir)
            path = tmp / "seg.docx"
            _make_doc_with_h4(
                path,
                "流程（REQ_9_1）",
                "引言",
                extra_paras=["a) 第一步", "f) 第六步", "1) 编号一"],
            )
            ast = build_ast(str(path))
            segs = set(ast.sections[0].segments.keys())
            self.assertIn("a", segs)
            self.assertIn("f", segs)
            self.assertIn("1", segs)


class LegacySectionMatchingSmoke(unittest.TestCase):
    """Keep core P0 behaviors after Phase-1 API shape change (4-tuple pairs)."""

    def test_footnote_title_noise(self):
        old = _ast([_section(4, "功能描述1", "模块A > 功能描述1", "旧正文内容足够长以便区分")])
        new = _ast([_section(4, "功能描述2", "模块A > 功能描述2", "新正文内容足够长以便区分")])
        pairs = build_section_pairs(old.sections, new.sections)
        self.assertEqual(1, len(pairs))
        self.assertIsNotNone(pairs[0][0])
        self.assertIsNotNone(pairs[0][1])
        changes = collect_changes(old, new)
        self.assertFalse(any(c["type"] in ("新增", "删除") for c in changes))
        self.assertFalse(any(c.get("seg") == "章节标题" for c in changes))

    def test_normalize_title_helper(self):
        self.assertEqual("功能描述", _normalize_title("功能描述1"))


if __name__ == "__main__":
    unittest.main()
