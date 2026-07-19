"""Uniform table render tests for change-order output."""
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.table import Table

from model.ast import Block, Segment
from render.change_order import render_change_order, _insert_uniform_table, _normalize_rows_matrix


class UniformTableRenderTests(unittest.TestCase):
    def test_normalize_pads_columns(self):
        rows = [["a", "b"], ["c"]]
        norm = _normalize_rows_matrix(rows)
        self.assertEqual([["a", "b"], ["c", ""]], norm)

    def test_render_uses_uniform_tables_not_missing(self):
        # synthetic table change without raw (simulates any source style)
        old_seg = Segment(
            "c",
            [
                Block(
                    text="名称\t标识\t类型\n该检测项\tv_new\tUint16",
                    block_type="table",
                    source="table",
                    raw=None,
                    path=("k", "c", 0),
                )
            ],
        )
        new_seg = Segment(
            "c",
            [
                Block(
                    text="名称\t标识\t类型\n该检测项\tv_new\tUint32",
                    block_type="table",
                    source="table",
                    raw=None,
                    path=("k", "c", 0),
                )
            ],
        )
        changes = [
            {
                "type": "修改",
                "key": "模块 > IFBIT（D/R_SDD01_001_003）",
                "seg": "c",
                "old": old_seg,
                "new": new_seg,
            }
        ]
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "out.docx"
            render_change_order(changes, str(out), metadata={"doc_no": "T-1", "version": "V1"})
            doc = Document(str(out))
            texts = [p.text for p in doc.paragraphs]
            self.assertFalse(any("表格缺失" in t for t in texts))
            self.assertGreaterEqual(len(doc.tables), 2)
            # header style: first row first cell has content
            self.assertEqual(doc.tables[0].rows[0].cells[0].text.strip(), "名称")
            # data row shows Uint16 / Uint32 respectively
            self.assertIn("Uint16", doc.tables[0].rows[1].cells[2].text)
            self.assertIn("Uint32", doc.tables[1].rows[1].cells[2].text)

    def test_uniform_table_insert(self):
        doc = Document()
        p = doc.add_paragraph("anchor")
        _insert_uniform_table(
            doc,
            p._p,
            [["名称", "类型"], ["A", "int"]],
            header_row=True,
        )
        self.assertEqual(1, len(doc.tables))
        self.assertEqual("名称", doc.tables[0].rows[0].cells[0].text.strip())


if __name__ == "__main__":
    unittest.main()
