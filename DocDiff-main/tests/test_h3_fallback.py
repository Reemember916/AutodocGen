import tempfile
import unittest
from pathlib import Path

from docx import Document

from canonical.normalize import build_ast
from diff.collect_changes import collect_changes


def _make_doc(path: Path, body_text: str, include_h4: bool = False):
    doc = Document()
    doc.add_heading("一级标题", level=1)
    doc.add_heading("二级标题", level=2)
    doc.add_heading("三级标题", level=3)
    if include_h4:
        doc.add_heading("四级标题", level=4)
    doc.add_paragraph(body_text)
    doc.save(path)


class H3FallbackTests(unittest.TestCase):
    def test_collects_changes_under_h3_when_no_h4_exists(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            old_path = tmp_path / "old.docx"
            new_path = tmp_path / "new.docx"
            _make_doc(old_path, "旧内容")
            _make_doc(new_path, "新内容")

            old_ast = build_ast(str(old_path))
            new_ast = build_ast(str(new_path))
            changes = collect_changes(old_ast, new_ast)

            self.assertEqual(1, len(old_ast.sections))
            self.assertEqual(3, old_ast.sections[0].level)
            self.assertEqual("一级标题 > 二级标题 > 三级标题", old_ast.sections[0].key)
            self.assertEqual(1, len(changes))
            self.assertEqual("修改", changes[0]["type"])
            self.assertEqual("_MAIN", changes[0]["seg"])
            self.assertEqual("一级标题 > 二级标题 > 三级标题", changes[0]["key"])

    def test_does_not_create_empty_h3_section_when_h4_is_the_leaf(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            doc_path = tmp_path / "single.docx"
            _make_doc(doc_path, "四级正文", include_h4=True)

            ast = build_ast(str(doc_path))

            self.assertEqual(1, len(ast.sections))
            self.assertEqual(4, ast.sections[0].level)
            self.assertEqual("一级标题 > 二级标题 > 三级标题 > 四级标题", ast.sections[0].key)


if __name__ == "__main__":
    unittest.main()
