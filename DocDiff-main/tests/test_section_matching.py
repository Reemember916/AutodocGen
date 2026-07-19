import tempfile
import unittest
from pathlib import Path

from docx import Document

from canonical.normalize import build_ast
from diff.collect_changes import (
    _normalize_key,
    _normalize_title,
    build_section_pairs,
    collect_changes,
)
from model.ast import Block, DocumentAST, Section, Segment


def _section(level, title, key, body: str, seg_id: str = "_MAIN") -> Section:
    sec = Section(level=level, title=title, key=key, segments={})
    sec.segments[seg_id] = Segment(
        seg_id=seg_id,
        blocks=[
            Block(
                text=body,
                block_type="para",
                source="body",
                raw=None,
                path=(key, seg_id, 0),
            )
        ],
    )
    return sec


def _ast(sections) -> DocumentAST:
    return DocumentAST(sections=list(sections))


def _make_doc_with_h4(path: Path, h4_title: str, body: str):
    doc = Document()
    doc.add_heading("一级标题", level=1)
    doc.add_heading("二级标题", level=2)
    doc.add_heading("三级标题", level=3)
    doc.add_heading(h4_title, level=4)
    doc.add_paragraph(body)
    doc.save(path)


class NormalizeHelpersTests(unittest.TestCase):
    def test_strips_trailing_footnote_digits(self):
        self.assertEqual("功能描述", _normalize_title("功能描述1"))
        self.assertEqual("功能描述", _normalize_title("功能描述12"))
        self.assertEqual("IFBITStateUpdate(D/R_SDD01_001_003)", _normalize_title("IFBITStateUpdate（D/R_SDD01_001_003）"))

    def test_normalize_key_path(self):
        self.assertEqual(
            "一级 > 二级 > 功能描述",
            _normalize_key("一级 > 二级 > 功能描述1"),
        )


class SectionMatchingTests(unittest.TestCase):
    def test_footnote_title_noise_pairs_as_modify_not_add_delete(self):
        """标题仅差脚注数字时，应配成同一章的修改，而不是整章删除+新增。"""
        old = _ast(
            [
                _section(
                    4,
                    "功能描述1",
                    "模块A > 功能描述1",
                    "旧正文内容足够长以便区分",
                )
            ]
        )
        new = _ast(
            [
                _section(
                    4,
                    "功能描述2",
                    "模块A > 功能描述2",
                    "新正文内容足够长以便区分",
                )
            ]
        )

        pairs = build_section_pairs(old.sections, new.sections)
        self.assertEqual(1, len(pairs))
        old_sec, new_sec, method, _score = pairs[0]
        self.assertIsNotNone(old_sec)
        self.assertIsNotNone(new_sec)
        self.assertIn(method, {"key", "fuzzy"})

        changes = collect_changes(old, new)
        types = [c["type"] for c in changes]
        self.assertNotIn("删除", types)
        self.assertNotIn("新增", types)
        self.assertTrue(any(c["type"] == "修改" and c["seg"] == "_MAIN" for c in changes))
        # 仅脚注数字差异：不应再单独报「章节标题」
        self.assertFalse(any(c.get("seg") == "章节标题" for c in changes))

    def test_unique_doc_id_pairs_despite_title_rename(self):
        old = _ast(
            [
                _section(
                    4,
                    "OldName（D/R_SDD01_001_003）",
                    "H1 > OldName（D/R_SDD01_001_003）",
                    "alpha body text for section",
                )
            ]
        )
        new = _ast(
            [
                _section(
                    4,
                    "NewName（D/R_SDD01_001_003）",
                    "H1 > NewName（D/R_SDD01_001_003）",
                    "beta body text for section",
                )
            ]
        )

        pairs = build_section_pairs(old.sections, new.sections)
        self.assertEqual(1, len(pairs))
        self.assertEqual("uid", pairs[0][2])
        self.assertIsNotNone(pairs[0][0])
        self.assertIsNotNone(pairs[0][1])

        changes = collect_changes(old, new)
        self.assertTrue(any(c.get("seg") == "章节标题" and c["type"] == "修改" for c in changes))
        self.assertTrue(any(c["seg"] == "_MAIN" and c["type"] == "修改" for c in changes))
        self.assertFalse(any(c["type"] in ("新增", "删除") for c in changes))

    def test_true_delete_not_forced_into_fuzzy_pair(self):
        old = _ast(
            [
                _section(4, "保留章", "H1 > 保留章", "same body for keep"),
                _section(4, "将删除章", "H1 > 将删除章", "unique delete only content xyz"),
            ]
        )
        new = _ast(
            [
                _section(4, "保留章", "H1 > 保留章", "same body for keep"),
            ]
        )

        pairs = build_section_pairs(old.sections, new.sections)
        methods = {
            (getattr(o, "title", None), getattr(n, "title", None), m)
            for o, n, m, _s in pairs
        }
        self.assertIn(("保留章", "保留章", "key"), methods)
        self.assertTrue(any(o is not None and n is None and m == "none" for o, n, m, _s in pairs))

        changes = collect_changes(old, new)
        deletes = [c for c in changes if c["type"] == "删除"]
        self.assertEqual(1, len(deletes))
        self.assertIn("将删除章", deletes[0]["key"])

    def test_fuzzy_pairs_when_parent_path_changes(self):
        """上级标题改了导致 key 路径变，正文仍高度相似 → fuzzy 配成修改。"""
        body_old = "本函数完成状态更新并写入共享内存缓冲区，错误码返回给调用方。"
        body_new = "本函数完成状态更新并写入共享内存缓冲区，错误码返回给上层调用方。"
        old = _ast(
            [
                _section(
                    4,
                    "状态更新",
                    "旧模块 > 子系统A > 状态更新",
                    body_old,
                )
            ]
        )
        new = _ast(
            [
                _section(
                    4,
                    "状态更新",
                    "新模块 > 子系统A > 状态更新",
                    body_new,
                )
            ]
        )

        pairs = build_section_pairs(old.sections, new.sections)
        self.assertEqual(1, len(pairs))
        old_sec, new_sec, method, _score = pairs[0]
        self.assertIsNotNone(old_sec)
        self.assertIsNotNone(new_sec)
        self.assertEqual("fuzzy", method)

        changes = collect_changes(old, new)
        self.assertTrue(any(c["type"] == "修改" for c in changes))
        self.assertFalse(any(c["type"] in ("新增", "删除") for c in changes))

    def test_end_to_end_footnote_titles_via_docx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp = Path(tmp_dir)
            old_path = tmp / "old.docx"
            new_path = tmp / "new.docx"
            _make_doc_with_h4(old_path, "接口说明1", "处理流程步骤一")
            _make_doc_with_h4(new_path, "接口说明2", "处理流程步骤二")

            changes = collect_changes(build_ast(str(old_path)), build_ast(str(new_path)))
            types = {c["type"] for c in changes}
            self.assertIn("修改", types)
            self.assertNotIn("新增", types)
            self.assertNotIn("删除", types)


if __name__ == "__main__":
    unittest.main()
