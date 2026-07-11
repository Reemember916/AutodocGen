from docx import Document
from docx.oxml.ns import qn


def _set_run_font(run, font_name: str):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str):
    style.font.name = font_name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_font_everywhere(doc: Document, font_name: str):
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            _set_style_font(doc.styles[style_name], font_name)
        except Exception:
            pass

    for p in doc.paragraphs:
        for r in p.runs:
            _set_run_font(r, font_name)


def _add_multiline_block(doc: Document, text: str):
    if not text:
        doc.add_paragraph("(无)")
        return
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)


def render_code_change_order(changes, output_path: str):
    doc = Document()
    doc.add_heading("软件代码更改说明书", 0)

    current_key = None

    for i, ch in enumerate(changes, 1):
        key = ch.get("key", "未命名文件")
        seg = ch.get("seg", "全局区域")
        ctype = ch.get("type", "修改")

        if key != current_key:
            doc.add_heading(f"{key} -", level=2)
            current_key = key

        if seg == "全局区域":
            doc.add_heading(f"（问题{i}）全局区域", level=3)
        elif seg == "头文件":
            doc.add_heading(f"（问题{i}）头文件", level=3)
        else:
            doc.add_heading(f"（问题{i}）{seg} 函数中", level=3)

        doc.add_paragraph("更改前：")
        if ctype == "新增":
            doc.add_paragraph("(本版本前不存在该文件/内容，为新增内容)")
        else:
            _add_multiline_block(doc, ch.get("old_text", ""))

        doc.add_paragraph("更改后：")
        if ctype == "删除":
            doc.add_paragraph("(该文件/内容已在新版本中删除)")
        elif ctype == "新增" and seg not in {"全局区域", "头文件"}:
            doc.add_paragraph("(新增函数，省略函数体)")
        else:
            _add_multiline_block(doc, ch.get("new_text", ""))

        doc.add_paragraph("")

    _apply_font_everywhere(doc, "宋体")
    doc.save(output_path)
