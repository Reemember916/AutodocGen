from __future__ import annotations

from copy import deepcopy
import difflib
from datetime import datetime
from typing import Any, Dict, List, Optional, Sequence, Set, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# 表头单元格中常见的“主键列”名称（命中则按该列对齐行，减少行重排误报）
_DEFAULT_KEY_HEADER_NAMES = (
    "字段名",
    "字段",
    "名称",
    "参数名",
    "参数",
    "标识",
    "标识符",
    "id",
    "ID",
    "Id",
    "变量名",
    "信号名",
    "接口名",
    "函数名",
    "编号",
    "序号",
    "name",
    "Name",
    "key",
    "Key",
)


def _set_run_font(run, font_name: str):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str):
    style.font.name = font_name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_font_everywhere(doc: Document, font_name: str):
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            _set_style_font(doc.styles[style_name], font_name)
        except Exception:
            pass

    for p in doc.paragraphs:
        for r in p.runs:
            _set_run_font(r, font_name)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_run_font(r, font_name)


def _norm_text(s: str) -> str:
    if not s:
        return ""
    return " ".join(str(s).split()).strip()


def _table_row_cells_from_block(block) -> List[List[str]]:
    raw = getattr(block, "raw", None)
    if raw is not None and getattr(raw, "rows", None) is not None:
        rows = []
        for row in raw.rows:
            rows.append([_norm_text(cell.text) for cell in row.cells])
        return rows

    text = getattr(block, "text", "") or ""
    rows = []
    for ln in str(text).splitlines():
        if not ln.strip():
            continue
        rows.append([_norm_text(c) for c in ln.split("\t")])
    return rows


def _table_row_sigs_from_block(block) -> List[str]:
    return ["\t".join(cells).rstrip() for cells in _table_row_cells_from_block(block)]


def _detect_key_column(header_cells: Sequence[str], key_header_names: Sequence[str]) -> Optional[int]:
    if not header_cells:
        return None
    names = {_norm_text(n).lower() for n in key_header_names if _norm_text(n)}
    for idx, cell in enumerate(header_cells):
        c = _norm_text(cell)
        if not c:
            continue
        if c.lower() in names or c in key_header_names:
            return idx
    # 宽松：表头包含关键词
    for idx, cell in enumerate(header_cells):
        c = _norm_text(cell).lower()
        for name in names:
            if name and name in c:
                return idx
    return None


def _row_key(cells: Sequence[str], key_col: Optional[int]) -> str:
    if key_col is not None and 0 <= key_col < len(cells):
        k = _norm_text(cells[key_col])
        if k:
            return f"k:{k}"
    return f"s:{chr(9).join(cells).rstrip()}"


def _diff_table_rows_by_key(
    old_rows: List[List[str]],
    new_rows: List[List[str]],
    key_col: int,
) -> Tuple[Set[int], Set[int]]:
    """按主键列对齐：同 key 内容变 → 修改；仅一侧有 → 增/删。返回需保留的行下标集合。"""
    old_keep: Set[int] = set()
    new_keep: Set[int] = set()

    def index_map(rows: List[List[str]]) -> Dict[str, List[int]]:
        m: Dict[str, List[int]] = {}
        for i, cells in enumerate(rows):
            if i == 0:
                continue  # skip header
            key = _row_key(cells, key_col)
            m.setdefault(key, []).append(i)
        return m

    old_map = index_map(old_rows)
    new_map = index_map(new_rows)
    all_keys = list(dict.fromkeys(list(old_map.keys()) + list(new_map.keys())))

    for key in all_keys:
        oi = old_map.get(key, [])
        nj = new_map.get(key, [])
        if oi and not nj:
            old_keep.update(oi)
        elif nj and not oi:
            new_keep.update(nj)
        else:
            # 同 key：比较整行签名
            used_new = set()
            for i in oi:
                old_sig = "\t".join(old_rows[i]).rstrip()
                matched = False
                for j in nj:
                    if j in used_new:
                        continue
                    new_sig = "\t".join(new_rows[j]).rstrip()
                    if old_sig == new_sig:
                        used_new.add(j)
                        matched = True
                        break
                if not matched:
                    old_keep.add(i)
            for j in nj:
                if j not in used_new:
                    # 可能与某旧行同 key 但内容不同
                    new_sig = "\t".join(new_rows[j]).rstrip()
                    pair_old = None
                    for i in oi:
                        if i in old_keep or "\t".join(old_rows[i]).rstrip() == new_sig:
                            continue
                        pair_old = i
                        break
                    if pair_old is not None and pair_old not in old_keep:
                        # 内容不同的同 key：两边都保留
                        old_keep.add(pair_old)
                    new_keep.add(j)

    return old_keep, new_keep


def _diff_table_rows(
    old_block,
    new_block,
    key_header_names: Sequence[str] = _DEFAULT_KEY_HEADER_NAMES,
    use_key_column: bool = True,
):
    old_cells = _table_row_cells_from_block(old_block)
    new_cells = _table_row_cells_from_block(new_block)
    old_rows = ["\t".join(c).rstrip() for c in old_cells]
    new_rows = ["\t".join(c).rstrip() for c in new_cells]

    old_keep: Set[int] = set()
    new_keep: Set[int] = set()
    key_col = None
    if use_key_column and old_cells and new_cells:
        key_col = _detect_key_column(old_cells[0], key_header_names)
        if key_col is None:
            key_col = _detect_key_column(new_cells[0], key_header_names)
        if key_col is not None:
            old_keep, new_keep = _diff_table_rows_by_key(old_cells, new_cells, key_col)

    if not old_keep and not new_keep:
        # 回退：整行序列 diff
        sm = difflib.SequenceMatcher(a=old_rows, b=new_rows, autojunk=False)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == "equal":
                continue
            if tag in ("replace", "delete"):
                old_keep.update(range(i1, i2))
            if tag in ("replace", "insert"):
                new_keep.update(range(j1, j2))

    # 表格整体被判定为“变更”但行级 diff 找不到时：退化为保留全部
    if not old_keep and old_rows:
        old_keep = set(range(len(old_rows)))
    if not new_keep and new_rows:
        new_keep = set(range(len(new_rows)))

    # 默认把表头带上，便于阅读（仍然只输出变更数据行）
    if old_rows and len(old_keep) < len(old_rows):
        old_keep.add(0)
    if new_rows and len(new_keep) < len(new_rows):
        new_keep.add(0)

    old_keep_sorted = sorted(old_keep)
    new_keep_sorted = sorted(new_keep)
    # rows_text：即使 raw 丢失（pickle/跨进程）也能用文本重建表格
    old_rows_text = [list(old_cells[i]) for i in old_keep_sorted if i < len(old_cells)]
    new_rows_text = [list(new_cells[i]) for i in new_keep_sorted if i < len(new_cells)]

    return (
        {
            "kind": "table_slice",
            "raw": getattr(old_block, "raw", None),
            "keep_rows": old_keep_sorted,
            "key_column": key_col,
            "rows_text": old_rows_text,
        },
        {
            "kind": "table_slice",
            "raw": getattr(new_block, "raw", None),
            "keep_rows": new_keep_sorted,
            "key_column": key_col,
            "rows_text": new_rows_text,
        },
    )


def _block_sig(block):
    block_type = getattr(block, "block_type", "") or ""
    if block_type == "table":
        return ("table", "\n".join(_table_row_sigs_from_block(block)))
    return ("para", _norm_text(getattr(block, "text", "") or ""))


def _diff_segment_items(old_seg, new_seg, use_table_key_column: bool = True):
    old_blocks = list(getattr(old_seg, "blocks", []) or [])
    new_blocks = list(getattr(new_seg, "blocks", []) or [])

    a = [_block_sig(b) for b in old_blocks]
    b = [_block_sig(b) for b in new_blocks]

    sm = difflib.SequenceMatcher(a=a, b=b, autojunk=False)
    old_items = []
    new_items = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue

        if tag == "replace":
            if (
                i2 - i1 == 1
                and j2 - j1 == 1
                and getattr(old_blocks[i1], "block_type", None) == "table"
                and getattr(new_blocks[j1], "block_type", None) == "table"
            ):
                old_slice, new_slice = _diff_table_rows(
                    old_blocks[i1],
                    new_blocks[j1],
                    use_key_column=use_table_key_column,
                )
                old_items.append(old_slice)
                new_items.append(new_slice)
            else:
                old_items.extend(old_blocks[i1:i2])
                new_items.extend(new_blocks[j1:j2])

        elif tag == "delete":
            old_items.extend(old_blocks[i1:i2])

        elif tag == "insert":
            new_items.extend(new_blocks[j1:j2])

    return old_items, new_items


def _insert_text_paragraph_after(doc: Document, anchor, text: str):
    p = doc.add_paragraph(text)
    anchor.addnext(p._p)
    return p._p


def _set_cell_shading(cell, fill_hex: str) -> None:
    """fill_hex like 'D9E2F3' (no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd
    for old in tcPr.xpath("./*[local-name()='shd']"):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "666666", sz: str = "4") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.xpath("./*[local-name()='tcBorders']"):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _set_cell_margins(cell, top=40, bottom=40, left=60, right=60) -> None:
    """dxa units (twips/20)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.xpath("./*[local-name()='tcMar']"):
        tcPr.remove(old)
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _write_cell_text(cell, text: str, *, bold: bool = False, font_name: str = "宋体", font_pt: float = 10.5):
    # clear default empty para content and write one clean paragraph
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text if text is not None else "")
    run.bold = bold
    run.font.size = Pt(font_pt)
    run.font.name = font_name
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _normalize_rows_matrix(rows_text) -> List[List[str]]:
    rows = [list(r) for r in (rows_text or []) if r is not None]
    if not rows:
        return []
    cols = max((len(r) for r in rows), default=1) or 1
    norm: List[List[str]] = []
    for r in rows:
        cells = ["" if c is None else str(c) for c in r]
        if len(cells) < cols:
            cells.extend([""] * (cols - len(cells)))
        elif len(cells) > cols:
            cells = cells[:cols]
        norm.append(cells)
    return norm


def _insert_uniform_table(doc: Document, anchor, rows_text, *, header_row: bool = True):
    """
    统一表格渲染：不依赖原文档表格样式/合并单元格。
    一律按「网格 + 表头底色 + 宋体」重建，避免原表复杂结构导致显示异常。
    """
    norm = _normalize_rows_matrix(rows_text)
    if not norm:
        return _insert_text_paragraph_after(doc, anchor, "(表格无内容)")

    cols = len(norm[0])
    tbl = doc.add_table(rows=len(norm), cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    try:
        tbl.autofit = True
    except Exception:
        pass

    for i, row_cells in enumerate(norm):
        is_header = bool(header_row and i == 0)
        for j, val in enumerate(row_cells):
            cell = tbl.rows[i].cells[j]
            _write_cell_text(cell, val, bold=is_header, font_name="宋体", font_pt=10.5)
            _set_cell_borders(cell, color="666666", sz="4")
            _set_cell_margins(cell)
            if is_header:
                _set_cell_shading(cell, "D9E2F3")  # 浅蓝灰表头

    # 限制表格宽度，避免超宽挤版
    try:
        tbl.autofit = False
        total = Cm(15.5)
        for row in tbl.rows:
            for cell in row.cells:
                cell.width = int(total) // cols
    except Exception:
        pass

    tbl_el = tbl._tbl
    tbl_el.getparent().remove(tbl_el)
    anchor.addnext(tbl_el)
    return tbl_el


def _rows_text_from_block(block, keep_rows=None):
    cells = _table_row_cells_from_block(block)
    if keep_rows is None:
        return [list(r) for r in cells]
    keep = set(keep_rows)
    return [list(cells[i]) for i in sorted(keep) if i < len(cells)]


def _resolve_rows_text(item_or_block, keep_rows=None) -> List[List[str]]:
    """优先 rows_text；否则从 block/raw 提取单元格文本。"""
    if isinstance(item_or_block, dict) and item_or_block.get("kind") == "table_slice":
        rows_text = item_or_block.get("rows_text")
        if rows_text:
            return _normalize_rows_matrix(rows_text)
        raw = item_or_block.get("raw")
        if raw is not None:
            # 临时 block 形态
            class _B:
                pass

            b = _B()
            b.raw = raw
            b.text = ""
            b.block_type = "table"
            kr = keep_rows if keep_rows is not None else item_or_block.get("keep_rows")
            return _rows_text_from_block(b, keep_rows=kr)
        return []

    block = item_or_block
    if getattr(block, "block_type", None) == "table" or getattr(block, "raw", None) is not None:
        return _rows_text_from_block(block, keep_rows=keep_rows)
    return []


def _insert_item_after(doc: Document, anchor, item):
    # 表格：一律统一格式重建，不拷贝原文档表格 XML（避免合并单元格/复杂边框异常）
    if isinstance(item, dict) and item.get("kind") == "table_slice":
        rows_text = _resolve_rows_text(item)
        if rows_text:
            return _insert_uniform_table(doc, anchor, rows_text, header_row=True)
        return _insert_text_paragraph_after(doc, anchor, "(表格无内容)")

    block = item
    if getattr(block, "block_type", None) == "table":
        rows_text = _resolve_rows_text(block)
        if rows_text:
            return _insert_uniform_table(doc, anchor, rows_text, header_row=True)
        t = (getattr(block, "text", "") or "").strip()
        return _insert_text_paragraph_after(doc, anchor, t or "(表格无内容)")

    raw = getattr(block, "raw", None)

    if raw is None:
        return _insert_text_paragraph_after(doc, anchor, getattr(block, "text", "") or "")

    if getattr(raw, "_p", None) is not None:
        new_p = deepcopy(raw._p)
        anchor.addnext(new_p)
        return new_p

    # 若误把 table 当 block 且 block_type 未标 table
    if getattr(raw, "_tbl", None) is not None:
        rows_text = _resolve_rows_text(block)
        if rows_text:
            return _insert_uniform_table(doc, anchor, rows_text, header_row=True)

    return _insert_text_paragraph_after(doc, anchor, getattr(block, "text", "") or "")


def _item_has_text(item) -> bool:
    if isinstance(item, dict) and item.get("kind") == "table_slice":
        if item.get("rows_text"):
            return any(any(str(c).strip() for c in row) for row in item["rows_text"])
        raw = item.get("raw")
        keep_rows = item.get("keep_rows") or []
        return bool(raw is not None and keep_rows)
    t = (getattr(item, "text", "") or "").strip()
    return bool(t)


def _items_has_any_text(items) -> bool:
    return any(_item_has_text(x) for x in (items or []))


def _append_delete_content(doc: Document, old_seg, old_items):
    doc.add_paragraph("删除：")
    anchor = doc.paragraphs[-1]._p
    if old_items:
        for item in old_items:
            anchor = _insert_item_after(doc, anchor, item)
    elif old_seg is not None:
        blocks = list(getattr(old_seg, "blocks", []) or [])
        if blocks:
            for b in blocks:
                anchor = _insert_item_after(doc, anchor, b)
        else:
            _insert_text_paragraph_after(doc, anchor, "(无内容)")
    else:
        _insert_text_paragraph_after(doc, anchor, "(无内容)")


def _segment_preview(seg, max_chars: int = 200) -> str:
    if seg is None:
        return ""
    parts = []
    for b in getattr(seg, "blocks", []) or []:
        t = (getattr(b, "text", "") or "").strip()
        if t:
            parts.append(t)
    text = "\n".join(parts).strip()
    if len(text) > max_chars:
        return text[: max_chars - 1] + "…"
    return text


def changes_to_jsonable(changes: List[dict]) -> List[dict]:
    """将 collect_changes 结果转为可 JSON 序列化结构（供台账/对接）。"""
    out = []
    for ch in changes or []:
        old_seg = ch.get("old")
        new_seg = ch.get("new")
        row = {
            "type": ch.get("type"),
            "key": ch.get("key"),
            "seg": ch.get("seg"),
            "match_method": ch.get("match_method"),
            "match_score": ch.get("match_score"),
            "problem_index": ch.get("problem_index"),
            "ticket_no": ch.get("ticket_no") or "",
            "ticket_title": ch.get("ticket_title") or "",
            "ticket_match_method": ch.get("ticket_match_method") or "",
            "ticket_match_score": ch.get("ticket_match_score"),
            "ticket_match_reason": ch.get("ticket_match_reason") or "",
            "old_preview": _segment_preview(old_seg),
            "new_preview": _segment_preview(new_seg),
        }
        out.append(row)
    return out


def _normalize_metadata(metadata: Optional[Dict[str, Any]]) -> Dict[str, str]:
    meta = dict(metadata or {})
    defaults = {
        "title": "软件文档更改说明书",
        "doc_no": "",
        "version": "",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "author": "",
        "old_path": "",
        "new_path": "",
        "security": "",
        "remark": "",
    }
    for k, v in defaults.items():
        if k not in meta or meta[k] is None:
            meta[k] = v
        else:
            meta[k] = str(meta[k]).strip()
    if not meta.get("date"):
        meta["date"] = datetime.now().strftime("%Y-%m-%d")
    return meta


def _add_metadata_block(doc: Document, metadata: Dict[str, str]) -> None:
    """在标题下写入元数据（文号/版本/日期等），便于归档。"""
    lines = []
    if metadata.get("doc_no"):
        lines.append(f"文号：{metadata['doc_no']}")
    if metadata.get("version"):
        lines.append(f"版本：{metadata['version']}")
    if metadata.get("date"):
        lines.append(f"编制日期：{metadata['date']}")
    if metadata.get("author"):
        lines.append(f"编制人：{metadata['author']}")
    if metadata.get("security"):
        lines.append(f"密级：{metadata['security']}")
    if metadata.get("old_path"):
        lines.append(f"旧版文档：{metadata['old_path']}")
    if metadata.get("new_path"):
        lines.append(f"新版文档：{metadata['new_path']}")
    if metadata.get("remark"):
        lines.append(f"备注：{metadata['remark']}")

    if not lines:
        # 至少写编制日期，避免空白封面
        lines.append(f"编制日期：{metadata.get('date') or datetime.now().strftime('%Y-%m-%d')}")

    for line in lines:
        doc.add_paragraph(line)
    doc.add_paragraph("")  # spacer


def render_change_order(
    changes,
    output_path,
    metadata: Optional[Dict[str, Any]] = None,
    use_table_key_column: bool = True,
    problem_start: int = 1,
    tickets: Optional[Dict[int, Any]] = None,
):
    """
    渲染软件文档更改说明书。

    metadata 可选字段：title, doc_no, version, date, author, old_path, new_path, security, remark
    problem_start: 问题编号起始（默认 1，便于续号）
    tickets: {序号: Ticket|dict}，序号与「问题N」对齐；可含 ticket_no / title
    """
    from tickets.tickets import Ticket, format_problem_heading

    meta = _normalize_metadata(metadata)
    doc = Document()
    title = meta.get("title") or "软件文档更改说明书"
    doc.add_heading(title, 0)
    _add_metadata_block(doc, meta)

    start = max(1, int(problem_start or 1))
    ticket_map = tickets or {}

    for offset, ch in enumerate(changes or []):
        i = int(ch.get("problem_index") or (start + offset))
        old_seg = ch.get("old")
        new_seg = ch.get("new")

        if old_seg is not None and new_seg is not None:
            old_items, new_items = _diff_segment_items(
                old_seg, new_seg, use_table_key_column=use_table_key_column
            )
        else:
            old_items = list(getattr(old_seg, "blocks", []) or []) if old_seg is not None else []
            new_items = list(getattr(new_seg, "blocks", []) or []) if new_seg is not None else []

        effective_type = ch.get("type", "修改")
        if effective_type == "修改" and _items_has_any_text(old_items) and not _items_has_any_text(new_items):
            effective_type = "删除"

        key = ch.get("key", "")
        seg = ch.get("seg", "")
        key_display = key.split(" > ")[-1] if key else ""
        seg_display = "逻辑流程图" if seg == "_MAIN" else seg

        ticket_no = (ch.get("ticket_no") or "").strip()
        ticket_title = (ch.get("ticket_title") or "").strip()
        if not ticket_no and not ticket_title:
            t = ticket_map.get(i)
            if isinstance(t, Ticket):
                ticket_no = t.display_no()
                ticket_title = t.display_title()
            elif isinstance(t, dict):
                ticket_no = str(t.get("ticket_no") or t.get("问题单编号") or "").strip()
                ticket_title = str(t.get("title") or t.get("问题") or "").strip()

        heading = format_problem_heading(
            i, effective_type, key_display, seg_display, ticket_no=ticket_no
        )
        doc.add_heading(heading, level=2)
        if ticket_title or ticket_no:
            parts = []
            if ticket_no:
                parts.append(f"问题单编号：{ticket_no}")
            if ticket_title:
                parts.append(f"问题：{ticket_title}")
            doc.add_paragraph("；".join(parts))

        if effective_type == "删除":
            _append_delete_content(doc, old_seg, old_items)
            doc.add_paragraph("-" * 40)
            continue

        p_old = doc.add_paragraph("更改前：")
        anchor = p_old._p
        if old_seg is None:
            anchor = _insert_text_paragraph_after(doc, anchor, "(本版本前不存在该小节，为新增内容)")
        elif not old_items:
            anchor = _insert_text_paragraph_after(doc, anchor, "(无)")
        else:
            for item in old_items:
                anchor = _insert_item_after(doc, anchor, item)

        p_new = doc.add_paragraph("更改后：")
        anchor = p_new._p
        if new_seg is None:
            anchor = _insert_text_paragraph_after(doc, anchor, "(该小节已在新版本中删除)")
        elif not new_items:
            anchor = _insert_text_paragraph_after(doc, anchor, "(无)")
        else:
            for item in new_items:
                anchor = _insert_item_after(doc, anchor, item)

        doc.add_paragraph("-" * 40)

    _apply_font_everywhere(doc, "宋体")
    doc.save(output_path)
