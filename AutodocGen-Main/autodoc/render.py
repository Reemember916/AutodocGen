"""Word rendering helpers for AutoDocGen."""

from __future__ import annotations

from dataclasses import replace
import os
import re
import string
import sys
import tempfile
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from ._legacy_support import legacy_backend
from . import utils
from .models import DesignModel, FunctionDesign


_UNIT_TABLE_HEADERS = ("序号", "软件单元名称", "函数原型", "唯一标识", "存放位置", "开发状态", "用途")
_UNIT_TABLE_HEADERS_LEGACY = ("序号", "软件单元名称", "唯一标识", "存放位置", "开发状态", "用途")
_CSU_TABLE_HEADERS = ("CSC 名称", "CSC 标识", "CSU 名称", "CSU 标识")
DEFAULT_REQ_ID_PREFIX = "D/R_SDD01"
_ID_BASE_RE = r"(?<![A-Za-z0-9_/])(?:[A-Za-z]+/[A-Za-z]+(?:_[A-Za-z0-9]+)+|[A-Za-z][A-Za-z0-9]*(?:_[A-Za-z0-9]+)*)"
_MOD_ID_RE = re.compile(rf"({_ID_BASE_RE}(?:_\d{{3}}){{0,4}})")
_CSU_ID_RE = re.compile(rf"({_ID_BASE_RE}(?:_\d{{3}}){{1,5}})")


def normalize_req_prefix(prefix: str) -> str:
    return (prefix or "").rstrip("_")


def _normalize_docx_output_path(path: str, *, ensure_parent_dir: bool) -> str:
    raw = (path or "").strip()
    if not raw:
        raise ValueError("输出路径为空")

    normalized = os.path.expanduser(raw)
    if sys.platform == "darwin" and (not os.path.isabs(normalized)) and normalized.startswith("Users/"):
        if not os.path.exists(os.path.join(os.getcwd(), "Users")):
            normalized = "/" + normalized

    base = os.path.basename(normalized)
    if not base:
        raise ValueError("输出路径无文件名，请指定 .docx 文件名")
    if os.path.isdir(normalized):
        raise ValueError("输出路径是目录，请指定 .docx 文件名")
    lower = base.lower()
    if lower.endswith("docx") and (not lower.endswith(".docx")) and ("." not in base):
        normalized = normalized[:-4] + ".docx"
    elif not lower.endswith(".docx"):
        normalized = normalized + ".docx"

    normalized = os.path.abspath(normalized)
    parent = os.path.dirname(normalized) or "."
    if ensure_parent_dir and parent and (not os.path.exists(parent)):
        os.makedirs(parent, exist_ok=True)
    return normalized


def _set_rfont(run_or_style, east_asia: str, *, ascii_font: str | None = None, hansi_font: str | None = None) -> None:
    rfonts = run_or_style._element.rPr.rFonts
    rfonts.set(qn("w:eastAsia"), east_asia)
    if ascii_font:
        rfonts.set(qn("w:ascii"), ascii_font)
    if hansi_font:
        rfonts.set(qn("w:hAnsi"), hansi_font)


def init_document(cfg):
    tpl = (getattr(cfg, "template_path", "") or "").strip()
    doc = None
    if tpl and os.path.exists(tpl):
        try:
            doc = Document(tpl)
            has_609 = all((f"609_{i}" in doc.styles) for i in (1, 2, 3, 4))
            if has_609:
                return doc
            utils.vlog(cfg, f"使用模板：{tpl}")
        except Exception as exc:
            utils.vlog(cfg, f"加载模板失败，改用默认样式：{exc}")

    if doc is None:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "宋体"
        _set_rfont(style, "宋体")
        style.font.size = Pt(12)
        normal_par = style.paragraph_format
        normal_par.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        normal_par.line_spacing = Pt(20)
        normal_par.space_before = Pt(0)
        normal_par.space_after = Pt(0)

        h4 = doc.styles["Heading 4"] if "Heading 4" in doc.styles else doc.styles.add_style("Heading 4", 1)
        h4.font.name = "宋体"
        _set_rfont(h4, "宋体")
        h4.font.size = Pt(12)
        h4.font.bold = True
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        h4._element.rPr.append(color)
        h4_par = h4.paragraph_format
        h4_par.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        h4_par.line_spacing = Pt(20)
        h4_par.space_before = Pt(0)
        h4_par.space_after = Pt(0)

    def ensure_heading(name: str, base: str | None = None) -> None:
        styles = doc.styles
        try:
            styles[name]
            return
        except KeyError:
            pass
        try:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            if base and base in styles:
                style.base_style = styles[base]
        except Exception:
            pass

    ensure_heading("Heading 1")
    ensure_heading("Heading 2", base="Heading 1")
    ensure_heading("Heading 3", base="Heading 2")
    ensure_heading("Heading 4", base="Heading 3")

    def set_heading_font(name: str) -> None:
        try:
            style = doc.styles[name]
        except KeyError:
            return
        style.font.name = "宋体"
        _set_rfont(style, "宋体")
        style.font.color.rgb = None
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        style._element.rPr.append(color)

    for heading_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        set_heading_font(heading_name)

    def ensure_heading_numbering() -> None:
        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            return
        numbering = numbering_part.element
        nsmap = numbering.nsmap

        def next_id(tag: str) -> int:
            ids: list[int] = []
            for element in numbering.findall(qn(f"w:{tag}"), nsmap):
                value = element.get(qn(f"w:{tag}Id"))
                if value and value.isdigit():
                    ids.append(int(value))
            return (max(ids) + 1) if ids else 1

        abs_id = next_id("abstractNum")
        num_id = next_id("num")

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)

        for level, text in enumerate(("%1.", "%1.%2.", "%1.%2.%3.", "%1.%2.%3.%4.")):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))

            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)

            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "decimal")
            lvl.append(num_fmt)

            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), text)
            lvl.append(lvl_text)

            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), "Segoe UI Symbol")
            rfonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
            rfonts.set(qn("w:cs"), "Segoe UI Symbol")
            rfonts.set(qn("w:eastAsia"), "宋体")
            rpr.append(rfonts)
            lvl.append(rpr)

            ppr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(360 * (level + 1)))
            ind.set(qn("w:hanging"), "360")
            ppr.append(ind)
            lvl.append(ppr)
            abstract.append(lvl)

        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abs_id))
        num.append(abstract_ref)
        numbering.append(num)

        for style_name, ilvl in (("Heading 1", 0), ("Heading 2", 1), ("Heading 3", 2), ("Heading 4", 3)):
            try:
                style = doc.styles[style_name]
            except KeyError:
                continue
            ppr = style.element.get_or_add_pPr()
            numpr = ppr.get_or_add_numPr()
            ilvl_el = numpr.get_or_add_ilvl()
            ilvl_el.val = ilvl
            numid_el = numpr.get_or_add_numId()
            numid_el.val = num_id

    ensure_heading_numbering()

    return doc


def safe_save_docx(doc, output: str) -> None:
    output = _normalize_docx_output_path(output, ensure_parent_dir=True)
    out_dir = os.path.dirname(output) or "."
    prefix = f".{os.path.splitext(os.path.basename(output))[0]}."
    fd, tmp_path = tempfile.mkstemp(prefix=prefix, suffix=".docx", dir=out_dir)
    os.close(fd)
    try:
        doc.save(tmp_path)
        os.replace(tmp_path, output)
    except Exception:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        raise


def add_seq_field(paragraph, seq_name: str) -> None:
    part = getattr(paragraph, "part", None)
    counters = getattr(part, "_autodoc_seq_counters", None) if part is not None else None
    if not isinstance(counters, dict):
        counters = {}
        if part is not None:
            try:
                setattr(part, "_autodoc_seq_counters", counters)
            except Exception:
                pass
    key = str(seq_name or "seq")
    counters[key] = int(counters.get(key, 0) or 0) + 1
    paragraph.add_run(str(counters[key]))


def pick_heading_style(doc, level: int) -> str:
    style_name = f"609_{level}"
    try:
        doc.styles[style_name]
        return style_name
    except KeyError:
        return f"Heading {level}"


def apply_table_style(table, doc, name: str = "Table Grid") -> None:
    try:
        table.style = name
    except Exception:
        try:
            table.style = doc.styles[name]
        except Exception:
            pass


def prevent_table_row_splitting(table) -> None:
    for row in getattr(table, "rows", ()) or ():
        try:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
        except Exception:
            continue


def find_content_placeholder(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "[[CONTENT]]" in text or "{{CONTENT}}" in text:
            return paragraph
    return None


def style_as_caption(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(20)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "宋体"
        _set_rfont(run, "宋体")
        run.font.size = Pt(12)
        run.bold = True


def style_as_normal_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(20)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "宋体"
        _set_rfont(run, "宋体")
        run.font.size = Pt(12)


def add_section_label(doc, text: str):
    paragraph = doc.add_paragraph(text)
    style_as_normal_paragraph(paragraph)
    return paragraph


def add_alpha_section_label(doc, title: str, idx: int, indent_pt: int = 24):
    letter = string.ascii_lowercase[idx % 26]
    paragraph = doc.add_paragraph(f"{letter}) {title}")
    paragraph.paragraph_format.left_indent = Pt(indent_pt)
    style_as_normal_paragraph(paragraph)
    return paragraph


def add_indented_text(doc, text: str, indent_pt: int = 24):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Pt(indent_pt)
    style_as_normal_paragraph(paragraph)
    return paragraph


def _logic_line_indent_level(line: str) -> tuple[int, str]:
    raw = str(line or "").rstrip()
    leading = len(raw) - len(raw.lstrip(" "))
    return max(0, leading // 4), raw.lstrip()


def add_logic_text(doc, text: str, base_indent_pt: int = 48, level_indent_pt: int = 18):
    level, clean = _logic_line_indent_level(text)
    return add_indented_text(doc, clean, indent_pt=base_indent_pt + level * level_indent_pt)


def add_ai_assist_banner(doc, cfg, *, backend_module=None) -> bool:
    backend = backend_module or legacy_backend()
    if not bool(getattr(cfg, "ai_assist", False)):
        return False
    head = "【AI 辅助已启用】AI 模型：%s；阈值：func=%.2f, symbol=%.2f%s" % (
        getattr(cfg, "ai_model", ""),
        getattr(cfg, "ai_conf_func", 0.0),
        getattr(cfg, "ai_conf_symbol", 0.0),
        "；（忽略置信度）" if getattr(cfg, "force_ai", False) else "",
    )
    if backend_module is not None:
        backend.add_section_label(doc, head)
    else:
        add_section_label(doc, head)
    return True


def init_generation_document(cfg, *, main_heading: str = "", heading_level: int = 1, backend_module=None) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    doc = backend.init_document(cfg) if backend_module is not None else init_document(cfg)
    placeholder = backend.find_content_placeholder(doc) if backend_module is not None else find_content_placeholder(doc)
    body_start_idx = len(doc._body._element)
    add_ai_assist_banner(doc, cfg, backend_module=backend)
    if main_heading:
        heading_style = backend.pick_heading_style(doc, heading_level) if backend_module is not None else pick_heading_style(doc, heading_level)
        doc.add_paragraph(main_heading, style=heading_style)
    return {
        "doc": doc,
        "placeholder": placeholder,
        "body_start_idx": body_start_idx,
    }


def add_module_section(doc, module_name: str, module_id: str, entries, *, backend_module=None):
    backend = backend_module or legacy_backend()
    # Guard: 截断过长或包含多个子句的模块标题（AI 长描述泄漏）
    sanitized = str(module_name or "").strip()
    if len(sanitized) > 30 or sanitized.count("，") >= 1 or sanitized.count("；") >= 1:
        sanitized = sanitized.split("，")[0].split("；")[0].strip()
    if not sanitized:
        sanitized = str(module_id or "").strip() or "(未命名模块)"
    heading_style = backend.pick_heading_style(doc, 3) if backend_module is not None else pick_heading_style(doc, 3)
    p = doc.add_paragraph(style=heading_style)
    p.add_run(f"{sanitized}（{module_id}）")
    return add_module_function_table(
        doc,
        sanitized,
        module_id,
        entries,
        backend_module=backend,
    )


def add_module_function_table(doc, module_name: str, module_id: str, entries, *, backend_module=None):
    backend = backend_module or legacy_backend()
    module_name = (module_name or "").strip() or "(未命名模块)"
    module_id = (module_id or "").strip()

    if backend_module is not None:
        backend.add_section_label(doc, "本模块存放的函数见下表。")
    else:
        add_section_label(doc, "本模块存放的函数见下表。")

    caption = doc.add_paragraph()
    caption.add_run("表 ")
    if backend_module is not None:
        backend.add_seq_field(caption, "表")
    else:
        add_seq_field(caption, "表")
    caption.add_run(f" {module_name}模块")
    if backend_module is not None:
        backend.style_as_caption(caption)
    else:
        style_as_caption(caption)

    count = max(1, len(entries or []))
    table = doc.add_table(rows=1 + count, cols=4)
    if backend_module is not None:
        backend.apply_table_style(table, doc)
    else:
        apply_table_style(table, doc)

    headers = ["CSC 名称", "CSC 标识", "CSU 名称", "CSU 标识"]
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    if not entries:
        row = table.rows[1].cells
        row[0].text = module_name
        row[1].text = module_id
        row[2].text = "无"
        row[3].text = ""
        prevent_table_row_splitting(table)
        return table

    for index, entry in enumerate(entries, start=1):
        row = table.rows[index].cells
        row[2].text = str((entry or {}).get("csu_name") or "")
        row[3].text = str((entry or {}).get("csu_id") or "")

    table.rows[1].cells[0].text = module_name
    table.rows[1].cells[1].text = module_id
    if count > 1:
        table.cell(1, 0).merge(table.cell(count, 0))
        table.cell(1, 1).merge(table.cell(count, 1))
    prevent_table_row_splitting(table)
    return table


def iter_doc_blocks(doc, *, backend_module=None):
    body = doc._body._element
    for element in list(body):
        tag = getattr(element, "tag", "")
        if tag.endswith("}p"):
            yield ("p", DocxParagraph(element, doc))
        elif tag.endswith("}tbl"):
            yield ("tbl", DocxTable(element, doc))


def _paragraph_style_names(paragraph) -> list[str]:
    names: list[str] = []
    try:
        style = paragraph.style
    except Exception:
        return names
    seen: set[int] = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        try:
            name = (style.name or "").strip()
        except Exception:
            name = ""
        if name:
            names.append(name)
        try:
            style = style.base_style
        except Exception:
            break
    return names


def is_heading(paragraph, level: int) -> bool:
    names = _paragraph_style_names(paragraph)
    if not names:
        return False
    for normalized in names:
        normalized = (normalized or "").strip()
        if not normalized:
            continue
        simple = re.sub(r"[\s_]+", "", normalized)
        level_str = str(level)
        if f"Heading {level}" in normalized:
            return True
        if f"Heading{level_str}" in simple:
            return True
        if f"标题 {level}" in normalized:
            return True
        if f"标题{level_str}" in simple:
            return True
        if normalized == f"609_{level}":
            return True
    normalized = (names[0] or "").strip()
    if not normalized:
        return False
    simple = re.sub(r"[\s_]+", "", normalized)
    return str(level) in simple and ("heading" in simple.lower() or "标题" in simple)


def get_heading_level(paragraph, max_level: int = 6) -> int | None:
    for level in range(1, max_level + 1):
        if is_heading(paragraph, level):
            return level
    return None


def extract_module_id(text: str, *, backend_module=None) -> str | None:
    matches = _MOD_ID_RE.findall(text or "")
    return max(matches, key=len) if matches else None


def extract_csu_id(text: str, *, backend_module=None) -> str | None:
    matches = _CSU_ID_RE.findall(text or "")
    return max(matches, key=len) if matches else None


def _replace_last_match(text: str, pattern, repl: str) -> str:
    if not text:
        return text
    last = None
    for match in pattern.finditer(text):
        last = match
    if not last:
        return text
    return f"{text[:last.start()]}{repl}{text[last.end():]}"


def collect_module_ids_in_doc(doc, *, backend_module=None) -> list[str]:
    backend = backend_module or legacy_backend()
    ids: list[str] = []
    try:
        for kind, obj in iter_doc_blocks(doc, backend_module=backend):
            if kind != "p":
                continue
            if not is_heading(obj, 3):
                continue
            module_id = extract_module_id(getattr(obj, "text", "") or "", backend_module=backend)
            if module_id:
                ids.append(module_id)
    except Exception:
        return ids
    return ids


def replace_csu_id_in_text(text: str, new_csu_id: str, *, backend_module=None) -> str:
    if not text:
        return text
    if _CSU_ID_RE.search(text):
        return _replace_last_match(text, _CSU_ID_RE, new_csu_id)
    if _MOD_ID_RE.search(text):
        return _replace_last_match(text, _MOD_ID_RE, new_csu_id)
    return text


def is_csu_table(table, *, backend_module=None) -> bool:
    try:
        if not table.rows:
            return False
        cells = table.rows[0].cells
        header = tuple((cell.text or "").strip() for cell in cells[:4])
        return header == _CSU_TABLE_HEADERS
    except Exception:
        return False


def replace_module_id_in_text(text: str, new_module_id: str, *, backend_module=None) -> str:
    if not text:
        return text
    if _MOD_ID_RE.search(text):
        return _replace_last_match(text, _MOD_ID_RE, new_module_id)
    return f"{text}（{new_module_id}）"


def update_module_function_table(table, module_name: str, module_id: str, csu_entries, *, backend_module=None) -> bool:
    backend = backend_module or legacy_backend()
    if table is None or (not is_csu_table(table, backend_module=backend)):
        return False

    module_name = (module_name or "").strip() or "(未命名模块)"
    module_id = (module_id or "").strip()
    count = len(csu_entries or [])

    try:
        while len(table.rows) < 1 + max(1, count):
            table.add_row()
    except Exception:
        pass

    content_rows = max(0, len(table.rows) - 1)
    upto = min(count, content_rows)
    for row_index in range(1, content_rows + 1):
        cells = table.rows[row_index].cells
        if row_index == 1:
            try:
                cells[0].text = module_name
                cells[1].text = module_id
            except Exception:
                pass
        if row_index <= upto:
            csu_name = str((csu_entries[row_index - 1] or {}).get("csu_name") or "")
            csu_id = f"{module_id}_{row_index:03d}" if module_id else ""
            cells[2].text = csu_name
            cells[3].text = csu_id
        else:
            cells[2].text = ""
            cells[3].text = ""

    try:
        if len(table.rows) > 2:
            last = len(table.rows) - 1
            c10 = table.cell(1, 0)
            c1l0 = table.cell(last, 0)
            if getattr(c10, "_tc", None) is not getattr(c1l0, "_tc", None):
                c10.merge(c1l0)
            c11 = table.cell(1, 1)
            c1l1 = table.cell(last, 1)
            if getattr(c11, "_tc", None) is not getattr(c1l1, "_tc", None):
                c11.merge(c1l1)
    except Exception:
        pass
    return True


def update_all_module_tables_and_headings(doc, modules, *, backend_module=None) -> dict[str, int]:
    backend = backend_module or legacy_backend()
    if not modules:
        return {"modules": 0, "headings": 0, "tables": 0}

    blocks = list(iter_doc_blocks(doc, backend_module=backend))
    module_index = 0
    total_headings = 0
    total_tables = 0
    block_index = 0
    while block_index < len(blocks) and module_index < len(modules):
        kind, obj = blocks[block_index]
        if kind != "p":
            block_index += 1
            continue
        paragraph = obj
        if not is_heading(paragraph, 3):
            block_index += 1
            continue

        info = modules[module_index] or {}
        module_id = str(info.get("module_id") or "")
        module_name = str(info.get("module_name") or "")
        csu_entries = info.get("csu_entries") or []

        new_text = replace_module_id_in_text(paragraph.text, module_id, backend_module=backend)
        if new_text != paragraph.text:
            paragraph.text = new_text

        table = None
        probe_index = block_index + 1
        while probe_index < len(blocks):
            probe_kind, probe_obj = blocks[probe_index]
            if probe_kind == "p" and is_heading(probe_obj, 3):
                break
            if probe_kind == "tbl" and is_csu_table(probe_obj, backend_module=backend):
                table = probe_obj
                break
            probe_index += 1
        if table is not None and update_module_function_table(
            table,
            module_name,
            module_id,
            csu_entries,
            backend_module=backend,
        ):
            total_tables += 1

        updated = 0
        heading_index = block_index + 1
        while heading_index < len(blocks):
            probe_kind, probe_obj = blocks[heading_index]
            if probe_kind == "p" and is_heading(probe_obj, 3):
                break
            if probe_kind != "p":
                heading_index += 1
                continue
            func_heading = probe_obj
            if not is_heading(func_heading, 4):
                heading_index += 1
                continue
            updated += 1
            if updated <= len(csu_entries):
                new_id = f"{module_id}_{updated:03d}"
                new_heading_text = replace_csu_id_in_text(func_heading.text, new_id, backend_module=backend)
                if new_heading_text != func_heading.text:
                    func_heading.text = new_heading_text
            heading_index += 1

        total_headings += min(updated, len(csu_entries))
        module_index += 1
        block_index = heading_index

    return {"modules": module_index, "headings": total_headings, "tables": total_tables}


def update_module_headings_only(doc, module_id: str, csu_count: int, *, backend_module=None) -> dict[str, int]:
    backend = backend_module or legacy_backend()
    module_id = (module_id or "").strip()
    if not module_id:
        return {"modules": 0, "headings": 0}

    total_headings = 0
    blocks = list(iter_doc_blocks(doc, backend_module=backend))
    for index, (kind, obj) in enumerate(blocks):
        if kind != "p":
            continue
        paragraph = obj
        if not is_heading(paragraph, 3):
            continue
        if extract_module_id(paragraph.text or "", backend_module=backend) != module_id:
            continue

        new_text = replace_module_id_in_text(paragraph.text or "", module_id, backend_module=backend)
        if new_text != (paragraph.text or ""):
            paragraph.text = new_text

        updated = 0
        for inner_index in range(index + 1, len(blocks)):
            inner_kind, inner_obj = blocks[inner_index]
            if inner_kind == "p" and is_heading(inner_obj, 3):
                break
            if inner_kind != "p":
                continue
            func_heading = inner_obj
            if not is_heading(func_heading, 4):
                continue
            updated += 1
            if updated > csu_count:
                continue
            new_id = f"{module_id}_{updated:03d}"
            new_heading_text = replace_csu_id_in_text(func_heading.text, new_id, backend_module=backend)
            if new_heading_text != func_heading.text:
                func_heading.text = new_heading_text
        total_headings += min(updated, csu_count)
        return {"modules": 1, "headings": total_headings}

    return {"modules": 0, "headings": 0}


def update_all_module_headings_only(doc, module_entries, *, backend_module=None) -> dict[str, int]:
    backend = backend_module or legacy_backend()
    if not module_entries:
        return {"modules": 0, "headings": 0}

    blocks = list(iter_doc_blocks(doc, backend_module=backend))
    module_index = 0
    total_headings = 0
    block_index = 0
    while block_index < len(blocks) and module_index < len(module_entries):
        kind, obj = blocks[block_index]
        if kind != "p":
            block_index += 1
            continue
        paragraph = obj
        if not is_heading(paragraph, 3):
            block_index += 1
            continue

        module_id, csu_count = module_entries[module_index]
        new_text = replace_module_id_in_text(paragraph.text or "", module_id, backend_module=backend)
        if new_text != (paragraph.text or ""):
            paragraph.text = new_text

        updated = 0
        heading_index = block_index + 1
        while heading_index < len(blocks):
            inner_kind, inner_obj = blocks[heading_index]
            if inner_kind == "p" and is_heading(inner_obj, 3):
                break
            if inner_kind != "p":
                heading_index += 1
                continue
            func_heading = inner_obj
            if not is_heading(func_heading, 4):
                heading_index += 1
                continue
            updated += 1
            if updated <= csu_count:
                new_id = f"{module_id}_{updated:03d}"
                new_heading_text = replace_csu_id_in_text(func_heading.text, new_id, backend_module=backend)
                if new_heading_text != func_heading.text:
                    func_heading.text = new_heading_text
            heading_index += 1

        total_headings += min(updated, csu_count)
        module_index += 1
        block_index = heading_index

    return {"modules": module_index, "headings": total_headings}


def update_module_csu_in_doc(doc, module_id: str, csu_entries, *, backend_module=None) -> dict[str, bool]:
    backend = backend_module or legacy_backend()
    module_id = (module_id or "").strip()
    if not module_id:
        return {"table": False, "headings": False}

    blocks = list(iter_doc_blocks(doc, backend_module=backend))
    for index, (kind, obj) in enumerate(blocks):
        if kind != "p":
            continue
        paragraph = obj
        if not is_heading(paragraph, 3):
            continue
        if extract_module_id(paragraph.text or "", backend_module=backend) != module_id:
            continue

        table = None
        table_index = None
        for inner_index in range(index + 1, len(blocks)):
            inner_kind, inner_obj = blocks[inner_index]
            if inner_kind == "p" and is_heading(inner_obj, 3):
                break
            if inner_kind == "tbl" and is_csu_table(inner_obj, backend_module=backend):
                table = inner_obj
                table_index = inner_index
                break

        if table is None:
            break

        try:
            module_name = (table.rows[1].cells[0].text or "").strip() if len(table.rows) > 1 else ""
        except Exception:
            module_name = ""
        table_updated = update_module_function_table(
            table,
            module_name,
            module_id,
            csu_entries,
            backend_module=backend,
        )

        updated = 0
        for inner_index in range((table_index or index) + 1, len(blocks)):
            inner_kind, inner_obj = blocks[inner_index]
            if inner_kind == "p" and is_heading(inner_obj, 3):
                break
            if inner_kind != "p":
                continue
            func_heading = inner_obj
            if not is_heading(func_heading, 4):
                continue
            updated += 1
            if updated > len(csu_entries or []):
                continue
            new_id = f"{module_id}_{updated:03d}"
            new_heading_text = replace_csu_id_in_text(func_heading.text, new_id, backend_module=backend)
            if new_heading_text != func_heading.text:
                func_heading.text = new_heading_text
        return {"table": table_updated, "headings": updated > 0}

    for kind, obj in blocks:
        if kind != "tbl":
            continue
        table = obj
        if not is_csu_table(table, backend_module=backend):
            continue
        flat = " ".join((cell.text or "") for row in table.rows for cell in row.cells)
        if module_id not in flat:
            continue
        return {
            "table": update_module_function_table(
                table,
                "",
                module_id,
                csu_entries,
                backend_module=backend,
            ),
            "headings": False,
        }

    return {"table": False, "headings": False}


def strip_trailing_id_suffix(text: str, *, backend_module=None) -> str:
    value = (text or "").strip()
    if not value:
        return value

    def _strip_by_parens(left: str, right: str) -> str | None:
        if left not in value or (not value.endswith(right)):
            return None
        index = value.rfind(left)
        if index < 0:
            return None
        inside = value[index + len(left) : -len(right)]
        if _MOD_ID_RE.search(inside) or _CSU_ID_RE.search(inside):
            return value[:index].strip()
        return None

    result = _strip_by_parens("（", "）")
    if result is not None:
        return result
    result = _strip_by_parens("(", ")")
    if result is not None:
        return result
    return value


def replace_or_append_csu_id_in_text(text: str, new_csu_id: str, *, backend_module=None) -> str:
    if not text:
        return f"（{new_csu_id}）"
    if _CSU_ID_RE.search(text) or _MOD_ID_RE.search(text):
        return replace_csu_id_in_text(text, new_csu_id)
    return f"{text}（{new_csu_id}）"


def detect_req_prefix_from_id(any_id: str, *, kind: str = "unknown", backend_module=None) -> str | None:
    value = (any_id or "").strip()
    if not value:
        return None

    def _strip_one(text: str) -> str:
        match = re.match(r"^(?P<prefix>.+)_\d{3}$", text)
        return match.group("prefix") if match else text

    if kind == "csu":
        return _strip_one(_strip_one(value))
    if kind == "module":
        return _strip_one(value)
    if re.search(r"_\d{3}_\d{3}_\d{3}$", value):
        return _strip_one(_strip_one(value))
    return _strip_one(value)


def update_csu_ids_in_design_chapter_by_headings(
    doc,
    cfg,
    *,
    chapter_keyword: str = "CSCI详细设计",
    backend_module=None,
) -> dict[str, int]:
    backend = backend_module or legacy_backend()
    blocks = list(iter_doc_blocks(doc, backend_module=backend))

    chapter_index = None
    chapter_level = None
    for index, (kind, obj) in enumerate(blocks):
        if kind != "p":
            continue
        paragraph = obj
        text = (paragraph.text or "").strip()
        if not text:
            continue
        if chapter_keyword in text:
            level = get_heading_level(paragraph, 6)
            if level is None:
                continue
            chapter_index = index
            chapter_level = level
            break

    if chapter_index is None or chapter_level is None:
        raise ValueError(f"未找到章节标题：{chapter_keyword}")

    end_index = len(blocks)
    for index in range(chapter_index + 1, len(blocks)):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        level = get_heading_level(obj, 6)
        if level is not None and level <= chapter_level:
            end_index = index
            break

    detected_prefix = None
    for index in range(chapter_index + 1, end_index):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        paragraph = obj
        if is_heading(paragraph, 3):
            module_id = extract_module_id(paragraph.text or "", backend_module=backend)
            if module_id:
                detected_prefix = detect_req_prefix_from_id(
                    module_id,
                    kind="module",
                    backend_module=backend,
                )
                break
        if is_heading(paragraph, 4):
            csu_id = extract_csu_id(paragraph.text or "", backend_module=backend)
            if csu_id:
                detected_prefix = detect_req_prefix_from_id(
                    csu_id,
                    kind="csu",
                    backend_module=backend,
                )
                break

    req_prefix = detected_prefix or normalize_req_prefix(
        getattr(cfg, "req_id_prefix", "") or DEFAULT_REQ_ID_PREFIX
    )
    req_prefix = normalize_req_prefix(req_prefix)

    modules_updated = 0
    functions_updated = 0
    tables_updated = 0

    block_index = chapter_index + 1
    module_counter = 1
    while block_index < end_index:
        kind, obj = blocks[block_index]
        if kind != "p":
            block_index += 1
            continue
        paragraph = obj
        if not is_heading(paragraph, 3):
            block_index += 1
            continue

        module_name = strip_trailing_id_suffix(paragraph.text or "", backend_module=backend)
        module_id = f"{req_prefix}_{module_counter:03d}"
        module_counter += 1

        new_text = replace_module_id_in_text(paragraph.text or "", module_id, backend_module=backend)
        if new_text != (paragraph.text or ""):
            paragraph.text = new_text
        elif extract_module_id(paragraph.text or "", backend_module=backend) != module_id:
            paragraph.text = replace_module_id_in_text(
                paragraph.text or "",
                module_id,
                backend_module=backend,
            )
        modules_updated += 1

        csu_entries: list[dict[str, str]] = []
        func_counter = 1
        next_index = block_index + 1
        table = None
        while next_index < end_index:
            inner_kind, inner_obj = blocks[next_index]
            if inner_kind == "p" and is_heading(inner_obj, 3):
                break
            if inner_kind == "tbl" and table is None and is_csu_table(inner_obj, backend_module=backend):
                table = inner_obj
                next_index += 1
                continue
            if inner_kind != "p":
                next_index += 1
                continue
            func_heading = inner_obj
            if not is_heading(func_heading, 4):
                next_index += 1
                continue

            func_title = strip_trailing_id_suffix(func_heading.text or "", backend_module=backend)
            csu_id = f"{module_id}_{func_counter:03d}"
            func_counter += 1

            func_heading.text = replace_or_append_csu_id_in_text(
                func_heading.text or func_title,
                csu_id,
                backend_module=backend,
            )
            functions_updated += 1
            csu_entries.append({"csu_name": func_title})
            next_index += 1

        if table is not None and csu_entries:
            if update_module_function_table(
                table,
                module_name,
                module_id,
                csu_entries,
                backend_module=backend,
            ):
                tables_updated += 1

        block_index = next_index

    return {"modules": modules_updated, "functions": functions_updated, "tables": tables_updated}


def find_unit_table_in_doc(doc, *, backend_module=None):
    backend = backend_module or legacy_backend()
    for kind, obj in iter_doc_blocks(doc, backend_module=backend):
        if kind != "tbl":
            continue
        table = obj
        try:
            if len(table.rows) < 1:
                continue
            if len(table.columns) >= len(_UNIT_TABLE_HEADERS):
                header = [cell.text.strip() for cell in table.rows[0].cells[: len(_UNIT_TABLE_HEADERS)]]
                if tuple(header) == _UNIT_TABLE_HEADERS:
                    return table
            if len(table.columns) >= len(_UNIT_TABLE_HEADERS_LEGACY):
                header = [cell.text.strip() for cell in table.rows[0].cells[: len(_UNIT_TABLE_HEADERS_LEGACY)]]
                if tuple(header) == _UNIT_TABLE_HEADERS_LEGACY:
                    return table
        except Exception:
            continue
    return None


def read_existing_unit_table_meta(unit_doc_path: str, *, backend_module=None) -> dict[str, dict[str, str]]:
    if not unit_doc_path or (not os.path.isfile(unit_doc_path)):
        return {}
    try:
        doc = Document(unit_doc_path)
    except Exception:
        return {}
    table = find_unit_table_in_doc(doc)
    if table is None:
        return {}
    result: dict[str, dict[str, str]] = {}
    header = [cell.text.strip() for cell in table.rows[0].cells]
    if header[: len(_UNIT_TABLE_HEADERS)] == list(_UNIT_TABLE_HEADERS):
        idx_req = 3
        idx_loc = 4
        idx_status = 5
    elif header[: len(_UNIT_TABLE_HEADERS_LEGACY)] == list(_UNIT_TABLE_HEADERS_LEGACY):
        idx_req = 2
        idx_loc = 3
        idx_status = 4
    else:
        return result
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) <= max(idx_req, idx_loc, idx_status):
            continue
        req_id = (cells[idx_req].text or "").strip()
        if not req_id:
            continue
        result[req_id] = {
            "location": (cells[idx_loc].text or "").strip(),
            "status": (cells[idx_status].text or "").strip(),
        }
    return result


def extract_purpose_from_function_blocks(blocks, start_idx: int, end_idx: int, *, backend_module=None) -> str:
    label_index = None
    for index in range(start_idx, end_idx):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        text = (obj.text or "").strip()
        if text == "b) 功能说明":
            label_index = index
            break
    if label_index is None:
        return ""
    for index in range(label_index + 1, end_idx):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        text = (obj.text or "").strip()
        if not text:
            continue
        if re.match(r"^[a-z]\)\s+", text):
            break
        if text in ("无。", "无", "略。", "略"):
            return ""
        return re.sub(r"[。;；]+$", "", text).strip()
    return ""


def extract_prototype_from_function_blocks(blocks, start_idx: int, end_idx: int, *, backend_module=None) -> str:
    label_index = None
    for index in range(start_idx, end_idx):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        text = (obj.text or "").strip()
        if text == "a) 函数原型":
            label_index = index
            break
    if label_index is None:
        return ""
    for index in range(label_index + 1, end_idx):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        text = (obj.text or "").strip()
        if not text:
            continue
        if re.match(r"^[a-z]\)\s+", text):
            break
        return text
    return ""


def collect_units_from_design_doc(
    doc,
    cfg,
    *,
    chapter_keyword: str = "CSCI详细设计",
    backend_module=None,
) -> list[dict]:
    backend = backend_module or legacy_backend()
    blocks = list(iter_doc_blocks(doc, backend_module=backend))

    chapter_index = None
    chapter_level = None
    for index, (kind, obj) in enumerate(blocks):
        if kind != "p":
            continue
        text = (obj.text or "").strip()
        if not text:
            continue
        if chapter_keyword in text:
            level = get_heading_level(obj, 6)
            if level is None:
                continue
            chapter_index = index
            chapter_level = level
            break
    if chapter_index is None or chapter_level is None:
        raise ValueError(f"未找到章节标题：{chapter_keyword}")

    end_index = len(blocks)
    for index in range(chapter_index + 1, len(blocks)):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        level = get_heading_level(obj, 6)
        if level is not None and level <= chapter_level:
            end_index = index
            break

    units: list[dict[str, str]] = []
    module_name = ""
    for index in range(chapter_index + 1, end_index):
        kind, obj = blocks[index]
        if kind != "p":
            continue
        paragraph = obj
        if is_heading(paragraph, 3):
            module_name = strip_trailing_id_suffix(paragraph.text or "", backend_module=backend)
            continue
        if not is_heading(paragraph, 4):
            continue
        func_title = strip_trailing_id_suffix(paragraph.text or "", backend_module=backend)
        csu_id = extract_csu_id(paragraph.text or "", backend_module=backend) or ""
        next_index = index + 1
        while next_index < end_index:
            inner_kind, inner_obj = blocks[next_index]
            if inner_kind == "p" and (is_heading(inner_obj, 3) or is_heading(inner_obj, 4)):
                break
            next_index += 1
        purpose = extract_purpose_from_function_blocks(
            blocks,
            index,
            next_index,
            backend_module=backend,
        ) or func_title
        prototype = extract_prototype_from_function_blocks(
            blocks,
            index,
            next_index,
            backend_module=backend,
        )
        units.append(
            {
                "module": module_name,
                "name": func_title,
                "req_id": csu_id,
                "purpose": purpose,
                "prototype": prototype,
            }
        )
    return units


def _resolve_unit_location(module_name: str, prev_location: str) -> str:
    prev = str(prev_location or "").strip()
    if prev.lower().endswith(".c"):
        return os.path.basename(prev)

    module = str(module_name or "").strip()
    safe = "".join(ch for ch in module if ch.isascii() and (ch.isalnum() or ch == "_"))
    if safe:
        return f"{safe}.c"
    return "xxx.c"


def update_software_unit_table_from_design_doc(
    doc,
    cfg,
    *,
    design_doc_path: str,
    chapter_keyword: str = "CSCI详细设计",
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    design_doc_path = os.path.abspath(design_doc_path or "")
    if not design_doc_path.lower().endswith(".docx"):
        raise ValueError("design_doc_path 必须是 .docx")

    unit_output = backend.derive_software_unit_output_path(design_doc_path)
    existing_meta = read_existing_unit_table_meta(unit_output, backend_module=backend)
    units = collect_units_from_design_doc(
        doc,
        cfg,
        chapter_keyword=chapter_keyword,
        backend_module=backend,
    )

    unit_rows: list[dict[str, Any]] = []
    by_module: dict[str, list[dict[str, Any]]] = {}
    for index, unit in enumerate(units, start=1):
        req_id = str((unit or {}).get("req_id") or "").strip()
        meta = existing_meta.get(req_id) if req_id else None
        status = (meta or {}).get("status") or "新研"
        module_name = str((unit or {}).get("module") or "").strip() or "未命名单元"
        location = _resolve_unit_location(
            module_name,
            str((meta or {}).get("location") or ""),
        )
        unit_rows.append(
            {
                "index": index,
                "name": str((unit or {}).get("name") or ""),
                "prototype": str((unit or {}).get("prototype") or ""),
                "req_id": req_id,
                "location": location,
                "status": status,
                "purpose": str((unit or {}).get("purpose") or ""),
            }
        )
        by_module.setdefault(module_name, []).append(
            {
                "index": len(by_module.get(module_name, [])) + 1,
                "name": str((unit or {}).get("name") or ""),
                "prototype": str((unit or {}).get("prototype") or ""),
            }
        )

    build_software_unit_table_doc(
        unit_rows,
        unit_output,
        backend_module=backend,
    )
    unit_func_tables = [
        {"unit_name": module_name, "func_rows": rows}
        for module_name, rows in by_module.items()
    ]
    func_output = backend.derive_unit_function_list_output_path(design_doc_path)
    build_unit_function_list_doc(
        unit_func_tables,
        func_output,
        backend_module=backend,
    )
    return {
        "units": len(unit_rows),
        "unit_output": unit_output,
        "unit_func_output": func_output,
    }


def build_software_unit_table_doc(unit_rows, output: str, *, backend_module=None) -> None:
    doc = Document()
    headers = ["序号", "软件单元名称", "函数原型", "唯一标识", "存放位置", "开发状态", "用途"]
    count = max(1, len(unit_rows or []))
    table = doc.add_table(rows=1 + count, cols=len(headers))
    if backend_module is not None:
        backend_module.apply_table_style(table, doc)
    else:
        apply_table_style(table, doc)
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    if not unit_rows:
        row = table.rows[1].cells
        row[0].text = "1"
        row[1].text = "无"
        row[2].text = ""
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""
        row[6].text = ""
        prevent_table_row_splitting(table)
        if backend_module is not None:
            backend_module.safe_save_docx(doc, output)
        else:
            safe_save_docx(doc, output)
        return

    for index, item in enumerate(unit_rows, start=1):
        row = table.rows[index].cells
        row[0].text = str(item.get("index") or index)
        row[1].text = str(item.get("name") or "")
        row[2].text = str(item.get("prototype") or "")
        row[3].text = str(item.get("req_id") or "")
        row[4].text = str(item.get("location") or "")
        row[5].text = str(item.get("status") or "")
        row[6].text = str(item.get("purpose") or "")

    prevent_table_row_splitting(table)
    if backend_module is not None:
        backend_module.safe_save_docx(doc, output)
    else:
        safe_save_docx(doc, output)


def build_unit_function_list_doc(unit_tables, output: str, *, backend_module=None) -> None:
    doc = Document()
    headers = ["序号", "函数名称", "函数原型"]

    if not unit_tables:
        doc.add_paragraph("本单元包含的函数列表")
        table = doc.add_table(rows=2, cols=len(headers))
        if backend_module is not None:
            backend_module.apply_table_style(table, doc)
        else:
            apply_table_style(table, doc)
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        row = table.rows[1].cells
        row[0].text = "1"
        row[1].text = "无"
        row[2].text = ""
        prevent_table_row_splitting(table)
        if backend_module is not None:
            backend_module.safe_save_docx(doc, output)
        else:
            safe_save_docx(doc, output)
        return

    for index, unit in enumerate(unit_tables, start=1):
        if index > 1:
            doc.add_page_break()
        unit_name = str((unit or {}).get("unit_name") or (unit or {}).get("name") or f"单元{index}")
        doc.add_paragraph(unit_name)
        doc.add_paragraph("本单元包含的函数列表")
        rows = list((unit or {}).get("func_rows") or [])
        count = max(1, len(rows))
        table = doc.add_table(rows=1 + count, cols=len(headers))
        if backend_module is not None:
            backend_module.apply_table_style(table, doc)
        else:
            apply_table_style(table, doc)
        for header_index, header in enumerate(headers):
            table.rows[0].cells[header_index].text = header
        if not rows:
            row = table.rows[1].cells
            row[0].text = "1"
            row[1].text = "无"
            row[2].text = ""
            prevent_table_row_splitting(table)
            continue
        for row_index, row_data in enumerate(rows, start=1):
            row = table.rows[row_index].cells
            row[0].text = str((row_data or {}).get("index") or row_index)
            row[1].text = str((row_data or {}).get("name") or "")
            row[2].text = str((row_data or {}).get("prototype") or "")
        prevent_table_row_splitting(table)

    if backend_module is not None:
        backend_module.safe_save_docx(doc, output)
    else:
        safe_save_docx(doc, output)


def to_design_model(design) -> DesignModel:
    backend = legacy_backend()
    if isinstance(design, DesignModel):
        return design
    return DesignModel(
        func_name="",
        func_cn_name=utils._safe_strip(getattr(design, "title", "")),
        desc="\n".join(getattr(design, "description_lines", ()) or ()),
        params=[{"name": e.name, "ident": e.ident, "c_type": e.c_type, "direction": e.direction} for e in (getattr(design, "io_elements", ()) or ())],
        locals=[{"name": e.name, "ident": e.ident, "c_type": e.c_type, "usage": e.usage} for e in (getattr(design, "local_elements", ()) or ())],
        logic_steps=list(getattr(design, "logic_lines", ()) or ()),
        file_context={},
    )


def render_table_or_none(doc, caption_suffix: str, headers, rows, *, backend_module=None):
    backend = backend_module or legacy_backend()
    if not rows:
        if backend_module is not None:
            return backend.add_indented_text(doc, "无。")
        return add_indented_text(doc, "无。")
    if backend_module is not None:
        return backend.add_captioned_table(doc, caption_suffix, headers, rows)
    caption = doc.add_paragraph()
    caption.add_run("表 ")
    add_seq_field(caption, "表")
    caption.add_run(f" {caption_suffix}")
    style_as_caption(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    apply_table_style(table, doc)
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = "" if header is None else str(header)
    for row_vals in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_vals):
            if index >= len(headers):
                break
            row[index].text = "" if value is None else str(value)
    prevent_table_row_splitting(table)
    return table


def _render_call_graph_section(doc, cfg, *, backend_module=None) -> None:
    try:
        from . import graph_visuals
    except Exception:
        return
    if not graph_visuals.graph_enabled(cfg):
        return
    func_data = getattr(cfg, "_current_render_func_data", None)
    if not isinstance(func_data, dict):
        return
    payload = graph_visuals.build_function_graph_payload(func_data, cfg)
    if payload.get("nodes"):
        graph_visuals.append_payload(cfg, payload)
    return


def render_project_graph_overview(doc, func_entries, cfg, *, root_dir: str = "", backend_module=None) -> bool:
    backend = backend_module or legacy_backend()
    try:
        from . import graph_visuals
    except Exception:
        return False
    if not graph_visuals.graph_enabled(cfg):
        return False
    payload = graph_visuals.build_project_overview_payload(func_entries or [], cfg, root_dir=root_dir)
    if payload.get("nodes"):
        graph_visuals.append_payload(cfg, payload)
    if not graph_visuals.payload_has_edges(payload):
        return False
    doc.add_paragraph("调用关系总览", style=backend.pick_heading_style(doc, 2))
    image_path = graph_visuals.render_payload_png(payload, cfg)
    if image_path and os.path.exists(image_path):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        style_as_normal_paragraph(p)
        try:
            p.add_run().add_picture(image_path, width=Inches(6.5))
            return True
        except Exception:
            pass
    rows = graph_visuals.fallback_rows(payload)
    if rows:
        render_table_or_none(doc, "调用关系总览", ["关系", "来源", "目标"], rows, backend_module=backend)
        return True
    return False


def _render_function_design_impl(doc, design, cfg, *, backend_module=None):
    backend = backend_module or legacy_backend()
    heading_style = backend.pick_heading_style(doc, 4) if backend_module is not None else pick_heading_style(doc, 4)
    paragraph = doc.add_paragraph(style=heading_style)
    paragraph.add_run(f"{design.title}（{design.req_id}）")

    if backend_module is not None:
        backend.add_alpha_section_label(doc, "函数原型", 0, indent_pt=24)
        backend.add_indented_text(doc, design.prototype, indent_pt=48)
    else:
        add_alpha_section_label(doc, "函数原型", 0, indent_pt=24)
        add_indented_text(doc, design.prototype, indent_pt=48)

    if backend_module is not None:
        backend.add_alpha_section_label(doc, "功能说明", 1, indent_pt=24)
    else:
        add_alpha_section_label(doc, "功能说明", 1, indent_pt=24)
    if design.description_lines:
        for line in design.description_lines:
            if backend_module is not None:
                backend.add_indented_text(doc, line, indent_pt=48)
            else:
                add_indented_text(doc, line, indent_pt=48)
    else:
        if backend_module is not None:
            backend.add_indented_text(doc, "无。", indent_pt=48)
        else:
            add_indented_text(doc, "无。", indent_pt=48)

    if backend_module is not None:
        backend.add_alpha_section_label(doc, "输入/输出元素", 2, indent_pt=24)
    else:
        add_alpha_section_label(doc, "输入/输出元素", 2, indent_pt=24)
    if design.io_none or (not design.io_elements):
        if backend_module is not None:
            backend.add_indented_text(doc, "无。", indent_pt=48)
        else:
            add_indented_text(doc, "无。", indent_pt=48)
    else:
        rows = [[item.name, item.ident, item.c_type, item.direction] for item in design.io_elements]
        render_table_or_none(doc, "输入/输出元素", ["名称", "标识", "类型", "输入/输出"], rows, backend_module=backend)

    if backend_module is not None:
        backend.add_alpha_section_label(doc, "局部数据元素", 3, indent_pt=24)
    else:
        add_alpha_section_label(doc, "局部数据元素", 3, indent_pt=24)
    if design.local_elements is None:
        if backend_module is not None:
            backend.add_indented_text(doc, "略。", indent_pt=48)
        else:
            add_indented_text(doc, "略。", indent_pt=48)
    elif not design.local_elements:
        if backend_module is not None:
            backend.add_indented_text(doc, "无。", indent_pt=48)
        else:
            add_indented_text(doc, "无。", indent_pt=48)
    else:
        rows = [[item.name, item.ident, item.c_type, item.usage] for item in design.local_elements]
        render_table_or_none(doc, "局部数据元素", ["名称", "标识", "类型", "用途"], rows, backend_module=backend)

    if backend_module is not None:
        backend.add_alpha_section_label(doc, "逻辑/流程图", 4, indent_pt=24)
    else:
        add_alpha_section_label(doc, "逻辑/流程图", 4, indent_pt=24)
    if design.logic_lines is None or not design.logic_lines:
        if backend_module is not None:
            backend.add_indented_text(doc, "略。", indent_pt=48)
        else:
            add_indented_text(doc, "略。", indent_pt=48)
    else:
        for line in design.logic_lines:
            if backend_module is not None:
                level, clean = _logic_line_indent_level(line)
                backend.add_indented_text(doc, clean, indent_pt=48 + level * 18)
            else:
                add_logic_text(doc, line, base_indent_pt=48, level_indent_pt=18)
    _render_call_graph_section(doc, cfg, backend_module=backend)


def render_design_model(doc, design_model: DesignModel, cfg) -> None:
    backend = legacy_backend()
    fd = FunctionDesign(
        title=design_model.func_cn_name or design_model.func_name,
        req_id="",
        prototype=design_model.func_name,
        description_lines=tuple((design_model.desc or "").splitlines() or ([design_model.desc] if design_model.desc else [])),
        io_elements=tuple(),
        io_none=not bool(design_model.params),
        local_elements=tuple(),
        logic_lines=tuple(design_model.logic_steps or []),
    )
    return _render_function_design_impl(doc, fd, cfg, backend_module=backend)


def render_function_design(doc, design, cfg):
    return _render_function_design_impl(doc, design, cfg)


def add_function_section(doc, func_data: dict, module_req_prefix: str, index: int, cfg):
    from . import pipeline as pipeline_utils

    design = pipeline_utils.build_function_design_impl(
        func_data,
        module_req_prefix,
        index,
        cfg,
    )
    try:
        cfg._current_render_func_data = func_data
        render_function_design(doc, design, cfg)
    finally:
        try:
            cfg._current_render_func_data = None
        except Exception:
            pass
    return design


def __getattr__(name: str) -> Any:
    return getattr(legacy_backend(), name)


__all__ = [
    "DesignModel",
    "add_ai_assist_banner",
    "add_function_section",
    "add_module_section",
    "init_document",
    "init_generation_document",
    "render_design_model",
    "render_function_design",
    "render_project_graph_overview",
    "render_table_or_none",
    "prevent_table_row_splitting",
    "safe_save_docx",
    "to_design_model",
    "detect_req_prefix_from_id",
    "build_software_unit_table_doc",
    "build_unit_function_list_doc",
    "extract_prototype_from_function_blocks",
    "extract_purpose_from_function_blocks",
    "find_unit_table_in_doc",
    "collect_units_from_design_doc",
    "read_existing_unit_table_meta",
    "replace_or_append_csu_id_in_text",
    "strip_trailing_id_suffix",
    "update_all_module_tables_and_headings",
    "update_csu_ids_in_design_chapter_by_headings",
    "update_module_function_table",
    "replace_csu_in_doc",
    "insert_csu_after_in_doc",
    "delete_csu_in_doc",
    "sync_module_function_table_for_module",
    "renumber_module_csu_ids",
    "update_software_unit_table_from_design_doc",
]


def replace_csu_in_doc(
    doc,
    csu_id: str,
    new_heading_text: str,
    new_body_elements,
    *,
    backend_module=None,
) -> dict[str, Any]:
    """Replace a single CSU's body in an existing document in-place.

    Locates the Heading 4 whose text ends with ``csu_id``, deletes all body
    elements from that heading up to (but not including) the next Heading 4 or
    Heading 3, then inserts ``new_body_elements`` at the same position.
    The heading paragraph itself is kept but its text is updated to
    ``new_heading_text`` (only when the title portion changed).

    Returns ``{"found": bool, "replaced": int, "old_title": str, "new_title": str}``.
    """
    backend = backend_module or legacy_backend()
    csu_id = (csu_id or "").strip()
    if not csu_id:
        return {"found": False, "replaced": 0, "old_title": "", "new_title": ""}

    body = doc._body._element
    elements = list(body)

    # 1. Locate the target Heading 4 by CSU id suffix.
    target_idx = None
    old_heading_text = ""
    for idx, elem in enumerate(elements):
        tag = getattr(elem, "tag", "")
        if not tag.endswith("}p"):
            continue
        para = DocxParagraph(elem, doc)
        if not is_heading(para, 4):
            continue
        if elem is None:
            continue
        text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
        if text.rstrip().endswith(csu_id) or f"（{csu_id}）" in text:
            target_idx = idx
            old_heading_text = text
            break

    if target_idx is None:
        return {"found": False, "replaced": 0, "old_title": "", "new_title": ""}

    # 2. Find the end of this CSU block = next Heading 4 or Heading 3.
    end_idx = len(elements)
    for idx in range(target_idx + 1, len(elements)):
        elem = elements[idx]
        tag = getattr(elem, "tag", "")
        if not tag.endswith("}p"):
            continue
        para = DocxParagraph(elem, doc)
        if is_heading(para, 4) or is_heading(para, 3):
            end_idx = idx
            break

    # 3. Delete elements (target_idx, end_idx) — the body after the heading.
    #    Keep the heading paragraph itself (target_idx).
    replaced_count = 0
    for idx in range(end_idx - 1, target_idx, -1):
        elem = elements[idx]
        body.remove(elem)
        replaced_count += 1

    # 4. Update heading text if the title changed (keep the same CSU id).
    #    new_heading_text already includes the id suffix, e.g. "新标题（D/R_SDD01_009_007）"
    new_title = new_heading_text or ""
    if new_title and new_title != old_heading_text:
        # Clear existing runs and set new text in the first run.
        p_elem = elements[target_idx]
        runs = p_elem.findall(qn("w:r"))
        if runs:
            # Set first run text, remove the rest.
            first_run = runs[0]
            t_nodes = first_run.findall(qn("w:t"))
            if t_nodes:
                t_nodes[0].text = new_title
                for extra_t in t_nodes[1:]:
                    first_run.remove(extra_t)
            else:
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = new_title
                first_run.append(t)
            for extra_run in runs[1:]:
                p_elem.remove(extra_run)
        else:
            # No runs at all — create one.
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = new_title
            r.append(t)
            p_elem.append(r)

    # 5. Insert new body elements right after the heading.
    anchor = elements[target_idx]
    for new_elem in reversed(new_body_elements or []):
        anchor.addnext(new_elem)

    return {
        "found": True,
        "replaced": replaced_count,
        "old_title": old_heading_text,
        "new_title": new_title,
    }


def insert_csu_after_in_doc(
    doc,
    after_csu_id: str,
    new_elements,
    *,
    backend_module=None,
) -> dict[str, Any]:
    """Insert a complete CSU section after an existing CSU block."""
    backend = backend_module or legacy_backend()
    after_csu_id = (after_csu_id or "").strip()
    if not after_csu_id:
        return {"found": False, "inserted": 0, "after_title": ""}

    body = doc._body._element
    elements = list(body)

    target_idx = None
    after_heading_text = ""
    for idx, elem in enumerate(elements):
        tag = getattr(elem, "tag", "")
        if not tag.endswith("}p"):
            continue
        para = DocxParagraph(elem, doc)
        if not is_heading(para, 4):
            continue
        text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
        if text.rstrip().endswith(after_csu_id) or f"（{after_csu_id}）" in text:
            target_idx = idx
            after_heading_text = text
            break

    if target_idx is None:
        return {"found": False, "inserted": 0, "after_title": ""}

    end_idx = len(elements)
    for idx in range(target_idx + 1, len(elements)):
        elem = elements[idx]
        tag = getattr(elem, "tag", "")
        if tag.endswith("}sectPr"):
            end_idx = idx
            break
        if not tag.endswith("}p"):
            continue
        para = DocxParagraph(elem, doc)
        if is_heading(para, 4) or is_heading(para, 3):
            end_idx = idx
            break

    insert_anchor = None
    for idx in range(end_idx - 1, target_idx - 1, -1):
        tag = getattr(elements[idx], "tag", "")
        if tag.endswith("}p") or tag.endswith("}tbl"):
            insert_anchor = elements[idx]
            break
    if insert_anchor is None:
        return {"found": False, "inserted": 0, "after_title": after_heading_text}

    insert_count = 0
    for new_elem in reversed(new_elements or []):
        insert_anchor.addnext(new_elem)
        insert_count += 1

    return {
        "found": True,
        "inserted": insert_count,
        "after_title": after_heading_text,
    }


def _module_id_from_csu_id(csu_id: str) -> str:
    match = re.match(r"^(.+?)_\d+$", (csu_id or "").strip())
    return match.group(1) if match else ""


def _cell_text(cell) -> str:
    try:
        return (cell.text or "").strip()
    except Exception:
        return ""


def sync_module_function_table_for_module(
    doc,
    module_id: str,
    *,
    backend_module=None,
) -> dict[str, Any]:
    """Synchronize one module's CSU table from actual H4 headings."""
    backend = backend_module or legacy_backend()
    module_id = (module_id or "").strip()
    if not module_id:
        return {"found": False, "updated": False, "entries": 0}

    blocks = list(iter_doc_blocks(doc, backend_module=backend))
    target_h3_index = None
    target_table = None
    module_name = ""

    for index, (kind, obj) in enumerate(blocks):
        if kind != "p" or not is_heading(obj, 3):
            continue
        end_index = len(blocks)
        for probe in range(index + 1, len(blocks)):
            probe_kind, probe_obj = blocks[probe]
            if probe_kind == "p" and is_heading(probe_obj, 3):
                end_index = probe
                break

        section_table = None
        section_has_module_csu = False
        for probe in range(index + 1, end_index):
            probe_kind, probe_obj = blocks[probe]
            if probe_kind == "tbl" and section_table is None and is_csu_table(probe_obj, backend_module=backend):
                section_table = probe_obj
                try:
                    if len(probe_obj.rows) > 1 and _cell_text(probe_obj.cell(1, 1)) == module_id:
                        section_has_module_csu = True
                except Exception:
                    pass
            if probe_kind == "p" and is_heading(probe_obj, 4):
                csu_id = extract_csu_id(probe_obj.text or "", backend_module=backend) or ""
                if _module_id_from_csu_id(csu_id) == module_id:
                    section_has_module_csu = True

        if section_has_module_csu:
            target_h3_index = index
            target_table = section_table
            module_name = strip_trailing_id_suffix(obj.text or "", backend_module=backend)
            break

    if target_h3_index is None or target_table is None:
        return {"found": False, "updated": False, "entries": 0}

    end_index = len(blocks)
    for probe in range(target_h3_index + 1, len(blocks)):
        probe_kind, probe_obj = blocks[probe]
        if probe_kind == "p" and is_heading(probe_obj, 3):
            end_index = probe
            break

    entries: list[dict[str, str]] = []
    for probe in range(target_h3_index + 1, end_index):
        probe_kind, probe_obj = blocks[probe]
        if probe_kind != "p" or not is_heading(probe_obj, 4):
            continue
        text = probe_obj.text or ""
        csu_id = extract_csu_id(text, backend_module=backend) or ""
        if _module_id_from_csu_id(csu_id) != module_id:
            continue
        entries.append({
            "csu_name": strip_trailing_id_suffix(text, backend_module=backend),
            "csu_id": csu_id,
        })

    try:
        while len(target_table.rows) < 1 + max(1, len(entries)):
            target_table.add_row()
    except Exception:
        pass

    row_count = max(0, len(target_table.rows) - 1)
    for row_index in range(1, row_count + 1):
        cells = target_table.rows[row_index].cells
        if row_index == 1:
            try:
                cells[0].text = module_name or _cell_text(cells[0]) or "(未命名模块)"
                cells[1].text = module_id
            except Exception:
                pass
        if row_index <= len(entries):
            cells[2].text = entries[row_index - 1]["csu_name"]
            cells[3].text = entries[row_index - 1]["csu_id"]
        else:
            cells[2].text = ""
            cells[3].text = ""

    return {"found": True, "updated": True, "entries": len(entries)}


def renumber_module_csu_ids(
    doc,
    module_id: str,
    *,
    backend_module=None,
) -> dict[str, Any]:
    """Renumber H4 CSU ids in one module by current document order."""
    backend = backend_module or legacy_backend()
    module_id = (module_id or "").strip()
    if not module_id:
        return {"found": False, "updated": 0, "mapping": []}

    blocks = list(iter_doc_blocks(doc, backend_module=backend))
    target_h3_index = None
    for index, (kind, obj) in enumerate(blocks):
        if kind != "p" or not is_heading(obj, 3):
            continue

        end_index = len(blocks)
        for probe in range(index + 1, len(blocks)):
            probe_kind, probe_obj = blocks[probe]
            if probe_kind == "p" and is_heading(probe_obj, 3):
                end_index = probe
                break

        section_has_module_csu = False
        for probe in range(index + 1, end_index):
            probe_kind, probe_obj = blocks[probe]
            if probe_kind != "p" or not is_heading(probe_obj, 4):
                continue
            csu_id = extract_csu_id(probe_obj.text or "", backend_module=backend) or ""
            if _module_id_from_csu_id(csu_id) == module_id:
                section_has_module_csu = True
                break

        if section_has_module_csu:
            target_h3_index = index
            break

    if target_h3_index is None:
        return {"found": False, "updated": 0, "mapping": []}

    end_index = len(blocks)
    for probe in range(target_h3_index + 1, len(blocks)):
        probe_kind, probe_obj = blocks[probe]
        if probe_kind == "p" and is_heading(probe_obj, 3):
            end_index = probe
            break

    mapping: list[dict[str, str]] = []
    updated = 0
    ordinal = 0
    for probe in range(target_h3_index + 1, end_index):
        probe_kind, probe_obj = blocks[probe]
        if probe_kind != "p" or not is_heading(probe_obj, 4):
            continue
        old_text = probe_obj.text or ""
        old_csu_id = extract_csu_id(old_text, backend_module=backend) or ""
        if _module_id_from_csu_id(old_csu_id) != module_id:
            continue
        ordinal += 1
        new_csu_id = f"{module_id}_{ordinal:03d}"
        new_text = replace_or_append_csu_id_in_text(old_text, new_csu_id, backend_module=backend)
        if new_text != old_text:
            probe_obj.text = new_text
            updated += 1
        mapping.append({
            "old_csu_id": old_csu_id,
            "new_csu_id": new_csu_id,
            "title": strip_trailing_id_suffix(old_text, backend_module=backend),
        })

    table_result = sync_module_function_table_for_module(doc, module_id, backend_module=backend)
    return {
        "found": True,
        "updated": updated,
        "mapping": mapping,
        "module_table": table_result,
    }


def delete_csu_in_doc(
    doc,
    csu_id: str,
    *,
    backend_module=None,
) -> dict[str, Any]:
    """Delete a single CSU section from an existing document in-place.

    Locates the Heading 4 whose text contains the CSU id, then deletes that
    heading and all following body elements up to (but not including) the next
    Heading 4 or Heading 3.
    """
    backend = backend_module or legacy_backend()
    csu_id = (csu_id or "").strip()
    if not csu_id:
        return {"found": False, "deleted": 0, "old_title": ""}

    body = doc._body._element
    elements = list(body)

    target_idx = None
    old_heading_text = ""
    for idx, elem in enumerate(elements):
        tag = getattr(elem, "tag", "")
        if not tag.endswith("}p"):
            continue
        para = DocxParagraph(elem, doc)
        if not is_heading(para, 4):
            continue
        text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
        if text.rstrip().endswith(csu_id) or f"（{csu_id}）" in text:
            target_idx = idx
            old_heading_text = text
            break

    if target_idx is None:
        return {"found": False, "deleted": 0, "old_title": ""}

    end_idx = len(elements)
    for idx in range(target_idx + 1, len(elements)):
        elem = elements[idx]
        tag = getattr(elem, "tag", "")
        if not tag.endswith("}p"):
            continue
        para = DocxParagraph(elem, doc)
        if is_heading(para, 4) or is_heading(para, 3):
            end_idx = idx
            break

    deleted_count = 0
    for idx in range(end_idx - 1, target_idx - 1, -1):
        body.remove(elements[idx])
        deleted_count += 1

    return {
        "found": True,
        "deleted": deleted_count,
        "old_title": old_heading_text,
    }
