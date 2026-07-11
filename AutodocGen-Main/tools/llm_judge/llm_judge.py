#!/usr/bin/env python3
"""
LLM-as-judge evaluator for AutodocGen AI mode output.

For each function:
  1. Read the source code (body only, no comments)
  2. Read the generated docx text (paragraphs + tables)
  3. Send to deepseek-v4-flash as judge
  4. Score on 5 dimensions (1-5 each)
  5. Aggregate: 0-25 total, mapped to A-F grade

Output: /tmp/llm_judge_<run>/report.json + report.md
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from autodoc import backend


JUDGE_PROMPT = """你是嵌入式 C 代码文档的评审专家. 请根据以下信息对生成的 Word 文档进行评分.

【源码(函数体)】
```c
{source}
```

【自动生成的文档】
```
{document}
```

请逐项打分 (1=极差, 5=优秀) 并给出简短依据 (≤20字/项):

1. **描述准确性** (accuracy): 文档中的功能描述、参数说明、返回值、变量名翻译是否与源码实际行为一致? 不允许出现与代码矛盾的虚构内容.
2. **信息量** (informativeness): 描述是否包含具体动作、操作对象、条件分支等具体信息? 还是只有 "执行功能处理" "完成相关更新" 等空泛模板句?
3. **逻辑步骤通顺** (logic_flow): 流程图/逻辑步骤是否按代码控制流 (if/while/for/switch) 完整呈现, 顺序与条件通顺? 是否漏掉关键分支?
4. **变量翻译** (var_translation): 形如 l_ii_u16, v_pData_t 这样的匈牙利命名变量, 翻译为中文时是否合理且符合上下文?
5. **无空泛占位** (no_placeholder): 是否完全没有出现 "执行操作" "执行功能处理并更新相关状态" "待人工修改" "完成相关处理" 等占位/模板句?

请以严格 JSON 格式返回 (不要任何额外文字):
{{
  "accuracy": <int 1-5>,
  "informativeness": <int 1-5>,
  "logic_flow": <int 1-5>,
  "var_translation": <int 1-5>,
  "no_placeholder": <int 1-5>,
  "reasons": {{
    "accuracy": "...",
    "informativeness": "...",
    "logic_flow": "...",
    "var_translation": "...",
    "no_placeholder": "..."
  }},
  "red_flags": ["..."],
  "total_25": <sum>,
  "grade": "A" | "B" | "C" | "D" | "F"
}}
"""


def read_docx_text(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_func_body(c_file: str, func_name: str, line_start: int) -> str:
    """Find function by name in raw source (line_start from include-expanded
    preprocessor is unreliable for our judge). Returns up to 200 lines starting
    from the first line containing the function name + opening brace."""
    text = backend.load_c_file(c_file)
    lines = text.splitlines()
    # find first line that contains "func_name" followed by " (" or just ends with (
    start_idx = -1
    for i, ln in enumerate(lines):
        stripped = ln.strip()
        if re.search(rf"\b{re.escape(func_name)}\s*\(", stripped):
            # verify it's a definition, not a call/declaration-without-body
            if "{" in ln or any("{" in lines[j] for j in range(i, min(i+5, len(lines)))):
                start_idx = i
                break
    if start_idx < 0:
        # fallback: take 100 lines from line_start, clipped
        idx = max(0, min(len(lines), line_start) - 1)
        return "\n".join(lines[idx: idx + 100])
    # walk braces
    body = []
    depth = 0
    started = False
    for j in range(start_idx, min(len(lines), start_idx + 300)):
        ln = lines[j]
        body.append(ln)
        depth += ln.count("{") - ln.count("}")
        if "{" in ln and not started:
            started = True
        if started and depth <= 0 and j > start_idx:
            break
    return "\n".join(body)


def call_judge(prompt: str, *, model: str, api_base: str, api_key: str) -> dict[str, Any]:
    import requests
    base = api_base.rstrip("/")
    if base.endswith("/chat/completions"):
        url = base
    elif base.endswith("/v1"):
        url = base + "/chat/completions"
    else:
        url = base + "/v1/chat/completions"
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是一个严格、专业的嵌入式 C 代码文档评审助手, 只输出有效 JSON."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.0,
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    resp = requests.post(url, json=payload, headers=headers, timeout=120)
    resp.raise_for_status()
    body = resp.json()
    text = body["choices"][0]["message"]["content"]
    m = re.search(r"\{[\s\S]*\}", text)
    if not m:
        raise RuntimeError(f"judge returned no JSON: {text[:200]}")
    return json.loads(m.group(0))


def main() -> None:
    p = argparse.ArgumentParser()
    p.add_argument("--input-report", required=True)
    p.add_argument("--output-dir", required=True)
    p.add_argument("--api-base", default="https://api.deepseek.com/v1")
    p.add_argument("--model", default="deepseek-v4-flash")
    p.add_argument("--api-key-env", default="DEEPSEEK_API_KEY")
    p.add_argument("--limit", type=int, default=10)
    p.add_argument("--seed", type=int, default=7)
    args = p.parse_args()

    api_key = os.environ.get(args.api_key_env, "")
    if not api_key:
        raise SystemExit(f"set env {args.api_key_env}")

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    report = json.load(open(args.input_report))
    import random
    rng = random.Random(args.seed)
    picks = rng.sample(report, min(args.limit, len(report)))

    results = []
    for idx, r in enumerate(picks, 1):
        sample = r["sample"]
        docx_path = r["output"]
        print(f"[{idx}/{len(picks)}] judging {sample['func_name']} <- {Path(sample['c_file']).name}", file=sys.stderr)
        try:
            doc_text = read_docx_text(docx_path)
            src_body = extract_func_body(sample["c_file"], sample["func_name"], sample["line_start"])
            if len(src_body) > 6000:
                src_body = src_body[:6000] + "\n/* ... truncated ... */"
            if len(doc_text) > 6000:
                doc_text = doc_text[:6000] + "\n... (truncated)"
            prompt = JUDGE_PROMPT.format(source=src_body, document=doc_text)
            t0 = time.time()
            judge = call_judge(prompt, model=args.model, api_base=args.api_base, api_key=api_key)
            elapsed = time.time() - t0
            entry = {
                "func_name": sample["func_name"],
                "c_file": sample["c_file"],
                "docx_score": r["score"],
                "judge": judge,
                "judge_elapsed": elapsed,
            }
            results.append(entry)
            print(f"  -> total={judge.get('total_25')}/25 grade={judge.get('grade')} ({elapsed:.1f}s)", file=sys.stderr)
        except Exception as exc:
            print(f"  ERROR: {exc}", file=sys.stderr)
            results.append({"func_name": sample["func_name"], "error": str(exc)})

    (out_dir / "report.json").write_text(json.dumps(results, ensure_ascii=False, indent=2))

    lines = [f"# LLM-Judge 评测报告 (样本 {len(results)})", ""]
    ok = [r for r in results if "judge" in r]
    if ok:
        totals = [r["judge"].get("total_25", 0) for r in ok]
        lines.append(f"- 平均分: {sum(totals)/len(totals):.1f}/25")
        lines.append(f"- 最低/最高: {min(totals)}/{max(totals)}")
        lines.append(f"- 25 分样本: {sum(1 for t in totals if t==25)}")
        lines.append(f"- 20+ 分样本: {sum(1 for t in totals if t>=20)}")
        lines.append(f"- 评分器100 vs judge>=20 一致率: {sum(1 for r in ok if r['docx_score']==100 and r['judge'].get('total_25',0)>=20)}/{len(ok)}")
        lines.append("")
        for r in ok:
            j = r["judge"]
            lines.append(f"## {r['func_name']} ({j.get('total_25')}/25, {j.get('grade')})")
            lines.append(f"- 文件: `{Path(r['c_file']).name}`")
            lines.append(f"- 评分器: {r['docx_score']}/100, judge: {j.get('total_25')}/25")
            reasons = j.get("reasons", {})
            for k in ["accuracy", "informativeness", "logic_flow", "var_translation", "no_placeholder"]:
                lines.append(f"  - {k}: {j.get(k)} - {reasons.get(k, '')}")
            flags = j.get("red_flags", [])
            if flags:
                lines.append(f"  - red_flags: {flags}")
            lines.append("")
    (out_dir / "report.md").write_text("\n".join(lines))
    print(f"saved {out_dir}/report.json report.md", file=sys.stderr)


if __name__ == "__main__":
    main()
