from __future__ import annotations

import argparse
import configparser
import json
import os
import random
import re
import sys
import time
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document

from autodoc import backend
from autodoc import lsp_facts
from autodoc import revision as revision_utils
from autodoc import quality_gate
from autodoc import effects as effects_utils
from autodoc.utils import resolve_api_key


@dataclass
class Sample:
    c_file: str
    func_name: str
    line_start: int
    ret_type: str = ""
    expected_outputs: tuple[str, ...] = ()
    expected_external_effects: tuple[str, ...] = ()
    expected_return_effects: tuple[str, ...] = ()
    expected_callee_effects: tuple[str, ...] = ()
    expected_unresolved_callee_effects: int = 0


@dataclass
class CheckResult:
    sample: dict[str, Any]
    output: str
    ok: bool
    elapsed_sec: float
    score: int
    warnings: list[str]
    error: str = ""
    docx_size: int = 0
    paragraphs: int = 0
    tables: int = 0
    text_excerpt: str = ""
    bad_terms: list[str] = field(default_factory=list)
    missing_expected_outputs: list[str] = field(default_factory=list)
    empty_type_count: int = 0
    quality_issues: list[dict[str, Any]] = field(default_factory=list)
    docx_leak_hits: list[str] = field(default_factory=list)
    judge_total_25: int = 0
    judge_grade: str = ""
    judge_red_flags: list[str] = field(default_factory=list)
    judge_elapsed_sec: float = 0.0
    judge_skipped: str = ""
    ai_quality_events: list[dict[str, Any]] = field(default_factory=list)
    effect_stats: dict[str, int] = field(default_factory=dict)


DOCX_LEAK_TERMS = (
    "logic_source_audit",
    "callee_comment",
    "state_update",
    "source_anchor",
    "provenance",
    "name_refs",
)

BAD_DOC_TERMS = (
    "待人工修改",
    "name 'sys' is not defined",
    "clangd 不可用",
    "生成失败",
    "Traceback",
    "本地模型",
    "返回 返回结果",
    "计算 返回结果",
    "执行操作",
    "有效 等于",
    "无效 等于",
    "IF NULL 等于",
    "计算 系统时间 = 系统时间",
    "完成相关处理",
    "科学中断",
)


def _logic_quality_issues_from_text(text: str) -> list[dict[str, Any]]:
    """Find only hard rendering defects, separate from general bad terms."""
    all_lines = str(text or "").splitlines()
    logic_start = next(
        (idx + 1 for idx, line in enumerate(all_lines) if re.search(r"(?:^|\s)(?:e\)\s*)?逻辑/流程图\s*$", line)),
        None,
    )
    # DOCX contains prototypes and table content with legitimate parentheses.
    # Audit only the rendered logic section when its heading is available.
    lines = all_lines[logic_start:] if logic_start is not None else all_lines
    logic_lines: list[str] = []
    for raw_line in lines:
        if re.match(r"^[a-z]\)\s*", raw_line.strip()):
            break
        logic_lines.append(raw_line)
    return [dict(item) for item in quality_gate.inspect_logic_lines(logic_lines, source="docx_logic")]


def _bool(value: Any, default: bool = False) -> bool:
    text = str(value if value is not None else "").strip().lower()
    if text in ("1", "true", "yes", "y", "on"):
        return True
    if text in ("0", "false", "no", "n", "off"):
        return False
    return default


def _int(value: Any, default: int = 0) -> int:
    try:
        return int(str(value).strip())
    except Exception:
        return default


def _ai_mode(value: Any, default: int = 0) -> int:
    return 1 if _int(value, default) > 0 else 0


def _float(value: Any, default: float = 0.0) -> float:
    try:
        return float(str(value).strip())
    except Exception:
        return default


def _lines(value: str) -> tuple[str, ...]:
    return tuple(x.strip() for x in str(value or "").replace("\\n", "\n").splitlines() if x.strip())


def _extra(raw: str) -> dict[str, Any]:
    text = str(raw or "").strip()
    if len(text) >= 2 and text[0] == text[-1] == '"':
        text = text[1:-1]
    text = text.replace('\\"', '"')
    try:
        value = json.loads(text) if text else {}
        return value if isinstance(value, dict) else {}
    except Exception:
        return {}


def _run_inline_judge(docx_path: Path, sample: Sample, args) -> tuple[int, str, list[str], float, str]:
    """Run LLM-judge on a single generated docx, return (total_25, grade, red_flags, elapsed, skip_reason).

    skip_reason is non-empty when the judge could not run (missing key, judge failed, etc.).
    Reuses llm_judge helper if available, else falls back to a minimal inline caller.
    """
    model = args.judge_model or args.ai_model
    api_base = args.judge_api_base or args.ai_api_base
    key_env = args.judge_api_key_env or args.ai_api_key_env
    if not model or not api_base or not key_env:
        return (0, "", [], 0.0, "judge 未配置 (缺 model/api_base/key_env)")
    api_key = os.environ.get(key_env, "")
    if not api_key:
        return (0, "", [], 0.0, f"env {key_env} 未设")
    try:
        from tools.llm_judge.llm_judge import read_docx_text, extract_func_body, call_judge, JUDGE_PROMPT
    except Exception as exc:
        return (0, "", [], 0.0, f"llm_judge 模块不可用: {exc}")
    try:
        t0 = time.time()
        doc_text = read_docx_text(str(docx_path))
        src_body = extract_func_body(sample.c_file, sample.func_name, sample.line_start)
        if len(src_body) > 6000:
            src_body = src_body[:6000] + "\n/* ... truncated ... */"
        if len(doc_text) > 6000:
            doc_text = doc_text[:6000] + "\n... (truncated)"
        prompt = JUDGE_PROMPT.format(source=src_body, document=doc_text)
        result = call_judge(prompt, model=model, api_base=api_base, api_key=api_key)
        elapsed = time.time() - t0
        total = int(result.get("total_25") or 0)
        grade = str(result.get("grade") or "")
        flags = [str(x) for x in (result.get("red_flags") or [])]
        return (total, grade, flags, elapsed, "")
    except Exception as exc:
        return (0, "", [], 0.0, f"judge 调用失败: {exc}")


def load_cfg(ini_path: Path, *, no_ai: bool, ai_workers: int | None) -> backend.GenConfig:
    cp = configparser.ConfigParser()
    cp.read(ini_path, encoding="utf-8")
    basic = cp["basic"] if cp.has_section("basic") else {}
    ai = cp["ai"] if cp.has_section("ai") else {}
    perf = cp["perf"] if cp.has_section("perf") else {}
    project = cp["project"] if cp.has_section("project") else {}
    adv = cp["advanced"] if cp.has_section("advanced") else {}
    extra = _extra(adv.get("extra_params_json", ""))
    ai_mode = 0 if no_ai else _ai_mode(ai.get("ai_mode", 0))
    cfg = backend.GenConfig(
        section_prefix=basic.get("section_prefix", "5.1.1."),
        req_id_prefix=basic.get("req_id_prefix", "D/R_SDD01_"),
        only_with_comment=_bool(basic.get("only_with_comment"), False),
        include_locals=_bool(basic.get("include_locals"), True),
        include_logic=_bool(basic.get("include_logic"), True),
        logic_use_comment=_bool(basic.get("logic_use_comment"), True),
        open_after_done=False,
        ai_assist=(ai_mode == 1),
        ai_mode=ai_mode,
        ai_provider=ai.get("ai_provider", "local"),
        ai_model=ai.get("ai_model", ""),
        ai_api_base=ai.get("ai_api_base", ""),
        ai_api_key=resolve_api_key(ai.get("ai_api_key", "")),
        ai_use_auth=bool(resolve_api_key(ai.get("ai_api_key", ""))),
        ai_num_ctx=_int(ai.get("ai_num_ctx", 0), 0),
        ai_read_timeout=_float(ai.get("ai_read_timeout", 40), 40),
        ai_workers=max(1, ai_workers if ai_workers is not None else _int(ai.get("ai_workers", 1), 1)),
        ai_max_tokens=_int(ai.get("ai_max_tokens", 16384), 16384),
        proxy=(ai.get("proxy", "") if _bool(ai.get("use_proxy"), False) else ""),
        no_proxy=_bool(ai.get("no_proxy"), False),
        ai_logic_format="json",
        ai_logic_policy="hybrid",
        ai_one_call=False,
        auto_disable_large_one_call=True,
        force_ai=False,
        verbose=_bool(ai.get("verbose"), True),
        preprocess_workers=max(0, _int(perf.get("preprocess_workers", 0), 0)),
        log_every_n=max(1, _int(perf.get("log_every_n", 5), 5)),
        prefilter_project_files=_bool(perf.get("prefilter_project_files"), True),
        extra_params=extra,
        exclude_dirs=_lines(project.get("exclude_dirs", ".git\n.settings\n.launches\ndebug\nrelease\n__pycache__")),
        mid_dir_keywords=_lines(project.get("mid_dir_keywords", "common")),
        drv_dir_keywords=_lines(project.get("drv_dir_keywords", "dspdriver")),
    )
    for key in ("ai_fail_policy", "ai_profile", "ai_retry_times"):
        if key in extra and hasattr(cfg, key):
            current = getattr(cfg, key)
            setattr(cfg, key, _int(extra[key], current) if isinstance(current, int) else str(extra[key]))
    if "wire_api" in extra and hasattr(cfg, "wire_api"):
        cfg.wire_api = str(extra["wire_api"] or "").strip() or cfg.wire_api
    return cfg


def _return_expr(body: str) -> str:
    match = re.search(r"\breturn\b(?P<expr>.*?);", str(body or ""), flags=re.S)
    if not match:
        return ""
    expr = re.sub(r"/\*.*?\*/", "", match.group("expr"), flags=re.S)
    expr = re.sub(r"//.*$", "", expr).strip()
    while expr.startswith("(") and expr.endswith(")"):
        expr = expr[1:-1].strip()
    return expr


def _root_ident(expr: str) -> str:
    value = str(expr or "").strip().lstrip("*&").strip()
    match = re.search(r"\b[A-Za-z_]\w*\b", value)
    return match.group(0) if match else ""


def _is_void_return_type(ret_type: str) -> bool:
    value = str(ret_type or "").strip()
    if not value:
        return False
    value = re.sub(r"\b(?:static|extern|inline|__inline|__inline__|const|volatile|register)\b", " ", value)
    value = re.sub(r"\s+", "", value).lower()
    return value == "void"


def _param_is_pointer_or_array(param: dict[str, Any]) -> bool:
    name = str((param or {}).get("name") or "").strip()
    ptype = str((param or {}).get("type") or "").strip()
    return bool("*" in ptype or "[" in ptype or name.startswith(("p_", "pp_", "vp_", "v_p_", "gp_", "lp_", "sp_", "cp_", "tp_")))


def _param_has_external_write(body: str, param: dict[str, Any]) -> bool:
    name = str((param or {}).get("name") or "").strip()
    if not name or not _param_is_pointer_or_array(param):
        return False
    try:
        _reads, writes = lsp_facts._collect_accesses(body)
    except Exception:
        return False
    ident = re.escape(name)
    patterns = (
        re.compile(rf"^\s*\*+\s*\(?\s*{ident}\s*\)?(?:\b|\s*(?:->|\.|\[))"),
        re.compile(rf"^\s*{ident}\s*(?:->|\[)"),
    )
    for item in writes:
        lhs = str(getattr(item, "lhs", "") or "").strip()
        if lhs and any(pattern.search(lhs) for pattern in patterns):
            return True
    return False


def expected_output_idents(fd: dict[str, Any], cfg: backend.GenConfig) -> tuple[str, ...]:
    fi = dict(fd.get("func_info") or {})
    comment_info = dict(fd.get("comment_info") or {})
    body = str(fd.get("body") or "")
    ret_type = str(fi.get("ret_type") or "").strip()
    outputs: list[str] = []
    if ret_type and not _is_void_return_type(ret_type):
        outputs.append(_return_expr(body) or "return")
    try:
        params = backend.parse_params_from_prototype(fi)
        out_map = backend.parse_param_desc(str(comment_info.get("output_desc") or ""))
        for param in params:
            name = str((param or {}).get("name") or "").strip()
            if name and (name in out_map or _param_has_external_write(body, param)):
                outputs.append(name)
    except Exception:
        pass
    return tuple(dict.fromkeys(x for x in outputs if x))


def expected_effect_idents(
    fd: dict[str, Any], cfg: backend.GenConfig | None = None, index: effects_utils.EffectIndex | None = None,
) -> tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...], int]:
    """Use the pipeline's deterministic effect rules as DOCX coverage oracle."""
    try:
        func_info = dict(fd.get("func_info") or {})
        params = backend.parse_params_from_prototype(func_info)
        locals_ = backend.parse_local_variables_from_body(str(fd.get("body") or ""))
        effects, returns = effects_utils.extract_direct_effects(fd, params=params, local_vars=locals_)
        external = tuple(item.target_ident for item in effects if item.kind == "global_write" and item.verified)
        result_values = tuple(item.target_ident for item in returns if item.verified)
        unresolved = 0
        callee_effects: tuple[str, ...] = ()
        if cfg is not None and index is not None:
            # Sample collection must remain fast and offline.  Reuse the
            # fallback call scanner instead of starting a per-function LSP
            # query merely to count unresolved one-hop callees.
            source_file = str((fd.get("file_context") or {}).get("source_file") or "")
            calls = lsp_facts._collect_calls(str(fd.get("body") or ""), {}, source_file)
            facts = {"calls": [
                {"callee": call.callee, "call_text": call.call_text, "range": vars(call.range)}
                for call in calls
            ]}
            mapped, issues = effects_utils.resolve_one_hop_effects(
                facts, index=index,
                source_file=source_file,
                source_function=str(func_info.get("func_name") or ""),
            )
            callee_effects = tuple(dict.fromkeys(item.target_ident for item in mapped if item.verified))
            unresolved = sum(1 for item in issues if item.get("code") == "callee_effect_unresolved")
        return tuple(dict.fromkeys(external)), tuple(dict.fromkeys(result_values)), callee_effects, unresolved
    except Exception:
        return (), (), (), 0


def iter_c_files(project_dir: Path, exclude_dirs: tuple[str, ...]) -> list[Path]:
    excluded = {x.lower() for x in exclude_dirs}
    files: list[Path] = []
    for root, dirs, names in os.walk(project_dir):
        dirs[:] = [d for d in dirs if d.lower() not in excluded and not d.startswith(".")]
        for name in names:
            if name.lower().endswith(".c"):
                files.append(Path(root) / name)
    return sorted(files)


def collect_samples(project_dir: Path, cfg: backend.GenConfig, *, count: int, seed: int, max_files: int) -> list[Sample]:
    rng = random.Random(seed)
    c_files = iter_c_files(project_dir, cfg.exclude_dirs)
    rng.shuffle(c_files)
    if max_files > 0:
        c_files = c_files[:max_files]
    effect_index = effects_utils.build_effect_index(str(project_dir), cfg)
    pool: list[Sample] = []
    for c_file in c_files:
        try:
            code = backend.load_c_file(str(c_file))
            funcs = backend.associate_comments_and_functions(code, file_context_extra={"source_file": str(c_file)})
        except Exception as exc:
            print(f"[WARN] 解析失败，跳过：{c_file}；{exc}", file=sys.stderr)
            continue
        for fd in funcs or []:
            fi = fd.get("func_info") or {}
            name = str(fi.get("func_name") or "").strip()
            if name:
                expected_effects, expected_returns, expected_callees, unresolved_callees = expected_effect_idents(fd, cfg, effect_index)
                pool.append(
                    Sample(
                        str(c_file),
                        name,
                        code.count("\n", 0, max(0, int(fi.get("start", 0) or 0))) + 1,
                        str(fi.get("ret_type") or "").strip(),
                        expected_output_idents(fd, cfg),
                        expected_effects,
                        expected_returns,
                        expected_callees,
                        unresolved_callees,
                    )
                )
    if not pool:
        raise SystemExit("没有找到可抽样函数")
    rng.shuffle(pool)
    return pool[: max(1, count)]


def doc_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _doc_table_rows(doc: Document) -> list[list[str]]:
    rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
    return rows


def _table_data_rows(doc: Document, headers: tuple[str, ...]) -> list[list[str]]:
    for table in doc.tables:
        if not table.rows:
            continue
        head = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if head == headers:
            return [[cell.text.strip() for cell in row.cells] for row in table.rows[1:]]
    return []


def _excerpt(text: str, warnings: list[str], *, max_chars: int = 700) -> str:
    if not text:
        return ""
    for warning in warnings:
        key = warning.split("：", 1)[-1].strip("` ")
        if key and key in text:
            idx = max(0, text.find(key) - 180)
            return text[idx: idx + max_chars]
    return text[:max_chars]


def _count_empty_type_rows(rows: list[list[str]]) -> int:
    count = 0
    for row in rows:
        if len(row) < 3:
            continue
        name, ident, c_type = row[0].strip(), row[1].strip(), row[2].strip()
        if not name or not ident:
            continue
        if name in {"名称", "CSC 名称", "CSU 名称"} or ident in {"标识", "CSC 标识", "CSU 标识"}:
            continue
        if not c_type:
            count += 1
    return count


# Suffixes/prefixes that legitimately appear inside interrupt-handler
# function names.  The audit should not flag ``ISR_*``, ``*_ISR``,
# ``*_INT`` etc.  as raw macros.
_INTERRUPT_TOKEN_HINTS = (
    "ISR", "INT", "IRQ", "NMI", "FIQ", "XINT", "RXINT", "TXINT",
    "ACK", "VECT", "HANDLER", "CALLBACK",
    # Common DSP / SoC peripherals whose names contain only uppercase letters
    "PIE", "ECAP", "EQEP", "EPWM", "I2C", "UART", "USB", "EMIF", "GPIO",
    "DMA", "ADC", "DAC", "TIMER", "WD", "WDT", "XINTF",
)


def _looks_like_interrupt_token(token: str) -> bool:
    if not token:
        return False
    parts = [p for p in token.split("_") if p]
    if not parts:
        return False
    # Direct match against the allow list
    if any(part in _INTERRUPT_TOKEN_HINTS for part in parts):
        return True
    # ``ISR_*`` / ``*_ISR`` patterns
    if parts[0] == "ISR" or parts[-1] == "ISR":
        return True
    return False


def _raw_macro_hits(text: str) -> list[str]:
    allow = {
        "CPU",
        "DSP",
        "ELSE",
        "END",
        "FPGA",
        "IF",
        "IFBIT",
        "PBIT",
        "PUBIT",
        "MBIT",
        "RIU",
        "ICD",
        "SCI",
        "SPI",
        "CAN",
        "CRC",
        "CASE",
        "DEFAULT",
        "NEXT",
        "SWITCH",
        "WHILE",
    }
    hits = []
    for token in re.findall(r"\b[A-Z][A-Z0-9_]{3,}\b", text or ""):
        if token in allow:
            continue
        if re.fullmatch(r"^(?:D[/_])?R_SDD\d+(?:_\d+)?$", token):
            continue
        if _looks_like_interrupt_token(token):
            continue
        hits.append(token)
    return sorted(set(hits))[:20]


def check_docx(path: Path, sample: Sample) -> tuple[int, list[str], int, int, int, str, dict[str, Any]]:
    warnings: list[str] = []
    details: dict[str, Any] = {
        "bad_terms": [],
        "missing_expected_outputs": [],
        "empty_type_count": 0,
        "quality_issues": [],
        "docx_leak_hits": [],
        "effect_stats": {},
    }
    if not path.exists():
        details["quality_issues"].append({"code": "docx_missing", "severity": "error", "message": "docx 不存在"})
        return 0, ["docx 不存在"], 0, 0, 0, "", details
    size = path.stat().st_size
    doc = Document(str(path))
    text = doc_text(doc)
    paragraph_text = "\n".join(p.text for p in doc.paragraphs if p.text)
    rows = _doc_table_rows(doc)
    row_text = "\n".join("\t".join(row) for row in rows)
    external_rows = _table_data_rows(doc, ("名称", "标识", "操作", "来源"))
    return_rows = _table_data_rows(doc, ("返回表达式", "含义", "成立条件"))
    direct_rows = [row for row in external_rows if len(row) > 3 and row[3] == sample.func_name]
    callee_rows = [row for row in external_rows if row not in direct_rows]
    details["effect_stats"] = {
        "document_direct_effects": len(direct_rows),
        "document_callee_effects": len(callee_rows),
        "document_return_semantics": len(return_rows),
        "expected_direct_effects": len(sample.expected_external_effects or ()),
        "expected_callee_effects": len(sample.expected_callee_effects or ()),
        "expected_return_effects": len(sample.expected_return_effects or ()),
        "unresolved_callee_effects": int(sample.expected_unresolved_callee_effects or 0),
    }
    score = 100
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    if size < 20000:
        warnings.append(f"docx 过小：{size} bytes")
        score -= 15
    if paragraphs < 10:
        warnings.append(f"段落数偏少：{paragraphs}")
        score -= 10
    if tables < 1:
        warnings.append(f"表格数偏少：{tables}")
        score -= 10
    if sample.func_name not in text:
        warnings.append("未找到原始函数名")
        details["quality_issues"].append({"code": "function_name_missing", "severity": "error", "message": "未找到原始函数名"})
        score -= 20
    docx_leak_hits = [term for term in DOCX_LEAK_TERMS if term in text]
    if docx_leak_hits:
        details["docx_leak_hits"] = docx_leak_hits
        warnings.append(f"包含内部审计字段：{', '.join(docx_leak_hits)}")
        details["quality_issues"].append({"code": "docx_internal_field_leak", "severity": "error", "message": ", ".join(docx_leak_hits)})
        score -= 35
    bad_terms = [term for term in BAD_DOC_TERMS if term in text]
    raw_macro_hits = _raw_macro_hits(paragraph_text)
    if raw_macro_hits:
        bad_terms.extend(f"raw_macro:{x}" for x in raw_macro_hits)
    details["bad_terms"] = list(dict.fromkeys(bad_terms))
    for bad in details["bad_terms"]:
        warnings.append(f"包含质量风险文本：{bad}")
        details["quality_issues"].append({"code": "bad_term", "severity": "warning", "message": bad})
        score -= 4 if bad.startswith("raw_macro:") else 18
    for issue in _logic_quality_issues_from_text(text):
        details["quality_issues"].append(issue)
        warnings.append(str(issue["message"]))
        score -= 25
    # === 新增: 语义盲点检测 (L1: 描述/代码 token 脱钩) ===
    try:
        from autodoc import backend as _backend
        c_src = _backend.load_c_file(sample.c_file)
        # 提取"功能说明"段: b) 之后到下一个 a)/c)/d) 之前
        m = re.search(r"b\)\s*功能说明\s*\n(.*?)(?=\n[cde]\)|\n表\s*\d|\Z)", text, re.S)
        desc = m.group(1).strip() if m else ""
        # 提取源码中所有 C 标识符 (>=4 字符), 过滤常见词
        c_idents = set(re.findall(r"\b[A-Za-z_][A-Za-z0-9_]{3,}\b", c_src))
        stop = {"void","static","const","unsigned","int","char","short","long","float","double","return","while","switch","case","break","continue","else","sizeof","typedef","struct","union","enum","extern","register"}
        c_idents -= stop
        # 描述里的中文必须能对到至少 1 个源码标识符的子串翻译
        # 简化: 看描述里是否有 c_idents 中任一标识符的连续 4+ 字符子串
        overlap = 0
        for ident in c_idents:
            if len(ident) >= 5 and ident in text:
                overlap += 1
            elif len(ident) >= 6:
                # 翻译可能拆开, 也允许部分匹配
                for chunk in (ident[:5], ident[-5:]):
                    if chunk in text:
                        overlap += 1
                        break
        # 描述段独立性更强: 描述里的汉字是否能拆出 c_ident 的子串
        if desc and c_idents:
            desc_overlap = sum(1 for ident in c_idents if len(ident) >= 5 and (ident in desc or any(ident[i:i+5] in desc for i in range(len(ident)-4))))
            if desc_overlap == 0 and len(desc) > 10:
                # 描述里完全没出现任何源码标识符 — 强信号 (可能 SciParitySet 类)
                # 但允许: 描述含明显中文术语 (>=3 个 CJK 短语)
                cjk_phrases = re.findall(r"[\u4e00-\u9fff]{3,}", desc)
                if len(cjk_phrases) >= 2:
                    # 进一步: 看描述里"汇总监测" "更新" 等通用动词 + 监测/状态 等通用名词
                    generic_markers = ("汇总监测", "监测结果", "执行功能", "完成相关", "执行操作", "处理并更新", "通用处理")
                    if any(m in desc for m in generic_markers):
                        warnings.append("功能描述疑似与代码无关 (含通用模板短语)")
                        details["quality_issues"].append({"code": "desc_generic_template", "severity": "error", "message": desc[:60]})
                        score -= 30
    except Exception:
        pass
    # === 新增: 循环索引变量名翻译一致性 (L2) ===
    # AI 改写时偶尔把 l_ii_u8 翻译为业务名词 (如"查询扇区号"), 与"循环索引"基础翻译冲突
    try:
        from autodoc.semantic_registry import lookup_var_role
        c_src_local2 = c_src if 'c_src' in dir() else _backend.load_c_file(sample.c_file)
        seen_drift = set()
        for var in set(re.findall(r"\b[lL]_[A-Za-z_][A-Za-z0-9_]*", c_src_local2)):
            role = lookup_var_role(var)
            if role != "loop_index" or len(var) < 4:
                continue
            bad_windows = []
            for m in re.finditer(re.escape(var), text):
                start = max(0, m.start() - 30)
                end = min(len(text), m.end() + 30)
                window = text[start:end]
                if re.search(r"[\u4e00-\u9fff]{2,}", window) and not re.search(
                    r"(索引|下标|计数|循环|遍历)", window
                ):
                    bad_windows.append(window)
            if bad_windows and var not in seen_drift:
                seen_drift.add(var)
                example = bad_windows[0]
                warnings.append(f"循环索引变量 {var} 翻译为业务名词: ...{example}...")
                details["quality_issues"].append({
                    "code": "loop_index_translation_drift",
                    "severity": "warning",
                    "message": f"{var}: {example}",
                })
                score -= 8
    except Exception:
        pass
    # === 新增: 逻辑步骤顺序与源码控制流一致性 (L3) ===
    try:
        # 提取源码中 if/for/while/switch 行的函数名/调用 (粗略标识)
        c_src_local = c_src if 'c_src' in dir() else _backend.load_c_file(sample.c_file)
        # 收集"形如 FuncName(...)" 调用和"形如 FUNC_NAME" 标识
        # 取函数体内前 200 行
        body_lines = []
        depth = 0
        started = False
        for ln in c_src_local.splitlines():
            depth += ln.count("{") - ln.count("}")
            if "{" in ln and not started:
                started = True
            if started and depth <= 0 and body_lines:
                break
            body_lines.append(ln)
        body_text = "\n".join(body_lines)
        # 提取 "FUNC_NAME(" 独立语句调用 (>=4 字符, 行首或 ; 之后开始)
        # 排除 #define 常量 (全大写 >=5)
        calls = []
        for line in body_text.splitlines():
            stripped = line.strip()
            # 跳过注释行/预处理行
            if not stripped or stripped.startswith(("//", "#", "/*", "*")):
                continue
            # 独立调用: 行首为 "  FooBar(" 或 "; FooBar(" 或 "{ FooBar("
            m = re.match(r"^[\{{;\s]*([A-Za-z_][A-Za-z0-9_]+)\s*\(", stripped)
            if m and not (m.group(1).isupper() and len(m.group(1)) >= 5):
                calls.append(m.group(1))
        calls = list(set(calls))
        # 提取"逻辑/流程图"段
        m_logic = re.search(r"e\)\s*逻辑/流程图\s*\n(.*?)(?=\n[A-Z]|\n表\s*\d|\Z)", text, re.S)
        if m_logic:
            logic_text = m_logic.group(1)
            logic_calls = set(re.findall(r"\b([A-Za-z_][A-Za-z0-9_]{4,})\b", logic_text))
            # 源码中关键调用 (在逻辑段中完全没出现) — 这才是盲点
            # 排除常见无关词
            stop = {"void","while","switch","sizeof","memset","memcpy","strcpy","strlen","printf","malloc","free","assert","return","const","static","unsigned","int"}
            meaningful_calls = [c for c in calls if c not in stop and not c.startswith(sample.func_name) and not c[0].isdigit()]
            missing = [c for c in meaningful_calls if c not in logic_calls and len(c) >= 5]
            # 阈值: missing >= 5 且 missing/meaningful >= 0.4 (占 40% 以上才算明显遗漏)
            if len(missing) >= 5 and len(meaningful_calls) >= 5 and len(missing) / max(1, len(meaningful_calls)) >= 0.4:
                warnings.append(f"逻辑步骤疑似遗漏 {len(missing)} 个源码调用: {missing[:3]}")
                details["quality_issues"].append({"code": "logic_flow_missing_calls", "severity": "warning", "message": str(missing[:5])})
                score -= min(20, len(missing) * 2)
    except Exception:
        pass
    if re.search(r"\b(?:\d*bit\d+|[A-Za-z_][A-Za-z0-9_]*bit\d+[A-Za-z0-9_]*)[\u4e00-\u9fff]", text, flags=re.IGNORECASE):
        warnings.append("疑似位号别名污染中文逻辑")
        details["quality_issues"].append({"code": "bit_alias_pollution", "severity": "warning", "message": "疑似位号别名污染中文逻辑"})
        score -= 18
    if re.search(r"\bIF\b[^\n]{0,120}(?:bit\d+\(|\d+bit\d+)", text, flags=re.IGNORECASE):
        warnings.append("IF 条件疑似残留 bit 别名")
        details["quality_issues"].append({"code": "bit_alias_in_condition", "severity": "warning", "message": "IF 条件疑似残留 bit 别名"})
        score -= 18
    if re.search(r"([\u4e00-\u9fff]{2,6})\1{1,}", text):
        warnings.append("疑似中文短语重复叠加")
        details["quality_issues"].append({"code": "repeated_cn_phrase", "severity": "warning", "message": "疑似中文短语重复叠加"})
        score -= 10
    table_nums = re.findall(r"表\s*([0-9]+)", text)
    if table_nums.count("1") > 1:
        warnings.append("疑似表编号重复：表 1 多次出现")
        score -= 6
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    if cjk < 100:
        warnings.append(f"中文内容偏少：{cjk} 字")
        score -= 15
    if "局部" not in text and "变量" not in text:
        warnings.append("未明显找到局部变量相关内容")
        score -= 8
    if "逻辑" not in text and "流程" not in text:
        warnings.append("未明显找到逻辑/流程相关内容")
        score -= 8
    if sample.expected_outputs:
        if re.search(r"输入/输出元素\s*\n\s*无。", text):
            warnings.append("源码存在输出事实，但文档输入/输出元素为无")
            details["quality_issues"].append({"code": "io_table_false_none", "severity": "error", "message": "源码存在输出事实，但文档输入/输出元素为无"})
            score -= 25
        for ident in sample.expected_outputs:
            root = _root_ident(ident)
            if ident not in row_text and (not root or root not in row_text):
                details["missing_expected_outputs"].append(ident)
                warnings.append(f"源码输出未进入输入/输出表：{ident}")
                details["quality_issues"].append({"code": "missing_expected_output", "severity": "error", "message": ident})
                score -= 14
    external_idents = {row[1] for row in external_rows if len(row) > 1}
    for ident in sample.expected_external_effects or ():
        if ident not in external_idents:
            warnings.append(f"源码外部副作用未进入副作用表：{ident}")
            details["quality_issues"].append({"code": "direct_effect_missing", "severity": "error", "message": ident})
            score -= 14
    for ident in sample.expected_callee_effects or ():
        if ident not in external_idents:
            warnings.append(f"已解析被调副作用未进入副作用表：{ident}")
            details["quality_issues"].append({"code": "callee_effect_missing", "severity": "error", "message": ident})
            score -= 14
    return_exprs = {row[0] for row in return_rows if row}
    for expr in sample.expected_return_effects or ():
        if expr not in return_exprs:
            warnings.append(f"源码返回分支未进入返回值语义表：{expr}")
            details["quality_issues"].append({"code": "return_semantic_missing", "severity": "error", "message": expr})
            score -= 14
    details["empty_type_count"] = _count_empty_type_rows(rows)
    if details["empty_type_count"] > 0:
        warnings.append(f"输入/输出或局部数据表存在空类型：{details['empty_type_count']}")
        details["quality_issues"].append({"code": "empty_type", "severity": "warning", "message": str(details["empty_type_count"])})
        score -= min(20, details["empty_type_count"] * 4)
    return max(0, score), warnings, size, paragraphs, tables, _excerpt(text, warnings), details


def apply_golden_audit(path: Path, sample: Sample, profile: dict[str, Any], warnings: list[str], details: dict[str, Any]) -> int:
    expectations = revision_utils.find_golden_expectations(profile, sample.c_file, sample.func_name)
    if not expectations or not path.exists():
        return 0
    try:
        text = doc_text(Document(str(path)))
    except Exception as exc:
        issue = {"code": "golden_docx_read_failed", "severity": "error", "message": repr(exc)}
        details.setdefault("quality_issues", []).append(issue)
        warnings.append(f"黄金样本审计读取失败：{exc}")
        return 25
    issues = revision_utils.audit_golden_text(text, expectations)
    if not issues:
        return 0
    details.setdefault("quality_issues", []).extend(dict(item) for item in issues)
    for item in issues:
        warnings.append(f"黄金样本审计失败：{item.get('message', '')}")
    return min(40, 12 * len(issues))


def write_reports(results: list[CheckResult], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    json_path = output_dir / "report.json"
    md_path = output_dir / "report.md"
    json_path.write_text(json.dumps([asdict(r) for r in results], ensure_ascii=False, indent=2), encoding="utf-8")
    lines = ["# 随机函数文档生成检查报告", ""]
    ok_count = sum(1 for r in results if r.ok)
    avg_score = round(sum(r.score for r in results) / max(1, len(results)), 1)
    issue_counts = {
        code: sum(1 for r in results for item in r.quality_issues if item.get("code") == code)
        for code in ("logic_placeholder", "logic_truncated", "bad_term")
    }
    effect_totals = {
        key: sum(int((r.effect_stats or {}).get(key) or 0) for r in results)
        for key in ("document_direct_effects", "document_callee_effects", "document_return_semantics", "expected_direct_effects", "expected_callee_effects", "expected_return_effects", "unresolved_callee_effects")
    }
    ai_event_counts = {
        action: sum(1 for r in results for item in r.ai_quality_events if item.get("action") == action)
        for action in ("ai_timeout", "targeted_ai_repaired", "line_deterministic_fallback")
    }
    lines.extend([
        f"- 样本数：{len(results)}",
        f"- 成功数：{ok_count}",
        f"- 平均评分：{avg_score}",
        f"- 逻辑占位：{issue_counts['logic_placeholder']}",
        f"- 逻辑截断：{issue_counts['logic_truncated']}",
        f"- 一般风险词：{issue_counts['bad_term']}",
        f"- AI 超时：{ai_event_counts['ai_timeout']}",
        f"- AI 定向修复：{ai_event_counts['targeted_ai_repaired']}",
        f"- 确定性逐行回退：{ai_event_counts['line_deterministic_fallback']}",
        f"- 直接副作用覆盖：{effect_totals['document_direct_effects']}/{effect_totals['expected_direct_effects']}",
        f"- 被调副作用覆盖：{effect_totals['document_callee_effects']}/{effect_totals['expected_callee_effects']}",
        f"- 返回语义覆盖：{effect_totals['document_return_semantics']}/{effect_totals['expected_return_effects']}",
        f"- 未确认被调副作用：{effect_totals['unresolved_callee_effects']}",
        "",
    ])
    for i, r in enumerate(results, 1):
        s = r.sample
        lines.extend([
            f"## {i}. `{s['func_name']}`",
            f"- 文件：`{s['c_file']}`",
            f"- 输出：`{r.output}`",
            f"- 结果：{'OK' if r.ok else 'FAIL'}，评分：{r.score}，耗时：{r.elapsed_sec:.1f}s",
            f"- docx：{r.docx_size} bytes，段落 {r.paragraphs}，表格 {r.tables}",
        ])
        if r.error:
            lines.append(f"- 错误：`{r.error}`")
        if r.warnings:
            lines.append("- 警告：")
            lines.extend(f"  - {w}" for w in r.warnings)
        if r.bad_terms:
            lines.append("- 坏词/风险文本：")
            lines.extend(f"  - `{x}`" for x in r.bad_terms)
        if r.docx_leak_hits:
            lines.append("- DOCX 内部字段泄露：")
            lines.extend(f"  - `{x}`" for x in r.docx_leak_hits)
        if r.missing_expected_outputs:
            lines.append("- 缺失源码输出：")
            lines.extend(f"  - `{x}`" for x in r.missing_expected_outputs)
        if r.effect_stats:
            lines.append(f"- 效果覆盖：直接 {r.effect_stats.get('document_direct_effects', 0)}/{r.effect_stats.get('expected_direct_effects', 0)}，被调 {r.effect_stats.get('document_callee_effects', 0)}/{r.effect_stats.get('expected_callee_effects', 0)}，返回 {r.effect_stats.get('document_return_semantics', 0)}/{r.effect_stats.get('expected_return_effects', 0)}")
        lines.append(f"- 空类型计数：{r.empty_type_count}")
        if r.quality_issues:
            lines.append("- 质量问题：")
            for item in r.quality_issues:
                code = str(item.get("code") or "")
                severity = str(item.get("severity") or "")
                message = str(item.get("message") or "")
                lines.append(f"  - `{severity}` `{code}` {message}")
        if r.ai_quality_events:
            lines.append("- AI 质量恢复：")
            lines.extend(f"  - `{item.get('action', '')}` 行 {item.get('lines', ())}" for item in r.ai_quality_events)
        if r.judge_total_25 or r.judge_grade or r.judge_skipped:
            if r.judge_skipped:
                lines.append(f"- LLM-judge: 跳过 ({r.judge_skipped})")
            else:
                judge_flag = " 🚩" if r.judge_total_25 < 20 else ""
                lines.append(f"- LLM-judge: {r.judge_total_25}/25, grade={r.judge_grade}, {r.judge_elapsed_sec:.1f}s{judge_flag}")
                if r.judge_red_flags:
                    lines.append("  - red_flags:")
                    lines.extend(f"    - {f}" for f in r.judge_red_flags[:5])
        expected_outputs = s.get("expected_outputs") or []
        if expected_outputs:
            lines.append("- 源码期望输出：")
            lines.extend(f"  - `{x}`" for x in expected_outputs)
        if r.text_excerpt:
            lines.extend(["- 文档摘录：", "```text", r.text_excerpt[:900], "```"])
        lines.append("")
    md_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"报告：{md_path}")
    print(f"JSON：{json_path}")


def main() -> int:
    parser = argparse.ArgumentParser(description="随机抽取真实项目 C 函数生成文档并做粗略质量检查")
    parser.add_argument("--project-dir", default=os.environ.get("AUTODOCGEN_TEST_PROJECT", ""),
                        help="测试 C 项目目录（可用 AUTODOCGEN_TEST_PROJECT 环境变量预设）")
    parser.add_argument("--ini", default=str(ROOT / "autodocgen.ini"))
    parser.add_argument("--output-dir", default=os.environ.get("AUTODOCGEN_CHECK_OUTPUT", "tmp/random_check"),
                        help="检查报告输出目录（可用 AUTODOCGEN_CHECK_OUTPUT 环境变量预设）")
    parser.add_argument("--samples", type=int, default=3)
    parser.add_argument("--seed", type=int, default=0)
    parser.add_argument("--max-files", type=int, default=0)
    parser.add_argument("--no-ai", action="store_true")
    parser.add_argument("--ai-workers", type=int, default=1,
                        help="AI 并发数（默认 1 避免频率限制）")
    parser.add_argument("--ai-api-base", default="", help="覆盖 AI API endpoint/base URL")
    parser.add_argument("--ai-model", default="", help="覆盖 AI model")
    parser.add_argument("--ai-api-key-env", default="", help="从环境变量读取 API key，避免写入命令行或报告")
    parser.add_argument("--wire-api", default="", help="覆盖 wire_api：chat_completions/completions/responses")
    parser.add_argument("--ai-mode", type=int, default=None, help="覆盖 AI 开关：0=无 AI，非零=开启 AI")
    parser.add_argument("--revision-profile", default="", help="函数级修订档案 JSON 路径，同时启用其中的 golden 审计")
    parser.add_argument("--judge", action="store_true", help="生成后对每个 docx 跑 LLM-judge 评估，写入 judge_total_25/grade/red_flags")
    parser.add_argument("--judge-model", default="", help="覆盖 judge 模型 (默认复用 --ai-model)")
    parser.add_argument("--judge-api-base", default="", help="覆盖 judge API base (默认复用 --ai-api-base)")
    parser.add_argument("--judge-api-key-env", default="", help="覆盖 judge API key env (默认复用 --ai-api-key-env)")
    parser.add_argument("--judge-min", type=int, default=20, help="judge_total_25 低于该值视为低质量, 入修复队列")
    args = parser.parse_args()
    project_dir = Path(args.project_dir).resolve()
    output_dir = Path(args.output_dir).resolve()
    cfg = load_cfg(Path(args.ini), no_ai=args.no_ai, ai_workers=args.ai_workers)
    if args.ai_api_base:
        cfg.ai_api_base = args.ai_api_base
    if args.ai_model:
        cfg.ai_model = args.ai_model
    if args.ai_api_key_env:
        cfg.ai_api_key = os.environ.get(args.ai_api_key_env, "")
        cfg.ai_use_auth = bool(cfg.ai_api_key)
    if args.wire_api and hasattr(cfg, "wire_api"):
        cfg.wire_api = args.wire_api
        cfg.extra_params = dict(cfg.extra_params or {})
        cfg.extra_params["wire_api"] = args.wire_api
    if args.ai_mode is not None:
        cfg.ai_mode = 0 if args.no_ai else _ai_mode(args.ai_mode)
        cfg.ai_assist = (cfg.ai_mode == 1)
    if args.revision_profile:
        cfg.extra_params = dict(cfg.extra_params or {})
        cfg.extra_params["revision_profile"] = args.revision_profile
    revision_profile = revision_utils.load_revision_profile(cfg)
    cfg.project_root = str(project_dir)
    cfg.open_after_done = False
    cfg.verbose = True
    active_ai_events: list[dict[str, Any]] = []

    def _gui_log(msg: str) -> None:
        print(f"[LOG] {msg}")
        if "timed out" in str(msg).lower() or "超时" in str(msg):
            active_ai_events.append({"action": "ai_timeout"})

    def _gui_event(event: dict[str, Any]) -> None:
        if isinstance(event, dict) and event.get("type") == "ai_quality_recovery":
            active_ai_events.append(dict(event))

    cfg.gui_log = _gui_log
    cfg.gui_event = _gui_event
    samples = collect_samples(project_dir, cfg, count=args.samples, seed=args.seed, max_files=args.max_files)
    output_dir.mkdir(parents=True, exist_ok=True)
    results: list[CheckResult] = []
    for idx, sample in enumerate(samples, 1):
        active_ai_events.clear()
        safe_func = re.sub(r"[^A-Za-z0-9_]+", "_", sample.func_name)[:80]
        out = output_dir / f"{idx:02d}_{Path(sample.c_file).stem}_{safe_func}.docx"
        print(f"[{idx}/{len(samples)}] 生成 {sample.func_name} <- {sample.c_file}")
        start = time.time()
        error = ""
        generated = False
        try:
            backend.generate_design_doc_for_single_function(sample.c_file, sample.func_name, str(out), cfg, project_root=str(project_dir))
            generated = True
        except Exception as exc:
            error = repr(exc)
        elapsed = time.time() - start
        if generated:
            score, warnings, size, paragraphs, tables, excerpt, details = check_docx(out, sample)
            penalty = apply_golden_audit(out, sample, revision_profile, warnings, details)
            if penalty:
                score = max(0, score - penalty)
            judge_total, judge_grade, judge_flags, judge_elapsed, judge_skip = (0, "", [], 0.0, "")
            if args.judge and generated and not error:
                judge_total, judge_grade, judge_flags, judge_elapsed, judge_skip = _run_inline_judge(
                    out, sample, args,
                )
        else:
            details = {
                "bad_terms": [],
                "missing_expected_outputs": [],
                "empty_type_count": 0,
                "quality_issues": [],
                "docx_leak_hits": [],
            }
            score, warnings, size, paragraphs, tables, excerpt = 0, [], 0, 0, 0, ""
        if error:
            warnings.append(error)
            details["quality_issues"].append({"code": "generation_error", "severity": "error", "message": error})
        hard_fail = bool(details.get("docx_leak_hits") or details.get("missing_expected_outputs"))
        hard_fail = hard_fail or any((item.get("severity") == "error") for item in (details.get("quality_issues") or []))
        results.append(
            CheckResult(
                asdict(sample),
                str(out),
                generated and score >= 60 and not hard_fail,
                elapsed,
                score,
                warnings,
                error,
                size,
                paragraphs,
                tables,
                excerpt,
                list(details.get("bad_terms") or []),
                list(details.get("missing_expected_outputs") or []),
                int(details.get("empty_type_count") or 0),
                list(details.get("quality_issues") or []),
                list(details.get("docx_leak_hits") or []),
                int(judge_total) if "judge_total" in dir() else 0,
                judge_grade if "judge_grade" in dir() else "",
                list(judge_flags) if "judge_flags" in dir() else [],
                float(judge_elapsed) if "judge_elapsed" in dir() else 0.0,
                judge_skip if "judge_skip" in dir() else "",
                ai_quality_events=list(active_ai_events),
                effect_stats=dict(details.get("effect_stats") or {}),
            )
        )
    write_reports(results, output_dir)
    failed = [r for r in results if not r.ok]
    if failed:
        print(f"失败/低分：{len(failed)}/{len(results)}")
    if args.judge:
        judged = [r for r in results if r.judge_total_25]
        if judged:
            avg = sum(r.judge_total_25 for r in judged) / len(judged)
            lowq = [r for r in judged if r.judge_total_25 < args.judge_min]
            print(f"LLM-judge: 平均 {avg:.1f}/25, 样本 {len(judged)}/{len(results)}")
            if lowq:
                print(f"judge 低质量(<{args.judge_min}/25): {len(lowq)}/{len(judged)}")
                for r in lowq:
                    print(f"  - {r.sample['func_name']:35s} judge={r.judge_total_25}/25 ({r.judge_grade})")
        return 1 if failed else 0
    if failed:
        return 1
    print("全部样本通过粗检")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
