from __future__ import annotations

import os
import json
import queue
import re
import time
import sys
from urllib.parse import urlsplit, urlunsplit
import requests
from typing import Optional

from PyQt5 import QtCore, QtGui, QtWidgets

try:
    from .runner import DocUpdateWorker, ExportFuncWorker, GenerateWorker, RegenerateCsuBatchWorker, RegenerateCsuWorker, StepDef, TaskSpec, TermTableWorker, UpdateCsuWorker
    from .settings_store import AppSettings, SettingsStore, LOCAL_LLM_API_BASE, normalize_ai_mode
except ImportError:  # allow running as a script: python qt_gui/main_window.py
    from qt_gui.runner import DocUpdateWorker, ExportFuncWorker, GenerateWorker, RegenerateCsuBatchWorker, RegenerateCsuWorker, StepDef, TaskSpec, TermTableWorker, UpdateCsuWorker
    from qt_gui.settings_store import AppSettings, SettingsStore, LOCAL_LLM_API_BASE, normalize_ai_mode


_STATUS_TEXT = {
    "pending": "等待",
    "running": "运行中",
    "success": "成功",
    "failed": "失败",
    "stopped": "已停止",
}


def _wrap_layout(layout: QtWidgets.QLayout) -> QtWidgets.QWidget:
    w = QtWidgets.QWidget()
    w.setLayout(layout)
    return w


def _asset_path(*parts: str) -> str:
    base_dir = os.path.dirname(__file__)
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        base_dir = os.path.join(sys._MEIPASS, "qt_gui")
    return os.path.join(base_dir, "assets", *parts)


class _NumericTableWidgetItem(QtWidgets.QTableWidgetItem):
    def __lt__(self, other) -> bool:
        try:
            left = float(self.data(QtCore.Qt.UserRole) or 0.0)
            right = float(other.data(QtCore.Qt.UserRole) or 0.0)
            return left < right
        except Exception:
            return super().__lt__(other)


class MainWindow(QtWidgets.QMainWindow):
    def __init__(self, *, backend, settings_store: SettingsStore) -> None:
        super().__init__()
        self.backend = backend
        self.settings_store = settings_store
        self.settings = settings_store.load()
        self._app_version = str(getattr(backend, "APP_VERSION", "") or "").strip()
        self._title_suffix = f" {self._app_version}" if self._app_version else ""
        self._store_unlocked = False

        self._thread: Optional[QtCore.QThread] = None
        self._worker: Optional["_QtWorker"] = None
        self._resume_state: Optional[dict] = None
        self._pending_continue: Optional[bool] = None
        self._last_task: Optional[TaskSpec] = None
        self._last_settings: Optional[AppSettings] = None
        self._progress_state = {"total": None, "done": 0, "start_ts": None}
        self._failed_functions: list[dict] = []  # 函数失败记录
        self._closing = False
        self._gui_event_queue: "queue.Queue[dict]" = queue.Queue()
        self._gui_log_queue: "queue.Queue[str]" = queue.Queue()

        self.setWindowTitle(f"CSCI 详细设计生成器{self._title_suffix}")
        self.resize(1280, 720)
        self.setMinimumWidth(1200)

        self._build_ui()
        # Settings 页较重，按需加载；这里不强制构建/应用
        self._set_running(False, note="就绪")

    def _set_store_unlocked(self, unlocked: bool) -> None:
        self._store_unlocked = bool(unlocked)
        try:
            self._nav_store.setVisible(self._store_unlocked)
        except Exception:
            pass
        if not self._store_unlocked:
            try:
                if self._stack.currentWidget() is self._page_store:
                    self._stack.setCurrentWidget(self._page_home)
            except Exception:
                pass

    def _check_store_secret(self) -> None:
        try:
            if isinstance(self.set_ai_model, QtWidgets.QComboBox):
                token = str(self.set_ai_model.currentText() or "").strip().lower()
            else:
                token = str(self.set_ai_model.text() or "").strip().lower()
        except Exception:
            return
        if token == "secret space":
            self._set_store_unlocked(True)
            try:
                if isinstance(self.set_ai_model, QtWidgets.QComboBox):
                    self.set_ai_model.setCurrentText("")
                else:
                    self.set_ai_model.setText("")
            except Exception:
                pass

    def _set_ai_model_items(self, items: list[str], *, preferred: str = "") -> None:
        if not hasattr(self, "set_ai_model"):
            return
        combo = self.set_ai_model
        if not isinstance(combo, QtWidgets.QComboBox):
            return
        current = preferred.strip() or str(combo.currentText() or "").strip() or str(getattr(self.settings, "ai_model", "") or "").strip()
        normalized = []
        seen = set()
        for raw in items or []:
            name = str(raw or "").strip()
            if not name or name in seen:
                continue
            seen.add(name)
            normalized.append(name)
        combo.blockSignals(True)
        combo.clear()
        combo.addItems(normalized)
        if current:
            if current not in normalized:
                combo.insertItem(0, current)
            combo.setCurrentText(current)
        combo.blockSignals(False)

    def _current_ai_base_url(self) -> str:
        if hasattr(self, "set_ai_base"):
            text = str(self.set_ai_base.text() or "").strip()
            if text:
                return text
        return str(getattr(self.settings, "ai_api_base", "") or LOCAL_LLM_API_BASE).strip()

    def _normalize_ai_base_for_gui(self, url: str) -> str:
        raw = str(url or "").strip().rstrip("/")
        if not raw:
            raw = LOCAL_LLM_API_BASE
        if raw and not re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*://", raw):
            raw = "http://" + raw
        return raw

    def _models_url_from_chat_url(self, url: str) -> str:
        raw = self._normalize_ai_base_for_gui(url)
        parts = urlsplit(raw)
        path = parts.path or ""
        if path.endswith("/chat/completions"):
            path = path[: -len("/chat/completions")] + "/models"
        elif path.endswith("/completions"):
            path = path[: -len("/completions")] + "/models"
        elif path.endswith("/models"):
            pass
        elif path.endswith("/v1") or path.endswith("/api/v1"):
            path = path.rstrip("/") + "/models"
        else:
            path = path.rstrip("/") + "/v1/models"
        return urlunsplit((parts.scheme or "http", parts.netloc, path, "", ""))

    def _fetch_local_model_names(self) -> list[str]:
        url = self._models_url_from_chat_url(self._current_ai_base_url())
        headers = {"Accept": "application/json"}
        api_key = ""
        if hasattr(self, "set_ai_key"):
            api_key = str(self.set_ai_key.text() or "").strip()
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        resp = requests.get(url, headers=headers, timeout=(3.0, 15.0))
        resp.raise_for_status()
        data = resp.json() if callable(getattr(resp, "json", None)) else {}
        items = []
        for item in (data.get("data") or []):
            if not isinstance(item, dict):
                continue
            model_id = str(item.get("id") or "").strip()
            if model_id:
                items.append(model_id)
        return items

    def _refresh_ai_models(self) -> None:
        if hasattr(self, "btn_refresh_ai_models"):
            self.btn_refresh_ai_models.setEnabled(False)
        try:
            models = self._fetch_local_model_names()
            self._set_ai_model_items(models)
            self._append_log(f"已刷新模型列表，共 {len(models)} 个")
        except Exception as e:
            self._append_log(f"刷新模型列表失败：{e}")
        finally:
            if hasattr(self, "btn_refresh_ai_models"):
                self.btn_refresh_ai_models.setEnabled(True)

    def _on_ai_mode_changed(self, index: int) -> None:
        enabled = normalize_ai_mode(index) == 1
        for name in (
            "set_ai_provider",
            "set_ai_model",
            "btn_refresh_ai_models",
            "set_ai_base",
            "set_ai_key",
            "set_ai_num_ctx",
            "set_ai_read_timeout",
            "set_ai_workers",
            "set_no_proxy",
        ):
            widget = getattr(self, name, None)
            if widget is not None:
                try:
                    widget.setEnabled(enabled)
                except Exception:
                    pass

    def _on_provider_changed(self, provider: str) -> None:
        provider = (provider or "").strip().lower()
        presets = {
            "local":   ("http://10.11.34.200:11434/v1", ""),
            "deepseek": ("https://api.deepseek.com/v1", ""),
            "openai":  ("https://api.openai.com/v1", ""),
            "openrouter": ("https://openrouter.ai/api/v1", ""),
            "anthropic": ("https://api.anthropic.com/v1", ""),
            "compshare": ("https://cp.compshare.cn/v1", ""),
        }
        if provider in presets:
            base, key = presets[provider]
            if not self.set_ai_base.text().strip() or self.set_ai_base.text().strip() in [v[0] for v in presets.values()]:
                self.set_ai_base.setText(base)
            if provider != "local":
                self.set_ai_key.setEchoMode(QtWidgets.QLineEdit.Password)
            else:
                self.set_ai_key.setEchoMode(QtWidgets.QLineEdit.Normal)
        if provider != "local":
            self._refresh_ai_models()

    @QtCore.pyqtSlot(str, str)
    def _on_step_event(self, step_id: str, status: str) -> None:
        self._set_step_status(step_id, status)

    @QtCore.pyqtSlot(str)
    def _on_log_event(self, text: str) -> None:
        self._append_log(text)

    def _enqueue_log_event(self, text: str) -> None:
        try:
            self._gui_log_queue.put(str(text or ""))
        except Exception:
            pass

    @QtCore.pyqtSlot(object)
    def _on_detail_event(self, payload) -> None:
        try:
            self._handle_detail_event(payload or {})
        except Exception:
            pass

    def _enqueue_detail_event(self, payload) -> None:
        try:
            self._gui_event_queue.put(dict(payload or {}))
        except Exception:
            pass

    @QtCore.pyqtSlot()
    def _flush_detail_events(self) -> None:
        for _ in range(200):
            try:
                text = self._gui_log_queue.get_nowait()
            except queue.Empty:
                break
            self._on_log_event(text)
        for _ in range(200):
            try:
                payload = self._gui_event_queue.get_nowait()
            except queue.Empty:
                break
            self._on_detail_event(payload)

    @QtCore.pyqtSlot(str)
    def _on_output_event(self, path: str) -> None:
        self.ed_output.setText(path)

    @QtCore.pyqtSlot(str, object, object)
    def _on_done_event(self, note: str, resume_state: Optional[dict], output_path: Optional[str]) -> None:
        self._resume_state = resume_state
        self._request_thread_finish()
        self._set_running(False, note=note or "就绪")
        if output_path:
            self._append_log(f"输出：{output_path}")
        QtCore.QTimer.singleShot(0, lambda: self._show_done_notification(note or "", output_path))

    def _show_done_notification(self, note: str, output_path: Optional[str]) -> None:
        if self._closing:
            return
        if self._failed_functions and not note.startswith("已停止"):
            self._show_failure_dialog(output_path)
        elif note.startswith("失败") or note.startswith("更新 CSU 标识失败"):
            QtWidgets.QMessageBox.critical(self, "失败", note)
        elif note.startswith("已停止"):
            msg = "生成已停止，可点击\"继续\"接力。"
            if self._failed_functions:
                msg = f"生成已停止（{len(self._failed_functions)} 个函数失败），可点击\"继续\"接力。"
            QtWidgets.QMessageBox.information(self, "已停止", msg)
        elif note.startswith("术语表已刷新"):
            QtWidgets.QMessageBox.information(self, "术语表", f"{note}\n\n{output_path or ''}".strip())
        else:
            QtWidgets.QMessageBox.information(self, "完成", note)

    def _show_failure_dialog(self, output_path: Optional[str]) -> None:
        """显示函数失败对话框。"""
        failures = self._failed_functions
        total = len(failures)

        # 构建失败信息
        msg = QtWidgets.QMessageBox(self)
        msg.setWindowTitle("生成完成（部分失败）")
        msg.setIcon(QtWidgets.QMessageBox.Warning)

        # 按错误类型分组统计
        error_types: dict[str, int] = {}
        for f in failures:
            et = f.get("error_type", "unknown")
            error_types[et] = error_types.get(et, 0) + 1

        type_names = {
            "ai_timeout": "AI 超时",
            "ai_parse_error": "AI 解析错误",
            "network_error": "网络错误",
            "ai_error": "AI 错误",
            "unknown": "未知错误",
        }

        type_summary = ", ".join(f"{type_names.get(k, k)}: {v}" for k, v in error_types.items())

        text = f"生成完成，{total} 个函数失败\n\n错误分布：{type_summary}\n\n失败函数：\n"
        for f in failures[:10]:  # 最多显示 10 个
            func_name = f.get("func_name", "未知")
            file_name = os.path.basename(f.get("file_path", "")) if f.get("file_path") else ""
            error_msg = f.get("error_message", "")[:50]
            text += f"  • {func_name}"
            if file_name:
                text += f" ({file_name})"
            text += f" - {error_msg}\n"
        if total > 10:
            text += f"  ... 还有 {total - 10} 个\n"

        msg.setText(text)

        # 添加按钮
        retry_btn = msg.addButton("重试失败函数", QtWidgets.QMessageBox.ActionRole)
        export_btn = msg.addButton("导出错误报告", QtWidgets.QMessageBox.ActionRole)
        ok_btn = msg.addButton("确定", QtWidgets.QMessageBox.AcceptRole)

        msg.exec()

        clicked = msg.clickedButton()
        if clicked == retry_btn:
            self._retry_failed_functions()
        elif clicked == export_btn:
            self._export_failure_report(output_path)

    def _retry_failed_functions(self) -> None:
        """重试失败的函数。"""
        if not self._failed_functions:
            return

        # 清空失败记录，准备重试
        failures = self._failed_functions
        self._failed_functions = []

        # 记录日志
        self._append_log(f"重试 {len(failures)} 个失败函数...")

        # TODO: 实现重试逻辑 - 需要重新构建任务并执行
        # 当前简化实现：提示用户重新生成
        QtWidgets.QMessageBox.information(
            self,
            "重试",
            f"将重试 {len(failures)} 个失败函数。\n请点击\"生成\"按钮重新执行。",
        )

    def _export_failure_report(self, output_path: Optional[str]) -> None:
        """导出失败报告。"""
        import json
        from datetime import datetime

        failures = self._failed_functions
        report = {
            "timestamp": datetime.now().isoformat(),
            "total_failures": len(failures),
            "failures": failures,
        }

        # 选择保存路径
        if output_path:
            default_path = output_path.replace(".docx", "_failures.json")
        else:
            default_path = "function_failures.json"

        path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self,
            "保存错误报告",
            default_path,
            "JSON 文件 (*.json)",
        )

        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(report, f, ensure_ascii=False, indent=2)
                self._append_log(f"错误报告已保存：{path}")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "保存失败", f"无法保存报告：{e}")

    def _build_ui(self) -> None:
        self._stack = QtWidgets.QStackedWidget()

        # Web-app style wrapper (sidebar + topbar + content)
        self._central = QtWidgets.QWidget()
        root = QtWidgets.QHBoxLayout(self._central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        self._sidebar = QtWidgets.QFrame()
        self._sidebar.setObjectName("sidebar")
        self._sidebar.setMinimumWidth(200)
        sb = QtWidgets.QVBoxLayout(self._sidebar)
        sb.setContentsMargins(14, 14, 14, 14)
        sb.setSpacing(6)
        title = QtWidgets.QLabel("AutoDocGen")
        title.setObjectName("sidebar_title")
        sb.addWidget(title)
        if self._app_version:
            ver = QtWidgets.QLabel(self._app_version)
            ver.setObjectName("sidebar_version")
            sb.addWidget(ver)

        self._nav_home = QtWidgets.QPushButton("主页")
        self._nav_home.setProperty("nav", True)
        self._nav_home.clicked.connect(lambda: self._stack.setCurrentWidget(self._page_home))
        self._nav_store = QtWidgets.QPushButton("应用市场")
        self._nav_store.setProperty("nav", True)
        self._nav_store.clicked.connect(lambda: self._stack.setCurrentWidget(self._page_store))
        self._nav_store.setVisible(False)
        self._nav_settings = QtWidgets.QPushButton("设置")
        self._nav_settings.setProperty("nav", True)
        self._nav_settings.clicked.connect(self._open_settings_page)
        self._nav_help = QtWidgets.QPushButton("帮助")
        self._nav_help.setProperty("nav", True)
        self._nav_help.clicked.connect(lambda: self._stack.setCurrentWidget(self._page_help))
        for b in (self._nav_home, self._nav_store, self._nav_settings, self._nav_help):
            b.setObjectName("nav_btn")
            b.setCursor(QtCore.Qt.PointingHandCursor)
            b.clicked.connect(self._sync_nav_active)
            sb.addWidget(b)

        sb.addStretch(1)

        self._main_area = QtWidgets.QWidget()
        main = QtWidgets.QVBoxLayout(self._main_area)
        main.setContentsMargins(16, 16, 16, 16)
        main.setSpacing(12)

        self._topbar = QtWidgets.QFrame()
        self._topbar.setObjectName("topbar")
        tb = QtWidgets.QHBoxLayout(self._topbar)
        tb.setContentsMargins(12, 10, 12, 10)
        tb.setSpacing(8)
        self._topbar_title = QtWidgets.QLabel("CSCI 详细设计生成器")
        self._topbar_title.setObjectName("topbar_title")
        tb.addWidget(self._topbar_title, 0)
        if self._app_version:
            ver_badge = QtWidgets.QLabel(self._app_version)
            ver_badge.setObjectName("version_badge")
            tb.addWidget(ver_badge, 0)
        tb.addStretch(1)

        self._status_label = QtWidgets.QLabel("就绪")
        self._status_label.setObjectName("status_pill")
        tb.addWidget(self._status_label, 0)

        self._progress = QtWidgets.QProgressBar()
        self._progress.setObjectName("progress_bar")
        self._progress.setFixedWidth(220)
        self._progress.setTextVisible(True)
        self._progress.setFormat("")
        self._progress.setRange(0, 1)
        self._progress.setValue(0)
        tb.addWidget(self._progress, 0)

        self._topbar_btn_start = QtWidgets.QToolButton()
        self._topbar_btn_stop = QtWidgets.QToolButton()
        self._topbar_btn_continue = QtWidgets.QToolButton()
        self._topbar_btn_update = QtWidgets.QToolButton()
        for w in (self._topbar_btn_start, self._topbar_btn_stop, self._topbar_btn_continue, self._topbar_btn_update):
            w.setToolButtonStyle(QtCore.Qt.ToolButtonTextOnly)
            w.setAutoRaise(False)
        tb.addWidget(self._topbar_btn_start)
        tb.addWidget(self._topbar_btn_stop)
        tb.addWidget(self._topbar_btn_continue)
        tb.addWidget(self._topbar_btn_update)

        top_splitter = QtWidgets.QSplitter(QtCore.Qt.Vertical)
        top_splitter.setObjectName("topbar_splitter")
        top_splitter.addWidget(self._topbar)
        top_splitter.addWidget(self._stack)
        top_splitter.setStretchFactor(0, 0)
        top_splitter.setStretchFactor(1, 1)
        top_splitter.setSizes([70, 600])
        main.addWidget(top_splitter, 1)

        splitter = QtWidgets.QSplitter(QtCore.Qt.Horizontal)
        splitter.setObjectName("layout_splitter")
        splitter.addWidget(self._sidebar)
        splitter.addWidget(self._main_area)
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        splitter.setSizes([220, 900])

        root.addWidget(splitter, 1)
        self.setCentralWidget(self._central)
        try:
            self.statusBar().hide()
        except Exception:
            pass

        self._project_tree = ProjectTreeView()
        self._project_tree.setHeaderHidden(True)
        self._project_tree.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self._project_tree.setUniformRowHeights(True)
        self._project_tree.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self._project_tree.customContextMenuRequested.connect(self._on_project_tree_menu)
        self._project_tree.removeRequested.connect(self._remove_selected_project_item)
        self._project_model = QtGui.QStandardItemModel()
        self._project_tree.setModel(self._project_model)
        self._init_project_model()

        self._step_tree = QtWidgets.QTreeView()
        self._step_tree.setHeaderHidden(True)
        self._step_tree.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self._step_tree.setUniformRowHeights(True)

        self._log_view = QtWidgets.QPlainTextEdit()
        self._log_view.setReadOnly(True)
        self._log_view.setMaximumBlockCount(5000)

        self._page_home = self._build_home_page()
        self._page_store = self._build_store_page()
        self._settings_page_built = False
        self._page_settings = QtWidgets.QWidget()
        ph = QtWidgets.QVBoxLayout(self._page_settings)
        ph.addStretch(1)
        ph.addWidget(QtWidgets.QLabel("设置页将在首次打开时加载..."))
        ph.addStretch(1)
        self._page_help = self._build_help_page()
        self._stack.addWidget(self._page_home)
        self._stack.addWidget(self._page_store)
        self._stack.addWidget(self._page_settings)
        self._stack.addWidget(self._page_help)
        self._stack.setCurrentWidget(self._page_home)
        self._sync_nav_active()
        self._stack.currentChanged.connect(lambda _i: self._sync_nav_active())

        self._build_actions()
        # Hide menu bar; topbar buttons still use the same actions.
        self.menuBar().setVisible(False)

        # Topbar uses the same actions (keeps enable/disable behavior consistent).
        self._topbar_btn_start.setDefaultAction(self.act_start)
        self._topbar_btn_stop.setDefaultAction(self.act_stop)
        self._topbar_btn_continue.setDefaultAction(self.act_continue)
        self._topbar_btn_update.setDefaultAction(self.act_update_csu)

        self._reset_steps([])
        QtCore.QTimer.singleShot(0, self._restore_window_state)

    def _sync_nav_active(self) -> None:
        cur = self._stack.currentWidget()
        try:
            for btn, page in (
                (self._nav_home, self._page_home),
                (self._nav_store, self._page_store),
                (self._nav_settings, self._page_settings),
                (self._nav_help, self._page_help),
            ):
                btn.setProperty("active", cur is page)
                btn.style().unpolish(btn)
                btn.style().polish(btn)
        except Exception:
            pass

    def _restore_window_state(self) -> None:
        try:
            geo, state = self.settings_store.load_window_layout("web")
            if geo:
                self.restoreGeometry(geo)
            if state:
                self.restoreState(state)
        except Exception:
            pass

    def _save_window_state(self) -> None:
        try:
            self.settings_store.save_window_layout("web", geometry=self.saveGeometry(), state=self.saveState())
        except Exception:
            pass

    def _init_project_model(self) -> None:
        self._project_model.clear()
        root = self._project_model.invisibleRootItem()

        self._proj_item_app = QtGui.QStandardItem("应用层")
        self._proj_item_mid = QtGui.QStandardItem("中间层")
        self._proj_item_drv = QtGui.QStandardItem("驱动层")

        for it in (self._proj_item_app, self._proj_item_mid, self._proj_item_drv):
            it.setEditable(False)
            it.setDragEnabled(False)
            it.setDropEnabled(True)
            it.setSelectable(False)
            root.appendRow(it)

        self._project_tree.expandAll()

    def _clear_project_children(self) -> None:
        for it in (self._proj_item_app, self._proj_item_mid, self._proj_item_drv):
            it.removeRows(0, it.rowCount())

    def _append_project_files(self, parent: QtGui.QStandardItem, files: list[str], *, base_dir: str) -> None:
        for path in files:
            abs_path = os.path.abspath(path)
            rel = None
            try:
                if base_dir:
                    rel = os.path.relpath(abs_path, base_dir)
            except Exception:
                rel = None
            text = rel if rel and (not rel.startswith("..")) else os.path.basename(abs_path)
            item = QtGui.QStandardItem(text)
            item.setEditable(False)
            item.setData(abs_path, QtCore.Qt.UserRole + 10)
            item.setData("file", QtCore.Qt.UserRole + 11)
            item.setDragEnabled(True)
            item.setDropEnabled(False)
            parent.appendRow(item)

    def _is_project_file_item(self, item: Optional[QtGui.QStandardItem]) -> bool:
        if item is None:
            return False
        return bool(item.data(QtCore.Qt.UserRole + 10)) and (item.data(QtCore.Qt.UserRole + 11) == "file")

    def _is_project_module_item(self, item: Optional[QtGui.QStandardItem]) -> bool:
        if item is None:
            return False
        return (item.data(QtCore.Qt.UserRole + 11) == "module") and (not item.data(QtCore.Qt.UserRole + 10))

    def _layer_root_item_for_index(self, idx: QtCore.QModelIndex) -> Optional[QtCore.QModelIndex]:
        if not idx.isValid():
            return None
        cur = idx
        while cur.parent().isValid():
            cur = cur.parent()
        return cur

    def _on_project_tree_menu(self, pos: QtCore.QPoint) -> None:
        idx = self._project_tree.indexAt(pos)
        if not idx.isValid():
            return
        if not idx.parent().isValid():
            return
        model = self._project_tree.model()
        if model is None:
            return
        key_idx = model.index(idx.row(), 0, idx.parent())
        path = model.data(key_idx, QtCore.Qt.UserRole + 10)
        kind = model.data(key_idx, QtCore.Qt.UserRole + 11)
        # module 节点没有 path

        menu = QtWidgets.QMenu(self)
        act_export = None
        act_regen = None
        act_regen_all = None
        if path:
            act_export = menu.addAction("导出单函数…")
            act_regen = menu.addAction("重新生成 CSU…")
            act_regen_all = menu.addAction("批量重新生成本文件 CSU…")

        act_merge = None
        if self._count_selected_file_items_in_same_layer() >= 2:
            act_merge = menu.addAction("合并为模块…")

        act_rename_module = None
        act_unmerge = None
        if kind == "module":
            act_rename_module = menu.addAction("重命名模块…")
            act_unmerge = menu.addAction("解散模块")

        act_remove = None
        if path:
            act_remove = menu.addAction("从列表移除")
        act = menu.exec_(self._project_tree.viewport().mapToGlobal(pos))
        if act_export is not None and act == act_export:
            self._export_single_function_from_file(str(path))
            return
        if act_regen is not None and act == act_regen:
            self._regenerate_csu_from_file(str(path))
            return
        if act_regen_all is not None and act == act_regen_all:
            self._regenerate_all_csu_from_file(str(path))
            return
        if act_merge is not None and act == act_merge:
            self._merge_selected_files_to_module()
            return
        if act_rename_module is not None and act == act_rename_module:
            self._rename_selected_module()
            return
        if act_unmerge is not None and act == act_unmerge:
            self._unmerge_selected_module()
            return
        if act_remove is not None and act == act_remove:
            self._remove_selected_project_item()
            return

    def _count_selected_file_items_in_same_layer(self) -> int:
        model = self._project_tree.model()
        sel = self._project_tree.selectionModel()
        if model is None or sel is None:
            return 0
        rows = []
        for idx in sel.selectedRows():
            if not idx.isValid() or (not idx.parent().isValid()):
                continue
            item = model.itemFromIndex(idx) if isinstance(model, QtGui.QStandardItemModel) else None
            if not self._is_project_file_item(item):
                continue
            rows.append(idx)
        if len(rows) < 2:
            return len(rows)
        layer0 = self._layer_root_item_for_index(rows[0])
        if layer0 is None:
            return 0
        for it in rows[1:]:
            if self._layer_root_item_for_index(it) != layer0:
                return 0
        return len(rows)

    def _merge_selected_files_to_module(self) -> None:
        if self._thread is not None:
            QtWidgets.QMessageBox.information(self, "提示", "任务运行中，暂不支持合并模块。")
            return
        model = self._project_tree.model()
        sel = self._project_tree.selectionModel()
        if model is None or sel is None or (not isinstance(model, QtGui.QStandardItemModel)):
            return

        selected: list[QtGui.QStandardItem] = []
        selected_indexes = list(sel.selectedRows())
        for idx in selected_indexes:
            if not idx.isValid() or (not idx.parent().isValid()):
                continue
            it = model.itemFromIndex(idx)
            if self._is_project_file_item(it):
                selected.append(it)

        if len(selected) < 2:
            return

        # 仅允许同一层合并（应用/中间/驱动）
        layer_idx0 = self._layer_root_item_for_index(selected_indexes[0])
        if layer_idx0 is None:
            return
        for idx in selected_indexes[1:]:
            if self._layer_root_item_for_index(idx) != layer_idx0:
                QtWidgets.QMessageBox.information(self, "提示", "仅支持同一层内合并为模块。")
                return

        default_name = "合并模块"
        name, ok = QtWidgets.QInputDialog.getText(self, "合并为模块", "模块名称：", text=default_name)
        name = (name or "").strip()
        if (not ok) or (not name):
            return

        layer_item = model.itemFromIndex(layer_idx0)
        if layer_item is None:
            return

        module_item = QtGui.QStandardItem(name)
        module_item.setEditable(False)
        module_item.setData("module", QtCore.Qt.UserRole + 11)
        module_item.setDragEnabled(True)
        module_item.setDropEnabled(True)
        layer_item.appendRow(module_item)

        # 移动选中文件到模块下；若来自已有模块且为空则清理
        parents_to_cleanup: set[QtGui.QStandardItem] = set()
        for it in selected:
            parent = it.parent()
            if parent is None:
                continue
            row = it.row()
            moved = parent.takeRow(row)
            if moved:
                module_item.appendRow(moved)
                if self._is_project_module_item(parent):
                    parents_to_cleanup.add(parent)

        for p in parents_to_cleanup:
            if p.rowCount() == 0:
                pp = p.parent()
                if pp is not None:
                    pp.removeRow(p.row())

        self._project_tree.expandAll()
        self._append_log(f"已合并为模块：{name}（{len(selected)} 个文件）")

    def _rename_selected_module(self) -> None:
        model = self._project_tree.model()
        idx = self._project_tree.currentIndex()
        if model is None or (not isinstance(model, QtGui.QStandardItemModel)) or (not idx.isValid()):
            return
        it = model.itemFromIndex(idx)
        if not self._is_project_module_item(it):
            return
        name, ok = QtWidgets.QInputDialog.getText(self, "重命名模块", "模块名称：", text=str(it.text() or ""))
        name = (name or "").strip()
        if (not ok) or (not name):
            return
        it.setText(name)
        self._append_log(f"模块已重命名为：{name}")

    def _unmerge_selected_module(self) -> None:
        if self._thread is not None:
            QtWidgets.QMessageBox.information(self, "提示", "任务运行中，暂不支持解散模块。")
            return
        model = self._project_tree.model()
        idx = self._project_tree.currentIndex()
        if model is None or (not isinstance(model, QtGui.QStandardItemModel)) or (not idx.isValid()):
            return
        it = model.itemFromIndex(idx)
        if not self._is_project_module_item(it):
            return
        parent = it.parent()
        if parent is None:
            return
        count = it.rowCount()
        # 把子文件移回层根节点末尾
        while it.rowCount() > 0:
            row = it.takeRow(0)
            if row:
                parent.appendRow(row)
        parent.removeRow(it.row())
        self._append_log(f"模块已解散：移回 {count} 个文件")


    def _export_single_function_from_file(self, c_file: str) -> None:
        if self._thread is not None:
            QtWidgets.QMessageBox.information(self, "提示", "任务运行中，暂不支持导出单函数。")
            return
        c_file = os.path.abspath((c_file or "").strip())
        if (not c_file) or (not os.path.isfile(c_file)):
            QtWidgets.QMessageBox.warning(self, "提示", "未找到 C 文件。")
            return
        if not c_file.lower().endswith(".c"):
            QtWidgets.QMessageBox.information(self, "提示", "目前仅支持从 .c 文件导出单函数。")
            return

        try:
            proj_root = (self.ed_project.text() or "").strip() or None
            tmp_cfg = self.backend.GenConfig(verbose=False)
            func_list, _ = self.backend.prepare_func_list_for_c_file(
                c_file,
                project_root=proj_root,
                cfg=tmp_cfg,
                prefilter=False,
                need_symbol_map=False,
            )
            names: list[str] = []
            for fd in func_list or []:
                fi = fd.get("func_info") or {}
                fn = (fi.get("func_name") or "").strip()
                if fn:
                    names.append(fn)
            names = sorted(set(names))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"读取函数列表失败：{e}")
            return

        if not names:
            QtWidgets.QMessageBox.information(self, "提示", "该文件未解析到任何函数。")
            return

        func_name, ok = QtWidgets.QInputDialog.getItem(self, "导出单函数", "选择要导出的函数：", names, 0, True)
        func_name = (func_name or "").strip()
        if (not ok) or (not func_name):
            return

        base_dir = ""
        try:
            if self.ed_output.text().strip():
                base_dir = os.path.dirname(os.path.abspath(self.ed_output.text().strip()))
        except Exception:
            base_dir = ""
        if not base_dir:
            base_dir = os.path.dirname(c_file)

        def _safe_name(s: str) -> str:
            s = (s or "").strip()
            s = re.sub(r"[\\\\/:*?\"<>|]+", "_", s)
            s = re.sub(r"\\s+", "_", s)
            return s or "function"

        default_name = f"{os.path.splitext(os.path.basename(c_file))[0]}_{_safe_name(func_name)}.docx"
        default_path = os.path.join(base_dir, default_name)
        out, _ = QtWidgets.QFileDialog.getSaveFileName(self, "选择导出路径", default_path, "Word 文档 (*.docx)")
        out = (out or "").strip()
        if not out:
            return

        self._save_settings_from_ui_silent()
        self._save_recent_inputs()
        settings = self.settings

        task = TaskSpec(
            mode="export_func",
            c_file=c_file,
            project_dir=self.ed_project.text().strip(),
            output=out,
            template_path=self.ed_template.text().strip(),
            project_file_order=None,
            func_name=func_name,
        )
        steps = [
            StepDef("validate", "校验输入"),
            StepDef("config", "准备配置"),
            StepDef("generate", "导出单函数"),
        ]
        self._reset_steps(steps)
        self._resume_state = None
        self._set_running(True, note=f"导出单函数：{func_name} ...")

        self._run_in_thread(
            impl=ExportFuncWorker(backend=self.backend, task=task, settings=settings),
            kind="export_func",
        )

    def _regenerate_csu_from_file(self, c_file: str) -> None:
        """重新生成单个 CSU，在已有文档中原位替换。"""
        if self._thread is not None:
            QtWidgets.QMessageBox.information(self, "提示", "任务运行中，暂不支持重新生成 CSU。")
            return
        c_file = os.path.abspath((c_file or "").strip())
        if (not c_file) or (not os.path.isfile(c_file)):
            QtWidgets.QMessageBox.warning(self, "提示", "未找到 C 文件。")
            return
        if not c_file.lower().endswith(".c"):
            QtWidgets.QMessageBox.information(self, "提示", "目前仅支持从 .c 文件重新生成 CSU。")
            return

        # 必须先有输出文档
        doc_path = (self.ed_output.text() or "").strip()
        if not doc_path or not os.path.isfile(doc_path):
            QtWidgets.QMessageBox.warning(self, "提示", "请先选择已生成的 Word 文档（输出路径）。")
            return

        # 解析函数列表
        try:
            proj_root = (self.ed_project.text() or "").strip() or None
            tmp_cfg = self.backend.GenConfig(verbose=False)
            func_list, _ = self.backend.prepare_func_list_for_c_file(
                c_file,
                project_root=proj_root,
                cfg=tmp_cfg,
                prefilter=False,
                need_symbol_map=False,
            )
            names: list[str] = []
            for fd in func_list or []:
                fi = fd.get("func_info") or {}
                fn = (fi.get("func_name") or "").strip()
                if fn:
                    names.append(fn)
            names = sorted(set(names))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"读取函数列表失败：{e}")
            return

        if not names:
            QtWidgets.QMessageBox.information(self, "提示", "该文件未解析到任何函数。")
            return

        # 选择函数
        func_name, ok = QtWidgets.QInputDialog.getItem(
            self, "重新生成 CSU", "选择要重新生成的函数：", names, 0, True
        )
        func_name = (func_name or "").strip()
        if (not ok) or (not func_name):
            return

        # 从已有文档中提取该文件对应模块的 CSU 列表，让用户选择 CSU 编号
        try:
            from docx import Document as _Doc
            doc = _Doc(doc_path)
            # 收集所有 Heading 4 的 (title, csu_id)
            csu_items: list[str] = []
            csu_ids: list[str] = []
            for p in doc.paragraphs:
                if p.style.name == "Heading 4":
                    import re as _re
                    m = _re.search(r"（(D/R_SDD01_\d+_\d+)）", p.text)
                    if m:
                        csu_ids.append(m.group(1))
                        csu_items.append(f"{p.text}  [{m.group(1)}]")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"读取文档 CSU 列表失败：{e}")
            return

        if not csu_items:
            QtWidgets.QMessageBox.information(self, "提示", "文档中未找到任何 CSU。")
            return

        # 选择 CSU
        sel_idx, ok2 = QtWidgets.QInputDialog.getItem(
            self, "重新生成 CSU", "选择要替换的 CSU：", csu_items, 0, True
        )
        if not ok2:
            return
        sel_idx = csu_items.index(sel_idx) if sel_idx in csu_items else 0
        csu_id = csu_ids[sel_idx] if sel_idx < len(csu_ids) else ""
        if not csu_id:
            return

        self._save_settings_from_ui_silent()
        self._save_recent_inputs()
        settings = self.settings

        task = TaskSpec(
            mode="regen_csu",
            c_file=c_file,
            project_dir=self.ed_project.text().strip(),
            output=doc_path,
            template_path=self.ed_template.text().strip(),
            project_file_order=None,
            func_name=func_name,
            csu_id=csu_id,
        )
        steps = [
            StepDef("validate", "校验输入"),
            StepDef("config", "准备配置"),
            StepDef("generate", "重新生成 CSU"),
        ]
        self._reset_steps(steps)
        self._resume_state = None
        self._set_running(True, note=f"重新生成 CSU：{func_name} → {csu_id} ...")

        self._run_in_thread(
            impl=RegenerateCsuWorker(backend=self.backend, task=task, settings=settings),
            kind="regen_csu",
        )

    def _regenerate_all_csu_from_file(self, c_file: str) -> None:
        """批量重新生成一个 .c 文件对应模块内的全部 CSU，逐个原位替换。"""
        if self._thread is not None:
            QtWidgets.QMessageBox.information(self, "提示", "任务运行中，暂不支持批量重新生成。")
            return
        c_file = os.path.abspath((c_file or "").strip())
        if (not c_file) or (not os.path.isfile(c_file)):
            QtWidgets.QMessageBox.warning(self, "提示", "未找到 C 文件。")
            return
        if not c_file.lower().endswith(".c"):
            QtWidgets.QMessageBox.information(self, "提示", "目前仅支持从 .c 文件重新生成 CSU。")
            return

        # 必须先有输出文档
        doc_path = (self.ed_output.text() or "").strip()
        if not doc_path or not os.path.isfile(doc_path):
            QtWidgets.QMessageBox.warning(self, "提示", "请先选择已生成的 Word 文档（输出路径）。")
            return

        # 解析源码函数列表
        try:
            proj_root = (self.ed_project.text() or "").strip() or None
            tmp_cfg = self.backend.GenConfig(verbose=False)
            func_list, _ = self.backend.prepare_func_list_for_c_file(
                c_file,
                project_root=proj_root,
                cfg=tmp_cfg,
                prefilter=False,
                need_symbol_map=False,
            )
            src_names: list[str] = []
            for fd in func_list or []:
                fi = fd.get("func_info") or {}
                fn = (fi.get("func_name") or "").strip()
                if fn:
                    src_names.append(fn)
            src_names = sorted(set(src_names))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"读取函数列表失败：{e}")
            return

        if not src_names:
            QtWidgets.QMessageBox.information(self, "提示", "该文件未解析到任何函数。")
            return

        # 从文档中读取该文件对应模块的 CSU 列表
        try:
            import re as _re
            from docx import Document as _Doc
            doc = _Doc(doc_path)
            # 收集所有 Heading 4: (title, csu_id, func_name)
            doc_csus: list[tuple[str, str, str]] = []
            for p in doc.paragraphs:
                if p.style.name == "Heading 4":
                    m = _re.search(r"(.+?)（(D/R_SDD01_\d+_\d+)）", p.text)
                    if m:
                        title = m.group(1)
                        csu_id = m.group(2)
                        # 从紧随的 a) 函数原型段提取函数名
                        func_name = ""
                        for i, pp in enumerate(doc.paragraphs):
                            if pp.text == p.text and pp.style.name == "Heading 4":
                                for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                                    if doc.paragraphs[j].text.strip() == "a) 函数原型":
                                        proto = doc.paragraphs[j + 1].text.strip() if j + 1 < len(doc.paragraphs) else ""
                                        fm = _re.search(r"(?:interrupt\s+)?[\w\s\*]+\s+(\w+)\s*\(", proto)
                                        if fm:
                                            func_name = fm.group(1)
                                        break
                                break
                        doc_csus.append((title, csu_id, func_name))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"读取文档 CSU 列表失败：{e}")
            return

        if not doc_csus:
            QtWidgets.QMessageBox.information(self, "提示", "文档中未找到任何 CSU。")
            return

        # 匹配: 文档中函数名在源码函数列表里的 CSU
        matched: list[tuple[str, str, str]] = []  # (title, csu_id, func_name)
        unmatched_doc: list[tuple[str, str]] = []  # (title, csu_id)
        for title, csu_id, fn in doc_csus:
            if fn and fn in src_names:
                matched.append((title, csu_id, fn))
            else:
                unmatched_doc.append((title, csu_id))

        if not matched:
            QtWidgets.QMessageBox.information(
                self, "提示",
                f"文档中的 CSU 没有匹配到 {os.path.basename(c_file)} 的函数。\n"
                f"文档 CSU: {len(doc_csus)} 个，源码函数: {len(src_names)} 个。"
            )
            return

        # 确认
        msg = (
            f"即将重新生成 {len(matched)} 个 CSU（来自 {os.path.basename(c_file)}）。\n"
            f"匹配的函数：\n"
        )
        for title, csu_id, fn in matched[:8]:
            msg += f"  • {fn} → {csu_id}（{title}）\n"
        if len(matched) > 8:
            msg += f"  ... 及其余 {len(matched) - 8} 个\n"
        msg += "\n文档将被原位更新，其他模块不受影响。是否继续？"
        ret = QtWidgets.QMessageBox.question(self, "批量重新生成 CSU", msg)
        if ret != QtWidgets.QMessageBox.Yes:
            return

        self._save_settings_from_ui_silent()
        self._save_recent_inputs()
        settings = self.settings

        task = TaskSpec(
            mode="regen_csu_batch",
            c_file=c_file,
            project_dir=self.ed_project.text().strip(),
            output=doc_path,
            template_path=self.ed_template.text().strip(),
            project_file_order=None,
            func_name="",  # 留空，worker 从 task._csu_pairs 读取
            csu_id="",
        )
        # 把匹配列表附在 task 上（TaskSpec 是 frozen dataclass，用 object.__setattr__）
        object.__setattr__(task, "_csu_pairs", [(fn, cid) for _, cid, fn in matched])

        steps = [
            StepDef("validate", "校验输入"),
            StepDef("config", "准备配置"),
            StepDef("generate", f"批量重新生成 {len(matched)} 个 CSU"),
        ]
        self._reset_steps(steps)
        self._resume_state = None
        self._set_running(True, note=f"批量重新生成 {len(matched)} 个 CSU …")

        self._run_in_thread(
            impl=RegenerateCsuBatchWorker(backend=self.backend, task=task, settings=settings),
            kind="regen_csu_batch",
        )

    def _remove_selected_project_item(self) -> None:
        model = self._project_tree.model()
        sel = self._project_tree.selectionModel()
        if model is None or sel is None:
            return

        indexes = list(sel.selectedRows())
        if not indexes:
            idx = self._project_tree.currentIndex()
            if idx.isValid():
                indexes = [idx]
        if not indexes:
            return

        # 仅移除"文件"项，且批量删除需要从后往前处理
        by_parent: dict[QtCore.QModelIndex, list[QtCore.QModelIndex]] = {}
        removed_names: list[str] = []
        for idx in indexes:
            if not idx.isValid() or (not idx.parent().isValid()):
                continue
            key_idx = model.index(idx.row(), 0, idx.parent())
            path = model.data(key_idx, QtCore.Qt.UserRole + 10)
            kind = model.data(key_idx, QtCore.Qt.UserRole + 11)
            if not path or kind != "file":
                continue
            by_parent.setdefault(idx.parent(), []).append(idx)
            removed_names.append(str(model.data(key_idx) or ""))

        if not by_parent:
            return

        for parent, rows in by_parent.items():
            for idx in sorted(rows, key=lambda x: x.row(), reverse=True):
                model.removeRow(idx.row(), parent)

        if removed_names:
            if len(removed_names) == 1:
                self._append_log(f"已从列表移除：{removed_names[0]}")
            else:
                self._append_log(f"已从列表移除：{len(removed_names)} 个文件")

    def _load_project_from_ui(self) -> None:
        proj = (self.ed_project.text() or "").strip()
        if not proj:
            return
        if not os.path.isdir(proj):
            QtWidgets.QMessageBox.warning(self, "提示", "工程目录不存在。")
            return
        if self._thread is not None:
            QtWidgets.QMessageBox.information(self, "提示", "任务运行中，暂不支持重载工程列表。")
            return
        try:
            cfg = self.backend.GenConfig(
                prefilter_project_files=bool(self.settings.prefilter_project_files),
                exclude_dirs=tuple([x.strip() for x in (self.settings.exclude_dirs or []) if str(x).strip()]),
                mid_dir_keywords=tuple([x.strip() for x in (self.settings.mid_dir_keywords or []) if str(x).strip()]),
                drv_dir_keywords=tuple([x.strip() for x in (self.settings.drv_dir_keywords or []) if str(x).strip()]),
            )
            src_dir, app_files, mid_files, drv_files = self.backend.collect_project_c_files_by_layer(proj, cfg)
            if not src_dir:
                QtWidgets.QMessageBox.information(self, "提示", "未找到 SRC 目录（不区分大小写）。")
                self._clear_project_children()
                return
            self._clear_project_children()
            self._append_project_files(self._proj_item_app, app_files, base_dir=src_dir)
            self._append_project_files(self._proj_item_mid, mid_files, base_dir=src_dir)
            self._append_project_files(self._proj_item_drv, drv_files, base_dir=src_dir)
            self._project_tree.expandAll()
            try:
                self._panel_tabs.setCurrentWidget(self._project_tree)
            except Exception:
                pass
            total = len(app_files) + len(mid_files) + len(drv_files)
            self._append_log(f"工程已加载：{total} 个 C 文件（应用层 {len(app_files)} / 中间层 {len(mid_files)} / 驱动层 {len(drv_files)}）")
            # 显示 clangd 状态
            self._check_and_show_clangd_status()
            self._check_compile_env(proj)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"加载工程失败：{e}")

    def _check_and_show_clangd_status(self) -> None:
        """检查并显示 clangd 状态"""
        import shutil
        import subprocess

        clangd_path = None
        clangd_version = None

        # 检查配置中指定的 clangd 路径
        try:
            extra_params = getattr(self.settings, 'extra_params', {}) or {}
            configured_path = extra_params.get('logic_lsp_clangd_path', '')
            if configured_path and os.path.isfile(configured_path):
                clangd_path = configured_path
        except Exception:
            pass

        # 如果配置路径无效，检查系统 PATH
        if not clangd_path:
            clangd_path = shutil.which('clangd')

        # 如果还是没找到，检查工具目录
        if not clangd_path:
            # 打包后的路径
            if getattr(sys, 'frozen', False):
                app_root = os.path.dirname(sys.executable)
            else:
                app_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            candidates = [
                os.path.join(app_root, 'tools', 'clangd', 'win7', 'llvm', 'bin', 'clangd.exe'),
                os.path.join(app_root, 'tools', 'clangd', 'win7', 'bin', 'clangd.exe'),
                os.path.join(app_root, 'tools', 'clangd', 'clangd.exe'),
            ]
            for candidate in candidates:
                if os.path.isfile(candidate):
                    clangd_path = candidate
                    break

        # 尝试获取版本
        if clangd_path and os.path.isfile(clangd_path):
            try:
                result = subprocess.run(
                    [clangd_path, '--version'],
                    capture_output=True,
                    text=True,
                    timeout=5
                )
                if result.returncode == 0:
                    # 解析版本号，格式通常是 "clangd version X.Y.Z"
                    version_line = result.stdout.strip().split('\n')[0] if result.stdout else ''
                    import re
                    match = re.search(r'clangd version (\d+\.\d+\.\d+)', version_line)
                    if match:
                        clangd_version = match.group(1)
            except Exception:
                pass

        # 显示状态
        if clangd_path and os.path.isfile(clangd_path):
            version_str = f" (版本 {clangd_version})" if clangd_version else ""
            self._append_log(f"✓ clangd 可用：{clangd_path}{version_str}")
        else:
            self._append_log("⚠ clangd 未找到，LSP 功能将不可用（可在设置中配置 clangd 路径）")

    def _check_compile_env(self, project_root: str) -> None:
        """检测项目编译环境并提示用户。"""
        root = project_root.strip()
        if not root or not os.path.isdir(root):
            return

        has_compile_commands = os.path.isfile(os.path.join(root, "compile_commands.json"))
        has_cproject = os.path.isfile(os.path.join(root, ".cproject"))
        has_iar = any(f.endswith(".ewp") for f in os.listdir(root))
        has_makefile = os.path.isfile(os.path.join(root, "Makefile"))

        compiler = "unknown"
        if has_iar:
            compiler = "iar"
        elif has_cproject:
            compiler = "ccs"
        elif has_makefile:
            try:
                with open(os.path.join(root, "Makefile"), "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read(4096)
                if "arm-none-eabi-gcc" in content:
                    compiler = "gcc-arm"
                elif "armcc" in content or "armclang" in content:
                    compiler = "armcc"
            except Exception:
                pass

        if has_compile_commands:
            self._append_log("✓ 检测到 compile_commands.json，LSP 精度最高")
        elif compiler == "ccs":
            self._append_log("⚠ 检测到 CCS 项目 (.cproject)，未找到 compile_commands.json。"
                             "建议在 CCS 中导出：Project → Properties → C/C++ Build → "
                             "Settings → Miscellaneous → Export compile_commands.json")
        elif compiler == "iar":
            self._append_log("⚠ 检测到 IAR 项目，clangd 对 IAR 支持有限，LSP 功能可能不完整。"
                             "建议在设置中关闭 LSP 或手动导出 compile_commands.json")
        elif compiler == "gcc-arm":
            self._append_log("ℹ 检测到 GCC ARM 项目，将使用自动推断的编译选项，精度可能降低")
        elif compiler == "armcc":
            self._append_log("⚠ 检测到 ARMCC 项目，clangd 对 ARMCC 支持有限，LSP 功能可能不完整")

    def _get_project_file_order(self) -> Optional[dict[str, list[object]]]:
        proj = (self.ed_project.text() or "").strip()
        if not proj:
            return None

        def children_entries(item: QtGui.QStandardItem) -> list[object]:
            out: list[object] = []
            for r in range(item.rowCount()):
                child = item.child(r)
                if child is None:
                    continue
                kind = child.data(QtCore.Qt.UserRole + 11)
                if kind == "module":
                    files: list[str] = []
                    for rr in range(child.rowCount()):
                        cc = child.child(rr)
                        if cc is None:
                            continue
                        p = cc.data(QtCore.Qt.UserRole + 10)
                        if p:
                            files.append(str(p))
                    if files:
                        out.append({"module": str(child.text() or "").strip(), "files": files})
                    continue
                p = child.data(QtCore.Qt.UserRole + 10)
                if p:
                    out.append(str(p))
            return out

        app = children_entries(self._proj_item_app)
        mid = children_entries(self._proj_item_mid)
        drv = children_entries(self._proj_item_drv)
        if not (app or mid or drv):
            return None
        return {"app": app, "mid": mid, "drv": drv, "_explicit": True}

    def _build_actions(self) -> None:
        self.act_start = QtWidgets.QAction("生成文档", self)
        self.act_start.triggered.connect(lambda: self._start_generate(is_resume=False))

        self.act_stop = QtWidgets.QAction("停止", self)
        self.act_stop.triggered.connect(self._stop_running)

        self.act_continue = QtWidgets.QAction("继续", self)
        self.act_continue.triggered.connect(lambda: self._start_generate(is_resume=True))

        self.act_update_csu = QtWidgets.QAction("更新CSU标识", self)
        self.act_update_csu.triggered.connect(self._start_update_csu)

        self.act_refresh_term_table = QtWidgets.QAction("刷新工程术语表", self)
        self.act_refresh_term_table.triggered.connect(self._start_refresh_term_table)

        self.act_show_term_table = QtWidgets.QAction("查看工程术语表", self)
        self.act_show_term_table.triggered.connect(self._show_project_term_table)

        self.act_exit = QtWidgets.QAction("退出", self)
        self.act_exit.triggered.connect(self.close)

        self.act_open_home = QtWidgets.QAction("主页", self)
        self.act_open_home.triggered.connect(lambda: self._stack.setCurrentWidget(self._page_home))

        self.act_open_settings = QtWidgets.QAction("设置", self)
        self.act_open_settings.triggered.connect(self._open_settings_page)

        self.act_open_help = QtWidgets.QAction("帮助", self)
        self.act_open_help.triggered.connect(lambda: self._stack.setCurrentWidget(self._page_help))

        self.act_about = QtWidgets.QAction("关于", self)
        self.act_about.triggered.connect(self._show_about)

        self.act_clear_log = QtWidgets.QAction("清空日志", self)
        self.act_clear_log.triggered.connect(lambda: self._log_view.clear())

        self.act_save_log = QtWidgets.QAction("导出日志...", self)
        self.act_save_log.triggered.connect(self._save_log)

        self.act_toggle_sidebar = QtWidgets.QAction("显示侧边栏", self)
        self.act_toggle_sidebar.setCheckable(True)
        self.act_toggle_sidebar.setChecked(True)
        self.act_toggle_sidebar.triggered.connect(lambda checked: self._sidebar.setVisible(bool(checked)))

        self.act_toggle_panels = QtWidgets.QAction("显示右侧面板", self)
        self.act_toggle_panels.setCheckable(True)
        self.act_toggle_panels.setChecked(True)
        self.act_toggle_panels.triggered.connect(lambda checked: self._right_panel.setVisible(bool(checked)))

    def _build_menus(self) -> None:
        mb = self.menuBar()
        m_nav = mb.addMenu("导航")
        m_nav.addAction(self.act_open_home)
        m_nav.addAction(self.act_open_settings)
        m_nav.addAction(self.act_open_help)

        m_file = mb.addMenu("文件")
        m_file.addAction(self.act_exit)

        m_run = mb.addMenu("运行")
        m_run.addAction(self.act_start)
        m_run.addAction(self.act_stop)
        m_run.addAction(self.act_continue)
        m_run.addSeparator()
        m_run.addAction(self.act_update_csu)
        m_run.addSeparator()
        m_run.addAction(self.act_refresh_term_table)
        m_run.addAction(self.act_show_term_table)

        m_view = mb.addMenu("视图")
        m_view.addAction(self.act_toggle_sidebar)
        m_view.addAction(self.act_toggle_panels)
        m_view.addSeparator()
        m_view.addAction(self.act_clear_log)
        m_view.addAction(self.act_save_log)

        m_settings = mb.addMenu("设置")
        m_settings.addAction(self.act_open_settings)

        m_help = mb.addMenu("帮助")
        m_help.addAction(self.act_open_help)
        m_help.addSeparator()
        m_help.addAction(self.act_about)

    def _build_store_page(self) -> QtWidgets.QWidget:
        def _apply_shadow(w: QtWidgets.QWidget, *, blur: int = 26, dy: int = 10, alpha: int = 45) -> None:
            try:
                eff = QtWidgets.QGraphicsDropShadowEffect(w)
                eff.setBlurRadius(blur)
                eff.setOffset(0, dy)
                eff.setColor(QtGui.QColor(15, 23, 42, alpha))
                w.setGraphicsEffect(eff)
            except Exception:
                pass

        def _make_store_card(*, kind: str, tone: str, title: str, subtitle: str, meta: str = "") -> QtWidgets.QFrame:
            card = QtWidgets.QFrame()
            card.setObjectName("store_card")
            card.setProperty("tone", tone)
            card.setCursor(QtCore.Qt.PointingHandCursor)
            card.setSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Fixed)
            card.setMinimumHeight(156)

            outer = QtWidgets.QVBoxLayout(card)
            outer.setContentsMargins(14, 14, 14, 14)
            outer.setSpacing(10)

            badge = QtWidgets.QLabel(kind)
            badge.setObjectName("store_badge")
            outer.addWidget(badge, 0, QtCore.Qt.AlignLeft)

            row = QtWidgets.QHBoxLayout()
            row.setContentsMargins(0, 0, 0, 0)
            row.setSpacing(12)
            outer.addLayout(row, 1)

            icon = QtWidgets.QLabel((title or "?")[:1])
            icon.setObjectName("store_icon")
            icon.setProperty("tone", tone)
            icon.setFixedSize(56, 56)
            icon.setAlignment(QtCore.Qt.AlignCenter)
            row.addWidget(icon, 0, QtCore.Qt.AlignTop)

            info = QtWidgets.QVBoxLayout()
            info.setContentsMargins(0, 0, 0, 0)
            info.setSpacing(4)
            row.addLayout(info, 1)

            title_lbl = QtWidgets.QLabel(title)
            title_lbl.setObjectName("store_title")
            subtitle_lbl = QtWidgets.QLabel(subtitle)
            subtitle_lbl.setObjectName("store_subtitle")
            subtitle_lbl.setWordWrap(True)
            info.addWidget(title_lbl)
            info.addWidget(subtitle_lbl, 1)

            right = QtWidgets.QVBoxLayout()
            right.setContentsMargins(0, 0, 0, 0)
            right.setSpacing(6)
            row.addLayout(right, 0)

            btn = QtWidgets.QPushButton("敬请期待")
            btn.setObjectName("store_get_btn")
            btn.setCursor(QtCore.Qt.PointingHandCursor)
            meta_lbl = QtWidgets.QLabel(meta or "")
            meta_lbl.setObjectName("store_meta")
            meta_lbl.setAlignment(QtCore.Qt.AlignRight)
            meta_lbl.setVisible(bool(meta))
            right.addWidget(btn, 0, QtCore.Qt.AlignRight)
            right.addWidget(meta_lbl, 0, QtCore.Qt.AlignRight)
            right.addStretch(1)

            _apply_shadow(card, blur=28, dy=10, alpha=30)
            return card

        def _make_today_card(*, tone: str, kicker: str, headline: str, caption: str, app: str, app_sub: str) -> QtWidgets.QFrame:
            card = QtWidgets.QFrame()
            card.setObjectName("store_today_card")
            card.setProperty("tone", tone)
            card.setCursor(QtCore.Qt.PointingHandCursor)
            card.setFixedWidth(340)
            card.setMinimumHeight(320)

            outer = QtWidgets.QVBoxLayout(card)
            outer.setContentsMargins(0, 0, 0, 0)
            outer.setSpacing(0)

            media = QtWidgets.QFrame()
            media.setObjectName("store_today_media")
            media.setProperty("tone", tone)
            media.setMinimumHeight(210)
            media_layout = QtWidgets.QVBoxLayout(media)
            media_layout.setContentsMargins(14, 14, 14, 14)
            media_layout.setSpacing(6)

            lab_kicker = QtWidgets.QLabel(kicker.upper())
            lab_kicker.setObjectName("store_today_kicker")
            lab_headline = QtWidgets.QLabel(headline)
            lab_headline.setObjectName("store_today_headline")
            lab_headline.setWordWrap(True)
            lab_caption = QtWidgets.QLabel(caption)
            lab_caption.setObjectName("store_today_caption")
            lab_caption.setWordWrap(True)

            media_layout.addWidget(lab_kicker)
            media_layout.addWidget(lab_headline)
            media_layout.addStretch(1)
            media_layout.addWidget(lab_caption)

            outer.addWidget(media)

            bottom = QtWidgets.QWidget()
            bottom.setObjectName("store_today_bottom")
            b = QtWidgets.QHBoxLayout(bottom)
            b.setContentsMargins(14, 12, 14, 12)
            b.setSpacing(10)

            icon = QtWidgets.QLabel((app or "?")[:1])
            icon.setObjectName("store_small_icon")
            icon.setProperty("tone", tone)
            icon.setFixedSize(44, 44)
            icon.setAlignment(QtCore.Qt.AlignCenter)
            b.addWidget(icon, 0)

            info = QtWidgets.QVBoxLayout()
            info.setContentsMargins(0, 0, 0, 0)
            info.setSpacing(2)
            app_title = QtWidgets.QLabel(app)
            app_title.setObjectName("store_today_app_title")
            app_subtitle = QtWidgets.QLabel(app_sub)
            app_subtitle.setObjectName("store_today_app_subtitle")
            app_subtitle.setWordWrap(True)
            info.addWidget(app_title)
            info.addWidget(app_subtitle)
            b.addLayout(info, 1)

            btn = QtWidgets.QPushButton("敬请期待")
            btn.setObjectName("store_get_btn")
            btn.setCursor(QtCore.Qt.PointingHandCursor)
            b.addWidget(btn, 0, QtCore.Qt.AlignTop)

            outer.addWidget(bottom)
            _apply_shadow(card, blur=34, dy=14, alpha=24)
            return card

        def _make_section_header(title_text: str) -> QtWidgets.QWidget:
            w = QtWidgets.QWidget()
            h = QtWidgets.QHBoxLayout(w)
            h.setContentsMargins(0, 0, 0, 0)
            h.setSpacing(10)
            title = QtWidgets.QLabel(title_text)
            title.setObjectName("store_section_title")
            more = QtWidgets.QLabel("查看全部")
            more.setObjectName("store_section_more")
            more.setCursor(QtCore.Qt.PointingHandCursor)
            h.addWidget(title, 1)
            h.addWidget(more, 0, QtCore.Qt.AlignRight)
            return w

        def _make_carousel(title_text: str, cards: list[QtWidgets.QWidget]) -> QtWidgets.QWidget:
            box = QtWidgets.QWidget()
            box.setObjectName("store_carousel")
            v = QtWidgets.QVBoxLayout(box)
            v.setContentsMargins(0, 0, 0, 0)
            v.setSpacing(10)

            v.addWidget(_make_section_header(title_text))

            scroll = QtWidgets.QScrollArea()
            scroll.setObjectName("store_carousel_scroll")
            scroll.setWidgetResizable(True)
            scroll.setFrameShape(QtWidgets.QFrame.NoFrame)
            scroll.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAsNeeded)
            scroll.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)

            try:
                QtWidgets.QScroller.grabGesture(scroll.viewport(), QtWidgets.QScroller.LeftMouseButtonGesture)
            except Exception:
                pass

            body = QtWidgets.QWidget()
            body.setObjectName("store_carousel_body")
            row = QtWidgets.QHBoxLayout(body)
            row.setContentsMargins(2, 0, 2, 0)
            row.setSpacing(14)
            for c in cards:
                row.addWidget(c, 0)
            row.addStretch(1)
            scroll.setWidget(body)

            v.addWidget(scroll)
            return box

        page = QtWidgets.QWidget()
        page.setObjectName("store_page")
        layout = QtWidgets.QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        scroll = QtWidgets.QScrollArea()
        scroll.setObjectName("store_scroll")
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QtWidgets.QFrame.NoFrame)
        layout.addWidget(scroll, 1)

        body = QtWidgets.QWidget()
        body.setObjectName("store_body")
        scroll.setWidget(body)
        body_layout = QtWidgets.QVBoxLayout(body)
        body_layout.setContentsMargins(18, 16, 18, 16)
        body_layout.setSpacing(16)

        header = QtWidgets.QWidget()
        header.setObjectName("store_header")
        header_layout = QtWidgets.QVBoxLayout(header)
        header_layout.setContentsMargins(2, 0, 2, 0)
        header_layout.setSpacing(6)
        title = QtWidgets.QLabel("应用市场")
        title.setObjectName("store_header_title")
        subtitle = QtWidgets.QLabel("发现好用的工具与插件（静态展示）")
        subtitle.setObjectName("store_header_subtitle")
        header_layout.addWidget(title)
        header_layout.addWidget(subtitle)
        body_layout.addWidget(header)

        featured = QtWidgets.QFrame()
        featured.setObjectName("store_featured")
        featured_layout = QtWidgets.QHBoxLayout(featured)
        featured_layout.setContentsMargins(16, 16, 16, 16)
        featured_layout.setSpacing(14)
        f_icon = QtWidgets.QLabel("A")
        f_icon.setObjectName("store_featured_icon")
        f_icon.setFixedSize(64, 64)
        f_icon.setAlignment(QtCore.Qt.AlignCenter)
        featured_layout.addWidget(f_icon, 0)
        f_info = QtWidgets.QVBoxLayout()
        f_info.setContentsMargins(0, 0, 0, 0)
        f_info.setSpacing(4)
        f_badge = QtWidgets.QLabel("编辑推荐")
        f_badge.setObjectName("store_badge")
        f_title = QtWidgets.QLabel("AutoDocGen Pro")
        f_title.setObjectName("store_featured_title")
        f_sub = QtWidgets.QLabel("结构化文档生成套件 · 模板 + 校验 + 目录索引")
        f_sub.setObjectName("store_featured_subtitle")
        f_sub.setWordWrap(True)
        f_info.addWidget(f_badge, 0, QtCore.Qt.AlignLeft)
        f_info.addWidget(f_title)
        f_info.addWidget(f_sub, 1)
        featured_layout.addLayout(f_info, 1)
        f_btns = QtWidgets.QVBoxLayout()
        f_btns.setContentsMargins(0, 0, 0, 0)
        f_btns.setSpacing(8)
        f_get = QtWidgets.QPushButton("敬请期待")
        f_get.setObjectName("store_primary_btn")
        f_get.setCursor(QtCore.Qt.PointingHandCursor)
        f_btns.addWidget(f_get, 0, QtCore.Qt.AlignRight)
        f_btns.addStretch(1)
        featured_layout.addLayout(f_btns, 0)
        _apply_shadow(featured, blur=34, dy=14, alpha=26)
        body_layout.addWidget(featured)

        body_layout.addWidget(
            _make_carousel(
                "Today 精选",
                [
                    _make_today_card(
                        tone="blue",
                        kicker="效率工具",
                        headline="把常用模板变成一键工作流",
                        caption="更少步骤，更快输出。",
                        app="批量模板库",
                        app_sub="企业模板与目录结构",
                    ),
                    _make_today_card(
                        tone="purple",
                        kicker="文档体验",
                        headline="结构化导航，让大型工程也清晰",
                        caption="从模块到函数，层层可追踪。",
                        app="结构化导航",
                        app_sub="模块/函数快速索引",
                    ),
                    _make_today_card(
                        tone="orange",
                        kicker="变更管理",
                        headline="差异比对与审阅记录，一页看懂",
                        caption="适合评审与归档。",
                        app="差异比对",
                        app_sub="变更点一目了然",
                    ),
                ],
            )
        )

        body_layout.addWidget(
            _make_carousel(
                "小工具",
                [
                    _make_store_card(kind="规范", tone="green", title="命名检查", subtitle="函数/变量命名规范扫描（静态演示）"),
                    _make_store_card(kind="集成", tone="pink", title="导出中心", subtitle="导出日志/单函数/章节（静态演示）"),
                    _make_store_card(kind="安全", tone="slate", title="权限模板", subtitle="按角色分配访问与导出权限（静态演示）"),
                    _make_store_card(kind="协作", tone="indigo", title="评审流转", subtitle="意见/批注合并与流转看板（静态演示）"),
                ],
            )
        )

        body_layout.addStretch(1)
        return page

    def _build_home_page(self) -> QtWidgets.QWidget:
        page = QtWidgets.QWidget()
        layout = QtWidgets.QHBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        splitter = QtWidgets.QSplitter(QtCore.Qt.Horizontal)
        splitter.setObjectName("workspace_splitter")
        layout.addWidget(splitter, 1)

        left = QtWidgets.QWidget()
        left_layout = QtWidgets.QVBoxLayout(left)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(12)

        card_src = QtWidgets.QFrame()
        card_src.setObjectName("card")
        src_outer = QtWidgets.QVBoxLayout(card_src)
        src_outer.setContentsMargins(14, 14, 14, 14)
        src_outer.setSpacing(10)
        src_title = QtWidgets.QLabel("源代码")
        src_title.setObjectName("card_title")
        src_outer.addWidget(src_title)

        mode_row = QtWidgets.QHBoxLayout()
        self.rb_single = QtWidgets.QRadioButton("单文件")
        self.rb_project = QtWidgets.QRadioButton("工程模式")
        self.rb_project.setToolTip("工程模式支持工程扫描、分层列表与文件顺序调整")
        self.rb_single.setChecked(True)
        self.rb_single.toggled.connect(self._sync_source_mode)
        self.rb_project.toggled.connect(self._sync_source_mode)
        mode_row.addWidget(self.rb_single)
        mode_row.addWidget(self.rb_project)
        mode_row.addStretch(1)
        src_outer.addLayout(mode_row)

        form = QtWidgets.QFormLayout()
        form.setLabelAlignment(QtCore.Qt.AlignLeft)
        form.setFormAlignment(QtCore.Qt.AlignTop)
        form.setHorizontalSpacing(10)
        form.setVerticalSpacing(10)

        c_row = QtWidgets.QHBoxLayout()
        self.ed_c_file = QtWidgets.QLineEdit()
        self.ed_c_file.setPlaceholderText("选择单个 .c 文件")
        self.btn_pick_c_file = QtWidgets.QPushButton("浏览…")
        self.btn_pick_c_file.clicked.connect(self._pick_c_file)
        c_row.addWidget(self.ed_c_file, 1)
        c_row.addWidget(self.btn_pick_c_file)
        form.addRow("C 文件", _wrap_layout(c_row))

        p_row = QtWidgets.QHBoxLayout()
        self.ed_project = QtWidgets.QLineEdit()
        self.ed_project.setPlaceholderText("选择工程根目录（包含 SRC）")
        self.btn_pick_project = QtWidgets.QPushButton("浏览…")
        self.btn_pick_project.clicked.connect(self._pick_project_dir)
        self.btn_load_project = QtWidgets.QPushButton("加载")
        self.btn_load_project.clicked.connect(self._load_project_from_ui)
        self.btn_refresh_term_table = QtWidgets.QPushButton("刷新术语表")
        self.btn_refresh_term_table.setToolTip("扫描当前工程，生成/更新 autodoc_term_table.json，并注入本次运行的符号字典")
        self.btn_refresh_term_table.clicked.connect(self._start_refresh_term_table)
        self.btn_open_term_table = QtWidgets.QPushButton("查看")
        self.btn_open_term_table.setToolTip("查看当前工程的 autodoc_term_table.json 摘要")
        self.btn_open_term_table.clicked.connect(self._show_project_term_table)
        p_row.addWidget(self.ed_project, 1)
        p_row.addWidget(self.btn_pick_project)
        p_row.addWidget(self.btn_load_project)
        p_row.addWidget(self.btn_refresh_term_table)
        p_row.addWidget(self.btn_open_term_table)
        form.addRow("工程目录", _wrap_layout(p_row))

        src_outer.addLayout(form)
        left_layout.addWidget(card_src)

        card_out = QtWidgets.QFrame()
        card_out.setObjectName("card")
        out_outer = QtWidgets.QVBoxLayout(card_out)
        out_outer.setContentsMargins(14, 14, 14, 14)
        out_outer.setSpacing(10)
        out_title = QtWidgets.QLabel("输出")
        out_title.setObjectName("card_title")
        out_outer.addWidget(out_title)

        out_form = QtWidgets.QFormLayout()
        out_form.setHorizontalSpacing(10)
        out_form.setVerticalSpacing(10)

        o_row = QtWidgets.QHBoxLayout()
        self.ed_output = QtWidgets.QLineEdit()
        self.ed_output.setPlaceholderText("选择输出 .docx 路径")
        self.btn_pick_output = QtWidgets.QPushButton("浏览…")
        self.btn_pick_output.clicked.connect(self._pick_output_docx)
        o_row.addWidget(self.ed_output, 1)
        o_row.addWidget(self.btn_pick_output)
        out_form.addRow("Word 输出", _wrap_layout(o_row))

        t_row = QtWidgets.QHBoxLayout()
        self.ed_template = QtWidgets.QLineEdit()
        self.ed_template.setPlaceholderText("可选：模板 .docx（不填使用默认模板）")
        self.btn_pick_template = QtWidgets.QPushButton("浏览…")
        self.btn_pick_template.clicked.connect(self._pick_template_docx)
        t_row.addWidget(self.ed_template, 1)
        t_row.addWidget(self.btn_pick_template)
        out_form.addRow("模板 Docx", _wrap_layout(t_row))
        out_outer.addLayout(out_form)
        left_layout.addWidget(card_out)

        card_update = QtWidgets.QFrame()
        card_update.setObjectName("card")
        update_outer = QtWidgets.QVBoxLayout(card_update)
        update_outer.setContentsMargins(14, 14, 14, 14)
        update_outer.setSpacing(10)
        update_title = QtWidgets.QLabel("文档增量更新")
        update_title.setObjectName("card_title")
        update_outer.addWidget(update_title)

        update_form = QtWidgets.QFormLayout()
        update_form.setHorizontalSpacing(10)
        update_form.setVerticalSpacing(10)

        old_code_row = QtWidgets.QHBoxLayout()
        self.ed_doc_old_code = QtWidgets.QLineEdit()
        self.ed_doc_old_code.setPlaceholderText("旧版本代码目录")
        self.btn_doc_old_code = QtWidgets.QPushButton("浏览…")
        self.btn_doc_old_code.clicked.connect(lambda: self._pick_dir_into(self.ed_doc_old_code, "选择旧代码目录"))
        old_code_row.addWidget(self.ed_doc_old_code, 1)
        old_code_row.addWidget(self.btn_doc_old_code)
        update_form.addRow("旧代码", _wrap_layout(old_code_row))

        new_code_row = QtWidgets.QHBoxLayout()
        self.ed_doc_new_code = QtWidgets.QLineEdit()
        self.ed_doc_new_code.setPlaceholderText("新版本代码目录")
        self.btn_doc_new_code = QtWidgets.QPushButton("浏览…")
        self.btn_doc_new_code.clicked.connect(lambda: self._pick_dir_into(self.ed_doc_new_code, "选择新代码目录"))
        new_code_row.addWidget(self.ed_doc_new_code, 1)
        new_code_row.addWidget(self.btn_doc_new_code)
        update_form.addRow("新代码", _wrap_layout(new_code_row))

        old_doc_row = QtWidgets.QHBoxLayout()
        self.ed_doc_old_doc = QtWidgets.QLineEdit()
        self.ed_doc_old_doc.setPlaceholderText("旧 Word 设计文档")
        self.btn_doc_old_doc = QtWidgets.QPushButton("浏览…")
        self.btn_doc_old_doc.clicked.connect(lambda: self._pick_file_into(self.ed_doc_old_doc, "选择旧 Word 文档", "Word 文档 (*.docx)"))
        old_doc_row.addWidget(self.ed_doc_old_doc, 1)
        old_doc_row.addWidget(self.btn_doc_old_doc)
        update_form.addRow("旧文档", _wrap_layout(old_doc_row))

        decision_row = QtWidgets.QHBoxLayout()
        self.ed_doc_review_decisions = QtWidgets.QLineEdit()
        self.ed_doc_review_decisions.setPlaceholderText("可选：review_decisions.json")
        self.btn_doc_review_decisions = QtWidgets.QPushButton("浏览…")
        self.btn_doc_review_decisions.clicked.connect(lambda: self._pick_file_into(self.ed_doc_review_decisions, "选择 review_decisions.json", "JSON 文件 (*.json);;所有文件 (*.*)"))
        decision_row.addWidget(self.ed_doc_review_decisions, 1)
        decision_row.addWidget(self.btn_doc_review_decisions)
        update_form.addRow("审查决策", _wrap_layout(decision_row))

        mode_row2 = QtWidgets.QHBoxLayout()
        self.cmb_doc_update_mode = QtWidgets.QComboBox()
        self.cmb_doc_update_mode.addItems(["plan-only", "apply-safe", "apply-review"])
        self.chk_doc_renumber_module_csu = QtWidgets.QCheckBox("重排模块 CSU 编号")
        self.chk_doc_renumber_module_csu.setToolTip("默认关闭；开启后 apply-review 会按模块内 H4 顺序重排 CSU ID")
        self.btn_start_doc_update = QtWidgets.QPushButton("运行增量更新")
        self.btn_start_doc_update.clicked.connect(self._start_doc_update)
        mode_row2.addWidget(self.cmb_doc_update_mode)
        mode_row2.addWidget(self.chk_doc_renumber_module_csu)
        mode_row2.addWidget(self.btn_start_doc_update)
        update_form.addRow("模式", _wrap_layout(mode_row2))

        update_outer.addLayout(update_form)
        left_layout.addWidget(card_update)

        card_recent = QtWidgets.QFrame()
        card_recent.setObjectName("card")
        recent_outer = QtWidgets.QVBoxLayout(card_recent)
        recent_outer.setContentsMargins(14, 14, 14, 14)
        recent_outer.setSpacing(6)
        recent_title = QtWidgets.QLabel("最近记录")
        recent_title.setObjectName("card_title")
        recent_outer.addWidget(recent_title)

        self._recent_c_file = QtWidgets.QLabel()
        self._recent_project = QtWidgets.QLabel()
        self._recent_output = QtWidgets.QLabel()
        self._recent_template = QtWidgets.QLabel()
        for lab in (self._recent_c_file, self._recent_project, self._recent_output, self._recent_template):
            lab.setObjectName("muted")
            lab.setWordWrap(True)
            recent_outer.addWidget(lab)

        recent_btns = QtWidgets.QHBoxLayout()
        btn_restore_recent = QtWidgets.QPushButton("恢复上次输入")
        btn_restore_recent.clicked.connect(self._restore_recent_inputs)
        recent_btns.addWidget(btn_restore_recent)
        recent_btns.addStretch(1)
        recent_outer.addLayout(recent_btns)
        left_layout.addWidget(card_recent)
        left_layout.addStretch(1)

        self._right_panel = QtWidgets.QFrame()
        self._right_panel.setObjectName("right_panel")
        right_layout = QtWidgets.QVBoxLayout(self._right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(0)

        self._panel_tabs = QtWidgets.QTabWidget()
        self._panel_tabs.setObjectName("panel_tabs")
        self._panel_tabs.setDocumentMode(True)
        self._panel_tabs.addTab(self._project_tree, "项目")
        self._panel_tabs.addTab(self._step_tree, "步骤")
        self._panel_tabs.addTab(self._log_view, "日志")
        # 术语一致性面板
        from .consistency_panel import ConsistencyPanel
        self._consistency_panel = ConsistencyPanel()
        self._panel_tabs.addTab(self._consistency_panel, "一致性")
        right_layout.addWidget(self._panel_tabs, 1)

        splitter.addWidget(left)
        splitter.addWidget(self._right_panel)
        splitter.setStretchFactor(0, 3)
        splitter.setStretchFactor(1, 2)
        splitter.setSizes([700, 420])

        self._refresh_recent_summary()
        self._sync_source_mode()
        return page

    def _sync_source_mode(self) -> None:
        is_project = bool(getattr(self, "rb_project", None) is not None and self.rb_project.isChecked())
        try:
            self.ed_c_file.setEnabled(not is_project)
            self.btn_pick_c_file.setEnabled(not is_project)
        except Exception:
            pass
        try:
            self.ed_project.setEnabled(is_project)
            self.btn_pick_project.setEnabled(is_project)
            self.btn_load_project.setEnabled(is_project)
            self.btn_refresh_term_table.setEnabled(is_project)
            self.btn_open_term_table.setEnabled(is_project)
        except Exception:
            pass

    def _refresh_recent_summary(self) -> None:
        try:
            data = self.settings_store.load_recent_inputs()
        except Exception:
            data = {}
        c_file = str((data or {}).get("c_file") or "").strip()
        project_dir = str((data or {}).get("project_dir") or "").strip()
        output = str((data or {}).get("output") or "").strip()
        template = str((data or {}).get("template") or "").strip()
        self._recent_c_file.setText(f"C 文件：{c_file or '（无）'}")
        self._recent_project.setText(f"工程目录：{project_dir or '（无）'}")
        self._recent_output.setText(f"Word 输出：{output or '（无）'}")
        self._recent_template.setText(f"模板 Docx：{template or '（无）'}")

    def _build_settings_page(self) -> QtWidgets.QWidget:
        w = QtWidgets.QWidget()
        outer = QtWidgets.QVBoxLayout(w)

        tabs = QtWidgets.QTabWidget()
        outer.addWidget(tabs, 1)

        tab_basic = QtWidgets.QWidget()
        tab_basic.setObjectName("settings_basic_tab")
        basic_outer = QtWidgets.QVBoxLayout(tab_basic)
        basic_outer.setContentsMargins(0, 0, 0, 0)
        basic_outer.setSpacing(0)

        basic_scroll = QtWidgets.QScrollArea()
        basic_scroll.setObjectName("settings_scroll")
        basic_scroll.setWidgetResizable(True)
        basic_scroll.setFrameShape(QtWidgets.QFrame.NoFrame)
        basic_outer.addWidget(basic_scroll, 1)

        basic_body = QtWidgets.QWidget()
        basic_body.setObjectName("settings_basic_body")
        basic_scroll.setWidget(basic_body)

        body_layout = QtWidgets.QVBoxLayout(basic_body)
        body_layout.setContentsMargins(16, 16, 16, 16)
        body_layout.setSpacing(12)

        def _make_settings_card(title_text: str):
            card = QtWidgets.QFrame()
            card.setObjectName("settings_card")
            v = QtWidgets.QVBoxLayout(card)
            v.setContentsMargins(14, 14, 14, 14)
            v.setSpacing(10)
            title = QtWidgets.QLabel(title_text)
            title.setObjectName("settings_card_title")
            v.addWidget(title)
            return card, v

        def _make_form() -> QtWidgets.QFormLayout:
            form = QtWidgets.QFormLayout()
            form.setLabelAlignment(QtCore.Qt.AlignLeft)
            form.setFormAlignment(QtCore.Qt.AlignTop)
            form.setHorizontalSpacing(12)
            form.setVerticalSpacing(10)
            return form

        self.set_req_prefix = QtWidgets.QLineEdit()
        self.set_open_after = QtWidgets.QCheckBox("完成后自动打开（仅 Windows）")
        card_common, common_layout = _make_settings_card("常用")
        common_form = _make_form()
        common_form.addRow("需求ID前缀：", self.set_req_prefix)
        common_form.addRow("", self.set_open_after)
        common_layout.addLayout(common_form)
        body_layout.addWidget(card_common)

        recent_row = QtWidgets.QHBoxLayout()
        btn_restore_recent = QtWidgets.QPushButton("恢复上次输入")
        btn_restore_recent.clicked.connect(self._restore_recent_inputs)
        btn_save_recent = QtWidgets.QPushButton("保存当前为上次输入")
        btn_save_recent.clicked.connect(self._save_recent_inputs)
        recent_row.addWidget(btn_restore_recent)
        recent_row.addWidget(btn_save_recent)
        recent_row.addStretch(1)
        card_recent, recent_layout = _make_settings_card("最近记录")
        recent_help = QtWidgets.QLabel("用于快速恢复/保存单文件、工程目录与输出路径等输入。")
        recent_help.setObjectName("muted")
        recent_help.setWordWrap(True)
        recent_layout.addWidget(recent_help)
        recent_layout.addLayout(recent_row)
        body_layout.addWidget(card_recent)

        self.set_exclude_dirs = QtWidgets.QPlainTextEdit()
        self.set_exclude_dirs.setPlaceholderText("每行一个目录名（大小写不敏感），例如：debug、release、.git")
        self.set_exclude_dirs.setFixedHeight(110)

        kw_row = QtWidgets.QHBoxLayout()
        self.set_mid_keywords = QtWidgets.QLineEdit()
        self.set_mid_keywords.setPlaceholderText("例如：common")
        self.set_drv_keywords = QtWidgets.QLineEdit()
        self.set_drv_keywords.setPlaceholderText("例如：dspdriver")
        kw_row.addWidget(QtWidgets.QLabel("中间层关键词："))
        kw_row.addWidget(self.set_mid_keywords, 1)
        kw_row.addWidget(QtWidgets.QLabel("驱动层关键词："))
        kw_row.addWidget(self.set_drv_keywords, 1)
        card_rules, rules_layout = _make_settings_card("工程规则")
        rules_form = _make_form()
        rules_form.addRow("排除目录：", self.set_exclude_dirs)
        rules_form.addRow("分层关键词：", _wrap_layout(kw_row))
        rules_layout.addLayout(rules_form)
        body_layout.addWidget(card_rules)

        self.set_domain_glossary = QtWidgets.QPlainTextEdit()
        self.set_domain_glossary.setPlaceholderText("每行一条：KEY=中文，例如：PBIT=飞行前BIT（支持 JSON dict）")
        self.set_domain_glossary.setFixedHeight(140)
        card_glossary, glossary_layout = _make_settings_card("术语表覆盖")
        glossary_layout.addWidget(self.set_domain_glossary)
        body_layout.addWidget(card_glossary)

        self.set_symbol_dict = QtWidgets.QPlainTextEdit()
        self.set_symbol_dict.setPlaceholderText("每行一条：符号=中文，例如：frceRst=强制复位（支持 JSON 分区字典）")
        self.set_symbol_dict.setFixedHeight(160)
        card_symbol_dict, symbol_dict_layout = _make_settings_card("符号字典覆盖")
        symbol_dict_layout.addWidget(self.set_symbol_dict)
        body_layout.addWidget(card_symbol_dict)

        body_layout.addStretch(1)
        tabs.addTab(tab_basic, "基本")

        tab_ai = QtWidgets.QWidget()
        ai = QtWidgets.QFormLayout(tab_ai)
        self.set_ai_mode = QtWidgets.QComboBox()
        self.set_ai_mode.addItems(["无 AI", "开启 AI"])
        self.set_ai_mode.currentIndexChanged.connect(self._on_ai_mode_changed)
        self.set_ai_provider = QtWidgets.QComboBox()
        self.set_ai_provider.addItems(["local", "deepseek", "openai", "openrouter", "anthropic", "compshare"])
        self.set_ai_provider.currentTextChanged.connect(self._on_provider_changed)
        self.set_ai_model = QtWidgets.QComboBox()
        self.set_ai_model.setEditable(True)
        self.set_ai_model.setInsertPolicy(QtWidgets.QComboBox.NoInsert)
        self.set_ai_model.lineEdit().editingFinished.connect(self._check_store_secret)
        self.set_ai_base = QtWidgets.QLineEdit()
        self.set_ai_key = QtWidgets.QLineEdit()
        self.set_ai_key.setEchoMode(QtWidgets.QLineEdit.Password)
        self.set_ai_key.setPlaceholderText("需要鉴权的 API 请填写 Key；刷新模型会使用该 Key")
        self.set_ai_num_ctx = QtWidgets.QSpinBox()
        self.set_ai_num_ctx.setRange(0, 1_000_000)
        self.set_ai_read_timeout = QtWidgets.QSpinBox()
        self.set_ai_read_timeout.setRange(5, 600)
        self.set_ai_read_timeout.setValue(40)
        self.set_ai_read_timeout.setSuffix(" s")
        self.set_ai_read_timeout.setToolTip("等待 AI 返回内容的读取超时，默认 40 秒。")
        self.set_ai_workers = QtWidgets.QSpinBox()
        self.set_ai_workers.setRange(1, 16)
        self.set_ai_workers.setValue(1)
        self.set_ai_workers.setToolTip("小模型默认串行生成以降低限流和质量波动；必要时再手动提高。")
        self.set_force_ai = QtWidgets.QCheckBox()
        self.set_verbose = QtWidgets.QCheckBox("详细日志（写入 tool.log）")
        self.set_no_proxy = QtWidgets.QCheckBox("禁用代理（直连 API）")
        self.set_no_proxy.setToolTip("勾选后跳过所有代理检测，直接连接 AI API。如果代理导致超时，请勾选此项。")
        self.set_ai_one_call = QtWidgets.QCheckBox()
        self.set_auto_disable_large_one_call = QtWidgets.QCheckBox()
        self.set_ai_logic_format = QtWidgets.QComboBox()
        self.set_ai_logic_format.addItems(["json"])
        self.set_ai_logic_policy = QtWidgets.QComboBox()
        self.set_ai_logic_policy.addItems(["hybrid"])
        self.set_logic_ignore_comment = QtWidgets.QCheckBox("忽略逻辑注释（仅按名称直译）")
        ai.addRow("AI：", self.set_ai_mode)
        ai.addRow("调用方式：", self.set_ai_provider)
        model_row = QtWidgets.QHBoxLayout()
        model_row.addWidget(self.set_ai_model, 1)
        self.btn_refresh_ai_models = QtWidgets.QPushButton("刷新模型")
        self.btn_refresh_ai_models.clicked.connect(self._refresh_ai_models)
        model_row.addWidget(self.btn_refresh_ai_models)
        ai.addRow("Model：", _wrap_layout(model_row))
        ai.addRow("固定地址：", self.set_ai_base)
        ai.addRow("API Key：", self.set_ai_key)
        ai.addRow("Context(num_ctx)：", self.set_ai_num_ctx)
        ai.addRow("Read Timeout：", self.set_ai_read_timeout)
        ai.addRow("AI 并发数：", self.set_ai_workers)
        ai.addRow("", self.set_verbose)
        ai.addRow("", self.set_no_proxy)
        ai.addRow("", self.set_logic_ignore_comment)

        ai_btns = QtWidgets.QHBoxLayout()
        btn_ai_defaults = QtWidgets.QPushButton("恢复 AI 默认")
        btn_ai_defaults.clicked.connect(self._reset_ai_defaults)
        ai_btns.addWidget(btn_ai_defaults)
        ai_btns.addStretch(1)
        ai.addRow("快捷操作：", _wrap_layout(ai_btns))
        tabs.addTab(tab_ai, "AI")

        tab_perf = QtWidgets.QWidget()
        perf = QtWidgets.QFormLayout(tab_perf)
        self.set_prefilter = QtWidgets.QCheckBox("工程扫描预筛选无注释/无函数文件")
        self.set_incremental = QtWidgets.QCheckBox("增量模式")
        self.set_incremental.setToolTip("仅工程生成生效：复用上次状态，只重新生成变更函数。默认关闭。")
        self.set_preprocess_workers = QtWidgets.QSpinBox()
        self.set_preprocess_workers.setRange(0, 256)
        self.set_log_every_n = QtWidgets.QSpinBox()
        self.set_log_every_n.setRange(1, 1_000_000)
        perf.addRow("", self.set_prefilter)
        perf.addRow("", self.set_incremental)
        perf.addRow("并行预处理线程(0=自动)：", self.set_preprocess_workers)
        perf.addRow("日志采样间隔(N)：", self.set_log_every_n)
        tabs.addTab(tab_perf, "性能")

        tab_adv = QtWidgets.QWidget()
        adv = QtWidgets.QVBoxLayout(tab_adv)
        self.tbl_extra_params = QtWidgets.QTableWidget(0, 2)
        self.tbl_extra_params.setHorizontalHeaderLabels(["Key", "Value"])
        self.tbl_extra_params.horizontalHeader().setStretchLastSection(True)
        self.tbl_extra_params.verticalHeader().setVisible(False)
        self.tbl_extra_params.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tbl_extra_params.setEditTriggers(QtWidgets.QAbstractItemView.DoubleClicked | QtWidgets.QAbstractItemView.EditKeyPressed)
        self.tbl_extra_params.setItemDelegateForColumn(1, ExtraParamsValueDelegate(self.tbl_extra_params))
        adv.addWidget(self.tbl_extra_params, 1)

        adv_btns = QtWidgets.QHBoxLayout()
        btn_add = QtWidgets.QPushButton("添加参数")
        btn_add.clicked.connect(self._add_extra_param_row)
        btn_del = QtWidgets.QPushButton("删除选中")
        btn_del.clicked.connect(self._delete_selected_extra_param_rows)
        btn_defaults = QtWidgets.QPushButton("恢复默认")
        btn_defaults.clicked.connect(self._reset_extra_params_defaults)
        adv_btns.addWidget(btn_add)
        adv_btns.addWidget(btn_del)
        adv_btns.addStretch(1)
        adv_btns.addWidget(btn_defaults)
        adv.addLayout(adv_btns)
        tabs.addTab(tab_adv, "高级")

        btns = QtWidgets.QHBoxLayout()
        btn_save = QtWidgets.QPushButton("保存设置")
        btn_save.clicked.connect(self._save_settings_from_ui)
        btns.addStretch(1)
        btns.addWidget(btn_save)
        outer.addLayout(btns)
        return w

    def _build_help_page(self) -> QtWidgets.QWidget:
        w = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(w)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        toolbar = QtWidgets.QFrame()
        toolbar.setObjectName("help_toolbar")
        tb = QtWidgets.QHBoxLayout(toolbar)
        tb.setContentsMargins(12, 8, 12, 8)
        tb.setSpacing(8)

        btn_back = QtWidgets.QToolButton()
        btn_back.setText("←")
        btn_back.setToolTip("后退")
        btn_back.clicked.connect(lambda: self._help_browser.backward())
        btn_fwd = QtWidgets.QToolButton()
        btn_fwd.setText("→")
        btn_fwd.setToolTip("前进")
        btn_fwd.clicked.connect(lambda: self._help_browser.forward())

        self._help_search = QtWidgets.QLineEdit()
        self._help_search.setPlaceholderText("搜索帮助内容...")
        self._help_search.setClearButtonEnabled(True)
        self._help_search.returnPressed.connect(self._on_help_search)
        self._help_search.setMinimumWidth(200)

        btn_zoom_in = QtWidgets.QToolButton()
        btn_zoom_in.setText("A+")
        btn_zoom_in.setToolTip("放大")
        btn_zoom_in.clicked.connect(lambda: self._help_browser.zoomIn(1))
        btn_zoom_out = QtWidgets.QToolButton()
        btn_zoom_out.setText("A-")
        btn_zoom_out.setToolTip("缩小")
        btn_zoom_out.clicked.connect(lambda: self._help_browser.zoomOut(1))

        for btn in (btn_back, btn_fwd, self._help_search, btn_zoom_in, btn_zoom_out):
            tb.addWidget(btn)

        layout.addWidget(toolbar)

        self._help_browser = QtWidgets.QTextBrowser()
        self._help_browser.setOpenExternalLinks(True)
        layout.addWidget(self._help_browser, 1)

        help_path = _asset_path("help.html")
        if os.path.exists(help_path):
            self._help_browser.setSource(QtCore.QUrl.fromLocalFile(help_path))
        else:
            self._help_browser.setHtml("<h2>帮助</h2><p>未找到本地帮助文档。</p>")
        return w

    def _on_help_search(self) -> None:
        keyword = self._help_search.text().strip()
        if not keyword:
            return
        if not self._help_browser.find(keyword):
            cursor = self._help_browser.textCursor()
            cursor.movePosition(QtGui.QTextCursor.Start)
            self._help_browser.setTextCursor(cursor)
            self._help_browser.find(keyword)

    def _apply_settings_to_ui(self) -> None:
        if not getattr(self, "_settings_page_built", False):
            return
        s = self.settings
        self.set_req_prefix.setText(s.req_id_prefix)
        self.set_open_after.setChecked(bool(s.open_after_done))
        self.set_exclude_dirs.setPlainText("\n".join([str(x).strip() for x in (s.exclude_dirs or []) if str(x).strip()]))
        self.set_mid_keywords.setText(",".join([str(x).strip() for x in (s.mid_dir_keywords or []) if str(x).strip()]))
        self.set_drv_keywords.setText(",".join([str(x).strip() for x in (s.drv_dir_keywords or []) if str(x).strip()]))
        try:
            self.set_domain_glossary.setPlainText(str(getattr(s, "domain_glossary_text", "") or ""))
        except Exception:
            pass
        try:
            self.set_symbol_dict.setPlainText(str(getattr(s, "symbol_dict_text", "") or ""))
        except Exception:
            pass

        self.set_ai_mode.setCurrentIndex(normalize_ai_mode(getattr(s, "ai_mode", 0)))
        self.set_ai_provider.setCurrentText(s.ai_provider or "local")
        self._set_ai_model_items([], preferred=s.ai_model)
        self.set_ai_base.setText(s.ai_api_base or LOCAL_LLM_API_BASE)
        self.set_ai_key.setText(s.ai_api_key)
        self.set_ai_num_ctx.setValue(int(s.ai_num_ctx))
        self.set_ai_read_timeout.setValue(max(5, min(600, int(getattr(s, "ai_read_timeout", 40) or 40))))
        self.set_ai_workers.setValue(max(1, min(16, int(getattr(s, "ai_workers", 2) or 2))))
        self.set_force_ai.setChecked(False)
        self.set_verbose.setChecked(bool(s.verbose))
        self.set_no_proxy.setChecked(bool(getattr(s, "no_proxy", False)))
        self.set_ai_one_call.setChecked(False)
        self.set_auto_disable_large_one_call.setChecked(True)
        self.set_ai_logic_format.setCurrentText("json")
        self.set_ai_logic_policy.setCurrentText("hybrid")
        try:
            self.set_logic_ignore_comment.setChecked(not bool(getattr(s, "logic_use_comment", True)))
        except Exception:
            pass

        self.set_prefilter.setChecked(bool(s.prefilter_project_files))
        self.set_incremental.setChecked(bool(getattr(s, "incremental", False)))
        self.set_preprocess_workers.setValue(int(s.preprocess_workers))
        self.set_log_every_n.setValue(int(s.log_every_n))

        self._load_extra_params_table(s.extra_params or {})
        self._on_ai_mode_changed(self.set_ai_mode.currentIndex())

    def _save_settings_from_ui(self) -> None:
        extra_params = self._read_extra_params_table()
        exclude_dirs = [x.strip() for x in self.set_exclude_dirs.toPlainText().replace(",", "\n").splitlines() if x.strip()]
        mid_keywords = [x.strip() for x in (self.set_mid_keywords.text() or "").replace("，", ",").replace(";", ",").split(",") if x.strip()]
        drv_keywords = [x.strip() for x in (self.set_drv_keywords.text() or "").replace("，", ",").replace(";", ",").split(",") if x.strip()]
        domain_glossary_text = ""
        try:
            domain_glossary_text = str(self.set_domain_glossary.toPlainText() or "")
        except Exception:
            domain_glossary_text = ""
        symbol_dict_text = ""
        try:
            symbol_dict_text = str(self.set_symbol_dict.toPlainText() or "")
        except Exception:
            symbol_dict_text = ""
        s = AppSettings(
            section_prefix="5.1.1.",
            req_id_prefix=self.set_req_prefix.text().strip() or "D/R_SDD01_",
            only_with_comment=False,
            include_locals=True,
            include_logic=True,
            open_after_done=bool(self.set_open_after.isChecked()),
            ai_mode=normalize_ai_mode(self.set_ai_mode.currentIndex()),
            ai_provider=str(self.set_ai_provider.currentText() or "local"),
            ai_model=str(self.set_ai_model.currentText() or "").strip(),
            ai_api_base=self._normalize_ai_base_for_gui(self.set_ai_base.text()),
            ai_api_key=str(self.set_ai_key.text() or "").strip(),
            ai_num_ctx=int(self.set_ai_num_ctx.value()),
            ai_read_timeout=int(self.set_ai_read_timeout.value()),
            ai_workers=int(self.set_ai_workers.value()),
            force_ai=False,
            verbose=bool(self.set_verbose.isChecked()),
            no_proxy=bool(self.set_no_proxy.isChecked()),
            ai_one_call=False,
            auto_disable_large_one_call=True,
            ai_logic_format="json",
            ai_logic_policy="hybrid",
            logic_use_comment=(not bool(self.set_logic_ignore_comment.isChecked())),
            prefilter_project_files=bool(self.set_prefilter.isChecked()),
            incremental=bool(self.set_incremental.isChecked()),
            preprocess_workers=int(self.set_preprocess_workers.value()),
            log_every_n=int(self.set_log_every_n.value()),
            extra_params=extra_params,
            exclude_dirs=exclude_dirs,
            mid_dir_keywords=mid_keywords,
            drv_dir_keywords=drv_keywords,
            domain_glossary_text=domain_glossary_text,
            symbol_dict_text=symbol_dict_text,
        )
        self.settings_store.save(s)
        self.settings = s
        QtWidgets.QMessageBox.information(self, "已保存", "设置已保存。")

    def _pick_c_file(self) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "选择 C 文件", "", "C 源文件 (*.c);;所有文件 (*.*)")
        if path:
            self.ed_c_file.setText(path)
            try:
                self.rb_single.setChecked(True)
            except Exception:
                pass

    def _pick_project_dir(self) -> None:
        path = QtWidgets.QFileDialog.getExistingDirectory(self, "选择工程目录")
        if path:
            self.ed_project.setText(path)
            try:
                self.rb_project.setChecked(True)
            except Exception:
                pass
            self._load_project_from_ui()

    def _pick_dir_into(self, line_edit: QtWidgets.QLineEdit, title: str) -> None:
        path = QtWidgets.QFileDialog.getExistingDirectory(self, title)
        if path:
            line_edit.setText(path)

    def _pick_file_into(self, line_edit: QtWidgets.QLineEdit, title: str, file_filter: str) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, title, "", file_filter)
        if path:
            line_edit.setText(path)

    def _pick_output_docx(self) -> None:
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "选择输出 Docx", "", "Word 文档 (*.docx)")
        if path:
            self.ed_output.setText(path)

    def _pick_template_docx(self) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "选择模板 Docx（可选）", "", "Word 文档 (*.docx);;所有文件 (*.*)")
        if path:
            self.ed_template.setText(path)

    def _show_about(self) -> None:
        QtWidgets.QMessageBox.about(
            self,
            "关于",
            "<h3>CSCI 详细设计生成器</h3>"
            f"<p>版本：{self._app_version or '未标注版本'}</p>"
            "<p>PyQt5 5.12 GUI</p>"
            "<p>支持：设置 / 帮助 / 关于 / 步骤级执行状态</p>"
            "<p>本版本已加入 AI 并发增强、连接复用与结果缓存优化。</p>",
        )

    def _save_log(self) -> None:
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "导出日志", "", "文本文件 (*.txt)")
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self._log_view.toPlainText())
            QtWidgets.QMessageBox.information(self, "完成", f"日志已导出：{path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"导出失败：{e}")

    def _restore_recent_inputs(self) -> None:
        data = self.settings_store.load_recent_inputs()
        if not any((data.get("c_file"), data.get("project_dir"), data.get("output"), data.get("template"))):
            QtWidgets.QMessageBox.information(self, "提示", "没有找到上次输入记录。")
            return
        if data.get("c_file"):
            self.ed_c_file.setText(data["c_file"])
            try:
                if not data.get("project_dir"):
                    self.rb_single.setChecked(True)
            except Exception:
                pass
        if data.get("project_dir"):
            self.ed_project.setText(data["project_dir"])
            self._load_project_from_ui()
            try:
                self.rb_project.setChecked(True)
            except Exception:
                pass
        if data.get("output"):
            self.ed_output.setText(data["output"])
        if data.get("template"):
            self.ed_template.setText(data["template"])
        try:
            self._sync_source_mode()
        except Exception:
            pass
        try:
            self._refresh_recent_summary()
        except Exception:
            pass
        self._stack.setCurrentWidget(self._page_home)
        QtWidgets.QMessageBox.information(self, "完成", "已恢复上次输入。")

    def _save_recent_inputs(self) -> None:
        self.settings_store.save_recent_inputs(
            {
                "c_file": self.ed_c_file.text().strip(),
                "project_dir": self.ed_project.text().strip(),
                "output": self.ed_output.text().strip(),
                "template": self.ed_template.text().strip(),
            }
        )
        try:
            self._refresh_recent_summary()
        except Exception:
            pass

    def _collect_task(self) -> TaskSpec:
        c_file = self.ed_c_file.text().strip()
        proj = self.ed_project.text().strip()
        mode = "project" if (getattr(self, "rb_project", None) is not None and self.rb_project.isChecked()) else "single"
        return TaskSpec(
            mode=mode,
            c_file=c_file,
            project_dir=proj,
            output=self.ed_output.text().strip(),
            template_path=self.ed_template.text().strip(),
            project_file_order=(self._get_project_file_order() if mode == "project" else None),
        )

    def _start_doc_update(self) -> None:
        if self._thread is not None:
            QtWidgets.QMessageBox.information(self, "提示", "任务运行中。")
            return
        self._save_settings_from_ui_silent()
        settings = self.settings
        task = TaskSpec(
            mode="doc_update",
            c_file="",
            project_dir="",
            output=self.ed_output.text().strip(),
            template_path=self.ed_template.text().strip(),
            old_code=self.ed_doc_old_code.text().strip(),
            new_code=self.ed_doc_new_code.text().strip(),
            old_doc=self.ed_doc_old_doc.text().strip(),
            review_decisions=self.ed_doc_review_decisions.text().strip(),
            doc_update_mode=str(self.cmb_doc_update_mode.currentText() or "plan-only"),
            renumber_module_csu=bool(self.chk_doc_renumber_module_csu.isChecked()),
        )
        steps = [
            StepDef("validate", "校验输入"),
            StepDef("diff", "对比代码"),
            StepDef("plan", "生成计划"),
            StepDef("apply", "应用更新"),
            StepDef("report", "生成报告"),
        ]
        self._reset_steps(steps)
        self._resume_state = None
        self._set_running(True, note="文档增量更新中...")
        self._run_in_thread(
            impl=DocUpdateWorker(backend=self.backend, task=task, settings=settings),
            kind="doc_update",
        )

    def _reset_steps(self, steps: list[StepDef]) -> None:
        model = QtGui.QStandardItemModel()
        self._step_model = model
        root = model.invisibleRootItem()
        self._step_items = {}
        self._func_items = {}

        by_parent = {}
        for s in steps:
            by_parent.setdefault(s.parent_id, []).append(s)

        def add_children(parent_item, parent_id):
            for sd in by_parent.get(parent_id, []):
                item = QtGui.QStandardItem()
                item.setData(sd.step_id, QtCore.Qt.UserRole + 1)
                item.setData(sd.name, QtCore.Qt.UserRole + 2)
                item.setData("pending", QtCore.Qt.UserRole + 3)
                item.setText(f"[{_STATUS_TEXT['pending']}] {sd.name}")
                parent_item.appendRow(item)
                self._step_items[sd.step_id] = item
                add_children(item, sd.step_id)

        add_children(root, None)
        self._step_tree.setModel(model)
        self._step_tree.expandAll()

    def _ensure_child_status_item(self, parent: QtGui.QStandardItem, key: str, title: str) -> QtGui.QStandardItem:
        item = parent.data(QtCore.Qt.UserRole + 50)
        by_key = item if isinstance(item, dict) else {}
        if not by_key:
            by_key = {}
            parent.setData(by_key, QtCore.Qt.UserRole + 50)
        if key in by_key:
            return by_key[key]
        child = QtGui.QStandardItem(f"[{_STATUS_TEXT['pending']}] {title}")
        child.setEditable(False)
        child.setData(title, QtCore.Qt.UserRole + 2)
        child.setData("pending", QtCore.Qt.UserRole + 3)
        parent.appendRow(child)
        by_key[key] = child
        return child

    def _set_item_status(self, item: QtGui.QStandardItem, status: str, *, name: Optional[str] = None) -> None:
        title = name if name is not None else (item.data(QtCore.Qt.UserRole + 2) or item.text())
        status_text = _STATUS_TEXT.get(status, str(status))
        item.setData(status, QtCore.Qt.UserRole + 3)
        item.setText(f"[{status_text}] {title}")
        f = item.font()
        f.setBold(status == "running")
        item.setFont(f)

    def _handle_detail_event(self, payload: dict) -> None:
        et = str(payload.get("type") or "")
        func_name = str(payload.get("func_name") or "")
        file_path = str(payload.get("file") or payload.get("source") or "")
        func_index = payload.get("func_index")
        func_pos = payload.get("func_pos")

        parent = self._step_items.get("generate")
        if parent is None:
            # fallback: attach to root
            model = self._step_tree.model()
            if isinstance(model, QtGui.QStandardItemModel):
                parent = model.invisibleRootItem()
            else:
                return

        if et == "func_start":
            key = f"{file_path}|{func_pos}|{func_index}"
            display = f"{func_name}" if func_name else f"func#{func_index}"
            if file_path:
                display = f"{display}  ({os.path.basename(file_path)})"
            item = QtGui.QStandardItem(f"[{_STATUS_TEXT['running']}] {display}")
            item.setEditable(False)
            item.setData(display, QtCore.Qt.UserRole + 2)
            item.setData("running", QtCore.Qt.UserRole + 3)
            item.setData(key, QtCore.Qt.UserRole + 60)
            parent.appendRow(item)
            self._func_items[key] = item

            self._ensure_child_status_item(item, "io", "参数提取")
            self._ensure_child_status_item(item, "locals", "局部变量提取")
            self._ensure_child_status_item(item, "logic", "流程语句润色")

            if parent.rowCount() > 500:
                parent.removeRow(0)

            self._step_tree.expand(parent.index())
            self._step_tree.expand(item.index())
            self._step_tree.scrollTo(item.index(), QtWidgets.QAbstractItemView.PositionAtCenter)
            return

        if et == "progress_init":
            try:
                total = int(payload.get("total") or 0)
            except Exception:
                total = 0
            self._progress_state["total"] = max(0, int(total))
            self._progress_state["done"] = 0
            self._progress_state["start_ts"] = time.monotonic()
            if total > 0:
                self._progress.setRange(0, int(total))
                self._progress.setValue(0)
            else:
                # 0/未知：保持忙碌样式
                self._progress.setRange(0, 0)
            self._refresh_progress_text()
            return

        if et == "ai_regression_start":
            key = f"{file_path}|{func_pos}|{func_index}"
            item = self._func_items.get(key)
            if item is None:
                for i in range(parent.rowCount() - 1, -1, -1):
                    it = parent.child(i)
                    if it is None:
                        continue
                    if (it.data(QtCore.Qt.UserRole + 3) or "") == "running":
                        item = it
                        break
            if item is None:
                return
            round_idx = int(payload.get("round") or 0)
            reg_item = self._ensure_child_status_item(item, "regression", "AI 回归补跑")
            title = f"AI 回归补跑（第 {round_idx} 轮）" if round_idx > 0 else "AI 回归补跑"
            self._set_item_status(reg_item, "running", name=title)
            return

        if et == "ai_regression_end":
            key = f"{file_path}|{func_pos}|{func_index}"
            item = self._func_items.get(key)
            if item is None:
                return
            round_idx = int(payload.get("round") or 0)
            improved = bool(payload.get("improved"))
            ok = bool(payload.get("ok"))
            reg_item = self._ensure_child_status_item(item, "regression", "AI 回归补跑")
            if not ok:
                self._set_item_status(reg_item, "failed", name=f"AI 回归补跑（第 {round_idx} 轮失败）")
            else:
                suffix = "已改善" if improved else "未改善"
                self._set_item_status(reg_item, "success", name=f"AI 回归补跑（第 {round_idx} 轮，{suffix}）")
            return

        if et == "ai_bundle_end":
            key = f"{file_path}|{func_pos}|{func_index}"
            current_item = self._func_items.get(key)
            if current_item is None and (not file_path) and (not func_pos) and (not func_index):
                for i in range(parent.rowCount() - 1, -1, -1):
                    it = parent.child(i)
                    if it is None:
                        continue
                    if (it.data(QtCore.Qt.UserRole + 3) or "") == "running":
                        current_item = it
                        break
            if current_item is None:
                return

            ok = bool(payload.get("ok"))
            params_expected = int(payload.get("params_expected") or 0)
            params_got = int(payload.get("params_got") or 0)
            locals_expected = int(payload.get("locals_expected") or 0)
            locals_got = int(payload.get("locals_got") or 0)
            logic_expected = int(payload.get("logic_expected") or 0)
            logic_got = int(payload.get("logic_got") or 0)

            io_item = self._ensure_child_status_item(current_item, "io", "参数提取")
            locals_item = self._ensure_child_status_item(current_item, "locals", "局部变量提取")
            logic_item = self._ensure_child_status_item(current_item, "logic", "流程语句润色")

            if not ok:
                self._set_item_status(io_item, "failed", name="参数提取（AI 失败）")
                self._set_item_status(locals_item, "failed", name="局部变量提取（AI 失败）")
                self._set_item_status(logic_item, "failed", name="流程语句润色（AI 失败）")
                return

            def status_for(expected: int, got: int) -> str:
                if expected <= 0:
                    return "success"
                return "success" if got > 0 else "failed"

            self._set_item_status(io_item, status_for(params_expected, params_got), name=f"参数提取（{params_got}/{params_expected}）")
            self._set_item_status(locals_item, status_for(locals_expected, locals_got), name=f"局部变量提取（{locals_got}/{locals_expected}）")
            self._set_item_status(logic_item, status_for(logic_expected, logic_got), name=f"流程语句润色（{logic_got}/{logic_expected}）")

            current_item.setData(
                {
                    "params_expected": params_expected,
                    "params_got": params_got,
                    "locals_expected": locals_expected,
                    "locals_got": locals_got,
                    "logic_expected": logic_expected,
                    "logic_got": logic_got,
                },
                QtCore.Qt.UserRole + 70,
            )
            return

        if et == "func_end":
            key = f"{file_path}|{func_pos}|{func_index}"
            item = self._func_items.get(key)
            if item is None:
                # fallback: last running
                for i in range(parent.rowCount() - 1, -1, -1):
                    it = parent.child(i)
                    if it is None:
                        continue
                    if (it.data(QtCore.Qt.UserRole + 3) or "") == "running":
                        item = it
                        break
            if item is None:
                return

            ok = bool(payload.get("ok"))
            self._set_item_status(item, "success" if ok else "failed")

            # 进度条推进（百分比 + ETA）
            try:
                total = self._progress_state.get("total")
                if total is not None:
                    self._progress_state["done"] = int(self._progress_state.get("done") or 0) + 1
                    if int(total) > 0:
                        self._progress.setRange(0, int(total))
                        self._progress.setValue(min(int(self._progress_state["done"]), int(total)))
                    self._refresh_progress_text()
            except Exception:
                pass

            io_ok = payload.get("io_ok")
            locals_ok = payload.get("locals_ok")
            logic_ok = payload.get("logic_ok")

            ai_stats = item.data(QtCore.Qt.UserRole + 70)
            ai_stats = ai_stats if isinstance(ai_stats, dict) else {}
            locals_expected = int(ai_stats.get("locals_expected") or 0)
            locals_got = int(ai_stats.get("locals_got") or 0)
            logic_expected = int(ai_stats.get("logic_expected") or 0)
            logic_got = int(ai_stats.get("logic_got") or 0)
            if io_ok is not None:
                self._set_item_status(self._ensure_child_status_item(item, "io", "参数提取"), "success" if bool(io_ok) else "failed")
            if locals_ok is not None:
                # 若 AI 明确"需要补全"但完全没返回，保留失败；否则"无局部变量"也视为成功
                if locals_expected > 0 and locals_got == 0:
                    pass
                else:
                    self._set_item_status(
                        self._ensure_child_status_item(item, "locals", "局部变量提取"),
                        "success" if bool(locals_ok) else "failed",
                    )
            if logic_ok is not None:
                # 若 AI 明确"需要润色"但完全没返回，保留失败；否则"无需润色"也视为成功
                if logic_expected > 0 and logic_got == 0:
                    pass
                else:
                    self._set_item_status(
                        self._ensure_child_status_item(item, "logic", "流程语句润色"),
                        "success" if bool(logic_ok) else "failed",
                    )
            return

        if et == "func_failure":
            # 函数失败记录
            task = dict(payload.get("task") or {})
            # Strip heavy file_context (424KB+ per func, shared across file)
            # to prevent failures.json from ballooning (44 funcs → 28MB).
            func_data = dict(task.get("func_data") or {})
            if "file_context" in func_data:
                func_data["file_context"] = {
                    "source_file": (func_data["file_context"] or {}).get("source_file", ""),
                    "_stripped": True,
                }
                task["func_data"] = func_data
            failure = {
                "func_name": func_name,
                "file_path": file_path,
                "error_type": str(payload.get("error_type") or "unknown"),
                "error_message": str(payload.get("error_message") or ""),
                "task": task,
            }
            self._failed_functions.append(failure)
            return

        if et == "consistency_report":
            # 术语一致性报告
            try:
                self._consistency_panel.update_report({
                    "score": float(payload.get("score") or 0),
                    "total_symbols": int(payload.get("total_symbols") or 0),
                    "consistent_symbols": int(payload.get("total_symbols") or 0) - int(payload.get("inconsistencies") or 0),
                    "inconsistencies": [],  # 简化版，完整报告需要额外数据
                    "symbol_dict_conflicts": [],
                })
                # 切换到一致性标签页
                self._panel_tabs.setCurrentWidget(self._consistency_panel)
            except Exception:
                pass
            return

    def _fmt_eta(self, seconds: Optional[float]) -> str:
        if seconds is None:
            return "--"
        try:
            s = int(round(max(0.0, float(seconds))))
        except Exception:
            return "--"
        h, rem = divmod(s, 3600)
        m, sec = divmod(rem, 60)
        if h > 0:
            return f"{h:02d}:{m:02d}:{sec:02d}"
        return f"{m:02d}:{sec:02d}"

    def _refresh_progress_text(self) -> None:
        total = self._progress_state.get("total")
        done = int(self._progress_state.get("done") or 0)
        if total is None or int(total) <= 0:
            self._progress.setFormat("")
            return
        total_i = int(total)
        done_i = min(done, total_i)
        pct = (done_i / total_i) * 100.0 if total_i > 0 else 0.0
        eta = None
        if done_i > 0 and self._progress_state.get("start_ts") is not None:
            elapsed = time.monotonic() - float(self._progress_state["start_ts"])
            eta = (elapsed / done_i) * max(0, total_i - done_i)
        self._progress.setFormat(f"{pct:4.1f}%  ({done_i}/{total_i})  剩余 {self._fmt_eta(eta)}")

    def _set_step_status(self, step_id: str, status: str) -> None:
        item = self._step_items.get(step_id)
        if item is None:
            return
        name = item.data(QtCore.Qt.UserRole + 2) or ""
        status_text = _STATUS_TEXT.get(status, str(status))
        item.setData(status, QtCore.Qt.UserRole + 3)
        item.setText(f"[{status_text}] {name}")
        f = item.font()
        f.setBold(status == "running")
        item.setFont(f)

    def _append_log(self, text: str) -> None:
        if not text:
            return
        self._log_view.appendPlainText(str(text).rstrip("\n"))

    def _set_running(self, running: bool, *, note: str) -> None:
        self.act_start.setEnabled(not running)
        self.act_update_csu.setEnabled(not running)
        self.act_refresh_term_table.setEnabled(not running)
        self.act_show_term_table.setEnabled(not running)
        self.act_stop.setEnabled(running)
        self.act_continue.setEnabled((not running) and bool(self._resume_state))
        self._status_label.setText(note)
        for w in (
            getattr(self, "ed_c_file", None),
            getattr(self, "ed_project", None),
            getattr(self, "ed_output", None),
            getattr(self, "ed_template", None),
            getattr(self, "btn_pick_c_file", None),
            getattr(self, "btn_pick_project", None),
            getattr(self, "btn_load_project", None),
            getattr(self, "btn_refresh_term_table", None),
            getattr(self, "btn_open_term_table", None),
            getattr(self, "btn_pick_output", None),
            getattr(self, "btn_pick_template", None),
            getattr(self, "rb_single", None),
            getattr(self, "rb_project", None),
            getattr(self, "ed_doc_old_code", None),
            getattr(self, "ed_doc_new_code", None),
            getattr(self, "ed_doc_old_doc", None),
            getattr(self, "ed_doc_review_decisions", None),
            getattr(self, "btn_doc_old_code", None),
            getattr(self, "btn_doc_new_code", None),
            getattr(self, "btn_doc_old_doc", None),
            getattr(self, "btn_doc_review_decisions", None),
            getattr(self, "cmb_doc_update_mode", None),
            getattr(self, "chk_doc_renumber_module_csu", None),
            getattr(self, "btn_start_doc_update", None),
        ):
            if w is None:
                continue
            try:
                w.setEnabled(not running)
            except Exception:
                pass
        if running:
            self._progress_state = {"total": None, "done": 0, "start_ts": None}
            self._progress.setFormat("")
            self._progress.setRange(0, 0)
        else:
            self._progress.setRange(0, 1)
            self._progress.setValue(0)
            self._progress.setFormat("")
            try:
                self._sync_source_mode()
            except Exception:
                pass

    def _stop_running(self) -> None:
        if self._worker is None:
            return
        try:
            self._worker.request_stop()
            self._append_log("已请求停止，等待当前步骤结束...")
        except Exception:
            pass

    def _start_generate(self, *, is_resume: bool) -> None:
        if self._thread is not None:
            self._pending_continue = is_resume
            return

        self._save_settings_from_ui_silent()
        self._save_recent_inputs()
        task = self._collect_task()
        if task.mode == "project" and not task.project_file_order:
            self._load_project_from_ui()
            task = self._collect_task()
        settings = self.settings
        resume_state = self._resume_state if is_resume else None

        steps = [
            StepDef("validate", "校验输入"),
            StepDef("config", "准备配置"),
            StepDef("generate", "生成文档"),
        ]
        self._reset_steps(steps)
        self._resume_state = None
        self._failed_functions = []  # 清空失败记录
        self._set_running(True, note="运行中...")
        self._last_task = task
        self._last_settings = settings

        self._run_in_thread(
            impl=GenerateWorker(backend=self.backend, task=task, settings=settings, resume_state=resume_state),
            kind="generate",
        )

    def _start_update_csu(self) -> None:
        if self._thread is not None:
            return

        self._save_settings_from_ui_silent()
        self._save_recent_inputs()
        task = self._collect_task()
        if (task.project_dir or "").strip() and not task.project_file_order:
            self._load_project_from_ui()
            task = self._collect_task()
        settings = self.settings

        steps = [
            StepDef("load_doc", "读取文档"),
            StepDef("update_ids", "更新标识"),
            StepDef("unit_table", "更新单元表"),
            StepDef("save_doc", "保存文档"),
        ]
        self._reset_steps(steps)
        self._set_running(True, note="更新 CSU 标识中...")

        self._run_in_thread(impl=UpdateCsuWorker(backend=self.backend, task=task, settings=settings), kind="update_csu")

    def _start_refresh_term_table(self) -> None:
        if self._thread is not None:
            return
        try:
            self.rb_project.setChecked(True)
        except Exception:
            pass
        self._save_settings_from_ui_silent()
        self._save_recent_inputs()
        task = self._collect_task()
        if not (task.project_dir or "").strip():
            QtWidgets.QMessageBox.warning(self, "提示", "请先选择工程目录。")
            return
        steps = [
            StepDef("scan", "扫描工程"),
            StepDef("prebuild", "预构建符号"),
            StepDef("functions", "解析函数"),
            StepDef("term_table", "生成术语总表"),
        ]
        self._reset_steps(steps)
        self._set_running(True, note="刷新术语表中...")
        self._panel_tabs.setCurrentWidget(self._step_tree)
        self._run_in_thread(
            impl=TermTableWorker(backend=self.backend, task=task, settings=self.settings),
            kind="term_table",
        )

    def _project_term_table_path(self) -> str:
        project_dir = (self.ed_project.text() or "").strip()
        if not project_dir:
            return ""
        try:
            return self.backend.default_term_table_path(project_dir)
        except Exception:
            return os.path.join(project_dir, "autodoc_term_table.json")

    def _show_project_term_table(self) -> None:
        path = self._project_term_table_path()
        if not path:
            QtWidgets.QMessageBox.warning(self, "提示", "请先选择工程目录。")
            return
        if not os.path.exists(path):
            QtWidgets.QMessageBox.information(self, "术语表", f"尚未生成术语表：\n{path}\n\n可点击“刷新术语表”。")
            return
        try:
            with open(path, "r", encoding="utf-8") as handle:
                data = json.load(handle)
            sections = ("functions", "symbols", "members", "macros")
            counts = {section: len((data.get(section) or {})) for section in sections}
            total = sum(counts.values())
            rows = []
            for section in sections:
                for ident, record in sorted((data.get(section) or {}).items(), key=lambda kv: str(kv[0]).lower()):
                    rec = record if isinstance(record, dict) else {"cn": str(record or "")}
                    rows.append(
                        {
                            "section": section,
                            "ident": str(ident or ""),
                            "cn": str(rec.get("cn") or ""),
                            "source": str(rec.get("source") or ""),
                            "confidence": float(rec.get("confidence", 0.0) or 0.0),
                            "locked": bool(rec.get("locked", False)),
                            "internal": bool(rec.get("internal", False)),
                            "scope": str(rec.get("scope") or ""),
                            "candidates": ", ".join(str((x or {}).get("cn") or "") for x in (rec.get("candidates") or []) if isinstance(x, dict) and str((x or {}).get("cn") or "").strip()),
                        }
                    )
            dlg = QtWidgets.QDialog(self)
            dlg.setWindowTitle("工程术语总表")
            dlg.resize(1080, 680)
            layout = QtWidgets.QVBoxLayout(dlg)
            summary = QtWidgets.QLabel(
                f"路径：{path}\n总计：{total} 项；函数 {counts['functions']}，变量 {counts['symbols']}，成员 {counts['members']}，宏 {counts['macros']}"
            )
            summary.setWordWrap(True)
            layout.addWidget(summary)
            filter_row = QtWidgets.QHBoxLayout()
            filter_kind = QtWidgets.QComboBox()
            filter_kind.addItems(["全部类型", "functions", "symbols", "members", "macros"])
            filter_conf = QtWidgets.QDoubleSpinBox()
            filter_conf.setRange(0.0, 1.0)
            filter_conf.setSingleStep(0.05)
            filter_conf.setDecimals(2)
            filter_conf.setValue(0.0)
            filter_conf.setToolTip("仅显示置信度不高于该值的术语；0 表示不按置信度过滤。")
            filter_text = QtWidgets.QLineEdit()
            filter_text.setPlaceholderText("搜索标识/中文名/来源/候选")
            filter_unlocked = QtWidgets.QCheckBox("仅未锁定")
            filter_show_internal = QtWidgets.QCheckBox("显示内部局部/参数")
            btn_low_conf = QtWidgets.QPushButton("低可信优先")
            btn_show_all = QtWidgets.QPushButton("显示全部")
            filter_count = QtWidgets.QLabel("")
            filter_row.addWidget(QtWidgets.QLabel("类型："))
            filter_row.addWidget(filter_kind)
            filter_row.addWidget(QtWidgets.QLabel("置信度≤"))
            filter_row.addWidget(filter_conf)
            filter_row.addWidget(filter_unlocked)
            filter_row.addWidget(filter_show_internal)
            filter_row.addWidget(filter_text, 1)
            filter_row.addWidget(btn_low_conf)
            filter_row.addWidget(btn_show_all)
            filter_row.addWidget(filter_count)
            layout.addLayout(filter_row)
            table = QtWidgets.QTableWidget(len(rows), 7)
            table.setHorizontalHeaderLabels(["类型", "标识", "中文名", "来源", "置信度", "锁定", "候选"])
            table.verticalHeader().setVisible(False)
            table.setSortingEnabled(False)
            table.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
            table.setEditTriggers(QtWidgets.QAbstractItemView.DoubleClicked | QtWidgets.QAbstractItemView.EditKeyPressed)
            table.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(2, QtWidgets.QHeaderView.Stretch)
            table.horizontalHeader().setSectionResizeMode(6, QtWidgets.QHeaderView.Stretch)
            for row_idx, row in enumerate(rows):
                table.setItem(row_idx, 0, QtWidgets.QTableWidgetItem(row["section"]))
                table.setItem(row_idx, 1, QtWidgets.QTableWidgetItem(row["ident"]))
                cn_item = QtWidgets.QTableWidgetItem(row["cn"])
                cn_item.setData(QtCore.Qt.UserRole + 10, row["cn"])
                table.setItem(row_idx, 2, cn_item)
                table.setItem(row_idx, 3, QtWidgets.QTableWidgetItem(row["source"]))
                conf_item = _NumericTableWidgetItem(f"{row['confidence']:.2f}")
                conf_item.setData(QtCore.Qt.UserRole, float(row["confidence"]))
                conf_item.setFlags(conf_item.flags() & ~QtCore.Qt.ItemIsEditable)
                table.setItem(row_idx, 4, conf_item)
                locked_item = QtWidgets.QTableWidgetItem("")
                locked_item.setFlags((locked_item.flags() | QtCore.Qt.ItemIsUserCheckable) & ~QtCore.Qt.ItemIsEditable)
                locked_item.setCheckState(QtCore.Qt.Checked if row["locked"] else QtCore.Qt.Unchecked)
                locked_item.setData(QtCore.Qt.UserRole + 10, bool(row["locked"]))
                table.setItem(row_idx, 5, locked_item)
                cand_item = QtWidgets.QTableWidgetItem(row["candidates"])
                cand_item.setData(QtCore.Qt.UserRole + 10, bool(row["internal"]))
                cand_item.setToolTip(f"scope={row['scope']}；internal={row['internal']}")
                cand_item.setFlags(cand_item.flags() & ~QtCore.Qt.ItemIsEditable)
                table.setItem(row_idx, 6, cand_item)
            table.setSortingEnabled(True)
            table.sortItems(4, QtCore.Qt.AscendingOrder)

            def _apply_term_filter() -> None:
                kind = str(filter_kind.currentText() or "")
                needle = str(filter_text.text() or "").strip().lower()
                max_conf = float(filter_conf.value())
                only_unlocked = bool(filter_unlocked.isChecked())
                show_internal = bool(filter_show_internal.isChecked())
                visible = 0
                for r in range(table.rowCount()):
                    section = (table.item(r, 0).text() if table.item(r, 0) else "")
                    conf = float((table.item(r, 4).data(QtCore.Qt.UserRole) if table.item(r, 4) else 0.0) or 0.0)
                    locked_now = bool(table.item(r, 5) is not None and table.item(r, 5).checkState() == QtCore.Qt.Checked)
                    internal_now = bool(table.item(r, 6) is not None and table.item(r, 6).data(QtCore.Qt.UserRole + 10))
                    row_text = " ".join((table.item(r, c).text() if table.item(r, c) else "") for c in (0, 1, 2, 3, 6)).lower()
                    hide = False
                    if internal_now and not show_internal:
                        hide = True
                    if kind != "全部类型" and section != kind:
                        hide = True
                    if max_conf > 0.0 and conf > max_conf:
                        hide = True
                    if needle and needle not in row_text:
                        hide = True
                    if only_unlocked and locked_now:
                        hide = True
                    table.setRowHidden(r, hide)
                    if not hide:
                        visible += 1
                filter_count.setText(f"显示 {visible}/{table.rowCount()}")

            filter_kind.currentTextChanged.connect(lambda _text: _apply_term_filter())
            filter_conf.valueChanged.connect(lambda _value: _apply_term_filter())
            filter_text.textChanged.connect(lambda _text: _apply_term_filter())
            filter_unlocked.stateChanged.connect(lambda _state: _apply_term_filter())
            filter_show_internal.stateChanged.connect(lambda _state: _apply_term_filter())
            btn_low_conf.clicked.connect(lambda: (filter_conf.setValue(0.80), filter_unlocked.setChecked(True), table.sortItems(4, QtCore.Qt.AscendingOrder), _apply_term_filter()))
            btn_show_all.clicked.connect(lambda: (filter_kind.setCurrentIndex(0), filter_conf.setValue(0.0), filter_unlocked.setChecked(False), filter_show_internal.setChecked(False), filter_text.clear(), _apply_term_filter()))
            _apply_term_filter()
            layout.addWidget(table, 1)
            btns = QtWidgets.QHBoxLayout()
            btn_save = QtWidgets.QPushButton("保存编辑")
            btn_export_locked = QtWidgets.QPushButton("导出锁定术语")
            btn_reload = QtWidgets.QPushButton("刷新术语表")
            btn_reload.clicked.connect(lambda: (dlg.accept(), self._start_refresh_term_table()))
            btn_open = QtWidgets.QPushButton("在系统中打开")
            btn_open.clicked.connect(lambda: QtGui.QDesktopServices.openUrl(QtCore.QUrl.fromLocalFile(path)))
            btn_close = QtWidgets.QPushButton("关闭")
            btn_close.clicked.connect(dlg.accept)
            btn_save.clicked.connect(lambda: self._save_term_table_from_widget(table, path))
            btn_export_locked.clicked.connect(lambda: self._export_locked_terms_from_widget(table, only_visible=False))
            btns.addWidget(btn_save)
            btns.addWidget(btn_export_locked)
            btns.addWidget(btn_reload)
            btns.addWidget(btn_open)
            btns.addStretch(1)
            btns.addWidget(btn_close)
            layout.addLayout(btns)
            dlg.exec_()
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"读取术语表失败：{e}")

    def _term_table_updates_from_widget(self, table: QtWidgets.QTableWidget) -> list[dict]:
        updates = []
        for row in range(table.rowCount()):
            section_item = table.item(row, 0)
            ident_item = table.item(row, 1)
            cn_item = table.item(row, 2)
            locked_item = table.item(row, 5)
            section = (section_item.text() if section_item else "").strip()
            ident = (ident_item.text() if ident_item else "").strip()
            cn = (cn_item.text() if cn_item else "").strip()
            if not section or not ident or not cn:
                continue
            original_cn = cn_item.data(QtCore.Qt.UserRole + 10) if cn_item is not None else None
            original_locked = locked_item.data(QtCore.Qt.UserRole + 10) if locked_item is not None else None
            current_locked = bool(locked_item is not None and locked_item.checkState() == QtCore.Qt.Checked)
            if original_cn is not None and str(original_cn) == cn and bool(original_locked) == current_locked:
                continue
            updates.append(
                {
                    "section": section,
                    "ident": ident,
                    "cn": cn,
                    "locked": current_locked,
                    "source": "manual_locked" if current_locked else "manual",
                    "confidence": 1.0 if current_locked else 0.98,
                }
            )
        return updates

    def _save_term_table_from_widget(self, table: QtWidgets.QTableWidget, path: str) -> None:
        project_dir = (self.ed_project.text() or "").strip()
        if not project_dir:
            QtWidgets.QMessageBox.warning(self, "提示", "请先选择工程目录。")
            return
        try:
            updates = self._term_table_updates_from_widget(table)
            self.backend.update_project_term_table_records(project_dir, updates, save=True)
            QtWidgets.QMessageBox.information(self, "已保存", f"已保存 {len(updates)} 条术语编辑：\n{path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"保存术语表失败：{e}")

    def _export_locked_terms_from_widget(self, table: QtWidgets.QTableWidget, *, only_visible: bool = False) -> None:
        locked = {"version": 2, "functions": {}, "symbols": {}, "members": {}, "macros": {}}
        for row in range(table.rowCount()):
            if only_visible and table.isRowHidden(row):
                continue
            locked_item = table.item(row, 5)
            if locked_item is None or locked_item.checkState() != QtCore.Qt.Checked:
                continue
            section = (table.item(row, 0).text() if table.item(row, 0) else "").strip()
            ident = (table.item(row, 1).text() if table.item(row, 1) else "").strip()
            cn = (table.item(row, 2).text() if table.item(row, 2) else "").strip()
            if section not in locked or not ident:
                continue
            locked[section][ident] = {
                "ident": ident,
                "cn": cn,
                "kind": section,
                "source": "manual_locked",
                "source_rank": 100,
                "confidence": 1.0,
                "locked": True,
            }
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "导出锁定术语", "locked_terms.json", "JSON (*.json)")
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as handle:
                json.dump(locked, handle, ensure_ascii=False, indent=2, sort_keys=True)
            QtWidgets.QMessageBox.information(self, "完成", f"已导出：{path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "失败", f"导出失败：{e}")

    def _run_in_thread(self, *, impl, kind: str) -> None:
        self._thread = QtCore.QThread(self)
        self._worker = _QtWorker(kind=kind, impl=impl)
        try:
            setattr(impl, "emit_log", self._enqueue_log_event)
            setattr(impl, "emit_detail", self._enqueue_detail_event)
        except Exception:
            pass
        self._worker.moveToThread(self._thread)
        self._worker.step.connect(self._on_step_event)
        self._worker.log.connect(self._on_log_event)
        self._worker.output.connect(self._on_output_event)
        self._worker.detail.connect(self._on_detail_event)
        self._worker.done.connect(self._thread.quit)
        self._worker.done.connect(self._on_done_event)
        self._thread.started.connect(self._worker.run)
        self._thread.finished.connect(self._worker.deleteLater)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.finished.connect(self._on_thread_finished)
        self._detail_timer = QtCore.QTimer(self)
        self._detail_timer.setInterval(50)
        self._detail_timer.timeout.connect(self._flush_detail_events)
        self._detail_timer.start()
        self._thread.start()

    def _request_thread_finish(self) -> None:
        thread = self._thread
        if thread is None:
            return
        try:
            thread.quit()
        except Exception:
            pass

    @QtCore.pyqtSlot()
    def _on_thread_finished(self) -> None:
        try:
            timer = getattr(self, "_detail_timer", None)
            if timer is not None:
                timer.stop()
                timer.deleteLater()
                self._detail_timer = None
        except Exception:
            pass
        self._flush_detail_events()
        self._thread = None
        self._worker = None
        if self._pending_continue is not None:
            is_resume = self._pending_continue
            self._pending_continue = None
            self._start_generate(is_resume=is_resume)

    def _cleanup_thread(self, wait_ms: int = 3000) -> bool:
        thread = self._thread
        if thread is None:
            return True
        try:
            if self._worker is not None:
                self._worker.request_stop()
            thread.quit()
            if not thread.wait(max(0, int(wait_ms))):
                return False
        except Exception:
            return False
        self._thread = None
        self._worker = None
        return True

    def closeEvent(self, event: QtGui.QCloseEvent) -> None:
        if self._thread is not None:
            ret = QtWidgets.QMessageBox.question(self, "确认退出", "任务运行中，是否先停止并退出？")
            if ret != QtWidgets.QMessageBox.Yes:
                event.ignore()
                return
            self._stop_running()
        if not self._cleanup_thread(wait_ms=5000):
            self._append_log("任务仍在收尾，请稍后再退出。")
            event.ignore()
            return
        self._closing = True
        self._save_window_state()
        event.accept()

    def _save_settings_from_ui_silent(self) -> None:
        try:
            if not getattr(self, "_settings_page_built", False):
                return
            extra_params = self._read_extra_params_table()
            exclude_dirs = [x.strip() for x in self.set_exclude_dirs.toPlainText().replace(",", "\n").splitlines() if x.strip()]
            mid_keywords = [x.strip() for x in (self.set_mid_keywords.text() or "").replace("，", ",").replace(";", ",").split(",") if x.strip()]
            drv_keywords = [x.strip() for x in (self.set_drv_keywords.text() or "").replace("，", ",").replace(";", ",").split(",") if x.strip()]
            domain_glossary_text = ""
            try:
                domain_glossary_text = str(self.set_domain_glossary.toPlainText() or "")
            except Exception:
                domain_glossary_text = ""
            symbol_dict_text = ""
            try:
                symbol_dict_text = str(self.set_symbol_dict.toPlainText() or "")
            except Exception:
                symbol_dict_text = ""
            s = AppSettings(
                section_prefix="5.1.1.",
                req_id_prefix=self.set_req_prefix.text().strip() or "D/R_SDD01_",
                only_with_comment=False,
                include_locals=True,
                include_logic=True,
                open_after_done=bool(self.set_open_after.isChecked()),
                ai_mode=normalize_ai_mode(self.set_ai_mode.currentIndex()),
                ai_provider=str(self.set_ai_provider.currentText() or "local"),
                ai_model=str(self.set_ai_model.currentText() or "").strip(),
                ai_api_base=self._normalize_ai_base_for_gui(self.set_ai_base.text()),
                ai_api_key=str(self.set_ai_key.text() or "").strip(),
                ai_num_ctx=int(self.set_ai_num_ctx.value()),
                ai_read_timeout=int(self.set_ai_read_timeout.value()),
                ai_workers=int(self.set_ai_workers.value()),
                force_ai=False,
                verbose=bool(self.set_verbose.isChecked()),
                no_proxy=bool(self.set_no_proxy.isChecked()),
                ai_one_call=False,
                auto_disable_large_one_call=True,
                ai_logic_format="json",
                ai_logic_policy="hybrid",
                logic_use_comment=(not bool(self.set_logic_ignore_comment.isChecked())),
                prefilter_project_files=bool(self.set_prefilter.isChecked()),
                incremental=bool(self.set_incremental.isChecked()),
                preprocess_workers=int(self.set_preprocess_workers.value()),
                log_every_n=int(self.set_log_every_n.value()),
                extra_params=extra_params,
                exclude_dirs=exclude_dirs,
                mid_dir_keywords=mid_keywords,
                drv_dir_keywords=drv_keywords,
                domain_glossary_text=domain_glossary_text,
                symbol_dict_text=symbol_dict_text,
            )
            self.settings_store.save(s)
            self.settings = s
        except Exception:
            pass

    def _open_settings_page(self) -> None:
        if not getattr(self, "_settings_page_built", False):
            idx = self._stack.indexOf(self._page_settings)
            self._stack.removeWidget(self._page_settings)
            try:
                self._page_settings.deleteLater()
            except Exception:
                pass
            self._page_settings = self._build_settings_page()
            self._stack.insertWidget(idx if idx >= 0 else 1, self._page_settings)
            self._settings_page_built = True
            self._apply_settings_to_ui()
            self._refresh_ai_models()
        self._stack.setCurrentWidget(self._page_settings)

    def _load_extra_params_table(self, data: dict[str, str]) -> None:
        self.tbl_extra_params.setRowCount(0)
        items = []
        for k, v in (data or {}).items():
            kk = str(k).strip()
            if not kk:
                continue
            items.append((kk, str(v)))
        items.sort(key=lambda kv: kv[0].lower())
        for k, v in items:
            r = self.tbl_extra_params.rowCount()
            self.tbl_extra_params.insertRow(r)
            self.tbl_extra_params.setItem(r, 0, QtWidgets.QTableWidgetItem(k))
            self.tbl_extra_params.setItem(r, 1, QtWidgets.QTableWidgetItem(v))

    def _read_extra_params_table(self) -> dict[str, str]:
        out: dict[str, str] = {}
        for r in range(self.tbl_extra_params.rowCount()):
            k_item = self.tbl_extra_params.item(r, 0)
            v_item = self.tbl_extra_params.item(r, 1)
            k = (k_item.text() if k_item else "").strip()
            v = (v_item.text() if v_item else "").strip()
            if not k:
                continue
            out[k] = v
        if "max_dist" not in out:
            out["max_dist"] = "2"
        if "min_ratio" not in out:
            out["min_ratio"] = "0.8"
        if "ai_retry_times" not in out:
            out["ai_retry_times"] = "0"
        if "ai_fail_policy" not in out:
            out["ai_fail_policy"] = "fallback"
        if "ai_regression_rounds" not in out:
            out["ai_regression_rounds"] = "1"
        if "ai_regression_force_one_call" not in out:
            out["ai_regression_force_one_call"] = "1"
        return out

    def _add_extra_param_row(self) -> None:
        r = self.tbl_extra_params.rowCount()
        self.tbl_extra_params.insertRow(r)
        self.tbl_extra_params.setItem(r, 0, QtWidgets.QTableWidgetItem(""))
        self.tbl_extra_params.setItem(r, 1, QtWidgets.QTableWidgetItem(""))
        self.tbl_extra_params.setCurrentCell(r, 0)

    def _delete_selected_extra_param_rows(self) -> None:
        sel = self.tbl_extra_params.selectionModel()
        if not sel:
            return
        rows = sorted({idx.row() for idx in sel.selectedRows()}, reverse=True)
        for r in rows:
            self.tbl_extra_params.removeRow(r)

    def _reset_extra_params_defaults(self) -> None:
        self._load_extra_params_table(
            {
                "max_dist": "2",
                "min_ratio": "0.8",
                "ai_profile": "small_model",
                "ai_retry_times": "0",
                "ai_fail_policy": "fallback",
                "ai_regression_rounds": "1",
                "structured_cond_ai": "0",
                "ai_regression_force_one_call": "0",
                "revision_profile": "",
            }
        )

    def _reset_ai_defaults(self) -> None:
        # Defaults are defined in AppSettings; use the same values here for UX.
        self.set_ai_mode.setCurrentIndex(0)
        self.set_ai_provider.setCurrentText("local")
        self._set_ai_model_items([], preferred="")
        self.set_ai_base.setText(LOCAL_LLM_API_BASE)
        self.set_ai_key.setText("")
        self.set_ai_num_ctx.setValue(0)
        self.set_ai_read_timeout.setValue(40)
        self.set_ai_workers.setValue(1)
        self.set_force_ai.setChecked(False)
        self.set_verbose.setChecked(True)
        self.set_no_proxy.setChecked(True)  # 默认禁用代理
        self.set_ai_one_call.setChecked(False)
        self.set_auto_disable_large_one_call.setChecked(True)
        self.set_ai_logic_format.setCurrentText("json")
        self.set_ai_logic_policy.setCurrentText("hybrid")
        try:
            self.set_logic_ignore_comment.setChecked(False)
        except Exception:
            pass
        self._on_ai_mode_changed(0)


class _QtWorker(QtCore.QObject):
    step = QtCore.pyqtSignal(str, str)
    log = QtCore.pyqtSignal(str)
    output = QtCore.pyqtSignal(str)
    detail = QtCore.pyqtSignal(object)
    done = QtCore.pyqtSignal(str, object, object)

    def __init__(self, *, kind: str, impl) -> None:
        super().__init__()
        self.kind = kind
        self.impl = impl

    def request_stop(self) -> None:
        try:
            self.impl.request_stop()
        except Exception:
            pass

    @QtCore.pyqtSlot()
    def run(self) -> None:
        if self.kind in ("generate", "export_func", "doc_update"):
            self.impl.run(
                emit_step=self.step.emit,
                emit_log=getattr(self.impl, "emit_log", self.log.emit),
                emit_output=self.output.emit,
                emit_done=self.done.emit,
                emit_detail=getattr(self.impl, "emit_detail", self.detail.emit),
            )
        else:
            self.impl.run(
                emit_step=self.step.emit,
                emit_log=getattr(self.impl, "emit_log", self.log.emit),
                emit_done=lambda note, out: self.done.emit(note, None, out),
            )


class ProjectTreeView(QtWidgets.QTreeView):
    """
    工程文件树：仅允许同一层级（应用/中间/驱动）内拖动改变顺序。
    """

    _ROLE_PATH = QtCore.Qt.UserRole + 10
    _ROLE_KIND = QtCore.Qt.UserRole + 11
    _LAYER_NAMES = ("应用层", "中间层", "驱动层")

    def __init__(self, parent: Optional[QtWidgets.QWidget] = None) -> None:
        super().__init__(parent)
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
        self.setDropIndicatorShown(True)
        self.setDragDropMode(QtWidgets.QAbstractItemView.InternalMove)
        self.setDefaultDropAction(QtCore.Qt.MoveAction)
        self.setSelectionMode(QtWidgets.QAbstractItemView.ExtendedSelection)

    removeRequested = QtCore.pyqtSignal()

    def _is_file_index(self, idx: QtCore.QModelIndex) -> bool:
        if not idx.isValid():
            return False
        model = idx.model()
        try:
            kind = model.data(idx, self._ROLE_KIND)
            path = model.data(idx, self._ROLE_PATH)
        except Exception:
            return False
        return bool(path) and (kind == "file")

    def _is_module_index(self, idx: QtCore.QModelIndex) -> bool:
        if not idx.isValid():
            return False
        model = idx.model()
        try:
            kind = model.data(idx, self._ROLE_KIND)
            path = model.data(idx, self._ROLE_PATH)
        except Exception:
            return False
        return (not path) and (kind == "module")

    def _is_layer_index(self, idx: QtCore.QModelIndex) -> bool:
        if not idx.isValid():
            return False
        if idx.parent().isValid():
            return False
        try:
            return str(idx.data() or "") in self._LAYER_NAMES
        except Exception:
            return False

    def _repair_orphan_top_level_files(self) -> None:
        """
        防御性修复：如果有"文件节点"被误拖成顶层（与 应用层/中间层/驱动层 同级），
        会导致后续无法拖动/无法右键删除；这里自动把它们塞回"应用层"。
        """
        model = self.model()
        if model is None or (not isinstance(model, QtGui.QStandardItemModel)):
            return
        root = model.invisibleRootItem()
        app_item = None
        for r in range(root.rowCount()):
            it = root.child(r, 0)
            if it is not None and it.text() == "应用层":
                app_item = it
                break
        if app_item is None:
            return

        for r in range(root.rowCount() - 1, -1, -1):
            it = root.child(r, 0)
            if it is None:
                continue
            if it.text() in self._LAYER_NAMES:
                continue
            # 顶层"file"节点：移动到"应用层"末尾
            if it.data(self._ROLE_KIND) == "file" and it.data(self._ROLE_PATH):
                moved = root.takeRow(r)
                if moved:
                    app_item.appendRow(moved)

    def keyPressEvent(self, event: QtGui.QKeyEvent) -> None:
        self._repair_orphan_top_level_files()
        if event.key() in (QtCore.Qt.Key_Delete, QtCore.Qt.Key_Backspace):
            idx = self.currentIndex()
            # 允许删除被误拖成顶层的文件节点（否则无法恢复）
            if idx.isValid() and (not idx.parent().isValid()) and self._is_file_index(idx):
                model = self.model()
                if isinstance(model, QtGui.QStandardItemModel):
                    model.removeRow(idx.row(), idx.parent())
                return
            self.removeRequested.emit()
            return
        super().keyPressEvent(event)

    def _layer_root_index(self, idx: QtCore.QModelIndex) -> Optional[QtCore.QModelIndex]:
        if not idx.isValid():
            return None
        cur = idx
        while cur.parent().isValid():
            cur = cur.parent()
        return cur

    def dragMoveEvent(self, event: QtGui.QDragMoveEvent) -> None:
        self._repair_orphan_top_level_files()
        src = self.currentIndex()
        if not src.isValid():
            event.ignore()
            return
        if (not src.parent().isValid()) and (not self._is_file_index(src)) and (not self._is_module_index(src)):
            # 不允许拖动层级节点本身
            event.ignore()
            return
        super().dragMoveEvent(event)

    def dropEvent(self, event: QtGui.QDropEvent) -> None:
        self._repair_orphan_top_level_files()

        model = self.model()
        if model is None or (not isinstance(model, QtGui.QStandardItemModel)):
            event.ignore()
            return

        selected = self.selectionModel().selectedRows(0) if self.selectionModel() else []
        selected = [i for i in selected if i.isValid() and (self._is_file_index(i) or self._is_module_index(i))]
        if not selected:
            event.ignore()
            return

        dest = self.indexAt(event.pos())

        # 仅允许两层结构：文件始终挂在某个层（应用/中间/驱动）下面，禁止拖成顶层/禁止拖成叶子子节点。
        # 这里不走 super().dropEvent，而是手动 takeRow/insertRow，避免 Qt 把文件插到顶层导致"无法再拖动"。
        root = model.invisibleRootItem()

        # 兜底：若拖到空白区域，则视为拖到"源层"末尾
        if not dest.isValid():
            dest_layer_idx = self._layer_root_index(selected[0])
            dest_layer_it = model.itemFromIndex(dest_layer_idx) if dest_layer_idx is not None else None
            if dest_layer_it is None:
                event.ignore()
                return
            target_parent = dest_layer_it
            target_row = target_parent.rowCount()
        else:
            # 确定目标层（顶层节点就是层；子节点则向上爬到层）
            dest_layer_idx = dest if (not dest.parent().isValid()) else self._layer_root_index(dest)
            dest_layer_it = model.itemFromIndex(dest_layer_idx) if dest_layer_idx is not None else None
            if dest_layer_it is None or dest_layer_it.text() not in self._LAYER_NAMES:
                event.ignore()
                return

            target_parent = dest_layer_it
            pos = self.dropIndicatorPosition()
            if not dest.parent().isValid():
                # 拖到层节点：Above=插到最前；Below/On=追加到末尾
                target_row = 0 if pos == QtWidgets.QAbstractItemView.AboveItem else target_parent.rowCount()
            else:
                # 拖到文件/模块节点：Above=插到该行前；Below/On=插到该行后
                base_row = dest.row()
                try:
                    if (dest.parent().isValid()
                            and model.data(dest.parent(), self._ROLE_KIND) == "module"
                            and (not self._is_module_index(dest))):
                        # 目标是模块内文件时，按模块在层中的行作为锚点
                        base_row = dest.parent().row()
                except Exception:
                    pass
                target_row = base_row if pos == QtWidgets.QAbstractItemView.AboveItem else (base_row + 1)

        # 收集待移动行（按其在各自父节点中的顺序稳定排序）
        # 注意：QStandardItem 在 PyQt 中不可 hash；这里用 id(parent_item) 做 key，并保留 item 引用。
        by_parent: dict[int, dict] = {}
        for idx in selected:
            parent_it = model.itemFromIndex(idx.parent())
            if parent_it is None:
                continue
            key = id(parent_it)
            slot = by_parent.get(key)
            if slot is None:
                slot = {"item": parent_it, "rows": []}
                by_parent[key] = slot
            slot["rows"].append(int(idx.row()))
        if not by_parent:
            event.ignore()
            return
        for k in list(by_parent.keys()):
            try:
                by_parent[k]["rows"] = sorted(set(by_parent[k].get("rows") or []))
            except Exception:
                by_parent[k]["rows"] = []

        # 目标行修正：若在同一父节点内移动，先减去"将被移走且位于目标行之前"的数量
        target_parent_key = id(target_parent)
        if target_parent_key in by_parent:
            rows = by_parent[target_parent_key].get("rows") or []
            shift = sum(1 for r in rows if int(r) < target_row)
            target_row = max(0, int(target_row) - shift)

        # 先 takeRow（从后往前），再按原顺序 insertRow
        moved_rows: list[list[QtGui.QStandardItem]] = []
        for slot in by_parent.values():
            parent_it = slot.get("item")
            rows = slot.get("rows") or []
            if parent_it is None:
                continue
            for r in sorted(rows, reverse=True):
                taken = parent_it.takeRow(int(r))
                if taken:
                    moved_rows.append(taken)
        if not moved_rows:
            event.ignore()
            return

        # 恢复原相对顺序（takeRow 是倒序取的）
        moved_rows.reverse()
        for row_items in moved_rows:
            target_parent.insertRow(target_row, row_items)
            target_row += 1

        try:
            self.expand(dest_layer_idx)
        except Exception:
            pass
        event.acceptProposedAction()
        return


class ExtraParamsValueDelegate(QtWidgets.QStyledItemDelegate):
    def createEditor(self, parent, option, index):  # type: ignore[override]
        if not index.isValid() or index.column() != 1:
            return super().createEditor(parent, option, index)
        model = index.model()
        key_idx = model.index(index.row(), 0)
        key = str(model.data(key_idx) or "").strip()
        if key == "ai_fail_policy":
            cb = QtWidgets.QComboBox(parent)
            cb.addItem("fallback（失败后继续纯规则）", "fallback")
            cb.addItem("skip_function（失败后跳过当前函数的 AI）", "skip_function")
            cb.addItem("circuit_fallback（失败后熔断后续 AI）", "circuit_fallback")
            return cb
        if key == "ai_regression_force_one_call":
            cb = QtWidgets.QComboBox(parent)
            cb.addItem("1（回归时强制 one_call）", "1")
            cb.addItem("0（沿用当前调用方式）", "0")
            return cb
        return super().createEditor(parent, option, index)

    def setEditorData(self, editor, index):  # type: ignore[override]
        if isinstance(editor, QtWidgets.QComboBox):
            cur = str(index.data() or "").strip()
            for i in range(editor.count()):
                if str(editor.itemData(i) or "") == cur:
                    editor.setCurrentIndex(i)
                    return
            editor.setCurrentIndex(0)
            return
        super().setEditorData(editor, index)

    def setModelData(self, editor, model, index):  # type: ignore[override]
        if isinstance(editor, QtWidgets.QComboBox):
            model.setData(index, str(editor.currentData() or "fallback"))
            return
        super().setModelData(editor, model, index)


def _load_backend_from_path(path: str):
    import importlib.util

    spec = importlib.util.spec_from_file_location("autodoc_backend", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"无法加载后端模块：{path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)  # type: ignore[attr-defined]
    return module


if __name__ == "__main__":
    import sys

    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
    if repo_root not in sys.path:
        sys.path.insert(0, repo_root)

    backend_path = os.path.join(repo_root, "AutoDocGen_V1.4.py")
    backend = _load_backend_from_path(backend_path)

    try:
        from qt_gui.app import run_qt_gui
    except Exception:
        from .app import run_qt_gui  # type: ignore

    raise SystemExit(run_qt_gui(backend=backend))
