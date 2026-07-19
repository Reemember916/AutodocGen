from __future__ import annotations

from docx import Document
import pytest

from autodoc.effects import EffectIndex, EffectSummary, extract_direct_effects, resolve_one_hop_effects
from autodoc.models import EffectFact, FunctionDesign, SourceRange
from autodoc.render import render_function_design
from autodoc.review_workspace import build_review_function
from autodoc.config import GenConfig
from autodoc.pipeline import build_function_design_impl
from tools.random_function_doccheck import Sample, check_docx


def _func(body: str, prototype: str = "void Demo(Uint16 *out, const Uint16 *input)"):
    params = prototype[prototype.find("(") + 1:prototype.rfind(")")]
    return {
        "func_info": {"func_name": "Demo", "prototype": prototype, "ret_type": prototype.split()[0], "params": params},
        "body": body,
        "file_context": {"source_file": "/tmp/demo.c"},
    }


def test_direct_effects_identify_pointer_member_array_and_global_writes():
    data = _func("""
    *out = 1U;
    out[1] = 2U;
    s_state.bit.ready = 1U;
    """)
    params = [{"name": "out", "type": "Uint16 *"}, {"name": "input", "type": "const Uint16 *"}]
    effects, returns = extract_direct_effects(data, params=params, local_vars=[], name_map={"s_state": "系统状态"})

    assert not returns
    assert [item.kind for item in effects] == ["param_write", "param_write", "global_write"]
    assert effects[-1].target_name == "系统状态"
    assert all(item.verified for item in effects)


def test_const_pointer_write_is_not_promoted_to_output_effect():
    data = _func("*input = 1U;")
    effects, _returns = extract_direct_effects(
        data,
        params=[{"name": "out", "type": "Uint16 *"}, {"name": "input", "type": "const Uint16 *"}],
        local_vars=[],
    )

    assert effects == ()


@pytest.mark.parametrize(
    ("lhs", "param", "local", "expected_kind"),
    [
        ("s_state_u16", None, None, "global_write"),
        ("s_state.bit.ready", None, None, "global_write"),
        ("s_state[1]", None, None, "global_write"),
        ("g_pack.word[0]", None, None, "global_write"),
        ("s_flag_u16", None, None, "global_write"),
        ("s_state->ready", None, None, "global_write"),
        ("*out", "Uint16 *", None, "param_write"),
        ("out[0]", "Uint16 *", None, "param_write"),
        ("out->ready", "State *", None, "param_write"),
        ("(*out).ready", "State *", None, "param_write"),
        ("**out", "Uint16 **", None, "param_write"),
        ("out[0].ready", "State *", None, "param_write"),
        ("*out", "const Uint16 *", None, None),
        ("out[0]", "const Uint16 *", None, None),
        ("l_temp_u16", None, "l_temp_u16", None),
        ("l_state.ready", None, "l_state", None),
        ("s_counter_u16", None, None, "global_write"),
        ("s_bits.bit.enabled", None, None, "global_write"),
        ("result[2]", "Uint16 result[]", None, "param_write"),
        ("result->field", "Result *", None, "param_write"),
        ("*result", "Result *", None, "param_write"),
        ("s_array[idx].field", None, None, "global_write"),
        ("s_value_u16", None, None, "global_write"),
        ("out.member", "State *", None, "param_write"),
    ],
)
def test_direct_effect_fixture_matrix(lhs, param, local, expected_kind):
    params = [{"name": "out" if "out" in lhs else "result", "type": param}] if param else []
    data = _func("", "void Demo(void)")
    pack = {"writes": [{"lhs": lhs, "rhs": "1U", "range": {"start_line": 1, "end_line": 1}, "verified": True, "confidence": 1.0}]}
    effects, _returns = extract_direct_effects(data, params=params, local_vars=([{"name": local, "type": "Uint16"}] if local else []), fact_pack=pack)

    assert (effects[0].kind if effects else None) == expected_kind


def test_return_effects_preserve_each_branch_and_condition():
    data = _func("""
    if (ok == VALID) { return STATUS_OK; }
    return STATUS_FAIL;
    """, "Uint16 Demo(Uint16 ok)")
    effects, returns = extract_direct_effects(data, params=[{"name": "ok", "type": "Uint16"}], local_vars=[], name_map={"STATUS_OK": "成功状态", "STATUS_FAIL": "失败状态"})

    assert not effects
    assert [item.target_ident for item in returns] == ["STATUS_OK", "STATUS_FAIL"]
    assert returns[0].condition == "ok == VALID"
    assert returns[0].target_name == "成功状态"


def test_one_hop_maps_callee_global_and_output_parameter_effects():
    inherited = (
        EffectFact(kind="global_write", target_ident="s_state_u16", target_name="状态", operation="写入", source_function="Helper", definition_source_file="/tmp/helper.c", definition_range=SourceRange(4, 4), verified=True),
        EffectFact(kind="param_write", target_ident="*out", target_name="输出", c_type="Uint16 *", operation="写入", source_function="Helper", definition_source_file="/tmp/helper.c", definition_range=SourceRange(5, 5), verified=True),
    )
    index = EffectIndex()
    index.add(EffectSummary("Helper", "/tmp/helper.c", ({"name": "out", "type": "Uint16 *"},), inherited))
    fact_pack = {"calls": [{"callee": "Helper", "call_text": "Helper(&s_result_u16)", "range": {"start_line": 9, "end_line": 9}}]}

    effects, issues = resolve_one_hop_effects(fact_pack, index=index, source_file="/tmp/caller.c", source_function="Caller", name_map={"s_result_u16": "调用结果"})

    assert not issues
    assert [item.target_ident for item in effects] == ["s_state_u16", "s_result_u16"]
    assert all(item.kind == "callee_effect" for item in effects)
    assert effects[-1].target_name == "调用结果"


def test_unresolved_and_recursive_calls_stay_as_audit_warnings():
    index = EffectIndex()
    fact_pack = {"calls": [
        {"callee": "External", "call_text": "External()", "range": {"start_line": 2, "end_line": 2}},
        {"callee": "Caller", "call_text": "Caller()", "range": {"start_line": 3, "end_line": 3}},
    ]}
    effects, issues = resolve_one_hop_effects(fact_pack, index=index, source_file="/tmp/caller.c", source_function="Caller")

    assert not effects
    assert [item["code"] for item in issues] == ["callee_effect_unresolved", "callee_effect_unresolved"]


def test_docx_and_review_workspace_render_effect_tables():
    effect = EffectFact(kind="global_write", target_ident="s_state_u16", target_name="系统状态", operation="写入", source_function="Demo", caller_source_file="/tmp/demo.c", caller_range=SourceRange(8, 8), verified=True)
    returned = EffectFact(kind="return", target_ident="STATUS_OK", target_name="成功状态", operation="返回", source_function="Demo", caller_source_file="/tmp/demo.c", caller_range=SourceRange(9, 9), verified=True, condition="有效时")
    design = FunctionDesign("测试", "D/R_001", "Uint16 Demo(void)", (), (), True, (), ("返回成功状态；",), effects=(effect,), return_effects=(returned,))
    doc = Document()
    render_function_design(doc, design, GenConfig())
    table_headers = [[cell.text for cell in table.rows[0].cells] for table in doc.tables]

    assert ["名称", "标识", "操作", "来源"] in table_headers
    assert ["返回表达式", "含义", "成立条件"] in table_headers
    review = build_review_function(design, {"func_info": {"func_name": "Demo"}, "source_file": "/tmp/demo.c"})
    assert {block.kind for block in review.blocks} >= {"effects_table", "return_semantics"}


def test_random_doccheck_requires_direct_effect_and_return_semantic_coverage(tmp_path):
    effect = EffectFact(kind="global_write", target_ident="s_state_u16", target_name="系统状态", operation="写入", source_function="Demo", verified=True)
    returned = EffectFact(kind="return", target_ident="STATUS_OK", target_name="成功状态", operation="返回", source_function="Demo", verified=True)
    design = FunctionDesign("测试", "D/R_001", "Uint16 Demo(void)", (), (), True, (), ("返回成功状态；",), effects=(effect,), return_effects=(returned,))
    path = tmp_path / "demo.docx"
    doc = Document()
    render_function_design(doc, design, GenConfig())
    doc.save(path)

    _score, _warnings, _size, _paragraphs, _tables, _excerpt, details = check_docx(
        path,
        Sample(c_file=str(tmp_path / "demo.c"), func_name="Demo", line_start=1, expected_external_effects=("s_state_u16",), expected_return_effects=("STATUS_OK",)),
    )
    assert not {item["code"] for item in details["quality_issues"]} & {"direct_effect_missing", "return_semantic_missing"}


def test_pipeline_attaches_effect_facts_without_changing_io_contract():
    data = _func("""
    *out = 1U;
    s_state_u16 = 2U;
    return;
    """)
    design = build_function_design_impl(data, "D/R_TEST", 1, GenConfig(ai_assist=False, effect_analysis_mode="direct"))

    assert any(item.kind == "param_write" for item in design.effects)
    assert any(item.kind == "global_write" for item in design.effects)
    assert any(item.ident == "out" and item.direction in {"输出", "输入/输出"} for item in design.io_elements)
