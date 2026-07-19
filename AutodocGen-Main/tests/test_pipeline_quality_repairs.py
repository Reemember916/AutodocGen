import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from autodoc.config import GenConfig
from autodoc import ai as ai_utils
from autodoc import backend
from autodoc import lsp_facts as lsp_fact_utils
from autodoc import logic as logic_utils
from autodoc import naming as naming_utils
from autodoc import revision as revision_utils
from autodoc.lsp_adapter import _collect_accesses as _collect_lsp_adapter_accesses
from autodoc.lsp_adapter import _scan_blocks as _scan_lsp_adapter_blocks
from autodoc.lsp_facts import _collect_accesses
from autodoc.lsp_facts import _collect_blocks
from autodoc.logic import _logic_cn_expr, _polish_logic_lines, _refresh_control_logic_line_idents, _render_humanized_compute_line, _render_structured_condition_cn, _sanitize_control_logic_line, generate_logic_from_semantic_pack
from autodoc.models import AIBuildMeta, FunctionDesign, SourceRange
from autodoc.parse import associate_comments_and_functions, extract_effective_comment_desc, parse_params_from_prototype, parse_single_comment_block
from autodoc.pipeline import (
    _fallback_structural_logic_lines,
    _logic_rendering_quality_issues,
    _meta_with_structural_quality,
    _lookup_io_display_name,
    compose_quality_feedback_text,
    prefer_regression_design,
    _repair_function_desc_by_domain,
    build_design_ai_meta,
    build_design_io_elements,
    build_design_name_map,
    build_design_text_sections,
    build_function_design_impl,
    build_logic_semantic_pack,
)
from autodoc.semantic_registry import classify_call_role, infer_local_semantic_label
from autodoc.render import add_module_function_table, render_function_design, render_table_or_none
from docx import Document
from docx.oxml.ns import qn
from tools.random_function_doccheck import Sample, check_docx, collect_samples


def _write_revision_profile(tmp_path, data):
    path = tmp_path / "revision_profile.json"
    import json

    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _base_ctx(*, ret_type="void", body="", local_vars=None, params=None, writes=None, symbol_map=None):
    symbol_map = symbol_map or {}
    return {
        "comment_info": {},
        "func_info": {
            "ret_type": ret_type,
            "func_name": "DemoFunc",
            "prototype": f"{ret_type} DemoFunc(void)",
        },
        "body": body,
        "params": list(params or []),
        "local_vars": list(local_vars or []),
        "in_map": {},
        "out_map": {},
        "param_ai_name_map": {},
        "var_cn_map": {},
        "global_symbol_map": dict(symbol_map),
        "file_context": {"symbol_map": dict(symbol_map), "glossary": {}},
        "family_prefix": "",
        "module_key": "",
        "owner_func": "DemoFunc",
        "source_file": "",
        "owner_ret_type": ret_type,
        "lsp_fact_pack": {"writes": list(writes or [])},
    }


def _text_sections_for_comment_and_body(raw_comment, *, func_name="DemoFunc", prototype="void DemoFunc(void)", body=""):
    comment_info = parse_single_comment_block(raw_comment)
    ctx = _base_ctx(body=body)
    ctx["comment_info"] = comment_info
    ctx["func_info"] = {
        "ret_type": prototype.split()[0],
        "func_name": func_name,
        "prototype": prototype,
    }
    return build_design_text_sections(ctx, "D/R_SDD01", 1, GenConfig(ai_assist=False))


def test_non_void_return_expression_is_output_element():
    ctx = _base_ctx(
        ret_type="Uint16",
        body="return s_ccdlTxRandData_u16[COMM_CCDL_SCI];",
        symbol_map={"s_ccdlTxRandData_u16": "基础帧历史兼容随机数字段发送值"},
    )

    io_elements, io_none, ret_var_name = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert not io_none
    assert ret_var_name is None
    assert len(io_elements) == 1
    assert io_elements[0].ident == "s_ccdlTxRandData_u16[COMM_CCDL_SCI]"
    assert io_elements[0].c_type == "Uint16"
    assert io_elements[0].direction == "输出"
    assert io_elements[0].name == "基础帧历史兼容随机数字段发送值"


def test_ternary_return_lists_possible_return_outputs():
    ctx = _base_ctx(
        ret_type="Uint16",
        body="return (VALID == l_found_u16) ? l_selectedID_u16 : l_preferredID_u16;",
        symbol_map={
            "VALID": "有效",
            "l_found_u16": "找到标志",
            "l_selectedID_u16": "最终通道",
            "l_preferredID_u16": "优先通道",
        },
    )
    ctx["comment_info"]["return_desc"] = "最终选中的RIU通道号"

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert not io_none
    assert [item.ident for item in io_elements] == ["l_selectedID_u16", "l_preferredID_u16"]
    assert [item.name for item in io_elements] == ["最终通道", "优先通道"]
    assert [item.c_type for item in io_elements] == ["Uint16", "Uint16"]
    assert all(item.direction == "输出" for item in io_elements)


def test_revision_profile_locks_names_and_function_text(tmp_path):
    source = tmp_path / "demo.c"
    source.write_text("void DemoFunc(void) { return; }\n", encoding="utf-8")
    profile_path = _write_revision_profile(
        tmp_path,
        {
            "version": 1,
            "functions": {
                "DemoFunc": {
                    "function_name": "人工锁定函数",
                    "description": "按人工审查结论输出状态。",
                    "locked_names": {"l_status_u16": "人工锁定状态"},
                }
            },
        },
    )
    func_data = {
        "comment_info": {},
        "func_info": {
            "func_name": "DemoFunc",
            "prototype": "void DemoFunc(void)",
            "ret_type": "void",
            "start": 1,
            "line_start": 1,
        },
        "body": """
        Uint16 l_status_u16 = INVALID;
        if (VALID == l_status_u16)
        {
            l_status_u16 = VALID;
        }
        return;
        """,
        "file_context": {
            "source_file": str(source),
            "symbol_map": {"VALID": "有效", "INVALID": "无效"},
            "glossary": {},
        },
    }
    cfg = GenConfig(ai_assist=False, extra_params={"revision_profile": str(profile_path)})

    design = build_function_design_impl(func_data, "D/R_TEST", 1, cfg)

    assert design.title == "人工锁定函数"
    assert "人工审查结论" in "\n".join(design.description_lines)
    assert any(item.ident == "l_status_u16" and item.name == "人工锁定状态" for item in (design.local_elements or ()))
    assert any("人工锁定状态" in line for line in (design.logic_lines or ()))


def test_revision_profile_logic_line_feedback_is_applied(tmp_path):
    profile_path = _write_revision_profile(
        tmp_path,
        {
            "version": 1,
            "functions": {
                "DemoFunc": {
                    "logic_replacements": [{"contains": "返回", "replace": "输出人工复核后的状态"}],
                    "logic_append": ["记录人工复核完成标志"],
                }
            },
        },
    )
    func_data = {
        "comment_info": {},
        "func_info": {
            "func_name": "DemoFunc",
            "prototype": "void DemoFunc(void)",
            "ret_type": "void",
            "start": 1,
            "line_start": 1,
        },
        "body": """
        if (VALID == l_flag_u16)
        {
            return;
        }
        """,
        "file_context": {
            "source_file": str(tmp_path / "demo.c"),
            "symbol_map": {"VALID": "有效", "l_flag_u16": "有效标志"},
            "glossary": {},
        },
    }
    cfg = GenConfig(ai_assist=False, extra_params={"revision_profile": str(profile_path)})

    design = build_function_design_impl(func_data, "D/R_TEST", 1, cfg)

    logic_text = "\n".join(design.logic_lines or ())
    assert "输出人工复核后的状态" in logic_text
    assert "记录人工复核完成标志" in logic_text


def test_revision_profile_prefers_file_specific_patch(tmp_path):
    source = tmp_path / "demo.c"
    source.write_text("void DemoFunc(void) {}\n", encoding="utf-8")
    profile = {
        "version": 1,
        "functions": {
            "DemoFunc": {"locked_names": {"l_status_u16": "通用状态"}},
            f"{source}::DemoFunc": {"locked_names": {"l_status_u16": "文件专用状态"}},
        },
    }

    patch = revision_utils.find_function_patch(profile, str(source), "DemoFunc")

    assert patch["locked_names"]["l_status_u16"] == "文件专用状态"


def test_revision_profile_golden_audit_reports_mismatch():
    issues = revision_utils.audit_golden_text(
        "函数说明中包含内部字段 logic_source_audit",
        [
            {
                "function": "DemoFunc",
                "must_contain": ["必须出现的逻辑"],
                "must_not_contain": ["logic_source_audit"],
            }
        ],
    )

    codes = {item["code"] for item in issues}
    assert "golden_must_contain_missing" in codes
    assert "golden_must_not_contain_hit" in codes


def test_ternary_return_macro_outputs_keep_status_suffix():
    ctx = _base_ctx(
        ret_type="Uint16",
        body="return (VALID == v_valid_u16) ? RIU_TX_SSM_VALID : RIU_TX_SSM_INVALID;",
        symbol_map={
            "VALID": "有效",
            "v_valid_u16": "有效数据",
            "RIU_TX_SSM_VALID": "RIU",
            "RIU_TX_SSM_INVALID": "RIU",
        },
    )

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert not io_none
    assert [item.ident for item in io_elements] == ["RIU_TX_SSM_VALID", "RIU_TX_SSM_INVALID"]
    assert [item.name for item in io_elements] == ["RIU发送SSM有效", "RIU发送SSM无效"]


def test_multiple_simple_returns_list_possible_return_outputs():
    ctx = _base_ctx(
        ret_type="unsigned int",
        body="""
        if (tdata != ARITHMETIC_TEST_VALUE)
        {
            return CPUTEST_ARITH_OP_ERR;
        }
        else
        {
            return CPUTEST_ARITH_OP_OK;
        }
        """,
        symbol_map={
            "CPUTEST_ARITH_OP_ERR": "CPU测试",
            "CPUTEST_ARITH_OP_OK": "CPU测试",
        },
    )
    ctx["comment_info"]["return_desc"] = "测试通过 ==> CPUTEST_ARITH_OP_OK\n测试未通过 ==> CPUTEST_ARITH_OP_ERR"

    io_elements, io_none, ret_var_name = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert not io_none
    assert ret_var_name is None
    assert [item.ident for item in io_elements] == ["CPUTEST_ARITH_OP_ERR", "CPUTEST_ARITH_OP_OK"]
    assert [item.name for item in io_elements] == ["CPU测试未通过", "CPU测试通过"]
    assert [item.c_type for item in io_elements] == ["unsigned int", "unsigned int"]
    assert all(item.direction == "输出" for item in io_elements)


def test_global_writes_are_not_function_io_elements():
    ctx = _base_ctx(
        body="""
        l_tmp_u16 = 1U;
        s_commCCDL422TxFlag_u16 = ON;
        s_ccdlSciTxActiveLen_u16 = COMM_CCDL_TX_FRAM_LEN;
        """,
        local_vars=[{"name": "l_tmp_u16", "type": "Uint16"}],
        writes=[
            {"lhs": "l_tmp_u16", "rhs": "1U"},
            {"lhs": "s_commCCDL422TxFlag_u16", "rhs": "ON"},
            {"lhs": "s_ccdlSciTxActiveLen_u16", "rhs": "COMM_CCDL_TX_FRAM_LEN"},
        ],
        symbol_map={
            "s_commCCDL422TxFlag_u16": "SCI链路CCDL发送忙闲标志",
            "s_ccdlSciTxActiveLen_u16": "当前正在发送帧的有效长度",
        },
    )

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert io_none
    assert io_elements == ()


def test_function_params_and_return_are_io_elements_without_global_write_expansion():
    ctx = _base_ctx(
        ret_type="Uint16",
        body="""
        *vp_valid_u16 = VALID;
        s_globalState_u16 = VALID;
        return l_result_u16;
        """,
        params=[
            {"name": "v_mode_u16", "type": "Uint16"},
            {"name": "vp_valid_u16", "type": "Uint16 *"},
        ],
        local_vars=[{"name": "l_result_u16", "type": "Uint16", "usage": "处理结果"}],
        writes=[
            {"lhs": "*vp_valid_u16", "rhs": "VALID"},
            {"lhs": "s_globalState_u16", "rhs": "VALID"},
        ],
        symbol_map={
            "s_globalState_u16": "全局状态",
            "l_result_u16": "处理结果",
        },
    )
    ctx["in_map"] = {"v_mode_u16": "工作模式", "vp_valid_u16": "有效标志指针"}
    ctx["out_map"] = {"vp_valid_u16": "有效标志指针"}

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert not io_none
    assert [item.ident for item in io_elements] == ["v_mode_u16", "vp_valid_u16", "l_result_u16"]
    assert [item.direction for item in io_elements] == ["输入", "输入/输出", "输出"]
    assert "s_globalState_u16" not in {item.ident for item in io_elements}


def test_pointer_param_write_marks_function_output_without_comment():
    ctx = _base_ctx(
        body="*vp_status_u16 = VALID;",
        params=[{"name": "vp_status_u16", "type": "Uint16 *"}],
        writes=[{"lhs": "*vp_status_u16", "rhs": "VALID"}],
    )

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert not io_none
    assert len(io_elements) == 1
    assert io_elements[0].ident == "vp_status_u16"
    assert io_elements[0].direction == "输出"


def test_global_member_writes_do_not_expand_function_io_table():
    ctx = _base_ctx(
        body="""
        s_RIUSendData_t.currState_u16 = RECEIVE_RIU_STATE_ACTIVE;
        s_synWhle_t[l_synStyleID_u16].costTimHdSk_u32[HANDSK_L_ID] = l_hadskTim_u32;
        """,
        writes=[
            {"lhs": "s_RIUSendData_t.currState_u16", "rhs": "RECEIVE_RIU_STATE_ACTIVE"},
            {"lhs": "s_synWhle_t[l_synStyleID_u16].costTimHdSk_u32[HANDSK_L_ID]", "rhs": "l_hadskTim_u32"},
        ],
    )

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert io_none
    assert io_elements == ()


def test_global_member_writes_from_typedef_do_not_expand_function_io_table():
    ctx = _base_ctx(
        body="s_demo_t.mode = DEMO_MODE_ACTIVE;",
        writes=[{"lhs": "s_demo_t.mode", "rhs": "DEMO_MODE_ACTIVE"}],
    )
    ctx["file_context"]["typedefs"] = [
        """
        typedef struct
        {
            Uint16 mode;
            float target;
        } Demo_t;
        """
    ]

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert io_none
    assert io_elements == ()


def test_static_void_return_type_does_not_create_return_output():
    ctx = _base_ctx(ret_type="static void", body="return;")

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert io_none
    assert io_elements == ()


def test_static_non_void_return_type_strips_storage_qualifier():
    ctx = _base_ctx(ret_type="static Uint16", body="return VALID;")

    io_elements, io_none, _ = build_design_io_elements(ctx, {"file_context": {}}, GenConfig(ai_assist=False))

    assert not io_none
    assert io_elements[0].ident == "VALID"
    assert io_elements[0].c_type == "Uint16"


def test_refresh_logic_does_not_replace_ascii_inside_chinese_label():
    line = "IF CCDL标识 小于 通道间通信数量 且 kzzz标识 小于 控制装置通信数量 时"
    name_map = {"CCDL": "bit2(CCDL)", "kzzz": "KZZZ"}

    refreshed = _refresh_control_logic_line_idents(line, name_map)

    assert "bit2" not in refreshed
    assert refreshed.startswith("IF CCDL标识 小于")


def test_refresh_logic_rejects_corrupt_bitfield_alias_for_plain_macro():
    line = "IF CCDL标识 小于 CCDL22_NUM时"
    name_map = {
        "CCDL22_NUM": "燃油系统指令1数据位域2燃油系统指令1数据位域2bit2",
    }

    refreshed = _refresh_control_logic_line_idents(line, name_map)

    assert "数据位域" not in refreshed
    assert "bit2" not in refreshed
    assert "CCDL22数量" in refreshed


def test_control_line_sanitizer_drops_ascii_hint_before_identifier_refresh():
    cleaned = _sanitize_control_logic_line(
        "FOR 遍历 循环计数(ii) 小于 10U",
        name_map={"ii": "循环索引(ii)"},
    )

    assert cleaned == "FOR 遍历 循环计数 小于 10U"
    assert _sanitize_control_logic_line("IF 读取数据 项 等于 闪存DEVICE标识时") == "IF 读取数据项 等于 闪存DEVICE标识时"
    assert _polish_logic_lines(["上电BIT自检数据按位或并更新构造第循环计数 位标志；"])[0] == "上电BIT自检数据按位或并更新构造第循环计数位标志；"


def test_semantic_logic_keeps_return_after_closed_if_block():
    pack = {
        "control_blocks": [
            {
                "id": "b1",
                "kind": "if",
                "condition": "v_id_u16 < MAX_NUM",
                "range": {"start_line": 10, "end_line": 12},
                "metadata": {"brace_depth_before": 1},
            }
        ],
        "state_updates": [
            {
                "kind": "control_compute",
                "lhs": "l_rslt_t",
                "rhs": "s_status_t[v_id_u16]",
                "range": {"start_line": 11, "end_line": 11},
            }
        ],
        "return_actions": [
            {"expr": "l_rslt_t", "range": {"start_line": 14, "end_line": 14}},
        ],
    }
    name_map = {
        "v_id_u16": "通道标识",
        "MAX_NUM": "通道数量",
        "l_rslt_t": "返回结果",
        "s_status_t": "状态表",
    }

    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False), name_map=name_map)
    lines = [line.strip() for line in logic.splitlines() if line.strip()]

    assert any(line.startswith("暂存") and "返回结果" in line for line in lines)
    assert lines.index("END IF") < lines.index("返回 处理结果")
    assert not any(line.startswith("返回 状态表") for line in lines)


def test_semantic_logic_does_not_duplicate_return_result_label_with_detail():
    pack = {
        "control_blocks": [],
        "state_updates": [],
        "return_actions": [
            {"expr": "lo_rData_f", "range": {"start_line": 1, "end_line": 1}},
        ],
    }
    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"lo_rData_f": "返回结果(数据)"},
    )

    assert "返回 处理结果" in logic
    assert "返回 返回结果" not in logic


def test_member_access_alias_prefers_known_member_name():
    ctx = _base_ctx(
        body="l_rslt_t.dataState_u16 = s_info_t[v_id_u16].rxState_u16;",
        symbol_map={"rxState_u16": "状态字"},
    )
    ctx["lsp_fact_pack"] = {
        "members": [
            {
                "access_text": "s_info_t[v_id_u16].rxState_u16",
                "member": "rxState_u16",
                "owner_type": "",
            }
        ],
        "writes": [],
        "blocks": [],
        "calls": [],
    }

    pack = build_logic_semantic_pack(ctx)

    assert pack["entity_aliases"]["s_info_t[v_id_u16].rxState_u16"] == "状态字"


def test_access_facts_skip_declaration_initializers():
    _reads, writes = _collect_accesses(
        """
        Uint16 l_buff_u16 = INSTRUCTION_WRITE_ENABLE;
        s_flag_u16 = ON;
        """
    )

    assert [item.lhs for item in writes] == ["s_flag_u16"]


def test_access_facts_skip_for_loop_initializers():
    _reads, writes = _collect_accesses(
        """
        for(l_ii_u16 = 0U; l_ii_u16 < v_len_u16; l_ii_u16++)
        {
            s_flag_u16 = ON;
        }
        """
    )

    assert [item.lhs for item in writes] == ["s_flag_u16"]


def test_access_facts_keep_compound_assignment_lhs_clean():
    _reads, writes = _collect_accesses("s_puBITData_u16 |= (Uint16)(0x01U << l_index_u16);")

    assert [(item.lhs, item.rhs) for item in writes] == [("s_puBITData_u16", "(Uint16)(0x01U << l_index_u16)")]
    assert writes[0].metadata["op"] == "|="


def test_access_facts_keep_inline_braced_if_assignments():
    _reads, writes = _collect_accesses(
        "if (l_valveData2_un32.bit.LPQD_state_u32 == VALVE_STATE_OPEN) { l_podValveOpen_u16 = VALID; }"
    )

    assert [(item.lhs, item.rhs) for item in writes] == [("l_podValveOpen_u16", "VALID")]


def test_lsp_adapter_access_facts_skip_declaration_initializers():
    source = """
void f(void)
{
    Uint16 l_buff_u16 = INSTRUCTION_WRITE_ENABLE;
    s_flag_u16 = ON;
}
"""
    _reads, writes = _collect_lsp_adapter_accesses(source, SourceRange(start_line=1, end_line=6))

    assert [item.lhs for item in writes] == ["s_flag_u16"]


def test_lsp_adapter_access_facts_skip_for_loop_initializers():
    source = """
void f(void)
{
    for(l_ii_u16 = 0U; l_ii_u16 < v_len_u16; l_ii_u16++)
    {
        s_flag_u16 = ON;
    }
}
"""
    _reads, writes = _collect_lsp_adapter_accesses(source, SourceRange(start_line=1, end_line=8))

    assert [item.lhs for item in writes] == ["s_flag_u16"]


def test_lsp_adapter_access_facts_keep_compound_assignment_lhs_clean():
    source = """
void f(void)
{
    s_puBITData_u16 |= (Uint16)(0x01U << l_index_u16);
}
"""
    _reads, writes = _collect_lsp_adapter_accesses(source, SourceRange(start_line=1, end_line=5))

    assert [(item.lhs, item.rhs) for item in writes] == [("s_puBITData_u16", "(Uint16)(0x01U << l_index_u16)")]
    assert writes[0].metadata["op"] == "|="


def test_lsp_adapter_access_facts_keep_inline_braced_if_assignments():
    source = """
void f(void)
{
    if (l_valveData2_un32.bit.LPQD_state_u32 == VALVE_STATE_OPEN) { l_podValveOpen_u16 = VALID; }
}
"""
    _reads, writes = _collect_lsp_adapter_accesses(source, SourceRange(start_line=1, end_line=5))

    assert [(item.lhs, item.rhs) for item in writes] == [("l_podValveOpen_u16", "VALID")]


def test_access_facts_strip_inline_case_labels():
    body = "case RIU429_MODE_LP:      l_newMode_u16 = WORK_MODE_LP_HELI; break;"
    _reads, writes = _collect_accesses(body)

    assert [(item.lhs, item.rhs) for item in writes] == [("l_newMode_u16", "WORK_MODE_LP_HELI")]


def test_lsp_adapter_blocks_strip_inline_case_action_from_condition():
    source = """
void f(void)
{
    switch (mode)
    {
        case RIU429_MODE_LP:      l_newMode_u16 = WORK_MODE_LP_HELI; break;
        default:                  l_newMode_u16 = WORK_MODE_STANDBY; break;
    }
}
"""
    blocks = _scan_lsp_adapter_blocks(source, SourceRange(start_line=1, end_line=9))
    case_blocks = [item for item in blocks if item.kind == "case"]

    assert case_blocks
    assert case_blocks[0].condition == "RIU429_MODE_LP"
    assert "=" not in case_blocks[0].condition


def test_fallback_blocks_strip_inline_case_action_from_condition():
    body = """
    switch (mode)
    {
        case RIU429_MODE_LP:      l_newMode_u16 = WORK_MODE_LP_HELI; break;
        default:                  l_newMode_u16 = WORK_MODE_STANDBY; break;
    }
"""
    blocks = _collect_blocks(body)
    case_blocks = [item for item in blocks if item.kind == "case"]

    assert case_blocks
    assert case_blocks[0].condition == "RIU429_MODE_LP"


def test_humanized_compute_does_not_treat_chinese_label_parentheses_as_call():
    text = _render_humanized_compute_line("新模式", "受油模式(ECIEVE)")

    assert text == "计算 新模式 = 受油模式(ECIEVE)"


def test_caption_numbers_increment_without_word_field_update():
    doc = Document()
    render_table_or_none(doc, "输入/输出元素", ["名称"], [["输入"]])
    render_table_or_none(doc, "局部数据元素", ["名称"], [["局部"]])

    captions = [p.text for p in doc.paragraphs if p.text.startswith("表 ")]

    assert captions == ["表 1 输入/输出元素", "表 2 局部数据元素"]


def test_rendered_table_rows_are_not_split_across_pages():
    doc = Document()
    table = render_table_or_none(
        doc,
        "输入/输出元素",
        ["名称", "标识"],
        [["标志位", "s_RIUSendData_t.ValveCtrl_t.bit.Pump0_Lcutoff_ctrl_u16"]],
    )

    assert table is not None
    for row in table.rows:
        tr_pr = row._tr.trPr
        assert tr_pr is not None
        assert tr_pr.find(qn("w:cantSplit")) is not None


def test_module_function_table_rows_are_not_split_across_pages():
    doc = Document()
    table = add_module_function_table(
        doc,
        "加油阶段预设处理",
        "D/R_SDD01",
        [{"csu_name": "加油模式发送开阀预位", "csu_id": "D/R_SDD01_001"}],
    )

    assert table is not None
    for row in table.rows:
        tr_pr = row._tr.trPr
        assert tr_pr is not None
        assert tr_pr.find(qn("w:cantSplit")) is not None


def test_parse_params_preserves_pointer_type():
    params = parse_params_from_prototype(
        {"params": "Uint16 *v_dBuff_u16, const Uint8 *vp_payload_u8, Uint8 v_len_u8"}
    )

    assert params == [
        {"name": "v_dBuff_u16", "type": "Uint16*"},
        {"name": "vp_payload_u8", "type": "const Uint8*"},
        {"name": "v_len_u8", "type": "Uint8"},
    ]


def test_fullwidth_function_comment_labels_are_parsed():
    parsed = parse_single_comment_block(
        """
         * 【函数名】:RefuelStagePreset
         *
         * 【功能描述】加油预位控制
         *             根据当前模式和供油目标打开相应吊舱阀、泵切断阀，并在超时内确认开到位。
         * 【输入参数说明】v_p_ConData_t：系统控制数据指针
         * 【输出参数说明】无
         * 【其他说明】       预位超时按阀位不到位处理
         * 【返回】          无
        """
    )

    assert parsed["func_name"] == "RefuelStagePreset"
    assert parsed["desc"].splitlines()[0] == "加油预位控制"
    assert "根据当前模式和供油目标" in parsed["desc"]
    assert parsed["input_desc"] == "v_p_ConData_t：系统控制数据指针"
    assert parsed["other_desc"] == "预位超时按阀位不到位处理"


def test_fullwidth_comment_without_colon_and_repeated_input_labels_are_parsed():
    parsed = parse_single_comment_block(
        """
        *【函数名】FrameSyn
        *
        *【功能描述】通道同步
        *
        *【输入参数说明】l_synStyleID_u16 - 同步类型
        *【输入参数说明】 l_frmSynTim_u32 -- 同步时间
        *【其他说明】高低握手，得到同步结果
        """
    )

    assert parsed["func_name"] == "FrameSyn"
    assert parsed["desc"] == "通道同步"
    assert "l_synStyleID_u16 - 同步类型" in parsed["input_desc"]
    assert "l_frmSynTim_u32 -- 同步时间" in parsed["input_desc"]
    assert parsed["desc"] != "【函数名】FrameSyn"



def test_comment_normalizer_extracts_project_same_line_fullwidth_desc():
    from autodoc.comment_normalizer import normalize_comment_block
    raw = """
    /**
     * 【函数名】:FdataAverage
     *
     * 【功能描述】浮点数求平均
     * \t 1、对一组浮点数，去除最大、最小值后，求平均值，当浮点数个数为零时，返回零；
     *   2、当浮点数个数大于零，小于三时，返回数组第一个数。
     *
     * 【输入参数说明】v_pBuff_f ---- 浮点数数组指针
     * \t\t\t   v_len_16  ---- 数据长度
     * 【输出参数说明】NONE
     * 【其他说明】       NONE
     * 【返回】:\t数组中浮点数的平均值
     */
    """

    normalized = normalize_comment_block(raw)

    assert normalized.func_name == "FdataAverage"
    assert "浮点数求平均" in normalized.desc
    assert "去除最大、最小值" in normalized.desc
    assert "小于三时，返回数组第一个数" in normalized.desc
    assert "v_pBuff_f" in normalized.input_desc
    assert "v_len_16" in normalized.input_desc
    assert normalized.output_desc == "NONE"
    assert normalized.return_desc == "数组中浮点数的平均值"


def test_comment_normalizer_extracts_project_bracket_desc_following_line():
    from autodoc.comment_normalizer import normalize_comment_block
    raw = """
    /**
     *    [函数名]\t Comm422FrameCheck
     *
     *    [功能描述]
     *    \t\t\t  检测对应通道接收缓冲区是否存在有效报文。
     *    [输入参数说明] v_commID_u16  ---- RS422通道ID，可能取值为:
     *              \tSCI_A_ID ---- SCIA接口
     *              \tSCI_B_ID ---- SCIB接口
     *              \tSCI_C_ID ---- SCIC接口
     *\t  [输出参数说明] NONE
     *    [其他说明]\t    NONE
     *    [返回]\t\t 返回有效报文首数据在缓冲区中索引，无有效报文时，返回:RS422_COMM_FRAM_NOT_EXIST
     */
    """

    normalized = normalize_comment_block(raw)

    assert normalized.func_name == "Comm422FrameCheck"
    assert normalized.desc == "检测对应通道接收缓冲区是否存在有效报文。"
    assert "v_commID_u16" in normalized.input_desc
    assert "SCI_A_ID" in normalized.input_desc
    assert normalized.output_desc == "NONE"
    assert "有效报文首数据" in normalized.return_desc
    assert "RS422_COMM_FRAM_NOT_EXIST" in normalized.return_desc



def test_comment_normalizer_accepts_plain_label_only_section_headers():
    from autodoc.comment_normalizer import normalize_comment_block
    raw = """
    /**
     * 函数名: PlainLabelOnly
     * 功能描述
     * 第一行功能。
     * 输入参数说明
     * arg ---- 参数说明
     * 输出参数说明
     * NONE
     * 其他说明
     * 无
     * 返回
     * 返回说明
     */
    """

    normalized = normalize_comment_block(raw)

    assert normalized.func_name == "PlainLabelOnly"
    assert normalized.desc == "第一行功能。"
    assert normalized.input_desc == "arg ---- 参数说明"
    assert normalized.output_desc == "NONE"
    assert normalized.other_desc == "无"
    assert normalized.return_desc == "返回说明"


def test_comment_normalizer_keeps_unknown_bracketed_lines_inside_section():
    from autodoc.comment_normalizer import normalize_comment_block
    raw = """
    /**
     * 函数名: UnknownBracketLine
     * [功能描述]
     * 第一行功能。
     * 【注意】保持状态
     * [提示]继续保留
     * 返回
     * NONE
     */
    """

    normalized = normalize_comment_block(raw)

    assert normalized.desc == "第一行功能。\n【注意】保持状态\n[提示]继续保留"
    assert normalized.return_desc == "NONE"


def test_comment_normalizer_preserves_explanation_placeholder_as_free_text():
    from autodoc.comment_normalizer import normalize_comment_block
    raw = """
    /**
     * 【说明】:SciTxRxEnable
     *
     * 该函数实现SCI端口接收、发送的使能和关闭功能。
     *
     * 【参数】:sciID SCI端口号
     */
    """

    normalized = normalize_comment_block(raw)

    assert normalized.desc == "【说明】:SciTxRxEnable"

def test_parse_single_comment_block_preserves_existing_dict_contract_with_normalized_desc():
    raw = """
    /**
     * 【函数名】:TimeCountInit
     *
     * 【功能描述】时间计数初始化
     * 【输入参数说明】NONE
     * 【输出参数说明】NONE
     * 【其他说明】       同步时间初始化为定时器1微秒值加上任务主周期时间，用于实现进入while周期后立刻执行一次同步！！！
     * 【返回】               NONE
     */
    """

    parsed = parse_single_comment_block(raw)

    assert set(parsed) == {
        "func_name",
        "func_cn_name",
        "desc",
        "input_desc",
        "output_desc",
        "other_desc",
        "return_desc",
    }
    assert parsed["func_name"] == "TimeCountInit"
    assert parsed["desc"] == "时间计数初始化"
    assert parsed["input_desc"] == "NONE"
    assert parsed["output_desc"] == "NONE"
    assert "同步时间初始化" in parsed["other_desc"]
    assert parsed["return_desc"] == "NONE"


def test_build_design_text_sections_uses_fdataaverage_comment_description():
    raw = """
    /**
     * 【函数名】:FdataAverage
     *
     * 【功能描述】浮点数求平均
     * \t 1、对一组浮点数，去除最大、最小值后，求平均值，当浮点数个数为零时，返回零；
     *   2、当浮点数个数大于零，小于三时，返回数组第一个数。
     *
     * 【输入参数说明】v_pBuff_f ---- 浮点数数组指针
     * \t\t\t   v_len_16  ---- 数据长度
     * 【输出参数说明】NONE
     * 【返回】:\t数组中浮点数的平均值
     */
    """

    sections = _text_sections_for_comment_and_body(
        raw,
        func_name="FdataAverage",
        prototype="float FdataAverage(float *v_pBuff_f, Uint16 v_len_16)",
    )

    desc = "\n".join(sections["description_lines"])
    assert "浮点数求平均" in desc
    assert "去除最大、最小值" in desc
    assert "小于三时" in desc
    assert desc != "无。"


def test_build_design_text_sections_uses_comm422_comment_description():
    raw = """
    /**
     *    [函数名]\t Comm422FrameCheck
     *
     *    [功能描述]
     *    \t\t\t  检测对应通道接收缓冲区是否存在有效报文。
     *    [输入参数说明] v_commID_u16  ---- RS422通道ID，可能取值为:
     *              \tSCI_A_ID ---- SCIA接口
     *              \tSCI_B_ID ---- SCIB接口
     *              \tSCI_C_ID ---- SCIC接口
     *\t  [输出参数说明] NONE
     *    [返回]\t\t 返回有效报文首数据在缓冲区中索引，无有效报文时，返回:RS422_COMM_FRAM_NOT_EXIST
     */
    """

    sections = _text_sections_for_comment_and_body(
        raw,
        func_name="Comm422FrameCheck",
        prototype="Uint16 Comm422FrameCheck(Uint16 v_commID_u16)",
    )

    desc = "\n".join(sections["description_lines"])
    assert desc == "检测对应通道接收缓冲区是否存在有效报文。"
    assert desc != "无。"


def test_function_comment_association_skips_decorative_separator():
    code = """
/**
 * 【函数名】:ChTypeRoundRobinCommitColdStartup
 *
 * 【功能描述】冷启动主备稳定建立后，提交下一次冷启动默认主通道轮值
 *             本次主控的对端作为下一次冷启动默认主通道
 * 【输入参数说明】无
 * 【输出参数说明】无
 * 【其他说明】       仅在冷启动板间资格成功后调用
 * 【返回】          无
 */
/* ***************************************************************** */
void ChTypeRoundRobinCommitColdStartup(void)
{
}
"""

    funcs = associate_comments_and_functions(code)
    comment_info = funcs[0]["comment_info"]

    assert funcs[0]["func_info"]["func_name"] == "ChTypeRoundRobinCommitColdStartup"
    assert comment_info["func_name"] == "ChTypeRoundRobinCommitColdStartup"
    assert "提交下一次冷启动默认主通道轮值" in comment_info["desc"]
    assert comment_info["other_desc"] == "仅在冷启动板间资格成功后调用"



def test_preceding_function_comment_keeps_project_fdataaverage_doc_past_inline_comment_and_separator():
    code = """
void PreviousHelper(void)
{
    /* 本地状态更新，不是下一个函数的说明 */
}

/**
 * 【函数名】:FdataAverage
 *
 * 【功能描述】浮点数求平均
 * \t 1、对一组浮点数，去除最大、最小值后，求平均值，当浮点数个数为零时，返回零；
 *   2、当浮点数个数大于零，小于三时，返回数组第一个数。
 *
 * 【输入参数说明】v_pBuff_f ---- 浮点数数组指针
 * \t\t\t   v_len_16  ---- 数据长度
 * 【输出参数说明】NONE
 * 【其他说明】       NONE
 * 【返回】:\t数组中浮点数的平均值
 */
/* ***************************************************************** */
float FdataAverage(float *v_pBuff_f, Uint16 v_len_16)
{
    return 0.0;
}
"""

    funcs = associate_comments_and_functions(code)
    previous = next(func for func in funcs if func["func_info"]["func_name"] == "PreviousHelper")
    item = next(func for func in funcs if func["func_info"]["func_name"] == "FdataAverage")

    assert previous["comment_info"] == {}
    assert item["comment_info"].get("func_name") == "FdataAverage"
    assert "浮点数求平均" in item["comment_info"].get("desc", "")
    assert "数组中浮点数的平均值" in item["comment_info"].get("return_desc", "")


def test_preceding_function_comment_keeps_project_comm422_doc_past_inline_comment_and_separator():
    code = """
void PreviousCommHelper(void)
{
    // 接收缓存巡检，不是下一个函数的说明
}

/**
 *    [函数名]\t Comm422FrameCheck
 *
 *    [功能描述]
 *    \t\t\t  检测对应通道接收缓冲区是否存在有效报文。
 *    [输入参数说明] v_commID_u16  ---- RS422通道ID，可能取值为:
 *              \tSCI_A_ID ---- SCIA接口
 *              \tSCI_B_ID ---- SCIB接口
 *              \tSCI_C_ID ---- SCIC接口
 *\t  [输出参数说明] NONE
 *    [其他说明]\t    NONE
 *    [返回]\t\t 返回有效报文首数据在缓冲区中索引，无有效报文时，返回:RS422_COMM_FRAM_NOT_EXIST
 */
/* ***************************************************************** */
Uint16 Comm422FrameCheck(Uint16 v_commID_u16)
{
    return 0;
}
"""

    funcs = associate_comments_and_functions(code)
    previous = next(func for func in funcs if func["func_info"]["func_name"] == "PreviousCommHelper")
    item = next(func for func in funcs if func["func_info"]["func_name"] == "Comm422FrameCheck")

    assert previous["comment_info"] == {}
    assert item["comment_info"].get("func_name") == "Comm422FrameCheck"
    assert "检测对应通道接收缓冲区是否存在有效报文" in item["comment_info"].get("desc", "")
    assert "RS422_COMM_FRAM_NOT_EXIST" in item["comment_info"].get("return_desc", "")


def test_preceding_function_comment_keeps_valid_desc_when_identifier_fields_look_non_semantic():
    code = """
void PreviousHelper(void)
{
    /* 上一个函数体内注释，不应关联到下一个函数 */
}

/**
 * 【函数名】:$Revision: 1.0 $
 *
 * 【功能描述】浮点数求平均
 * 【输入参数说明】$Revision: 1.0 $ ---- 版本标识，不是功能说明
 * 【输出参数说明】NONE
 * 【其他说明】NONE
 * 【返回】: 数组中浮点数的平均值
 */
/* ***************************************************************** */
float FdataAverage(float *v_pBuff_f, Uint16 v_len_16)
{
    return 0.0;
}
"""

    funcs = associate_comments_and_functions(code)
    item = next(func for func in funcs if func["func_info"]["func_name"] == "FdataAverage")

    assert "浮点数求平均" in item["comment_info"].get("desc", "")


def test_preceding_function_comment_rejects_placeholder_param_blocks_without_desc_after_separator():
    for marker in ("TODO", "FIXME", "TBD", "XXX"):
        code = f"""
/** 【输入参数说明】{marker}: fill in params */
/* ***************************************************************** */
void DemoFunc(void)
{{
}}
"""

        funcs = associate_comments_and_functions(code)
        item = next(func for func in funcs if func["func_info"]["func_name"] == "DemoFunc")

        assert item["comment_info"] == {}, marker


def test_comment_normalizer_association_keeps_project_doc_block_with_param_separators():
    code = """
/**
 * 【函数名】:FdataAverage
 *
 * 【功能描述】浮点数求平均
 * 【输入参数说明】v_pBuff_f ---- 浮点数数组指针
 *                 v_len_16  ---- 数据长度
 * 【输出参数说明】NONE
 * 【返回】: 数组中浮点数的平均值
 */
/* ***************************************************************** */
float FdataAverage(float *v_pBuff_f, Uint16 v_len_16)
{
    return 0.0;
}
"""

    funcs = associate_comments_and_functions(code)
    comment_info = funcs[0]["comment_info"]

    assert funcs[0]["func_info"]["func_name"] == "FdataAverage"
    assert comment_info["func_name"] == "FdataAverage"
    assert "浮点数求平均" in comment_info["desc"]
    assert "v_pBuff_f ---- 浮点数数组指针" in comment_info["input_desc"]
    assert "v_len_16  ---- 数据长度" in comment_info["input_desc"]


def test_function_comment_association_keeps_decorative_bordered_doc_block():
    code = """
/***********************************************************************************************************************
 *
 *【函数名】FrameSyn
 *
 *【功能描述】通道同步
 *
 *【输入参数说明】l_synStyleID_u16 - 同步类型
 *【输入参数说明】 l_frmSynTim_u32 -- 同步时间
 *
 *【其他说明】高低握手，得到同步结果
 *
 *【返回】 NONE
 **************************************************************************************************************************/
void FrameSyn(Uint16 l_synStyleID_u16,Uint32 l_frmSynTim_u32)
{
}
"""

    funcs = associate_comments_and_functions(code)
    comment_info = funcs[0]["comment_info"]

    assert funcs[0]["func_info"]["func_name"] == "FrameSyn"
    assert comment_info["func_name"] == "FrameSyn"
    assert comment_info["desc"] == "通道同步"
    assert "l_frmSynTim_u32 -- 同步时间" in comment_info["input_desc"]


def test_function_title_prefers_compact_first_comment_desc_line():
    assert backend.get_function_chinese_name(
        {"desc": "加油预位控制\n根据当前模式和供油目标打开相应吊舱阀"},
        {"func_name": "RefuelStagePreset"},
    ) == "加油预位控制"
    assert backend.get_function_chinese_name(
        {"desc": "通道同步"},
        {"func_name": "FrameSyn"},
    ) == "通道同步"


def test_round_robin_cold_start_title_is_not_numeric_rounding():
    title = backend.get_function_chinese_name(
        {"desc": "冷启动主备稳定建立后，提交下一次冷启动默认主通道轮值\n本次主控的对端作为下一次冷启动默认主通道"},
        {"func_name": "ChTypeRoundRobinCommitColdStartup"},
    )

    assert title == "冷启动主通道轮值提交"
    assert "四舍五入" not in title


def test_symbol_memory_does_not_override_compact_comment_title():
    old_symbols = dict(backend.SYMBOL_DICTIONARY_RUNTIME)
    try:
        backend.SYMBOL_DICTIONARY_RUNTIME["RefuelStagePreset"] = "模式预设写入"
        comment_info = {"func_cn_name": "", "desc": "加油预位控制\n根据当前模式打开相应阀门"}

        backend._seed_symbol_memory_into_scope(
            comment_info,
            {"func_name": "RefuelStagePreset"},
            [],
            [],
            {},
            {},
        )

        assert comment_info["func_cn_name"] == ""
        assert backend.get_function_chinese_name(
            comment_info,
            {"func_name": "RefuelStagePreset"},
        ) == "加油预位控制"
    finally:
        backend.SYMBOL_DICTIONARY_RUNTIME.clear()
        backend.SYMBOL_DICTIONARY_RUNTIME.update(old_symbols)


def test_header_member_comments_parse_typedef_struct_on_next_line():
    member_map = backend._extract_member_symbol_map_from_header_code(
        """
        typedef struct
        {
            Uint16 targetTank_u16;           /* 目标供油油箱选择 (0:未决/1:0号/2:2_3号/3:LRP) */
            Uint16 commandSent_u16;          /* 预位指令已发送标志 */
            Uint16 presetReady_u16;          /* 燃油预位完成标志 */
        } RefuelModeContext_t;
        """
    )

    assert member_map["targetTank_u16"] == "目标供油油箱选择"
    assert member_map["commandSent_u16"] == "预位指令已发送标志"
    assert member_map["presetReady_u16"] == "燃油预位完成标志"
    assert member_map["RefuelModeContext_t.presetReady_u16"] == "燃油预位完成标志"


def test_header_member_comments_keep_owner_qualified_duplicate_fields():
    member_map = backend._extract_member_symbol_map_from_header_code(
        """
        typedef struct
        {
            Uint16 presetReady_u16;          /* 预设准备完成标志 */
        } ReceiveModeContext_t;

        typedef struct
        {
            Uint16 presetReady_u16;          /* 燃油预位完成标志 */
        } RefuelModeContext_t;
        """
    )

    assert member_map["ReceiveModeContext_t.presetReady_u16"] == "预设准备完成标志"
    assert member_map["RefuelModeContext_t.presetReady_u16"] == "燃油预位完成标志"


def test_io_display_name_prefers_member_comment_over_generic_global_alias():
    ctx = {
        "global_symbol_map": {"s_refuelCtx_t.targetTank_u16": "缓存值"},
        "file_context": {"member_symbol_map": {"targetTank_u16": "目标供油油箱选择"}},
    }

    assert _lookup_io_display_name(
        ctx,
        "s_refuelCtx_t.targetTank_u16",
        fallback="全局状态",
    ) == "目标供油油箱选择"


def test_io_display_name_prefers_instance_member_comment_for_duplicate_field():
    ctx = {
        "global_symbol_map": {"s_refuelCtx_t.presetReady_u16": "缓存值"},
        "file_context": {
            "member_symbol_map": {
                "presetReady_u16": "预设准备完成标志",
                "s_refuelCtx_t.presetReady_u16": "燃油预位完成标志",
            }
        },
    }

    assert _lookup_io_display_name(
        ctx,
        "s_refuelCtx_t.presetReady_u16",
        fallback="全局状态",
    ) == "燃油预位完成标志"


def test_io_display_name_uses_owner_for_union_container_members():
    ctx = {
        "global_symbol_map": {
            "all": "燃油系统指令1数据",
            "bit": "燃油系统指令1数据位域",
        },
        "file_context": {
            "member_symbol_map": {
                "s_sysConData_t.commDataSourse_un16": "通信数据来源",
                "s_sysConData_t.CHVIn_un16": "通道有效输入信号",
                "CHVIn_un16.myCHV_u16": "bit0:本端运行期授权使用的本地CHV有效位",
            }
        },
    }

    comm = _lookup_io_display_name(ctx, "s_sysConData_t.commDataSourse_un16.all", fallback="全局状态")
    chv = _lookup_io_display_name(ctx, "s_sysConData_t.CHVIn_un16.bit.myCHV_u16", fallback="全局状态")

    assert comm == "通信数据来源"
    assert chv == "本端运行期授权使用的本地CHV有效位"
    assert "燃油系统指令1数据" not in comm + chv


def test_logic_uses_full_member_alias_before_generic_owner_alias():
    logic, _ = backend.generate_logic_from_body(
        """
        if (INVALID == s_refuelCtx_t.presetReady_u16)
        {
            s_refuelCtx_t.presetReady_u16 = VALID;
        }
        """,
        [],
        GenConfig(ai_assist=False),
        name_map={
            "s_refuelCtx_t": "缓存值",
            "presetReady_u16": "预设准备完成标志",
            "s_refuelCtx_t.presetReady_u16": "燃油预位完成标志",
            "INVALID": "无效",
            "VALID": "有效",
        },
    )

    assert "燃油预位完成标志" in logic
    assert "缓存值的预设准备完成标志" not in logic


def test_control_role_macro_comments_are_kept_and_used_in_logic():
    symbol_map = backend._extract_symbol_map_from_header_code(
        """
        #define ROLE_BACKUP             (0U)          /* 当前不持有控制权 */
        #define ROLE_MASTER             (1U)          /* 当前持有控制权 */
        """
    )
    assert symbol_map["ROLE_MASTER"] == "当前持有控制权"
    assert symbol_map["ROLE_BACKUP"] == "当前不持有控制权"

    logic, _ = backend.generate_logic_from_body(
        """
        if (ROLE_MASTER == s_sysConData_t.runtimeRole_u16)
        {
            ;
        }
        """,
        [],
        GenConfig(ai_assist=False),
        name_map={
            **symbol_map,
            "s_sysConData_t.runtimeRole_u16": "动态控制权归属",
        },
    )

    assert "当前持有控制权" in logic
    assert "ROLE_MASTER" not in logic


def test_assignment_logic_translates_common_refuel_macros_without_ai():
    logic, _ = backend.generate_logic_from_body(
        """
        s_refuelCtx_t.commandSent_u16 = INVALID;
        s_refuelCtx_t.targetTank_u16 = REFUEL_TARGET_LRP_ALL;
        s_RIUSendData_t.ValveCtrl_t.bit.LPQD_ctrl_u16 = REFUEL_VALVE_CMD_OPEN;
        """,
        [],
        GenConfig(ai_assist=False),
        name_map={
            "s_refuelCtx_t.commandSent_u16": "预位指令已发送标志",
            "s_refuelCtx_t.targetTank_u16": "目标供油箱选择",
            "s_RIUSendData_t.ValveCtrl_t.bit.LPQD_ctrl_u16": "左吊舱切断阀控制指令",
        },
    )

    compact_logic = logic.replace(" ", "")
    assert "置预位指令已发送标志为无效" in compact_logic
    assert "将左右吊舱全部目标写入目标供油箱选择" in compact_logic
    assert "将开阀指令写入左吊舱切断阀控制指令" in compact_logic
    assert "INVALID" not in logic
    assert "REFUEL_TARGET_LRP_ALL" not in logic
    assert "REFUEL_VALVE_CMD_OPEN" not in logic


def test_array_write_target_does_not_inherit_index_macro_alias():
    ctx = _base_ctx(
        body="s_puBITInfo_u16[PUBIT_INDEX_FLASH] = PUBIT_TEST_ERR;",
        symbol_map={
            "s_puBITInfo_u16": "上电BIT检测项信息",
            "PUBIT_TEST_ERR": "上电自检结果异常",
        },
    )
    ctx["lsp_fact_pack"] = {
        "writes": [
            {
                "lhs": "s_puBITInfo_u16[PUBIT_INDEX_FLASH]",
                "rhs": "PUBIT_TEST_ERR",
                "range": {"start_line": 1, "end_line": 1},
                "source": "references",
                "confidence": 0.84,
                "verified": True,
            }
        ],
        "blocks": [],
        "calls": [],
        "members": [],
    }

    pack = build_logic_semantic_pack(ctx)
    assert pack["entity_aliases"].get("s_puBITInfo_u16[PUBIT_INDEX_FLASH]") != "循环索引"

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={
            "s_puBITInfo_u16": "上电BIT检测项信息",
            "PUBIT_TEST_ERR": "上电自检结果异常",
        },
    )

    assert "循环索引" not in logic
    assert "上电BIT检测项信息" in logic
    assert "上电自检结果异常" in logic


def test_semantic_logic_renders_compound_assignment_operator():
    ctx = _base_ctx(
        body="s_puBITData_u16 |= (Uint16)(0x01U << l_index_u16);",
        symbol_map={
            "s_puBITData_u16": "上电BIT自检数据",
            "l_index_u16": "循环计数",
        },
    )
    ctx["lsp_fact_pack"] = {
        "writes": [
            {
                "lhs": "s_puBITData_u16",
                "rhs": "(Uint16)(0x01U << l_index_u16)",
                "metadata": {"op": "|="},
                "range": {"start_line": 1, "end_line": 1},
                "source": "references",
                "confidence": 0.84,
                "verified": True,
            }
        ],
        "blocks": [],
        "calls": [],
        "members": [],
    }

    pack = build_logic_semantic_pack(ctx)
    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))

    assert "上电BIT自检数据按位或并更新构造第循环计数位标志" in logic
    assert "上电BIT自检数据 |" not in logic


def test_design_name_map_keeps_source_comment_over_buffer_alias():
    ctx = _base_ctx(
        body="Uint16 l_readBuff_u16[3] = {0U, 0U, 0U};",
        local_vars=[
            {
                "name": "l_readBuff_u16",
                "type": "Uint16",
                "cn_name": "读取数据",
                "comment_cn_name": "读取数据",
            }
        ],
    )
    ctx["lsp_fact_pack"] = {
        "locals": [
            {
                "name": "l_readBuff_u16",
                "decl_type": "Uint16[3]",
                "source": "documentSymbol",
                "confidence": 0.74,
                "verified": False,
            }
        ],
        "writes": [],
        "blocks": [],
        "calls": [],
        "members": [],
    }
    ctx["logic_semantic_pack"] = build_logic_semantic_pack(ctx)

    name_map = build_design_name_map(ctx)

    assert name_map["l_readBuff_u16"] == "读取数据"


def test_local_declaration_comment_cn_name_beats_dataflow_profile():
    body = """
    float l_tank0Vol_f = 0.0F; /* 0号箱油量，用于决定单吊舱预位路径。 */
    l_tank0Vol_f = RedunDataGet(TANK0).fData_f;
    """
    local_vars = backend.parse_local_variables_from_body(body)
    item = next(v for v in local_vars if v["name"] == "l_tank0Vol_f")

    backend._repair_local_cn_name_with_profile(
        item,
        body=body,
        neighbor_symbols=[],
        comment_desc="加油预位控制",
        cfg=GenConfig(ai_assist=False),
    )

    assert item["cn_name"] == "0号箱油量"
    assert backend._select_local_display_name(item) == "0号箱油量"


def test_parse_local_variables_splits_multi_declarator_comment_labels():
    body = """
    float  l_min_f = 0.0,l_max_f = 0.0; /* 最小值，最大值 */
    double l_sum_f = 0.0; /* 数据和值 */
    """

    locals_ = backend.parse_local_variables_from_body(body)
    by_name = {item["name"]: item for item in locals_}

    assert by_name["l_min_f"]["cn_name"] == "最小值"
    assert by_name["l_min_f"].get("name_source") == "inline_comment_split"
    assert by_name["l_max_f"]["cn_name"] == "最大值"
    assert by_name["l_max_f"].get("name_source") == "inline_comment_split"
    assert by_name["l_sum_f"]["cn_name"] in {"数据和值", "累加和", "数据和"}


def test_parse_local_variables_preserves_pointer_attached_type_declarations():
    body = """
    Uint8* l_buf_pu8 = NULL; /* 缓冲区 */
    const Foo* state; /* 状态 */
    char* l_min_ptr = NULL;
    """

    locals_ = backend.parse_local_variables_from_body(body)
    by_name = {item["name"]: item for item in locals_}

    assert by_name["l_buf_pu8"]["type"] == "Uint8 *"
    assert by_name["l_buf_pu8"]["comment_hint"] == "缓冲区"
    assert by_name["state"]["type"] == "const Foo *"
    assert by_name["state"]["comment_hint"] == "状态"
    assert by_name["l_min_ptr"]["type"] == "char *"
    assert by_name["l_min_ptr"]["cn_name"] == "最小值"


def test_parse_local_variables_keeps_custom_typedef_pointer_locals():
    body = """
    FooState* state; /* 状态指针 */
    const FooConfig* config; /* 配置指针 */
    total*scale;
    """

    locals_ = backend.parse_local_variables_from_body(body)
    by_name = {item["name"]: item for item in locals_}

    assert by_name["state"]["type"] in {"FooState*", "FooState *"}
    assert by_name["state"]["cn_name"] or by_name["state"].get("comment_hint") == "状态指针"
    assert "config" in by_name
    assert "scale" not in by_name


def test_parse_local_variables_ignores_uppercase_multiplication_expressions():
    body = """
    LIMIT*scale;
    Foo*bar;
    FooState* state; /* 状态指针 */
    const FooConfig* config; /* 配置指针 */
    """

    locals_ = backend.parse_local_variables_from_body(body)
    by_name = {item["name"]: item for item in locals_}

    assert "scale" not in by_name
    assert "bar" not in by_name
    assert by_name["state"]["type"] in {"FooState*", "FooState *"}
    assert by_name["state"].get("comment_hint") == "状态指针"
    assert by_name["config"]["type"] in {"const FooConfig*", "const FooConfig *"}


def test_parse_local_variables_ignores_commented_multiplication_expression():
    body = """
    Foo*bar; /* scale */
    Foo*bar; /* scale variable */
    """

    locals_ = backend.parse_local_variables_from_body(body)
    names = {item["name"] for item in locals_}

    assert "bar" not in names


def test_parse_local_variables_ignores_adjacent_multiplication_expression():
    body = """
    Uint16 total = 1U;
    total*scale;
    float l_tank0Vol_f = 1.0;
    l_tank0Vol_f*scale;
    Uint8* l_buf_pu8;
    """

    locals_ = backend.parse_local_variables_from_body(body)
    names = {item["name"] for item in locals_}

    assert "total" in names
    assert "l_tank0Vol_f" in names
    assert "l_buf_pu8" in names
    assert "scale" not in names


def test_parse_local_variables_falls_back_to_identifier_semantics():
    body = """
    Uint16 l_ii_u16 = 0U;
    Uint16 l_jj_u16 = 0U;
    Uint16 l_count_u16 = 0U;
    double l_sum_f = 0.0;
    """

    locals_ = backend.parse_local_variables_from_body(body)
    by_name = {item["name"]: item for item in locals_}

    ii_label = by_name["l_ii_u16"]["cn_name"]
    jj_label = by_name["l_jj_u16"]["cn_name"]

    assert ii_label in {"循环索引ii", "索引ii"}
    assert jj_label in {"循环索引jj", "索引jj"}
    assert ii_label != jj_label
    assert "ii" in ii_label
    assert "jj" in jj_label
    assert by_name["l_count_u16"]["cn_name"] == "计数"
    assert by_name["l_sum_f"]["cn_name"] in {"累加和", "数据和", "数据和值"}



def test_final_quality_preserves_loop_index_labels_through_design_path():
    body = """
    Uint16 l_ii_u16 = 0U; /* 循环索引ii */
    Uint16 l_jj_u16 = 0U; /* 循环索引jj */
    Uint16 l_count_u16 = 2U; /* 循环次数 */
    for (l_ii_u16 = 0U; l_ii_u16 < l_count_u16; l_ii_u16++)
    {
        for (l_jj_u16 = 0U; l_jj_u16 < l_count_u16; l_jj_u16++)
        {
            l_count_u16 = l_count_u16;
        }
    }
    """
    func_data = {
        "comment_info": {"func_cn_name": "循环索引检查", "desc": "循环索引检查"},
        "func_info": {
            "func_name": "LoopIndexLabelCheck",
            "prototype": "void LoopIndexLabelCheck(void)",
            "ret_type": "void",
            "start": 1,
            "line_start": 1,
        },
        "body": body,
        "file_context": {"source_file": "", "symbol_map": {}, "glossary": {}},
    }

    design = build_function_design_impl(func_data, "D/R_SDD01", 1, GenConfig(ai_assist=False))
    locals_by_ident = {item.ident: item for item in (design.local_elements or ())}

    assert locals_by_ident["l_ii_u16"].name != locals_by_ident["l_jj_u16"].name
    assert "ii" in locals_by_ident["l_ii_u16"].name
    assert "jj" in locals_by_ident["l_jj_u16"].name



def test_parse_local_variables_identifier_fallback_avoids_substring_matches():
    body = """
    Uint8 l_summary_u8 = 0U;
    Uint16 l_discount_u16 = 0U;
    Uint8 indexedMode = 0U;
    """

    locals_ = backend.parse_local_variables_from_body(body)
    by_name = {item["name"]: item for item in locals_}

    assert by_name["l_summary_u8"]["cn_name"] != "累加和"
    assert by_name["l_discount_u16"]["cn_name"] != "计数"
    assert by_name["indexedMode"]["cn_name"] != "循环索引"


def test_fdataaverage_design_locals_bind_min_and_max_correctly():
    body = """
    Uint16 l_ii_u16 = 0U; /* 循环索引 */
    float  l_min_f = 0.0,l_max_f = 0.0; /* 最小值，最大值 */
    double l_sum_f = 0.0; /* 数据和值 */
    float  l_fData_f = 0.0;  /* 平均值 */
    return l_fData_f;
    """
    func_data = {
        "comment_info": {"func_cn_name": "浮点数求平均", "desc": "浮点数求平均"},
        "func_info": {
            "func_name": "FdataAverage",
            "prototype": "float FdataAverage(float *v_pBuff_f, Uint16 v_len_16)",
            "ret_type": "float",
            "start": 1,
            "line_start": 1,
        },
        "body": body,
        "file_context": {"source_file": "", "symbol_map": {}, "glossary": {}},
    }

    design = build_function_design_impl(func_data, "D/R_SDD01", 1, GenConfig(ai_assist=False))
    locals_by_ident = {item.ident: item for item in (design.local_elements or ())}

    assert locals_by_ident["l_min_f"].name == "最小值"
    assert locals_by_ident["l_min_f"].usage != "最大值"
    assert locals_by_ident["l_max_f"].name == "最大值"
    assert locals_by_ident["l_sum_f"].name in {"数据和值", "累加和", "数据和"}


def test_indexed_member_io_display_distinguishes_handshake_slots():
    ctx = {
        "global_symbol_map": {
            "s_synWhle_t": "同步整体的信息",
            "costTimHdSk_u32": "高、低握手消耗的时间",
            "HANDSK_L_ID": "低握手",
            "HANDSK_H_ID": "高握手",
        },
        "file_context": {
            "member_symbol_map": {"costTimHdSk_u32": "高、低握手消耗的时间"},
            "symbol_map": {
                "HANDSK_L_ID": "低握手",
                "HANDSK_H_ID": "高握手",
            },
        },
    }

    assert _lookup_io_display_name(
        ctx,
        "s_synWhle_t[l_synStyleID_u16].costTimHdSk_u32[HANDSK_L_ID]",
    ) == "低握手消耗的时间"
    assert _lookup_io_display_name(
        ctx,
        "s_synWhle_t[l_synStyleID_u16].costTimHdSk_u32[HANDSK_H_ID]",
    ) == "高握手消耗的时间"


def test_statement_comment_replaces_generic_raw_function_call():
    logic, _ = backend.generate_logic_from_body(
        """
        /* 通道间 帧同步监控 , 不能放在紧挨着握手结束后 */
        SynMonitor(l_synStyleID_u16);
        """,
        [],
        GenConfig(ai_assist=False),
        name_map={},
    )

    assert "通道间 帧同步监控" in logic
    assert "调用SynMonitor函数" not in logic


def test_statement_comment_replaces_mapped_simple_function_call():
    logic, _ = backend.generate_logic_from_body(
        """
        /* 通道间 帧同步监控 , 不能放在紧挨着握手结束后 */
        SynMonitor(l_synStyleID_u16);
        """,
        [],
        GenConfig(ai_assist=False),
        name_map={"SynMonitor": "模块间同步监控"},
    )

    assert "通道间 帧同步监控" in logic
    assert "调用模块间同步监控函数" not in logic


def test_semantic_call_role_uses_statement_comment_hint():
    pack = {
        "call_roles": [
            {
                "callee": "SynMonitor",
                "role": "相关处理",
                "range": {"start_line": 2, "end_line": 2},
                "comment_hints": [{"kind": "action", "text": "通道间 帧同步监控", "confidence": 0.84}],
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"SynMonitor": "模块间同步监控"},
    )

    assert "通道间 帧同步监控" in logic
    assert "完成相关处理" not in logic


def test_semantic_generic_call_role_falls_back_to_function_action_only():
    pack = {
        "call_roles": [
            {
                "callee": "delayUs",
                "role": "相关处理",
                "range": {"start_line": 2, "end_line": 2},
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"delayUs": "延时函数"},
    )

    assert "等待微秒延时" in logic
    assert "完成相关处理" not in logic
    assert "执行操作" not in logic


def test_semantic_generic_select_role_uses_callee_action_name():
    pack = {
        "call_roles": [
            {
                "callee": "ControlCCDLActiveSourceSelect",
                "role": "选择处理",
                "range": {"start_line": 2, "end_line": 2},
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))

    assert "选择控制CCDL活动来源" in logic
    assert "ActiveSource" not in logic
    assert "执行操作" not in logic


def test_void_cast_function_call_is_rendered_as_named_action():
    logic = logic_utils.heuristic_logic_line(
        "(void)ControlCCDLActiveSourceSelect(&l_ccdlID_u16, &l_ccdlValid_u16);",
        name_map={
            "l_ccdlID_u16": "CCDL标识",
            "l_ccdlValid_u16": "CCDL有效标志",
        },
    )

    assert logic == "选择控制CCDL活动来源"


def test_ai_circuit_preserves_rule_logic_and_acronym_local_names():
    func_data = {
        "comment_info": {},
        "func_info": {
            "func_name": "DemoFunc",
            "prototype": "void DemoFunc(void)",
            "ret_type": "void",
            "start": 1,
            "line_start": 1,
        },
        "body": """
        Uint16 l_ccdlID_u16 = COMM_CCDL_SCI;
        Uint16 l_ccdlValid_u16 = INVALID;
        Uint16 l_kzzzLeftOk_u16 = (RX429_STATE_OK == Comm429KZZZRxStateGet(COMM429_KZZZ_1).rxState_u16) ? VALID : INVALID;
        Uint16 l_kzzzRightOk_u16 = (RX429_STATE_OK == Comm429KZZZRxStateGet(COMM429_KZZZ_2).rxState_u16) ? VALID : INVALID;
        Uint16 l_kzzzPeerOk_u16 = INVALID;
        if (((VALID == l_kzzzLeftOk_u16) || (VALID == l_kzzzRightOk_u16)) && (VALID == l_ccdlValid_u16))
        {
            l_kzzzPeerOk_u16 = VALID;
        }
        """,
        "file_context": {
            "source_file": "/tmp/demo.c",
            "symbol_map": {
                "COMM_CCDL_SCI": "通道间SCI通信",
                "RX429_STATE_OK": "429接收状态正常",
                "COMM429_KZZZ_1": "控制装置1",
                "COMM429_KZZZ_2": "控制装置2",
                "VALID": "有效",
                "INVALID": "无效",
            },
        },
    }
    cfg = GenConfig(ai_assist=True, ai_mode=1, include_logic=True, include_locals=True)
    cfg.ai_circuit_break = True
    cfg._skip_ai_current_func = True
    cfg.enhanced_single_func_pseudocode = True

    design = build_function_design_impl(func_data, "D/R_TEST", 1, cfg)

    logic_text = "\n".join(design.logic_lines or ())
    assert "执行操作" not in logic_text
    assert "设置CCDL标识 = 通道间SCI通信" in logic_text
    assert "设置CCDL有效标志 = 无效" in logic_text
    assert "KZZZ左正常标志有效" in logic_text
    assert "KZZZ右正常标志有效" in logic_text
    assert any(item.ident == "l_ccdlID_u16" and item.name == "CCDL标识" for item in (design.local_elements or ()))
    assert any(item.ident == "l_ccdlValid_u16" and item.name == "CCDL有效标志" for item in (design.local_elements or ()))
    assert any(item.ident == "l_kzzzLeftOk_u16" and item.name == "KZZZ左正常标志" for item in (design.local_elements or ()))
    assert any(item.ident == "l_kzzzRightOk_u16" and item.name == "KZZZ右正常标志" for item in (design.local_elements or ()))
    assert any(item.ident == "l_kzzzPeerOk_u16" and item.name == "KZZZ对端正常标志" for item in (design.local_elements or ()))


def test_stale_symbol_memory_yields_to_specific_identifier_guess():
    left = backend.SymbolEvidence(
        symbol="l_kzzzLeftOk_u16",
        kind="symbols",
        decl_type="Uint16",
        usage_patterns=("assign_lhs", "condition"),
        consumer_patterns=("used_in_condition",),
        dataflow_roles=("member_snapshot", "state_snapshot", "state_value"),
        memory_cn="上拍状态",
    )
    peer = backend.SymbolEvidence(
        symbol="l_kzzzPeerOk_u16",
        kind="symbols",
        decl_type="Uint16",
        usage_patterns=("assign_lhs", "condition"),
        consumer_patterns=("used_in_condition",),
        memory_cn="标志位",
    )

    assert backend._infer_symbol_semantics_rule(left).candidate_cn == "KZZZ左正常标志"
    assert backend._infer_symbol_semantics_rule(peer).candidate_cn == "KZZZ对端正常标志"


def test_delay_call_role_is_specific_without_name_map():
    assert classify_call_role("delayUs") == "等待微秒延时"

    logic, _ = generate_logic_from_semantic_pack(
        {
            "call_roles": [
                {
                    "callee": "delayUs",
                    "role": classify_call_role("delayUs"),
                    "range": {"start_line": 1, "end_line": 1},
                }
            ]
        },
        GenConfig(ai_assist=False),
    )

    assert "等待微秒延时" in logic
    assert "执行操作" not in logic


def test_semantic_call_role_uses_actionable_condition_hint_for_generic_call():
    pack = {
        "call_roles": [
            {
                "callee": "RefuelModeExitToTaskEnd",
                "role": "相关处理",
                "range": {"start_line": 2, "end_line": 2},
                "comment_hints": [
                    {
                        "kind": "condition",
                        "text": "预位阶段若上位撤销加油模式，直接结束本轮加油流程",
                        "confidence": 0.78,
                    }
                ],
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"RefuelModeExitToTaskEnd": "把加油链统一切入任务结束态"},
    )

    assert "预位阶段若上位撤销加油模式，直接结束本轮加油流程" in logic
    assert "完成相关处理" not in logic


def test_semantic_call_role_uses_definition_comment_for_generic_call():
    pack = {
        "call_roles": [
            {
                "callee": "RefuelModeLowPressureFaultApply",
                "role": "相关处理",
                "definition_comment": "按当前加油路径对低压故障做统一收口\n置对应泵故障位、上报故障态，并切入任务结束态以禁止继续自动加油",
                "range": {"start_line": 2, "end_line": 2},
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"RefuelModeLowPressureFaultApply": "加油路径低压故障"},
    )

    assert "按当前加油路径对低压故障做统一收口" in logic
    assert "禁止继续自动加油" in logic
    assert "完成相关处理" not in logic
    assert "调用加油路径低压故障函数" not in logic


def test_logic_semantic_pack_attaches_statement_comment_hints_to_calls():
    ctx = _base_ctx(
        body="""
        /* 通道间 帧同步监控 , 不能放在紧挨着握手结束后 */
        SynMonitor(l_synStyleID_u16);
        """,
    )
    ctx["lsp_fact_pack"] = {
        "writes": [],
        "blocks": [],
        "members": [],
        "calls": [
            {
                "callee": "SynMonitor",
                "range": {"start_line": 3, "end_line": 3},
                "source": "callHierarchy",
                "confidence": 0.82,
                "verified": True,
                "definition_comment": "监控通道间同步状态",
            }
        ],
    }

    pack = build_logic_semantic_pack(ctx)

    assert pack["call_roles"][0]["comment_hints"][0]["text"] == "通道间 帧同步监控"
    assert pack["call_roles"][0]["definition_comment"] == "监控通道间同步状态"


def test_logic_semantic_pack_backfills_call_definition_comment_from_file_context():
    ctx = _base_ctx(
        body="""
        RefuelModeLowPressureFaultApply(v_p_ConData_t, l_fuelPump_un8, REFUEL_TARGET_TANK0);
        """,
    )
    ctx["file_context"]["func_comment_map"] = {
        "RefuelModeLowPressureFaultApply": "按当前加油路径对低压故障做统一收口\n置对应泵故障位、上报故障态，并切入任务结束态以禁止继续自动加油",
    }
    ctx["lsp_fact_pack"] = {
        "writes": [],
        "blocks": [],
        "members": [],
        "calls": [
            {
                "callee": "RefuelModeLowPressureFaultApply",
                "range": {"start_line": 2, "end_line": 2},
                "source": "callHierarchy",
                "confidence": 0.82,
                "verified": True,
                "definition_comment": "→ void",
            }
        ],
    }

    pack = build_logic_semantic_pack(ctx)
    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"RefuelModeLowPressureFaultApply": "加油路径低压故障"},
    )

    assert pack["call_roles"][0]["definition_comment"].startswith("按当前加油路径")
    assert "按当前加油路径对低压故障做统一收口" in logic
    assert "完成相关处理" not in logic


def test_statement_comment_still_precedes_definition_comment_for_generic_call():
    pack = {
        "call_roles": [
            {
                "callee": "RefuelModeLowPressureFaultApply",
                "role": "相关处理",
                "definition_comment": "按当前加油路径对低压故障做统一收口\n置对应泵故障位、上报故障态，并切入任务结束态以禁止继续自动加油",
                "range": {"start_line": 2, "end_line": 2},
                "comment_hints": [
                    {
                        "kind": "condition",
                        "text": "0号路径低压时，预位阶段直接按禁止自动加油故障收口",
                        "confidence": 0.8,
                    }
                ],
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"RefuelModeLowPressureFaultApply": "加油路径低压故障"},
    )

    assert "0号路径低压时，预位阶段直接按禁止自动加油故障收口" in logic
    assert "按当前加油路径对低压故障做统一收口" not in logic


def test_refined_function_title_rejects_compressed_ai_phrase():
    accepted = backend._should_accept_refined_function_title(
        "SPI闪存写禁用",
        "发送写禁关闪写保",
        func_name="SpiFlashWriteDis",
        comment_desc="发送写禁止指令，禁止后续写入",
        examples=[],
    )

    assert accepted is False
    assert backend._normalize_function_cn_title(
        "发送写禁指令禁写",
        func_name="SpiFlashWriteDis",
        comment_desc="SPI-FLASH写禁止",
    ) == "SPI闪存写禁用"
    assert _repair_function_desc_by_domain(
        "SpiFlashWriteDis",
        "发送写禁止指令以关闭FLASH写保护",
        current_desc="SPI-FLASH写禁止",
    ) == "发送写禁止指令，禁止后续FLASH写入"


def test_refined_function_title_rejects_acronym_translation():
    accepted = backend._should_accept_refined_function_title(
        "SCI中断应答",
        "科学中断确认",
        func_name="SciISRAck",
        comment_desc="",
        examples=[],
    )

    assert accepted is False
    assert naming_utils.rerank_function_title_candidates(
        "SciISRAck",
        "",
        "科学中断确认",
        ["SCI中断应答"],
        [],
    ) == "SCI中断应答"
    assert backend._normalize_function_cn_title(
        "科学中断确认",
        func_name="SciISRAck",
        comment_desc="",
    ) == "SCI中断应答"
    assert backend._is_strict_symbol_candidate_rejected("科学端口号", raw_ident="sciID")
    assert not backend._is_strict_symbol_candidate_rejected("SCI端口号", raw_ident="sciID")


def test_function_title_rejects_return_type_fragments_for_u_width_funcs():
    assert backend._normalize_function_cn_title(
        "Uint32",
        func_name="CommMaintAddrConcatU32",
        comment_desc="",
    ) == "CommMaintAddrConcatU32"
    assert backend._normalize_function_cn_title(
        "通信地址32",
        func_name="CommMaintAddrConcatU32",
        comment_desc="",
    ) == "CommMaintAddrConcatU32"
    assert backend._normalize_function_cn_title(
        "RS422接收",
        func_name="CommMaintAddrConcatU32",
        comment_desc="",
    ) == "RS422接收"


def test_isr_int_title_treats_int_as_interrupt():
    assert backend.get_function_chinese_name(
        {"desc": ""},
        {"func_name": "ISR_XNMIInt"},
    ) == "XNMI中断响应"


def test_pseudo_function_description_is_treated_as_missing():
    body = """
    if( sciID < SCI_PORT_NUM )
    {
        mySciPorts[sciID].pSci->SCIFFRX.bit.RXFFINTCLR = 1U;
        PieCtrlRegs.PIEACK.all = PIEACK_GROUP9;
    }
    """

    assert backend._is_missing_gap_text("【说明】:循环索引")
    assert ai_utils._looks_like_codeish_description("【说明】:循环索引")
    assert ai_utils._fallback_function_description(
        {"func_name": "SciISRAck"},
        body,
        current_desc="【说明】:循环索引",
    ) == "清除SCI接收FIFO中断标志并应答PIE控制器"

    ctx = _base_ctx(body=body)
    ctx["func_info"] = {
        "ret_type": "void",
        "func_name": "SciISRAck",
        "prototype": "void SciISRAck(Uint8 sciID)",
    }
    ctx["comment_info"]["desc"] = "【说明】:循环索引"
    sections = build_design_text_sections(ctx, "D/R_SDD01", 1, GenConfig(ai_assist=False))
    assert tuple(sections["description_lines"]) == ("清除SCI接收FIFO中断标志并应答PIE控制器",)


def test_placeholder_definition_comment_uses_following_real_description():
    raw_comment = """
    /**
     * 【说明】:SciTxRxEnable
     *
     * 该函数实现SCI端口接收、发送的使能和关闭功能。
     *
     * 【参数】:sciID ---- SCI口ID
     */
    """

    assert parse_single_comment_block(raw_comment)["desc"] == "【说明】:SciTxRxEnable"
    assert extract_effective_comment_desc(
        raw_comment,
        parsed_desc="【说明】:SciTxRxEnable",
        func_name="SciTxRxEnable",
    ) == "该函数实现SCI端口接收、发送的使能和关闭功能"
    assert classify_call_role("SciTxRxEnable", "【说明】:SciTxRxEnable") != "【说明】:SciTxRxEnable"


def test_call_logic_does_not_emit_placeholder_definition_comment():
    pack = {
        "call_roles": [
            {
                "callee": "SciTxRxEnable",
                "role": "【说明】:SCI接收",
                "definition_comment": "【说明】:SciTxRxEnable",
                "range": {"start_line": 1, "end_line": 1},
            }
        ]
    }

    logic, _meta = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))

    assert "执行【说明】" not in logic
    assert "【说明】" not in logic
    assert "SciTxRxEnable" not in logic


def test_header_macro_comment_keeps_acronym_qualified_tail():
    symbol_map = backend._extract_symbol_map_from_header_code(
        """
        #define SCI_A_ID           (0U)  /* SCI A 端口ID */
        #define SCI_PORT_NUM       (3U)  /* SCI 端口数量 */
        #define PLL_DIVIDE         (2)   //PLL分频系数，只能为：1,2,4其中一个
        """
    )

    assert symbol_map["SCI_A_ID"] == "SCIA端口ID"
    assert symbol_map["SCI_PORT_NUM"] == "SCI端口数量"
    assert symbol_map["PLL_DIVIDE"] == "PLL分频系数"


def test_spi_flash_datatrans_is_not_numeric_conversion():
    assert classify_call_role("spiFlashDataTrans") == "SPI Flash数据传输"
    assert classify_call_role("SPI_FLASH_DATATRANS") == "执行SPI Flash数据传输"
    assert backend._normalize_function_cn_title(
        "FLASH数据转换",
        func_name="spiFlashDataTrans",
        comment_desc="SPI-FLASH数据交互",
    ) == "SPI闪存数据传输"
    assert _repair_function_desc_by_domain(
        "spiFlashDataTrans",
        "通过SPI接口完成FLASH数据的转换传输",
    ) == "通过SPI接口完成FLASH数据交互传输"


def test_refuel_stage_preset_desc_keeps_main_control_behavior():
    assert _repair_function_desc_by_domain(
        "RefuelStagePreset",
        "校验模式有效性并初始化预设状态",
        current_desc="加油预位控制",
    ) == "根据加油模式和目标油箱发送开阀预位命令，并在阀位和泵低压检查后切换加油执行或故障结束状态"


def test_refuel_stage_preset_final_text_sections_repair_desc():
    ctx = _base_ctx(body="s_refuelCtx_t.commandSent_u16 = VALID;")
    ctx["func_info"] = {
        "ret_type": "static void",
        "func_name": "RefuelStagePreset",
        "prototype": "static void RefuelStagePreset(ConData_t *v_p_ConData_t)",
    }
    ctx["comment_info"] = {
        "func_cn_name": "校验模式写入油箱命令预设",
        "desc": "校验当前模式有效性后，写入目标油箱、命令发送及预设就绪标志",
    }

    text_sections = build_design_text_sections(
        ctx,
        "D/R_SDD01",
        1,
        GenConfig(ai_assist=False),
    )

    assert text_sections["description_lines"] == (
        "根据加油模式和目标油箱发送开阀预位命令，并在阀位和泵低压检查后切换加油执行或故障结束状态",
    )


def test_semantic_logic_includes_spi_flash_statement_macros():
    ctx = _base_ctx(
        body="""
        SPI_FLASH_CS_LOW;
        NOP;NOP;NOP;NOP;
        SPI_FLASH_DATATRANS(v_dBuff_u16, v_len_u8);
        NOP;NOP;NOP;NOP;
        SPI_FLASH_CS_HIGH;
        """,
    )
    ctx["lsp_fact_pack"] = {
        "writes": [],
        "blocks": [],
        "members": [],
        "calls": [
            {
                "callee": "SPI_FLASH_DATATRANS",
                "range": {"start_line": 4, "end_line": 4},
                "source": "structured",
                "confidence": 0.82,
                "verified": True,
            }
        ],
    }

    pack = build_logic_semantic_pack(ctx)
    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))
    lines = [line.strip() for line in logic.splitlines() if line.strip()]

    assert "拉低SPI Flash片选" in lines
    assert "执行SPI Flash数据传输" in lines
    assert "拉高SPI Flash片选" in lines
    assert any(line == "等待片选建立/保持时序裕量" for line in lines)
    assert not any("数值转换" in line for line in lines)


def test_semantic_logic_keeps_local_declaration_initializer_before_call():
    ctx = _base_ctx(
        body="Uint16 l_buff_u16 = INSTRUCTION_WRITE_DISABLE;\nspiFlashDataTrans(&l_buff_u16,1U);",
        local_vars=[{"name": "l_buff_u16", "type": "Uint16", "cn_name": "写禁用指令"}],
        symbol_map={"INSTRUCTION_WRITE_DISABLE": "写禁止指令"},
    )
    ctx["lsp_fact_pack"] = {
        "writes": [],
        "blocks": [],
        "members": [],
        "calls": [
            {
                "callee": "spiFlashDataTrans",
                "range": {"start_line": 2, "end_line": 2},
                "source": "structured",
                "confidence": 0.82,
                "verified": True,
            }
        ],
    }

    pack = build_logic_semantic_pack(ctx)
    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))
    lines = [line.strip() for line in logic.splitlines() if line.strip()]

    assert lines[0] == "设置写禁用指令 = 写禁止指令"
    assert "执行SPI Flash数据传输" in lines[1]


def test_logic_polish_repairs_recieve_typo_in_receive_mode_label():
    lines = _polish_logic_lines([
        "CASE 分支 受油模式ECIEVE",
        "计算 新模式 = 受油模式(ECIEVE)",
        "返回 返回结果(数据)",
        "IF 有效有效 等于 状态快照时",
        "返回 有效(有效(有效(有效)))",
        "将 加油阀打开命令 写入 标志位(0(0(bit0(RIU)))Send数据)的LPQD控制；",
        "将RECEIVE_RIU_STATE_ACTIVE写入标志位(RIUSendData_t)的当前值；",
        "IF 通道1 等于 存放l_nextMasterChId_u16 或 通道2 等于 存放l_nextMasterChId_u16时",
    ])

    assert lines == [
        "CASE 分支 受油模式",
        "计算 新模式 = 受油模式",
        "返回 处理结果",
        "IF 有效 等于 状态快照时",
        "返回 有效",
        "将加油阀打开命令写入RIU发送数据的LPQD控制；",
        "将RECEIVE_RIU_STATE_ACTIVE写入RIU发送数据的当前值；",
        "IF 通道1 等于 nextMasterCh标识 或 通道2 等于 nextMasterCh标识时",
    ]


def test_logic_polish_repairs_null_condition_and_system_time_call():
    lines = _polish_logic_lines([
        "IF 空 等于 系统控制数据指针时",
        "IF NULL 等于 系统控制数据指针时",
        "IF 系统控制数据指针 不等于 空时",
        "计算 系统时间 = 系统时间()；",
    ])

    assert lines == [
        "IF 系统控制数据指针为空时",
        "IF 系统控制数据指针不为空时",
        "记录当前系统时间到系统时间；",
    ]


def test_logic_polish_naturalizes_valid_invalid_conditions():
    lines = _polish_logic_lines([
        "IF 有效 等于 吊舱阀到位结果 且 有效 等于 供油通路到位结果时",
        "IF 无效 等于 预位指令已发送标志时",
        "IF 有效 不等于 工作模式有效性校验结果时",
        "IF 无效 不等于 燃油预位完成标志 或 燃油预位完成标志 等于 有效时",
    ])

    assert lines == [
        "IF 吊舱阀到位结果有效 且 供油通路到位结果有效时",
        "IF 预位指令已发送标志无效时",
        "IF 工作模式有效性校验结果无效时",
        "IF 燃油预位完成标志有效 或 燃油预位完成标志有效时",
    ]


def test_logic_condition_renders_check_call_argument_without_name_gluing():
    cond_cn, _ = _render_structured_condition_cn(
        "VALID != RefuelModeTaskValidCheck(v_p_ConData_t->workMode_u16)",
        (),
        {
            "VALID": "有效",
            "RefuelModeTaskValidCheck": "模式有效校验",
            "v_p_ConData_t.workMode_u16": "工作模式",
        },
        GenConfig(ai_assist=False),
    )
    lines = _polish_logic_lines([f"IF {cond_cn}时"])

    assert lines == ["IF 工作模式有效性校验结果无效时"]


def test_structured_condition_puts_variable_before_enum_literals():
    cond_cn, _ = _render_structured_condition_cn(
        "(WORK_MODE_LRP_FIXEDWING == v_p_ConData_t->workMode_u16) || "
        "(WORK_MODE_LRP_HELI == v_p_ConData_t->workMode_u16)",
        (),
        {
            "WORK_MODE_LRP_FIXEDWING": "左右吊舱固定翼加油模式",
            "WORK_MODE_LRP_HELI": "左右吊舱直升机加油模式",
            "v_p_ConData_t.workMode_u16": "工作模式",
        },
        GenConfig(ai_assist=False),
    )
    lines = _polish_logic_lines([f"IF {cond_cn}时"])

    assert lines == [
        "IF 工作模式 等于 左右吊舱固定翼加油模式 或 工作模式 等于 左右吊舱直升机加油模式时"
    ]


def test_simple_condition_puts_variable_before_role_macro():
    cond_cn, _ = _render_structured_condition_cn(
        "ROLE_MASTER == s_sysConData_t.runtimeRole_u16",
        (),
        {
            "ROLE_MASTER": "当前持有控制权",
            "s_sysConData_t.runtimeRole_u16": "动态控制权归属",
        },
        GenConfig(ai_assist=False),
    )
    lines = _polish_logic_lines([f"IF {cond_cn}时"])

    assert lines == ["IF 动态控制权归属 等于 当前持有控制权时"]


def test_ai_condition_translation_does_not_override_locked_or_relation(monkeypatch):
    monkeypatch.setattr(
        logic_utils,
        "_ai_structured_condition_cn",
        lambda *args, **kwargs: "维护指令执行结束等于执行状态 或 维护指令执行未结束等于执行状态",
    )

    cond_cn, _ = _render_structured_condition_cn(
        "MAINT_CMD_EXE_DONE == v_exeState_u16 || MAINT_CMD_EXE_NEW == v_exeState_u16",
        (),
        {
            "MAINT_CMD_EXE_DONE": "执行完成",
            "MAINT_CMD_EXE_NEW": "收到新指令",
            "v_exeState_u16": "执行状态",
        },
        GenConfig(ai_assist=True),
    )

    assert "执行状态 等于 维护指令执行结束" in cond_cn
    assert "执行状态 等于 维护指令执行未结束且有新指令" in cond_cn
    assert "维护指令执行结束等于执行状态" not in cond_cn


def test_structured_condition_ai_is_locked_even_when_legacy_flag_enabled(monkeypatch):
    called = {"value": False}

    def fake_ai_condition(*args, **kwargs):
        called["value"] = True
        return "被AI改写的条件"

    monkeypatch.setattr(logic_utils, "_ai_structured_condition_cn", fake_ai_condition)
    cfg = GenConfig(ai_assist=True, extra_params={"structured_cond_ai": "1"})

    cond_cn, _ = _render_structured_condition_cn(
        "v_ccdlID_u16 < COMMDRI_422_NUM",
        (),
        {"v_ccdlID_u16": "CCDL模块ID", "COMMDRI_422_NUM": "CCDL端口数量"},
        cfg,
    )

    assert called["value"] is False
    assert cond_cn == "CCDL模块ID 小于 CCDL端口数量"


def test_common_control_macros_render_without_raw_leakage():
    logic, _ = generate_logic_from_semantic_pack(
        {
            "control_blocks": [
                {
                    "kind": "if",
                    "condition": "s_sysConData_t.sysWorkTimeSum_u16 > TMIE_WORK_SUM_MAX",
                    "range": {"start_line": 1, "end_line": 3},
                }
            ],
            "state_updates": [
                {
                    "kind": "control_compute",
                    "lhs": "l_preferredID_u16",
                    "rhs": "COMM429_RIU_2",
                    "range": {"start_line": 2, "end_line": 2},
                },
                {
                    "kind": "control_compute",
                    "lhs": "l_IDData_u16",
                    "rhs": "HARD_XINT_UINT16(CPLD_ADDR_R_HKA_DATA1)",
                    "range": {"start_line": 3, "end_line": 3},
                }
            ],
        },
        GenConfig(ai_assist=False),
        name_map={
            "s_sysConData_t.sysWorkTimeSum_u16": "系统累计工作时间",
            "l_preferredID_u16": "优先通道ID",
            "l_IDData_u16": "通道编码",
        },
    )

    assert "系统累计工作时间 大于 系统累计工作时间上限" in logic
    assert "备份通道任务计算机通信SCI口" in logic
    assert "硬件16位输入寄存器读取结果" in logic
    assert "TMIE_WORK_SUM_MAX" not in logic
    assert "COMM429_RIU_2" not in logic
    assert "HARD_XINT_UINT16" not in logic


def test_pointer_output_assignment_uses_pointed_variable_wording():
    logic, _ = generate_logic_from_semantic_pack(
        {
            "state_updates": [
                {
                    "kind": "control_compute",
                    "lhs": "*vp_commID_u16",
                    "rhs": "(VALID == l_found_u16) ? l_selectedID_u16 : l_preferredID_u16",
                    "range": {"start_line": 1, "end_line": 1},
                },
                {
                    "kind": "control_compute",
                    "lhs": "*vp_valid_u16",
                    "rhs": "l_found_u16",
                    "range": {"start_line": 2, "end_line": 2},
                },
            ],
        },
        GenConfig(ai_assist=False),
        name_map={
            "vp_commID_u16": "通信ID指针",
            "vp_valid_u16": "有效标志指针",
            "VALID": "有效",
            "l_found_u16": "找到标志位",
            "l_selectedID_u16": "最终通道ID",
            "l_preferredID_u16": "优先通道ID",
        },
    )

    assert "写入通信ID指针指向的变量" in logic
    assert "将找到标志位写入有效标志指针指向的变量" in logic
    assert "作为*通信ID指针" not in logic
    assert "计算 *有效标志指针" not in logic


def test_specific_cross_file_call_roles_reduce_generic_processing_text():
    pack = {
        "call_roles": [
            {"callee": "memset", "role": "设置处理", "range": {"start_line": 1, "end_line": 1}},
            {"callee": "CommDataSourceUpdate", "role": "状态更新", "range": {"start_line": 2, "end_line": 2}},
            {"callee": "WorkModeDataObtain", "role": "相关处理", "range": {"start_line": 3, "end_line": 3}},
            {"callee": "ControlFaultDebounceReset", "role": "设置处理", "range": {"start_line": 1, "end_line": 1}},
            {"callee": "ControlModeDebounceReset", "role": "设置处理", "range": {"start_line": 2, "end_line": 2}},
            {"callee": "ControlModeReentryLatchReset", "role": "设置处理", "range": {"start_line": 3, "end_line": 3}},
            {"callee": "CHVConDataObtain", "role": "相关处理", "range": {"start_line": 4, "end_line": 4}},
            {"callee": "SpeDataGet", "role": "读取数据", "range": {"start_line": 5, "end_line": 5}},
            {"callee": "SysControlOut", "role": "相关处理", "range": {"start_line": 6, "end_line": 6}},
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))

    assert "初始化内存区域" in logic
    assert "更新通信数据来源" in logic
    assert "获取工作模式数据" in logic
    assert "复位控制故障防抖状态" in logic
    assert "复位控制模式防抖状态" in logic
    assert "复位控制模式重入锁存状态" in logic
    assert "采集CHV控制数据" in logic
    assert "读取专项存储数据" in logic
    assert "下发系统控制输出" in logic
    assert "设置处理" not in logic
    assert "完成相关处理" not in logic


def test_semantic_pack_collects_memset_args_for_zero_init_logic():
    ctx = _base_ctx(
        ret_type="KZZZ429InfoData_t",
        body="""
        KZZZ429InfoData_t l_rslt_t;
        memset(&l_rslt_t, 0, sizeof(l_rslt_t));
        return l_rslt_t;
        """,
        local_vars=[{"name": "l_rslt_t", "type": "KZZZ429InfoData_t", "cn_name": "接收结果数据"}],
        symbol_map={"l_rslt_t": "接收结果数据"},
    )

    pack = build_logic_semantic_pack(ctx, backend_module=backend)

    memset_calls = [item for item in pack["call_roles"] if item["callee"] == "memset"]
    assert len(memset_calls) == 1
    assert memset_calls[0]["args"] == ["&l_rslt_t", "0", "sizeof(l_rslt_t)"]


def test_semantic_pack_merges_memset_args_into_existing_lsp_call_fact():
    ctx = _base_ctx(
        ret_type="KZZZ429InfoData_t",
        body="""
        KZZZ429InfoData_t l_rslt_t;
        memset(&l_rslt_t, 0, sizeof(l_rslt_t));
        return l_rslt_t;
        """,
        local_vars=[{"name": "l_rslt_t", "type": "KZZZ429InfoData_t", "cn_name": "接收结果数据"}],
        symbol_map={"l_rslt_t": "接收结果数据"},
    )
    ctx["lsp_fact_pack"] = {
        "calls": [
            {
                "callee": "memset",
                "role": "初始化内存区域",
                "range": {"start_line": 3, "end_line": 3},
                "source": "definition",
                "confidence": 0.78,
                "verified": True,
            }
        ]
    }

    pack = build_logic_semantic_pack(ctx, backend_module=backend)

    memset_calls = [item for item in pack["call_roles"] if item["callee"] == "memset"]
    assert len(memset_calls) == 1
    assert memset_calls[0]["args"] == ["&l_rslt_t", "0", "sizeof(l_rslt_t)"]


def test_semantic_logic_renders_memset_zero_init_target():
    pack = {
        "call_roles": [
            {
                "callee": "memset",
                "role": "内存设置",
                "args": ["&l_rslt_t", "0", "sizeof(l_rslt_t)"],
                "range": {"start_line": 1, "end_line": 1},
            }
        ],
        "return_actions": [
            {"expr": "l_rslt_t", "range": {"start_line": 2, "end_line": 2}},
        ],
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"l_rslt_t": "接收结果数据"},
    )

    assert "清零接收结果数据" in logic
    assert "返回 接收结果数据" in logic


def test_function_call_name_is_not_polluted_by_member_alias():
    text = _logic_cn_expr(
        "Comm429RIURxStateGet(l_preferredID_u16)",
        name_map={
            "RIU": "bit0(RIU)",
            "l_preferredID_u16": "优先通道",
        },
    )

    assert "读取RIU429接收状态结果(优先通道)" == text
    assert "bit0" not in text
    assert "Comm4290" not in text


def test_riu_data_get_call_name_is_not_polluted_by_bit_alias():
    name_map = {
        "RIU": "bit0(RIU)",
        "KZZZ": "bit4",
        "l_commID_u16": "RIU通道号",
    }

    rx_text = _logic_cn_expr("Comm429RIURxDataGet(l_commID_u16)", name_map=name_map)
    orig_text = _logic_cn_expr("Comm429RIUOrigDataGet(l_commID_u16)", name_map=name_map)

    assert "Comm429RIU接收数据获取结果(RIU通道号)" == rx_text
    assert "Comm429RIU原始数据获取结果(RIU通道号)" == orig_text
    assert "bit0" not in rx_text + orig_text
    assert "bit4" not in rx_text + orig_text


def test_kzzz_ccdl_data_get_name_is_not_polluted_by_bit_alias():
    text = _logic_cn_expr(
        "Comm429KZZZCcdlExtDataGet(l_ccdlID_u16, COMM429_KZZZ_1)",
        name_map={
            "KZZZ": "bit4",
            "l_ccdlID_u16": "CCDL链路号",
            "COMM429_KZZZ_1": "控制装置1通信",
        },
    )

    assert "读取KZZZ CCDL镜像数据结果" in text
    assert "bit4" not in text


def test_ascii_parenthetical_symbol_hint_is_removed_from_logic_label():
    cond = _render_structured_condition_cn(
        "RX429_STATE_OK == l_rxState_t.rxState_u16",
        (),
        {
            "RX429_STATE_OK": "状态正常",
            "l_rxState_t.rxState_u16": "接收状态(rxState)",
        },
        None,
    )[0]

    assert cond == "接收状态 等于 状态正常"
    assert "rx" not in cond.lower()


def test_adc_sequence_macro_condition_prefers_selector_over_action_comment():
    cond = _render_structured_condition_cn(
        "ADC_SEQ1 == v_seqnum_u8",
        (),
        {
            "ADC_SEQ1": "启动ADC",
            "v_seqnum_u8": "序列号",
        },
        None,
    )[0]

    assert cond == "序列号 等于 ADC序列器1"
    assert "启动ADC" not in cond


def test_kzzz_time_request_side_macro_has_stable_side_label():
    cond = _render_structured_condition_cn(
        "0U != (l_currTimeAsk_u16 & KZZZ_TIME_REQUEST_SIDE_LEFT)",
        (),
        {
            "KZZZ": "bit4",
            "l_currTimeAsk_u16": "时间请求位图",
        },
        None,
    )[0]

    assert cond == "时间请求位图 且 左吊舱当前时间请求 不等于 0U"
    assert "bit4" not in cond
    assert "REQUESTSIDE" not in cond


def test_condition_function_member_keeps_subject_and_member_label():
    cond = _render_structured_condition_cn(
        "RX429_STATE_OK == Comm429KZZZRxStateGet(COMM429_KZZZ_1).rxState_u16",
        (),
        {
            "RX429_STATE_OK": "状态正常",
            "Comm429KZZZRxStateGet": "Comm429bit4接收状态获取",
            "COMM429_KZZZ_1": "控制装置1通信",
            "rxState_u16": "接收状态",
        },
        None,
    )[0]

    assert cond == "控制装置1通信接收状态 等于 状态正常"
    assert "bit4" not in cond


def test_local_data_state_semantic_label_preserves_left_right_side():
    assert infer_local_semantic_label("l_leftDataState_u16") == "左侧数据状态"
    assert infer_local_semantic_label("l_rightDataState_u16") == "右侧数据状态"

    pack = {
        "state_updates": [
            {
                "kind": "local_init",
                "lhs": "l_leftDataState_u16",
                "rhs": "REDUN_DATA_STATE_ERR",
                "range": {"start_line": 1, "end_line": 1},
            },
            {
                "kind": "local_init",
                "lhs": "l_rightDataState_u16",
                "rhs": "REDUN_DATA_STATE_ERR",
                "range": {"start_line": 2, "end_line": 2},
            },
        ],
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"REDUN_DATA_STATE_ERR": "数据状态异常"},
    )

    assert "设置左侧数据状态 = 数据状态异常" in logic
    assert "设置右侧数据状态 = 数据状态异常" in logic
    assert "设置状态 = 数据状态异常" not in logic


def test_kzzz_event_cache_semantic_labels_preserve_last_and_side():
    assert infer_local_semantic_label("l_lpPreFuel_u16") == "左吊舱预选油量"
    assert infer_local_semantic_label("l_rpPreFuel_u16") == "右吊舱预选油量"
    assert infer_local_semantic_label("l_lifeLeft_u16") == "左吊舱寿命信息请求"
    assert infer_local_semantic_label("lastLifeRight_u16") == "上一周期右吊舱寿命信息请求"
    assert infer_local_semantic_label("lastOilResetLeft_u16") == "上一周期左吊舱油量清零请求"

    cond = _render_structured_condition_cn(
        "s_kzzzTxCache_t.lastLifeLeft_u16 != l_lifeLeft_u16",
        (),
        {
            "s_kzzzTxCache_t.lastLifeLeft_u16": "上一周期值",
            "l_lifeLeft_u16": "缓存值",
        },
        None,
    )[0]

    assert cond == "上一周期左吊舱寿命信息请求 不等于 左吊舱寿命信息请求"
    assert "上一周期值 不等于 上一周期值" not in cond


def test_complex_function_call_condition_keeps_call_as_result():
    cond = _render_structured_condition_cn(
        "VALID == RoleConfirmUpdate(&s_masterLossConfirmCtx_t, (VALID == l_localHealthy_u16) ? INVALID : VALID, CONTROL_OWNER_HOLD_MS)",
        (),
        {
            "VALID": "有效",
            "INVALID": "无效",
            "RoleConfirmUpdate": "更新角色相关确认窗口",
            "s_masterLossConfirmCtx_t": "数据指针",
            "l_localHealthy_u16": "本端健康状态",
            "CONTROL_OWNER_HOLD_MS": "控制权切换确认保持窗口",
        },
        None,
    )[0]

    assert cond == "更新角色相关确认窗口结果 等于 有效"
    assert "&" not in cond
    assert "?" not in cond
    assert "有效 等于" not in cond



def test_c_expr_renders_low_byte_mask():
    from autodoc.c_expr import parse_c_expression, render_expr_cn
    expr = parse_c_expression("s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16] & 0xFFU")

    rendered = render_expr_cn(
        expr,
        {
            "s_rs422CommBuff_t": "接收数据缓冲区",
            "v_commID_u16": "RS422通道ID",
            "commBuff_u16": "接收数据",
            "l_ii_u16": "候选帧起始索引",
        },
    )

    assert rendered.text
    assert "低8位" in rendered.text or "低 8 位" in rendered.text
    assert "候选帧起始索引" in rendered.text
    assert "且 0xFF" not in rendered.text



def test_c_expr_prefers_full_reference_alias_before_splitting_low_byte():
    from autodoc.c_expr import parse_c_expression, render_expr_cn

    expr = parse_c_expression("l_peerBase_t.ctrlInfo_u16 & 0xFFU")

    rendered = render_expr_cn(expr, {"l_peerBase_t.ctrlInfo_u16": "基础帧控制信息字"})

    assert "基础帧控制信息字" in rendered.text
    assert "低8位" in rendered.text or "低 8 位" in rendered.text
    assert "ctrlInfo" not in rendered.text

def test_c_expr_renders_twos_complement_checksum():
    from autodoc.c_expr import parse_c_expression, render_expr_cn
    expr = parse_c_expression("(((~l_sum_u16) + 1U) & 0xFFU)")

    rendered = render_expr_cn(expr, {"l_sum_u16": "数据和"})

    assert "补码校验和" in rendered.text
    assert "低8位" in rendered.text or "低 8 位" in rendered.text
    assert "&" not in rendered.text


def test_logic_c_expr_helper_renders_nested_bitwise_expression():
    rendered = logic_utils._render_supported_c_expr_cn(
        "flags & (value & 0xFFU)",
        {"flags": "标志", "value": "数据"},
    )

    assert rendered
    assert "按位与" in rendered
    assert "&" not in rendered


def test_logic_c_expr_helper_renders_mixed_bitwise_expression():
    rendered = logic_utils._render_supported_c_expr_cn(
        "flags ^ value & 0xFFU",
        {"flags": "标志", "value": "数据"},
    )

    assert rendered
    assert "^" not in rendered
    assert "&" not in rendered


def test_supported_c_expr_renders_shift_under_byte_mask():
    rendered = logic_utils._render_supported_c_expr_cn("(value >> 8) & 0xFFU", {"value": "数据"})
    assert rendered
    assert "右移" in rendered
    assert "&" not in rendered


def test_supported_c_expr_renders_multiply_inside_masked_sum():
    rendered = logic_utils._render_supported_c_expr_cn("(value * scale + 1U) & 0xFFU", {"value": "数据", "scale": "系数"})
    assert rendered
    assert "低8位" in rendered
    assert "*" not in rendered


def test_condition_semantic_infers_and_renders_byte_mask_frame_header():
    from autodoc.semantic_elements import ConditionSemantic, infer_condition_semantic, render_condition_semantic

    semantic = infer_condition_semantic(
        "RS422_COMM_FRAME_HEAD_1 == (s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16] & 0xFFU)",
        {
            "RS422_COMM_FRAME_HEAD_1": "RS422帧头1",
            "s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16]": "报文头",
            "v_commID_u16": "RS422通道号",
            "commBuff_u16": "接收数据",
            "l_ii_u16": "候选帧起始索引",
        },
    )

    assert semantic == ConditionSemantic(
        left_label="报文头低8位",
        relation="equals",
        right_label="RS422帧头1",
    )
    assert render_condition_semantic(semantic) == "报文头低8位等于RS422帧头1"


def test_condition_semantic_derives_rs422_frame_head_macro_label_without_name_map():
    from autodoc.semantic_elements import infer_condition_semantic, render_condition_semantic

    semantic = infer_condition_semantic(
        "RS422_COMM_FRAME_HEAD_1 == (s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16] & 0xFFU)",
        {
            "s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16]": "报文头",
            "v_commID_u16": "RS422通道号",
            "commBuff_u16": "接收数据",
            "l_ii_u16": "候选帧起始索引",
        },
    )

    assert semantic is not None
    assert semantic.right_label == "RS422通信接收报文帧头1"
    assert render_condition_semantic(semantic) == "报文头低8位等于RS422通信接收报文帧头1"


def test_condition_semantic_renders_ordered_channel_bound():
    from autodoc.semantic_elements import infer_condition_semantic, render_condition_semantic

    semantic = infer_condition_semantic(
        "v_commID_u16 < COMM422_ID_NUM",
        {"v_commID_u16": "RS422通道号", "COMM422_ID_NUM": "RS422通道数量"},
    )

    assert semantic is not None
    assert render_condition_semantic(semantic) == "RS422通道号小于RS422通道数量"


def test_condition_semantic_inverts_relation_when_operands_swap():
    from autodoc.semantic_elements import infer_condition_semantic, render_condition_semantic

    semantic = infer_condition_semantic(
        "COMM422_ID_NUM > v_commID_u16",
        {"v_commID_u16": "RS422通道号", "COMM422_ID_NUM": "RS422通道数量"},
    )

    assert semantic is not None
    assert semantic.relation == "less_than"
    assert render_condition_semantic(semantic) == "RS422通道号小于RS422通道数量"


def test_condition_semantic_rejects_nested_function_call_operands():
    from autodoc.semantic_elements import infer_condition_semantic

    semantic = infer_condition_semantic(
        "RS422_COMM_FRAME_HEAD_1 == (GetFrameByte(v_commID_u16) & 0xFFU)",
        {"RS422_COMM_FRAME_HEAD_1": "RS422帧头1", "v_commID_u16": "RS422通道号"},
    )

    assert semantic is None


def test_condition_semantic_rejects_function_call_inside_raw_reference_operand():
    from autodoc.semantic_elements import infer_condition_semantic

    semantic = infer_condition_semantic(
        "RS422_COMM_FRAME_HEAD_1 == (s_rs422CommBuff_t[GetFrameByte(v_commID_u16)].commBuff_u16 & 0xFFU)",
        {"RS422_COMM_FRAME_HEAD_1": "RS422帧头1", "s_rs422CommBuff_t": "通信缓冲区"},
    )

    assert semantic is None


def test_structured_condition_uses_condition_semantic_for_byte_mask_frame_header():
    cond, _ = _render_structured_condition_cn(
        "RS422_COMM_FRAME_HEAD_1 == (s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16] & 0xFFU)",
        (),
        {
            "RS422_COMM_FRAME_HEAD_1": "RS422帧头1",
            "s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16]": "报文头",
            "v_commID_u16": "RS422通道号",
            "commBuff_u16": "接收数据",
            "l_ii_u16": "候选帧起始索引",
        },
        GenConfig(ai_assist=False),
    )

    assert cond == "报文头低8位等于RS422帧头1"
    assert "逻辑" not in cond
    assert "且 0xFF" not in cond


def test_structured_condition_locked_ai_does_not_override_condition_semantic(monkeypatch):
    called = {"value": False}

    def fake_ai_condition(*args, **kwargs):
        called["value"] = True
        return "被AI改写的报文头条件"

    monkeypatch.setattr(logic_utils, "_ai_structured_condition_cn", fake_ai_condition)
    cfg = GenConfig(ai_assist=True, extra_params={"structured_cond_ai": "1", "lock_structured_conditions": "1"})

    cond, _ = _render_structured_condition_cn(
        "RS422_COMM_FRAME_HEAD_1 == (s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16] & 0xFFU)",
        (),
        {
            "RS422_COMM_FRAME_HEAD_1": "RS422帧头1",
            "s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16]": "报文头",
            "v_commID_u16": "RS422通道号",
            "commBuff_u16": "接收数据",
            "l_ii_u16": "候选帧起始索引",
        },
        cfg,
    )

    assert called["value"] is False
    assert cond == "报文头低8位等于RS422帧头1"


def test_structured_condition_renders_byte_mask_without_logical_and_text():
    cond, _ = _render_structured_condition_cn(
        "RS422_COMM_FRAME_HEAD_1 != (s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16] & 0xFFU)",
        (),
        {
            "RS422_COMM_FRAME_HEAD_1": "RS422第一帧头",
            "s_rs422CommBuff_t": "接收数据缓冲区",
            "v_commID_u16": "RS422通道ID",
            "commBuff_u16": "接收数据",
            "l_ii_u16": "候选帧起始索引",
        },
        GenConfig(ai_assist=False),
    )

    assert "RS422第一帧头" in cond
    assert "低8位" in cond or "低 8 位" in cond
    assert "候选帧起始索引" in cond
    assert "且 0xFF" not in cond


def test_comm422_frame_check_logic_uses_low_byte_and_checksum_language():
    body = """
    Uint16 l_ii_u16 = 0U;
    Uint16 l_jj_u16 = 0U;
    Uint16 l_sum_u16 = 0U;
    Uint16 l_rData_u16 = RS422_COMM_FRAM_NOT_EXIST;
    if (v_commID_u16 < COMM422_ID_NUM)
    {
        for (l_ii_u16 = 0U; l_ii_u16 < l_count_u16; l_ii_u16++)
        {
            if (RS422_COMM_FRAME_HEAD_1 != (s_rs422CommBuff_t[v_commID_u16].commBuff_u16[l_ii_u16] & 0xFFU))
            {
                l_headErrCnt_u16 = l_headErrCnt_u16 + 1U;
            }
            l_sum_u16 = (((~l_sum_u16) + 1U) & 0xFFU);
        }
    }
    return l_rData_u16;
    """
    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={
            "v_commID_u16": "RS422通道ID",
            "COMM422_ID_NUM": "422通信数量",
            "RS422_COMM_FRAME_HEAD_1": "RS422第一帧头",
            "s_rs422CommBuff_t": "接收数据缓冲区",
            "commBuff_u16": "接收数据",
            "l_ii_u16": "候选帧起始索引",
            "l_count_u16": "候选帧数量",
            "l_sum_u16": "数据和",
            "l_rData_u16": "检测结果",
        },
    )

    assert "且 0xFF" not in logic_text
    assert "低8位" in logic_text or "低 8 位" in logic_text
    assert "补码校验和" in logic_text

def test_bit_shift_mask_condition_does_not_become_less_than_less_than():
    cond = _render_structured_condition_cn(
        "0U != (l_peerBase_t.ctrlInfo_u16 & (0x01U << COMM_CCDL_CTRLINFO_OWNER_BIT))",
        (),
        {
            "l_peerBase_t.ctrlInfo_u16": "基础帧控制信息字",
            "COMM_CCDL_CTRLINFO_OWNER_BIT": "控制权位",
        },
        None,
    )[0]

    assert "小于 小于" not in cond
    assert "按位与" in cond
    assert "左移" in cond
    assert cond.endswith("不等于 0U")


def test_fallback_logic_line_preserves_shift_condition_text():
    line = logic_utils.fallback_logic_line("if (flags << bit)", {"flags": "标志", "bit": "位"})

    assert "左移" in line
    assert "小于 小于" not in line



def test_fallback_logic_line_does_not_treat_raw_bitwise_mask_as_logical_and():
    line = logic_utils.fallback_logic_line("if (flags & 0x01U)", {"flags": "标志"})

    assert "且 0x01" not in line

def test_fallback_logic_line_preserves_while_shift_condition_text():
    line = logic_utils.fallback_logic_line("while (flags >> bit)", {"flags": "标志", "bit": "位"})

    assert "右移" in line
    assert "大于 大于" not in line


def test_function_title_strips_member_bit_alias_pollution():
    title = backend._normalize_function_cn_title(
        "Comm429bit0(RIU)Sim16",
        func_name="Comm429RIUSimU16",
    )

    assert title == "Comm429RIUSim16"
    assert "bit0" not in title


def test_peer_function_title_repairs_dropped_peer_prefix():
    title = backend._normalize_function_cn_title(
        "端控制权上报校验",
        func_name="PeerControlOwnerSeenCheck",
        comment_desc="检查对端基础帧是否上报当前控制权",
    )

    assert title == "对端控制权上报校验"


def test_ai_symbol_candidate_rejects_return_result_for_temp_local():
    assert backend._is_strict_symbol_candidate_rejected("返回结果", raw_ident="l_temp_u16")
    assert not backend._is_strict_symbol_candidate_rejected("返回结果", raw_ident="lo_rData_u16")


def test_project_symbol_memory_filters_stale_return_result_for_temp_local():
    payload = {
        "symbols": {
            "l_temp_u16": {"cn": "返回结果", "confidence": 0.99},
            "lo_rData_u16": {"cn": "返回结果", "confidence": 0.99},
        }
    }

    flat = backend._flatten_project_symbol_memory(payload)

    assert "l_temp_u16" not in flat
    assert flat["lo_rData_u16"] == "返回结果"


def test_symbol_memory_seed_filters_stale_return_result_for_temp_local():
    old_symbols = dict(backend.SYMBOL_DICTIONARY_RUNTIME)
    try:
        backend.SYMBOL_DICTIONARY_RUNTIME["l_temp_u16"] = "返回结果"
        backend.SYMBOL_DICTIONARY_RUNTIME["lo_rData_u16"] = "返回结果"
        local_vars = [
            {"name": "l_temp_u16", "cn_name": ""},
            {"name": "lo_rData_u16", "cn_name": ""},
        ]

        backend._seed_symbol_memory_into_scope({}, {}, local_vars, [], {}, {})

        assert local_vars[0].get("cn_name") == ""
        assert local_vars[1].get("cn_name") == "返回结果"
    finally:
        backend.SYMBOL_DICTIONARY_RUNTIME.clear()
        backend.SYMBOL_DICTIONARY_RUNTIME.update(old_symbols)


def test_project_symbol_memory_filters_flag_name_for_length_parameter():
    payload = {
        "symbols": {
            "v_len_u16": {"cn": "标志位", "confidence": 0.99},
            "l_found_u16": {"cn": "找到标志位", "confidence": 0.99},
        }
    }

    flat = backend._flatten_project_symbol_memory(payload)

    assert "v_len_u16" not in flat
    assert flat["l_found_u16"] == "找到标志位"


def test_bit_number_header_comment_keeps_description_tail():
    assert backend._normalize_header_comment_cn("bit0: 对端当前是否持有控制权") == "对端当前是否持有控制权"


def test_identifier_guess_handles_length_and_flash_macro_tokens():
    assert backend._guess_cn_from_ident("v_len_u16") == "长度"
    assert backend._guess_cn_from_ident("ADDR_EMPTY_CHECK_LEN_MAX") == "地址空检查长度最大"
    assert backend._guess_cn_from_ident("FLASH_SECTOR_FIRST") == "闪存扇区起始"
    assert backend._guess_cn_from_ident("RIU_TX_SSM_VALID") == "RIU发送SSM有效"
    assert backend._guess_cn_from_ident("PLL_MULMAX") == "PLL倍频最大值"


def test_condition_render_uses_macro_identifier_guess_fallback():
    cond_cn, _ = _render_structured_condition_cn(
        "v_len_u16 > ADDR_EMPTY_CHECK_LEN_MAX",
        (),
        {"v_len_u16": "长度"},
        GenConfig(ai_assist=False),
    )

    assert cond_cn == "长度 大于 地址空检查长度最大"


def test_semantic_call_uses_source_comment_hint_for_generic_read_action():
    pack = {
        "call_roles": [
            {
                "callee": "STORE_DATAREAD_DRI",
                "role": "读取结果",
                "range": {"start_line": 1, "end_line": 1},
                "comment_hints": [
                    {"kind": "condition", "text": "读取一条记录长度数据", "confidence": 0.62}
                ],
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))

    assert "读取一条记录长度数据" in logic
    assert "读取结果" not in logic


def test_semantic_pack_repairs_multiline_assignment_rhs_from_source():
    ctx = _base_ctx(
        body="""
        l_otherChvInvalidConfirmed_u16 = RoleConfirmUpdate(&s_otherChvConfirmCtx_t,
                                                          (CHV_INVALID == s_sysConData_t.CHVIn_un16.bit.otherCHV_u16) ? VALID : INVALID,
                                                          CONTROL_OWNER_HOLD_MS);
        """,
        writes=[
            {
                "lhs": "l_otherChvInvalidConfirmed_u16",
                "rhs": "RoleConfirmUpdate(&s_otherChvConfirmCtx_t,",
                "range": {"start_line": 2, "end_line": 2},
                "source": "references",
            }
        ],
    )
    ctx["lsp_fact_pack"]["function"] = {"range": {"start_line": 1}}

    pack = build_logic_semantic_pack(ctx)
    update = pack["state_updates"][0]

    assert update["rhs"].startswith("RoleConfirmUpdate(")
    assert "CONTROL_OWNER_HOLD_MS" in update["rhs"]
    assert "? VALID : INVALID" in update["rhs"]
    assert not update["rhs"].endswith(",")


def test_semantic_pack_preserves_repeated_pointer_pulse_writes():
    lhs = "*(volatile Uint16 *)(s_CCDL422RegConfs_t[v_ccdlID_u16].WReg_resetRFifo_u16)"
    ctx = _base_ctx(
        body="""
        if( v_ccdlID_u16 < COMMDRI_422_NUM )
        {
            *(volatile Uint16 *)(s_CCDL422RegConfs_t[v_ccdlID_u16].WReg_resetRFifo_u16) = DRI422_RFIFO_RESET_EN_VALID;
            *(volatile Uint16 *)(s_CCDL422RegConfs_t[v_ccdlID_u16].WReg_resetRFifo_u16) = DRI422_RFIFO_RESET_EN_INVALID;
        }
        """,
        params=[{"name": "v_ccdlID_u16", "type": "Uint16"}],
        symbol_map={
            "COMMDRI_422_NUM": "CCDL422数量",
            "DRI422_RFIFO_RESET_EN_VALID": "有效",
            "DRI422_RFIFO_RESET_EN_INVALID": "无效",
            "WReg_resetRFifo_u16": "复位接收FIFO",
        },
        writes=[
            {
                "lhs": lhs,
                "rhs": "DRI422_RFIFO_RESET_EN_VALID",
                "range": {"start_line": 4, "end_line": 4},
                "source": "references",
            },
            {
                "lhs": lhs,
                "rhs": "DRI422_RFIFO_RESET_EN_INVALID",
                "range": {"start_line": 5, "end_line": 5},
                "source": "references",
            },
        ],
    )
    ctx["lsp_fact_pack"]["function"] = {"range": {"start_line": 1}}

    pack = build_logic_semantic_pack(ctx)
    rhs_values = [item["rhs"] for item in pack["state_updates"]]
    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))

    assert rhs_values == ["DRI422_RFIFO_RESET_EN_VALID", "DRI422_RFIFO_RESET_EN_INVALID"]
    assert "将有效写入" in logic
    assert "将无效写入" in logic


def test_lsp_fact_pack_falls_back_when_payload_misses_obvious_structure(monkeypatch):
    def fake_lsp_payload(func_data, cfg=None, *, backend_module=None):
        return {
            "locals": [
                {
                    "name": "DemoPoorLsp",
                    "decl_type": "Uint16",
                    "scope": "local",
                    "source": "hover",
                    "confidence": 0.82,
                    "verified": True,
                }
            ],
            "metadata": {"provider": "lsp"},
        }

    monkeypatch.setattr(lsp_fact_utils, "_try_build_lsp_fact_pack", fake_lsp_payload)
    func_data = {
        "func_info": {
            "func_name": "DemoPoorLsp",
            "ret_type": "Uint16",
            "prototype": "Uint16 DemoPoorLsp(Uint16 v_flag_u16)",
        },
        "body": """
            Uint16 l_rData_u16 = 0U;
            if (VALID == v_flag_u16)
            {
                l_rData_u16 = DemoRead();
            }
            return l_rData_u16;
        """,
        "file_context": {"source_file": ""},
    }

    pack = lsp_fact_utils.build_function_fact_pack(func_data, GenConfig(ai_assist=False), backend_module=backend)

    assert pack.get("metadata", {}).get("lsp_degraded") is True
    assert pack.get("blocks")
    assert pack.get("writes")
    assert pack.get("calls")


def test_fallback_fact_ranges_are_offset_to_function_lines(tmp_path):
    body = """
        Uint16 l_rData_u16 = 0U;
        if (VALID == v_flag_u16)
        {
            l_rData_u16 = DemoRead();
        }
        return l_rData_u16;
    """
    source = tmp_path / "demo_offset.c"
    source.write_text(
        "\n" * 10
        + "Uint16 DemoOffset(Uint16 v_flag_u16)\n"
        + "{\n"
        + body
        + "\n}\n",
        encoding="utf-8",
    )
    func_data = {
        "func_info": {
            "func_name": "DemoOffset",
            "ret_type": "Uint16",
            "prototype": "Uint16 DemoOffset(Uint16 v_flag_u16)",
        },
        "body": body,
        "file_context": {"source_file": str(source)},
    }

    pack = lsp_fact_utils._try_build_fallback_fact_pack(
        func_data,
        GenConfig(ai_assist=False),
        backend_module=backend,
    )
    func_start = int(pack["function"]["range"]["start_line"])

    assert func_start == 11
    assert min(item["range"]["start_line"] for item in pack["blocks"]) > func_start
    assert min(item["range"]["start_line"] for item in pack["writes"]) > func_start
    assert min(item["range"]["start_line"] for item in pack["calls"]) > func_start


def test_semantic_logic_translates_receive_riu_state_macros():
    pack = {
        "state_updates": [
            {
                "kind": "state_sync",
                "lhs": "s_RIUSendData_t.currState_u16",
                "rhs": "RECEIVE_RIU_STATE_ACTIVE",
                "range": {"start_line": 1, "end_line": 1},
            },
            {
                "kind": "state_sync",
                "lhs": "s_RIUSendData_t.checkState_u16",
                "rhs": "RECEIVE_RIU_REASON_VALVE_TIMEOUT",
                "range": {"start_line": 2, "end_line": 2},
            },
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={
            "s_RIUSendData_t.currState_u16": "数据当前状态",
            "s_RIUSendData_t.checkState_u16": "数据检查状态",
        },
    )

    assert "将RIU受油执行状态写入数据当前状态" in logic
    assert "将阀位超时原因写入数据检查状态" in logic
    assert "RECEIVE_RIU" not in logic


def test_semantic_pack_renders_byte_mask_before_bitwise_fallback():
    pack = {
        "state_updates": [
            {
                "kind": "control_compute",
                "lhs": "l_byte_u16",
                "rhs": "buf[i] & 0xFFU",
                "range": {"start_line": 1, "end_line": 1},
            }
        ]
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={"l_byte_u16": "字节值", "buf": "缓冲区", "i": "索引"},
    )

    assert "低8位" in logic or "低 8 位" in logic
    assert "按位与" not in logic
    assert "& 0xFF" not in logic
    assert "&0xFF" not in logic

def test_ai_meta_carries_logic_source_audit_without_rendering_it():
    low_pressure_action = "按当前加油路径对低压故障做统一收口，置对应泵故障位、上报故障态，并切入任务结束态以禁止继续自动加油"
    ctx = _base_ctx()
    ctx["initial_gaps"] = {}
    ctx["logic_semantic_pack"] = {
        "control_blocks": [
            {
                "kind": "if",
                "condition": "NULL == v_p_ConData_t",
                "range": {"start_line": 1, "end_line": 3},
            }
        ],
        "state_updates": [
            {
                "kind": "state_sync",
                "lhs": "s_RIUSendData_t.currState_u16",
                "rhs": "RECEIVE_RIU_STATE_ACTIVE",
                "range": {"start_line": 4, "end_line": 4},
            }
        ],
        "call_roles": [
            {
                "callee": "RefuelModeLowPressureFaultApply",
                "role": "相关处理",
                "definition_comment": "按当前加油路径对低压故障做统一收口\n置对应泵故障位、上报故障态，并切入任务结束态以禁止继续自动加油",
                "range": {"start_line": 5, "end_line": 5},
            }
        ],
        "return_actions": [{"expr": "", "range": {"start_line": 6, "end_line": 6}}],
    }
    logic_lines = (
        "IF 系统控制数据指针为空时",
        "将RIU受油执行状态写入数据当前状态",
        low_pressure_action,
        "返回",
    )
    quality_inputs = {
        "logic_lines": logic_lines,
        "logic_placeholders": 0,
        "post_missing_params": [],
        "post_missing_locals": [],
        "quality_report": {
            "unresolved_locals": [],
            "unresolved_params": [],
            "unresolved_logic_symbols": [],
            "generic_logic_count": 0,
            "comment_leak_count": 0,
            "term_drift_count": 0,
            "over_translation_count": 0,
            "bad_symbol_guess_count": 0,
        },
    }

    meta = build_design_ai_meta(ctx, GenConfig(ai_assist=False), quality_inputs)
    audit = list(meta.logic_source_audit)

    assert [item["source"] for item in audit] == [
        "control_block",
        "state_update",
        "callee_comment",
        "return_action",
    ]
    assert audit[0]["refinements"] == ("null_condition_polish",)
    assert "macro_display_name" in audit[1]["refinements"]

    design = FunctionDesign(
        title="审计不渲染",
        req_id="D/R_TEST_001",
        prototype="void Demo(void)",
        description_lines=("验证审计不进入文档",),
        io_elements=(),
        io_none=True,
        local_elements=(),
        logic_lines=logic_lines,
        ai_meta=meta,
    )
    doc = Document()
    render_function_design(doc, design, GenConfig(ai_assist=False))
    text = "\n".join(par.text for par in doc.paragraphs)

    assert low_pressure_action in text
    assert "callee_comment" not in text
    assert "logic_source_audit" not in text


def test_function_call_graph_section_is_not_rendered_to_word_doc():
    design = FunctionDesign(
        title="图谱隐藏",
        req_id="D/R_TEST_002",
        prototype="void Demo(void)",
        description_lines=("验证详细设计隐藏图谱小节",),
        io_elements=(),
        io_none=True,
        local_elements=(),
        logic_lines=("调用辅助函数；",),
    )
    cfg = GenConfig(ai_assist=False)
    cfg.graph_output = "both"
    cfg._autodoc_graph_configured = True
    cfg._autodoc_graph_payloads = []
    cfg._current_render_func_data = {
        "func_info": {"func_name": "Demo"},
        "file_context": {
            "source_file": "/tmp/demo.c",
            "codegraph_callees": [{"name": "Helper", "filePath": "/tmp/helper.c", "startLine": 10}],
        },
    }
    doc = Document()

    render_function_design(doc, design, cfg)
    text = "\n".join(par.text for par in doc.paragraphs)

    assert "调用关系图" not in text
    assert "无可展示的调用关系" not in text
    assert any(payload.get("title") == "Demo" for payload in cfg._autodoc_graph_payloads)


def test_logic_alias_lookup_cleans_corrupt_riu_bit_alias():
    text = _logic_cn_expr(
        "s_RIUSendData_t.currState_u16",
        {"s_RIUSendData_t.currState_u16": "标志位(0(bit0(RIU))Send数据)的当前值"},
    )

    assert text == "RIU发送数据的当前值"


def test_ai_usage_sanitizer_preserves_refuel_supply_word():
    assert backend._sanitize_ai_usage_text("记录供油通路阀开到位") == "记录供油通路阀开到位"
    assert backend._sanitize_ai_usage_text("记录状态供后续读取") == "记录状态"


def test_semantic_logic_renders_ternary_assignment():
    pack = {
        "control_blocks": [],
        "state_updates": [
            {
                "kind": "control_compute",
                "lhs": "l_nextMasterChId_u16",
                "rhs": "(SYS_CH_ID_1 == s_sysConData_t.myChID_u16) ? SYS_CH_ID_2 : SYS_CH_ID_1",
                "range": {"start_line": 1, "end_line": 1},
            }
        ],
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={
            "l_nextMasterChId_u16": "下一主机通道ID",
            "SYS_CH_ID_1": "通道1",
            "SYS_CH_ID_2": "通道2",
        },
    )

    assert "下一主机通道ID" in logic
    assert "通道2" in logic
    assert "通道1" in logic
    assert "? " not in logic


def test_semantic_logic_renders_ternary_return():
    pack = {
        "return_actions": [
            {
                "expr": "(VALID == l_found_u16) ? l_selectedID_u16 : l_preferredID_u16",
                "range": {"start_line": 1, "end_line": 1},
            }
        ],
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={
            "VALID": "有效",
            "l_found_u16": "选中标志",
            "l_selectedID_u16": "选中通道",
            "l_preferredID_u16": "优先通道",
        },
    )

    assert "根据选中标志 等于 有效选择选中通道，否则选择优先通道作为返回值" in logic
    assert "?" not in logic


def test_semantic_logic_renders_specific_ternary_return_macro_names():
    pack = {
        "return_actions": [
            {
                "expr": "(VALID == v_valid_u16) ? RIU_TX_SSM_VALID : RIU_TX_SSM_INVALID",
                "range": {"start_line": 1, "end_line": 1},
            }
        ],
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={
            "VALID": "有效",
            "v_valid_u16": "有效数据",
            "RIU_TX_SSM_VALID": "RIU",
            "RIU_TX_SSM_INVALID": "RIU",
        },
    )

    assert "根据有效数据 等于 有效选择RIU发送SSM有效，否则选择RIU发送SSM无效作为返回值" in logic
    assert "选择RIU，否则选择RIU" not in logic
    assert "?" not in logic


def test_semantic_logic_rejects_bit_alias_for_status_return_macro():
    pack = {
        "return_actions": [
            {
                "expr": "(VALID == v_valid_u16) ? RIU_TX_SSM_VALID : RIU_TX_SSM_INVALID",
                "range": {"start_line": 1, "end_line": 1},
            }
        ],
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={
            "VALID": "有效",
            "v_valid_u16": "有效数据",
            "RIU_TX_SSM_VALID": "bit0发送SSM标号有效",
            "RIU_TX_SSM_INVALID": "bit0发送SSM标号无效",
        },
    )

    assert "RIU发送SSM有效" in logic
    assert "RIU发送SSM无效" in logic
    assert "bit0" not in logic


def test_semantic_logic_preserves_additive_assignment_operands():
    pack = {
        "state_updates": [
            {
                "kind": "control_compute",
                "lhs": "s_sysConData_t.sysWorkTimeSum_u16",
                "rhs": "s_l_baseTime_u16 + s_sysConData_t.sysWorkTime_u16",
                "range": {"start_line": 1, "end_line": 1},
            }
        ],
    }

    logic, _ = generate_logic_from_semantic_pack(
        pack,
        GenConfig(ai_assist=False),
        name_map={
            "s_sysConData_t.sysWorkTimeSum_u16": "系统累计工作时间",
            "s_l_baseTime_u16": "系统基准工作时间",
            "s_sysConData_t.sysWorkTime_u16": "系统单次工作时间",
        },
    )

    assert "将系统基准工作时间与系统单次工作时间之和写入系统累计工作时间" in logic


def test_semantic_pack_keeps_break_and_continue_flow_actions():
    ctx = _base_ctx(
        body="""
        for (l_idx_u16 = 0U; l_idx_u16 < MAX_CH_NUM; l_idx_u16++)
        {
            if (l_idx_u16 == l_preferredID_u16) { continue; }
            if (VALID == l_found_u16) { break; }
        }
        """,
        symbol_map={
            "l_idx_u16": "候选通道",
            "MAX_CH_NUM": "最大通道数量",
            "l_preferredID_u16": "优先通道",
            "l_found_u16": "选中标志",
            "VALID": "有效",
        },
    )

    pack = build_logic_semantic_pack(ctx)
    logic, _ = generate_logic_from_semantic_pack(pack, GenConfig(ai_assist=False))

    assert [item["kind"] for item in pack["flow_actions"]] == ["continue", "break"]
    assert "跳过本轮循环，进入下一轮循环" in logic
    assert "退出当前循环" in logic or "退出当前循环或分支" in logic


def test_union_container_member_alias_does_not_leak_from_other_union():
    name_map = {
        "all": "燃油系统指令1数据",
        "bit": "燃油系统指令1数据位域",
        "RIU": "bit0",
        "s_sysConData_t.commDataSourse_un16": "通信数据来源",
        "s_sysConData_t.CHVIn_un16": "通道有效输入信号",
        "CHVIn_un16.myCHV_u16": "bit0:本端运行期授权使用的本地CHV有效位",
    }

    comm = _logic_cn_expr("s_sysConData_t.commDataSourse_un16.all", name_map)
    chv = _logic_cn_expr("s_sysConData_t.CHVIn_un16.bit.myCHV_u16", name_map)

    assert comm == "通信数据来源"
    assert _logic_cn_expr("s_sysConData_t.commDataSourse_un16.bit.RIU", name_map) == "通信数据来源的RIU"
    assert "通道有效输入信号" in chv
    assert "本端运行期授权使用的本地CHV有效位" in chv
    assert "燃油系统指令1数据" not in comm + chv


def test_unified_name_resolver_precedence_and_shape():
    ctx = _base_ctx(
        local_vars=[{"name": "l_state_u16", "type": "Uint16", "cn_name": "本地状态"}],
        params=[{"name": "v_mode_u16", "type": "Uint16"}],
        symbol_map={"l_state_u16": "符号表状态", "VALID": "有效宏"},
    )
    ctx["in_map"] = {"v_mode_u16": "输入工作模式"}
    ctx["file_context"]["member_symbol_map"] = {"ConData_t.workMode_u16": "工作模式"}

    local_ref = naming_utils.resolve_symbol_display("l_state_u16", ctx=ctx)
    param_ref = naming_utils.resolve_symbol_display("v_mode_u16", ctx=ctx)
    member_ref = naming_utils.resolve_symbol_display("ConData_t.workMode_u16", ctx=ctx)
    macro_ref = naming_utils.resolve_symbol_display("INVALID", ctx=ctx)

    assert local_ref == {
        "raw": "l_state_u16",
        "display": "本地状态",
        "kind": "symbol",
        "source": "source_comment",
        "confidence": 0.95,
        "locked": True,
    }
    assert param_ref["display"] == "输入工作模式"
    assert param_ref["source"] == "source_comment"
    assert member_ref["display"] == "工作模式"
    assert member_ref["source"] == "struct_member"
    assert macro_ref["display"] == "无效"
    assert macro_ref["source"] == "macro_rule"


def test_logic_semantic_pack_v2_carries_source_anchor_provenance_and_name_refs():
    ctx = _base_ctx(
        body="""
        if (VALID != RefuelModeTaskValidCheck(v_p_ConData_t->workMode_u16))
        {
            s_RIUSendData_t.currState_u16 = RECEIVE_RIU_STATE_FAULT;
            return;
        }
        """,
        params=[{"name": "v_p_ConData_t", "type": "ConData_t *"}],
        symbol_map={
            "VALID": "有效",
            "RefuelModeTaskValidCheck": "模式有效校验",
            "v_p_ConData_t.workMode_u16": "工作模式",
            "s_RIUSendData_t.currState_u16": "数据当前状态",
        },
    )
    ctx["lsp_fact_pack"] = {
        "metadata": {"provider": "structured", "source_file": "/tmp/demo.c"},
        "blocks": [
            {
                "id": "if1",
                "kind": "if",
                "condition": "VALID != RefuelModeTaskValidCheck(v_p_ConData_t->workMode_u16)",
                "range": {"start_line": 2, "end_line": 2},
                "source": "structured",
                "confidence": 0.88,
                "verified": True,
            }
        ],
        "writes": [
            {
                "lhs": "s_RIUSendData_t.currState_u16",
                "rhs": "RECEIVE_RIU_STATE_FAULT",
                "range": {"start_line": 4, "end_line": 4},
                "source": "references",
                "confidence": 0.8,
                "verified": True,
            }
        ],
        "calls": [],
        "members": [],
    }

    pack = build_logic_semantic_pack(ctx)
    block = pack["control_blocks"][0]
    update = pack["state_updates"][0]

    assert pack["semantic_pack_version"] == 2
    assert block["source_anchor"]["file"] == "/tmp/demo.c"
    assert block["source_anchor"]["start_line"] == 2
    assert "RefuelModeTaskValidCheck" in block["source_anchor"]["raw_code"]
    assert block["provenance"]["source"] == "structured"
    assert any(ref["raw"] == "VALID" and ref["display"] == "有效" for ref in block["name_refs"])
    assert any(ref["raw"] == "s_RIUSendData_t.currState_u16" for ref in update["name_refs"])
    assert pack["quality_summary"]["item_count"] >= 2
    assert "resolver_stats" in pack


def test_controlled_ai_candidate_rejects_fact_mutation():
    semantic_pack = {
        "name_map": {"s_out_u16": "输出状态"},
        "control_blocks": [
            {
                "condition": "s_out_u16 != VALID",
                "range": {"start_line": 7, "end_line": 7},
                "name_refs": [{"raw": "s_out_u16", "display": "输出状态", "locked": True}],
            }
        ],
    }
    result = ai_utils.validate_controlled_ai_candidate(
        {
            "summary": "测试",
            "logic_line_suggestions": [{"source_line": 7, "text": "IF 输出状态 等于 有效时"}],
            "name_suggestions": [{"raw": "s_out_u16", "display": "新的输出状态"}],
            "risk_notes": [],
            "extra": "not allowed",
        },
        semantic_pack=semantic_pack,
        locked_names={"s_out_u16": "输出状态"},
    )

    assert not result["accepted"]
    codes = {item["code"] for item in result["issues"]}
    assert "locked_name_override" in codes
    assert "condition_relation_flip" in codes
    assert "unexpected_key" in codes


def test_random_doccheck_reports_structured_quality_fields(tmp_path):
    doc_path = tmp_path / "demo.docx"
    doc = Document()
    doc.add_paragraph("void DemoFunc(void)")
    doc.add_paragraph("逻辑/流程图")
    doc.add_paragraph("SWITCH 根据 模式 分支处理")
    doc.add_paragraph("CASE 分支 自动")
    doc.add_paragraph("DEFAULT 默认分支")
    doc.add_paragraph("NEXT")
    doc.add_paragraph("logic_source_audit")
    table = doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "名称"
    table.rows[0].cells[1].text = "标识"
    table.rows[0].cells[2].text = "类型"
    table.rows[0].cells[3].text = "输入/输出"
    table.rows[1].cells[0].text = "输出状态"
    table.rows[1].cells[1].text = "s_out_u16"
    table.rows[1].cells[2].text = ""
    table.rows[1].cells[3].text = "输出"
    doc.save(doc_path)

    score, warnings, _size, _paragraphs, _tables, _excerpt, details = check_docx(
        doc_path,
        Sample(c_file="/tmp/demo.c", func_name="DemoFunc", line_start=1, expected_outputs=("s_missing_u16",)),
    )

    assert score < 100
    assert "logic_source_audit" in details["docx_leak_hits"]
    assert not any(str(term).startswith("raw_macro:SWITCH") for term in details["bad_terms"])
    assert not any(str(term).startswith("raw_macro:CASE") for term in details["bad_terms"])
    assert not any(str(term).startswith("raw_macro:DEFAULT") for term in details["bad_terms"])
    assert not any(str(term).startswith("raw_macro:NEXT") for term in details["bad_terms"])
    assert "s_missing_u16" in details["missing_expected_outputs"]
    assert details["empty_type_count"] == 1
    assert any(item["code"] == "docx_internal_field_leak" for item in details["quality_issues"])
    assert warnings


def test_random_doccheck_flags_embedded_bit_alias_pollution(tmp_path):
    doc_path = tmp_path / "bit_alias.docx"
    doc = Document()
    doc.add_paragraph("void DemoFunc(void)")
    doc.add_paragraph("计算 左吊舱最终数据 = Comm429bit4CCDLExt数据获取(CCDL链路号, 控制装置1通信)")
    table = doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "名称"
    table.rows[0].cells[1].text = "标识"
    table.rows[0].cells[2].text = "类型"
    table.rows[0].cells[3].text = "输入/输出"
    table.rows[1].cells[0].text = "输入"
    table.rows[1].cells[1].text = "v_in_u16"
    table.rows[1].cells[2].text = "Uint16"
    table.rows[1].cells[3].text = "输入"
    doc.save(doc_path)

    _score, warnings, _size, _paragraphs, _tables, _excerpt, details = check_docx(
        doc_path,
        Sample(c_file="/tmp/demo.c", func_name="DemoFunc", line_start=1),
    )

    assert any(item["code"] == "bit_alias_pollution" for item in details["quality_issues"])
    assert any("位号别名污染" in warning for warning in warnings)


def test_logic_quality_gate_distinguishes_placeholder_and_truncation():
    issues = _logic_rendering_quality_issues(
        ("待人工修改", "IF 值 &&", "将结果写入状态")
    )
    assert {item["code"] for item in issues} == {"logic_placeholder", "logic_truncated"}
    assert all(item["severity"] == "error" for item in issues)
    assert not _logic_rendering_quality_issues(("IF 状态等于有效时", "将结果写入状态"))


def test_structural_quality_feedback_is_line_addressable_and_blocks_regression_candidate():
    from autodoc.quality_gate import inspect_logic_lines

    bad_lines = ("将关闭标志(写入状态；",)
    issues = inspect_logic_lines(bad_lines, source_anchors=({"idx": 1, "source": "valve_flags"},))
    bad_meta = AIBuildMeta(
        regression_needed=True,
        regression_reasons=("logic_truncated",),
        logic_source_audit=({"idx": 1, "source": "valve_flags"},),
        quality_issues=issues,
    )
    feedback = compose_quality_feedback_text(bad_meta)
    assert "logic_truncated" in feedback
    assert "逻辑第 1 行" in feedback
    assert "valve_flags" in feedback

    good = FunctionDesign("好", "D/R_001", "void Demo(void)", (), (), True, (), ("将关闭标志写入状态；",), AIBuildMeta())
    bad = FunctionDesign("坏", "D/R_001", "void Demo(void)", (), (), True, (), bad_lines, bad_meta)
    assert not prefer_regression_design(good, bad)
    assert prefer_regression_design(bad, good)


def test_build_meta_promotes_structural_logic_error_to_regression_reason():
    ctx = _base_ctx()
    ctx.update({
        "initial_gaps": {},
        "logic_semantic_pack": {},
        "scope_inference_log": {},
        "file_symbol_inference_log": {},
    })
    quality_inputs = {
        "logic_lines": ("将关闭标志(写入状态；",),
        "logic_placeholders": 0,
        "post_missing_params": (),
        "post_missing_locals": (),
        "quality_report": {
            "unresolved_locals": (), "unresolved_params": (), "unresolved_logic_symbols": (),
            "generic_logic_count": 0, "comment_leak_count": 0, "term_drift_count": 0,
            "over_translation_count": 0, "bad_symbol_guess_count": 0,
        },
    }
    meta = build_design_ai_meta(ctx, GenConfig(ai_assist=True), quality_inputs)
    assert "logic_truncated" in meta.regression_reasons
    assert any(item["code"] == "logic_truncated" for item in meta.quality_issues)


def test_structural_line_fallback_preserves_other_ai_lines():
    from autodoc.quality_gate import inspect_logic_lines

    bad_lines = ("AI 改进的说明；", "将关闭标志(写入状态；")
    anchors = ({"idx": 1, "source": "first"}, {"idx": 2, "source": "valve_flags"})
    meta = AIBuildMeta(logic_source_audit=anchors, quality_issues=inspect_logic_lines(bad_lines, source_anchors=anchors))
    design = FunctionDesign("测试", "D/R_001", "void Demo(void)", (), (), True, (), bad_lines, meta)
    baseline = FunctionDesign("测试", "D/R_001", "void Demo(void)", (), (), True, (), ("确定性说明；", "将关闭标志写入状态；"), AIBuildMeta())

    recovered = _fallback_structural_logic_lines(design, baseline)
    assert recovered.logic_lines == ("AI 改进的说明；", "将关闭标志写入状态；")
    assert not _logic_rendering_quality_issues(recovered.logic_lines)
    assert recovered.ai_meta.quality_recovery[-1]["action"] == "line_deterministic_fallback"


def test_ai_name_and_usage_reject_structural_pollution():
    from autodoc.naming import is_strict_symbol_candidate_rejected, sanitize_ai_usage_text

    assert is_strict_symbol_candidate_rejected("关闭标志(", raw_ident="RCV1_Close_u16")
    assert sanitize_ai_usage_text("执行操作（if(value &&") == ""
    assert not is_strict_symbol_candidate_rejected("阀门关闭标志", raw_ident="RCV1_Close_u16")


def test_random_doccheck_makes_logic_placeholder_and_truncation_hard_failures(tmp_path):
    doc_path = tmp_path / "logic_defect.docx"
    doc = Document()
    doc.add_paragraph("void DemoFunc(void)")
    doc.add_paragraph("逻辑/流程图")
    doc.add_paragraph("待人工修改")
    doc.add_paragraph("IF 值 &&")
    table = doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "名称"
    table.rows[0].cells[1].text = "标识"
    table.rows[0].cells[2].text = "类型"
    table.rows[0].cells[3].text = "输入/输出"
    table.rows[1].cells[0].text = "状态"
    table.rows[1].cells[1].text = "state"
    table.rows[1].cells[2].text = "Uint16"
    table.rows[1].cells[3].text = "输入"
    doc.save(doc_path)

    _score, _warnings, _size, _paragraphs, _tables, _excerpt, details = check_docx(
        doc_path, Sample(c_file="/tmp/demo.c", func_name="DemoFunc", line_start=1)
    )
    codes = {item["code"] for item in details["quality_issues"]}
    assert "logic_placeholder" in codes
    assert "logic_truncated" in codes
    assert any(item["severity"] == "error" for item in details["quality_issues"])


def test_random_doccheck_accepts_chinese_loop_index_label(tmp_path):
    source = tmp_path / "demo.c"
    source.write_text("void DemoFunc(void) { int l_ii_u16 = 0; }\n", encoding="utf-8")
    doc_path = tmp_path / "loop.docx"
    doc = Document()
    doc.add_paragraph("void DemoFunc(void)")
    doc.add_paragraph("油箱索引 l_ii_u16")
    table = doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "名称"
    table.rows[0].cells[1].text = "标识"
    table.rows[0].cells[2].text = "类型"
    table.rows[0].cells[3].text = "输入/输出"
    table.rows[1].cells[0].text = "油箱索引"
    table.rows[1].cells[1].text = "l_ii_u16"
    table.rows[1].cells[2].text = "Uint16"
    table.rows[1].cells[3].text = "局部"
    doc.save(doc_path)

    _score, _warnings, _size, _paragraphs, _tables, _excerpt, details = check_docx(
        doc_path, Sample(c_file=str(source), func_name="DemoFunc", line_start=1)
    )
    assert not any(item["code"] == "loop_index_translation_drift" for item in details["quality_issues"])


def test_random_doccheck_converts_function_offset_to_source_line(tmp_path):
    source = tmp_path / "sample.c"
    source.write_text("\n\nvoid DemoFunc(void)\n{\n}\n", encoding="utf-8")
    samples = collect_samples(tmp_path, GenConfig(ai_assist=False), count=1, seed=1, max_files=0)
    assert samples[0].func_name == "DemoFunc"
    assert samples[0].line_start == 3


def test_short_loop_index_token_recognizes_real_indices():
    """`ii`/`jj`/`kk` should match only as standalone tokens, not as
    substrings of unrelated names like `sciinfo`.
    """
    from autodoc.semantic_registry import (
        _has_short_loop_index_token,
        infer_local_semantic_label,
    )

    # Real loop indices still detected
    for name in ("l_ii", "l_ii2", "l_ii_t", "l_jj1", "kk0", "ii2"):
        assert _has_short_loop_index_token(name), f"{name!r} should be a loop index"
        assert infer_local_semantic_label(name) == "循环索引"

    # Names that happen to contain "ii" as a substring are NOT loop indices
    for name in ("sciinfo", "l_sciInfo_t", "l_sciInfo", "iidata", "aii", "biii"):
        assert not _has_short_loop_index_token(name), f"{name!r} must not match"
        assert infer_local_semantic_label(name) != "循环索引"


def test_struct_field_copy_with_linfo_no_longer_falls_back_to_pending_review():
    """Regression: assigning a struct field from a local ``l_*Info*`` variable
    used to be flagged as "待人工修改" because the local was mis-translated
    as "循环索引" (the ``ii`` in ``sciinfo`` matched the loop-index rule).
    The structural copy line should now produce a non-generic translation.
    """
    from autodoc.logic import heuristic_logic_line
    from autodoc import backend

    line = (
        "s_CCDLCommInfo_t[COMM_CCDL_SCI].rxBytesCount_u16 = "
        "l_sciInfo_t.rxBytesCount_u16;"
    )
    translation = heuristic_logic_line(line, name_map={})
    assert translation
    assert "循环索引" not in translation
    # Should not be flagged as a generic "将 X 写入 Y" filler once the
    # operands carry real, distinct semantic content.
    lhs_cn = backend._logic_cn_expr(
        "s_CCDLCommInfo_t[COMM_CCDL_SCI].rxBytesCount_u16", name_map={}
    )
    rhs_cn = backend._logic_cn_expr(
        "l_sciInfo_t.rxBytesCount_u16", name_map={}
    )
    assert lhs_cn != rhs_cn


def test_merge_multiline_continuation_ends_with_equals():
    """A multi-line assignment ending in ``=`` (typical for
    struct-field copies that wrap) must be merged into a single line
    so the downstream IR builder can recognize it as ``assign``.

    Regression: previously only lines beginning with ``&``/``|``/``&&``/``||``
    were merged, so

        s_xxx.f1 =
            l_xxx.f1;

    was split into two ``raw`` IR nodes, neither of which the simple-action
    renderer could translate.  The whole statement then fell back to
    ``待人工修改`` even when both sides carried real, distinct cn names.
    """
    from autodoc.logic import (
        _build_logic_ir_node,
        _merge_multiline_expression_line_infos,
    )

    line_infos = [
        {"code": "", "raw": "", "comments": []},
        {"code": "l_sciInfo_t = Comm422CommInfoGet(COMM422_CCDL_ID);", "raw": "l_sciInfo_t = Comm422CommInfoGet(COMM422_CCDL_ID);", "comments": []},
        {"code": "s_CCDLCommInfo_t[COMM_CCDL_SCI].rxBytesCount_u16 =", "raw": "s_CCDLCommInfo_t[COMM_CCDL_SCI].rxBytesCount_u16 =", "comments": []},
        {"code": "    l_sciInfo_t.rxBytesCount_u16;", "raw": "    l_sciInfo_t.rxBytesCount_u16;", "comments": []},
    ]
    merged = _merge_multiline_expression_line_infos(line_infos)
    codes = [li["code"] for li in merged if li["code"]]
    # The two-line struct-field copy is now a single line.
    assert "s_CCDLCommInfo_t[COMM_CCDL_SCI].rxBytesCount_u16 = l_sciInfo_t.rxBytesCount_u16;" in codes

    # And the merged line yields an ``assign`` IR node, not ``raw``.
    merged_code = codes[-1]
    ir = _build_logic_ir_node(merged_code, attached=[], name_map={}, cfg=None, use_cond_comment=False)
    assert ir["kind"] == "assign"
    assert ir["lhs"].endswith("rxBytesCount_u16")
    assert ir["rhs"] == "l_sciInfo_t.rxBytesCount_u16"


def test_merge_multiline_does_not_swallow_return_value():
    """A ``return`` header already starts a return statement; merging
    the following expression into it would corrupt the IR.  Verify the
    rule still leaves such continuations alone.
    """
    from autodoc.logic import _merge_multiline_expression_line_infos

    line_infos = [
        {"code": "return foo(", "raw": "return foo(", "comments": []},
        {"code": "  bar);", "raw": "  bar);", "comments": []},
    ]
    merged = _merge_multiline_expression_line_infos(line_infos)
    codes = [li["code"] for li in merged if li["code"]]
    # ``return`` is a control header — must NOT swallow the continuation.
    assert "return foo(" in codes
    # Continuation line keeps its raw spacing; only the merge step would strip it.
    assert any(c.strip() == "bar);" for c in codes)


def test_multiline_application_expressions_render_without_placeholders_or_raw_ops():
    """Regression fixtures for WorkModeDataObtain, CommCCDL and Refuel paths."""
    body = """
    l_newMode_RIU_u16 = WorkModeRIUDataCheck(
        l_oilCMD_RIU_un32.bit.fuelObject_u8,
        l_oilCMD_RIU_un32.bit.fuelMode_u8);
    if((s_info[l_id_u16].checkTime_u32 -
        s_info[l_id_u16].rxBytesTime_u32) >= TIMEOUT_MS)
    {
        l_ok_u16 = 1U;
    }
    l_valveClosed_u16 = (l_cmd_un.bit.valveA_u8 & l_state_un.bit.valveA_u8 & l_enable_un.bit.valveA_u8) &&
        (l_cmd_un.bit.valveB_u8 == l_state_un.bit.valveB_u8);
    """
    logic_text, unknowns = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={
            "l_newMode_RIU_u16": "新模式RIU",
            "WorkModeRIUDataCheck": "工作模式RIU数据检查",
            "l_oilCMD_RIU_un32": "燃油命令RIU",
            "fuelObject_u8": "燃油对象",
            "fuelMode_u8": "燃油模式",
            "s_info": "通信信息",
            "l_id_u16": "通道索引",
            "checkTime_u32": "检查时间",
            "rxBytesTime_u32": "接收时间",
            "TIMEOUT_MS": "超时阈值",
            "l_ok_u16": "检查结果",
            "l_valveClosed_u16": "阀门关闭状态",
            "l_cmd_un": "阀门命令",
            "l_state_un": "阀门状态",
            "l_enable_un": "阀门使能",
            "valveA_u8": "阀门A",
            "valveB_u8": "阀门B",
        },
    )

    assert not unknowns
    assert "结果写入 新模式RIU" in logic_text
    assert "大于等于 超时阈值" in logic_text
    assert "按位与结果且" in logic_text
    assert not any(token in logic_text for token in ("待人工修改", "if(", "&&", "执行操作（"))
    assert not _logic_rendering_quality_issues(logic_text.splitlines())


def test_bitwise_operand_keeps_disambiguation_parentheses_balanced():
    rendered = logic_utils._simplify_bitwise_operand_text("关闭标志(RCV1_Close)")
    assert rendered == "关闭标志(RCV1_Close)"
    assert not _logic_rendering_quality_issues((f"将{rendered}写入状态；",))


def test_raw_macro_hits_allow_interrupt_function_names():
    """``ISR_*``/``*_ISR`` function names and other interrupt-handler
    naming conventions must not be flagged as raw macros in the
    docx audit.  The audit's allow-list previously only knew a small
    set of acronyms and missed the obvious interrupt-handler family.
    """
    import importlib.util
    import sys
    from pathlib import Path

    spec = importlib.util.spec_from_file_location(
        "random_function_doccheck",
        str(Path(__file__).resolve().parent.parent / "tools" / "random_function_doccheck.py"),
    )
    mod = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = mod
    spec.loader.exec_module(mod)

    # All these are legitimate interrupt-handler / peripheral function
    # names that should NOT show up in raw_macro hits.
    for text in (
        "interrupt void ISR_SCIB_RXINT(void)",
        "interrupt void rsvd_ISR(void)",
        "interrupt void SCIA_TXINT(void)",
        "interrupt void XINT1_IRQ(void)",
        "interrupt void NMI_Handler(void)",
        "void PIE_ACK_GROUP9(void)",
        "interrupt void EPWM1_INT(void)",
        "void I2C_SLAVE_CALLBACK(void)",
    ):
        assert mod._raw_macro_hits(text) == [], (
            f"unexpected hits for {text!r}: {mod._raw_macro_hits(text)!r}"
        )

    # Requirement ID prefixes must also be allowed, in both slash and
    # underscore forms.
    for text in (
        "D/R_SDD01",
        "D/R_SDD01_001",
        "R_SDD01",
        "R_SDD01_001",
        "D_R_SDD01_001",
    ):
        assert mod._raw_macro_hits(text) == [], (
            f"SDD prefix {text!r} should be allowed but got {mod._raw_macro_hits(text)!r}"
        )

    # Genuinely suspicious ALL_CAPS tokens still get caught.
    hits = mod._raw_macro_hits("FOO_BAR_BAZ DEBUG_LOG")
    assert set(hits) == {"FOO_BAR_BAZ", "DEBUG_LOG"}


def test_logic_cleanup_removes_empty_else_branch_from_no_deal_comment():
    body = """
    if (v_len_16 > 3U)
    {
        l_fData_f = v_pBuff_f[0];
    }
    else
    {
        /* no deal to do */
    }
    return l_fData_f;
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={"v_len_16": "长度", "l_fData_f": "平均值", "v_pBuff_f": "缓冲"},
    )

    lines = [line.strip() for line in logic_text.splitlines() if line.strip()]
    assert "ELSE" not in lines
    assert "END IF" in lines


def test_fdataaverage_logic_contains_no_bare_empty_else_branch():
    body = """
    if (v_pBuff_f != NULL)
    {
        if (v_len_16 > 3U)
        {
            l_fData_f = v_pBuff_f[0];
        }
        else if (v_len_16 > 0U)
        {
            l_fData_f = v_pBuff_f[0];
        }
        else
        {
            /* no deal to do */
        }
    }
    return l_fData_f;
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={"v_pBuff_f": "缓冲", "v_len_16": "长度", "l_fData_f": "平均值"},
    )

    lines = [line.strip() for line in logic_text.splitlines() if line.strip()]
    assert "ELSE" not in lines
    assert any(line.startswith("ELSE IF") for line in lines)


def test_switch_case_indentation_survives_final_logic_cleanup():
    body = """
    switch (mode) {
    case 1:
        result = value;
        break;
    default:
        result = 0U;
        break;
    }
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={"mode": "模式", "result": "结果", "value": "输入值"},
    )

    lines = [line for line in logic_text.splitlines() if line.strip()]
    stripped = [line.strip() for line in lines]
    case_idx = next(i for i, line in enumerate(stripped) if line.startswith("CASE 分支"))
    default_idx = next(i for i, line in enumerate(stripped) if line.startswith("DEFAULT"))
    assert any(line.startswith("CASE 分支") and "1" in line for line in stripped)
    assert any(line.startswith("DEFAULT") for line in stripped)

    case_indent = len(lines[case_idx]) - len(lines[case_idx].lstrip(" "))
    default_indent = len(lines[default_idx]) - len(lines[default_idx].lstrip(" "))
    case_body = lines[case_idx + 1 : default_idx]
    default_body = lines[default_idx + 1 :]
    assert any("结果" in line and "输入值" in line for line in case_body)
    assert any("结果" in line for line in default_body)
    assert all(len(line) - len(line.lstrip(" ")) > case_indent for line in case_body)
    assert all(
        len(line) - len(line.lstrip(" ")) > default_indent
        for line in default_body
        if line.strip() != "END SWITCH"
    )

def test_ai_non_structured_unknown_indexes_follow_final_cleanup():
    body = """
    if (flag)
    {
        value = input;
    }
    else
    {
        /* no deal to do */
    }
    status + value;
    """

    logic_text, unknowns = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=True, ai_mode=2, ai_logic_policy="ai_non_structured"),
        name_map={"flag": "标志", "value": "数值", "input": "输入", "status": "状态"},
    )

    lines = logic_text.splitlines()
    assert unknowns
    for unknown in unknowns:
        idx = int(unknown["idx"])
        assert idx < len(lines)
        assert lines[idx].strip() == "待人工修改"
        lines[idx] = "AI填充"
        assert lines[idx] == "AI填充"
    assert "END IF" in logic_text.splitlines()
    assert "ELSE" not in [line.strip() for line in logic_text.splitlines()]


def test_logic_cleanup_collapses_duplicate_setup_but_preserves_loop_reset():
    body = """
    Uint16 l_headErrCnt_u16 = 0U;
    l_headErrCnt_u16 = 0U;
    l_headErrCnt_u16 = 0U;
    for (l_ii_u16 = 0U; l_ii_u16 < l_count_u16; l_ii_u16++)
    {
        l_headErrCnt_u16 = 0U;
        if (l_headErrCnt_u16 == 0U)
        {
            l_rData_u16 = 1U;
        }
    }
    return l_rData_u16;
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={
            "l_headErrCnt_u16": "帧头错误计数",
            "l_ii_u16": "候选帧起始索引",
            "l_count_u16": "候选帧数量",
            "l_rData_u16": "检测结果",
        },
    )

    lines = [line.strip() for line in logic_text.splitlines() if line.strip()]
    reset_lines = [line for line in lines if "帧头错误计数" in line and ("清零" in line or "设置" in line or "初始化" in line)]
    assert len(reset_lines) <= 2
    for_index = next(i for i, line in enumerate(lines) if line.startswith("FOR"))
    assert any("帧头错误计数" in line and ("清零" in line or "设置" in line or "初始化" in line) for line in lines[for_index + 1:])


def test_comm422_setup_reset_preserves_loop_reset_without_repeated_top_level_setup():
    body = """
    Uint16 l_headErrCnt_u16 = 0U;
    l_headErrCnt_u16 = 0U;
    l_headErrCnt_u16 = 0U;
    Uint16 l_ii_u16 = 0U;
    Uint16 l_count_u16 = 4U;
    Uint16 l_rData_u16 = RS422_COMM_FRAM_NOT_EXIST;
    for (l_ii_u16 = 0U; l_ii_u16 < l_count_u16; l_ii_u16++)
    {
        l_headErrCnt_u16 = 0U;
        if (l_headErrCnt_u16 == 0U)
        {
            l_rData_u16 = l_ii_u16;
        }
    }
    return l_rData_u16;
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={
            "l_headErrCnt_u16": "帧头错误计数",
            "l_ii_u16": "候选帧起始索引",
            "l_count_u16": "候选帧数量",
            "l_rData_u16": "检测结果",
            "RS422_COMM_FRAM_NOT_EXIST": "422通信帧不存在",
        },
    )

    lines = [line.strip() for line in logic_text.splitlines() if line.strip()]
    for_index = next(i for i, line in enumerate(lines) if line.startswith("FOR"))
    frame_head_setup_before_for = [
        line
        for line in lines[:for_index]
        if "帧头错误计数" in line and ("清零" in line or "设置" in line or "初始化" in line)
    ]
    frame_head_reset_after_for = [
        line
        for line in lines[for_index + 1:]
        if "帧头错误计数" in line and ("清零" in line or "设置" in line or "初始化" in line)
    ]

    assert frame_head_setup_before_for
    assert len(frame_head_setup_before_for) <= 1
    assert frame_head_reset_after_for


def test_logic_cleanup_preserves_trailing_assignment_plain_write():
    body = """
    l_result_u16 = l_value_u16;
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={"l_result_u16": "结果", "l_value_u16": "输入值"},
    )

    assert "将 输入值 写入 结果" in logic_text.splitlines()


def test_logic_cleanup_preserves_do_while_body_nesting():
    body = """
    do
    {
        l_count_u16 = l_count_u16 + 1U;
    } while (l_count_u16 < 3U);
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={"l_count_u16": "计数"},
    )

    lines = logic_text.splitlines()
    do_index = next(i for i, line in enumerate(lines) if line.startswith("DO WHILE"))
    end_index = next(i for i, line in enumerate(lines) if line == "END DO WHILE")
    body_indices = [i for i, line in enumerate(lines) if line.startswith("    ") and "计数" in line and ("写入" in line or "计算" in line)]

    assert do_index < end_index
    assert any(do_index < i < end_index for i in body_indices)


def test_validate_control_blocks_ignores_orphan_end_if_inside_do_while():
    lines = logic_utils._validate_control_blocks(
        ["DO WHILE 条件 时", "    更新计数", "END IF", "END DO WHILE"]
    )

    assert lines == ["DO WHILE 条件时", "    更新计数", "END DO WHILE"]



def test_logic_cleanup_keeps_same_target_setup_lines_with_different_values():
    body = """
    l_status_u16 = l_initial_u16;
    l_status_u16 = l_normal_u16;
    """

    logic_text, _ = logic_utils.generate_logic_from_body(
        body,
        [],
        GenConfig(ai_assist=False),
        name_map={"l_status_u16": "状态", "l_initial_u16": "初始化值", "l_normal_u16": "正常值"},
    )

    assert logic_text.splitlines() == ["将 初始化值 写入 状态", "将 正常值 写入 状态"]

    setup_lines = ["设置状态 = 初始化值", "设置状态 = 正常值"]
    assert logic_utils._collapse_duplicate_setup_lines(setup_lines) == setup_lines
    assert logic_utils._collapse_duplicate_setup_lines([setup_lines[0], setup_lines[0]]) == [setup_lines[0]]


def test_review_workspace_serializes_and_renders_escaped_html(tmp_path):
    from autodoc.review_workspace import (
        ReviewBlock,
        ReviewBundle,
        ReviewFunction,
        render_review_html,
        review_block_id,
        review_bundle_to_dict,
    )

    block_id = review_block_id("Comm422FrameCheck", "logic", 1)
    assert block_id == "Comm422FrameCheck.logic.001"

    bundle = ReviewBundle(
        schema_version=1,
        project_root="/project",
        output_docx="/project/out.docx",
        functions=(
            ReviewFunction(
                function_id="Comm422FrameCheck",
                name="Comm422FrameCheck",
                title="422 <Frame> & Check",
                source_file="Comm422.c",
                source_hash="abc123",
                blocks=(
                    ReviewBlock(
                        block_id=block_id,
                        function_id="Comm422FrameCheck",
                        kind="logic_line",
                        title="Logic <1>",
                        text="if (x < y) & dangerous <script>alert(1)</script>",
                        confidence=0.75,
                    ),
                ),
            ),
        ),
    )

    data = review_bundle_to_dict(bundle)
    assert data["schema_version"] == 1
    assert data["functions"][0]["blocks"][0]["block_id"] == block_id
    assert data["functions"][0]["blocks"][0]["confidence"] == 0.75

    html = render_review_html(bundle)
    assert "data-block-id=\"Comm422FrameCheck.logic.001\"" in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html
    assert "<script>alert(1)</script>" not in html
    assert "review_bundle.json" in html



def test_review_workspace_renders_bundle_quality_flags_escaped():
    from autodoc.review_workspace import ReviewBundle, ReviewQualityFlag, render_review_html

    bundle = ReviewBundle(
        quality_flags=(
            ReviewQualityFlag(
                code="bundle-html",
                severity="warning",
                message="Bundle <flag> & <script>alert(1)</script>",
            ),
        ),
    )

    html = render_review_html(bundle)

    assert "warning: Bundle &lt;flag&gt; &amp; &lt;script&gt;alert(1)&lt;/script&gt;" in html
    assert "Bundle <flag> & <script>alert(1)</script>" not in html
def test_review_workspace_block_id_sanitizes_non_identifier_names():
    from autodoc.review_workspace import review_block_id

    assert review_block_id("模块/函数 名", "params.input", 12) == "_____.params.input.012"
    assert review_block_id("", "logic", 2) == "function.logic.002"


def test_build_review_function_creates_summary_table_and_logic_blocks():
    from autodoc.models import FunctionDesign, IOElement, LocalDataElement
    from autodoc.review_workspace import build_review_function

    design = FunctionDesign(
        title="帧检查",
        req_id="D/R_SDD01_001",
        prototype="Uint16 Comm422FrameCheck(Uint16 v_commID_u16)",
        description_lines=("检测接收缓冲区是否存在有效报文。",),
        io_elements=(IOElement("通道号", "v_commID_u16", "Uint16", "输入"),),
        io_none=False,
        local_elements=(LocalDataElement("候选帧索引", "l_ii_u16", "Uint16", "局部变量"),),
        logic_lines=("IF 通道号有效时", "遍历候选帧。"),
    )
    func_data = {
        "func_info": {"func_name": "Comm422FrameCheck", "prototype": design.prototype},
        "source_file": "Comm422.c",
        "body": "return 0;\n",
    }

    review_fn = build_review_function(design, func_data)

    assert review_fn.function_id == "Comm422FrameCheck"
    block_ids = [b.block_id for b in review_fn.blocks]
    assert "Comm422FrameCheck.summary.001" in block_ids
    assert "Comm422FrameCheck.prototype.001" in block_ids
    assert "Comm422FrameCheck.io.001" in block_ids
    assert "Comm422FrameCheck.locals.001" in block_ids
    assert "Comm422FrameCheck.logic.001" in block_ids
    assert "Comm422FrameCheck.logic.002" in block_ids
    logic_texts = [b.text for b in review_fn.blocks if b.kind == "logic_line"]
    assert logic_texts == ["IF 通道号有效时", "遍历候选帧。"]
    assert review_fn.source_hash


def test_build_review_function_preserves_logic_line_text_while_skipping_blank_lines():
    from autodoc.models import FunctionDesign
    from autodoc.review_workspace import build_review_function

    design = FunctionDesign(
        title="缩进逻辑",
        req_id="D/R_SDD01_002",
        prototype="void Demo(void)",
        description_lines=(),
        io_elements=(),
        io_none=True,
        local_elements=(),
        logic_lines=("  IF 条件成立时", "   ", "\t执行缩进动作"),
    )

    review_fn = build_review_function(design, {"func_info": {"func_name": "Demo"}})

    logic_texts = [b.text for b in review_fn.blocks if b.kind == "logic_line"]
    assert logic_texts == ["  IF 条件成立时", "\t执行缩进动作"]


def test_build_review_function_uses_comment_return_desc_when_design_has_no_return_field():
    from autodoc.models import FunctionDesign
    from autodoc.review_workspace import build_review_function

    design = FunctionDesign(
        title="返回说明",
        req_id="D/R_SDD01_003",
        prototype="Uint16 Demo(void)",
        description_lines=(),
        io_elements=(),
        io_none=True,
        local_elements=(),
        logic_lines=(),
    )
    func_data = {
        "func_info": {"func_name": "Demo", "prototype": design.prototype},
        "comment_info": {"return_desc": "返回检测结果。"},
    }

    review_fn = build_review_function(design, func_data)

    return_blocks = [b for b in review_fn.blocks if b.kind == "return"]
    assert [b.text for b in return_blocks] == ["返回检测结果。"]


def test_write_review_workspace_creates_json_and_html(tmp_path):
    import json
    from autodoc.review_workspace import ReviewBlock, ReviewBundle, ReviewFunction, write_review_workspace

    bundle = ReviewBundle(
        output_docx=str(tmp_path / "out.docx"),
        functions=(
            ReviewFunction(
                function_id="DemoFunc",
                name="DemoFunc",
                blocks=(ReviewBlock("DemoFunc.logic.001", "DemoFunc", "logic_line", text="执行处理"),),
            ),
        ),
    )

    out_dir = write_review_workspace(bundle, str(tmp_path / "review"))

    assert out_dir == str(tmp_path / "review")
    data = json.loads((tmp_path / "review" / "review_bundle.json").read_text(encoding="utf-8"))
    assert data["functions"][0]["function_id"] == "DemoFunc"
    assert "data-block-id=\"DemoFunc.logic.001\"" in (tmp_path / "review" / "index.html").read_text(encoding="utf-8")



def test_write_review_workspace_overwrites_existing_bundle_on_full_run(tmp_path):
    import json
    from autodoc.review_workspace import ReviewBlock, ReviewBundle, ReviewFunction, write_review_workspace

    review_dir = tmp_path / "review"
    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="OldFunc",
                    name="OldFunc",
                    source_file="old.c",
                    blocks=(ReviewBlock("OldFunc.logic.001", "OldFunc", "logic_line", text="旧函数"),),
                ),
            ),
        ),
        str(review_dir),
    )

    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="NewFunc",
                    name="NewFunc",
                    source_file="new.c",
                    blocks=(ReviewBlock("NewFunc.logic.001", "NewFunc", "logic_line", text="新函数"),),
                ),
            ),
        ),
        str(review_dir),
    )

    data = json.loads((review_dir / "review_bundle.json").read_text(encoding="utf-8"))
    assert [fn["name"] for fn in data["functions"]] == ["NewFunc"]

def test_write_review_workspace_merges_resume_bundle_and_dedupes_current_function(tmp_path):
    import json
    from autodoc.review_workspace import ReviewBlock, ReviewBundle, ReviewFunction, write_review_workspace

    review_dir = tmp_path / "review"
    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="OldFunc",
                    name="OldFunc",
                    source_file="old.c",
                    blocks=(ReviewBlock("OldFunc.logic.001", "OldFunc", "logic_line", text="保留旧函数"),),
                ),
                ReviewFunction(
                    function_id="DemoFunc",
                    name="DemoFunc",
                    source_file="demo.c",
                    blocks=(ReviewBlock("DemoFunc.logic.001", "DemoFunc", "logic_line", text="旧内容"),),
                ),
            ),
        ),
        str(review_dir),
    )

    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="DemoFunc",
                    name="DemoFunc",
                    source_file="demo.c",
                    blocks=(ReviewBlock("DemoFunc.logic.001", "DemoFunc", "logic_line", text="新内容"),),
                ),
            ),
        ),
        str(review_dir),
        merge_existing=True,
    )

    data = json.loads((review_dir / "review_bundle.json").read_text(encoding="utf-8"))
    functions = data["functions"]
    assert [fn["name"] for fn in functions] == ["OldFunc", "DemoFunc"]
    assert functions[0]["blocks"][0]["text"] == "保留旧函数"
    assert functions[1]["blocks"] == [
        {
            "block_id": "DemoFunc.logic.001",
            "function_id": "DemoFunc",
            "kind": "logic_line",
            "title": "",
            "text": "新内容",
            "rows": [],
            "source_range": {"file": "", "start_line": 0, "end_line": 0},
            "evidence": [],
            "quality_flags": [],
            "confidence": 1.0,
            "editable": True,
        }
    ]


def test_write_review_workspace_disambiguates_resumed_duplicate_names_by_source(tmp_path):
    import json
    from autodoc.review_workspace import ReviewBlock, ReviewBundle, ReviewFunction, write_review_workspace

    review_dir = tmp_path / "review"
    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="Demo",
                    name="Demo",
                    source_file="src/a/file_a.c",
                    blocks=(ReviewBlock("Demo.logic.001", "Demo", "logic_line", text="A 文件逻辑"),),
                ),
            ),
        ),
        str(review_dir),
    )

    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="Demo",
                    name="Demo",
                    source_file="src/b/file_b.c",
                    blocks=(ReviewBlock("Demo.logic.001", "Demo", "logic_line", text="B 文件逻辑"),),
                ),
            ),
        ),
        str(review_dir),
        merge_existing=True,
    )

    data = json.loads((review_dir / "review_bundle.json").read_text(encoding="utf-8"))
    functions = data["functions"]
    assert [fn["name"] for fn in functions] == ["Demo", "Demo"]
    assert functions[0]["function_id"] == "Demo"
    assert functions[1]["function_id"].startswith("Demo_file_b")
    assert functions[1]["blocks"][0]["function_id"] == functions[1]["function_id"]
    assert functions[1]["blocks"][0]["block_id"].startswith(functions[1]["function_id"] + ".logic.")


def test_write_review_workspace_preserves_existing_disambiguated_id_on_replacement(tmp_path):
    import json
    from autodoc.review_workspace import ReviewBlock, ReviewBundle, ReviewFunction, write_review_workspace

    review_dir = tmp_path / "review"
    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="Demo",
                    name="Demo",
                    source_file="src/a/file_a.c",
                    blocks=(ReviewBlock("Demo.logic.001", "Demo", "logic_line", text="A 文件逻辑"),),
                ),
                ReviewFunction(
                    function_id="Demo_file_b",
                    name="Demo",
                    source_file="src/b/file_b.c",
                    blocks=(ReviewBlock("Demo_file_b.logic.001", "Demo_file_b", "logic_line", text="B 旧逻辑"),),
                ),
            ),
        ),
        str(review_dir),
    )

    write_review_workspace(
        ReviewBundle(
            functions=(
                ReviewFunction(
                    function_id="Demo",
                    name="Demo",
                    source_file="src/b/file_b.c",
                    blocks=(ReviewBlock("Demo.logic.001", "Demo", "logic_line", text="B 新逻辑"),),
                ),
            ),
        ),
        str(review_dir),
        merge_existing=True,
    )

    data = json.loads((review_dir / "review_bundle.json").read_text(encoding="utf-8"))
    functions = data["functions"]
    assert [fn["name"] for fn in functions] == ["Demo", "Demo"]
    assert [fn["function_id"] for fn in functions] == ["Demo", "Demo_file_b"]
    assert functions[1]["blocks"][0]["text"] == "B 新逻辑"
    assert functions[1]["blocks"][0]["function_id"] == "Demo_file_b"
    assert functions[1]["blocks"][0]["block_id"] == "Demo_file_b.logic.001"


def test_review_collection_disambiguates_duplicate_function_names_by_source():
    from autodoc.models import FunctionDesign
    from autodoc.pipeline import _collect_review_function

    cfg = GenConfig(ai_assist=False, extra_params={"review_output": "html"})
    design = FunctionDesign(
        title="重复函数",
        req_id="D/R_SDD01_004",
        prototype="void Demo(void)",
        description_lines=("处理。",),
        io_elements=(),
        io_none=True,
        local_elements=(),
        logic_lines=("执行处理",),
    )

    _collect_review_function(cfg, design, {"func_name": "Demo", "source_file": "src/a/file_a.c", "func_data": {"func_info": {"func_name": "Demo"}, "source_file": "src/a/file_a.c"}})
    _collect_review_function(cfg, design, {"func_name": "Demo", "source_file": "src/b/file_b.c", "func_data": {"func_info": {"func_name": "Demo"}, "source_file": "src/b/file_b.c"}})

    functions = cfg._review_workspace_functions
    assert [fn.name for fn in functions] == ["Demo", "Demo"]
    assert functions[0].function_id == "Demo"
    assert functions[1].function_id.startswith("Demo_file_b")
    assert functions[1].blocks[0].function_id == functions[1].function_id
    assert functions[1].blocks[0].block_id.startswith(functions[1].function_id + ".summary.")


def test_review_workspace_config_defaults_off():
    from autodoc.review_workspace import review_output_enabled, review_output_dir

    cfg = GenConfig(ai_assist=False)
    assert review_output_enabled(cfg) is False
    assert review_output_dir(cfg, "/tmp/out.docx").endswith("out_review")


def test_cli_doc_accepts_review_output_flags(monkeypatch):
    from autodoc.cli import parse_args

    monkeypatch.setattr(
        sys,
        "argv",
        [
            "AutoDocGen_V1.4.py",
            "doc",
            "-f",
            "demo.c",
            "-o",
            "out.docx",
            "--review-output",
            "html",
            "--review-dir",
            "review-out",
        ],
    )
    args = parse_args()
    assert args.review_output == "html"
    assert args.review_dir == "review-out"
