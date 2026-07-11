import os
import sys
import json

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from autodoc.render import delete_csu_in_doc, insert_csu_after_in_doc, renumber_module_csu_ids, sync_module_function_table_for_module
from qt_gui.runner import DocUpdateWorker, TaskSpec
from qt_gui.settings_store import AppSettings
from tools.update_doc_from_code_diff import (
    PlannedItem,
    allocate_next_csu_id,
    apply_review_decisions,
    attach_alignment_to_items,
    build_csu_index,
    build_doc_code_alignment_index,
    classify_changes,
)
from tools import update_doc_from_code_diff as updater
from tools.render_update_review_html import render_review_html


def _write_two_csu_doc(path):
    doc = Document()
    doc.add_heading("DemoModule（D/R_SDD01_001）", level=3)
    table = doc.add_table(rows=3, cols=4)
    for index, header in enumerate(["CSC 名称", "CSC 标识", "CSU 名称", "CSU 标识"]):
        table.rows[0].cells[index].text = header
    table.rows[1].cells[0].text = "DemoModule"
    table.rows[1].cells[1].text = "D/R_SDD01_001"
    table.rows[1].cells[2].text = "OldFunc"
    table.rows[1].cells[3].text = "D/R_SDD01_001_001"
    table.rows[2].cells[0].text = "DemoModule"
    table.rows[2].cells[1].text = "D/R_SDD01_001"
    table.rows[2].cells[2].text = "KeepFunc"
    table.rows[2].cells[3].text = "D/R_SDD01_001_002"
    doc.add_heading("OldFunc（D/R_SDD01_001_001）", level=4)
    doc.add_paragraph("a) 函数原型")
    doc.add_paragraph("void OldFunc(void);")
    doc.add_heading("KeepFunc（D/R_SDD01_001_002）", level=4)
    doc.add_paragraph("a) 函数原型")
    doc.add_paragraph("void KeepFunc(void);")
    doc.save(path)


def _doc_texts(path):
    doc = Document(path)
    return [(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()]


def _module_table_csu_ids(path):
    doc = Document(path)
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = [(cell.text or "").strip() for cell in table.rows[0].cells[:4]]
        if header != ["CSC 名称", "CSC 标识", "CSU 名称", "CSU 标识"]:
            continue
        return [(row.cells[2].text or "").strip() + ":" + (row.cells[3].text or "").strip() for row in table.rows[1:]]
    return []


def test_insert_csu_after_in_doc_inserts_complete_block_before_next_csu(tmp_path):
    path = tmp_path / "target.docx"
    _write_two_csu_doc(path)

    doc = Document(path)
    new_heading = doc.add_heading("NewFunc（D/R_SDD01_001_003）", level=4)
    new_proto_label = doc.add_paragraph("a) 函数原型")
    new_proto = doc.add_paragraph("void NewFunc(void);")
    body = doc._body._element
    new_elements = [new_heading._element, new_proto_label._element, new_proto._element]
    for elem in new_elements:
        body.remove(elem)

    result = insert_csu_after_in_doc(doc, "D/R_SDD01_001_001", new_elements)
    sync_result = sync_module_function_table_for_module(doc, "D/R_SDD01_001")
    doc.save(path)

    assert result["found"] is True
    assert sync_result["updated"] is True
    texts = _doc_texts(path)
    assert texts.index("OldFunc（D/R_SDD01_001_001）") < texts.index("NewFunc（D/R_SDD01_001_003）")
    assert texts.index("NewFunc（D/R_SDD01_001_003）") < texts.index("KeepFunc（D/R_SDD01_001_002）")
    assert "void NewFunc(void);" in texts
    assert any("D/R_SDD01_001_003" in text for text in _module_table_csu_ids(path))


def test_allocate_next_csu_id_uses_same_module_max_suffix(tmp_path):
    path = tmp_path / "target.docx"
    doc = Document()
    doc.add_heading("Func1（D/R_SDD01_001_001）", level=4)
    doc.add_heading("Func4（D/R_SDD01_001_004）", level=4)
    doc.add_heading("Other1（D/R_SDD01_002_001）", level=4)
    doc.save(path)

    assert allocate_next_csu_id(str(path), "D/R_SDD01_001_001") == "D/R_SDD01_001_005"


def test_build_csu_index_ignores_module_h3_ids(tmp_path):
    path = tmp_path / "target.docx"
    _write_two_csu_doc(path)

    index = build_csu_index(str(path))

    assert [entry["csu_id"] for entry in index["OldFunc"]] == ["D/R_SDD01_001_001"]
    assert "D/R_SDD01_001" not in [entry["csu_id"] for entries in index.values() for entry in entries]


def test_build_doc_code_alignment_index_matches_unique_code_function(tmp_path):
    doc_path = tmp_path / "target.docx"
    _write_two_csu_doc(doc_path)
    code_root = tmp_path / "code"
    code_root.mkdir()
    (code_root / "demo.c").write_text(
        "void OldFunc(void)\n{\n}\n\nvoid KeepFunc(void)\n{\n}\n",
        encoding="utf-8",
    )

    alignment = build_doc_code_alignment_index(str(doc_path), str(code_root))

    old = alignment["D/R_SDD01_001_001"]
    assert old["status"] == "matched_high"
    assert old["matched_function"] == "OldFunc"
    assert old["rel_path"] == "demo.c"
    assert old["confidence"] >= 0.9
    assert old["evidence"]["function_name"] == "exact_unique"


def test_build_doc_code_alignment_index_marks_duplicate_code_function_ambiguous(tmp_path):
    doc_path = tmp_path / "target.docx"
    _write_two_csu_doc(doc_path)
    code_root = tmp_path / "code"
    code_root.mkdir()
    (code_root / "a.c").write_text("void OldFunc(void)\n{\n}\n", encoding="utf-8")
    (code_root / "b.c").write_text("void OldFunc(void)\n{\n}\n", encoding="utf-8")

    alignment = build_doc_code_alignment_index(str(doc_path), str(code_root))

    old = alignment["D/R_SDD01_001_001"]
    assert old["status"] == "ambiguous"
    assert old["confidence"] == 0.0
    assert old["evidence"]["candidate_count"] == 2


def test_attach_alignment_to_items_adds_plan_evidence(tmp_path):
    doc_path = tmp_path / "target.docx"
    _write_two_csu_doc(doc_path)
    code_root = tmp_path / "code"
    code_root.mkdir()
    (code_root / "demo.c").write_text("void OldFunc(void)\n{\n}\n", encoding="utf-8")
    item = PlannedItem(
        action="modified_function",
        status="safe",
        rel_path="demo.c",
        func_name="OldFunc",
        csu_id="D/R_SDD01_001_001",
    )

    attach_alignment_to_items([item], build_doc_code_alignment_index(str(doc_path), str(code_root)))

    payload = item.to_dict()
    assert payload["alignment"]["status"] == "matched_high"
    assert payload["alignment"]["rel_path"] == "demo.c"


def test_load_alignment_decisions_from_review_export_json(tmp_path):
    decisions_path = tmp_path / "review_decisions.json"
    decisions_path.write_text(
        json.dumps({
            "schema_version": 1,
            "decisions": [{"item_index": 0, "decision": "skip"}],
            "alignment_decisions": [{
                "alignment_index": 3,
                "csu_id": "D/R_SDD01_001_001",
                "doc_func_name": "OldFunc",
                "manual_function": "NewFunc",
                "manual_rel_path": "demo.c",
                "status": "unmatched",
                "notes": "人工确认改名",
            }],
        }, ensure_ascii=False),
        encoding="utf-8",
    )

    loaded = updater._load_alignment_decisions(str(decisions_path))

    assert len(loaded) == 1
    assert loaded[0]["manual_function"] == "NewFunc"
    assert loaded[0]["manual_rel_path"] == "demo.c"


def test_apply_alignment_decisions_overrides_unmatched_and_csu_index(tmp_path):
    doc_path = tmp_path / "target.docx"
    _write_two_csu_doc(doc_path)
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    old_code.mkdir()
    new_code.mkdir()
    (new_code / "demo.c").write_text("void NewFunc(void)\n{\n}\n", encoding="utf-8")

    alignment = build_doc_code_alignment_index(str(doc_path), str(old_code))
    csu_index = build_csu_index(str(doc_path))

    updater.apply_alignment_decisions(
        alignment,
        csu_index,
        [{
            "csu_id": "D/R_SDD01_001_001",
            "manual_function": "NewFunc",
            "manual_rel_path": "demo.c",
            "notes": "人工确认改名",
        }],
        code_roots=[str(old_code), str(new_code)],
    )

    mapped = alignment["D/R_SDD01_001_001"]
    assert mapped["status"] == "manual_matched"
    assert mapped["matched_function"] == "NewFunc"
    assert mapped["rel_path"] == "demo.c"
    assert mapped["evidence"] == {
        "source": "alignment_decisions",
        "manual_function": "NewFunc",
        "manual_rel_path": "demo.c",
        "previous_status": "unmatched",
        "notes": "人工确认改名",
    }
    assert [entry["csu_id"] for entry in csu_index["NewFunc"]] == ["D/R_SDD01_001_001"]


def test_manual_alignment_allows_new_function_to_target_existing_csu(tmp_path):
    doc_path = tmp_path / "target.docx"
    _write_two_csu_doc(doc_path)
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    old_code.mkdir()
    new_code.mkdir()
    (new_code / "demo.c").write_text("void NewFunc(void)\n{\n}\n", encoding="utf-8")
    csu_index = build_csu_index(str(doc_path))
    alignment = build_doc_code_alignment_index(str(doc_path), str(old_code))
    updater.apply_alignment_decisions(
        alignment,
        csu_index,
        [{
            "csu_id": "D/R_SDD01_001_001",
            "manual_function": "NewFunc",
            "manual_rel_path": "demo.c",
        }],
        code_roots=[str(old_code), str(new_code)],
    )

    items = classify_changes(
        [{
            "change_kind": "new_function",
            "language": "c",
            "key": "demo.c",
            "function_name": "NewFunc",
            "new_function_name": "NewFunc",
        }],
        new_code=str(new_code),
        csu_index=csu_index,
    )
    attach_alignment_to_items(items, alignment)

    assert len(items) == 1
    assert items[0].action == "new_function"
    assert items[0].status == "review"
    assert items[0].csu_id == "D/R_SDD01_001_001"
    assert items[0].alignment["status"] == "manual_matched"


def test_custom_heading4_base_style_is_indexed_and_deletable(tmp_path):
    path = tmp_path / "target.docx"
    doc = Document()
    module_style = doc.styles.add_style("ModuleCustom", WD_STYLE_TYPE.PARAGRAPH)
    module_style.base_style = doc.styles["Heading 3"]
    csu_style = doc.styles.add_style("CsuCustom", WD_STYLE_TYPE.PARAGRAPH)
    csu_style.base_style = doc.styles["Heading 4"]
    doc.add_paragraph("DemoModule（D/R_SDD01_001）", style=module_style)
    doc.add_paragraph("OldFunc（D/R_SDD01_001_001）", style=csu_style)
    doc.add_paragraph("a) 函数原型")
    doc.add_paragraph("void OldFunc(void);")
    doc.add_paragraph("KeepFunc（D/R_SDD01_001_002）", style=csu_style)
    doc.add_paragraph("a) 函数原型")
    doc.add_paragraph("void KeepFunc(void);")
    doc.save(path)

    index = build_csu_index(str(path))
    assert [entry["csu_id"] for entry in index["OldFunc"]] == ["D/R_SDD01_001_001"]
    assert [entry["csu_id"] for entry in index["KeepFunc"]] == ["D/R_SDD01_001_002"]

    doc = Document(path)
    result = delete_csu_in_doc(doc, "D/R_SDD01_001_001")
    doc.save(path)

    assert result["found"] is True
    texts = _doc_texts(path)
    assert not any("OldFunc" in text for text in texts)
    assert any("KeepFunc" in text for text in texts)


def test_renumber_module_csu_ids_updates_h4_order_and_module_table(tmp_path):
    path = tmp_path / "target.docx"
    _write_two_csu_doc(path)

    doc = Document(path)
    new_heading = doc.add_heading("NewFunc（D/R_SDD01_001_003）", level=4)
    new_proto = doc.add_paragraph("void NewFunc(void);")
    body = doc._body._element
    new_elements = [new_heading._element, new_proto._element]
    for elem in new_elements:
        body.remove(elem)
    insert_csu_after_in_doc(doc, "D/R_SDD01_001_001", new_elements)

    result = renumber_module_csu_ids(doc, "D/R_SDD01_001")
    doc.save(path)

    assert result["found"] is True
    assert result["updated"] == 2
    assert result["mapping"] == [
        {"old_csu_id": "D/R_SDD01_001_001", "new_csu_id": "D/R_SDD01_001_001", "title": "OldFunc"},
        {"old_csu_id": "D/R_SDD01_001_003", "new_csu_id": "D/R_SDD01_001_002", "title": "NewFunc"},
        {"old_csu_id": "D/R_SDD01_001_002", "new_csu_id": "D/R_SDD01_001_003", "title": "KeepFunc"},
    ]
    texts = _doc_texts(path)
    assert "OldFunc（D/R_SDD01_001_001）" in texts
    assert "NewFunc（D/R_SDD01_001_002）" in texts
    assert "KeepFunc（D/R_SDD01_001_003）" in texts
    table_ids = _module_table_csu_ids(path)
    assert table_ids == [
        "OldFunc:D/R_SDD01_001_001",
        "NewFunc:D/R_SDD01_001_002",
        "KeepFunc:D/R_SDD01_001_003",
    ]


def test_classify_changes_detects_possible_function_rename(tmp_path):
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "demo.c").write_text(
        "void NewFunc(void)\\n{\\n    int status = 1;\\n    (void)status;\\n}\\n",
        encoding="utf-8",
    )
    changes = [
        {
            "change_kind": "deleted_function",
            "language": "c",
            "key": "demo.c",
            "old_function_name": "OldFunc",
            "old_text": "void OldFunc(void)\\n{\\n    int status = 1;\\n    (void)status;\\n}\\n",
        },
        {
            "change_kind": "new_function",
            "language": "c",
            "key": "demo.c",
            "new_function_name": "NewFunc",
            "new_text": "void NewFunc(void)\\n{\\n    int status = 1;\\n    (void)status;\\n}\\n",
        },
    ]

    items = classify_changes(
        changes,
        new_code=str(new_code),
        csu_index={"OldFunc": [{"csu_id": "D/R_SDD01_001_001", "title": "OldFunc", "heading": ""}]},
    )

    assert len(items) == 1
    assert items[0].action == "renamed_function"
    assert items[0].status == "review"
    assert items[0].func_name == "NewFunc"
    assert items[0].csu_id == "D/R_SDD01_001_001"
    assert items[0].change["old_function_name"] == "OldFunc"
    assert items[0].change["new_function_name"] == "NewFunc"
    assert items[0].change["rename_similarity"] >= 0.78


def test_classify_changes_adds_header_impacted_functions(tmp_path):
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "demo.h").write_text("#define DEMO_VALUE 1\n", encoding="utf-8")
    (new_code / "demo.c").write_text(
        '#include "demo.h"\n\n'
        "void HeaderUser(void)\n"
        "{\n"
        "    int status = DEMO_VALUE;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )

    items = classify_changes(
        [{
            "change_kind": "header_changed",
            "language": "header",
            "key": "demo.h",
        }],
        new_code=str(new_code),
        csu_index={"HeaderUser": [{"csu_id": "D/R_SDD01_001_001", "title": "HeaderUser", "heading": ""}]},
    )

    impacted = [item for item in items if item.action == "header_impacted_function"]
    assert len(impacted) == 1
    assert impacted[0].status == "review"
    assert impacted[0].func_name == "HeaderUser"
    assert impacted[0].rel_path == "demo.c"
    assert impacted[0].csu_id == "D/R_SDD01_001_001"
    assert impacted[0].change["header_file"] == "demo.h"


def test_classify_changes_adds_transitive_header_impacted_functions(tmp_path):
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "demo.h").write_text("#define DEMO_VALUE 1\n", encoding="utf-8")
    (new_code / "wrapper.h").write_text('#include "demo.h"\n', encoding="utf-8")
    (new_code / "demo.c").write_text(
        '#include "wrapper.h"\n\n'
        "void HeaderUser(void)\n"
        "{\n"
        "    int status = DEMO_VALUE;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )

    items = classify_changes(
        [{
            "change_kind": "header_changed",
            "language": "header",
            "key": "demo.h",
        }],
        new_code=str(new_code),
        csu_index={"HeaderUser": [{"csu_id": "D/R_SDD01_001_001", "title": "HeaderUser", "heading": ""}]},
    )

    impacted = [item for item in items if item.action == "header_impacted_function"]
    assert len(impacted) == 1
    assert impacted[0].status == "review"
    assert impacted[0].func_name == "HeaderUser"
    assert impacted[0].change["header_file"] == "demo.h"
    assert impacted[0].change["matched_header_file"] == "wrapper.h"
    assert "wrapper.h" in impacted[0].change["impacted_headers"]


def test_classify_changes_dedupes_header_impacts_for_same_function(tmp_path):
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "a.h").write_text("#define A_VALUE 1\n", encoding="utf-8")
    (new_code / "b.h").write_text("#define B_VALUE 1\n", encoding="utf-8")
    (new_code / "demo.c").write_text(
        '#include "a.h"\n'
        '#include "b.h"\n\n'
        "void HeaderUser(void)\n"
        "{\n"
        "    int status = A_VALUE + B_VALUE;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )

    items = classify_changes(
        [
            {"change_kind": "header_changed", "language": "header", "key": "a.h"},
            {"change_kind": "header_changed", "language": "header", "key": "b.h"},
        ],
        new_code=str(new_code),
        csu_index={"HeaderUser": [{"csu_id": "D/R_SDD01_001_001", "title": "HeaderUser", "heading": ""}]},
    )

    impacted = [item for item in items if item.action == "header_impacted_function"]
    assert len(impacted) == 1
    assert impacted[0].func_name == "HeaderUser"
    assert impacted[0].change["impacted_by_headers"] == ["a.h", "b.h"]
    assert impacted[0].reason == "function impacted by 2 changed headers"


def test_classify_changes_expands_file_level_new_c_file(tmp_path):
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "added.c").write_text(
        "void AddedOne(void)\n"
        "{\n"
        "}\n\n"
        "static void AddedTwo(void)\n"
        "{\n"
        "}\n",
        encoding="utf-8",
    )

    items = classify_changes(
        [{
            "change_kind": "new_file",
            "language": "c",
            "key": "added.c",
            "new_text": "",
        }],
        new_code=str(new_code),
        csu_index={},
    )

    assert [(item.action, item.status, item.func_name) for item in items] == [
        ("new_function", "review", "AddedOne"),
        ("new_function", "review", "AddedTwo"),
    ]


def test_classify_changes_expands_file_level_deleted_c_file(tmp_path):
    new_code = tmp_path / "new"
    new_code.mkdir()

    items = classify_changes(
        [{
            "change_kind": "deleted_file",
            "language": "c",
            "key": "removed.c",
            "old_text": (
                "void RemovedOne(void)\n"
                "{\n"
                "}\n\n"
                "void RemovedTwo(void)\n"
                "{\n"
                "}\n"
            ),
        }],
        new_code=str(new_code),
        csu_index={
            "RemovedOne": [{"csu_id": "D/R_SDD01_001_001", "title": "RemovedOne", "heading": ""}],
        },
    )

    assert [(item.action, item.status, item.func_name, item.csu_id) for item in items] == [
        ("deleted_function", "review", "RemovedOne", "D/R_SDD01_001_001"),
        ("deleted_function", "review", "RemovedTwo", ""),
    ]
    assert items[1].reason == "deleted function has no CSU match"


def test_doc_update_worker_runs_plan_only_with_gui_callbacks(tmp_path, monkeypatch):
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    old_code.mkdir()
    new_code.mkdir()
    old_doc = tmp_path / "old.docx"
    Document().save(old_doc)
    out_doc = tmp_path / "out.docx"

    from tools import update_doc_from_code_diff as updater

    monkeypatch.setattr(updater, "_run_docdiff", lambda **kwargs: None)
    monkeypatch.setattr(updater, "_load_changes", lambda path: [])
    monkeypatch.setattr(updater, "build_csu_index", lambda path: {})
    monkeypatch.setattr(updater, "classify_changes", lambda changes, **kwargs: [])
    monkeypatch.setattr(updater, "write_reports", lambda **kwargs: None)
    monkeypatch.setattr(updater, "write_review_html", lambda **kwargs: None)

    task = TaskSpec(
        mode="doc_update",
        c_file="",
        project_dir="",
        output=str(out_doc),
        template_path="",
        old_code=str(old_code),
        new_code=str(new_code),
        old_doc=str(old_doc),
        doc_update_mode="plan-only",
    )
    worker = DocUpdateWorker(backend=None, task=task, settings=AppSettings())
    steps = []
    logs = []
    outputs = []
    done = []

    worker.run(
        emit_step=lambda step, status: steps.append((step, status)),
        emit_log=logs.append,
        emit_output=outputs.append,
        emit_done=lambda note, resume, output: done.append((note, resume, output)),
    )

    assert ("report", "success") in steps
    assert outputs == [str(out_doc)]
    assert done and done[0][0] == "文档增量更新完成"


def test_render_review_html_supports_persistence_and_import(tmp_path):
    out_html = tmp_path / "review.html"
    render_review_html(
        {
            "schema_version": 1,
            "metadata": {"old_doc": "old.docx", "new_code": "new"},
            "items": [{
                "action": "new_function",
                "status": "review",
                "rel_path": "demo.c",
                "func_name": "NewFunc",
                "csu_id": "",
                "reason": "new function insertion needs section position and CSU ID",
                "change": {"new_text": "void NewFunc(void) {}"},
                "result": {},
            }],
            "alignment_items": [{
                "csu_id": "D/R_SDD01_001_001",
                "doc_title": "OldFunc",
                "doc_func_name": "OldFunc",
                "status": "manual_matched",
                "matched_function": "NewFunc",
                "rel_path": "demo.c",
                "confidence": 1.0,
                "evidence": {"source": "alignment_decisions"},
            }],
        },
        str(out_html),
    )

    html = out_html.read_text(encoding="utf-8")
    assert "localStorage.setItem(storageKey" in html
    assert "loadPersistedDecisions()" in html
    assert "importDecisionPayload" in html
    assert 'id="importBtn"' in html
    assert 'id="importFile"' in html
    assert 'id="viewMode"' in html
    assert "CSU 映射" in html
    assert "alignment_decisions" in html
    assert "manual_function" in html
    assert "manual_matched" in html
    assert "人工映射已应用" in html
    assert "decisions[index].csu_id" in html


def test_update_doc_cli_apply_review_end_to_end(tmp_path, monkeypatch):
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    docdiff_root = tmp_path / "docdiff"
    old_code.mkdir()
    new_code.mkdir()
    docdiff_root.mkdir()
    out_doc = tmp_path / "out.docx"
    old_doc = tmp_path / "old.docx"
    decisions_path = tmp_path / "review_decisions.json"

    _write_two_csu_doc(old_doc)
    (old_code / "demo.c").write_text(
        "void OldFunc(void)\n"
        "{\n"
        "}\n",
        encoding="utf-8",
    )
    (new_code / "demo.c").write_text(
        "void OldFunc(void)\n"
        "{\n"
        "    int status = 2;\n"
        "    (void)status;\n"
        "}\n\n"
        "void NewFunc(void)\n"
        "{\n"
        "    int status = 1;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )
    decisions_path.write_text(
        json.dumps({
            "schema_version": 1,
            "decisions": [{
                "item_index": 1,
                "decision": "insert_after_csu",
                "insert_after_csu_id": "D/R_SDD01_001_001",
                "func_name": "NewFunc",
                "rel_path": "demo.c",
            }],
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    (docdiff_root / "cli.py").write_text(
        "import argparse, json, pathlib\n"
        "parser = argparse.ArgumentParser()\n"
        "parser.add_argument('--mode')\n"
        "parser.add_argument('--old')\n"
        "parser.add_argument('--new')\n"
        "parser.add_argument('--out')\n"
        "parser.add_argument('--json-out')\n"
        "args = parser.parse_args()\n"
        "pathlib.Path(args.out).write_bytes(b'fake code change docx')\n"
        "changes = [\n"
        "  {'change_kind': 'modified_function', 'language': 'c', 'key': 'demo.c', 'function_name': 'OldFunc', 'old_function_name': 'OldFunc', 'new_function_name': 'OldFunc', 'old_signature': 'void OldFunc(void)', 'new_signature': 'void OldFunc(void)', 'old_text': '', 'new_text': ''},\n"
        "  {'change_kind': 'new_function', 'language': 'c', 'key': 'demo.c', 'function_name': 'NewFunc', 'old_function_name': '', 'new_function_name': 'NewFunc', 'old_signature': '', 'new_signature': 'void NewFunc(void)', 'old_text': '', 'new_text': ''},\n"
        "]\n"
        "pathlib.Path(args.json_out).write_text(json.dumps(changes, ensure_ascii=False), encoding='utf-8')\n",
        encoding="utf-8",
    )

    monkeypatch.setattr(sys, "argv", [
        "update_doc_from_code_diff.py",
        "--old-code", str(old_code),
        "--new-code", str(new_code),
        "--old-doc", str(old_doc),
        "--out", str(out_doc),
        "--docdiff-root", str(docdiff_root),
        "--mode", "apply-review",
        "--review-decisions", str(decisions_path),
    ])

    assert updater.main() == 0

    plan_path = tmp_path / "out.update_plan.json"
    report_path = tmp_path / "out.update_report.md"
    review_html = tmp_path / "out.update_review.html"
    assert out_doc.exists()
    assert plan_path.exists()
    assert report_path.exists()
    assert review_html.exists()
    plan = json.loads(plan_path.read_text(encoding="utf-8"))
    assert plan["summary"] == {"applied": 1, "applied_review": 1}
    assert plan["metadata"]["mode"] == "apply-review"
    assert plan["metadata"]["review_decisions"] == str(decisions_path)
    assert plan["metadata"]["alignment"]["total_csu"] == 2
    assert any(item["csu_id"] == "D/R_SDD01_001_001" and item["status"] == "matched_high" for item in plan["alignment_items"])
    texts = _doc_texts(out_doc)
    joined = "\n".join(texts)
    assert "void OldFunc(void)" in texts
    assert "void NewFunc(void)" in texts
    assert "设置状态 = 2；" in joined
    assert "设置状态 = 1；" in joined
    old_heading = next(text for text in texts if text.endswith("D/R_SDD01_001_001）"))
    inserted_heading = next(text for text in texts if text.endswith("D/R_SDD01_001_003）"))
    assert texts.index(old_heading) < texts.index(inserted_heading)
    assert texts.index(inserted_heading) < texts.index("KeepFunc（D/R_SDD01_001_002）")
    assert any("D/R_SDD01_001_003" in text for text in _module_table_csu_ids(out_doc))


def test_update_doc_cli_replays_alignment_decisions_into_plan(tmp_path, monkeypatch):
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    docdiff_root = tmp_path / "docdiff"
    old_code.mkdir()
    new_code.mkdir()
    docdiff_root.mkdir()
    old_doc = tmp_path / "old.docx"
    out_doc = tmp_path / "out.docx"
    decisions_path = tmp_path / "review_decisions.json"
    _write_two_csu_doc(old_doc)
    (new_code / "demo.c").write_text("void NewFunc(void)\n{\n}\n", encoding="utf-8")
    decisions_path.write_text(
        json.dumps({
            "schema_version": 1,
            "alignment_decisions": [{
                "csu_id": "D/R_SDD01_001_001",
                "doc_title": "OldFunc",
                "doc_func_name": "OldFunc",
                "manual_function": "NewFunc",
                "manual_rel_path": "demo.c",
                "status": "unmatched",
                "notes": "人工确认 OldFunc 对应 NewFunc",
            }],
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    (docdiff_root / "cli.py").write_text(
        "import argparse, json, pathlib\n"
        "parser = argparse.ArgumentParser()\n"
        "parser.add_argument('--mode')\n"
        "parser.add_argument('--old')\n"
        "parser.add_argument('--new')\n"
        "parser.add_argument('--out')\n"
        "parser.add_argument('--json-out')\n"
        "args = parser.parse_args()\n"
        "pathlib.Path(args.out).write_bytes(b'fake code change docx')\n"
        "changes = [{'change_kind': 'new_function', 'language': 'c', 'key': 'demo.c', 'function_name': 'NewFunc', 'new_function_name': 'NewFunc'}]\n"
        "pathlib.Path(args.json_out).write_text(json.dumps(changes, ensure_ascii=False), encoding='utf-8')\n",
        encoding="utf-8",
    )

    monkeypatch.setattr(sys, "argv", [
        "update_doc_from_code_diff.py",
        "--old-code", str(old_code),
        "--new-code", str(new_code),
        "--old-doc", str(old_doc),
        "--out", str(out_doc),
        "--docdiff-root", str(docdiff_root),
        "--mode", "plan-only",
        "--alignment-decisions", str(decisions_path),
    ])

    assert updater.main() == 0

    plan = json.loads((tmp_path / "out.update_plan.json").read_text(encoding="utf-8"))
    assert plan["metadata"]["alignment"]["manual_matched"] == 1
    assert any(item["csu_id"] == "D/R_SDD01_001_001" and item["status"] == "manual_matched" for item in plan["alignment_items"])
    update_item = plan["items"][0]
    assert update_item["func_name"] == "NewFunc"
    assert update_item["csu_id"] == "D/R_SDD01_001_001"
    assert update_item["alignment"]["status"] == "manual_matched"


def test_update_doc_cli_apply_review_uses_alignment_decisions_from_review_export(tmp_path, monkeypatch):
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    docdiff_root = tmp_path / "docdiff"
    old_code.mkdir()
    new_code.mkdir()
    docdiff_root.mkdir()
    old_doc = tmp_path / "old.docx"
    out_doc = tmp_path / "out.docx"
    decisions_path = tmp_path / "review_decisions.json"
    _write_two_csu_doc(old_doc)
    (new_code / "demo.c").write_text(
        "void NewFunc(void)\n"
        "{\n"
        "    int status = 7;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )
    decisions_path.write_text(
        json.dumps({
            "schema_version": 1,
            "alignment_decisions": [{
                "csu_id": "D/R_SDD01_001_001",
                "doc_title": "OldFunc",
                "doc_func_name": "OldFunc",
                "manual_function": "NewFunc",
                "manual_rel_path": "demo.c",
                "status": "unmatched",
                "notes": "人工确认 OldFunc 对应 NewFunc",
            }],
            "decisions": [{
                "item_index": 0,
                "decision": "replace_csu",
                "target_csu_id": "D/R_SDD01_001_001",
                "func_name": "NewFunc",
                "rel_path": "demo.c",
            }],
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    (docdiff_root / "cli.py").write_text(
        "import argparse, json, pathlib\n"
        "parser = argparse.ArgumentParser()\n"
        "parser.add_argument('--mode')\n"
        "parser.add_argument('--old')\n"
        "parser.add_argument('--new')\n"
        "parser.add_argument('--out')\n"
        "parser.add_argument('--json-out')\n"
        "args = parser.parse_args()\n"
        "pathlib.Path(args.out).write_bytes(b'fake code change docx')\n"
        "changes = [{'change_kind': 'new_function', 'language': 'c', 'key': 'demo.c', 'function_name': 'NewFunc', 'new_function_name': 'NewFunc'}]\n"
        "pathlib.Path(args.json_out).write_text(json.dumps(changes, ensure_ascii=False), encoding='utf-8')\n",
        encoding="utf-8",
    )

    monkeypatch.setattr(sys, "argv", [
        "update_doc_from_code_diff.py",
        "--old-code", str(old_code),
        "--new-code", str(new_code),
        "--old-doc", str(old_doc),
        "--out", str(out_doc),
        "--docdiff-root", str(docdiff_root),
        "--mode", "apply-review",
        "--review-decisions", str(decisions_path),
    ])

    assert updater.main() == 0

    plan = json.loads((tmp_path / "out.update_plan.json").read_text(encoding="utf-8"))
    assert plan["metadata"]["alignment"]["manual_matched"] == 1
    item = plan["items"][0]
    assert item["status"] == "applied_review"
    assert item["csu_id"] == "D/R_SDD01_001_001"
    assert item["alignment"]["status"] == "manual_matched"
    assert item["result"]["review_result"]["found"] is True
    texts = _doc_texts(out_doc)
    joined = "\n".join(texts)
    assert "OldFunc（D/R_SDD01_001_001）" not in texts
    assert any(text.endswith("D/R_SDD01_001_001）") for text in texts)
    assert "void NewFunc(void)" in texts
    assert "设置状态 = 7；" in joined


def test_update_doc_cli_apply_review_can_renumber_module_csu(tmp_path, monkeypatch):
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    docdiff_root = tmp_path / "docdiff"
    old_code.mkdir()
    new_code.mkdir()
    docdiff_root.mkdir()
    out_doc = tmp_path / "out_renumber.docx"
    old_doc = tmp_path / "old.docx"
    decisions_path = tmp_path / "review_decisions.json"

    _write_two_csu_doc(old_doc)
    (new_code / "demo.c").write_text(
        "void NewFunc(void)\n"
        "{\n"
        "    int status = 1;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )
    decisions_path.write_text(
        json.dumps({
            "schema_version": 1,
            "decisions": [{
                "item_index": 0,
                "decision": "insert_after_csu",
                "insert_after_csu_id": "D/R_SDD01_001_001",
                "func_name": "NewFunc",
                "rel_path": "demo.c",
            }],
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    (docdiff_root / "cli.py").write_text(
        "import argparse, json, pathlib\n"
        "parser = argparse.ArgumentParser()\n"
        "parser.add_argument('--mode')\n"
        "parser.add_argument('--old')\n"
        "parser.add_argument('--new')\n"
        "parser.add_argument('--out')\n"
        "parser.add_argument('--json-out')\n"
        "args = parser.parse_args()\n"
        "pathlib.Path(args.out).write_bytes(b'fake code change docx')\n"
        "changes = [\n"
        "  {'change_kind': 'new_function', 'language': 'c', 'key': 'demo.c', 'function_name': 'NewFunc', 'old_function_name': '', 'new_function_name': 'NewFunc'},\n"
        "]\n"
        "pathlib.Path(args.json_out).write_text(json.dumps(changes, ensure_ascii=False), encoding='utf-8')\n",
        encoding="utf-8",
    )

    monkeypatch.setattr(sys, "argv", [
        "update_doc_from_code_diff.py",
        "--old-code", str(old_code),
        "--new-code", str(new_code),
        "--old-doc", str(old_doc),
        "--out", str(out_doc),
        "--docdiff-root", str(docdiff_root),
        "--mode", "apply-review",
        "--review-decisions", str(decisions_path),
        "--renumber-module-csu",
    ])

    assert updater.main() == 0

    plan = json.loads((tmp_path / "out_renumber.update_plan.json").read_text(encoding="utf-8"))
    assert plan["summary"] == {"applied_review": 1}
    assert plan["metadata"]["renumber_module_csu"] is True
    item = plan["items"][0]
    assert item["csu_id"] == "D/R_SDD01_001_002"
    assert item["result"]["review_result"]["auto_allocated_csu_id"] == "D/R_SDD01_001_003"
    assert item["result"]["review_result"]["renumber_module_csu"]["updated"] == 2

    texts = _doc_texts(out_doc)
    assert texts.index("OldFunc（D/R_SDD01_001_001）") < texts.index("新功能（D/R_SDD01_001_002）")
    assert texts.index("新功能（D/R_SDD01_001_002）") < texts.index("KeepFunc（D/R_SDD01_001_003）")
    assert _module_table_csu_ids(out_doc) == [
        "OldFunc:D/R_SDD01_001_001",
        "新功能:D/R_SDD01_001_002",
        "KeepFunc:D/R_SDD01_001_003",
    ]


def test_doc_update_worker_runs_apply_review_end_to_end(tmp_path):
    old_code = tmp_path / "old"
    new_code = tmp_path / "new"
    docdiff_root = tmp_path / "docdiff"
    old_code.mkdir()
    new_code.mkdir()
    docdiff_root.mkdir()
    out_doc = tmp_path / "gui_out.docx"
    old_doc = tmp_path / "old.docx"
    decisions_path = tmp_path / "review_decisions.json"

    _write_two_csu_doc(old_doc)
    (new_code / "demo.c").write_text(
        "void OldFunc(void)\n"
        "{\n"
        "    int status = 2;\n"
        "    (void)status;\n"
        "}\n\n"
        "void NewFunc(void)\n"
        "{\n"
        "    int status = 1;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )
    decisions_path.write_text(
        json.dumps({
            "schema_version": 1,
            "decisions": [{
                "item_index": 1,
                "decision": "insert_after_csu",
                "insert_after_csu_id": "D/R_SDD01_001_001",
                "func_name": "NewFunc",
                "rel_path": "demo.c",
            }],
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    (docdiff_root / "cli.py").write_text(
        "import argparse, json, pathlib\n"
        "parser = argparse.ArgumentParser()\n"
        "parser.add_argument('--mode')\n"
        "parser.add_argument('--old')\n"
        "parser.add_argument('--new')\n"
        "parser.add_argument('--out')\n"
        "parser.add_argument('--json-out')\n"
        "args = parser.parse_args()\n"
        "pathlib.Path(args.out).write_bytes(b'fake code change docx')\n"
        "changes = [\n"
        "  {'change_kind': 'modified_function', 'language': 'c', 'key': 'demo.c', 'function_name': 'OldFunc', 'old_function_name': 'OldFunc', 'new_function_name': 'OldFunc'},\n"
        "  {'change_kind': 'new_function', 'language': 'c', 'key': 'demo.c', 'function_name': 'NewFunc', 'old_function_name': '', 'new_function_name': 'NewFunc'},\n"
        "]\n"
        "pathlib.Path(args.json_out).write_text(json.dumps(changes, ensure_ascii=False), encoding='utf-8')\n",
        encoding="utf-8",
    )

    task = TaskSpec(
        mode="doc_update",
        c_file="",
        project_dir="",
        output=str(out_doc),
        template_path="",
        old_code=str(old_code),
        new_code=str(new_code),
        old_doc=str(old_doc),
        review_decisions=str(decisions_path),
        doc_update_mode="apply-review",
        docdiff_root=str(docdiff_root),
    )
    worker = DocUpdateWorker(backend=None, task=task, settings=AppSettings())
    steps = []
    logs = []
    outputs = []
    done = []

    worker.run(
        emit_step=lambda step, status: steps.append((step, status)),
        emit_log=logs.append,
        emit_output=outputs.append,
        emit_done=lambda note, resume, output: done.append((note, resume, output)),
    )

    assert ("apply", "success") in steps
    assert ("report", "success") in steps
    assert outputs == [str(out_doc)]
    assert done and done[0][0] == "文档增量更新完成"
    plan = json.loads((tmp_path / "gui_out.update_plan.json").read_text(encoding="utf-8"))
    assert plan["summary"] == {"applied": 1, "applied_review": 1}
    assert plan["metadata"]["docdiff_root"] == str(docdiff_root)
    assert (tmp_path / "gui_out.update_review.html").exists()
    texts = _doc_texts(out_doc)
    joined = "\n".join(texts)
    assert "void OldFunc(void)" in texts
    assert "void NewFunc(void)" in texts
    assert "设置状态 = 2；" in joined
    assert "设置状态 = 1；" in joined


def test_delete_csu_in_doc_removes_target_block_only(tmp_path):
    path = tmp_path / "target.docx"
    _write_two_csu_doc(path)

    doc = Document(path)
    result = delete_csu_in_doc(doc, "D/R_SDD01_001_001")
    doc.save(path)

    assert result["found"] is True
    texts = _doc_texts(path)
    assert not any("OldFunc" in text for text in texts)
    assert not any("void OldFunc" in text for text in texts)
    assert any("KeepFunc" in text for text in texts)
    assert any("void KeepFunc" in text for text in texts)


def test_apply_review_delete_csu_decision_updates_doc_and_plan(tmp_path):
    out_doc = tmp_path / "reviewed.docx"
    new_code = tmp_path / "new"
    new_code.mkdir()
    _write_two_csu_doc(out_doc)

    item = PlannedItem(
        action="deleted_function",
        status="review",
        rel_path="demo.c",
        func_name="OldFunc",
        csu_id="D/R_SDD01_001_001",
        reason="deleted function removal needs confirmation",
    )

    apply_review_decisions(
        [item],
        out_doc=str(out_doc),
        new_code=str(new_code),
        review_decisions=[{
            "item_index": 0,
            "decision": "delete_csu",
            "target_csu_id": "D/R_SDD01_001_001",
            "func_name": "OldFunc",
            "rel_path": "demo.c",
        }],
        ai_assist=False,
        template_path="",
    )

    assert item.status == "applied_review", item.reason
    assert item.result["review_result"]["found"] is True
    assert item.result["review_result"]["module_table"]["updated"] is True
    texts = _doc_texts(out_doc)
    assert not any("OldFunc" in text for text in texts)
    assert any("KeepFunc" in text for text in texts)
    table_ids = _module_table_csu_ids(out_doc)
    assert not any("D/R_SDD01_001_001" in text for text in table_ids)
    assert any("D/R_SDD01_001_002" in text for text in table_ids)


def test_apply_review_decision_uses_stable_fields_when_index_is_stale(tmp_path):
    out_doc = tmp_path / "reviewed.docx"
    new_code = tmp_path / "new"
    new_code.mkdir()
    _write_two_csu_doc(out_doc)

    wrong_item = PlannedItem(
        action="deleted_function",
        status="review",
        rel_path="demo.c",
        func_name="OtherFunc",
        csu_id="D/R_SDD01_999_001",
        reason="different row",
    )
    target_item = PlannedItem(
        action="deleted_function",
        status="review",
        rel_path="demo.c",
        func_name="OldFunc",
        csu_id="D/R_SDD01_001_001",
        reason="deleted function removal needs confirmation",
    )

    apply_review_decisions(
        [wrong_item, target_item],
        out_doc=str(out_doc),
        new_code=str(new_code),
        review_decisions=[{
            "item_index": 0,
            "decision": "delete_csu",
            "action": "deleted_function",
            "rel_path": "demo.c",
            "func_name": "OldFunc",
            "csu_id": "D/R_SDD01_001_001",
            "target_csu_id": "D/R_SDD01_001_001",
        }],
        ai_assist=False,
        template_path="",
    )

    assert wrong_item.status == "review"
    assert target_item.status == "applied_review", target_item.reason
    texts = _doc_texts(out_doc)
    assert not any("OldFunc" in text for text in texts)
    assert any("KeepFunc" in text for text in texts)


def test_apply_review_replace_csu_uses_real_pipeline(tmp_path):
    out_doc = tmp_path / "reviewed.docx"
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "demo.c").write_text(
        "void OldFunc(void)\n"
        "{\n"
        "    int status = 2;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )
    _write_two_csu_doc(out_doc)
    doc = Document(out_doc)
    doc.paragraphs[2].text = "old body marker"
    doc.save(out_doc)

    item = PlannedItem(
        action="modified_function",
        status="review",
        rel_path="demo.c",
        func_name="OldFunc",
        csu_id="D/R_SDD01_001_001",
        reason="manual replace confirmation",
    )

    apply_review_decisions(
        [item],
        out_doc=str(out_doc),
        new_code=str(new_code),
        review_decisions=[{
            "item_index": 0,
            "decision": "replace_csu",
            "target_csu_id": "D/R_SDD01_001_001",
            "func_name": "OldFunc",
            "rel_path": "demo.c",
        }],
        ai_assist=False,
        template_path="",
    )

    assert item.status == "applied_review", item.reason
    assert item.result["review_result"]["found"] is True
    texts = _doc_texts(out_doc)
    joined = "\n".join(texts)
    assert "old body marker" not in joined
    assert "void OldFunc(void)" in texts
    assert "设置状态 = 2；" in texts
    assert any(text.endswith("D/R_SDD01_001_001）") for text in texts)


def test_apply_review_manual_and_skip_work_with_manual_alignment_csu(tmp_path):
    out_doc = tmp_path / "reviewed.docx"
    new_code = tmp_path / "new"
    new_code.mkdir()
    _write_two_csu_doc(out_doc)

    manual_alignment = {
        "status": "manual_matched",
        "csu_id": "D/R_SDD01_001_001",
        "matched_function": "NewFunc",
        "rel_path": "demo.c",
        "evidence": {
            "source": "alignment_decisions",
            "manual_function": "NewFunc",
            "manual_rel_path": "demo.c",
        },
    }
    manual_item = PlannedItem(
        action="new_function",
        status="review",
        rel_path="demo.c",
        func_name="NewFunc",
        csu_id="D/R_SDD01_001_001",
        reason="new function has manual CSU alignment; review replace_csu target",
        alignment=manual_alignment,
    )
    skip_item = PlannedItem(
        action="deleted_function",
        status="review",
        rel_path="demo.c",
        func_name="OldFunc",
        csu_id="D/R_SDD01_001_002",
        reason="deleted function has manual CSU alignment",
        alignment={**manual_alignment, "csu_id": "D/R_SDD01_001_002"},
    )

    apply_review_decisions(
        [manual_item, skip_item],
        out_doc=str(out_doc),
        new_code=str(new_code),
        review_decisions=[
            {
                "item_index": 0,
                "decision": "manual",
                "csu_id": "D/R_SDD01_001_001",
                "func_name": "NewFunc",
                "rel_path": "demo.c",
            },
            {
                "item_index": 1,
                "decision": "skip",
                "csu_id": "D/R_SDD01_001_002",
                "func_name": "OldFunc",
                "rel_path": "demo.c",
            },
        ],
        ai_assist=False,
        template_path="",
    )

    assert manual_item.status == "manual_review"
    assert manual_item.reason == "review decision: manual"
    assert manual_item.csu_id == "D/R_SDD01_001_001"
    assert manual_item.alignment["status"] == "manual_matched"
    assert skip_item.status == "skipped_review"
    assert skip_item.reason == "review decision: skip"
    assert skip_item.csu_id == "D/R_SDD01_001_002"


def test_apply_review_insert_after_csu_decision_updates_doc_and_plan(tmp_path, monkeypatch):
    out_doc = tmp_path / "reviewed.docx"
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "demo.c").write_text(
        "void NewFunc(void)\\n{\\n    int status = 1;\\n    (void)status;\\n}\\n",
        encoding="utf-8",
    )
    _write_two_csu_doc(out_doc)

    item = PlannedItem(
        action="new_function",
        status="review",
        rel_path="demo.c",
        func_name="NewFunc",
        reason="new function insertion needs section position and CSU ID",
    )
    captured = {}

    def fake_insert_csu_after_in_doc(
        doc_path,
        source,
        func_name,
        csu_id,
        after_csu_id,
        cfg,
        *,
        project_root=None,
    ):
        captured.update({
            "doc_path": doc_path,
            "source": source,
            "func_name": func_name,
            "csu_id": csu_id,
            "after_csu_id": after_csu_id,
            "project_root": project_root,
        })
        doc = Document(doc_path)
        new_heading = doc.add_heading("NewFunc（D/R_SDD01_001_003）", level=4)
        new_proto = doc.add_paragraph("void NewFunc(void)")
        body = doc._body._element
        new_elements = [new_heading._element, new_proto._element]
        for elem in new_elements:
            body.remove(elem)
        result = insert_csu_after_in_doc(doc, after_csu_id, new_elements)
        result["module_table"] = sync_module_function_table_for_module(doc, "D/R_SDD01_001")
        doc.save(doc_path)
        return result

    import autodoc.pipeline as pipeline

    monkeypatch.setattr(pipeline, "insert_csu_after_in_doc", fake_insert_csu_after_in_doc)

    apply_review_decisions(
        [item],
        out_doc=str(out_doc),
        new_code=str(new_code),
        review_decisions=[{
            "item_index": 0,
            "decision": "insert_after_csu",
            "insert_after_csu_id": "D/R_SDD01_001_001",
            "func_name": "NewFunc",
            "rel_path": "demo.c",
        }],
        ai_assist=False,
        template_path="",
    )

    assert item.status == "applied_review", item.reason
    assert item.csu_id == "D/R_SDD01_001_003"
    assert item.result["review_result"]["found"] is True
    assert item.result["review_result"]["auto_allocated_csu_id"] == "D/R_SDD01_001_003"
    assert item.result["review_result"]["module_table"]["updated"] is True
    assert captured["func_name"] == "NewFunc"
    assert captured["csu_id"] == "D/R_SDD01_001_003"
    assert captured["after_csu_id"] == "D/R_SDD01_001_001"
    texts = _doc_texts(out_doc)
    new_heading = next(text for text in texts if "D/R_SDD01_001_003" in text)
    assert "New" in new_heading or "新" in new_heading
    assert texts.index("OldFunc（D/R_SDD01_001_001）") < texts.index(new_heading)
    assert texts.index(new_heading) < texts.index("KeepFunc（D/R_SDD01_001_002）")
    table_ids = _module_table_csu_ids(out_doc)
    assert any("D/R_SDD01_001_003" in text for text in table_ids)


def test_apply_review_insert_after_csu_uses_real_pipeline(tmp_path):
    out_doc = tmp_path / "reviewed.docx"
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "demo.c").write_text(
        "void OldFunc(void)\n"
        "{\n"
        "    int old_value = 0;\n"
        "    (void)old_value;\n"
        "}\n\n"
        "void NewFunc(void)\n"
        "{\n"
        "    int status = 1;\n"
        "    (void)status;\n"
        "}\n",
        encoding="utf-8",
    )
    _write_two_csu_doc(out_doc)

    item = PlannedItem(
        action="new_function",
        status="review",
        rel_path="demo.c",
        func_name="NewFunc",
        reason="new function insertion needs section position and CSU ID",
    )

    apply_review_decisions(
        [item],
        out_doc=str(out_doc),
        new_code=str(new_code),
        review_decisions=[{
            "item_index": 0,
            "decision": "insert_after_csu",
            "insert_after_csu_id": "D/R_SDD01_001_001",
            "func_name": "NewFunc",
            "rel_path": "demo.c",
        }],
        ai_assist=False,
        template_path="",
    )

    assert item.status == "applied_review", item.reason
    assert item.csu_id == "D/R_SDD01_001_003"
    assert item.result["review_result"]["found"] is True
    assert item.result["review_result"]["module_table"]["entries"] == 3
    texts = _doc_texts(out_doc)
    inserted_heading = next(text for text in texts if text.endswith("D/R_SDD01_001_003）"))
    assert "NewFunc" in "\n".join(texts)
    assert "void NewFunc(void)" in texts
    assert texts.index("OldFunc（D/R_SDD01_001_001）") < texts.index(inserted_heading)
    assert texts.index(inserted_heading) < texts.index("KeepFunc（D/R_SDD01_001_002）")
    table_ids = _module_table_csu_ids(out_doc)
    assert any("D/R_SDD01_001_003" in text for text in table_ids)


def test_apply_review_insert_can_renumber_module_csu_ids(tmp_path, monkeypatch):
    out_doc = tmp_path / "reviewed.docx"
    new_code = tmp_path / "new"
    new_code.mkdir()
    (new_code / "demo.c").write_text(
        "void NewFunc(void)\\n{\\n    int status = 1;\\n    (void)status;\\n}\\n",
        encoding="utf-8",
    )
    _write_two_csu_doc(out_doc)

    item = PlannedItem(
        action="new_function",
        status="review",
        rel_path="demo.c",
        func_name="NewFunc",
        reason="new function insertion needs section position and CSU ID",
    )

    def fake_insert_csu_after_in_doc(
        doc_path,
        source,
        func_name,
        csu_id,
        after_csu_id,
        cfg,
        *,
        project_root=None,
    ):
        doc = Document(doc_path)
        new_heading = doc.add_heading(f"{func_name}（{csu_id}）", level=4)
        new_proto = doc.add_paragraph(f"void {func_name}(void)")
        body = doc._body._element
        new_elements = [new_heading._element, new_proto._element]
        for elem in new_elements:
            body.remove(elem)
        result = insert_csu_after_in_doc(doc, after_csu_id, new_elements)
        result["module_table"] = sync_module_function_table_for_module(doc, "D/R_SDD01_001")
        doc.save(doc_path)
        return result

    import autodoc.pipeline as pipeline

    monkeypatch.setattr(pipeline, "insert_csu_after_in_doc", fake_insert_csu_after_in_doc)

    apply_review_decisions(
        [item],
        out_doc=str(out_doc),
        new_code=str(new_code),
        review_decisions=[{
            "item_index": 0,
            "decision": "insert_after_csu",
            "insert_after_csu_id": "D/R_SDD01_001_001",
            "func_name": "NewFunc",
            "rel_path": "demo.c",
        }],
        ai_assist=False,
        template_path="",
        renumber_module_csu=True,
    )

    assert item.status == "applied_review", item.reason
    assert item.csu_id == "D/R_SDD01_001_002"
    result = item.result["review_result"]
    assert result["auto_allocated_csu_id"] == "D/R_SDD01_001_003"
    assert result["renumber_module_csu"]["updated"] == 2
    texts = _doc_texts(out_doc)
    assert texts.index("OldFunc（D/R_SDD01_001_001）") < texts.index("NewFunc（D/R_SDD01_001_002）")
    assert texts.index("NewFunc（D/R_SDD01_001_002）") < texts.index("KeepFunc（D/R_SDD01_001_003）")
    assert _module_table_csu_ids(out_doc) == [
        "OldFunc:D/R_SDD01_001_001",
        "NewFunc:D/R_SDD01_001_002",
        "KeepFunc:D/R_SDD01_001_003",
    ]
