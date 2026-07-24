# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import ast
import glob
import json
import shlex
import string
import argparse
import copy
from dataclasses import dataclass, fields, replace
from collections import Counter, OrderedDict, defaultdict
from typing import Any, Literal, Optional, Callable, Sequence
import concurrent.futures
from pathlib import Path
try:
    import requests  # 新增：Win7 友好的 HTTP 客户端
except ImportError:  # 生成文档时未用到 AI，可缺省安装 requests
    requests = None
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

import threading
import sys
import time
import tempfile

from ._legacy_support import legacy_backend
from . import utils
from . import text as text_utils
from . import cli as cli_utils
from .models import CommentHint as StableCommentHint
from .models import FunctionDesign as StableFunctionDesign
from .models import IOElement as StableIOElement
from .models import LocalDataElement as StableLocalDataElement
from .models import SymbolEvidence as StableSymbolEvidence
from .models import SymbolInference as StableSymbolInference

if __name__ != "autodoc.backend" and sys.modules.get(__name__) is not None:
    sys.modules["autodoc.backend"] = sys.modules[__name__]
    if sys.modules.get("autodoc") is not None:
        setattr(sys.modules["autodoc"], "backend", sys.modules[__name__])
if sys.modules.get(__name__) is not None:
    sys.modules["_autodoc_legacy_backend"] = sys.modules[__name__]

APP_NAME = "AutoDocGen"
APP_VERSION = "V2.00"

def _ensure_stdio():
    """
    某些环境（比如 pythonw / PyInstaller --noconsole）下 sys.stdout/sys.stderr 可能为 None，
    argparse / traceback / print 会直接崩；这里兜底到 devnull。
    """
    if sys.stdout is None:
        sys.stdout = open(os.devnull, "w", encoding="utf-8", errors="ignore")
    if sys.stderr is None:
        sys.stderr = sys.stdout if sys.stdout is not None else open(
            os.devnull, "w", encoding="utf-8", errors="ignore"
        )


_ensure_stdio()


def _utils():
    """Resolve the utils module — works in both normal import and exec() mode."""
    try:
        from . import utils as _m
        return _m
    except ImportError:
        import autodoc.utils as _m
        return _m


def _text():
    """Resolve the text module — works in both normal import and exec() mode."""
    try:
        from . import text as _m
        return _m
    except ImportError:
        import autodoc.text as _m
        return _m


def _cli():
    """Resolve the cli module — works in both normal import and exec() mode."""
    try:
        from . import cli as _m
        return _m
    except ImportError:
        import autodoc.cli as _m
        return _m



import os
os.environ.pop('HTTPS_PROXY', None)
os.environ.pop('HTTP_PROXY', None)

import datetime
import traceback
import hashlib


_AI_RESPONSE_CACHE_MAX = 512
_AI_RESPONSE_CACHE_LOCK = threading.Lock()
_AI_RESPONSE_CACHE: "OrderedDict[str, Any]" = OrderedDict()
_AI_HTTP_LOCAL = threading.local()
_CACHE_MISS = object()
_SYMBOL_MEMORY_LOCK = threading.Lock()
_NAMING_INDEX_LOCK = threading.Lock()
_PROJECT_SYMBOL_MEMORY_DATA: dict[str, Any] = {}
_PROJECT_SYMBOL_MEMORY_PATH = ""
_PROJECT_TITLE_INDEX_DATA: dict[str, Any] = {}
_PROJECT_TITLE_INDEX_PATH = ""
_PROJECT_SYMBOL_INDEX_DATA: dict[str, Any] = {}
_PROJECT_SYMBOL_INDEX_PATH = ""
_PROJECT_SEMANTIC_INDEX_DATA: dict[str, Any] = {}
_PROJECT_SEMANTIC_INDEX_PATH = ""
_PROJECT_TERM_TABLE_DATA: dict[str, Any] = {}
_PROJECT_TERM_TABLE_PATH = ""


def reset_project_index_state() -> None:
    """Reset all project-level index state. For test isolation."""
    global _PROJECT_SYMBOL_MEMORY_DATA, _PROJECT_SYMBOL_MEMORY_PATH
    global _PROJECT_TITLE_INDEX_DATA, _PROJECT_TITLE_INDEX_PATH
    global _PROJECT_SYMBOL_INDEX_DATA, _PROJECT_SYMBOL_INDEX_PATH
    global _PROJECT_SEMANTIC_INDEX_DATA, _PROJECT_SEMANTIC_INDEX_PATH
    global _PROJECT_TERM_TABLE_DATA, _PROJECT_TERM_TABLE_PATH
    with _SYMBOL_MEMORY_LOCK:
        _PROJECT_SYMBOL_MEMORY_DATA.clear()
        _PROJECT_SYMBOL_MEMORY_PATH = ""
    with _NAMING_INDEX_LOCK:
        _PROJECT_TITLE_INDEX_DATA.clear()
        _PROJECT_TITLE_INDEX_PATH = ""
        _PROJECT_SYMBOL_INDEX_DATA.clear()
        _PROJECT_SYMBOL_INDEX_PATH = ""
        _PROJECT_SEMANTIC_INDEX_DATA.clear()
        _PROJECT_SEMANTIC_INDEX_PATH = ""
        _PROJECT_TERM_TABLE_DATA.clear()
        _PROJECT_TERM_TABLE_PATH = ""
    SYMBOL_DICTIONARY_RUNTIME.clear()
    SESSION_SYMBOL_DICTIONARY.clear()


def _safe_text(value: Any, *, default: str = "") -> str:
    return utils._safe_text(value, default=default)


def _safe_strip(value: Any, *, default: str = "") -> str:
    return utils._safe_strip(value, default=default)


def _safe_textish(value: Any) -> str:
    from . import ai as ai_utils
    return ai_utils._safe_textish(value)


def _looks_like_codeish_description(text: str) -> bool:
    from . import ai as ai_utils
    return ai_utils._looks_like_codeish_description(text)


def _fallback_function_description(func_info, body, *, current_desc=""):
    from . import ai as ai_utils
    return ai_utils._fallback_function_description(func_info, body, current_desc=current_desc)


def _make_ai_cache_key(kind: str, prompt: str, cfg: "GenConfig", provider: str, url: str) -> str:
    from . import ai as ai_utils

    return ai_utils._make_ai_cache_key(kind, prompt, cfg, provider, url)


def _ai_cache_get(key: str, *, clone: bool) -> Any:
    from . import ai as ai_utils

    return ai_utils._ai_cache_get(key, clone=clone)


def _ai_cache_set(key: str, value: Any, *, clone: bool) -> None:
    from . import ai as ai_utils

    ai_utils._ai_cache_set(key, value, clone=clone)


def _clear_ai_runtime_state() -> None:
    from . import ai as ai_utils

    ai_utils._clear_ai_runtime_state()


def _get_http_session(cfg: Optional["GenConfig"] = None):
    from . import ai as ai_utils

    return ai_utils._get_http_session(cfg)


def _normalize_proxy_url(raw: Any) -> str:
    from . import ai as ai_utils

    return ai_utils._normalize_proxy_url(
        raw,
        backend_module=sys.modules[__name__],
    )


def _default_proxy_candidate_ports(cfg: Optional["GenConfig"] = None) -> tuple[str, ...]:
    from . import ai as ai_utils

    return ai_utils._default_proxy_candidate_ports(
        cfg,
        backend_module=sys.modules[__name__],
    )


def _resolve_proxy_candidates(
    cfg: Optional["GenConfig"],
    *,
    provider: str = "",
    url: str = "",
) -> list[tuple[str, str]]:
    from . import ai as ai_utils

    return ai_utils._resolve_proxy_candidates(
        cfg,
        provider=provider,
        url=url,
        backend_module=sys.modules[__name__],
    )


def _proxy_dict_for_request(proxy_url: str) -> dict[str, str]:
    from . import ai as ai_utils

    return ai_utils._proxy_dict_for_request(
        proxy_url,
        backend_module=sys.modules[__name__],
    )


def _post_with_proxy_fallback(
    *,
    session: Any,
    url: str,
    data: dict[str, Any],
    headers: dict[str, str],
    timeout: Any,
    cfg: Optional["GenConfig"],
    provider: str,
    attempts: int,
    log_label: str,
) -> tuple[Any, str, dict[str, Any]]:
    from . import ai as ai_utils

    return ai_utils._post_with_proxy_fallback(
        session=session,
        url=url,
        data=data,
        headers=headers,
        timeout=timeout,
        cfg=cfg,
        provider=provider,
        attempts=attempts,
        log_label=log_label,
    )




# ================= 配置区域（默认值） =================

DEFAULT_SECTION_PREFIX = '5.1.1.'          # 小节号前缀，例如 "5.1.1."
DEFAULT_REQ_ID_PREFIX = 'D/R_SDD01'        # 需求/设计ID前缀（不带末尾编号，函数/模块处自动拼三位序号）

# 航空类常用 BIT 术语表（可按需扩展/修改）
DOMAIN_GLOSSARY_BASE = {
    "PuBIT": "上电BIT",
    "IFBIT": "周期BIT",
    "PBIT": "飞行前BIT",
    "MBIT": "维护BIT",
}
DOMAIN_GLOSSARY = dict(DOMAIN_GLOSSARY_BASE)


def build_project_glossary(
    file_symbols: Optional[dict[str, str]] = None,
) -> dict[str, str]:
    from . import semantic as semantic_utils

    return semantic_utils.build_project_glossary(
        file_symbols=file_symbols,
        backend_module=sys.modules[__name__],
    )


SYMBOL_DICTIONARY_BASE: dict[str, str] = {}
SYMBOL_DICTIONARY_RUNTIME: dict[str, str] = {}
SESSION_SYMBOL_DICTIONARY: dict[str, dict[str, Any]] = {}


def app_root() -> Path:
    from ._legacy_support import app_root as _app_root

    return Path(_app_root())


def parse_domain_glossary_text(text: str) -> dict[str, str]:
    from . import naming as naming_utils

    return naming_utils.parse_domain_glossary_text(text)


def _default_symbol_dictionary_path() -> str:
    from . import naming as naming_utils

    return naming_utils._default_symbol_dictionary_path(backend_module=sys.modules[__name__])


def _normalize_symbol_dictionary_payload(data: Any) -> dict[str, str]:
    from . import naming as naming_utils

    return naming_utils._normalize_symbol_dictionary_payload(data)


def parse_symbol_dictionary_text(text: str) -> dict[str, str]:
    from . import naming as naming_utils

    return naming_utils.parse_symbol_dictionary_text(text)


def load_symbol_dictionary_file(path: str) -> dict[str, str]:
    from . import naming as naming_utils

    return naming_utils.load_symbol_dictionary_file(path)


def _default_project_symbol_memory_path(project_root: str) -> str:
    from . import naming as naming_utils

    return naming_utils._default_project_symbol_memory_path(project_root)


def _normalize_symbol_memory_record(record: Any, *, raw_ident: str = "", section: str = "") -> dict[str, Any]:
    from . import naming as naming_utils

    return naming_utils._normalize_symbol_memory_record(
        record,
        raw_ident=raw_ident,
        section=section,
        backend_module=sys.modules[__name__],
    )


def _normalize_project_symbol_memory_payload(data: Any) -> dict[str, Any]:
    from . import naming as naming_utils

    return naming_utils._normalize_project_symbol_memory_payload(data, backend_module=sys.modules[__name__])


def load_project_symbol_memory(project_root: str) -> tuple[str, dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils.load_project_symbol_memory(project_root, backend_module=sys.modules[__name__])


def _flatten_project_symbol_memory(data: Optional[dict[str, Any]]) -> dict[str, str]:
    from . import naming as naming_utils

    return naming_utils._flatten_project_symbol_memory(data, backend_module=sys.modules[__name__])


def save_project_symbol_memory() -> None:
    from . import naming as naming_utils

    return naming_utils.save_project_symbol_memory(backend_module=sys.modules[__name__])


def init_project_symbol_memory(project_root: str, cfg: Optional["GenConfig"] = None, overrides: Optional[dict[str, str]] = None) -> None:
    from . import naming as naming_utils

    return naming_utils.init_project_symbol_memory(
        project_root,
        cfg,
        overrides,
        backend_module=sys.modules[__name__],
    )


def finalize_project_symbol_memory(cfg: Optional["GenConfig"] = None) -> None:
    from . import naming as naming_utils

    return naming_utils.finalize_project_symbol_memory(cfg, backend_module=sys.modules[__name__])


def default_term_table_path(project_root: str) -> str:
    from . import term_table as term_table_utils

    return term_table_utils.default_term_table_path(project_root)


def load_project_term_table(project_root: str) -> tuple[str, dict[str, Any]]:
    from . import term_table as term_table_utils

    return term_table_utils.load_term_table(project_root, backend_module=sys.modules[__name__])


def save_project_term_table(path: str, payload: dict[str, Any]) -> None:
    from . import term_table as term_table_utils

    return term_table_utils.save_term_table(path, payload, backend_module=sys.modules[__name__])


def flatten_project_term_table(payload: Optional[dict[str, Any]]) -> dict[str, str]:
    from . import term_table as term_table_utils

    return term_table_utils.flatten_term_table(payload, backend_module=sys.modules[__name__])


def apply_project_term_table(project_root: str) -> dict[str, str]:
    from . import term_table as term_table_utils

    flat = term_table_utils.apply_term_table_to_runtime(project_root, backend_module=sys.modules[__name__])
    with _NAMING_INDEX_LOCK:
        global _PROJECT_TERM_TABLE_PATH, _PROJECT_TERM_TABLE_DATA
        _PROJECT_TERM_TABLE_PATH, _PROJECT_TERM_TABLE_DATA = term_table_utils.load_term_table(
            project_root,
            backend_module=sys.modules[__name__],
        )
    return flat


def build_project_term_table(
    project_root: str,
    cfg: Optional["GenConfig"] = None,
    *,
    prebuilt: Optional[dict[str, dict[str, str]]] = None,
    func_entries: Optional[Sequence[dict]] = None,
    save: bool = True,
) -> dict[str, Any]:
    from . import term_table as term_table_utils

    payload = term_table_utils.build_project_term_table(
        project_root,
        cfg,
        prebuilt=prebuilt,
        func_entries=func_entries,
        save=save,
        backend_module=sys.modules[__name__],
    )
    with _NAMING_INDEX_LOCK:
        global _PROJECT_TERM_TABLE_PATH, _PROJECT_TERM_TABLE_DATA
        _PROJECT_TERM_TABLE_PATH = term_table_utils.default_term_table_path(project_root)
        _PROJECT_TERM_TABLE_DATA = payload
    return payload


def update_project_term_table_records(
    project_root: str,
    updates: Sequence[dict[str, Any]],
    *,
    save: bool = True,
) -> dict[str, Any]:
    from . import term_table as term_table_utils

    payload = term_table_utils.update_term_table_records(
        project_root,
        updates,
        save=save,
        backend_module=sys.modules[__name__],
    )
    with _NAMING_INDEX_LOCK:
        global _PROJECT_TERM_TABLE_PATH, _PROJECT_TERM_TABLE_DATA
        _PROJECT_TERM_TABLE_PATH = term_table_utils.default_term_table_path(project_root)
        _PROJECT_TERM_TABLE_DATA = payload
    return payload


def _default_project_title_index_path(project_root: str) -> str:
    from . import naming as naming_utils

    return naming_utils._default_project_title_index_path(project_root)


def _default_project_symbol_index_path(project_root: str) -> str:
    from . import naming as naming_utils

    return naming_utils._default_project_symbol_index_path(project_root)


def _default_project_semantic_index_path(project_root: str) -> str:
    root = os.path.abspath(os.path.expanduser(str(project_root or "").strip())) if project_root else ""
    if not root:
        return os.path.abspath("autodoc_semantic_index.json")
    return os.path.join(root, "autodoc_semantic_index.json")


def _normalize_title_index_record(record: Any) -> dict[str, Any]:
    from . import naming as naming_utils

    return naming_utils._normalize_title_index_record(record, backend_module=sys.modules[__name__])


def _normalize_symbol_index_record(record: Any) -> dict[str, Any]:
    from . import naming as naming_utils

    return naming_utils._normalize_symbol_index_record(record, backend_module=sys.modules[__name__])


def _normalize_title_index_payload(data: Any) -> dict[str, Any]:
    from . import naming as naming_utils

    return naming_utils._normalize_title_index_payload(data, backend_module=sys.modules[__name__])


def _normalize_symbol_index_payload(data: Any) -> dict[str, Any]:
    from . import naming as naming_utils

    return naming_utils._normalize_symbol_index_payload(data, backend_module=sys.modules[__name__])


def _normalize_semantic_symbol_profile(record: Any) -> dict[str, Any]:
    if not isinstance(record, dict):
        return {}
    out = {
        "name": _safe_strip(record.get("name")),
        "scope": _safe_strip(record.get("scope")),
        "decl_type": _safe_strip(record.get("decl_type")),
        "role": _safe_strip(record.get("role")),
        "direction": _safe_strip(record.get("direction")),
        "producer_call": _safe_strip(record.get("producer_call")),
        "producer_arg_tags": tuple(
            _safe_strip(x) for x in (record.get("producer_arg_tags") or ()) if _safe_strip(x)
        ),
        "consumer_patterns": tuple(
            _safe_strip(x) for x in (record.get("consumer_patterns") or ()) if _safe_strip(x)
        ),
        "sink_patterns": tuple(
            _safe_strip(x) for x in (record.get("sink_patterns") or ()) if _safe_strip(x)
        ),
        "dataflow_roles": tuple(
            _safe_strip(x) for x in (record.get("dataflow_roles") or ()) if _safe_strip(x)
        ),
        "usage_patterns": tuple(
            _safe_strip(x) for x in (record.get("usage_patterns") or ()) if _safe_strip(x)
        ),
        "paired_symbols": tuple(
            _safe_strip(x) for x in (record.get("paired_symbols") or ()) if _safe_strip(x)
        ),
    }
    if not out["name"]:
        return {}
    return out


def _normalize_semantic_index_record(record: Any) -> dict[str, Any]:
    if not isinstance(record, dict):
        return {}
    out = {
        "id": _safe_strip(record.get("id")),
        "source_file": _safe_strip(record.get("source_file")),
        "module_key": _safe_strip(record.get("module_key")),
        "func_name": _safe_strip(record.get("func_name")),
        "family_prefix": _safe_strip(record.get("family_prefix")),
        "action_suffix": _safe_strip(record.get("action_suffix")),
        "ret_type": _safe_strip(record.get("ret_type")),
        "comment_desc": _safe_strip(record.get("comment_desc")),
        "callee_names": tuple(
            _safe_strip(x) for x in (record.get("callee_names") or ()) if _safe_strip(x)
        ),
        "macro_refs": tuple(
            _safe_strip(x) for x in (record.get("macro_refs") or ()) if _safe_strip(x)
        ),
        "condition_signatures": tuple(
            _safe_strip(x) for x in (record.get("condition_signatures") or ()) if _safe_strip(x)
        ),
        "member_accesses": tuple(
            _safe_strip(x) for x in (record.get("member_accesses") or ()) if _safe_strip(x)
        ),
        "return_exprs": tuple(
            _safe_strip(x) for x in (record.get("return_exprs") or ()) if _safe_strip(x)
        ),
        "return_symbols": tuple(
            _safe_strip(x) for x in (record.get("return_symbols") or ()) if _safe_strip(x)
        ),
        "written_params": tuple(
            _safe_strip(x) for x in (record.get("written_params") or ()) if _safe_strip(x)
        ),
        "read_params": tuple(
            _safe_strip(x) for x in (record.get("read_params") or ()) if _safe_strip(x)
        ),
        "symbol_profiles": [],
    }
    for item in (record.get("symbol_profiles") or ()):
        profile = _normalize_semantic_symbol_profile(item)
        if profile:
            out["symbol_profiles"].append(profile)
    if not out["id"]:
        return {}
    return out


def _normalize_semantic_index_payload(data: Any) -> dict[str, Any]:
    out = {"version": 1, "items": []}
    if not isinstance(data, dict):
        return out
    items = []
    for item in (data.get("items") or []):
        rec = _normalize_semantic_index_record(item)
        if rec:
            items.append(rec)
    out["items"] = items
    if data.get("updated_at"):
        out["updated_at"] = str(data.get("updated_at"))
    return out


def load_project_title_index(project_root: str) -> tuple[str, dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils.load_project_title_index(project_root, backend_module=sys.modules[__name__])


def load_project_symbol_index(project_root: str) -> tuple[str, dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils.load_project_symbol_index(project_root, backend_module=sys.modules[__name__])


def load_project_semantic_index(project_root: str) -> tuple[str, dict[str, Any]]:
    path = _default_project_semantic_index_path(project_root)
    if not os.path.isfile(path):
        return path, _normalize_semantic_index_payload({})
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return path, _normalize_semantic_index_payload({})
    return path, _normalize_semantic_index_payload(data)


def _save_json_sidecar(path: str, payload: dict[str, Any], prefix: str) -> None:
    from . import naming as naming_utils

    return naming_utils._save_json_sidecar(path, payload, prefix)


def _collect_project_source_mtime(project_root: str, *, max_files: int = 0) -> float:
    from . import naming as naming_utils

    return naming_utils._collect_project_source_mtime(project_root, max_files=max_files)


def _should_refresh_naming_indexes(
    project_root: str,
    title_path: str,
    symbol_path: str,
    cfg: Optional["GenConfig"],
) -> bool:
    from . import naming as naming_utils

    return naming_utils._should_refresh_naming_indexes(
        project_root,
        title_path,
        symbol_path,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _identifier_family_prefix(name: str) -> str:
    tokens = _split_ident_tokens(name)
    if not tokens:
        return ""
    if len(tokens) >= 2 and tokens[1].upper() == "BIT":
        return f"{tokens[0]}{tokens[1]}"
    return tokens[0]


def _identifier_action_suffix(name: str) -> str:
    tokens = _split_ident_tokens(name)
    if not tokens:
        return ""
    known = {
        "Get", "Set", "Init", "Update", "Check", "Test", "Handle", "Pack",
        "Send", "Recv", "Receive", "Write", "Read", "Calc", "Clear", "Reset",
    }
    last = tokens[-1]
    if last in known:
        return last
    if len(tokens) >= 2 and f"{tokens[-2]}{tokens[-1]}" in {"StateUpdate", "ResultGet"}:
        return f"{tokens[-2]}{tokens[-1]}"
    return last


def _module_key_for_source(path: str) -> str:
    return os.path.splitext(os.path.basename(_safe_strip(path)))[0]


def _text_terms(text: str) -> set[str]:
    s = _safe_strip(text)
    if not s:
        return set()
    tokens = {tok.lower() for tok in re.split(r"[\s,，;；。、:/()（）]+", s) if tok}
    if _contains_cjk(s):
        compact = re.sub(r"\s+", "", s)
        for idx in range(len(compact) - 1):
            gram = compact[idx:idx + 2]
            if _contains_cjk(gram):
                tokens.add(gram)
    return {tok for tok in tokens if tok}


def _remember_ai_symbol(
    name: str,
    cn: str,
    *,
    kind: str,
    confidence: float,
    evidence_kinds: int = 2,
    persist_scope: str = "graded",
    cfg: Optional["GenConfig"] = None,
    source: str = "ai",
) -> bool:
    from . import naming as naming_utils

    return naming_utils._remember_ai_symbol(
        name,
        cn,
        kind=kind,
        confidence=confidence,
        evidence_kinds=evidence_kinds,
        persist_scope=persist_scope,
        cfg=cfg,
        source=source,
        backend_module=sys.modules[__name__],
    )


def apply_symbol_dictionary_overrides(
    overrides: Optional[dict[str, str]],
    file_path: str = "",
    project_memory: Optional[dict[str, str]] = None,
) -> None:
    from . import naming as naming_utils

    return naming_utils.apply_symbol_dictionary_overrides(
        overrides,
        file_path=file_path,
        project_memory=project_memory,
        backend_module=sys.modules[__name__],
    )


def _lookup_symbol_dictionary(name: str) -> str:
    from . import naming as naming_utils

    return naming_utils._lookup_symbol_dictionary(name, backend_module=sys.modules[__name__])


def _lookup_session_symbol_record(name: str) -> dict[str, Any]:
    from . import naming as naming_utils

    return naming_utils._lookup_session_symbol_record(name, backend_module=sys.modules[__name__])


def _lookup_session_symbol(name: str) -> str:
    from . import naming as naming_utils

    return naming_utils._lookup_session_symbol(name, backend_module=sys.modules[__name__])


def _remember_inferred_symbol(
    name: str,
    cn: str,
    *,
    kind: str,
    confidence: float,
    evidence_kinds: int,
    cfg: Optional["GenConfig"] = None,
    source: str = "infer",
) -> bool:
    from . import naming as naming_utils

    return naming_utils._remember_inferred_symbol(
        name,
        cn,
        kind=kind,
        confidence=confidence,
        evidence_kinds=evidence_kinds,
        cfg=cfg,
        source=source,
        backend_module=sys.modules[__name__],
    )


def _collect_preferred_symbol_names(names: Sequence[str], *, limit: int = 24) -> dict[str, str]:
    from . import naming as naming_utils

    return naming_utils.collect_preferred_symbol_names(
        names,
        limit=limit,
        backend_module=sys.modules[__name__],
    )


_GENERIC_CN_NAMES = {
    "初始化",
    "处理",
    "配置",
    "函数",
    "模块",
    "数据",
    "变量",
    "状态",
    "标志",
    "用途",
}
_GENERIC_LOCAL_CN_NAMES = {
    "中间量",
    "中间变量",
    "临时",
    "临时值",
    "临时变量",
    "临时数据",
    "计数器",
    "缓存值",
    "缓存缓存值",
    "当前值",
    "变量值",
    "数据值",
    "数据指针",
    "指针",
    "宏定义",
    "临时8位整型",
    "临时16位整型",
    "临时32位整型",
    "临时64位整型",
}

_PURPOSE_MARKERS = ("用于", "以便", "防止", "避免", "保证", "确保", "供")
_HISTORY_MARKERS = ("修改记录", "发布日期", "版本记录", "改为", "TODO", "FIXME", "TESTONLY")
_STRICT_SYMBOL_ROLES = (
    "返回值", "缓存值", "指针", "索引", "计数器", "上一周期值", "当前值", "标志",
    "模式", "状态", "阈值", "中间量", "换算系数", "偏移", "时间阈值", "寄存器", "位标志", "宏定义",
)
_ASCII_WORD_RE = re.compile(r"[A-Za-z]+")
_STRICT_CN_PURPOSE_TAIL_RE = re.compile(r"(?:用于|以便|供(?!油)).+$")
_GENERIC_MACRO_FALLBACK_NAMES = {
    "宏定义",
    "位标志",
    "换算系数",
    "偏移量",
    "超时阈值",
    "时间阈值",
    "寄存器",
}


def _strip_allowed_cn_acronyms(text: str) -> str:
    s = _safe_strip(text)
    s = re.sub(r"[A-Z]{1,6}\d+", "", s)
    s = re.sub(r"\d+(?:kHz|MHz|Hz|MS|ms|V|A|W)(?=$|[^A-Za-z])", "", s)
    s = re.sub(r"(?<=[\u4e00-\u9fff0-9])[ABCXYZ](?=$)", "", s)
    s = re.sub(r"[A-Z]{2,6}", "", s)
    return s


def _looks_like_bad_canonical_name(text: str, *, raw_ident: str = "") -> bool:
    s = _safe_strip(text)
    if not s:
        return True
    compact = re.sub(r"\s+", "", s)
    if compact in _GENERIC_CN_NAMES:
        return True
    if compact in _GENERIC_LOCAL_CN_NAMES:
        return True
    if compact.upper() in ("TODO", "NONE", "N/A", "NA"):
        return True
    if not _contains_cjk(compact):
        raw = _safe_strip(raw_ident)
        return (not raw) or compact == raw
    ascii_removed = _strip_allowed_cn_acronyms(compact)
    if re.search(r"[A-Za-z]", ascii_removed):
        return True
    if any(mark in compact for mark in _HISTORY_MARKERS):
        return True
    if len(compact) <= 2 and compact in {"变量", "数据", "函数"}:
        return True
    return False


def _looks_like_low_quality_member_cn(text: str) -> bool:
    compact = re.sub(r"\s+", "", _safe_strip(text))
    if not compact:
        return True
    return compact in {"状态快照", "上拍状态", "结果快照", "上拍结果"}


def _looks_like_memberish_bitfield(name: str) -> bool:
    ident = _safe_strip(name)
    if not ident:
        return False
    if re.search(r"_b\d+$", ident, flags=re.IGNORECASE):
        return True
    return bool(re.fullmatch(r"[A-Z0-9_]+", ident) and "_" in ident)


def _is_complex_project_macro(name: str) -> bool:
    ident = _safe_strip(name).upper()
    if not _is_macro_identifier(ident):
        return False
    parts = [part for part in ident.split("_") if part]
    if len(parts) >= 3:
        return True
    if any(len(part) <= 1 for part in parts):
        return True
    return False


def _should_preserve_macro_token(name: str, replacement: str) -> bool:
    ident = _safe_strip(name)
    mapped = re.sub(r"\s+", "", _safe_strip(replacement))
    if (not ident) or (not _is_macro_identifier(ident)):
        return False
    if not mapped:
        return True
    if mapped in _GENERIC_MACRO_FALLBACK_NAMES and _is_complex_project_macro(ident):
        return True
    return False


def _default_macro_cn_for_role(name: str, role: str) -> str:
    ident = _safe_strip(name).upper()
    if (not ident) or _is_complex_project_macro(ident):
        return ""
    parts = [part for part in ident.split("_") if part]
    if role == "时间阈值" and any(ch.isdigit() for ch in ident) and len(parts) <= 2:
        return "超时阈值"
    if role == "换算系数" and len(parts) <= 2:
        return "换算系数"
    if role == "偏移" and len(parts) <= 2:
        return "偏移量"
    if role == "寄存器" and len(parts) <= 1:
        return "寄存器"
    if role == "位标志" and len(parts) <= 2:
        return "位标志"
    return ""


def _strict_symbol_persist_scope(confidence: float, evidence_kinds: int) -> str:
    conf_val = float(confidence or 0.0)
    evidence_count = max(0, int(evidence_kinds or 0))
    if conf_val < 0.60:
        return "off"
    if conf_val < 0.82:
        return "session_only"
    if evidence_count >= 2:
        return "graded"
    return "session_only"


def _is_strict_symbol_candidate_rejected(text: str, *, raw_ident: str = "") -> bool:
    s = _safe_strip(text)
    if _looks_like_bad_canonical_name(s, raw_ident=raw_ident):
        return True
    from . import naming as naming_utils

    if naming_utils.title_violates_required_acronyms(s, raw_ident):
        return True
    return naming_utils.is_strict_symbol_candidate_rejected(_strip_allowed_cn_acronyms(s), raw_ident=raw_ident)


def _sanitize_ai_usage_text(text: str) -> str:
    from . import naming as naming_utils

    s = naming_utils.sanitize_ai_usage_text(text)
    s = _STRICT_CN_PURPOSE_TAIL_RE.sub("", s).strip()
    return s


def _normalize_symbol_hint_text(text: str) -> str:
    s = re.sub(r"[。;；]+$", "", _safe_strip(text))
    if (not s) or (not _contains_cjk(s)):
        return ""
    compact = re.sub(r"\s+", "", s)
    if len(compact) > 12:
        return ""
    if any(mark in compact for mark in ("用于", "以便", "表示", "范围", "单位", "说明", "对应", "默认")):
        return ""
    if any(mark in compact for mark in ("保存", "记录", "存放", "比较", "刷新", "更新")) and len(compact) >= 5:
        return ""
    if compact in _GENERIC_LOCAL_CN_NAMES:
        return ""
    if compact.endswith("索引"):
        return "索引"
    if "检测结果" in compact:
        return "检测结果"
    if "状态值" in compact:
        return "状态值"
    if compact in ("临时数据", "临时值"):
        return ""
    return compact


def _call_source_name_from_patterns(patterns: Sequence[str]) -> str:
    for pattern in (patterns or ()):
        text = str(pattern or "").strip()
        if text.startswith("call_source:"):
            return text.split(":", 1)[1].strip()
    return ""


def _derive_candidate_cn_from_evidence(evidence: "SymbolEvidence") -> str:
    if ("returned" in set(evidence.usage_patterns or ()) or "returned_directly" in set(evidence.consumer_patterns or ())) and "results_bit32" not in set(evidence.producer_arg_tags or ()):
        return ""
    concepts = _candidate_concepts_from_evidence(evidence)
    if concepts:
        return concepts[0]

    if evidence.normalized_comment_hint:
        return evidence.normalized_comment_hint

    call_source = _call_source_name_from_patterns(evidence.usage_patterns or ())
    call_lower = call_source.lower()
    symbol_lower = _safe_strip(evidence.symbol).lower()
    if call_source and any(tag in call_lower for tag in ("stateget", "statusget", "state", "status")):
        return "状态值"
    if (
        ("result" in symbol_lower or "results" in symbol_lower or "ret" in symbol_lower)
        and any(tag in call_lower for tag in ("check", "test", "resultget", "result"))
    ):
        return "检测结果"
    if "result" in symbol_lower or "results" in symbol_lower:
        return "结果"
    return ""


def _looks_like_state_machine_body(body: str) -> bool:
    text = _safe_text(body)
    if not text:
        return False
    mode_hits = len(re.findall(r"(?:mode|Mode)", text))
    branch_hits = len(re.findall(r"\b(?:if|else\s+if)\b", text))
    return mode_hits >= 3 and branch_hits >= 2


def _build_state_machine_logic_summary(body: str) -> str:
    if not _looks_like_state_machine_body(body):
        return ""
    example_modes: list[str] = []
    for hit in re.findall(r"==\s*([A-Za-z_]\w*)", _safe_text(body)):
        name = _safe_strip(hit)
        if not name or name in example_modes:
            continue
        example_modes.append(name)
        if len(example_modes) >= 3:
            break
    examples = "、".join(example_modes) if example_modes else "不同模式"
    return f"根据模式判定当前系统状态，例如切换到{examples}，并同步上一模式记录"


def _looks_like_control_loop_body(body: str) -> bool:
    text = _safe_text(body)
    if not text:
        return False
    score = 0
    if re.search(r"(?:PosErr|Err_f|误差)", text):
        score += 1
    if re.search(r"(?:VelCmd|速度指令|PidCal)", text):
        score += 1
    if re.search(r"(?:CurUff|前馈|Uff)", text):
        score += 1
    if re.search(r"(?:DataTransFToInt|outMax|outMin)", text):
        score += 1
    return score >= 3


def _build_control_loop_logic_summary(body: str) -> str:
    if not _looks_like_control_loop_body(body):
        return ""
    return "围绕控制误差执行分支计算，生成速度指令并换算相关前馈量"


def _filter_glossary_for_prompt(glossary: Optional[dict[str, str]], texts: Sequence[str], *, limit: int = 18) -> dict[str, str]:
    data = glossary if isinstance(glossary, dict) else {}
    if not data:
        return {}
    haystack = "\n".join(_safe_text(t).lower() for t in (texts or []) if _safe_text(t))
    if not haystack:
        return dict(list(data.items())[:limit])
    picked: dict[str, str] = {}
    for k, v in data.items():
        key = _safe_strip(k)
        if (not key) or key.lower() not in haystack:
            continue
        picked[key] = _safe_strip(v)
        if len(picked) >= limit:
            break
    if picked:
        return picked
    return dict(list(data.items())[: min(limit, 6)])


def _symbol_kind_for_name(name: str, default: str = "symbols") -> str:
    ident = _safe_strip(name)
    if (not ident) or default == "function":
        return default
    if _is_macro_identifier(ident):
        return "macros"
    return default


def resolve_canonical_symbol_name(
    name: str,
    *,
    kind: str = "symbols",
    comment_cn: str = "",
    fallback: str = "",
    allow_guess: bool = True,
) -> str:
    from . import naming as naming_utils

    return naming_utils.resolve_canonical_symbol_name(
        name,
        kind=kind,
        comment_cn=comment_cn,
        fallback=fallback,
        allow_guess=allow_guess,
        backend_module=sys.modules[__name__],
    )


def _safe_relpath(path: str, root: str) -> str:
    try:
        return os.path.relpath(path, root)
    except Exception:
        return _safe_strip(path)


def _project_title_index_items() -> list[dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils._project_title_index_items(backend_module=sys.modules[__name__])


def _project_symbol_index_items() -> list[dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils._project_symbol_index_items(backend_module=sys.modules[__name__])


def _project_semantic_index_items() -> list[dict[str, Any]]:
    with _NAMING_INDEX_LOCK:
        data = copy.deepcopy(_PROJECT_SEMANTIC_INDEX_DATA)
    return list((_normalize_semantic_index_payload(data) or {}).get("items") or [])


def _project_semantic_record_maps() -> tuple[dict[str, dict[str, Any]], dict[tuple[str, str], dict[str, Any]], dict[str, list[dict[str, Any]]]]:
    by_id: dict[str, dict[str, Any]] = {}
    by_file_func: dict[tuple[str, str], dict[str, Any]] = {}
    by_func: dict[str, list[dict[str, Any]]] = {}
    for record in _project_semantic_index_items():
        record_id = _safe_strip(record.get("id"))
        if record_id:
            by_id[record_id] = record
        source_file = _safe_strip(record.get("source_file"))
        func_name = _safe_strip(record.get("func_name"))
        if source_file and func_name:
            by_file_func[(os.path.abspath(source_file), func_name)] = record
        if func_name:
            by_func.setdefault(func_name, []).append(record)
    return by_id, by_file_func, by_func


def _lookup_project_semantic_record(
    *,
    record_id: str = "",
    source_file: str = "",
    func_name: str = "",
    module_key: str = "",
    semantic_maps: Optional[tuple[dict[str, dict[str, Any]], dict[tuple[str, str], dict[str, Any]], dict[str, list[dict[str, Any]]]]] = None,
) -> dict[str, Any]:
    by_id, by_file_func, by_func = semantic_maps or _project_semantic_record_maps()
    record_id = _safe_strip(record_id)
    if record_id and record_id in by_id:
        return dict(by_id[record_id])

    source_file = _safe_strip(source_file)
    func_name = _safe_strip(func_name)
    module_key = _safe_strip(module_key)
    if source_file and func_name:
        hit = by_file_func.get((os.path.abspath(source_file), func_name))
        if hit:
            return dict(hit)

    if func_name:
        candidates = list(by_func.get(func_name) or [])
        if module_key:
            module_hit = next((item for item in candidates if _safe_strip(item.get("module_key")) == module_key), None)
            if module_hit:
                return dict(module_hit)
        if len(candidates) == 1:
            return dict(candidates[0])
    return {}


def _lightweight_semantic_record_from_body(
    *,
    func_name: str = "",
    source_file: str = "",
    module_key: str = "",
    family_prefix: str = "",
    ret_type: str = "",
    comment_desc: str = "",
    body: str = "",
    cfg: Optional["GenConfig"] = None,
) -> dict[str, Any]:
    func_name = _safe_strip(func_name)
    body = _safe_text(body)
    source_file = _safe_strip(source_file)
    project_root = _safe_strip(getattr(cfg, "project_root", "") if cfg is not None else "")
    return_exprs, return_symbols = _collect_return_semantics_from_body(body)
    if source_file:
        abs_source = os.path.abspath(source_file)
    else:
        abs_source = ""
    record_id = ""
    if abs_source and func_name:
        rel_root = project_root or os.path.dirname(abs_source)
        record_id = f"{_safe_relpath(abs_source, rel_root)}::{func_name}"
    return {
        "id": record_id,
        "source_file": abs_source,
        "module_key": _safe_strip(module_key),
        "func_name": func_name,
        "family_prefix": _safe_strip(family_prefix),
        "action_suffix": _identifier_action_suffix(func_name),
        "ret_type": _safe_strip(ret_type),
        "comment_desc": _safe_strip(comment_desc),
        "callee_names": _collect_callee_names_from_body(body),
        "macro_refs": _collect_unresolved_macro_candidates(body, known_names=set()),
        "condition_signatures": _collect_condition_signatures_from_body(body),
        "member_accesses": _collect_member_access_signatures_from_body(body),
        "return_exprs": return_exprs,
        "return_symbols": return_symbols,
        "written_params": tuple(),
        "read_params": tuple(),
        "symbol_profiles": [],
    }


def _resolve_current_function_semantic_record(func_data: dict, cfg: Optional["GenConfig"] = None) -> dict[str, Any]:
    record = _normalize_semantic_index_record((func_data or {}).get("semantic_record"))
    if record:
        return record

    func_info = (func_data or {}).get("func_info") or {}
    file_context = (func_data or {}).get("file_context") or {}
    func_name = _safe_strip(func_info.get("func_name"))
    source_file = _safe_strip(file_context.get("source_file"))
    module_key = _safe_strip(file_context.get("module_key"))
    semantic_maps = _project_semantic_record_maps()
    body = _safe_text((func_data or {}).get("body"))
    if body and func_name:
        built = {}
        if source_file:
            project_root = _safe_strip(getattr(cfg, "project_root", "") if cfg is not None else "")
            rel_root = project_root or os.path.dirname(source_file)
            built = _build_function_semantic_record(
                rel_root,
                source_file,
                {
                    "comment_info": dict((func_data or {}).get("comment_info") or {}),
                    "func_info": dict(func_info or {}),
                    "body": body,
                    "file_context": dict(file_context or {}),
                },
                cfg,
            ) or {}
        if built:
            return _normalize_semantic_index_record(built)
        return _lightweight_semantic_record_from_body(
            func_name=func_name,
            source_file=source_file,
            module_key=module_key,
            family_prefix=_safe_strip(file_context.get("family_prefix")),
            ret_type=_safe_strip(func_info.get("ret_type")),
            comment_desc=_safe_strip(((func_data or {}).get("comment_info") or {}).get("desc")),
            body=body,
            cfg=cfg,
        )

    return _lookup_project_semantic_record(
        source_file=source_file,
        func_name=func_name,
        module_key=module_key,
        semantic_maps=semantic_maps,
    )


def _lookup_semantic_symbol_profile(record: Optional[dict[str, Any]], symbol: str) -> dict[str, Any]:
    ident = _safe_strip(symbol)
    if not ident or not isinstance(record, dict):
        return {}
    for item in (record.get("symbol_profiles") or ()):
        profile = _normalize_semantic_symbol_profile(item)
        if profile and _safe_strip(profile.get("name")) == ident:
            return profile
    return {}


def _resolve_symbol_owner_semantic_record(symbol_record: dict, cfg: Optional["GenConfig"] = None) -> dict[str, Any]:
    owner_semantic = _normalize_semantic_index_record((symbol_record or {}).get("owner_semantic"))
    if owner_semantic:
        return owner_semantic

    source_file = _safe_strip((symbol_record or {}).get("source_file"))
    owner_func = _safe_strip((symbol_record or {}).get("owner_func"))
    module_key = _safe_strip((symbol_record or {}).get("module_key"))
    body = _safe_text((symbol_record or {}).get("body"))
    if body:
        return _lightweight_semantic_record_from_body(
            func_name=owner_func,
            source_file=source_file,
            module_key=module_key,
            family_prefix=_safe_strip((symbol_record or {}).get("family_prefix")),
            ret_type=_safe_strip((symbol_record or {}).get("owner_ret_type")),
            comment_desc=_safe_strip((symbol_record or {}).get("comment_desc")),
            body=body,
            cfg=cfg,
        )
    return _lookup_project_semantic_record(source_file=source_file, func_name=owner_func, module_key=module_key)


def _resolve_current_symbol_semantic_profile(symbol_record: dict) -> dict[str, Any]:
    profile = _normalize_semantic_symbol_profile((symbol_record or {}).get("symbol_profile"))
    if profile:
        return profile
    return _normalize_semantic_symbol_profile(
        {
            "name": _safe_strip((symbol_record or {}).get("symbol")),
            "scope": _safe_strip((symbol_record or {}).get("scope")),
            "decl_type": _safe_strip((symbol_record or {}).get("decl_type")),
            "role": _safe_strip((symbol_record or {}).get("role")),
            "direction": _safe_strip((symbol_record or {}).get("direction")),
            "producer_call": _safe_strip((symbol_record or {}).get("producer_call")),
            "producer_arg_tags": tuple((symbol_record or {}).get("producer_arg_tags") or ()),
            "consumer_patterns": tuple((symbol_record or {}).get("consumer_patterns") or ()),
            "usage_patterns": tuple((symbol_record or {}).get("usage_patterns") or ()),
            "paired_symbols": tuple((symbol_record or {}).get("paired_symbols") or ()),
        }
    )


def _dedupe_title_records(records: Sequence[tuple[int, dict[str, Any]]]) -> list[dict[str, Any]]:
    best: dict[str, tuple[int, dict[str, Any]]] = {}
    for score, record in records:
        title = _safe_strip(record.get("resolved_title"))
        key = title or _safe_strip(record.get("id"))
        cur = best.get(key)
        if cur is None or score > cur[0]:
            item = dict(record)
            item["score"] = int(score)
            best[key] = (score, item)
    return [item for _score, item in sorted(best.values(), key=lambda x: (-x[0], _safe_strip(x[1].get("resolved_title"))))]


def _dedupe_symbol_records(records: Sequence[tuple[int, dict[str, Any]]]) -> list[dict[str, Any]]:
    best: dict[str, tuple[int, dict[str, Any]]] = {}
    for score, record in records:
        key = _safe_strip(record.get("existing_cn")) or _safe_strip(record.get("symbol")) or _safe_strip(record.get("id"))
        cur = best.get(key)
        if cur is None or score > cur[0]:
            item = dict(record)
            item["score"] = int(score)
            best[key] = (score, item)
    return [item for _score, item in sorted(best.values(), key=lambda x: (-x[0], _safe_strip(x[1].get("existing_cn"))))]


def _normalize_condition_signature(text: str) -> str:
    s = _strip_balanced_outer_parens(_safe_strip(text)).replace("->", ".")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _collect_callee_names_from_body(body: str) -> tuple[str, ...]:
    hits: list[str] = []
    for raw in _join_c_line_continuations(body or "").splitlines():
        code, _ = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt:
            continue
        for match in re.finditer(r"\b([A-Za-z_]\w*)\s*\(", stmt):
            name = _safe_strip(match.group(1))
            if (not name) or (name in _C_KEYWORDS) or _is_macro_identifier(name):
                continue
            hits.append(name)
    return tuple(dict.fromkeys(hits))


def _collect_condition_signatures_from_body(body: str) -> tuple[str, ...]:
    hits: list[str] = []
    for raw in _join_c_line_continuations(body or "").splitlines():
        code, _ = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt:
            continue
        match = re.match(r"^(?:if|while|switch)\s*\((.+)\)\s*$", stmt)
        if match:
            sig = _normalize_condition_signature(match.group(1))
            if sig:
                hits.append(sig)
            continue
        match = re.match(r"^for\s*\(([^;]*);([^;]*);", stmt)
        if match:
            sig = _normalize_condition_signature(match.group(2))
            if sig:
                hits.append(sig)
    return tuple(dict.fromkeys(hits))


def _collect_member_access_signatures_from_body(body: str, *, limit: int = 24) -> tuple[str, ...]:
    hits: list[str] = []
    for raw in _join_c_line_continuations(body or "").splitlines():
        code, _ = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt:
            continue
        for match in _MEMBER_CHAIN_RE.finditer(stmt):
            base = _safe_strip(match.group("base")).replace("->", ".")
            rest = _safe_strip(match.group("rest")).replace("->", ".")
            if not base or not rest:
                continue
            sig = f"{base}.{rest}"
            if sig not in hits:
                hits.append(sig)
            if len(hits) >= max(1, limit):
                return tuple(hits)
    return tuple(hits)


def _collect_return_semantics_from_body(body: str) -> tuple[tuple[str, ...], tuple[str, ...]]:
    exprs: list[str] = []
    symbols: list[str] = []
    for raw in _join_c_line_continuations(body or "").splitlines():
        code, _ = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt:
            continue
        match = re.match(r"^return(?:\s+(.+?))?;?$", stmt)
        if not match:
            continue
        expr = _normalize_condition_signature(match.group(1))
        if expr:
            exprs.append(expr)
            if re.fullmatch(r"[A-Za-z_]\w*", expr) and expr not in _C_KEYWORDS:
                symbols.append(expr)
    return tuple(dict.fromkeys(exprs)), tuple(dict.fromkeys(symbols))


def _infer_param_direction_from_body(name: str, decl_type: str, body: str) -> str:
    ident = _safe_strip(name)
    decl = _safe_strip(decl_type)
    if not ident:
        return "input"
    pointer_like = ("*" in decl) or ident.startswith(("p_", "pp_", "vp_", "v_p_", "gp_", "lp_", "sp_", "cp_", "tp_"))
    read_hit = False
    write_hit = False
    for raw in _join_c_line_continuations(body or "").splitlines():
        code, _ = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt:
            continue
        if re.search(rf"\*\s*{re.escape(ident)}\s*=", stmt):
            write_hit = True
        if re.search(rf"\b{re.escape(ident)}\s*(?:->|\[)[^;=]*=", stmt):
            write_hit = True
        if re.search(rf"\b(?:memset|memcpy)\s*\(\s*{re.escape(ident)}\b", stmt):
            write_hit = True
        if re.search(rf"\b{re.escape(ident)}\b", stmt):
            read_hit = True
    if write_hit and read_hit:
        return "input_output"
    if write_hit:
        return "output"
    if not pointer_like:
        return "input"
    return "input"


def _build_function_semantic_record(
    project_root: str,
    c_path: str,
    func_data: dict,
    cfg: Optional["GenConfig"],
) -> Optional[dict[str, Any]]:
    func_info = (func_data or {}).get("func_info") or {}
    comment_info = (func_data or {}).get("comment_info") or {}
    body = _safe_text((func_data or {}).get("body"))
    file_context = (func_data or {}).get("file_context") or {}
    func_name = _safe_strip(func_info.get("func_name"))
    if not func_name:
        return None

    module_key = _safe_strip(file_context.get("module_key")) or _module_key_for_source(c_path)
    family_prefix = _safe_strip(file_context.get("family_prefix")) or _identifier_family_prefix(func_name)
    params = parse_params_from_prototype(func_info)
    local_vars = parse_local_variables_from_body(body)
    local_vars = _filter_local_vars_against_params(local_vars, params, cfg=cfg, func_name=func_name)
    symbol_items = list(params or []) + list(local_vars or [])
    neighbor_symbols = [
        _safe_strip((item or {}).get("name"))
        for item in symbol_items
        if _safe_strip((item or {}).get("name"))
    ]

    symbol_profiles: list[dict[str, Any]] = []
    for item in symbol_items:
        symbol = _safe_strip((item or {}).get("name"))
        if not symbol:
            continue
        scope = "param" if item in params else "local"
        evidence = collect_symbol_evidence(
            symbol,
            kind="symbols",
            body=body,
            decl_type=_safe_strip((item or {}).get("type")),
            neighbor_symbols=[x for x in neighbor_symbols if x and x != symbol],
            source_comment_hints=[_safe_strip((item or {}).get("comment_hint"))],
        )
        inference = _infer_symbol_semantics_rule(evidence)
        profile = {
            "name": symbol,
            "scope": scope,
            "decl_type": _safe_strip((item or {}).get("type")),
            "role": _safe_strip(inference.role),
            "direction": _infer_param_direction_from_body(symbol, _safe_strip((item or {}).get("type")), body) if scope == "param" else "local",
            "producer_call": _safe_strip(evidence.producer_call),
            "producer_arg_tags": tuple(evidence.producer_arg_tags or ()),
            "consumer_patterns": tuple(evidence.consumer_patterns or ()),
            "sink_patterns": tuple(evidence.sink_patterns or ()),
            "dataflow_roles": tuple(evidence.dataflow_roles or ()),
            "usage_patterns": tuple(evidence.usage_patterns or ()),
            "paired_symbols": tuple(evidence.paired_symbols or ()),
        }
        symbol_profiles.append(profile)

    written_params = tuple(
        profile["name"]
        for profile in symbol_profiles
        if profile.get("scope") == "param" and profile.get("direction") in {"output", "input_output"}
    )
    read_params = tuple(
        profile["name"]
        for profile in symbol_profiles
        if profile.get("scope") == "param" and profile.get("direction") in {"input", "input_output"}
    )
    return_exprs, return_symbols = _collect_return_semantics_from_body(body)

    comment_desc = _safe_strip(comment_info.get("desc"))
    if _is_noop_comment(comment_desc) or _looks_like_logic_noise_comment(comment_desc):
        comment_desc = ""
    return {
        "id": f"{_safe_relpath(c_path, project_root)}::{func_name}",
        "source_file": os.path.abspath(c_path),
        "module_key": module_key,
        "func_name": func_name,
        "family_prefix": family_prefix,
        "action_suffix": _identifier_action_suffix(func_name),
        "ret_type": _safe_strip(func_info.get("ret_type")),
        "comment_desc": comment_desc,
        "callee_names": _collect_callee_names_from_body(body),
        "macro_refs": _collect_unresolved_macro_candidates(body, known_names=set()),
        "condition_signatures": _collect_condition_signatures_from_body(body),
        "member_accesses": _collect_member_access_signatures_from_body(body),
        "return_exprs": return_exprs,
        "return_symbols": return_symbols,
        "written_params": written_params,
        "read_params": read_params,
        "symbol_profiles": symbol_profiles,
    }


def _rebuild_project_naming_indexes(project_root: str, cfg: Optional["GenConfig"]) -> tuple[dict[str, Any], dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils._rebuild_project_naming_indexes(project_root, cfg, backend_module=sys.modules[__name__])


def _should_refresh_semantic_index(
    project_root: str,
    semantic_path: str,
    cfg: Optional["GenConfig"],
) -> bool:
    mode = utils.cfg_get_str(cfg, "semantic_index_refresh", "auto").lower()
    if mode == "always":
        return True
    if mode == "off":
        return False
    if not os.path.isfile(semantic_path):
        return True
    max_files = max(0, utils.cfg_get_int(cfg, "semantic_index_max_files", 0))
    latest_src = _collect_project_source_mtime(project_root, max_files=max_files)
    try:
        semantic_mtime = os.path.getmtime(semantic_path)
    except Exception:
        return True
    return latest_src > semantic_mtime


def _rebuild_project_semantic_index(project_root: str, cfg: Optional["GenConfig"]) -> dict[str, Any]:
    items: list[dict[str, Any]] = []
    if not project_root or not os.path.isdir(project_root):
        return _normalize_semantic_index_payload({})

    worker_cfg = _clone_cfg(cfg, ai_assist=False) if isinstance(cfg, GenConfig) else GenConfig(ai_assist=False)
    source_files = _get_ordered_project_c_files(project_root, worker_cfg)
    max_files = max(0, utils.cfg_get_int(cfg, "semantic_index_max_files", 0))
    if max_files > 0:
        source_files = source_files[:max_files]

    for c_path in source_files:
        func_list, skip_reason = prepare_func_list_for_c_file(
            c_path,
            project_root=project_root,
            cfg=worker_cfg,
            prefilter=False,
        )
        if skip_reason or not func_list:
            continue
        for func_data in func_list:
            record = _build_function_semantic_record(project_root, c_path, func_data, worker_cfg)
            if record:
                items.append(record)

    return _normalize_semantic_index_payload({
        "version": 1,
        "items": items,
        "updated_at": datetime.datetime.now().isoformat(timespec="seconds"),
    })


def init_project_semantic_index(project_root: str, cfg: Optional["GenConfig"] = None) -> None:
    from . import semantic as semantic_utils

    semantic_utils.init_project_semantic_index(project_root, cfg)


def init_project_naming_indexes(project_root: str, cfg: Optional["GenConfig"] = None) -> None:
    from . import naming as naming_utils

    return naming_utils.init_project_naming_indexes(project_root, cfg, backend_module=sys.modules[__name__])


def retrieve_function_title_context(func_data: dict, cfg: Optional["GenConfig"] = None) -> list[dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils.retrieve_function_title_context(func_data, cfg, backend_module=sys.modules[__name__])


def retrieve_symbol_context(symbol_record: dict, cfg: Optional["GenConfig"] = None) -> list[dict[str, Any]]:
    from . import naming as naming_utils

    return naming_utils.retrieve_symbol_context(symbol_record, cfg, backend_module=sys.modules[__name__])


def _extract_symbol_usage_patterns(body: str, symbol: str) -> tuple[str, ...]:
    ident = _safe_strip(symbol)
    if (not ident) or (not body):
        return ()

    pat = re.compile(rf"\b{re.escape(ident)}\b")
    patterns: set[str] = set()
    for raw in _join_c_line_continuations(body).splitlines():
        code, _comments = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt or not pat.search(stmt):
            continue
        if re.search(rf"\breturn\s+{re.escape(ident)}\b", stmt):
            patterns.add("returned")
        if re.search(rf"\bif\s*\([^)]*\b{re.escape(ident)}\b", stmt) or re.search(rf"\bwhile\s*\([^)]*\b{re.escape(ident)}\b", stmt):
            patterns.add("condition")
        if re.search(rf"\b{re.escape(ident)}\s*=", stmt) and not re.search(rf"\b{re.escape(ident)}\s*==", stmt):
            patterns.add("assign_lhs")
            call_m = re.search(rf"\b{re.escape(ident)}\s*=\s*([A-Za-z_]\w*)\s*\(", stmt)
            if call_m:
                patterns.add("assigned_from_call")
                patterns.add(f"call_source:{call_m.group(1)}")
        assign_rhs = re.search(r"=\s*(.+)$", stmt)
        if assign_rhs and pat.search(assign_rhs.group(1) or ""):
            patterns.add("assign_rhs")
        # 只把单个 '&ident' 视为取地址，避免把 '&& ident' 误判成 address_of。
        if re.search(rf"(?<!&)&(?!&)\s*{re.escape(ident)}\b", stmt):
            patterns.add("address_of")
        if re.search(rf"\b{re.escape(ident)}\b\s*(\+\+|--)", stmt) or re.search(rf"\b{re.escape(ident)}\b\s*=\s*\b{re.escape(ident)}\b\s*[+\-]", stmt):
            patterns.add("accumulate")
        if re.search(rf"\b{re.escape(ident)}\b[^;]*[&|^~<>]", stmt):
            patterns.add("bitop")
        if re.search(rf"\b[A-Za-z_]\w*\s*\([^)]*\b{re.escape(ident)}\b", stmt):
            patterns.add("call_arg")
        if re.search(rf"\b{re.escape(ident)}\b\s*\[", stmt) or re.search(rf"\[\s*{re.escape(ident)}\s*\]", stmt):
            patterns.add("array")
    return tuple(sorted(patterns))


def _extract_symbol_sink_patterns(body: str, symbol: str) -> tuple[str, ...]:
    ident = _safe_strip(symbol)
    if (not ident) or (not body):
        return ()
    patterns: set[str] = set()
    pat = re.compile(rf"\b{re.escape(ident)}\b")
    joined = _join_c_line_continuations(body)
    clamp_re = re.compile(
        rf"if\s*\([^)]*\b{re.escape(ident)}\b\s*(?:>|>=|<|<=)\s*([A-Za-z_]\w*)[^)]*\)"
        rf"[\s\S]{{0,220}}?\b{re.escape(ident)}\b\s*=\s*\1\s*;",
        re.MULTILINE,
    )
    if clamp_re.search(joined):
        patterns.add("clamped_to_peer")
    for raw in joined.splitlines():
        code, _comments = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt or not pat.search(stmt):
            continue
        if re.search(rf"[*]\s*[A-Za-z_]\w*\s*=\s*{re.escape(ident)}\b", stmt):
            patterns.add("pointer_write")
        if re.search(
            rf"(?:\b[A-Za-z_]\w*(?:\[[^\]]+\])?\s*(?:\.|->)\s*[A-Za-z_]\w+)\s*=\s*{re.escape(ident)}\b",
            stmt,
        ):
            patterns.add("member_write")
            lowered = stmt.lower()
            if any(token in lowered for token in (".outmax", "->outmax", ".outmin", "->outmin")):
                patterns.add("pid_output_limit")
            if any(token in lowered for token in (".state", "->state", ".status", "->status", ".flag", "->flag", ".bit_", "->bit_")):
                patterns.add("state_member_write")
        if re.search(rf"\breturn\b[^;]*\b{re.escape(ident)}\b", stmt):
            patterns.add("return_flow")
    return tuple(sorted(patterns))


def _expr_pattern(expr: str) -> str:
    tokens = [re.escape(tok) for tok in re.split(r"\s+", _safe_strip(expr)) if tok]
    if not tokens:
        return ""
    return r"\s*".join(tokens)


def _extract_symbol_member_sources(body: str, symbol: str) -> tuple[str, ...]:
    ident = _safe_strip(symbol)
    if (not ident) or (not body):
        return ()
    joined = _join_c_line_continuations(body)
    pattern = re.compile(rf"\b{re.escape(ident)}\b\s*(?<![!<>=])=(?!=)\s*([^;]+);", re.MULTILINE)
    sources: list[str] = []
    for match in pattern.finditer(joined):
        expr = _safe_strip(match.group(1))
        if (not expr) or (ident in expr):
            continue
        if ("." not in expr) and ("->" not in expr):
            continue
        if expr not in sources:
            sources.append(expr)
    return tuple(sources[:4])


def _body_compares_symbol_with_expr(body: str, symbol: str, expr: str) -> tuple[bool, bool]:
    ident = _safe_strip(symbol)
    expr_pattern = _expr_pattern(expr)
    if (not ident) or (not expr_pattern) or (not body):
        return False, False
    joined = _join_c_line_continuations(body)
    neq = re.search(
        rf"\b{re.escape(ident)}\b\s*!=\s*{expr_pattern}|{expr_pattern}\s*!=\s*\b{re.escape(ident)}\b",
        joined,
        re.MULTILINE,
    )
    eq = re.search(
        rf"\b{re.escape(ident)}\b\s*==\s*{expr_pattern}|{expr_pattern}\s*==\s*\b{re.escape(ident)}\b",
        joined,
        re.MULTILINE,
    )
    return bool(neq), bool(eq)


def _extract_symbol_dataflow_roles(
    symbol: str,
    *,
    body: str,
    decl_type: str,
    usage_patterns: Sequence[str],
    consumer_patterns: Sequence[str],
    sink_patterns: Sequence[str],
    producer_call: str,
    producer_arg_tags: Sequence[str],
    paired_symbols: Sequence[str],
) -> tuple[str, ...]:
    ident = _safe_strip(symbol).lower()
    ident_raw = _safe_strip(symbol)
    decl_lower = _safe_strip(decl_type).lower()
    joined = _join_c_line_continuations(body)
    usage = set(usage_patterns or ())
    consumers = set(consumer_patterns or ())
    sinks = set(sink_patterns or ())
    tags = set(producer_arg_tags or ())
    roles: set[str] = set()
    if "pid_output_limit" in sinks:
        roles.add("output_limit")
    if "clamped_to_peer" in sinks:
        roles.add("clamp_result")
    if "pointer_write" in sinks:
        roles.add("output_value")
    if "state_member_write" in sinks:
        roles.add("state_output")
    if "return_flow" in sinks or "returned" in usage or "returned_directly" in consumers:
        roles.add("return_value")
    if producer_call and "assign_lhs" in usage:
        roles.add("call_result")
    if "results_bit32" in tags:
        roles.add("bit_result")
        if "compared_to_static_prev" in consumers or "used_in_change_detection" in consumers:
            roles.add("snapshot")
    if paired_symbols and ident.startswith(("l_s_", "s_")):
        roles.add("previous_snapshot")
    member_sources = _extract_symbol_member_sources(body, symbol)
    if member_sources:
        roles.add("member_snapshot")
        for expr in member_sources:
            expr_lower = expr.lower()
            neq_match, eq_match = _body_compares_symbol_with_expr(joined, symbol, expr)
            if neq_match or eq_match:
                roles.add("snapshot")
            if neq_match:
                roles.add("previous_snapshot")
                roles.add("change_baseline")
            if any(tag in expr_lower for tag in (".bit_t.", ".state", "->state", ".status", "->status", ".flag", "->flag", "_v_b1", "_state_", "_status_")):
                roles.add("state_snapshot")
                roles.add("state_value")
    if "accumulate" in usage:
        if any(tag in decl_lower for tag in ("uint", "int", "short", "long", "char")):
            roles.add("counter_value")
        elif any(tag in decl_lower for tag in ("float", "double", "iq", "fq")):
            roles.add("accumulator")
    if re.search(rf"\b{re.escape(ident_raw)}\b\s*=\s*0(?:[uUlLfF]*)\b", joined):
        if "counter_value" in roles:
            roles.add("resettable_counter")
    if re.search(rf"\b{re.escape(ident_raw)}\b\s*(?:>|>=|<|<=)\s*\d", joined) and any(tag in decl_lower for tag in ("uint", "int", "short", "long", "char")):
        roles.add("counter_value")
    if "state_field_read" in consumers or "state" in tags or "status" in tags:
        roles.add("state_value")
    if any(tag in ident for tag in ("limit", "limt", "lmt")) and ("output_limit" in roles or "clamp_result" in roles):
        roles.add("limited_value")
    if any(tag in ident for tag in ("brk", "break")) and ("output_limit" in roles or "clamp_result" in roles):
        roles.add("brake_limit")
    if "eval" in decl_lower:
        roles.add("evaluated_value")
    return tuple(sorted(roles))


def _infer_symbol_role(evidence: SymbolEvidence) -> str:
    ident = _safe_strip(evidence.symbol).lower()
    patterns = set(evidence.usage_patterns or ())
    dataflow_roles = set(evidence.dataflow_roles or ())
    decl_type = _safe_strip(evidence.decl_type).lower()
    if "counter_value" in dataflow_roles or "resettable_counter" in dataflow_roles:
        return "计数器"
    if "state_snapshot" in dataflow_roles or "previous_snapshot" in dataflow_roles:
        return "上一周期值"
    if "previous_snapshot" in dataflow_roles:
        return "上一周期值"
    if "returned" in patterns:
        return "返回值"
    if "output_limit" in dataflow_roles or "limited_value" in dataflow_roles:
        return "阈值"
    if "state_value" in dataflow_roles or "state_output" in dataflow_roles:
        return "状态"
    if "snapshot" in dataflow_roles:
        return "缓存值"
    call_source = _call_source_name_from_patterns(evidence.usage_patterns or ())
    call_lower = call_source.lower()
    if call_source and "assign_lhs" in patterns:
        if any(tag in call_lower for tag in ("stateget", "statusget", "state", "status")):
            return "状态"
        if any(tag in call_lower for tag in ("resultget", "check", "test", "result")):
            return "返回值"
        return "缓存值"
    if "*" in decl_type or "address_of" in patterns:
        return "指针"
    if any(tag in ident for tag in ("idx", "index", "_ii_", "_jj_", "_kk_")):
        return "索引"
    # ii/jj/kk 作为独立变量名或后缀
    if re.search(r"(?:^|_)(ii|jj|kk)(?:_|$)", ident):
        return "索引"
    if any(tag in ident for tag in ("cnt", "count", "tick")):
        return "计数器"
    if any(tag in ident for tag in ("last", "prev", "pre")):
        return "上一周期值"
    if any(tag in ident for tag in ("curr", "current", "cur")):
        return "当前值"
    if any(tag in ident for tag in ("flag", "enable", "en", "ok", "pass")):
        return "标志"
    if any(tag in ident for tag in ("mode",)):
        return "模式"
    if any(tag in ident for tag in ("state", "stat", "sta", "status")):
        return "状态"
    if any(tag in ident for tag in ("limit", "lmt", "threshold", "thd")):
        return "阈值"
    if any(tag in ident for tag in ("result", "results", "ret")):
        return "返回值"
    if any(tag in ident for tag in ("data", "value", "val")):
        return "缓存值"
    if any(tag in ident for tag in ("buff", "buf", "rx", "tx")):
        return "缓存值"
    if any(tag in ident for tag in ("sum", "chk", "crc")):
        return "校验值"
    if any(tag in ident for tag in ("addr", "offset")):
        return "偏移"
    if any(tag in ident for tag in ("temp", "tmp")):
        return "中间量"
    return "中间量"


def _infer_macro_role(name: str) -> str:
    ident = _safe_strip(name).upper()
    if any(token in ident for token in ("RATIO", "SCALE", "GAIN")):
        return "换算系数"
    if any(token in ident for token in ("BIAS", "OFFSET", "ADDR", "ADDRESS")):
        return "偏移"
    if any(token in ident for token in ("TIMEOUT", "TIME", "TICK", "_MS", "_US", "_NS")):
        return "时间阈值"
    if any(token in ident for token in ("REG", "FPGA", "XINT")):
        return "寄存器"
    if any(token in ident for token in ("BIT", "FLAG", "MASK", "_EN", "FAULT", "ERR")):
        return "位标志"
    return "宏定义"


def _default_cn_for_role(role: str) -> str:
    return {
        "返回值": "返回结果",
        "缓存值": "缓存值",
        "指针": "数据指针",
        "索引": "索引",
        "计数器": "计数器",
        "上一周期值": "上一周期值",
        "当前值": "当前值",
        "标志": "标志位",
        "模式": "模式字",
        "状态": "状态字",
        "阈值": "阈值",
        "中间量": "",
        "换算系数": "换算系数",
        "偏移": "偏移量",
        "时间阈值": "超时阈值",
        "寄存器": "寄存器",
        "位标志": "位标志",
        "宏定义": "宏定义",
    }.get(role, "")


def _extract_call_assignment_profile(body: str, symbol: str) -> tuple[str, str, tuple[str, ...]]:
    ident = _safe_strip(symbol)
    if (not ident) or (not body):
        return "", "", ()
    pat = re.compile(rf"\b{re.escape(ident)}\b\s*(?<![!<>=])=(?!=)\s*([A-Za-z_]\w*)\s*\((.*?)\)\s*;?")
    for raw in _join_c_line_continuations(body).splitlines():
        code, _ = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt:
            continue
        m = pat.search(stmt)
        if not m:
            continue
        call_name = _safe_strip(m.group(1))
        args_raw = _safe_strip(m.group(2))
        args = tuple(
            _safe_strip(part)
            for part in re.split(r",(?![^()]*\))", args_raw)
            if _safe_strip(part)
        )
        if call_name:
            return "call_return", call_name, args
    return "", "", ()


def _producer_arg_tags(call_name: str, args: Sequence[str]) -> tuple[str, ...]:
    tags: set[str] = set()
    call_upper = _safe_strip(call_name).upper()
    for arg in (args or ()):
        text = _safe_strip(arg)
        upper = text.upper()
        if not upper:
            continue
        if "RESULTS_BIT32" in upper:
            tags.add("results_bit32")
        if "RESULTS" in upper:
            tags.add("results")
        if "FLEVEL" in upper or "FAULT_LEVEL" in upper:
            tags.add("fault_level")
        if "STATE" in upper or "STATUS" in upper:
            tags.add("state")
        if "RIU" in upper:
            tags.add("riu")
        if "CCDL" in upper:
            tags.add("ccdl")
        if "KZZZ" in upper:
            tags.add("kzzz")
        if "LEFT" in upper or "_L_" in upper:
            tags.add("left")
        if "RIGHT" in upper or "_R_" in upper:
            tags.add("right")
        if "IFBIT" in upper:
            tags.add("ifbit")
        if "MBIT" in upper:
            tags.add("mbit")
        if "PUBIT" in upper or "PUBIT" in upper:
            tags.add("pubit")
    if "STATEGET" in call_upper or "STATUSGET" in call_upper:
        tags.add("state")
    if "RESULTGET" in call_upper:
        tags.add("results")
    return tuple(sorted(tags))


def _extract_symbol_consumer_patterns(body: str, symbol: str) -> tuple[str, ...]:
    ident = _safe_strip(symbol)
    if (not ident) or (not body):
        return ()
    patterns: set[str] = set()
    pat = re.compile(rf"\b{re.escape(ident)}\b")
    prev_pair_re = re.compile(rf"\b([A-Za-z_]\w*)\b\s*!=\s*\b{re.escape(ident)}\b|\b{re.escape(ident)}\b\s*!=\s*\b([A-Za-z_]\w*)\b")
    for raw in _join_c_line_continuations(body).splitlines():
        code, _ = _split_code_and_comments_for_symbol(raw)
        stmt = _safe_strip(code)
        if not stmt or not pat.search(stmt):
            continue
        if re.search(rf"\bif\s*\([^)]*\b{re.escape(ident)}\b", stmt):
            patterns.add("used_in_condition")
        if re.search(rf"\bif\s*\([^)]*(?:!=|==)[^)]*\b{re.escape(ident)}\b", stmt):
            patterns.add("used_in_change_detection")
        mm = prev_pair_re.search(stmt)
        if mm:
            other = _safe_strip(mm.group(1) or mm.group(2))
            if other.startswith(("l_s_", "s_")):
                patterns.add("compared_to_static_prev")
        if re.search(rf"\breturn\s+{re.escape(ident)}\b", stmt):
            patterns.add("returned_directly")
        if re.search(rf"\b{re.escape(ident)}\b\s*\.\s*dataState_u16\b|\b{re.escape(ident)}\b\s*->\s*dataState_u16\b", stmt):
            patterns.add("state_field_read")
    return tuple(sorted(patterns))


def _extract_paired_symbols(symbol: str, neighbor_symbols: Sequence[str]) -> tuple[str, ...]:
    ident = _safe_strip(symbol)
    if not ident:
        return ()
    stem = re.sub(r"^(?:l_s_|l_|s_)", "", ident)
    if not stem:
        return ()
    pairs: list[str] = []
    for raw in (neighbor_symbols or ()):
        other = _safe_strip(raw)
        if not other or other == ident:
            continue
        other_stem = re.sub(r"^(?:l_s_|l_|s_)", "", other)
        if other_stem == stem:
            pairs.append(other)
    return tuple(pairs[:4])


def _bit_family_cn_from_text(text: str) -> str:
    upper = _safe_strip(text).upper()
    if "IFBIT" in upper:
        return "周期自检"
    if "MBIT" in upper:
        return "维护自检"
    if "PUBIT" in upper or "PBIT" in upper:
        return "上电自检"
    return ""


def _candidate_concepts_from_evidence(evidence: SymbolEvidence) -> tuple[str, ...]:
    from . import semantic as semantic_utils

    return semantic_utils.candidate_concepts_from_evidence(
        evidence,
        backend_module=sys.modules[__name__],
    )


def collect_symbol_evidence(
    symbol: str,
    *,
    kind: str = "symbols",
    body: str = "",
    decl_type: str = "",
    owner_type: str = "",
    neighbor_symbols: Optional[Sequence[str]] = None,
    source_comment_hints: Optional[Sequence[str]] = None,
) -> SymbolEvidence:
    from . import semantic as semantic_utils

    return semantic_utils.collect_symbol_evidence(
        symbol,
        kind=kind,
        body=body,
        decl_type=decl_type,
        owner_type=owner_type,
        neighbor_symbols=neighbor_symbols,
        source_comment_hints=source_comment_hints,
        backend_module=sys.modules[__name__],
    )


def _infer_symbol_semantics_rule(evidence: SymbolEvidence) -> SymbolInference:
    from . import semantic as semantic_utils

    return semantic_utils.infer_symbol_semantics_rule(
        evidence,
        backend_module=sys.modules[__name__],
    )


def _build_symbol_inference_prompt(evidence: SymbolEvidence, cfg: Optional["GenConfig"] = None) -> str:
    from . import semantic as semantic_utils

    return semantic_utils.build_symbol_inference_prompt(
        evidence,
        cfg,
        backend_module=sys.modules[__name__],
    )


def infer_symbol_semantics(
    evidence: SymbolEvidence,
    cfg: Optional["GenConfig"] = None,
) -> SymbolInference:
    from . import semantic as semantic_utils

    return semantic_utils.infer_symbol_semantics(
        evidence,
        cfg,
        backend_module=sys.modules[__name__],
    )


def infer_scope_symbol_names(
    local_vars: Sequence[dict],
    params: Sequence[dict],
    *,
    body: str,
    func_info: Optional[dict],
    comment_info: Optional[dict],
    in_map: dict[str, str],
    out_map: dict[str, str],
    cfg: Optional["GenConfig"],
) -> dict[str, SymbolInference]:
    from . import semantic as semantic_utils

    return semantic_utils.infer_scope_symbol_names(
        local_vars,
        params,
        body=body,
        func_info=func_info,
        comment_info=comment_info,
        in_map=in_map,
        out_map=out_map,
        cfg=cfg,
        backend_module=sys.modules[__name__],
    )


def build_local_symbol_profile(
    item: dict,
    *,
    body: str,
    neighbor_symbols: Sequence[str] = (),
    scope: str = "local",
    comment_desc: str = "",
    cfg: Optional["GenConfig"] = None,
):
    from . import semantic as semantic_utils

    return semantic_utils.build_local_symbol_profile(
        item,
        body=body,
        neighbor_symbols=neighbor_symbols,
        scope=scope,
        comment_desc=comment_desc,
        cfg=cfg,
        backend_module=sys.modules[__name__],
    )


def build_function_local_symbol_profiles(
    local_vars: Sequence[dict],
    params: Sequence[dict] = (),
    *,
    body: str,
    comment_desc: str = "",
    cfg: Optional["GenConfig"] = None,
):
    from . import semantic as semantic_utils

    return semantic_utils.build_function_local_symbol_profiles(
        local_vars,
        params,
        body=body,
        comment_desc=comment_desc,
        cfg=cfg,
        backend_module=sys.modules[__name__],
    )


def _seed_symbol_memory_into_scope(
    comment_info: dict,
    func_info: dict,
    local_vars: Sequence[dict],
    params: Sequence[dict],
    in_map: dict,
    out_map: dict,
    param_ai_name_map: Optional[dict[str, str]] = None,
) -> dict[str, str]:
    seeded_params = param_ai_name_map if isinstance(param_ai_name_map, dict) else {}

    func_name = _safe_strip((func_info or {}).get("func_name"))
    if func_name and _is_missing_gap_text((comment_info or {}).get("func_cn_name") or ""):
        from . import naming as naming_utils

        comment_desc = _safe_strip((comment_info or {}).get("desc"))
        compact_comment_title = naming_utils._extract_compact_function_title(comment_desc)
        if not compact_comment_title:
            func_cn = _lookup_symbol_dictionary(func_name)
            if func_cn:
                comment_info["func_cn_name"] = func_cn

    for v in (local_vars or []):
        name = _safe_strip((v or {}).get("name"))
        if (not name) or _safe_strip((v or {}).get("cn_name")):
            continue
        cn = _lookup_symbol_dictionary(name)
        if cn:
            v["cn_name"] = cn

    for p in (params or []):
        name = _safe_strip((p or {}).get("name"))
        if (not name) or _safe_strip(in_map.get(name) or out_map.get(name)):
            continue
        cn = _lookup_symbol_dictionary(name)
        if not cn:
            continue
        in_map[name] = cn
        seeded_params[name] = cn
    return seeded_params


def apply_domain_glossary_overrides(overrides: Optional[dict[str, str]]) -> None:
    """
    重置/覆盖运行时 DOMAIN_GLOSSARY（不影响默认逻辑：默认=DOMAIN_GLOSSARY_BASE）。
    注意：调用后会清空中文名缓存，确保新术语立即生效。
    """
    try:
        DOMAIN_GLOSSARY.clear()
        DOMAIN_GLOSSARY.update(DOMAIN_GLOSSARY_BASE)
        if isinstance(overrides, dict) and overrides:
            for k, v in overrides.items():
                kk = str(k).strip()
                vv = str(v).strip()
                if kk and vv:
                    DOMAIN_GLOSSARY[kk] = vv
    except Exception:
        pass
    try:
        _FUNC_CN_CACHE.clear()
    except Exception:
        pass


apply_symbol_dictionary_overrides({})


#
# ----------------------
#  解析/渲染解耦：模型与异常
# ----------------------

# =============== 异常类与配置（从 config.py 导入） ===============
from .config import (
    GenConfig,
    NoDataError,
    ParseError,
    RenderError,
    SkipFunctionError,
    SourceReadError,
    ToolError,
    _normalize_ai_profile_label,
)

# =============== 数据类（models.py 回退） ===============


if StableIOElement is None:
    @dataclass(frozen=True)
    class IOElement:
        name: str
        ident: str
        c_type: str
        direction: Literal["输入", "输出", "输入/输出"]
else:
    IOElement = StableIOElement


if StableLocalDataElement is None:
    @dataclass(frozen=True)
    class LocalDataElement:
        name: str
        ident: str
        c_type: str
        usage: str
else:
    LocalDataElement = StableLocalDataElement


try:
    from .models import AIBuildMeta as _StableAIBuildMeta
except ImportError:
    _StableAIBuildMeta = None

if _StableAIBuildMeta is not None:
    AIBuildMeta = _StableAIBuildMeta
else:
    @dataclass(frozen=True)
    class AIBuildMeta:
        ai_enabled: bool = False
        ai_failed: bool = False
        regression_needed: bool = False
        regression_round: int = 0
        regression_reasons: tuple[str, ...] = ()
        logic_placeholders: int = 0
        unresolved_local_symbols: tuple[str, ...] = ()
        unresolved_param_symbols: tuple[str, ...] = ()
        unresolved_logic_symbols: tuple[str, ...] = ()
        generic_logic_count: int = 0
        comment_leak_count: int = 0
        term_drift_count: int = 0
        over_translation_count: int = 0
        bad_symbol_guess_count: int = 0
        raw_func_title: str = ""
        pre_rerank_func_title: str = ""
        title_candidates: tuple[str, ...] = ()
        title_pattern: str = ""
        title_rerank_changed: bool = False
        title_fallback_used: bool = False
        title_model_confidence: float = 0.0
        title_retry_used: bool = False
        title_stage_debug: tuple[dict[str, Any], ...] = ()
        logic_source_audit: tuple[dict[str, Any], ...] = ()
        quality_issues: tuple[dict[str, Any], ...] = ()
        quality_recovery: tuple[dict[str, Any], ...] = ()


if StableCommentHint is None:
    @dataclass(frozen=True)
    class CommentHint:
        kind: Literal["action", "condition", "purpose", "constraint", "history", "debug", "noise"]
        text: str
        confidence: float = 0.0
else:
    CommentHint = StableCommentHint


if StableSymbolEvidence is None:
    @dataclass(frozen=True)
    class SymbolEvidence:
        symbol: str
        kind: str
        decl_type: str = ""
        owner_type: str = ""
        usage_patterns: tuple[str, ...] = ()
        consumer_patterns: tuple[str, ...] = ()
        sink_patterns: tuple[str, ...] = ()
        dataflow_roles: tuple[str, ...] = ()
        neighbor_symbols: tuple[str, ...] = ()
        paired_symbols: tuple[str, ...] = ()
        source_comment_hints: tuple[str, ...] = ()
        normalized_comment_hint: str = ""
        producer_kind: str = ""
        producer_call: str = ""
        producer_args: tuple[str, ...] = ()
        producer_arg_tags: tuple[str, ...] = ()
        preferred_cn: str = ""
        memory_cn: str = ""
else:
    SymbolEvidence = StableSymbolEvidence


if StableSymbolInference is None:
    @dataclass(frozen=True)
    class SymbolInference:
        symbol: str
        kind: str
        candidate_cn: str = ""
        role: str = ""
        confidence: float = 0.0
        evidence_kinds: int = 0
        persist_scope: str = "off"
        reason: str = ""
else:
    SymbolInference = StableSymbolInference


if StableFunctionDesign is None:
    @dataclass(frozen=True)
    class FunctionDesign:
        title: str
        req_id: str
        prototype: str
        description_lines: tuple[str, ...]
        io_elements: tuple[IOElement, ...]
        io_none: bool
        local_elements: Optional[tuple[LocalDataElement, ...]]  # None=略；()=无
        logic_lines: Optional[tuple[str, ...]]                  # None=略/无
        ai_meta: AIBuildMeta = AIBuildMeta()
else:
    FunctionDesign = StableFunctionDesign

def normalize_req_prefix(prefix: str) -> str:
    """去掉用户输入前缀尾部多余的下划线，避免出现双下划线。"""
    return (prefix or "").rstrip("_")

def normalize_docx_output_path(path: str, *, ensure_parent_dir: bool) -> str:
    """
    规范化输出 docx 路径：
    - 修正常见的绝对路径漏写前导 /（macOS：Users/... -> /Users/...）
    - 自动补齐 .docx 扩展名（1docx -> 1.docx；xxx -> xxx.docx）
    - 可选：自动创建父目录（ensure_parent_dir=True）
    """
    raw = (path or "").strip()
    if not raw:
        raise ValueError("输出路径为空")

    p = os.path.expanduser(raw)

    # macOS 常见误输入：Users/...（漏了前导 /）
    if sys.platform == "darwin" and (not os.path.isabs(p)) and p.startswith("Users/"):
        # 若当前目录下真的存在 Users/，则尊重相对路径
        if not os.path.exists(os.path.join(os.getcwd(), "Users")):
            p = "/" + p

    # 自动补齐扩展名
    base = os.path.basename(p)
    if not base:
        raise ValueError("输出路径无文件名，请指定 .docx 文件名")
    if os.path.isdir(p):
        raise ValueError("输出路径是目录，请指定 .docx 文件名")
    lower = base.lower()
    if lower.endswith("docx") and (not lower.endswith(".docx")) and ("." not in base):
        p = p[:-4] + ".docx"
    elif not lower.endswith(".docx"):
        p = p + ".docx"

    p = os.path.abspath(p)
    parent = os.path.dirname(p) or "."
    if ensure_parent_dir and parent and (not os.path.exists(parent)):
        os.makedirs(parent, exist_ok=True)
    return p


def derive_software_unit_output_path(main_output: str) -> str:
    """
    根据主输出路径生成"软件单元清单"的输出路径（ASCII 命名）。
    例：/a/b/out.docx -> /a/b/out_unit_table.docx
    """
    base, ext = os.path.splitext(main_output)
    if not ext:
        ext = ".docx"
    return base + "_unit_table" + ext


def derive_unit_function_list_output_path(main_output: str) -> str:
    """
    根据主输出路径生成"单元函数列表"的输出路径（ASCII 命名）。
    例：/a/b/out.docx -> /a/b/out_unit_func_list.docx
    """
    base, ext = os.path.splitext(main_output)
    if not ext:
        ext = ".docx"
    return base + "_unit_func_list" + ext


def _get_ai_capability_profile(cfg: Optional[GenConfig]) -> str:
    from . import ai as ai_utils

    return ai_utils._get_ai_capability_profile(cfg)


def _small_model_prompt_mode(cfg: Optional[GenConfig]) -> str:
    from . import ai as ai_utils

    return ai_utils._small_model_prompt_mode(
        cfg,
        backend_module=sys.modules[__name__],
    )


def _is_small_model_strict_mode(cfg: Optional[GenConfig]) -> bool:
    from . import ai as ai_utils

    return ai_utils._is_small_model_strict_mode(
        cfg,
        backend_module=sys.modules[__name__],
    )


def _small_model_bool(cfg: Optional[GenConfig], key: str, default: int) -> bool:
    from . import ai as ai_utils

    return ai_utils._small_model_bool(
        cfg,
        key,
        default,
        backend_module=sys.modules[__name__],
    )

# =============== 公共工具 ===============

import datetime

# =============================
#      日志函数（核心）
# =============================
def vlog(cfg: GenConfig, *args):
    return utils.vlog(cfg, *args)


def should_log_step(cfg: GenConfig, step: int) -> bool:
    return utils.should_log_step(cfg, step)


def write_error_log(event: str, payload: dict) -> None:
    return utils.write_error_log(event, payload)






def ai_debug_log(cfg: GenConfig, event: str, payload: dict):
    return utils.ai_debug_log(cfg, event, payload)

def stop_requested(cfg: GenConfig) -> bool:
    return utils.stop_requested(cfg)


def gui_event(cfg: Optional[GenConfig], payload: dict) -> None:
    return utils.gui_event(cfg, payload)





def load_c_file(filepath: str) -> str:
    """按多种编码尝试读取 C 源文件。"""
    if not os.path.exists(filepath):
        raise SourceReadError(f"找不到文件：{filepath}")

    encodings = ['utf-8', 'gb18030', 'gbk', 'cp936']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                content = f.read()
            # print(f"成功使用 {enc} 编码读取文件。")
            return content
        except UnicodeDecodeError:
            continue
    raise SourceReadError(f"无法识别文件编码：{filepath}")


def add_seq_field(paragraph, seq_name: str):
    from . import render as render_utils

    return render_utils.add_seq_field(paragraph, seq_name)

def add_field_instr(paragraph, instr: str, placeholder: str = "1"):
    """
    在段落中插入一个域，instr 直接写完整指令，例如：
      'SEQ MOD \\# "000"'
      'SEQ MOD \\c \\# "000"'
      'SEQ FUNC \\s 3 \\# "000"'
    """
    r = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fld_char_begin)

    r = paragraph.add_run()
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f" {instr} "
    r._r.append(instr_text)

    r = paragraph.add_run()
    fld_char_sep = OxmlElement('w:fldChar')
    fld_char_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fld_char_sep)

    paragraph.add_run(placeholder)

    r = paragraph.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_char_end)


# =============== 注释与函数解析 ===============

def find_comment_blocks(code: str):
    from . import parse as parse_utils

    return parse_utils.find_comment_blocks(code)


def parse_single_comment_block(raw: str) -> dict:
    from . import parse as parse_utils

    return parse_utils.parse_single_comment_block(raw)


def _parse_line_comment_block(raw_lines: Sequence[str]) -> dict:
    from . import parse as parse_utils

    return parse_utils._parse_line_comment_block(list(raw_lines or []))


def _find_line_comment_blocks(code: str) -> list[dict]:
    from . import parse as parse_utils

    return parse_utils._find_line_comment_blocks(code)


def _clean_comment_line(line: str) -> str:
    from . import parse as parse_utils

    return parse_utils._clean_comment_line(line)


def _extract_file_header_info(code: str) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils._extract_file_header_info(code)


def _extract_module_cn_from_header(code: str) -> str:
    from . import parse as parse_utils

    return parse_utils._extract_module_cn_from_header(code)


def _derive_module_display_name(c_path: str, code: str) -> str:
    from . import parse as parse_utils

    return parse_utils._derive_module_display_name(c_path, code)


_PP_LITERAL_EXPR_RE = re.compile(r"^(0|1)(?:[uUlL]+)?$")


def _mask_non_newline_chars(text: str) -> str:
    return re.sub(r"[^\r\n]", " ", text or "")


def _strip_inactive_preprocessor_regions_keep_layout(code: str) -> str:
    """
    对可静态判定的预处理分支做"保长度消隐"：
    - 处理 #if 0 / #if 1 / #elif 0 / #elif 1 / #else / #endif
    - 仅消隐不可达分支和条件控制行本身，保留行号与字符位置
    - 对 #ifdef / #ifndef / #if SOME_MACRO 等未知条件保持原样，不做激进裁剪
    """
    if not code:
        return ""

    def _strip_expr_comments(expr: str) -> str:
        out = re.sub(r"/\*.*?\*/", "", expr or "")
        out = re.sub(r"//.*", "", out)
        return out.strip()

    def _trim_outer_parens(expr: str) -> str:
        s = re.sub(r"\s+", "", expr or "")
        while s.startswith("(") and s.endswith(")"):
            depth = 0
            ok = True
            for idx, ch in enumerate(s):
                if ch == "(":
                    depth += 1
                elif ch == ")":
                    depth -= 1
                    if depth < 0:
                        ok = False
                        break
                    if depth == 0 and idx != len(s) - 1:
                        ok = False
                        break
            if not ok or depth != 0:
                break
            s = s[1:-1].strip()
        return s

    def _parse_literal_flag(expr: str) -> Optional[bool]:
        compact = _trim_outer_parens(_strip_expr_comments(expr))
        if not compact:
            return None
        m = _PP_LITERAL_EXPR_RE.fullmatch(compact)
        if not m:
            return None
        return bool(int(m.group(1)))

    lines = code.splitlines(keepends=True)
    out_lines: list[str] = []
    stack: list[dict[str, Any]] = []

    for raw in lines:
        stripped = raw.lstrip()
        m = re.match(r"#\s*(if|elif|else|endif)\b(.*)", stripped)
        if m:
            directive = (m.group(1) or "").strip()
            expr = (m.group(2) or "").strip()

            if directive == "if":
                parent_active = stack[-1]["current_active"] if stack else True
                literal = _parse_literal_flag(expr)
                if literal is None:
                    stack.append({
                        "parent_active": parent_active,
                        "current_active": parent_active,
                        "eval_known": False,
                        "taken": False,
                    })
                else:
                    stack.append({
                        "parent_active": parent_active,
                        "current_active": parent_active and literal,
                        "eval_known": True,
                        "taken": literal,
                    })
            elif directive == "elif":
                if stack:
                    frame = stack[-1]
                    if frame.get("eval_known"):
                        literal = _parse_literal_flag(expr)
                        if literal is None:
                            frame["eval_known"] = False
                            frame["current_active"] = frame.get("parent_active", True)
                        else:
                            branch_active = (not frame.get("taken", False)) and literal
                            frame["current_active"] = frame.get("parent_active", True) and branch_active
                            frame["taken"] = bool(frame.get("taken", False) or literal)
                    else:
                        frame["current_active"] = frame.get("parent_active", True)
            elif directive == "else":
                if stack:
                    frame = stack[-1]
                    if frame.get("eval_known"):
                        branch_active = not frame.get("taken", False)
                        frame["current_active"] = frame.get("parent_active", True) and branch_active
                        frame["taken"] = True
                    else:
                        frame["current_active"] = frame.get("parent_active", True)
            elif directive == "endif":
                if stack:
                    stack.pop()

            out_lines.append(_mask_non_newline_chars(raw))
            continue

        current_active = stack[-1]["current_active"] if stack else True
        out_lines.append(raw if current_active else _mask_non_newline_chars(raw))

    return "".join(out_lines)


def _strip_c_comments_keep_layout(code: str) -> str:
    """
    删除 C/CPP 注释，但尽量保留原始换行和字符宽度，避免影响后续正则定位。
    """
    if not code:
        return ""

    def _repl_block(m: re.Match) -> str:
        text = m.group(0)
        return re.sub(r"[^\n]", " ", text)

    out = re.sub(r"/\*[\s\S]*?\*/", _repl_block, code)
    out = re.sub(r"//[^\n\r]*", lambda m: " " * len(m.group(0)), out)
    return out


def find_function_prototypes(code: str):
    from . import parse as parse_utils

    return parse_utils.find_function_prototypes(code)


def extract_function_body(code: str, brace_start_index: int) -> str:
    from . import parse as parse_utils

    return parse_utils.extract_function_body(code, brace_start_index)


def extract_nearby_typedefs(code: str, before_index: int, max_blocks: int = 3):
    from . import parse as parse_utils

    return parse_utils.extract_nearby_typedefs(code, before_index, max_blocks=max_blocks)


def _extract_typedef_blocks_from_code(code: str, max_blocks: int = 24) -> list[str]:
    from . import parse as parse_utils

    return parse_utils._extract_typedef_blocks_from_code(code, max_blocks=max_blocks)


def extract_nearby_macros(code: str, before_index: int, max_items: int = 6):
    from . import parse as parse_utils

    return parse_utils.extract_nearby_macros(code, before_index, max_items=max_items)


def _parse_c_file_base(code: str) -> list[dict]:
    from . import parse as parse_utils

    return parse_utils._parse_c_file_base(code)


def _clone_func_item(item: dict, file_context_extra: Optional[dict[str, Any]] = None) -> dict:
    file_context = dict(item.get("file_context") or {})
    if file_context_extra:
        file_context.update(file_context_extra)
    return {
        "comment_info": dict(item.get("comment_info") or {}),
        "func_info": dict(item.get("func_info") or {}),
        "body": item.get("body") or "",
        "file_context": file_context,
    }


def associate_comments_and_functions(code: str, file_context_extra: Optional[dict[str, Any]] = None):
    from . import parse as parse_utils

    return parse_utils.associate_comments_and_functions(code, file_context_extra=file_context_extra)


def get_cached_func_list_for_c_file(
    c_path: str,
    code: str,
    file_context_extra: Optional[dict[str, Any]] = None,
) -> list[dict]:
    from . import parse as parse_utils

    return parse_utils.get_cached_func_list_for_c_file(
        c_path,
        code,
        file_context_extra=file_context_extra,
    )


def prepare_func_list_for_c_file(
    c_path: str,
    project_root: Optional[str],
    cfg: GenConfig,
    prefilter: bool,
    need_symbol_map: bool = True,
) -> tuple[list[dict], Optional[str]]:
    from . import parse as parse_utils

    return parse_utils.prepare_func_list_for_c_file(
        c_path,
        project_root=project_root,
        cfg=cfg,
        prefilter=prefilter,
        need_symbol_map=need_symbol_map,
    )


# =============== 参数 / 局部变量 / 返回值解析 ===============

def parse_params_from_prototype(func_info: dict):
    from . import parse as parse_utils

    return parse_utils.parse_params_from_prototype(func_info)


_PAREN_SUFFIX_RE = re.compile(r"\s*[\(（][^)\）]*[\)）]\s*$")


def _strip_trailing_paren_content(text: str) -> str:
    from . import parse as parse_utils

    return parse_utils._strip_trailing_paren_content(text)


def parse_param_desc(desc_text: str, *, strip_paren_content: bool = False):
    from . import parse as parse_utils

    return parse_utils.parse_param_desc(
        desc_text,
        strip_paren_content=strip_paren_content,
    )


def _split_cn_name_and_usage_from_comment(comment: str) -> tuple[str, str]:
    from . import parse as parse_utils

    return parse_utils._split_cn_name_and_usage_from_comment(comment)


def _looks_like_compact_cn_label(text: str) -> bool:
    from . import parse as parse_utils

    return parse_utils._looks_like_compact_cn_label(text)


def _split_short_label_and_tail(text: str) -> tuple[str, str]:
    from . import parse as parse_utils

    return parse_utils._split_short_label_and_tail(text)


def _shorten_element_display_name(text: str, fallback: str = "") -> str:
    from . import parse as parse_utils

    return parse_utils._shorten_element_display_name(text, fallback=fallback)


def parse_local_variables_from_body(body: str):
    from . import parse as parse_utils

    return parse_utils.parse_local_variables_from_body(body)


def _filter_local_vars_against_params(
    local_vars: Sequence[dict],
    params: Sequence[dict],
    *,
    cfg: Optional["GenConfig"] = None,
    func_name: str = "",
) -> list[dict]:
    from . import parse as parse_utils

    return parse_utils._filter_local_vars_against_params(
        list(local_vars or []),
        list(params or []),
        cfg=cfg,
        func_name=func_name,
    )


def parse_return_var_from_body(body: str):
    from . import parse as parse_utils

    return parse_utils.parse_return_var_from_body(body)


# =============== .h 全局符号注释解析（用于逻辑条件替换） ===============

_HEADER_INDEX_CACHE: dict[str, tuple[float, dict[str, list[str]]]] = {}
_HEADER_SYMBOL_CACHE: dict[str, tuple[float, dict[str, str]]] = {}
_HEADER_TYPEDEF_CACHE: dict[str, tuple[float, list[str]]] = {}
_HEADER_MEMBER_MAP_CACHE: dict[str, tuple[float, dict[str, str]]] = {}
_SYMBOL_MAP_CACHE: dict[str, tuple[float, dict[str, str]]] = {}
_RELATED_TYPEDEF_CACHE: dict[str, tuple[float, list[str], dict[str, str]]] = {}
_C_FILE_PARSE_CACHE: dict[str, tuple[float, list[dict]]] = {}
_FUNC_CN_CACHE: dict[tuple[str, str, str], str] = {}


def _get_file_mtime(path: str) -> float:
    from . import parse as parse_utils

    return parse_utils._get_file_mtime(path)


def _quick_scan_c_code(code: str) -> tuple[bool, bool]:
    from . import parse as parse_utils

    return parse_utils._quick_scan_c_code(code)


def _get_cached_scan_result(c_path: str, code: str) -> tuple[bool, bool]:
    from . import parse as parse_utils

    return parse_utils._get_cached_scan_result(c_path, code)


def _extract_includes(code: str) -> list[str]:
    from . import parse as parse_utils

    return parse_utils._extract_includes(code)


def _build_header_index(root_dir: str, exclude_dirs: Optional[Sequence[str]] = None) -> dict[str, list[str]]:
    from . import parse as parse_utils

    return parse_utils._build_header_index(root_dir, exclude_dirs=list(exclude_dirs or []))


def _get_header_index(root_dir: str, exclude_dirs: Optional[Sequence[str]] = None) -> dict[str, list[str]]:
    from . import parse as parse_utils

    return parse_utils._get_header_index(root_dir, exclude_dirs=list(exclude_dirs or []))


def _split_code_and_comments_for_symbol(line: str) -> tuple[str, list[str]]:
    from . import parse as parse_utils

    return parse_utils._split_code_and_comments_for_symbol(line)


def _join_c_line_continuations(text: str) -> str:
    from . import parse as parse_utils

    return parse_utils._join_c_line_continuations(text)


def _clean_symbol_comment_text(text: str) -> str:
    from . import parse as parse_utils

    return parse_utils._clean_symbol_comment_text(text)


def _looks_like_noise_symbol_comment(text: str) -> bool:
    from . import parse as parse_utils

    return parse_utils._looks_like_noise_symbol_comment(text)


def _is_noop_comment(text: str) -> bool:
    from . import parse as parse_utils

    return parse_utils._is_noop_comment(text)


def _looks_like_logic_noise_comment(text: str) -> bool:
    from . import parse as parse_utils

    return parse_utils._looks_like_logic_noise_comment(text)


def _is_non_semantic_comment(text: str) -> bool:
    from . import parse as parse_utils

    return parse_utils._is_non_semantic_comment(text)


def _get_logic_comment_mode(cfg: Optional[GenConfig]) -> str:
    from . import parse as parse_utils

    return parse_utils._get_logic_comment_mode(cfg)


def classify_comment_hint(text: str) -> CommentHint:
    from . import parse as parse_utils

    return parse_utils.classify_comment_hint(text)


def extract_statement_hints(code_line: str, comments: Sequence[str]) -> list[CommentHint]:
    from . import parse as parse_utils

    return parse_utils.extract_statement_hints(code_line, list(comments or []))


def _is_macro_identifier(name: str) -> bool:
    s = (name or "").strip()
    return bool(s) and bool(re.fullmatch(r"[A-Z][A-Z0-9_]*", s))


def _looks_like_sentence_cn(text: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._looks_like_sentence_cn(text)


_LOGIC_LABEL_ACTION_PREFIXES = (
    "读取", "判断", "检查", "获取", "更新", "计算", "清除", "设置", "写入",
    "拷贝", "比较", "确认", "处理", "执行", "检测", "刷新", "发送", "接收",
)
_LOGIC_LABEL_PURPOSE_MARKERS = ("用于", "以便", "供", "表示", "表示为", "默认", "注意", "说明", "例如", "如下")


def _normalize_short_logic_label_comment(text: str, *, strip_action_prefix: bool = False) -> str:
    from . import parse as parse_utils

    return parse_utils._normalize_short_logic_label_comment(text, strip_action_prefix=strip_action_prefix)


def apply_comment_hints_to_logic(action_text: str, hints: Sequence[CommentHint], mode: str = "hint_only") -> str:
    from . import logic as logic_utils

    return logic_utils.apply_comment_hints_to_logic(
        action_text,
        hints,
        mode=mode,
        backend_module=sys.modules[__name__],
    )


def _should_keep_symbol_cn(name: str, cn: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._should_keep_symbol_cn(
        name,
        cn,
        backend_module=sys.modules[__name__],
    )


def _shorten_header_cn_comment(text: str) -> str:
    from . import parse as parse_utils

    return parse_utils._shorten_header_cn_comment(text)


def _normalize_header_comment_cn(cn_name: str, usage: str = "") -> str:
    from . import parse as parse_utils

    return parse_utils._normalize_header_comment_cn(cn_name, usage)


def _extract_symbol_map_from_header_code(header_code: str) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils._extract_symbol_map_from_header_code(header_code)


def _extract_member_symbol_map_from_typedefs(typedef_blocks: Sequence[str]) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils._extract_member_symbol_map_from_typedefs(list(typedef_blocks or []))


def _extract_member_symbol_map_from_header_code(header_code: str) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils._extract_member_symbol_map_from_header_code(header_code)


def _strip_function_bodies_keep_layout(code: str) -> str:
    from . import parse as parse_utils

    return parse_utils._strip_function_bodies_keep_layout(code)


def _load_header_symbol_map(header_path: str, cfg: GenConfig) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils._load_header_symbol_map(header_path, cfg)


def _load_header_typedef_blocks(header_path: str) -> list[str]:
    from . import parse as parse_utils

    return parse_utils._load_header_typedef_blocks(header_path)


def _load_header_member_symbol_map(header_path: str) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils._load_header_member_symbol_map(header_path)


def _iter_parent_dirs(start_dir: str, max_levels: int = 6) -> list[str]:
    from . import parse as parse_utils

    return parse_utils._iter_parent_dirs(start_dir, max_levels=max_levels)


def _guess_project_root_for_source(source_path: str, max_levels: int = 8) -> str:
    from . import parse as parse_utils

    return parse_utils._guess_project_root_for_source(source_path, max_levels=max_levels)


def _build_candidate_include_dirs(
    c_dir: str,
    project_root: str,
    exclude_dirs: Optional[Sequence[str]] = None,
    include_subdir_depth: int = 6,
) -> list[str]:
    from . import parse as parse_utils

    return parse_utils._build_candidate_include_dirs(
        c_dir,
        project_root,
        exclude_dirs=list(exclude_dirs or []),
        include_subdir_depth=include_subdir_depth,
    )


def _resolve_header_path(include: str, search_dirs: Sequence[str], header_index: Optional[dict[str, list[str]]]) -> Optional[str]:
    from . import parse as parse_utils

    return parse_utils._resolve_header_path(include, list(search_dirs or []), header_index)


def _resolve_header_path_with_reason(
    include: str,
    search_dirs: Sequence[str],
    header_index: Optional[dict[str, list[str]]],
) -> tuple[Optional[str], str]:
    from . import parse as parse_utils

    return parse_utils._resolve_header_path_with_reason(
        include,
        list(search_dirs or []),
        header_index,
    )


def _collect_transitive_headers(
    start_headers: Sequence[str],
    search_dirs: Sequence[str],
    header_index: Optional[dict[str, list[str]]],
    cfg: GenConfig,
) -> list[str]:
    from . import parse as parse_utils

    return parse_utils._collect_transitive_headers(
        list(start_headers or []),
        list(search_dirs or []),
        header_index,
        cfg,
    )


def _collect_related_header_paths_for_c_file(
    c_path: str,
    c_code: str,
    project_root: Optional[str],
    cfg: GenConfig,
) -> list[str]:
    from . import parse as parse_utils

    return parse_utils._collect_related_header_paths_for_c_file(
        c_path,
        c_code,
        project_root,
        cfg,
    )


def build_related_header_context_for_c_file(
    c_path: str,
    c_code: str,
    project_root: Optional[str],
    cfg: GenConfig,
) -> tuple[list[str], dict[str, str]]:
    from . import parse as parse_utils

    return parse_utils.build_related_header_context_for_c_file(
        c_path,
        c_code,
        project_root,
        cfg,
        backend_module=sys.modules[__name__],
    )


def build_global_symbol_map_for_c_file(c_path: str, c_code: str, project_root: Optional[str], cfg: GenConfig) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils.build_global_symbol_map_for_c_file(
        c_path,
        c_code,
        project_root,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _classify_symbol_kind(name: str) -> str:
    from . import parse as parse_utils

    return parse_utils._classify_symbol_kind(name)


def _extract_all_define_names(header_code: str) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils._extract_all_define_names(header_code)


def _group_symbols_by_prefix(
    symbols: Sequence[str], max_group_size: int = 50
) -> list[list[str]]:
    from . import parse as parse_utils

    return parse_utils._group_symbols_by_prefix(list(symbols or []), max_group_size=max_group_size)


def batch_translate_symbols(
    symbols: Sequence[str],
    kind: str = "macros",
    cfg: Optional["GenConfig"] = None,
) -> dict[str, str]:
    from . import parse as parse_utils

    return parse_utils.batch_translate_symbols(list(symbols or []), kind=kind, cfg=cfg)


def merge_prebuilt_symbols_into_runtime(prebuilt: dict) -> None:
    from . import parse as parse_utils

    parse_utils.merge_prebuilt_symbols_into_runtime(prebuilt)


def prebuild_project_symbol_db(project_dir: str, cfg: Optional[GenConfig] = None) -> dict:
    from . import parse as parse_utils

    return parse_utils.prebuild_project_symbol_db(project_dir, cfg=cfg)


# =============== 逻辑/流程伪代码生成 ===============

def is_decorative_comment(text: str) -> bool:
    from . import logic as logic_utils

    return logic_utils.is_decorative_comment(text)


def is_noop_statement(code: str) -> bool:
    from . import logic as logic_utils

    return logic_utils.is_noop_statement(code)


def is_declaration_line(code: str) -> bool:
    from . import logic as logic_utils

    return logic_utils.is_declaration_line(code)


def detect_increment_action(code: str, local_var_usages: dict):
    from . import logic as logic_utils

    return logic_utils.detect_increment_action(code, local_var_usages)


def _normalize_simple_int_literal(text: str) -> str:
    from . import logic as logic_utils

    return logic_utils._normalize_simple_int_literal(text)


def _render_simple_statement_action(
    code: str,
    name_map: Optional[dict[str, str]] = None,
    local_var_usages: Optional[dict[str, str]] = None,
) -> Optional[str]:
    from . import logic as logic_utils

    return logic_utils._render_simple_statement_action(
        code,
        name_map=name_map,
        local_var_usages=local_var_usages,
        backend_module=sys.modules[__name__],
    )


def _split_plain_assignment(code: str) -> Optional[tuple[str, str]]:
    from . import logic as logic_utils

    return logic_utils._split_plain_assignment(code)


def _extract_txpack_bias_name(text: str) -> str:
    from . import logic as logic_utils

    return logic_utils._extract_txpack_bias_name(text, backend_module=sys.modules[__name__])


def _is_txpack_dest_expr(expr: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._is_txpack_dest_expr(expr, backend_module=sys.modules[__name__])


def _is_txpack_temp_assignment(lhs: str, rhs: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._is_txpack_temp_assignment(lhs, rhs, backend_module=sys.modules[__name__])


def _humanize_bulk_owner(
    owner: str,
    name_map: Optional[dict[str, str]] = None,
    local_var_usages: Optional[dict[str, str]] = None,
) -> str:
    from . import logic as logic_utils

    return logic_utils._humanize_bulk_owner(
        owner,
        name_map=name_map,
        local_var_usages=local_var_usages,
        backend_module=sys.modules[__name__],
    )


def _extract_assignment_owner(expr: str) -> str:
    from . import logic as logic_utils

    return logic_utils._extract_assignment_owner(expr)


def _extract_assignment_group_key(code: str) -> str:
    from . import logic as logic_utils

    return logic_utils._extract_assignment_group_key(code, backend_module=sys.modules[__name__])


def _summarize_bulk_assignment_run(
    items: Sequence[dict[str, str]],
    name_map: Optional[dict[str, str]] = None,
    local_var_usages: Optional[dict[str, str]] = None,
) -> Optional[str]:
    from . import logic as logic_utils

    return logic_utils._summarize_bulk_assignment_run(
        list(items or []),
        name_map=name_map,
        local_var_usages=local_var_usages,
        backend_module=sys.modules[__name__],
    )


def _is_complex_condition(cond: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._is_complex_condition(cond)


def _split_top_level_comparison(cond: str) -> Optional[tuple[str, str, str]]:
    from . import logic as logic_utils

    return logic_utils._split_top_level_comparison(cond)


def _strip_balanced_outer_parens(text: str) -> str:
    from . import logic as logic_utils

    return logic_utils._strip_balanced_outer_parens(text)


def _split_top_level_logical(cond: str) -> Optional[tuple[str, str, str]]:
    from . import logic as logic_utils

    return logic_utils._split_top_level_logical(cond)


def _cn_compare_op(op: str) -> str:
    from . import logic as logic_utils

    return logic_utils._cn_compare_op(op)


def _prettify_logic_expr_text(text: str) -> str:
    from . import logic as logic_utils

    return logic_utils._prettify_logic_expr_text(text)


_STRUCTURED_COND_AI_CACHE: dict[str, str] = {}


def _ai_structured_condition_cn(cond: str, attached: Sequence[str], name_map: Optional[dict[str, str]], cfg: GenConfig) -> str:
    from . import logic as logic_utils

    return logic_utils._ai_structured_condition_cn(
        cond,
        attached,
        name_map,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _render_structured_condition_cn(cond: str, attached: Sequence[str], name_map: Optional[dict[str, str]], cfg: GenConfig) -> tuple[str, bool]:
    from . import logic as logic_utils

    return logic_utils._render_structured_condition_cn(
        cond,
        attached,
        name_map,
        cfg,
        backend_module=sys.modules[__name__],
    )

def _extract_condition_hint_from_attached(attached: Sequence[str]) -> str:
    from . import logic as logic_utils

    return logic_utils._extract_condition_hint_from_attached(
        attached,
        backend_module=sys.modules[__name__],
    )


def _sanitize_ai_logic_action(text: str) -> str:
    from . import logic as logic_utils

    return logic_utils._sanitize_ai_logic_action(text)


def render_logic_action_from_code(
    code_line: str,
    *,
    name_map: Optional[dict[str, str]] = None,
    local_var_usages: Optional[dict[str, str]] = None,
    hints: Sequence[CommentHint] = (),
    comment_mode: str = "hint_only",
    literal: bool = False,
) -> str:
    from . import logic as logic_utils

    return logic_utils.render_logic_action_from_code(
        code_line,
        name_map=name_map,
        local_var_usages=local_var_usages,
        hints=hints,
        comment_mode=comment_mode,
        literal=literal,
        backend_module=sys.modules[__name__],
    )


def generate_logic_from_body(body: str, local_vars, cfg: GenConfig, name_map: Optional[dict[str, str]] = None):
    from . import logic as logic_utils

    return logic_utils.generate_logic_from_body(
        body,
        local_vars,
        cfg,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def generate_logic_from_semantic_pack(logic_semantic_pack: dict[str, Any], cfg: GenConfig, name_map: Optional[dict[str, str]] = None):
    from . import logic as logic_utils

    return logic_utils.generate_logic_from_semantic_pack(
        logic_semantic_pack,
        cfg,
        backend_module=sys.modules[__name__],
        name_map=name_map,
    )


def select_ai_logic_polish_unknowns(logic: str, *, max_items: int = 12):
    from . import logic as logic_utils

    return logic_utils.select_ai_logic_polish_unknowns(
        logic,
        max_items=max_items,
        backend_module=sys.modules[__name__],
    )


def _collect_statement_like_units(body: str) -> list[str]:
    from . import logic as logic_utils

    return logic_utils._collect_statement_like_units(
        body,
        backend_module=sys.modules[__name__],
    )


def _build_pack_block_summary(
    codes: Sequence[str],
    *,
    name_map: Optional[dict[str, str]] = None,
    local_var_usages: Optional[dict[str, str]] = None,
) -> Optional[tuple[str, str]]:
    from . import logic as logic_utils

    return logic_utils._build_pack_block_summary(
        codes,
        name_map=name_map,
        local_var_usages=local_var_usages,
        backend_module=sys.modules[__name__],
    )


def _merge_pack_block_summaries(items: Sequence[tuple[str, str]]) -> list[str]:
    from . import logic as logic_utils

    return logic_utils._merge_pack_block_summaries(items)


def _build_enhanced_single_function_logic(
    body: str,
    local_vars,
    *,
    name_map: Optional[dict[str, str]] = None,
) -> Optional[str]:
    from . import logic as logic_utils

    return logic_utils._build_enhanced_single_function_logic(
        body,
        local_vars,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def _suggest_enhanced_single_function_desc(
    func_info: dict,
    body: str,
    current_desc: str = "",
) -> Optional[str]:
    from . import logic as logic_utils

    return logic_utils._suggest_enhanced_single_function_desc(
        func_info,
        body,
        current_desc=current_desc,
        backend_module=sys.modules[__name__],
    )


def _normalize_function_design_texts(design: FunctionDesign, name_map: Optional[dict[str, str]] = None) -> FunctionDesign:
    from . import logic as logic_utils

    return logic_utils._normalize_function_design_texts(
        design,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def _looks_like_output_label_text(text: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._looks_like_output_label_text(
        text,
        backend_module=sys.modules[__name__],
    )


def _looks_like_stable_usage_text(text: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._looks_like_stable_usage_text(
        text,
        backend_module=sys.modules[__name__],
    )

# =============== Word 样式选择 ===============
def pick_heading_style(doc: Document, level: int) -> str:
    from . import render as render_utils

    return render_utils.pick_heading_style(doc, level)


# =============== Word 生成（含格式设置） ===============

def init_document(cfg: GenConfig) -> Document:
    from . import render as render_utils

    return render_utils.init_document(cfg)


def apply_table_style(table, doc: Document, name: str = "Table Grid"):
    from . import render as render_utils

    return render_utils.apply_table_style(table, doc, name=name)


def find_content_placeholder(doc: Document):
    from . import render as render_utils

    return render_utils.find_content_placeholder(doc)


def relocate_generated_blocks(doc: Document, start_idx: int, anchor_para):
    """
    将生成内容块移动到占位符位置前，然后删除占位符。
    - start_idx：生成开始前文档 body 元素数量。
    - anchor_para：占位符段落。
    """
    if anchor_para is None:
        return
    body = doc._body._element
    new_elems = list(body)[start_idx:]
    if not new_elems:
        body.remove(anchor_para._p)
        return
    for el in new_elems:
        body.remove(el)
    try:
        anchor_idx = list(body).index(anchor_para._p)
    except ValueError:
        return
    for el in new_elems:
        body.insert(anchor_idx, el)
        anchor_idx += 1
    body.remove(anchor_para._p)


def style_as_caption(p):
    from . import render as render_utils

    return render_utils.style_as_caption(p)


def style_as_normal_paragraph(p):
    from . import render as render_utils

    return render_utils.style_as_normal_paragraph(p)


def add_section_label(doc: Document, text: str):
    from . import render as render_utils

    return render_utils.add_section_label(doc, text)


def add_alpha_section_label(doc: Document, title: str, idx: int, indent_pt: int = 24):
    from . import render as render_utils

    return render_utils.add_alpha_section_label(doc, title, idx, indent_pt=indent_pt)



def add_indented_text(doc: Document, text: str, indent_pt: int = 24):
    from . import render as render_utils

    return render_utils.add_indented_text(doc, text, indent_pt=indent_pt)


def add_captioned_table(
    doc: Document,
    caption_suffix: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
):
    caption = doc.add_paragraph()
    caption.add_run("表 ")
    add_seq_field(caption, "表")
    caption.add_run(f" {caption_suffix}")
    style_as_caption(caption)

    table = doc.add_table(rows=1, cols=len(headers))
    apply_table_style(table, doc)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)

    for row_vals in rows:
        row = table.add_row().cells
        for i, v in enumerate(row_vals):
            if i >= len(headers):
                break
            row[i].text = "" if v is None else str(v)
    from . import render as render_utils

    render_utils.prevent_table_row_splitting(table)
    return table


def render_table_or_none(
    doc: Document,
    caption_suffix: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
):
    from . import render as render_utils

    return render_utils.render_table_or_none(
        doc,
        caption_suffix,
        headers,
        rows,
        backend_module=sys.modules[__name__],
    )


def add_module_function_table(
    doc: Document,
    module_name: str,
    module_id: str,
    entries: Sequence[dict],
):
    from . import render as render_utils

    return render_utils.add_module_function_table(
        doc,
        module_name,
        module_id,
        entries,
        backend_module=sys.modules[__name__],
    )


def build_software_unit_table_doc(unit_rows: Sequence[dict], output: str) -> None:
    from . import render as render_utils

    return render_utils.build_software_unit_table_doc(
        unit_rows,
        output,
        backend_module=sys.modules[__name__],
    )


def _format_func_prototype(func_info: dict) -> str:
    proto = str((func_info or {}).get("prototype") or "").strip()
    if proto.endswith(";"):
        proto = proto[:-1].strip()
    if proto:
        return proto
    name = str((func_info or {}).get("func_name") or (func_info or {}).get("name") or "").strip()
    if not name:
        return ""
    ret = str((func_info or {}).get("ret_type") or "void").strip()
    params = str((func_info or {}).get("params") or "void").strip()
    return f"{ret} {name}({params})".strip()


def build_unit_function_list_doc(unit_tables: Sequence[dict], output: str) -> None:
    from . import render as render_utils

    return render_utils.build_unit_function_list_doc(
        unit_tables,
        output,
        backend_module=sys.modules[__name__],
    )

_CSU_TABLE_HEADERS = ("CSC 名称", "CSC 标识", "CSU 名称", "CSU 标识")
# 兼容两类标识：
# - 工程模式：D/R_SDD01_005（模块） + D/R_SDD01_005_002（CSU）
# - 单文件：  D/R_SDD01（模块）      + D/R_SDD01_002（CSU）
#
# 也兼容"需求前缀包含项目号"的形式（例如 req_id_prefix=D/R_SDD01_609）：
# - 模块：D/R_SDD01_609_001
# - CSU： D/R_SDD01_609_001_002
#
# 注意：这里的正则用于"替换/提取"而不是严格校验，因此允许多个 _NNN 段。
_ID_BASE_RE = r"(?<![A-Za-z0-9_/])(?:[A-Za-z]+/[A-Za-z]+(?:_[A-Za-z0-9]+)+|[A-Za-z][A-Za-z0-9]*(?:_[A-Za-z0-9]+)*)"
_MOD_ID_RE = re.compile(rf"({_ID_BASE_RE}(?:_\d{{3}}){{0,4}})")
_CSU_ID_RE = re.compile(rf"({_ID_BASE_RE}(?:_\d{{3}}){{1,5}})")


def _iter_doc_blocks(doc: Document):
    from . import render as render_utils

    yield from render_utils.iter_doc_blocks(
        doc,
        backend_module=sys.modules[__name__],
    )


def _is_heading(p: DocxParagraph, level: int) -> bool:
    from . import render as render_utils

    return render_utils.is_heading(p, level)


def _extract_module_id(text: str) -> Optional[str]:
    from . import render as render_utils

    return render_utils.extract_module_id(
        text,
        backend_module=sys.modules[__name__],
    )


def _extract_csu_id(text: str) -> Optional[str]:
    from . import render as render_utils

    return render_utils.extract_csu_id(
        text,
        backend_module=sys.modules[__name__],
    )


def _replace_last_match(text: str, pattern: re.Pattern, repl: str) -> str:
    if not text:
        return text
    last = None
    for m in pattern.finditer(text):
        last = m
    if not last:
        return text
    return f"{text[:last.start()]}{repl}{text[last.end():]}"


def collect_module_ids_in_doc(doc: Document) -> list[str]:
    from . import render as render_utils

    return render_utils.collect_module_ids_in_doc(
        doc,
        backend_module=sys.modules[__name__],
    )


def _replace_csu_id_in_text(text: str, new_csu_id: str) -> str:
    from . import render as render_utils

    return render_utils.replace_csu_id_in_text(
        text,
        new_csu_id,
        backend_module=sys.modules[__name__],
    )


def _is_csu_table(tbl: DocxTable) -> bool:
    from . import render as render_utils

    return render_utils.is_csu_table(
        tbl,
        backend_module=sys.modules[__name__],
    )


def _replace_module_id_in_text(text: str, new_module_id: str) -> str:
    from . import render as render_utils

    return render_utils.replace_module_id_in_text(
        text,
        new_module_id,
        backend_module=sys.modules[__name__],
    )


def update_module_headings_only(
    doc: Document,
    module_id: str,
    csu_count: int,
) -> dict[str, int]:
    from . import render as render_utils

    return render_utils.update_module_headings_only(
        doc,
        module_id,
        csu_count,
        backend_module=sys.modules[__name__],
    )


def update_all_module_headings_only(
    doc: Document,
    module_entries: Sequence[tuple[str, int]],
) -> dict[str, int]:
    from . import render as render_utils

    return render_utils.update_all_module_headings_only(
        doc,
        module_entries,
        backend_module=sys.modules[__name__],
    )


def update_module_csu_in_doc(
    doc: Document,
    module_id: str,
    csu_entries: Sequence[dict],
) -> dict[str, bool]:
    from . import render as render_utils

    return render_utils.update_module_csu_in_doc(
        doc,
        module_id,
        csu_entries,
        backend_module=sys.modules[__name__],
    )


def compute_project_module_id(project_dir: str, c_file: str, cfg: GenConfig) -> Optional[str]:
    """
    使用与 generate_design_doc_for_project 相同的文件排序逻辑，计算指定 c_file 的 module_id。
    """
    from . import pipeline as pipeline_utils

    return pipeline_utils.compute_project_module_id(
        project_dir,
        c_file,
        cfg,
        backend_module=sys.modules[__name__],
    )


def collect_project_c_files_by_layer(project_dir: str, cfg: GenConfig) -> tuple[str, list[str], list[str], list[str]]:
    """
    扫描工程，返回：(src_dir, app_files, mid_files, drv_files)。
    分层规则与 generate_design_doc_for_project 一致。
    """
    from . import pipeline as pipeline_utils

    return pipeline_utils.collect_project_c_files_by_layer(
        project_dir,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _apply_project_file_order_override(
    app_files: list[str],
    mid_files: list[str],
    drv_files: list[str],
    order: Optional[dict[str, list[Any]]],
) -> tuple[list[str], list[str], list[str]]:
    from . import pipeline as pipeline_utils

    return pipeline_utils.apply_project_file_order_override(
        app_files,
        mid_files,
        drv_files,
        order,
    )


def _build_project_modules_from_order(
    *,
    app_files: list[str],
    mid_files: list[str],
    drv_files: list[str],
    order: Optional[dict[str, list[Any]]],
) -> tuple[list[dict], list[dict], list[dict]]:
    """
    将 GUI 的 project_file_order 转为"模块列表"：
    - entry 为 str：单文件模块
    - entry 为 dict：{"module": "...", "files": [..]} 多文件合并模块
    支持跨层移动：entry 出现在哪一层，就归属到哪一层（仅保留工程内文件）。
    """
    from . import pipeline as pipeline_utils

    return pipeline_utils.build_project_modules_from_order(
        app_files=app_files,
        mid_files=mid_files,
        drv_files=drv_files,
        order=order,
    )


def _get_ordered_project_c_files(project_dir: str, cfg: GenConfig) -> list[str]:
    """
    返回工程中 .c 文件的顺序列表（与 generate_design_doc_for_project 保持一致）。
    """
    from . import pipeline as pipeline_utils

    return pipeline_utils.get_ordered_project_c_files(
        project_dir,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _collect_project_module_entries_and_units(
    project_dir: str,
    cfg: GenConfig,
    prefilter: bool,
) -> tuple[list[tuple[str, int]], list[dict]]:
    """
    收集工程中各模块的 (module_id, csu_count) 以及软件单元表行。
    """
    from . import pipeline as pipeline_utils

    return pipeline_utils.collect_project_module_entries_and_units(
        project_dir,
        cfg,
        prefilter,
        backend_module=sys.modules[__name__],
    )


def _collect_project_module_tables_data(
    project_dir: str,
    cfg: GenConfig,
    prefilter: bool,
) -> tuple[list[dict], list[dict]]:
    """
    收集工程中各模块的模块表数据（用于 CSU 更新）。
    返回：(modules, unit_rows)
      modules: [{"module_id": str, "module_name": str, "csu_entries": [{"csu_name": str}]}]
      unit_rows: 软件单元表行（同 _collect_project_module_entries_and_units 的行为）
    """
    from . import pipeline as pipeline_utils

    return pipeline_utils.collect_project_module_tables_data(
        project_dir,
        cfg,
        prefilter,
        backend_module=sys.modules[__name__],
    )


def update_all_module_tables_and_headings(
    doc: Document,
    modules: Sequence[dict],
) -> dict[str, int]:
    """
    按文档中 Heading 3 的顺序，依次更新模块标题(module_id)、模块表(CSC/CSU)、以及后续 Heading 4 的 CSU 标识。
    modules: [{"module_id": str, "module_name": str, "csu_entries": [{"csu_name": str}]}]
    返回：{"modules": int, "headings": int, "tables": int}
    """
    from . import render as render_utils

    return render_utils.update_all_module_tables_and_headings(
        doc,
        modules,
        backend_module=sys.modules[__name__],
    )


def _get_heading_level(p: DocxParagraph, max_level: int = 6) -> Optional[int]:
    from . import render as render_utils

    return render_utils.get_heading_level(p, max_level)


def update_csu_ids_in_design_chapter_by_headings(
    doc: Document,
    cfg: GenConfig,
    *,
    chapter_keyword: str = "CSCI详细设计",
) -> dict[str, int]:
    from . import render as render_utils

    return render_utils.update_csu_ids_in_design_chapter_by_headings(
        doc,
        cfg,
        chapter_keyword=chapter_keyword,
        backend_module=sys.modules[__name__],
    )


_UNIT_TABLE_HEADERS = ("序号", "软件单元名称", "函数原型", "唯一标识", "存放位置", "开发状态", "用途")
_UNIT_TABLE_HEADERS_LEGACY = ("序号", "软件单元名称", "唯一标识", "存放位置", "开发状态", "用途")


def collect_units_from_design_doc(
    doc: Document,
    cfg: GenConfig,
    *,
    chapter_keyword: str = "CSCI详细设计",
) -> list[dict]:
    """
    仅从 doc 文档中提取软件单元清单（用于更新 unit_table）：
    - 读取 Heading3（模块名）与 Heading4（函数名+CSU 标识）
    - 用函数段落内的 "b) 功能说明" 第一段作为用途（否则回退为函数名）
    """
    from . import render as render_utils

    return render_utils.collect_units_from_design_doc(
        doc,
        cfg,
        chapter_keyword=chapter_keyword,
        backend_module=sys.modules[__name__],
    )


def update_software_unit_table_from_design_doc(
    doc: Document,
    cfg: GenConfig,
    *,
    design_doc_path: str,
    chapter_keyword: str = "CSCI详细设计",
) -> dict[str, Any]:
    from . import render as render_utils

    return render_utils.update_software_unit_table_from_design_doc(
        doc,
        cfg,
        design_doc_path=design_doc_path,
        chapter_keyword=chapter_keyword,
        backend_module=sys.modules[__name__],
    )


def get_function_chinese_name(comment_info: dict, func_info: dict) -> str:
    """
    优先顺序：
    1. [函数中文名]
    2. [函数名]（如果含中文）
    3. [功能描述] 提取短名
    4. C 原型中的函数名
    """
    key = (
        (comment_info.get("func_cn_name") or "").strip(),
        (comment_info.get("func_name") or "").strip(),
        (comment_info.get("desc") or "").strip(),
        (func_info.get("func_name") or "").strip(),
    )
    cached = _FUNC_CN_CACHE.get(key)
    if cached is not None:
        return cached

    from . import naming as naming_utils

    result = naming_utils.get_function_chinese_name(
        dict(comment_info or {}),
        dict(func_info or {}),
        resolve_canonical_name=resolve_canonical_symbol_name,
    )

    _FUNC_CN_CACHE[key] = result
    if len(_FUNC_CN_CACHE) > 4096:
        _FUNC_CN_CACHE.clear()
    return result


def get_function_chinese_name_rich(func_data: dict, *, cfg: Optional["GenConfig"] = None) -> str:
    """LLM-first function Chinese naming with rich body context.

    See autodoc.naming.get_function_chinese_name_rich for full documentation.
    """
    from . import naming as naming_utils

    return naming_utils.get_function_chinese_name_rich(
        func_data,
        cfg=cfg,
        resolve_canonical_name=resolve_canonical_symbol_name,
    )


# 生成"函数标识符 -> 中文名"映射（用于逻辑语句替换）
def _build_func_cn_map(func_list: Sequence[dict]) -> dict[str, str]:
    out: dict[str, str] = {}
    for item in func_list or []:
        fi = (item or {}).get("func_info") or {}
        if not fi:
            continue
        name = (fi.get("func_name") or "").strip()
        if not name:
            continue
        cn = get_function_chinese_name(item.get("comment_info") or {}, fi)
        if cn and cn != name:
            out[name] = cn
    return out


_VERBOSE_FUNC_TITLE_PREFIXES = (
    "根据",
    "按",
    "按照",
    "遍历",
    "读取",
    "获取",
    "更新",
    "检查",
    "判断",
    "完成",
    "处理",
    "设置",
    "执行",
    "计算",
    "轮询",
    "查询",
    "汇总",
    "生成",
    "记录",
)


def _extract_compact_function_title(text: str) -> str:
    s = _safe_strip(text)
    if not s:
        return ""
    label, _tail = _split_short_label_and_tail(s)
    if label:
        return label
    if _looks_like_compact_cn_label(s):
        return s
    return ""


def _apply_domain_title_hint(title: str, func_name: str = "") -> str:
    s = _safe_strip(title)
    ident = _safe_strip(func_name).upper()
    if not s:
        return ""
    if "IFBIT" in ident:
        if s.startswith("周期BIT"):
            return "周期自检" + s[len("周期BIT"):]
        if s.startswith("周期") and not s.startswith("周期自检"):
            return "周期自检" + s[len("周期"):]
    if "PUBIT" in ident:
        if s.startswith("上电BIT"):
            return "上电自检" + s[len("上电BIT"):]
        if s.startswith("上电") and not s.startswith("上电自检"):
            return "上电自检" + s[len("上电"):]
    return s


def _function_domain_prefix(func_name: str) -> str:
    ident = _safe_strip(func_name).upper()
    if "IFBIT" in ident:
        return "周期自检"
    if "PUBIT" in ident:
        return "上电自检"
    if "MBIT" in ident:
        return "维护自检"
    return ""


def _clean_function_title_tail(text: str) -> str:
    s = _safe_strip(text)
    if not s:
        return ""
    s = s.replace("周期BIT", "").replace("IFBIT", "").replace("PuBIT", "").replace("MBIT", "")
    s = s.replace("BIT", "")
    s = s.replace("检测结果获取", "结果获取")
    s = s.replace("检测项状态信息", "状态")
    s = s.replace("状态信息", "状态")
    s = s.replace("结果数据", "结果")
    s = re.sub(r"^(?:周期自检|上电自检|维护自检|周期|上电|维护)", "", s)
    s = re.sub(r"^(?:根据|按|按照|用于|对)\s*", "", s)
    s = re.sub(r"[的地得]$", "", s)
    s = re.sub(r"\s+", "", s)
    return s


def _suffix_action_hint(func_name: str) -> str:
    ident = _safe_strip(func_name)
    if not ident:
        return ""
    if re.search(r"(?:Result)?Get$", ident, re.IGNORECASE):
        return "获取"
    if re.search(r"(?:State)?Update$", ident, re.IGNORECASE):
        return "更新"
    if re.search(r"(?:Check|Test)$", ident, re.IGNORECASE):
        return "检测"
    if re.search(r"Init$", ident, re.IGNORECASE):
        return "初始化"
    return ""


def _compose_short_function_title(func_name: str, comment_desc: str, current_title: str) -> str:
    domain_prefix = _function_domain_prefix(func_name)
    compact_desc = _extract_compact_function_title(comment_desc)
    normalized_desc = _apply_domain_title_hint(compact_desc, func_name) if compact_desc else ""
    normalized_title = _apply_domain_title_hint(current_title, func_name) if current_title else ""
    suffix_action = _suffix_action_hint(func_name)

    tail_source = normalized_desc or normalized_title
    tail = _clean_function_title_tail(tail_source)
    if domain_prefix and tail:
        if suffix_action == "获取":
            if "结果" in tail and not tail.endswith("获取"):
                tail = "结果获取"
        elif suffix_action == "更新":
            if "状态" in tail and not tail.endswith("更新"):
                tail = "状态更新"
        elif suffix_action == "检测":
            if tail.endswith("检查"):
                tail = tail[:-2] + "检测"
            elif not tail.endswith("检测"):
                tail = tail + "检测"

        tail = re.sub(r"^(?:自检)+", "", tail)
        candidate = domain_prefix + tail
        if len(candidate) <= 12 and _contains_cjk(candidate):
            return candidate

    return normalized_title or normalized_desc


def _looks_like_verbose_function_cn_title(text: str, desc: str = "") -> bool:
    s = _safe_strip(text)
    desc_s = _safe_strip(desc)
    if (not s) or (not _contains_cjk(s)):
        return False
    if desc_s and s == desc_s and len(s) >= 10:
        return True
    if any(s.startswith(prefix) for prefix in _VERBOSE_FUNC_TITLE_PREFIXES) and len(s) >= 6:
        return True
    if len(s) >= 10 and any(token in s for token in ("并", "后", "然后", "用于", "以便")):
        return True
    return False


def _normalize_function_cn_title(text: str, *, func_name: str = "", comment_desc: str = "") -> str:
    from . import naming as naming_utils

    return naming_utils.normalize_function_cn_title(
        text,
        func_name=func_name,
        comment_desc=comment_desc,
    )


def _function_title_needs_ai_upgrade(comment_info: Optional[dict], func_info: Optional[dict]) -> bool:
    current_title = _safe_strip((comment_info or {}).get("func_cn_name"))
    comment_desc = _safe_strip((comment_info or {}).get("desc"))
    func_name = _safe_strip((func_info or {}).get("func_name"))
    if _is_missing_gap_text(current_title):
        return False
    compact = re.sub(r"\s+", "", current_title)
    normalized = _normalize_function_cn_title(current_title, func_name=func_name, comment_desc=comment_desc)
    normalized_compact = re.sub(r"\s+", "", normalized)
    if not _contains_cjk(current_title):
        return True
    if len(compact) > 12:
        return True
    if normalized_compact and normalized_compact != compact:
        return True
    if _looks_like_verbose_function_cn_title(current_title, comment_desc):
        return True
    if comment_desc and current_title == comment_desc and len(compact) >= 8:
        return True
    return False


def _title_ai_force_refine_enabled(cfg: Optional["GenConfig"]) -> bool:
    mode = utils.cfg_get_str(cfg, "title_ai_force_refine", "auto").strip().lower()
    if mode in {"0", "false", "off", "no"}:
        return False
    if mode in {"1", "true", "on", "always"}:
        return True
    authority = utils.cfg_get_str(cfg, "naming_authority", "ai_first").strip().lower()
    return authority == "ai_first"


def _should_force_ai_title_refine(
    cfg: Optional["GenConfig"],
    comment_info: Optional[dict],
    func_info: Optional[dict],
) -> bool:
    if not bool(getattr(cfg, "ai_assist", False)):
        return False
    if not _title_ai_force_refine_enabled(cfg):
        return False
    current_title = _safe_strip((comment_info or {}).get("func_cn_name"))
    func_name = _safe_strip((func_info or {}).get("func_name"))
    if _is_missing_gap_text(current_title) or not func_name:
        return False
    if not _contains_cjk(current_title):
        return True
    compact = re.sub(r"\s+", "", current_title)
    if not compact:
        return False
    if len(compact) <= 16:
        return True
    return _function_title_needs_ai_upgrade(comment_info, func_info)


def _should_accept_refined_function_title(
    current_title: str,
    candidate_title: str,
    *,
    func_name: str,
    comment_desc: str,
    examples: Sequence[dict[str, Any]],
) -> bool:
    current = _safe_strip(current_title)
    candidate = _safe_strip(candidate_title)
    if not candidate:
        return False
    from . import naming as naming_utils

    if naming_utils.title_violates_required_acronyms(candidate, func_name):
        return False
    if not current:
        return True
    current_score = _rank_function_title_candidate(
        current,
        func_name=func_name,
        comment_desc=comment_desc,
        examples=examples,
    )
    candidate_score = _rank_function_title_candidate(
        candidate,
        func_name=func_name,
        comment_desc=comment_desc,
        examples=examples,
    )
    if candidate_score > current_score:
        return True
    current_compact = re.sub(r"\s+", "", _safe_strip(current))
    candidate_compact = re.sub(r"\s+", "", _safe_strip(candidate))
    if candidate_score == current_score and candidate_compact and current_compact and len(candidate_compact) < len(current_compact):
        return True
    return False


# =============== AI 辅助：缺口探测 / 调用 / 合并 ===============

def detect_gaps(comment_info, locals_, params, in_map, out_map, func_info=None, cfg: Optional["GenConfig"] = None, logic_semantic_pack: Optional[dict[str, Any]] = None, body: str = ""):
    from . import ai as ai_utils

    return ai_utils.detect_gaps(
        comment_info,
        locals_,
        params,
        in_map,
        out_map,
        func_info=func_info,
        cfg=cfg,
        logic_semantic_pack=logic_semantic_pack,
        body=body,
        _runtime_module=sys.modules.get(__name__),
    )


def _is_missing_gap_text(text: str) -> bool:
    if not text:
        return True
    s = str(text or "").strip()
    return s in ("NONE", "None", "none", "无", "待人工修改") or _looks_like_pseudo_function_desc(s)


def _looks_like_pseudo_function_desc(text: str) -> bool:
    s = _safe_strip(text)
    if not s:
        return False
    compact = re.sub(r"\s+", "", s)
    if re.match(r"^【?说明】?\s*[:：]", s):
        return True
    if compact in {"循环索引", "临时变量", "临时数据", "临时值", "计数值", "返回值"}:
        return True
    if compact.endswith("索引") and len(compact) <= 6:
        return True
    return False


def _looks_like_generic_local_cn_name(text: str) -> bool:
    s = _safe_strip(text)
    if not s:
        return False
    compact = re.sub(r"\s+", "", s)
    if compact in _GENERIC_LOCAL_CN_NAMES:
        return True
    if _looks_like_bad_canonical_name(s):
        return True
    if _looks_like_low_quality_symbol_cn(s, raw_ident=s):
        return True
    if compact in {"数据指针", "指针", "缓存缓存值", "存放变量值", "存放数据值"}:
        return True
    if compact.startswith(("存放", "缓存", "记录")) and len(compact) <= 8:
        return True
    if compact.endswith("指针") and len(compact) <= 6:
        return True
    return False


def _local_var_needs_ai_upgrade(item: Optional[dict]) -> bool:
    entry = item or {}
    ident = _safe_strip(entry.get("name"))
    cn_name = _safe_strip(entry.get("cn_name"))
    usage = _safe_strip(entry.get("usage"))
    if _is_missing_gap_text(cn_name) or _looks_like_generic_local_cn_name(cn_name) or _looks_like_low_quality_symbol_cn(cn_name, raw_ident=ident):
        return True
    if _is_missing_gap_text(usage):
        return True
    if _should_replace_local_usage_with_ai(usage, cn_name):
        return True
    return False


def _preferred_local_cn_hint(item: Optional[dict]) -> str:
    entry = item or {}
    ident = _safe_strip(entry.get("name"))
    comment_cn_name = _safe_strip(entry.get("comment_cn_name"))
    profile_cn = _safe_strip(entry.get("profile_cn_candidate"))
    cn_name = _safe_strip(entry.get("cn_name"))
    comment_hint = _safe_strip(entry.get("comment_hint"))
    if comment_cn_name and (not _looks_like_generic_local_cn_name(comment_cn_name)) and (not _looks_like_low_quality_symbol_cn(comment_cn_name, raw_ident=ident)):
        return comment_cn_name
    if profile_cn and (not _looks_like_generic_local_cn_name(profile_cn)) and (not _looks_like_low_quality_symbol_cn(profile_cn, raw_ident=ident)):
        return profile_cn
    if cn_name and (not _looks_like_generic_local_cn_name(cn_name)) and (not _looks_like_low_quality_symbol_cn(cn_name, raw_ident=ident)):
        return cn_name
    guessed = _safe_strip(_guess_cn_from_ident(ident))
    if guessed and guessed != ident and not _looks_like_generic_local_cn_name(guessed) and not _looks_like_low_quality_symbol_cn(guessed, raw_ident=ident):
        return guessed
    normalized_hint = _normalize_symbol_hint_text(comment_hint)
    if normalized_hint:
        return normalized_hint
    return cn_name


def _derive_local_cn_from_usage(usage_text: str, ident: str = "") -> str:
    usage = _safe_strip(usage_text)
    if not usage:
        return ""
    s = usage
    s = re.sub(r"^(?:用于|以便|供|存放|缓存|记录|指向|更新|获取|读取|写入|遍历|计算|判断)", "", s).strip()
    s = re.sub(r"^根据[^，,；;。]*?(?:计算|获取|更新|读取|判断|生成)", "", s).strip()
    s = re.sub(r"(?:变量|数据|结果)?$", lambda m: m.group(0), s)
    for suffix in ("状态值", "检测结果", "结果", "索引", "状态", "计数", "标志", "信息"):
        if s.endswith(suffix):
            return s
    if "检测结果" in usage:
        return "检测结果"
    if "状态值" in usage:
        return "状态值"
    if "状态" in usage:
        return "状态"
    if "索引" in usage or "遍历检测项" in usage:
        return "索引"
    compact = re.sub(r"^(?:更新|获取|读取|写入|遍历|计算|判断)", "", s).strip()
    if _looks_like_compact_cn_label(compact) and not _looks_like_generic_local_cn_name(compact):
        return compact
    guessed = _safe_strip(_guess_cn_from_ident(ident))
    if guessed and not _looks_like_generic_local_cn_name(guessed):
        return guessed
    return ""


def _derive_local_cn_from_item(item: Optional[dict]) -> str:
    from . import semantic_registry

    entry = item or {}
    ident = _safe_strip(entry.get("name")).lower()
    decl_type = _safe_strip(entry.get("type")).lower()
    if not ident:
        return ""
    registry_hint = semantic_registry.local_name_hint(entry)
    if registry_hint:
        return registry_hint
    if any(token in ident for token in ("ratio", "scale")):
        return "换算系数"
    if "gain" in ident:
        return "增益系数"
    packet_match = re.search(r"(?:pmfl|data)(\d{3})", decl_type or ident)
    if packet_match and any(token in decl_type for token in ("1553b", "pmfl", "revdef")):
        return f"{packet_match.group(1)}字打包缓存"
    if "compat" in ident:
        if ("act" in ident) and ("flt" in decl_type):
            return "作动器故障兼容字"
        if "flt" in decl_type:
            return "故障兼容字"
        if "mode" in ident and "1553b" in decl_type:
            return "模式源字"
    if ident.endswith("srcerr_u16"):
        return "源有效性错误标志"
    if ident.endswith("modeerr_u16"):
        return "模式错误标志"
    if "redundata_t" in decl_type:
        if "riu" in ident:
            return "RIU链路状态"
        if "ccdl" in ident:
            return "CCDL链路状态"
        if "kzzz" in ident and "left" in ident:
            return "左吊舱链路状态"
        if "kzzz" in ident and "right" in ident:
            return "右吊舱链路状态"
        if "kzzz" in ident:
            return "吊舱链路状态"
    return ""


def _specific_loop_index_label_for_ident(label: str, ident: str) -> str:
    compact = re.sub(r"\s+", "", _safe_strip(label)).lower()
    if not compact:
        return ""
    match = re.fullmatch(r"(?:循环)?索引(?P<suffix>ii|jj)", compact)
    if not match:
        return ""
    suffix = match.group("suffix")
    ident_tokens = {token.lower() for token in re.split(r"[^A-Za-z0-9]+", _safe_strip(ident)) if token}
    if suffix not in ident_tokens:
        return ""
    return _safe_strip(label)


def _select_local_display_name(item: dict) -> str:
    ident = _safe_strip((item or {}).get("name"))
    comment_cn_name = _safe_strip((item or {}).get("comment_cn_name"))
    cn_name = _safe_strip((item or {}).get("cn_name"))
    comment_hint = _safe_strip((item or {}).get("comment_hint"))
    specific_loop_label = (
        _specific_loop_index_label_for_ident(comment_cn_name, ident)
        or _specific_loop_index_label_for_ident(cn_name, ident)
        or _specific_loop_index_label_for_ident(comment_hint, ident)
    )
    if specific_loop_label:
        return _shorten_element_display_name(specific_loop_label, fallback=ident)
    if comment_cn_name and not _looks_like_generic_local_cn_name(comment_cn_name) and not _looks_like_low_quality_symbol_cn(comment_cn_name, raw_ident=ident):
        return _shorten_element_display_name(comment_cn_name, fallback=ident)
    from_item = _derive_local_cn_from_item(item)
    if from_item and not _looks_like_generic_local_cn_name(from_item) and not _looks_like_low_quality_symbol_cn(from_item, raw_ident=ident):
        return _shorten_element_display_name(from_item, fallback=ident)
    preferred_cn = _preferred_local_cn_hint(item)
    from_usage = _derive_local_cn_from_usage(_safe_strip((item or {}).get("usage")), ident)
    if (
        preferred_cn
        and from_usage.endswith(("临时量", "临时值", "当前值", "缓存值"))
        and preferred_cn.endswith(("限幅值", "限流值", "快照", "状态值", "状态输出值"))
        and len(re.sub(r"\s+", "", from_usage)) <= len(re.sub(r"\s+", "", preferred_cn))
    ):
        return _shorten_element_display_name(preferred_cn, fallback=ident)
    preferred_cov = _candidate_ident_semantic_coverage(preferred_cn, ident)
    usage_cov = _candidate_ident_semantic_coverage(from_usage, ident)
    if (
        preferred_cn
        and from_usage
        and from_usage != preferred_cn
        and (
            from_usage.endswith(preferred_cn)
            or preferred_cn in {"结果", "状态", "信息", "计数器", "标志"}
            or usage_cov > preferred_cov
        )
        and not _looks_like_generic_local_cn_name(from_usage)
        and not _looks_like_low_quality_symbol_cn(from_usage, raw_ident=ident)
    ):
        return _shorten_element_display_name(from_usage, fallback=ident)
    if preferred_cn and not _looks_like_generic_local_cn_name(preferred_cn) and not _looks_like_low_quality_symbol_cn(preferred_cn, raw_ident=ident):
        return _shorten_element_display_name(preferred_cn, fallback=ident)
    if from_usage and not _looks_like_generic_local_cn_name(from_usage) and not _looks_like_low_quality_symbol_cn(from_usage, raw_ident=ident):
        return _shorten_element_display_name(from_usage, fallback=ident)
    resolved = _shorten_element_display_name(
        resolve_canonical_symbol_name(
            ident,
            kind="symbols",
            comment_cn=preferred_cn or (item or {}).get("usage") or "",
            fallback=ident,
        ),
        fallback=ident,
    )
    if resolved and not _looks_like_generic_local_cn_name(resolved):
        return resolved
    guessed = _safe_strip(_guess_cn_from_ident(ident))
    if guessed and not _looks_like_generic_local_cn_name(guessed) and not _looks_like_low_quality_symbol_cn(guessed, raw_ident=ident):
        return _shorten_element_display_name(guessed, fallback=ident)
    return _shorten_element_display_name(ident, fallback=ident)


def _ident_cn_fragments_for_scoring(symbol: str) -> tuple[str, ...]:
    fragments: list[str] = []
    for token in _split_ident_tokens(symbol):
        cn = _safe_strip(_IDENT_CN_MAP.get(token.lower()))
        if len(cn) < 2:
            continue
        if cn in {"临时", "当前", "前值", "上", "下", "类型", "参数", "信息"}:
            continue
        if cn not in fragments:
            fragments.append(cn)
    return tuple(fragments[:8])


def _candidate_ident_semantic_coverage(candidate: str, symbol: str) -> int:
    compact = re.sub(r"\s+", "", _safe_strip(candidate))
    if not compact:
        return 0
    return sum(1 for frag in _ident_cn_fragments_for_scoring(symbol) if frag in compact)


def _looks_like_low_quality_symbol_cn(text: str, *, raw_ident: str = "") -> bool:
    value = _safe_strip(text)
    ident = _safe_strip(raw_ident)
    if (not value) or (not ident):
        return False
    compact = re.sub(r"\s+", "", value)
    coverage = _candidate_ident_semantic_coverage(compact, ident)
    fragments = _ident_cn_fragments_for_scoring(ident)
    if _looks_like_memberish_bitfield(ident) and compact in {"状态快照", "当前值", "位标志", "状态值", "状态"}:
        return True
    if (
        ("数据位域" in compact or re.search(r"bit\d+", compact, flags=re.IGNORECASE))
        and not re.search(r"(?:bit|mask|field|ctrlinfo|bit\d+)", ident, flags=re.IGNORECASE)
    ):
        return True
    if compact.count("数据位域") >= 2:
        return True
    if compact.endswith(("临时量", "临时值", "缓存值", "当前值", "模型")) and len(fragments) >= 3 and coverage <= 1:
        return True
    return False


def _prefer_more_specific_local_cn_candidate(candidate: str, current_cn: str, evidence: SymbolEvidence) -> bool:
    from . import naming as naming_utils

    return naming_utils._prefer_more_specific_local_cn_candidate(
        candidate,
        current_cn,
        evidence,
        backend_module=sys.modules[__name__],
    )


def _score_profile_local_cn(candidate: str, *, current_cn: str, backup_cn: str, evidence: SymbolEvidence) -> int:
    from . import naming as naming_utils

    return naming_utils._score_profile_local_cn(
        candidate,
        current_cn=current_cn,
        backup_cn=backup_cn,
        evidence=evidence,
        backend_module=sys.modules[__name__],
    )


def _retrieve_local_symbol_name_candidates(
    item: dict,
    *,
    body: str,
    neighbor_symbols: Sequence[str],
    comment_desc: str = "",
    cfg: Optional["GenConfig"] = None,
    evidence: Optional[SymbolEvidence] = None,
) -> tuple[str, ...]:
    from . import naming as naming_utils

    return naming_utils._retrieve_local_symbol_name_candidates(
        item,
        body=body,
        neighbor_symbols=neighbor_symbols,
        comment_desc=comment_desc,
        cfg=cfg,
        evidence=evidence,
        backend_module=sys.modules[__name__],
    )


def _repair_local_cn_name_with_profile(
    item: dict,
    *,
    body: str,
    neighbor_symbols: Sequence[str],
    comment_desc: str = "",
    backup_cn: str = "",
    cfg: Optional["GenConfig"] = None,
) -> None:
    from . import naming as naming_utils

    ident = _safe_strip((item or {}).get("name"))
    current_before = _safe_strip((item or {}).get("cn_name"))
    comment_cn_name = _safe_strip((item or {}).get("comment_cn_name"))
    comment_hint = _safe_strip((item or {}).get("comment_hint"))
    specific_loop_label = (
        _specific_loop_index_label_for_ident(current_before, ident)
        or _specific_loop_index_label_for_ident(comment_cn_name, ident)
        or _specific_loop_index_label_for_ident(comment_hint, ident)
    )
    if specific_loop_label:
        item["cn_name"] = specific_loop_label
        return
    if (
        comment_cn_name
        and current_before == comment_cn_name
        and not _looks_like_generic_local_cn_name(comment_cn_name)
        and not _looks_like_low_quality_symbol_cn(comment_cn_name, raw_ident=ident)
    ):
        return

    naming_utils.repair_local_cn_name_with_profile(
        item,
        body=body,
        neighbor_symbols=neighbor_symbols,
        comment_desc=comment_desc,
        backup_cn=backup_cn,
        cfg=cfg,
        backend_module=sys.modules[__name__],
    )
    current_cn = _safe_strip((item or {}).get("cn_name"))
    comment_cn_name = _safe_strip((item or {}).get("comment_cn_name"))
    try:
        profile = build_local_symbol_profile(
            item,
            body=body,
            neighbor_symbols=neighbor_symbols,
            scope=_safe_strip((item or {}).get("scope")) or "local",
            comment_desc=comment_desc,
            cfg=cfg,
        )
    except Exception:
        return
    candidate = _safe_strip(getattr(profile, "suggested_cn", ""))
    reason = _safe_strip(getattr(profile, "suggestion_reason", ""))
    if (
        not candidate
        or candidate == current_cn
        or (
            comment_cn_name
            and current_cn == comment_cn_name
            and not _looks_like_generic_local_cn_name(comment_cn_name)
            and not _looks_like_low_quality_symbol_cn(comment_cn_name, raw_ident=_safe_strip((item or {}).get("name")))
            and reason != "声明注释"
        )
        or _looks_like_bad_canonical_name(candidate, raw_ident=_safe_strip((item or {}).get("name")))
        or _looks_like_generic_local_cn_name(candidate)
        or _looks_like_low_quality_symbol_cn(candidate, raw_ident=_safe_strip((item or {}).get("name")))
    ):
        return
    should_replace = (
        _is_missing_gap_text(current_cn)
        or _looks_like_generic_local_cn_name(current_cn)
        or _looks_like_low_quality_symbol_cn(current_cn, raw_ident=_safe_strip((item or {}).get("name")))
        or reason in {"声明注释", "数据流候选", "赋值来源"}
        and current_cn.endswith(("临时值", "临时量", "缓存值", "当前值", "结果值"))
    )
    if not should_replace:
        return
    item["profile_cn_candidate"] = candidate
    item["profile_cn_reason"] = reason
    item["cn_name"] = candidate


def _should_accept_refined_local_cn(
    candidate: str,
    *,
    current_cn: str,
    item: dict,
    body: str,
    neighbor_symbols: Sequence[str],
    comment_desc: str = "",
    cfg: Optional["GenConfig"] = None,
) -> bool:
    from . import naming as naming_utils

    return naming_utils.should_accept_refined_local_cn(
        candidate,
        current_cn=current_cn,
        item=item,
        body=body,
        neighbor_symbols=neighbor_symbols,
        comment_desc=comment_desc,
        cfg=cfg,
        backend_module=sys.modules[__name__],
    )


def _local_cn_needs_ai_refine(
    item: Optional[dict],
    *,
    body: str,
    neighbor_symbols: Sequence[str],
    comment_desc: str = "",
    cfg: Optional["GenConfig"] = None,
) -> bool:
    from . import naming as naming_utils

    return naming_utils.local_cn_needs_ai_refine(
        item,
        body=body,
        neighbor_symbols=neighbor_symbols,
        comment_desc=comment_desc,
        cfg=cfg,
        backend_module=sys.modules[__name__],
    )


def _normalize_key_text(key: str) -> str:
    return re.sub(r"\s+", "", str(key or "").strip().lower())


def _edit_distance(a: str, b: str) -> int:
    if a == b:
        return 0
    if not a:
        return len(b)
    if not b:
        return len(a)
    prev = list(range(len(b) + 1))
    for i, ca in enumerate(a, start=1):
        curr = [i]
        for j, cb in enumerate(b, start=1):
            cost = 0 if ca == cb else 1
            curr.append(min(prev[j] + 1, curr[j - 1] + 1, prev[j - 1] + cost))
        prev = curr
    return prev[-1]


def _similarity_ratio(a: str, b: str) -> float:
    if not a and not b:
        return 1.0
    max_len = max(len(a), len(b))
    if max_len == 0:
        return 1.0
    return 1.0 - (_edit_distance(a, b) / max_len)


def _best_key_match(key: str, candidates: Sequence[str], max_dist: int = 2,
                    min_ratio: float = 0.8) -> Optional[str]:
    key_norm = _normalize_key_text(key)
    best = None
    best_dist = None
    tie = False
    for cand in candidates:
        cand_norm = _normalize_key_text(cand)
        dist = _edit_distance(key_norm, cand_norm)
        ratio = _similarity_ratio(key_norm, cand_norm)
        if ratio < min_ratio:
            continue
        if dist <= max_dist:
            if best is None or dist < best_dist:
                best = cand
                best_dist = dist
                tie = False
            elif dist == best_dist:
                tie = True
    if tie:
        return None
    return best


def _is_empty_value(val: Any) -> bool:
    if val is None:
        return True
    if isinstance(val, str):
        return not val.strip()
    if isinstance(val, (dict, list, tuple, set)):
        return len(val) == 0
    return False


def _merge_value(existing, new):
    if _is_empty_value(existing) and not _is_empty_value(new):
        return new
    return existing if existing is not None else new


def _coerce_dict_keys(data: Any, expected_keys: Sequence[str],
                      aliases: Optional[dict[str, str]] = None,
                      max_dist: int = 2,
                      min_ratio: float = 0.8):
    if not isinstance(data, dict):
        return data
    expected = [str(k) for k in expected_keys]
    alias_map = aliases or {}
    out = {}
    for k, v in data.items():
        key = str(k).strip()
        if key in expected:
            out[key] = _merge_value(out.get(key), v)
            continue
        alias = alias_map.get(key)
        if alias and alias in expected:
            out[alias] = _merge_value(out.get(alias), v)
            continue
        best = _best_key_match(key, expected, max_dist=max_dist, min_ratio=min_ratio)
        if best:
            out[best] = _merge_value(out.get(best), v)
        else:
            out[key] = v
    return out


def _unwrap_named_result_dict(
    data: Any,
    expected_keys: Sequence[str],
    aliases: Optional[dict[str, str]] = None,
) -> Any:
    if not isinstance(data, dict):
        return data
    expected = {str(k).strip() for k in (expected_keys or []) if str(k).strip()}
    alias_keys = {str(k).strip() for k in (aliases or {}).keys() if str(k).strip()}
    wrapper_keys = (
        "func",
        "result",
        "data",
        "output",
        "response",
        "answer",
        "payload",
    )

    def _has_expected_keys(candidate: Any) -> bool:
        if not isinstance(candidate, dict):
            return False
        keys = {str(k).strip() for k in candidate.keys()}
        return bool(keys & (expected | alias_keys))

    if _has_expected_keys(data):
        return data
    for key in wrapper_keys:
        nested = data.get(key)
        if _has_expected_keys(nested):
            return nested
    for value in data.values():
        if _has_expected_keys(value):
            return value
    return data


def _normalize_ai_var_keys(ai_map: Any, expected_names: Sequence[str], max_dist: int = 2,
                           min_ratio: float = 0.8):
    if not isinstance(ai_map, dict):
        return ai_map
    expected = [str(n).strip() for n in (expected_names or []) if str(n).strip()]
    if not expected:
        return ai_map
    out = {}
    for k, v in ai_map.items():
        key = str(k).strip()
        if key in expected:
            out[key] = _merge_value(out.get(key), v)
            continue
        best = _best_key_match(key, expected, max_dist=max_dist, min_ratio=min_ratio)
        if best:
            out[best] = _merge_value(out.get(best), v)
        else:
            out[key] = v
    return out


def _sanitize_one_call_section_map(
    section_map: Any,
    *,
    allowed_keys: Sequence[str],
    peer_keys: Sequence[str] = (),
) -> tuple[Any, dict[str, Any]]:
    """
    对 one-call 返回的 locals/params/logic 做严格白名单裁剪。
    仅保留 allowed_keys 中的条目，其余全部丢弃并返回 guard 信息。
    """
    if not isinstance(section_map, dict):
        return section_map, {
            "raw_keys": [],
            "allowed_keys": [],
            "dropped_keys": [],
            "peer_hits": [],
            "unexpected_nonempty": False,
        }

    allowed = {
        _safe_strip(k)
        for k in (allowed_keys or [])
        if _safe_strip(k)
    }
    peer = {
        _safe_strip(k)
        for k in (peer_keys or [])
        if _safe_strip(k)
    }

    raw_keys: list[str] = []
    dropped_keys: list[str] = []
    peer_hits: list[str] = []
    kept: dict[str, Any] = {}

    for key, value in section_map.items():
        name = _safe_strip(key)
        if not name:
            continue
        raw_keys.append(name)
        if name in allowed:
            kept[name] = value
            continue
        dropped_keys.append(name)
        if name in peer:
            peer_hits.append(name)

    return kept, {
        "raw_keys": raw_keys,
        "allowed_keys": sorted(allowed),
        "dropped_keys": dropped_keys,
        "peer_hits": peer_hits,
        "unexpected_nonempty": bool(raw_keys and not allowed),
    }








def _capture_title_call_debug(cfg: Optional[GenConfig], stage: str, prompt: str, result: Any) -> dict[str, Any]:
    debug = utils._get_last_llm_json_debug(cfg)
    payload: dict[str, Any] = {
        "stage": _safe_strip(stage),
        "prompt_len": len(prompt or ""),
        "prompt_sha256": _safe_strip(debug.get("prompt_sha256")),
        "from_cache": bool(debug.get("from_cache")),
        "error": _safe_strip(debug.get("error")),
        "proxy_source": _safe_strip(debug.get("proxy_source")),
        "proxy_url": _safe_strip(debug.get("proxy_url")),
        "raw_content": _safe_strip(debug.get("raw_content")),
        "content": _safe_strip(debug.get("content")),
        "parsed_type": _safe_strip(debug.get("parsed_type")) or type(result).__name__,
        "parsed_keys": tuple(str(x) for x in (debug.get("parsed_keys") or ()) if _safe_strip(x)),
        "parsed_preview": _safe_strip(debug.get("parsed_preview")),
        "result_preview": utils._debug_preview_json(result),
    }
    if isinstance(result, dict):
        payload["func_cn_name"] = _safe_strip(result.get("func_cn_name"))
        payload["desc"] = _safe_strip(result.get("desc"))
        payload["candidates"] = tuple(
            _safe_strip(x) for x in (result.get("candidates") or ()) if _safe_strip(x)
        )
    return payload


_CJK_RE = re.compile(r"[\u4e00-\u9fff]")
_IDENT_TOKEN_RE = re.compile(r"[A-Z]+(?=[A-Z][a-z])|[A-Z]?[a-z]+|[A-Z]+|\d+")
_IDENT_SKIP_TOKENS = {
    "u8", "u16", "u32", "u64",
    "i8", "i16", "i32", "i64",
    "l", "g", "s", "v",  # 常见局部/全局/静态变量前缀
    "ls", "gs", "ss", "vs", "ps", "vp", "gp", "lp", "sp", "cp", "tp",
}
_IDENT_CN_MAP = text_utils._IDENT_CN_MAP


def _contains_cjk(text: str) -> bool:
    return text_utils._contains_cjk(text)


def _split_ident_tokens(name: str) -> list[str]:
    return text_utils._split_ident_tokens(name)


def _guess_cn_from_ident(name: str, glossary: Optional[dict[str, str]] = None) -> str:
    return text_utils._guess_cn_from_ident(name, glossary=glossary)


def collect_usage_snippets(body: str, symbol: str, max_hits: int = 6):
    """
    收集 symbol 的使用片段，优先选择"信息密度高"的行：条件判断、函数调用、赋值、结构体成员访问等。
    """
    symbol = (symbol or "").strip()
    if not body or not symbol or max_hits <= 0:
        return []

    # 仅按标识符边界匹配，避免 foo 命中 foobar
    ident_re = re.compile(rf"(?<![A-Za-z0-9_]){re.escape(symbol)}(?![A-Za-z0-9_])")

    def score(line: str) -> int:
        s = line.strip()
        if not s:
            return 0
        if re.match(r"^\s*//", s):
            return 0
        if "/*" in s and "*/" not in s:
            # 避免多行注释开头的噪声
            return 0

        sc = 1
        # 条件/控制结构
        if re.search(r"\b(if|else\s+if|while|for|switch)\b", s):
            sc += 6
        # 返回值/错误码/状态检查
        if re.search(r"\breturn\b", s):
            sc += 3
        # 赋值：symbol 在左侧更关键
        if "=" in s and not re.search(r"(==|!=|>=|<=)", s):
            lhs, rhs = s.split("=", 1)
            if ident_re.search(lhs):
                sc += 6
            elif ident_re.search(rhs):
                sc += 3
        # 作为函数调用参数/参与调用
        if "(" in s and ")" in s:
            sc += 3
        # 结构体成员/数组访问通常更能体现语义
        if re.search(r"(\.|->)\s*\w+", s):
            sc += 3
        if "[" in s and "]" in s:
            sc += 2
        # 常见语义函数
        if re.search(r"\b(memset|memcpy|memmove|strcpy|strncpy|strcmp|strncmp)\b", s, flags=re.I):
            sc += 4

        # 越短越不容易塞满 prompt，但不能太短
        if len(s) > 200:
            sc -= 2
        return sc

    candidates = []
    for i, line in enumerate(_join_c_line_continuations(body).splitlines(), start=1):
        if not ident_re.search(line):
            continue
        s = line.strip()
        if not s:
            continue
        if len(s) > 220:
            s = s[:220] + "..."
        candidates.append((score(line), i, s))

    # 先按分数降序，行号升序（保证稳定）
    candidates.sort(key=lambda t: (-t[0], t[1]))

    hits = []
    seen = set()
    for sc, i, s in candidates:
        if s in seen:
            continue
        seen.add(s)
        hits.append({"line": i, "code": s})
        if len(hits) >= max_hits:
            break
    return hits



def _is_local_provider(cfg: Optional[GenConfig]) -> bool:
    return str(getattr(cfg, "ai_provider", "local") or "local").strip().lower() == "local"


def _trim_text_chars(text: str, max_chars: int) -> str:
    s = _safe_text(text)
    max_chars = max(0, int(max_chars or 0))
    if max_chars <= 0 or len(s) <= max_chars:
        return s
    half = max(32, max_chars // 2)
    return s[:half] + "\n/* ... 省略 ... */\n" + s[-half:]


def _trim_body_for_ai(body: str, cfg: Optional[GenConfig], *, one_call: bool = False) -> str:
    text = text_utils.trim_body(body, max_lines=utils.cfg_get_int(cfg, "ai_body_max_lines", 300))
    if _is_local_provider(cfg):
        max_lines = utils.cfg_get_int(
            cfg,
            "local_one_call_body_max_lines" if one_call else "local_body_max_lines",
            180 if one_call else 220,
        )
        text = text_utils.trim_body(text, max_lines=max_lines)
        max_chars = utils.cfg_get_int(
            cfg,
            "local_one_call_body_max_chars" if one_call else "local_body_max_chars",
            12000 if one_call else 16000,
        )
        text = _trim_text_chars(text, max_chars)
    return text


def _trim_unknowns_for_ai(unknowns: Sequence[dict], cfg: Optional[GenConfig]) -> list[dict]:
    items = list(unknowns or [])
    if not _is_local_provider(cfg):
        return items
    limit = utils.cfg_get_int(cfg, "local_one_call_unknown_limit", 48)
    if limit > 0 and len(items) > limit:
        head_n = max(1, limit // 2)
        tail_n = max(1, limit - head_n)
        items = items[:head_n] + items[-tail_n:]
    code_chars = utils.cfg_get_int(cfg, "local_one_call_unknown_code_max_chars", 240)
    trimmed: list[dict] = []
    for item in items:
        item2 = dict(item or {})
        item2["code"] = _trim_text_chars(_safe_text(item2.get("code")), code_chars)
        item2["code_cn"] = _trim_text_chars(_safe_text(item2.get("code_cn")), code_chars)
        trimmed.append(item2)
    return trimmed


def _should_disable_local_one_call(body: str, unknowns: Sequence[dict], cfg: Optional[GenConfig]) -> bool:
    if not bool(getattr(cfg, "auto_disable_large_one_call", True)):
        return False
    if not _is_local_provider(cfg):
        return False
    if _is_small_model_strict_mode(cfg) and _looks_like_state_machine_body(body):
        return True
    body_chars = len(_safe_text(body))
    body_lines = len(_safe_text(body).splitlines())
    unknown_count = len(unknowns or [])
    if body_chars >= utils.cfg_get_int(cfg, "local_one_call_disable_body_chars", 20000):
        return True
    if body_lines >= utils.cfg_get_int(cfg, "local_one_call_disable_body_lines", 450):
        return True
    if unknown_count >= utils.cfg_get_int(cfg, "local_one_call_disable_unknowns", 80):
        return True
    return False


def build_func_prompt(func_info, body, comment_info, cfg: Optional[GenConfig] = None):
    from . import ai as ai_utils

    return ai_utils.build_func_prompt(func_info, body, comment_info, cfg)


def build_func_title_prompt(func_info, body, comment_info, cfg: Optional[GenConfig] = None):
    from . import ai as ai_utils

    return ai_utils.build_func_title_prompt(func_info, body, comment_info, cfg)


def build_func_title_retry_prompt(func_info, body, comment_info, cfg: Optional[GenConfig] = None):
    from . import ai as ai_utils

    return ai_utils.build_func_title_retry_prompt(func_info, body, comment_info, cfg)


def build_llm_evidence_pack(func_data: dict[str, Any], cfg: Optional[GenConfig] = None, task: str = "func"):
    from . import ai as ai_utils

    return ai_utils.build_llm_evidence_pack(func_data, cfg, task=task)


def build_local_naming_prompt(
    payload: Sequence[dict[str, Any]],
    *,
    func_cn_name: str,
    func_desc: str,
    function_semantic_summary: dict[str, Any],
    prompt_glossary: dict[str, str],
    remembered_terms: dict[str, str],
    cfg: Optional[GenConfig] = None,
) -> str:
    from . import ai as ai_utils

    return ai_utils.build_local_naming_prompt(
        payload,
        func_cn_name=func_cn_name,
        func_desc=func_desc,
        function_semantic_summary=function_semantic_summary,
        prompt_glossary=prompt_glossary,
        remembered_terms=remembered_terms,
        cfg=cfg,
    )



def build_symbol_prompt(symbol, vtype, body, func_cn_name: str = "", func_desc: str = ""):
    from . import ai as ai_utils

    return ai_utils.build_symbol_prompt(
        symbol,
        vtype,
        body,
        func_cn_name=func_cn_name,
        func_desc=func_desc,
        backend_module=sys.modules[__name__],
    )


def ai_suggest_for_locals_batch(missing_vars, locals_all, body: str, cfg: GenConfig,
                                func_cn_name: str = "", func_desc: str = "", glossary=None):
    """批量为缺失局部变量建议中文名与用途。"""
    from . import ai as ai_utils

    return ai_utils.ai_suggest_for_locals_batch(
        missing_vars,
        locals_all,
        body,
        cfg,
        func_cn_name=func_cn_name,
        func_desc=func_desc,
        glossary=glossary,
        _runtime_module=sys.modules.get(__name__),
    )


# -------- 单函数一次调用：函数级 + 局部/参数 + 逻辑润色 --------

def ai_suggest_bundle_one_call(
    func_info: dict,
    body: str,
    comment_info: dict,
    local_vars: Sequence[dict],
    params: Sequence[dict],
    missing_locals: Sequence[str],
    missing_params: Sequence[str],
    unknowns: Sequence[dict],
    cfg: GenConfig,
    file_context: Optional[dict] = None,
    evidence_pack: Optional[dict[str, Any]] = None,
):
    """
    单次调用本地模型，返回一个 JSON：
    {
      "func": {"func_cn_name": "...", "desc": "..."},
      "locals": {"var": {"cn_name": "...", "usage": "..."} },
      "params": {"param": {"cn_name": "..."} },
      "return": {"cn_name": "..." },
      "logic": {"idx": "..."}  # 只允许替换给定 idx 的动作行
    }
    """
    file_context = file_context or {}
    evidence_pack = dict(evidence_pack or {})
    # 检查停止信号
    if stop_requested(cfg):
        return {}
    glossary = file_context.get("glossary") or DOMAIN_GLOSSARY
    body_for_prompt = _trim_body_for_ai(body, cfg, one_call=True)
    prompt_glossary = _filter_glossary_for_prompt(glossary, [body_for_prompt, json.dumps(func_info or {}, ensure_ascii=False)], limit=12)
    missing_local_set = {str(x).strip() for x in (missing_locals or []) if str(x).strip()}
    missing_param_set = {str(x).strip() for x in (missing_params or []) if str(x).strip()}

    def _collect(name: str) -> str:
        hits = collect_usage_snippets(body, name, 4)
        return "\n".join([f"{h['line']}: {h['code']}" for h in hits]) or "(无明显片段)"

    type_map = {v.get("name"): (v.get("type") or "") for v in (local_vars or []) if v.get("name")}
    local_lookup = {v.get("name"): v for v in (local_vars or []) if v.get("name")}
    regression_rename_set = set(getattr(cfg, "_ai_regression_allow_rename", set()) or set())
    local_payload = []
    for name in (missing_locals or []):
        item = local_lookup.get(name) or {}
        local_payload.append({
            "name": name,
            "type": type_map.get(name, ""),
            "context": _collect(name),
            "locked_cn": _safe_strip(item.get("cn_name")) if name not in regression_rename_set else "",
            "role_hint": _safe_strip(item.get("role_hint") or item.get("comment_hint")),
            "allow_rename": name in regression_rename_set,
        })

    param_payload = []
    for p in (params or []):
        n = (p or {}).get("name") or ""
        if not n:
            continue
        if n in set(missing_params or []):
            param_payload.append({
                "name": n,
                "type": (p or {}).get("type") or "",
                "context": _collect(n),
                "locked_cn": "" if n in regression_rename_set else "",
                "role_hint": "",
                "allow_rename": n in regression_rename_set,
            })

    unknown_payload = [{
        "idx": u.get("idx"),
        "code": (u.get("code") or "").strip(),
        "code_cn": (u.get("code_cn") or "").strip(),
        "comment_hints": list(u.get("comment_hints") or []),
        "polish_only": bool(u.get("polish_only")),
    } for u in _trim_unknowns_for_ai(unknowns, cfg)]
    remembered_terms = _collect_preferred_symbol_names(
        [func_info.get("func_name", "")]
        + [v.get("name", "") for v in (local_vars or [])]
        + [p.get("name", "") for p in (params or [])],
        limit=32,
    )
    quality_feedback = _safe_strip(getattr(cfg, "_ai_quality_feedback", ""))
    quality_focus_symbols = list(getattr(cfg, "_ai_quality_focus_symbols", ()) or ())
    quality_block = ""
    if quality_feedback or quality_focus_symbols:
        lines = ["质量回归要求："]
        if quality_feedback:
            lines.append(f"- {quality_feedback}")
        if quality_focus_symbols:
            lines.append(f"- 以下标识符不要直接保留英文符号名：{json.dumps(quality_focus_symbols, ensure_ascii=False)}")
        lines.append("- 若逻辑说明中涉及这些标识符，优先改写为中文名称或中文动作短语。")
        quality_block = "\n".join(lines)
    logic_required_rule = ""
    if unknown_payload:
        logic_required_rule = (
            f"\n- 【必须】logic 必须覆盖以下 idx："
            f"{json.dumps([str(item.get('idx')) for item in unknown_payload if item.get('idx') is not None], ensure_ascii=False)}；"
            "即使不确定，也要基于原 code_cn 给出保守中文动作，不得返回空 logic。"
        )

    if _is_small_model_strict_mode(cfg):
        prompt = f"""只输出单个JSON对象。
规则:
1. func 只做轻量润色，locals/params 不得重命名 locked_cn 项。
2. locals/params 中 allow_rename 为 false 时，不要改名；不确定返回空字符串。
3. logic 只替换给定 idx 的动作行；不新增/删除流程行，不改 IF/FOR/NEXT/END IF 等结构。
4. 输出字段必须固定为 func/locals/params/return/logic。
5. 缺失项返回空字符串或空对象。
6. 若 unknowns 非空，logic 不得为空，必须按 idx 返回每一项的动作短句。
{quality_block}
evidence:{json.dumps(evidence_pack or {}, ensure_ascii=False, separators=(",", ":"))}
函数:{json.dumps(func_info or {}, ensure_ascii=False, separators=(",", ":"))}
注释:{json.dumps(comment_info or {}, ensure_ascii=False, separators=(",", ":"))}
术语:{json.dumps(prompt_glossary, ensure_ascii=False, separators=(",", ":"))}
已确认:{json.dumps(remembered_terms, ensure_ascii=False, separators=(",", ":"))}
locals:{json.dumps(local_payload, ensure_ascii=False, separators=(",", ":"))}
params:{json.dumps(param_payload, ensure_ascii=False, separators=(",", ":"))}
unknowns:{json.dumps(unknown_payload, ensure_ascii=False, separators=(",", ":"))}
code:{json.dumps(body_for_prompt, ensure_ascii=False)}
输出:{{"func":{{"func_cn_name":"","desc":""}},"locals":{{}},"params":{{}},"return":{{"cn_name":""}},"logic":{{}}}}"""
    else:
        prompt = f"""你是严谨的嵌入式软件详细设计说明撰写助手。
请基于给定 C 函数代码与上下文，完成以下任务，并严格输出 JSON（不要包含解释、不要代码块）：

任务A：函数中文名与功能说明（简洁、贴合语义，避免泛化）。
任务B：为缺失的局部变量/参数给出中文名称与用途（用途 10~20 字，不确定处标注"(推测)"）。
任务C：为需要补全或润色的流程语句条目输出更自然的中文动作说明（建议 8~15 字，不要句号）。

规则：
- 不要根据标识符后缀（_u16/_u32/i16/i32 等类型后缀）推断含义；后缀仅表示类型。
- 术语表优先：{json.dumps(prompt_glossary, ensure_ascii=False)}。
- 已确认术语优先：{json.dumps(remembered_terms, ensure_ascii=False)}。若某个标识符已存在中文名，必须沿用，不得改名。
- 优先使用 evidence.authoritative 和 evidence.high_confidence；derived 可参考；low_confidence 只能辅助，不得照抄。
- 输出必须是单个 JSON 对象，字段缺失则给空字符串/空对象。
- 对赋值语句（含结构体成员赋值）避免机械重复"设置/更新变量"，优先用更具体动词：写入/赋给/清零/置位/累加/清除位/拷贝/计算/获取。
- 【重要】TaskC 只能改写给定 idx 对应的动作行；不得新增、删除、重排流程，不得输出 IF/FOR/NEXT/END IF 等结构行。
- 【重要】TaskC 的动作说明只写"动作本身"，不要追加目的/用途/解释性尾巴（例如包含"供…/用于…/以便…"），也不要写"并存储…写入"这种冗余复述。
- TaskC 可参考 comment_hints，但 history/debug/purpose 类提示绝不能直接出现在最终动作说明中。
- locals/params payload 中 allow_rename=false 的项不得重命名；locked_cn 只读。
{logic_required_rule}
{quality_block}

函数信息：
{json.dumps(func_info or {}, ensure_ascii=False, indent=2)}

已有注释信息（可能为空）：
{json.dumps(comment_info or {}, ensure_ascii=False, indent=2)}

函数代码（上下文）：
{body_for_prompt}

需要补全的局部变量：
{json.dumps(local_payload, ensure_ascii=False, indent=2)}

需要补全的参数：
{json.dumps(param_payload, ensure_ascii=False, indent=2)}

需要补全/润色的流程语句条目（idx 对应静态流程图行号；polish_only=true 表示只能润色原动作，不得改变事实）：
{json.dumps(unknown_payload, ensure_ascii=False, indent=2)}

分级证据包：
{json.dumps(evidence_pack or {}, ensure_ascii=False, indent=2)}

    输出 JSON（严格，字段名必须完全一致，不要代码块）：
    {{
      "func": {{"func_cn_name": "函数中文名", "desc": "功能说明"}},
      "locals": {{
        "局部变量标识符": {{"cn_name": "中文名称", "usage": "用途描述(10~20字)"}}
      }},
      "params": {{
        "参数标识符": {{"cn_name": "中文名称"}}
      }},
      "return": {{"cn_name": "返回值中文名"}},
      "logic": {{
        "0": "润色后的动作说明",
        "5": "润色后的动作说明"
      }}
    }}
    说明：
    - "locals" 的 value 只能用 "cn_name"/"usage"（不要用 name/purpose）。
    - "params" 的 value 只能用 "cn_name"。
    - "logic" 必须是 "idx -> string" 的映射（不要用 steps 数组）。"""
    func_name = (func_info or {}).get("func_name") or (func_info or {}).get("name") or ""
    bundle_event_base = {
        "file": str(getattr(cfg, "_current_file", "") or ""),
        "func_name": func_name,
        "func_index": int(getattr(cfg, "_current_func_index", 0) or 0),
        "func_pos": int(getattr(cfg, "_current_func_pos", 0) or 0),
    }

    def _emit_bundle_event(event_type: str, **extra) -> None:
        payload = {"type": event_type}
        payload.update(bundle_event_base)
        payload.update(extra)
        gui_event(cfg, payload)

    def _has_bad_logic_residual_ident(text: str) -> bool:
        value = _safe_strip(text)
        if not value:
            return False
        if value.startswith("#") or any(ch in value for ch in "{}"):
            return True
        return bool(
            re.search(
                r"\b(?:[lgsvcp]_|[A-Za-z_]\w*_(?:u|i|f)(?:8|16|32|64)?\b|[A-Z][A-Z0-9_]{2,})",
                value,
            )
        )

    def _logic_specificity_score(text: str) -> int:
        value = re.sub(r"\s+", "", _safe_strip(text))
        if not value:
            return 0
        generic_terms = ("相关处理", "相关计算", "相关更新", "执行处理", "更新数据", "处理数据", "执行操作")
        score = 0
        if _contains_cjk(value):
            score += min(8, len(re.findall(r"[\u4e00-\u9fff]", value)) // 2)
        for term in ("写入", "计算", "清零", "置位", "拷贝", "打包", "组装", "同步", "记录", "比较", "滤波", "校验"):
            if term in value:
                score += 2
        for term in ("状态", "故障", "模式", "输出", "输入", "缓存", "结果", "数据字", "通道", "作动器"):
            if term in value:
                score += 1
        for term in generic_terms:
            if term in value:
                score -= 4
        return score

    unknown_by_idx = {
        str(u.get("idx")): _safe_strip(u.get("code_cn") or u.get("code") or "")
        for u in (unknowns or [])
        if u.get("idx") is not None
    }

    def _is_logic_polish_regression(idx_v: Any, text: str) -> bool:
        old_text = unknown_by_idx.get(str(idx_v), "")
        if not old_text:
            return False
        old_score = _logic_specificity_score(old_text)
        new_score = _logic_specificity_score(text)
        return old_score >= 5 and new_score + 3 < old_score

    def _sanitize_logic_part(logic_part) -> tuple[dict[str, str], dict]:
        raw_logic_map: dict[str, str] = {}
        allowed_logic_keys = [str(u.get("idx")) for u in (unknowns or []) if u.get("idx") is not None]

        if isinstance(logic_part, dict):
            if isinstance(logic_part.get("steps"), list):
                logic_iter = logic_part.get("steps") or []
            else:
                logic_iter = list((logic_part or {}).items())
        elif isinstance(logic_part, list):
            logic_iter = logic_part
        else:
            logic_iter = []

        if logic_iter and isinstance(logic_iter[0], tuple):
            for idx_v, value in logic_iter:
                if isinstance(value, str):
                    txt = _sanitize_ai_logic_action(_safe_strip(value))
                elif isinstance(value, dict):
                    txt = _sanitize_ai_logic_action(_safe_strip(value.get("action") or value.get("text") or value.get("desc")))
                else:
                    txt = ""
                if txt and _is_control_logic_line(txt):
                    continue
                if txt and _has_bad_logic_residual_ident(txt):
                    continue
                if txt and _is_logic_polish_regression(idx_v, txt):
                    continue
                if idx_v is None or not txt:
                    continue
                raw_logic_map[str(idx_v)] = txt
        else:
            for item in (logic_iter or []):
                if not isinstance(item, dict):
                    continue
                idx_v = item.get("idx")
                txt = _sanitize_ai_logic_action(_safe_strip(item.get("action") or item.get("text") or item.get("desc")))
                if txt and _is_control_logic_line(txt):
                    continue
                if txt and _has_bad_logic_residual_ident(txt):
                    continue
                if txt and _is_logic_polish_regression(idx_v, txt):
                    continue
                if idx_v is None or not txt:
                    continue
                raw_logic_map[str(idx_v)] = txt

        strict_logic, logic_guard = _sanitize_one_call_section_map(
            raw_logic_map,
            allowed_keys=allowed_logic_keys,
        )
        return strict_logic, logic_guard

    _emit_bundle_event(
        "ai_bundle_start",
        locals_expected=len(local_payload),
        params_expected=len(param_payload),
        logic_expected=len(unknown_payload),
    )
    try:
        js = call_llm_json(
            prompt,
            cfg,
            log_title=f"AI 合并润色：原始输出 func={func_name}",
            log_full_output=False,
        )
    except Exception as e:
        _emit_bundle_event("ai_bundle_end", ok=False, error=str(e))
        raise
    # 用户取消时快速返回，不再处理 AI 结果
    if getattr(cfg, "_user_cancelled", False) or stop_requested(cfg):
        _emit_bundle_event("ai_bundle_end", ok=False, error="用户取消")
        return {}
    if isinstance(js, dict):
        ai_debug_log(cfg, "one_call_response_keys", {
            "keys": sorted(list(js.keys())),
            "has_func": isinstance(js.get("func"), dict),
            "locals_keys": len(js.get("locals") or {}) if isinstance(js.get("locals"), dict) else -1,
            "params_keys": len(js.get("params") or {}) if isinstance(js.get("params"), dict) else -1,
            "logic_keys": len(js.get("logic") or {}) if isinstance(js.get("logic"), dict) else -1,
        })
    if not isinstance(js, dict):
        _emit_bundle_event("ai_bundle_end", ok=False, error="AI 返回非 JSON 对象")
        return {}
    js = _coerce_dict_keys(
        js,
        ("func", "locals", "params", "return", "logic"),
        aliases={
            "local": "locals",
            "param": "params",
            "returns": "return",
        },
        max_dist=utils.cfg_get_int(cfg, "max_dist", 2),
        min_ratio=utils.cfg_get_float(cfg, "min_ratio", 0.8),
    )
    try:
        utils._tool_log_write_block(cfg, f"AI 合并润色/func func={func_name}", json.dumps(js.get("func") or {}, ensure_ascii=False, indent=2))
        utils._tool_log_write_block(cfg, f"AI 合并润色/locals func={func_name}", json.dumps(js.get("locals") or {}, ensure_ascii=False, indent=2))
        utils._tool_log_write_block(cfg, f"AI 合并润色/params func={func_name}", json.dumps(js.get("params") or {}, ensure_ascii=False, indent=2))
        utils._tool_log_write_block(cfg, f"AI 合并润色/return func={func_name}", json.dumps(js.get("return") or {}, ensure_ascii=False, indent=2))
        utils._tool_log_write_block(cfg, f"AI 合并润色/logic func={func_name}", json.dumps(js.get("logic") or {}, ensure_ascii=False, indent=2))
    except Exception:
        pass
    func_part = js.get("func")
    if isinstance(func_part, dict):
        js["func"] = _coerce_dict_keys(
            func_part,
            ("func_cn_name", "desc", "confidence"),
            aliases={
                "func_cn": "func_cn_name",
                "func_name_cn": "func_cn_name",
                "func_desc": "desc",
                "function_desc": "desc",
                "description": "desc",
            },
            max_dist=utils.cfg_get_int(cfg, "max_dist", 2),
            min_ratio=utils.cfg_get_float(cfg, "min_ratio", 0.8),
        )
    locals_part = js.get("locals")
    if isinstance(locals_part, dict):
        expected_local_names = [v.get("name") for v in (local_vars or []) if v.get("name")]
        locals_part = _normalize_ai_var_keys(
            locals_part,
            expected_local_names,
            max_dist=utils.cfg_get_int(cfg, "max_dist", 2),
            min_ratio=utils.cfg_get_float(cfg, "min_ratio", 0.8),
        )
        cleaned_locals = {}
        for name, item in locals_part.items():
            if isinstance(item, dict):
                item = _coerce_dict_keys(
                    item,
                    ("cn_name", "usage", "confidence"),
                    aliases={
                        "name": "cn_name",
                        "cn": "cn_name",
                        "cname": "cn_name",
                        "cnname": "cn_name",
                        "purpose": "usage",
                        "desc": "usage",
                        "description": "usage",
                    },
                    max_dist=utils.cfg_get_int(cfg, "max_dist", 2),
                    min_ratio=utils.cfg_get_float(cfg, "min_ratio", 0.8),
                )
            elif isinstance(item, str):
                item = {"cn_name": item}
            locked_cn = _safe_strip((local_lookup.get(name) or {}).get("cn_name"))
            if isinstance(item, dict):
                cn_name = _safe_strip(item.get("cn_name"))
                usage = _sanitize_ai_usage_text(item.get("usage"))
                if locked_cn:
                    cn_name = ""
                elif cn_name and _is_strict_symbol_candidate_rejected(cn_name, raw_ident=name):
                    cn_name = ""
                item["cn_name"] = cn_name
                item["usage"] = usage
            cleaned_locals[name] = item
        strict_locals, locals_guard = _sanitize_one_call_section_map(
            cleaned_locals,
            allowed_keys=missing_locals,
            peer_keys=[p.get("name") for p in (params or []) if p.get("name")],
        )
        js["locals"] = strict_locals
    else:
        locals_guard = {
            "raw_keys": [],
            "allowed_keys": [],
            "dropped_keys": [],
            "peer_hits": [],
            "unexpected_nonempty": False,
        }
    params_part = js.get("params")
    if isinstance(params_part, dict):
        expected_param_names = [p.get("name") for p in (params or []) if p.get("name")]
        params_part = _normalize_ai_var_keys(
            params_part,
            expected_param_names,
            max_dist=utils.cfg_get_int(cfg, "max_dist", 2),
            min_ratio=utils.cfg_get_float(cfg, "min_ratio", 0.8),
        )
        cleaned_params = {}
        for name, item in params_part.items():
            if isinstance(item, dict):
                item = _coerce_dict_keys(
                    item,
                    ("cn_name",),
                    aliases={"name": "cn_name"},
                    max_dist=utils.cfg_get_int(cfg, "max_dist", 2),
                    min_ratio=utils.cfg_get_float(cfg, "min_ratio", 0.8),
                )
            elif isinstance(item, str):
                item = {"cn_name": item}
            if isinstance(item, dict):
                cn_name = _safe_strip(item.get("cn_name"))
                if cn_name and _is_strict_symbol_candidate_rejected(cn_name, raw_ident=name):
                    cn_name = ""
                item["cn_name"] = cn_name
            cleaned_params[name] = item
        strict_params, params_guard = _sanitize_one_call_section_map(
            cleaned_params,
            allowed_keys=missing_params,
            peer_keys=[v.get("name") for v in (local_vars or []) if v.get("name")],
        )
        js["params"] = strict_params
    else:
        params_guard = {
            "raw_keys": [],
            "allowed_keys": [],
            "dropped_keys": [],
            "peer_hits": [],
            "unexpected_nonempty": False,
        }
    return_part = js.get("return")
    if isinstance(return_part, dict):
        js["return"] = _coerce_dict_keys(
            return_part,
            ("cn_name",),
            aliases={"name": "cn_name"},
            max_dist=utils.cfg_get_int(cfg, "max_dist", 2),
            min_ratio=utils.cfg_get_float(cfg, "min_ratio", 0.8),
        )
    logic_part = js.get("logic")
    strict_logic, logic_guard = _sanitize_logic_part(logic_part)
    js["logic"] = strict_logic
    guard_reasons: list[str] = []
    if locals_guard.get("peer_hits"):
        guard_reasons.append("locals_hit_param_names")
    if locals_guard.get("unexpected_nonempty"):
        guard_reasons.append("locals_unexpected_nonempty")
    if params_guard.get("peer_hits"):
        guard_reasons.append("params_hit_local_names")
    js["_guard"] = {
        "force_fallback": bool(guard_reasons),
        "reasons": tuple(dict.fromkeys(guard_reasons)),
        "locals": locals_guard,
        "params": params_guard,
    }
    if guard_reasons:
        vlog(cfg, f"函数 {func_name} 的 one-call 回包触发严格校验: {guard_reasons}")
        ai_debug_log(cfg, "one_call_strict_guard", {
            "func_name": func_name,
            "reasons": guard_reasons,
            "locals_guard": locals_guard,
            "params_guard": params_guard,
        })
    if logic_guard.get("dropped_keys"):
        vlog(cfg, f"函数 {func_name} 的 one-call logic 已过滤无效 idx: {logic_guard.get('dropped_keys')}")
        ai_debug_log(cfg, "one_call_logic_strict_guard", {
            "func_name": func_name,
            "allowed_keys": logic_guard.get("allowed_keys"),
            "dropped_keys": logic_guard.get("dropped_keys"),
        })
    _emit_bundle_event(
        "ai_bundle_end",
        ok=True,
        locals_expected=len(local_payload),
        locals_got=(len(js.get("locals") or {}) if isinstance(js.get("locals"), dict) else 0),
        params_expected=len(param_payload),
        params_got=(len(js.get("params") or {}) if isinstance(js.get("params"), dict) else 0),
        logic_expected=len(unknown_payload),
        logic_got=(len(js.get("logic") or {}) if isinstance(js.get("logic"), dict) else 0),
    )
    return js


def _clone_cfg(cfg: GenConfig, **overrides) -> GenConfig:
    from . import pipeline as pipeline_utils

    return pipeline_utils.clone_cfg(cfg, **overrides)


class FunctionBuildTaskError(Exception):
    """函数设计构建失败，并携带对应任务元数据。"""

    def __init__(self, task: dict, cause: Exception):
        super().__init__(str(cause))
        self.task = task
        self.cause = cause


def _make_regression_cfg(cfg: GenConfig, *, round_idx: int, meta: Optional[AIBuildMeta] = None) -> GenConfig:
    from . import pipeline as pipeline_utils

    return pipeline_utils.make_regression_cfg(
        cfg,
        round_idx=round_idx,
        meta=meta,
        backend_module=sys.modules[__name__],
    )


def _run_function_design_task(task: dict, cfg: GenConfig):
    from . import pipeline as pipeline_utils
    return pipeline_utils.run_function_design_task(
        task,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _build_function_design_task(func_data, module_req_prefix, index, cfg):
    from . import pipeline as pipeline_utils
    return pipeline_utils.build_function_design_task(
        func_data, module_req_prefix, index, cfg,
        backend_module=sys.modules[__name__],
    )


def _iter_function_design_results(tasks: Sequence[dict], cfg: GenConfig, *, on_submit: Optional[Callable[[dict], None]] = None):
    from . import pipeline as pipeline_utils

    yield from pipeline_utils._iter_function_design_results(
        tasks,
        cfg,
        on_submit=on_submit,
        backend_module=sys.modules[__name__],
    )

# -------- OpenAI/OpenRouter 客户端 & 调用 --------


import ast  # 文件顶部如果还没有，记得加这一行

_THINK_BLOCK_RE = re.compile(r"(?is)<think>.*?</think>\s*")
_URL_SCHEME_RE = re.compile(r"^[a-zA-Z][a-zA-Z0-9+.-]*://")
_HOST_LIKE_URL_RE = re.compile(r"^(?:localhost|[\w.-]+)(?::\d+)?(?:/.*)?$")

def strip_think_blocks(text: str) -> str:
    from . import ai as ai_utils

    return ai_utils.strip_think_blocks(text, backend_module=sys.modules[__name__])

def safe_json_loads(text: str):
    from . import ai as ai_utils

    return ai_utils.safe_json_loads(text, backend_module=sys.modules[__name__])


def _looks_like_utf8_mojibake(text: str) -> bool:
    from . import ai as ai_utils

    return ai_utils._looks_like_utf8_mojibake(text, backend_module=sys.modules[__name__])


def _repair_mojibake_text(text: Any) -> Any:
    from . import ai as ai_utils

    return ai_utils._repair_mojibake_text(text, backend_module=sys.modules[__name__])


def _parse_response_json_robust(resp) -> Any:
    from . import ai as ai_utils

    return ai_utils._parse_response_json_robust(resp, backend_module=sys.modules[__name__])


def normalize_chat_completion_url(u: str, *, default_scheme: str = "http") -> str:
    from . import ai as ai_utils

    return ai_utils.normalize_chat_completion_url(u, default_scheme=default_scheme)


def safe_save_docx(doc: Document, output: str) -> None:
    from . import render as render_utils

    return render_utils.safe_save_docx(doc, output)


def _get_ai_http_timeout(cfg: GenConfig, provider: str) -> tuple[float, float]:
    from . import ai as ai_utils

    return ai_utils._get_ai_http_timeout(cfg, provider)


def _ai_retry_sleep_seconds(attempt: int, attempts: int, cfg: GenConfig) -> float:
    from . import ai as ai_utils

    return ai_utils._ai_retry_sleep_seconds(
        attempt,
        attempts,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _build_curl_repro_command(url: str, headers: dict, payload_path: str) -> str:
    from . import ai as ai_utils

    return ai_utils._build_curl_repro_command(url, headers, payload_path)


def _write_ai_repro_bundle(
    cfg: Optional[GenConfig],
    *,
    provider: str,
    url: str,
    headers: dict,
    data: dict,
    prompt_sha: str,
    reason: str,
    tag: str,
) -> dict[str, str]:
    from . import ai as ai_utils

    return ai_utils._write_ai_repro_bundle(
        cfg,
        provider=provider,
        url=url,
        headers=headers,
        data=data,
        prompt_sha=prompt_sha,
        reason=reason,
        tag=tag,
        backend_module=sys.modules[__name__],
    )




def call_llm_json(
    prompt: str,
    cfg: GenConfig,
    *,
    log_title: str = "LLM 输出(完整)",
    log_preview: bool = True,
    log_full_output: bool = True,
    **kwargs,
):
    from . import ai as ai_utils

    return ai_utils.call_llm_json(
        prompt,
        cfg,
        log_title=log_title,
        log_preview=log_preview,
        log_full_output=log_full_output,
        _runtime_module=sys.modules.get(__name__),
        **kwargs,
    )


def validate_controlled_ai_candidate(*args, **kwargs):
    from . import ai as ai_utils

    return ai_utils.validate_controlled_ai_candidate(*args, **kwargs)


def _rank_function_title_candidate(candidate: str, *, func_name: str, comment_desc: str, examples: Sequence[dict[str, Any]]) -> int:
    from . import naming as naming_utils

    return naming_utils.rank_function_title_candidate(
        candidate,
        func_name=func_name,
        comment_desc=comment_desc,
        examples=examples,
    )


def ai_suggest_for_func(func_info, body, comment_info, cfg: GenConfig,
                        params=None, locals_=None, in_map=None, out_map=None,
                        file_context=None):
    from . import ai as ai_utils

    return ai_utils.ai_suggest_for_func(
        func_info,
        body,
        comment_info,
        cfg,
        params=params,
        locals_=locals_,
        in_map=in_map,
        out_map=out_map,
        file_context=file_context,
        _runtime_module=sys.modules.get(__name__),
    )


_C_IDENT_RE = re.compile(r"\b[A-Za-z_][A-Za-z0-9_]*\b", re.ASCII)
_C_KEYWORDS = {
    "if", "else", "for", "while", "do", "switch", "case", "default",
    "break", "continue", "goto", "return", "sizeof",
    "struct", "union", "enum", "typedef",
    "static", "extern", "const", "volatile", "register",
    "true", "false", "NULL",
}


def _replace_idents_for_logic(text: str, name_map: Optional[dict[str, str]]) -> str:
    from . import logic as logic_utils

    return logic_utils._replace_idents_for_logic(
        text,
        name_map,
        backend_module=sys.modules[__name__],
    )


_IDENT_TYPE_SUFFIX_RE = re.compile(r"_(?:u|i)(?:8|16|32|64|6)\b", re.IGNORECASE)
_IDENT_SCOPE_PREFIX_RE = re.compile(r"^(?:gc|gs|sc|gp|sp|lp|vp|fp|cp|tp|g|s|l|v|p)_")
_IDENT_NAME_SUFFIX_RE = re.compile(r"_(?:tp|pt|ptr|buf|arr|list|tbl|table|t|p|un)\b", re.IGNORECASE)


def _prettify_logic_ident(ident: str, name_map: Optional[dict[str, str]]) -> str:
    from . import logic as logic_utils

    return logic_utils._prettify_logic_ident(
        ident,
        name_map,
        backend_module=sys.modules[__name__],
    )


def _replace_idents_for_logic_ex(text: str, name_map: Optional[dict[str, str]], allow_member: bool) -> str:
    from . import logic as logic_utils

    return logic_utils._replace_idents_for_logic_ex(
        text,
        name_map,
        allow_member,
        backend_module=sys.modules[__name__],
    )


def _map_func_ident(func: str, name_map: Optional[dict[str, str]]) -> str:
    from . import logic as logic_utils

    return logic_utils._map_func_ident(
        func,
        name_map,
        backend_module=sys.modules[__name__],
    )


_MEMBER_PREFIX_RE = re.compile(
    r"[A-Za-z_\u4e00-\u9fff][A-Za-z0-9_\u4e00-\u9fff]*(?:\[[^\]]*\])?\s*\.(?!\d)\s*"
)

_MEMBER_CHAIN_RE = re.compile(
    r"(?P<base>[A-Za-z_]\w*(?:\s*\[[^\]]*\])?)\s*(?:\.|->)\s*"
    r"(?P<rest>[A-Za-z_]\w*(?:\s*(?:\.|->)\s*[A-Za-z_]\w*)*)"
)


def _collect_unresolved_macro_candidates(body: str, *, known_names: Optional[set[str]] = None) -> list[str]:
    from . import logic as logic_utils

    return logic_utils._collect_unresolved_macro_candidates(
        body,
        known_names=known_names,
        backend_module=sys.modules[__name__],
    )


def _collect_member_access_candidates(body: str, *, known_members: Optional[set[str]] = None) -> list[tuple[str, str]]:
    from . import logic as logic_utils

    return logic_utils._collect_member_access_candidates(
        body,
        known_members=known_members,
        backend_module=sys.modules[__name__],
    )


def _replace_member_chain_with_owner(text: str, name_map: Optional[dict[str, str]]) -> str:
    from . import logic as logic_utils

    return logic_utils._replace_member_chain_with_owner(
        text,
        name_map,
        backend_module=sys.modules[__name__],
    )


def _build_canonical_file_symbol_map(
    file_context: Optional[dict],
    body: str,
    local_vars: Sequence[dict],
    params: Sequence[dict],
    cfg: Optional[GenConfig],
) -> tuple[dict[str, str], dict[str, SymbolInference]]:
    from . import logic as logic_utils

    return logic_utils._build_canonical_file_symbol_map(
        file_context,
        body,
        local_vars,
        params,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _simplify_member_access(text: str) -> str:
    if not text:
        return text
    cur = text
    for _ in range(20):
        new_cur = _MEMBER_PREFIX_RE.sub("", cur)
        if new_cur == cur:
            break
        cur = new_cur
    return cur


def _cleanup_generated_logic_text(text: str) -> str:
    from . import logic as logic_utils

    return logic_utils._cleanup_generated_logic_text(
        text,
        backend_module=sys.modules[__name__],
    )


def _normalize_explanatory_text_for_output(text: str, name_map: Optional[dict[str, str]] = None) -> str:
    from . import logic as logic_utils

    return logic_utils._normalize_explanatory_text_for_output(
        text,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def _normalize_logic_line_for_output(text: str, name_map: Optional[dict[str, str]] = None) -> str:
    from . import logic as logic_utils

    return logic_utils._normalize_logic_line_for_output(
        text,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def fallback_logic_line(code_line: str, name_map: Optional[dict[str, str]] = None) -> str:
    from . import logic as logic_utils

    return logic_utils.fallback_logic_line(
        code_line,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def _split_c_call_args(arg_text: str) -> list[str]:
    """
    轻量拆分 C 函数调用参数（按顶层逗号切分），尽量兼容括号/中括号/花括号与字符串。
    """
    s = (arg_text or "").strip()
    if not s:
        return []

    args: list[str] = []
    cur: list[str] = []
    depth_paren = 0
    depth_brack = 0
    depth_brace = 0
    in_squote = False
    in_dquote = False
    escape = False

    for ch in s:
        if escape:
            cur.append(ch)
            escape = False
            continue
        if ch == "\\":
            cur.append(ch)
            escape = True
            continue

        if in_squote:
            cur.append(ch)
            if ch == "'":
                in_squote = False
            continue
        if in_dquote:
            cur.append(ch)
            if ch == '"':
                in_dquote = False
            continue

        if ch == "'":
            in_squote = True
            cur.append(ch)
            continue
        if ch == '"':
            in_dquote = True
            cur.append(ch)
            continue

        if ch == "(":
            depth_paren += 1
            cur.append(ch)
            continue
        if ch == ")":
            depth_paren = max(0, depth_paren - 1)
            cur.append(ch)
            continue
        if ch == "[":
            depth_brack += 1
            cur.append(ch)
            continue
        if ch == "]":
            depth_brack = max(0, depth_brack - 1)
            cur.append(ch)
            continue
        if ch == "{":
            depth_brace += 1
            cur.append(ch)
            continue
        if ch == "}":
            depth_brace = max(0, depth_brace - 1)
            cur.append(ch)
            continue

        if ch == "," and depth_paren == 0 and depth_brack == 0 and depth_brace == 0:
            part = "".join(cur).strip()
            if part:
                args.append(part)
            cur = []
            continue

        cur.append(ch)

    last = "".join(cur).strip()
    if last:
        args.append(last)
    return args


def _logic_cn_expr(expr: str, name_map: Optional[dict[str, str]] = None) -> str:
    from . import logic as logic_utils

    return logic_utils._logic_cn_expr(
        expr,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def _format_call_expr(expr: str, name_map: Optional[dict[str, str]] = None) -> Optional[tuple[str, str]]:
    from . import logic as logic_utils

    return logic_utils._format_call_expr(
        expr,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def heuristic_logic_line(code_line: str, name_map: Optional[dict[str, str]] = None, *, literal: bool = False) -> Optional[str]:
    from . import logic as logic_utils

    return logic_utils.heuristic_logic_line(
        code_line,
        name_map=name_map,
        literal=literal,
        backend_module=sys.modules[__name__],
    )


def _build_logic_ir_node(
    code_line: str,
    *,
    attached: Optional[Sequence[str]] = None,
    name_map: Optional[dict[str, str]] = None,
    cfg: Optional[GenConfig] = None,
    use_cond_comment: bool = False,
) -> Optional[dict[str, Any]]:
    from . import logic as logic_utils
    return logic_utils._build_logic_ir_node(
        code_line,
        attached=attached,
        name_map=name_map,
        cfg=cfg,
        use_cond_comment=use_cond_comment,
        backend_module=sys.modules[__name__],
    )


def _render_logic_ir_node(
    node: dict[str, Any],
    *,
    name_map: Optional[dict[str, str]] = None,
    local_var_usages: Optional[dict[str, str]] = None,
    literal: bool = False,
) -> str:
    from . import logic as logic_utils
    return logic_utils._render_logic_ir_node(
        node,
        name_map=name_map,
        local_var_usages=local_var_usages,
        literal=literal,
        backend_module=sys.modules[__name__],
    )


def ai_refine_logic_unknowns(unknown_list, code_context: str, cfg: GenConfig):
    from . import ai as ai_utils
    return ai_utils.ai_refine_logic_unknowns(unknown_list, code_context, cfg)



def call_llm_text(prompt: str, cfg: GenConfig, **kwargs) -> str:
    """
    调用 LLM 并返回纯文本，适用于不需要 JSON 结构化输出的场景。
    会自动清理模型输出中的代码块、解释性文字等，保证返回更干净的文本。
    """

    from . import ai as ai_utils

    return ai_utils.call_llm_text(
        prompt,
        cfg,
        _runtime_module=sys.modules.get(__name__),
        **kwargs,
    )


def enrich_with_ai(func_data: dict, cfg: GenConfig):
    """补全截断的 AI 逻辑。"""
    from . import ai as ai_utils

    return ai_utils.enrich_with_ai(
        func_data,
        cfg,
        _runtime_module=sys.modules.get(__name__),
    )


def _symbol_memory_kind_for_name(name: str) -> str:
    from . import ai as ai_utils

    return ai_utils._symbol_memory_kind_for_name(
        name,
        backend_module=sys.modules[__name__],
    )


def _collect_unresolved_body_symbol_candidates(
    body: str,
    *,
    local_names: Optional[set[str]] = None,
    param_names: Optional[set[str]] = None,
    known_names: Optional[set[str]] = None,
) -> list[str]:
    from . import ai as ai_utils

    return ai_utils._collect_unresolved_body_symbol_candidates(
        body,
        local_names=local_names,
        param_names=param_names,
        known_names=known_names,
        backend_module=sys.modules[__name__],
    )


def _collect_symbol_memory_warmup_candidates(func_entries: Sequence[dict], cfg: Optional[GenConfig] = None) -> list[dict]:
    from . import ai as ai_utils

    return ai_utils._collect_symbol_memory_warmup_candidates(
        func_entries,
        cfg,
        backend_module=sys.modules[__name__],
    )


def _warmup_symbol_memory_once(func_entries: Sequence[dict], cfg: GenConfig, *, scope_label: str) -> None:
    from . import ai as ai_utils

    return ai_utils._warmup_symbol_memory_once(
        func_entries,
        cfg,
        scope_label=scope_label,
        backend_module=sys.modules[__name__],
    )


def _flatten_preprocessed_func_entries(preprocessed: dict[str, dict]) -> list[dict]:
    from . import ai as ai_utils

    return ai_utils._flatten_preprocessed_func_entries(preprocessed)


# =============== 章节生成 ===============

def _clean_description_lines(desc: str) -> tuple[str, ...]:
    from . import logic as logic_utils

    return logic_utils._clean_description_lines(desc)


def _is_control_logic_line(text: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._is_control_logic_line(text)


def _refresh_control_logic_line_idents(text: str, name_map: Optional[dict[str, str]]) -> str:
    from . import logic as logic_utils

    return logic_utils._refresh_control_logic_line_idents(
        text,
        name_map,
        backend_module=sys.modules[__name__],
    )


def _normalize_local_usage(cfg: GenConfig, usage_text: str, var_name: str, cn_name: str) -> str:
    from . import ai as ai_utils

    return ai_utils._normalize_local_usage(
        cfg,
        usage_text,
        var_name,
        cn_name,
        backend_module=sys.modules[__name__],
    )


def _needs_ai_local_usage_refine(item: dict) -> bool:
    from . import ai as ai_utils

    return ai_utils._needs_ai_local_usage_refine(
        item,
        backend_module=sys.modules[__name__],
    )


def _should_replace_local_usage_with_ai(old_usage: str, cn_name: str = "") -> bool:
    from . import ai as ai_utils

    return ai_utils._should_replace_local_usage_with_ai(
        old_usage,
        cn_name,
        backend_module=sys.modules[__name__],
    )


def _looks_like_generic_local_usage(text: str, cn_name: str = "") -> bool:
    from . import ai as ai_utils

    return ai_utils._looks_like_generic_local_usage(
        text,
        cn_name,
        backend_module=sys.modules[__name__],
    )


def _looks_like_too_generic_usage_text(text: str) -> bool:
    from . import ai as ai_utils

    return ai_utils._looks_like_too_generic_usage_text(
        text,
        backend_module=sys.modules[__name__],
    )


def _strip_previous_prefix(text: str) -> str:
    from . import logic as logic_utils

    return logic_utils._strip_previous_prefix(
        text,
        backend_module=sys.modules[__name__],
    )


def _select_local_usage_text(item: dict, *, body: str, comment_desc: str = "", cfg: Optional[GenConfig] = None) -> str:
    from . import logic as logic_utils

    return logic_utils._select_local_usage_text(
        item,
        body=body,
        comment_desc=comment_desc,
        cfg=cfg,
        backend_module=sys.modules[__name__],
    )


def _build_local_param_symbol_map(
    local_vars: Sequence[dict],
    params: Sequence[dict],
    in_map: dict[str, str],
    out_map: dict[str, str],
    param_ai_name_map: dict[str, str],
) -> dict[str, str]:
    from . import logic as logic_utils

    return logic_utils._build_local_param_symbol_map(
        local_vars,
        params,
        in_map,
        out_map,
        param_ai_name_map,
        backend_module=sys.modules[__name__],
    )


def build_function_fact_pack(func_data: dict[str, Any], cfg: Optional["GenConfig"] = None):
    from . import lsp_facts as fact_utils

    return fact_utils.build_function_fact_pack(func_data, cfg, backend_module=sys.modules[__name__])


def register_external_snapshot(source_file: str, text: str, version: int, project_root: str = "") -> None:
    from . import lsp_gateway as gateway_utils

    gateway_utils.register_external_snapshot(source_file, text, version, project_root=project_root or None)


def shutdown_lsp_idle_sessions(idle_seconds: int = 120) -> None:
    from . import lsp_gateway as gateway_utils

    gateway_utils.get_lsp_gateway(backend_module=sys.modules[__name__]).shutdown_idle_sessions(idle_seconds=idle_seconds)


def build_logic_semantic_pack(ctx: dict[str, Any]):
    from . import pipeline as pipeline_utils

    return pipeline_utils.build_logic_semantic_pack(ctx, backend_module=sys.modules[__name__])


def _count_logic_placeholder_lines(logic: Any) -> int:
    from . import logic as logic_utils

    return logic_utils._count_logic_placeholder_lines(
        logic,
        backend_module=sys.modules[__name__],
    )


_GENERIC_LOGIC_PHRASES = (
    "执行操作",
    "处理数据",
    "处理信息",
    "进行处理",
    "更新变量",
    "设置变量",
    "完成相关处理",
    "完成相关计算",
    "完成相关更新",
)


def _is_resolved_symbol_text(symbol_name: str, text: str) -> bool:
    from . import logic as logic_utils

    return logic_utils._is_resolved_symbol_text(
        symbol_name,
        text,
        backend_module=sys.modules[__name__],
    )


def repair_unresolved_logic_lines(
    logic_lines: Sequence[str],
    symbol_map: dict[str, str],
) -> list[str]:
    from . import logic as logic_utils

    return logic_utils.repair_unresolved_logic_lines(logic_lines, symbol_map)


def _replace_generic_call_with_verb(
    line: str, func_name: str, name_map: Optional[dict[str, str]] = None
) -> str:
    from . import logic as logic_utils

    return logic_utils._replace_generic_call_with_verb(
        line,
        func_name,
        name_map=name_map,
    )


def repair_generic_logic_calls(
    logic_lines: Sequence[str],
    body: str = "",
    name_map: Optional[dict[str, str]] = None,
) -> list[str]:
    from . import logic as logic_utils

    return logic_utils.repair_generic_logic_calls(
        logic_lines,
        body=body,
        name_map=name_map,
        backend_module=sys.modules[__name__],
    )


def expand_thin_logic(
    logic_lines: Sequence[str],
    body: str,
    name_map: Optional[dict[str, str]] = None,
    cfg: Optional["GenConfig"] = None,
) -> list[str]:
    from . import logic as logic_utils

    return logic_utils.expand_thin_logic(
        logic_lines,
        body,
        name_map=name_map,
        cfg=cfg,
    )


def _collect_function_quality_report(
    local_vars: Sequence[dict],
    params: Sequence[dict],
    in_map: dict[str, str],
    out_map: dict[str, str],
    param_ai_name_map: dict[str, str],
    logic_lines: Optional[Sequence[str]],
    name_map: Optional[dict[str, str]],
    inferences: Sequence[SymbolInference] = (),
) -> dict[str, Any]:
    from . import ai as ai_utils

    return ai_utils._collect_function_quality_report(
        local_vars,
        params,
        in_map,
        out_map,
        param_ai_name_map,
        logic_lines,
        name_map,
        inferences,
        backend_module=sys.modules[__name__],
    )


def build_function_design(
    func_data: dict,
    module_req_prefix: str,
    index: int,
    cfg: GenConfig,
    *,
    cached_logic_lines: Optional[dict] = None,
    changed_statement_lines: Optional[set] = None,
) -> FunctionDesign:
    try:
        from .pipeline import build_function_design_impl
    except Exception:
        from .pipeline import build_function_design_impl
    return build_function_design_impl(
        func_data,
        module_req_prefix,
        index,
        cfg,
        backend_module=sys.modules[__name__],
        cached_logic_lines=cached_logic_lines,
        changed_statement_lines=changed_statement_lines,
    )

def render_function_design(doc: Document, design: FunctionDesign, cfg: GenConfig):
    from . import render as render_utils

    return render_utils.render_function_design(doc, design, cfg)



def add_function_section(doc: Document, func_data: dict,
                         module_req_prefix: str,
                         index: int,
                         cfg: GenConfig):
    design = build_function_design(func_data, module_req_prefix=module_req_prefix, index=index, cfg=cfg)
    render_function_design(doc, design, cfg)
    return design


#=============== 单文件/工程生成 ===============

def generate_design_doc_from_file(source: str, output: str, cfg: GenConfig, resume_state: Optional[dict] = None):
    from . import pipeline as pipeline_utils
    from . import render as render_utils
    from . import runtime as runtime_utils
    return pipeline_utils.run_single_file_generation(
        source,
        output,
        cfg,
        resume_state=resume_state,
        backend_module=sys.modules[__name__],
        render_module=render_utils,
        runtime_module=runtime_utils,
    )


def generate_design_doc_for_single_function(
    source: str,
    func_name: str,
    output: str,
    cfg: GenConfig,
    *,
    project_root: Optional[str] = None,
) -> str:
    from . import pipeline as pipeline_utils
    from . import render as render_utils
    from . import runtime as runtime_utils
    return pipeline_utils.run_single_export_generation(
        source,
        func_name,
        output,
        cfg,
        project_root=project_root,
        backend_module=sys.modules[__name__],
        render_module=render_utils,
        runtime_module=runtime_utils,
    )



def generate_design_doc_for_project(root_dir: str, output: str, cfg: GenConfig, resume_state: Optional[dict] = None):
    from . import pipeline as pipeline_utils
    from . import runtime as runtime_utils
    return pipeline_utils.run_project_generation(
        root_dir,
        output,
        cfg,
        resume_state=resume_state,
        incremental=bool(getattr(cfg, "incremental", False)),
        backend_module=sys.modules[__name__],
        runtime_module=runtime_utils,
    )


# ================= 命令行主程序 =================


_legacy_build_function_design_impl = build_function_design
_legacy_render_function_design_impl = render_function_design
_legacy_add_function_section_impl = add_function_section
_legacy_generate_design_doc_from_file_impl = generate_design_doc_from_file
_legacy_generate_design_doc_for_single_function_impl = generate_design_doc_for_single_function
_legacy_generate_design_doc_for_project_impl = generate_design_doc_for_project




if __name__ == "__main__":
    cli_utils.main()
