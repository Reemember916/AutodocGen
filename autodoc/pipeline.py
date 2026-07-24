"""Document-generation pipeline wrappers and model conversion."""

from __future__ import annotations

import concurrent.futures
import copy
import json
import os
import re
from dataclasses import fields, replace
from typing import Any, Optional, Sequence

from ._legacy_support import legacy_backend
from . import logic as logic_utils
from . import scanner as scanner_utils
from . import semantic_registry
from . import utils as utils_module
from . import text as text_utils
from . import lsp_facts as lsp_fact_utils
from . import parse as parse_utils
from . import naming as naming_utils
from . import revision as revision_utils
from . import quality_gate
from . import effects as effects_utils
from .models import AIBuildMeta, DesignModel, FunctionBuildResult, FunctionDesign, IOElement, LocalDataElement, ProjectResumeState

# P0#3 Evidence 旁路采集（shadow mode，可选开关：extra_params["evidence_output"]）
# Switches (extra_params):
#   evidence_output = off|0|false  → disabled (default when empty)
#                     on|1|true|yes → write <stem>_evidence/evidence_report.json
#                     /path/to/file.json or dir → custom report path
#   logic_step_ir   = shadow|on|1 (default when evidence on) | off|0
_EVIDENCE_ENABLED = False
try:
    from .evidence import record_function_evidence, clear_recorded_evidence, write_evidence_report
    from .logic_step_ir import clear_untranslated_idents, auto_suggest_symbol_translations
    _EVIDENCE_ENABLED = True
except Exception:
    _EVIDENCE_ENABLED = False
    clear_untranslated_idents = lambda: None  # type: ignore
    auto_suggest_symbol_translations = lambda *a, **kw: {}  # type: ignore


def _truthy_flag(value: Any) -> bool:
    return str(value or "").strip().lower() in ("1", "true", "on", "yes", "shadow")


def _falsey_flag(value: Any) -> bool:
    return str(value or "").strip().lower() in ("0", "false", "off", "no", "none", "")


def evidence_output_enabled(cfg) -> bool:
    """Whether evidence shadow collection/report is enabled for this run."""
    extra = getattr(cfg, "extra_params", None) or {}
    if not isinstance(extra, dict):
        return False
    raw = str(extra.get("evidence_output", "") or "").strip()
    if not raw or _falsey_flag(raw):
        return False
    if _truthy_flag(raw):
        return True
    # Custom path counts as enabled
    return True


def logic_step_ir_enabled(cfg) -> bool:
    """logic_step_ir=shadow|on|primary by default when evidence is on; explicit off disables."""
    extra = getattr(cfg, "extra_params", None) or {}
    if not isinstance(extra, dict):
        return False
    raw = str(extra.get("logic_step_ir", "") or "").strip()
    if raw:
        return _truthy_flag(raw) or str(raw).strip().lower() in (
            "primary", "main", "replace", "render",
        )
    return evidence_output_enabled(cfg)


def logic_step_ir_primary(cfg) -> bool:
    """When true, LogicStep IR replaces generate_logic_from_body text for docx."""
    extra = getattr(cfg, "extra_params", None) or {}
    if not isinstance(extra, dict):
        return False
    raw = str(extra.get("logic_step_ir", "") or "").strip().lower()
    return raw in ("primary", "main", "replace", "render")


def resolve_evidence_report_path(output: str, cfg) -> str:
    """
    Resolve evidence JSON path.

    - on/1/true → ``<output_stem>_evidence/evidence_report.json``
    - absolute/relative path ending with .json → that file
    - directory path → ``<dir>/evidence_report.json``
    """
    extra = getattr(cfg, "extra_params", None) or {}
    raw = str((extra or {}).get("evidence_output", "") or "").strip()
    if not raw or _falsey_flag(raw):
        return ""
    if _truthy_flag(raw):
        stem = os.path.splitext(os.path.abspath(output or "out.docx"))[0]
        return stem + "_evidence" + os.sep + "evidence_report.json"
    path = os.path.abspath(os.path.expanduser(raw))
    if path.lower().endswith(".json"):
        return path
    return os.path.join(path, "evidence_report.json")


def maybe_write_evidence_report(output: str, cfg, *, backend_module=None) -> str:
    """Write evidence report if enabled; return path or empty string."""
    if not _EVIDENCE_ENABLED or not evidence_output_enabled(cfg):
        return ""
    ev_path = resolve_evidence_report_path(output, cfg)
    if not ev_path:
        return ""
    try:
        write_evidence_report(ev_path)
        backend = backend_module or legacy_backend()
        backend.vlog(cfg, f"[Evidence] 旁路报告已生成：{ev_path}")
        return ev_path
    except Exception as exc:
        try:
            backend = backend_module or legacy_backend()
            backend.vlog(cfg, f"[Evidence] 报告写入失败：{exc}")
        except Exception:
            pass
        return ""


def _evidence_shadow_collect(func_data: dict, ctx: dict, cfg) -> None:
    """旁路采集函数 evidence（shadow mode）。

    在 build_function_design_impl 的 prepare_design_context 后调用。
    不影响 docx 输出；clang 不可用时 fact_pack 为 None，自动降级。
    """
    if not _EVIDENCE_ENABLED:
        return
    if not evidence_output_enabled(cfg):
        return
    if not logic_step_ir_enabled(cfg):
        return
    body = ctx.get("body") or (func_data.get("body") or "")
    if not body:
        return
    try:
        from .logic_step_ir import build_logic_steps
        logic_steps = build_logic_steps(body, ctx.get("local_vars"), cfg, name_map=ctx.get("global_symbol_map"))
    except Exception:
        return
    # 构造 evidence 用的 func_data（补充 ctx 中已有的字段）
    ev_func_data = dict(func_data)
    ev_func_data.setdefault("local_vars", ctx.get("local_vars") or [])
    ev_func_data.setdefault("params", ctx.get("params") or [])
    ev_func_data.setdefault("file_context", ctx.get("file_context") or {})
    name_map = ctx.get("global_symbol_map") or {}
    fact_pack = ctx.get("lsp_fact_pack")
    record_function_evidence(
        ev_func_data, logic_steps, name_map,
        lsp_fact_pack=fact_pack,
    )


def clone_cfg(cfg, **overrides):
    cfg_type = type(cfg)
    try:
        cfg_fields = fields(cfg_type)
    except TypeError:
        legacy = legacy_backend()
        cfg_type = legacy.GenConfig
        cfg_fields = fields(cfg_type)
    data = {item.name: getattr(cfg, item.name) for item in cfg_fields}
    data.update(overrides)
    new_cfg = cfg_type(**data)
    # 传播动态取消标记
    if getattr(cfg, "_user_cancelled", False):
        try:
            new_cfg._user_cancelled = True
        except Exception:
            pass
    return new_cfg


def _review_collection(cfg) -> list:
    items = getattr(cfg, "_review_workspace_functions", None)
    if not isinstance(items, list):
        items = []
        try:
            cfg._review_workspace_functions = items
        except Exception:
            pass
    return items


def _reset_review_collection(cfg) -> None:
    try:
        cfg._review_workspace_functions = []
    except Exception:
        pass

def _design_workspace_pairs(cfg) -> list:
    items = getattr(cfg, "_design_workspace_pairs", None)
    if not isinstance(items, list):
        items = []
        try:
            cfg._design_workspace_pairs = items
        except Exception:
            pass
    return items


def _reset_design_workspace_pairs(cfg) -> None:
    try:
        cfg._design_workspace_pairs = []
    except Exception:
        pass


def _collect_design_workspace_pair(cfg, design, task) -> None:
    try:
        if design is None or task is None:
            return
        pairs = _design_workspace_pairs(cfg)
        pairs.append((design, task))
    except Exception:
        pass


def _write_design_workspace_if_enabled(cfg, output: str, *, project_root: str = "", merge_existing: bool = False) -> None:
    try:
        from . import design_workspace
        extra = getattr(cfg, "extra_params", {}) or {}
        enabled = str(extra.get("design_workspace") or getattr(cfg, "design_workspace", "") or "off").strip().lower()
        if enabled not in ("1", "on", "yes", "true"):
            return
        pairs = tuple(getattr(cfg, "_design_workspace_pairs", []) or ())
        if not pairs:
            return
        bundle = design_workspace.build_workspace_bundle(
            pairs,
            project_root=str(project_root or ""),
            output_docx=str(output or ""),
        )
        explicit_dir = str(extra.get("design_workspace_dir") or "").strip()
        if explicit_dir:
            out_path = os.path.join(explicit_dir, "design_workspace.json")
        else:
            out_path = str(output or "output") + ".design_workspace.json"
        design_workspace.write_workspace(bundle, out_path, merge_existing=merge_existing)
        legacy_backend().vlog(cfg, f"Design workspace 已生成：{out_path}")
    except Exception as exc:
        legacy_backend().vlog(cfg, f"Design workspace 生成失败: {exc}")


def _collect_review_function(cfg, design, task) -> None:
    try:
        from . import review_workspace

        if not review_workspace.review_output_enabled(cfg):
            return
        func_data = dict((task or {}).get("func_data") or {})
        func_data.setdefault("source_file", (task or {}).get("source_file", ""))
        items = _review_collection(cfg)
        review_fn = review_workspace.build_review_function(design, func_data, cfg)
        items.append(review_workspace.disambiguate_review_function(review_fn, items))
    except Exception as exc:
        backend = legacy_backend()
        backend.vlog(cfg, f"Review workspace collection skipped: {exc}")


def _review_project_root(project_root: str = "", *, source: str = "") -> str:
    """Normalize review workspace project_root for stable function keys."""
    candidate = str(project_root or source or "").strip()
    if not candidate:
        return ""
    try:
        abs_path = os.path.abspath(os.path.expanduser(candidate))
    except Exception:
        return candidate
    if os.path.isfile(abs_path) or abs_path.lower().endswith((".c", ".h", ".cpp", ".cc", ".cxx")):
        return os.path.dirname(abs_path) or abs_path
    return abs_path


def _registered_function_title(func_data: dict[str, Any], *, backend_module=None) -> str:
    """Return a project-registry title before legacy per-file name maps."""
    backend = backend_module or legacy_backend()
    file_context = dict((func_data or {}).get("file_context") or {})
    registered = utils_module._safe_strip(file_context.get("function_title"))
    if registered:
        return registered
    func_info = dict((func_data or {}).get("func_info") or {})
    comment_info = dict((func_data or {}).get("comment_info") or {})
    func_name = utils_module._safe_strip(func_info.get("func_name"))
    mapped = utils_module._safe_strip(dict(file_context.get("func_cn_map") or {}).get(func_name))
    return mapped or backend.get_function_chinese_name(comment_info, func_info)


def apply_project_function_title_registry(
    func_entries: Sequence[dict[str, Any]],
    project_root: str,
    cfg,
    *,
    backend_module=None,
) -> tuple[dict[str, Any], ...]:
    """Give colliding Chinese function titles stable, project-wide qualifiers.

    Function identity is always ``relative_source_file::c_function_name``.
    The first deterministic identity keeps the concise title; later collisions
    receive a source-stem qualifier, with the C name as a last-resort suffix.
    """
    backend = backend_module or legacy_backend()
    entries = [item for item in (func_entries or ()) if isinstance(item, dict)]
    candidates: list[tuple[str, str, dict[str, Any], str]] = []
    for item in entries:
        func_info = dict(item.get("func_info") or {})
        comment_info = dict(item.get("comment_info") or {})
        file_context = dict(item.get("file_context") or {})
        func_name = utils_module._safe_strip(func_info.get("func_name"))
        source_file = utils_module._safe_strip(file_context.get("source_file"))
        if not func_name:
            continue
        try:
            rel_source = os.path.relpath(source_file, project_root).replace(os.sep, "/") if source_file else "unknown"
        except Exception:
            rel_source = source_file or "unknown"
        identity = f"{rel_source}::{func_name}"
        title = backend.get_function_chinese_name(comment_info, func_info)
        title = backend._normalize_function_cn_title(
            title, func_name=func_name, comment_desc=utils_module._safe_strip(comment_info.get("desc")),
        )
        candidates.append((utils_module._safe_strip(title) or func_name, identity, item, func_name))

    by_title: dict[str, list[tuple[str, str, dict[str, Any], str]]] = {}
    for record in candidates:
        by_title.setdefault(record[0], []).append(record)
    assigned: dict[str, str] = {}
    collisions: list[dict[str, Any]] = []
    used: set[str] = set()
    for base_title, records in sorted(by_title.items()):
        for pos, (_title, identity, item, func_name) in enumerate(sorted(records, key=lambda value: value[1])):
            final = base_title
            if pos > 0 or final in used:
                source_file = utils_module._safe_strip((item.get("file_context") or {}).get("source_file"))
                qualifier = os.path.splitext(os.path.basename(source_file))[0] or "source"
                final = f"{base_title}（{qualifier}）"
                if final in used:
                    final = f"{base_title}（{qualifier}_{func_name}）"
                suffix = 2
                while final in used:
                    final = f"{base_title}（{qualifier}_{suffix}）"
                    suffix += 1
                collisions.append({"title": base_title, "identity": identity, "resolved_title": final})
            used.add(final)
            assigned[identity] = final
            # Preserve the original nested objects when possible.  Project
            # preprocessing later creates shallow copies of function entries,
            # so replacing these dicts would leave the actual render inputs
            # without the registered title.
            comment_info = item.get("comment_info")
            if not isinstance(comment_info, dict):
                comment_info = {}
                item["comment_info"] = comment_info
            comment_info["func_cn_name"] = final
            file_context = item.get("file_context")
            if not isinstance(file_context, dict):
                file_context = {}
                item["file_context"] = file_context
            file_context["function_title"] = final
            file_context["function_title_key"] = identity
    try:
        cfg.function_title_registry = dict(assigned)
        cfg.function_title_collisions = tuple(collisions)
    except Exception:
        pass
    if collisions:
        backend.vlog(cfg, f"[title_registry] 已消歧 {len(collisions)} 个函数中文标题冲突")
    return tuple(collisions)


def apply_preprocessed_function_title_registry(
    preprocessed: dict[str, dict],
    project_root: str,
    cfg,
    *,
    backend_module=None,
) -> tuple[dict[str, Any], ...]:
    """Register titles on the original project entries used for rendering."""
    entries: list[dict[str, Any]] = []
    for c_path, pre in (preprocessed or {}).items():
        for fd in (pre or {}).get("func_list") or ():
            if not isinstance(fd, dict):
                continue
            file_context = fd.get("file_context")
            if not isinstance(file_context, dict):
                file_context = {}
                fd["file_context"] = file_context
            file_context["source_file"] = c_path
            entries.append(fd)
    return apply_project_function_title_registry(
        entries,
        project_root,
        cfg,
        backend_module=backend_module,
    )


def _write_review_workspace_if_enabled(cfg, output: str, *, project_root: str = "", merge_existing: bool = False) -> None:
    from . import review_workspace

    if not review_workspace.review_output_enabled(cfg):
        return
    functions = tuple(getattr(cfg, "_review_workspace_functions", []) or ())
    bundle = review_workspace.ReviewBundle(
        project_root=_review_project_root(project_root),
        output_docx=str(output or ""),
        functions=functions,
    )
    out_dir = review_workspace.review_output_dir(cfg, output)
    review_workspace.write_review_workspace(bundle, out_dir, merge_existing=merge_existing)
    legacy_backend().vlog(cfg, f"Review workspace 已生成：{out_dir}")


def _future_result_until_stop(future, cfg, backend, *, poll_interval: float = 0.2):
    while True:
        if backend.stop_requested(cfg):
            raise TimeoutError("stopped")
        try:
            return future.result(timeout=poll_interval)
        except concurrent.futures.TimeoutError:
            continue


def _prepare_func_list_for_c_file(
    c_path: str,
    *,
    project_root: Optional[str],
    cfg,
    prefilter: bool,
    need_symbol_map: bool = True,
    backend_module=None,
):
    backend = backend_module or legacy_backend()
    # Prefer parse module as default implementation source.
    # When an explicit backend is injected, always honor its hook for compatibility/tests.
    if backend_module is not None and hasattr(backend, "prepare_func_list_for_c_file"):
        return backend.prepare_func_list_for_c_file(
            c_path,
            project_root=project_root,
            cfg=cfg,
            prefilter=prefilter,
            need_symbol_map=need_symbol_map,
        )
    return parse_utils.prepare_func_list_for_c_file(
        c_path,
        project_root=project_root,
        cfg=cfg,
        prefilter=prefilter,
        need_symbol_map=need_symbol_map,
    )


def _prebuild_project_symbols_into_runtime(project_root: str, *, cfg, backend_module=None) -> dict[str, dict[str, str]]:
    backend = backend_module or legacy_backend()
    if backend.stop_requested(cfg):
        return {"macros": {}, "members": {}, "symbols": {}}
    if backend_module is not None and hasattr(backend, "prebuild_project_symbol_db"):
        prebuilt = backend.prebuild_project_symbol_db(project_root, cfg=cfg)
        untranslated = prebuilt.pop("untranslated_macros", [])
        # 宏翻译推迟到函数级按需进行（默认）；仅当显式开启 prebuild_translate_macros 时才在预构建阶段批量翻译
        if utils_module.cfg_get_int(cfg, "prebuild_translate_macros", 0) and cfg.ai_assist and untranslated and (not backend.stop_requested(cfg)):
            translated = backend.batch_translate_symbols(untranslated, kind="macros", cfg=cfg)
            for key, value in translated.items():
                prebuilt.setdefault("macros", {})[key] = value
        else:
            # 未翻译宏名单存入 runtime，供函数级按需翻译参考（不调 AI）
            prebuilt.setdefault("untranslated_macros", untranslated)
        backend.merge_prebuilt_symbols_into_runtime(prebuilt)
        return prebuilt
    prebuilt = parse_utils.prebuild_project_symbol_db(project_root, cfg=cfg)
    untranslated = prebuilt.pop("untranslated_macros", [])
    # 宏翻译推迟到函数级按需进行（默认）；仅当显式开启 prebuild_translate_macros 时才在预构建阶段批量翻译
    if utils_module.cfg_get_int(cfg, "prebuild_translate_macros", 0) and cfg.ai_assist and untranslated and (not backend.stop_requested(cfg)):
        translated = parse_utils.batch_translate_symbols(untranslated, kind="macros", cfg=cfg)
        for key, value in translated.items():
            prebuilt.setdefault("macros", {})[key] = value
    else:
        prebuilt.setdefault("untranslated_macros", untranslated)
    parse_utils.merge_prebuilt_symbols_into_runtime(prebuilt)
    return prebuilt


def should_parallel_build_design(cfg) -> bool:
    if not cfg:
        return False
    return bool(
        getattr(cfg, "ai_assist", False)
        and int(getattr(cfg, "ai_workers", 1) or 1) > 1
    )


def build_function_design_task(
    func_data: dict,
    module_req_prefix: str,
    index: int,
    cfg,
    *,
    backend_module=None,
):
    backend = backend_module or legacy_backend()
    worker_cfg = clone_cfg(cfg)
    try:
        for attr in (
            "_ai_cache_salt",
            "_ai_regression_round",
            "_ai_quality_feedback",
            "_ai_quality_focus_symbols",
            "_ai_regression_allow_rename",
            "_current_file",
            "_current_func_index",
            "_current_func_pos",
        ):
            if hasattr(cfg, attr):
                setattr(worker_cfg, attr, getattr(cfg, attr))
        worker_cfg._in_func_context = True
        worker_cfg._current_func_ai_failed = False
        worker_cfg._skip_ai_current_func = False
    except Exception:
        pass
    # 提取语句级增量缓存
    cached_logic_lines = func_data.pop("_cached_logic_lines", None)
    changed_statement_lines = func_data.pop("_changed_statement_lines", None)
    return backend.build_function_design(
        func_data,
        module_req_prefix=module_req_prefix,
        index=index,
        cfg=worker_cfg,
        cached_logic_lines=cached_logic_lines,
        changed_statement_lines=changed_statement_lines,
    )


def cfg_with_function_task_context(cfg, task: Optional[dict]):
    task_cfg = clone_cfg(cfg)
    try:
        for attr in (
            "_ai_cache_salt", "_ai_regression_round", "_ai_quality_feedback",
            "_ai_quality_focus_symbols", "_ai_regression_allow_rename",
        ):
            if hasattr(cfg, attr):
                setattr(task_cfg, attr, getattr(cfg, attr))
        if isinstance(task, dict):
            task_cfg._current_file = str(task.get("file") or "")
            task_cfg._current_func_index = int(task.get("index") or 0)
            task_cfg._current_func_pos = int(task.get("func_pos") or 0)
    except Exception:
        pass
    return task_cfg


def design_regression_rounds(cfg, *, backend_module=None) -> int:
    backend = backend_module or legacy_backend()
    if not cfg or (not getattr(cfg, "ai_assist", False)):
        return 0
    default_rounds = int(getattr(cfg, "ai_regression_rounds", 2) or 2)
    return max(0, int(utils_module.cfg_get_int(cfg, "ai_regression_rounds", default_rounds)))


def design_regression_force_one_call(cfg, *, backend_module=None) -> bool:
    backend = backend_module or legacy_backend()
    return bool(utils_module.cfg_get_int(cfg, "ai_regression_force_one_call", 1))


def score_ai_meta(meta: Optional[Any], *, backend_module=None) -> int:
    backend = backend_module or legacy_backend()
    if not isinstance(meta, AIBuildMeta):
        return 0
    score = len(meta.regression_reasons)
    if meta.ai_failed:
        score += 3
    score += max(0, int(meta.logic_placeholders or 0))
    score += len(meta.unresolved_local_symbols or ())
    score += len(meta.unresolved_param_symbols or ())
    score += len(meta.unresolved_logic_symbols or ())
    score += max(0, int(meta.generic_logic_count or 0))
    # Structural defects dominate every soft naming/wording score.  A retry
    # that introduces one must never be selected merely because it reduces
    # over-translation or unresolved-symbol counts.
    score += 1000 * sum(
        1 for item in (meta.quality_issues or ())
        if str((item or {}).get("code") or "") in quality_gate.STRUCTURAL_LOGIC_CODES
        and str((item or {}).get("severity") or "").lower() == "error"
    )
    return score


def prefer_regression_design(current: Any, retry: Any, *, backend_module=None) -> bool:
    backend = backend_module or legacy_backend()
    if not isinstance(current, FunctionDesign):
        return True
    if not isinstance(retry, FunctionDesign):
        return False
    cur_meta = current.ai_meta if isinstance(current.ai_meta, AIBuildMeta) else AIBuildMeta()
    new_meta = retry.ai_meta if isinstance(retry.ai_meta, AIBuildMeta) else AIBuildMeta()
    cur_hard = quality_gate.has_structural_logic_error(cur_meta.quality_issues)
    new_hard = quality_gate.has_structural_logic_error(new_meta.quality_issues)
    if cur_hard != new_hard:
        return not new_hard
    if new_hard:
        return False
    cur_score = score_ai_meta(cur_meta, backend_module=backend)
    new_score = score_ai_meta(new_meta, backend_module=backend)
    if (not new_meta.regression_needed) and cur_meta.regression_needed:
        return True
    if cur_meta.ai_failed and (not new_meta.ai_failed):
        return True
    if new_score < cur_score:
        return True
    return False


def compose_quality_feedback_text(meta: Optional[Any], *, backend_module=None) -> str:
    backend = backend_module or legacy_backend()
    if not isinstance(meta, AIBuildMeta):
        return ""

    parts: list[str] = []
    if meta.unresolved_local_symbols:
        parts.append("局部变量仍未收口为中文：" + "、".join(meta.unresolved_local_symbols[:8]))
    if meta.unresolved_param_symbols:
        parts.append("参数名称仍未收口为中文：" + "、".join(meta.unresolved_param_symbols[:8]))
    if meta.unresolved_logic_symbols:
        parts.append("逻辑说明仍残留变量符号名：" + "、".join(meta.unresolved_logic_symbols[:10]))
    if int(meta.generic_logic_count or 0) > 0:
        parts.append(f"逻辑说明中仍有 {int(meta.generic_logic_count or 0)} 行过于泛化")
    if int(meta.comment_leak_count or 0) > 0:
        parts.append(f"逻辑说明中仍有 {int(meta.comment_leak_count or 0)} 处注释泄漏")
    if int(meta.term_drift_count or 0) > 0:
        parts.append(f"同一轮生成中仍有 {int(meta.term_drift_count or 0)} 处术语漂移")
    if int(meta.over_translation_count or 0) > 0:
        parts.append(f"仍有 {int(meta.over_translation_count or 0)} 处低置信硬翻")
    if int(meta.bad_symbol_guess_count or 0) > 0:
        parts.append(f"仍有 {int(meta.bad_symbol_guess_count or 0)} 处高置信命名冲突")
    if int(meta.logic_placeholders or 0) > 2:
        parts.append("逻辑步骤过度压缩，请保留关键中间计算步骤，不要省略数据转换/同步等中间动作")
    hard_issues = [
        item for item in (meta.quality_issues or ())
        if str((item or {}).get("code") or "") in quality_gate.STRUCTURAL_LOGIC_CODES
        and str((item or {}).get("severity") or "").lower() == "error"
    ]
    for item in hard_issues[:6]:
        line_no = item.get("logic_line") or "?"
        text = utils_module._safe_strip(item.get("logic_text") or item.get("message"))
        anchor = dict(item.get("source_anchor") or {})
        anchor_text = utils_module._safe_strip(
            anchor.get("raw_code") or anchor.get("source") or anchor.get("text") or ""
        )
        detail = f"结构硬错误 {item.get('code')}：逻辑第 {line_no} 行“{text[:80]}”"
        if anchor_text:
            detail += f"；源码锚点“{anchor_text[:80]}”"
        detail += "。仅修复该行，禁止输出 C 运算符、if(、&&、|| 或未闭合括号"
        parts.append(detail)

    # LSP 精确事实注入：为未收口符号补充类型/成员/调用上下文
    lsp_snap = dict(getattr(meta, "lsp_fact_snapshot", {}) or {})
    if lsp_snap:
        type_facts = lsp_snap.get("type_facts") or {}
        if type_facts:
            type_lines = [f"{name}: {t}" for name, t in list(type_facts.items())[:8]]
            parts.append("LSP类型事实：" + "；".join(type_lines))
        member_facts = lsp_snap.get("member_facts") or []
        if member_facts:
            member_lines = [
                f"{m['symbol']}.{m['member']} (owner: {m['owner_type']})"
                for m in member_facts[:8]
                if m.get("symbol") and m.get("member")
            ]
            if member_lines:
                parts.append("LSP成员访问：" + "；".join(member_lines))
        call_facts = lsp_snap.get("call_facts") or []
        if call_facts:
            call_lines = [c.get("call_text", "") for c in call_facts[:6] if c.get("call_text")]
            if call_lines:
                parts.append("LSP调用上下文：" + "；".join(call_lines))
        block_facts = lsp_snap.get("block_facts") or []
        if block_facts:
            block_lines = [
                f"{b['kind']}({b['condition']})"
                for b in block_facts[:6]
                if b.get("kind")
            ]
            if block_lines:
                parts.append("LSP控制结构：" + "；".join(block_lines))

    return "；".join(parts)


def make_regression_cfg(cfg, *, round_idx: int, meta: Optional[Any] = None, backend_module=None):
    backend = backend_module or legacy_backend()
    extra = dict(getattr(cfg, "extra_params", {}) or {})
    retry_times = max(
        1,
        int(utils_module.cfg_get_int(cfg, "ai_retry_times", int(getattr(cfg, "ai_retry_times", 0) or 0))),
    )
    extra["ai_retry_times"] = str(retry_times)
    extra["ai_fail_policy"] = "fallback"
    retry_cfg = clone_cfg(
        cfg,
        ai_workers=1,
        ai_circuit_break=False,
        extra_params=extra,
    )
    if design_regression_force_one_call(cfg, backend_module=backend):
        retry_cfg.ai_one_call = True
    try:
        retry_cfg._ai_cache_salt = f"regression:{round_idx}"
        retry_cfg._ai_regression_round = int(round_idx)
        retry_cfg._in_func_context = True
        retry_cfg._current_func_ai_failed = False
        retry_cfg._skip_ai_current_func = False
        retry_cfg._ai_quality_feedback = compose_quality_feedback_text(meta, backend_module=backend)
        focus_symbols = tuple(
            dict.fromkeys(
                list((meta.unresolved_local_symbols if isinstance(meta, AIBuildMeta) else ()) or ())
                + list((meta.unresolved_param_symbols if isinstance(meta, AIBuildMeta) else ()) or ())
                + list((meta.unresolved_logic_symbols if isinstance(meta, AIBuildMeta) else ()) or ())
            )
        )
        retry_cfg._ai_quality_focus_symbols = focus_symbols[:12]
        # 回归轮对 focus_symbols 放开改名：回归的目的就是纠错
        retry_cfg._ai_regression_allow_rename = set(focus_symbols[:24])
    except Exception:
        pass
    return retry_cfg


def _meta_with_structural_quality(
    meta: AIBuildMeta,
    logic_lines: Sequence[str] | None,
    recovery: Sequence[dict[str, Any]] = (),
) -> AIBuildMeta:
    """Refresh only structural metadata after a targeted line repair."""
    anchors = tuple(meta.logic_source_audit or ())
    structural = quality_gate.inspect_logic_lines(logic_lines, source_anchors=anchors)
    retained = tuple(
        item for item in (meta.quality_issues or ())
        if str((item or {}).get("code") or "") not in quality_gate.STRUCTURAL_LOGIC_CODES
    )
    reasons = [
        reason for reason in (meta.regression_reasons or ())
        if reason not in quality_gate.STRUCTURAL_LOGIC_CODES
    ]
    reasons.extend(str(item.get("code") or "") for item in structural)
    merged_recovery = tuple(meta.quality_recovery or ()) + tuple(recovery or ())
    return replace(
        meta,
        logic_placeholders=sum("待人工修改" in str(line or "") for line in (logic_lines or ())),
        quality_issues=retained + structural,
        regression_reasons=tuple(dict.fromkeys(reason for reason in reasons if reason)),
        regression_needed=bool(reasons),
        quality_recovery=merged_recovery,
    )


def _emit_quality_recovery_event(cfg, task: dict, records: Sequence[dict[str, Any]]) -> None:
    for record in records or ():
        try:
            payload = {
                "type": "ai_quality_recovery",
                "file": str(task.get("file") or ""),
                "func_name": str(task.get("func_name") or ""),
                "func_index": int(task.get("index") or 0),
                "func_pos": int(task.get("func_pos") or 0),
            }
            payload.update(dict(record or {}))
            legacy_backend().gui_event(cfg, payload)
        except Exception:
            pass


def _try_targeted_structural_repair(task: dict, design: FunctionDesign, cfg, *, backend_module=None):
    """Ask AI only for the structurally broken rendered lines.

    The suggestion is accepted solely after the shared gate passes; it can
    never replace a valid control-flow line or alter unrelated logic steps.
    """
    backend = backend_module or legacy_backend()
    meta = design.ai_meta if isinstance(design.ai_meta, AIBuildMeta) else AIBuildMeta()
    issues = tuple(
        item for item in (meta.quality_issues or ())
        if str((item or {}).get("code") or "") in quality_gate.STRUCTURAL_LOGIC_CODES
    )
    if not issues or not design.logic_lines:
        return design, False
    body = str((task.get("func_data") or {}).get("body") or "")
    items: list[dict[str, Any]] = []
    for issue in issues:
        try:
            idx = int(issue.get("logic_line")) - 1
        except (TypeError, ValueError):
            continue
        if idx < 0 or idx >= len(design.logic_lines):
            continue
        current = str(design.logic_lines[idx] or "")
        if logic_utils._is_control_logic_line(current, backend_module=backend):
            continue
        items.append({
            "idx": idx,
            "code": str((issue.get("source_anchor") or {}).get("raw_code") or current),
            "code_cn": current,
            "indent": current[: len(current) - len(current.lstrip())],
            "polish_only": True,
        })
    if not items:
        return design, False
    try:
        suggestions = backend.ai_refine_logic_unknowns(items, body, cfg)
    except Exception as exc:
        recovery = ({"action": "targeted_ai_failed", "error": str(exc)[:160]},)
        return replace(design, ai_meta=_meta_with_structural_quality(meta, design.logic_lines, recovery)), False
    candidate = list(design.logic_lines)
    repaired_lines: list[int] = []
    for item in items:
        idx = int(item["idx"])
        text = utils_module._safe_strip((suggestions or {}).get(str(idx)))
        if not text or not quality_gate.is_safe_ai_text(text):
            continue
        if logic_utils._is_control_logic_line(text, backend_module=backend):
            continue
        if not text.endswith("；"):
            text += "；"
        candidate[idx] = item["indent"] + text
        repaired_lines.append(idx + 1)
    candidate_issues = quality_gate.inspect_logic_lines(candidate, source_anchors=meta.logic_source_audit)
    if candidate_issues or not repaired_lines:
        recovery = ({
            "action": "targeted_ai_rejected",
            "lines": tuple(repaired_lines),
            "remaining_codes": tuple(dict.fromkeys(item.get("code") for item in candidate_issues)),
        },)
        return replace(design, ai_meta=_meta_with_structural_quality(meta, design.logic_lines, recovery)), False
    recovery = ({"action": "targeted_ai_repaired", "lines": tuple(repaired_lines)},)
    repaired_meta = _meta_with_structural_quality(meta, tuple(candidate), recovery)
    return replace(design, logic_lines=tuple(candidate), ai_meta=repaired_meta), True


def _build_deterministic_baseline(task: dict, cfg, build_task_fn, task_cfg_fn, *, backend_module=None):
    """Build an isolated no-AI design used only for line-level recovery."""
    backend = backend_module or legacy_backend()
    baseline_cfg = clone_cfg(cfg, ai_assist=False, ai_mode=0, ai_workers=1)
    baseline_cfg._in_func_context = True
    baseline_task = dict(task)
    baseline_task["func_data"] = copy.deepcopy(task.get("func_data") or {})
    baseline_task_cfg = task_cfg_fn(baseline_cfg, baseline_task)
    if build_task_fn is build_function_design_task:
        return build_task_fn(
            baseline_task["func_data"], baseline_task["module_req_prefix"],
            int(baseline_task["index"]), baseline_task_cfg, backend_module=backend,
        )
    return build_task_fn(
        baseline_task["func_data"], baseline_task["module_req_prefix"],
        int(baseline_task["index"]), baseline_task_cfg,
    )


def _fallback_structural_logic_lines(design: FunctionDesign, baseline: Optional[FunctionDesign]):
    meta = design.ai_meta if isinstance(design.ai_meta, AIBuildMeta) else AIBuildMeta()
    issues = quality_gate.inspect_logic_lines(design.logic_lines, source_anchors=meta.logic_source_audit)
    if not issues:
        return design
    result = list(design.logic_lines or ())
    restored: list[int] = []
    omitted: list[int] = []
    for issue in issues:
        try:
            idx = int(issue.get("logic_line")) - 1
        except (TypeError, ValueError):
            continue
        if idx < 0 or idx >= len(result):
            continue
        fallback = ""
        if baseline is not None and idx < len(baseline.logic_lines or ()):
            fallback = str(baseline.logic_lines[idx] or "")
        if fallback and not quality_gate.inspect_logic_lines((fallback,)):
            result[idx] = fallback
            restored.append(idx + 1)
        else:
            # The deterministic renderer can itself be unable to describe a
            # statement.  The contract is still stronger than completeness:
            # never send the corrupt AI line to DOCX.  Other accepted AI lines
            # remain intact and the omission is auditable in review metadata.
            result[idx] = ""
            omitted.append(idx + 1)
    recovery = ({
        "action": "line_deterministic_fallback",
        "lines": tuple(restored),
        "omitted_lines": tuple(omitted),
        "baseline_available": baseline is not None,
    },)
    return replace(design, logic_lines=tuple(result), ai_meta=_meta_with_structural_quality(meta, tuple(result), recovery))


def maybe_regress_function_design(task: dict, design: Any, cfg, *, backend_module=None):
    backend = backend_module or legacy_backend()
    if (not isinstance(design, FunctionDesign)) or (not getattr(cfg, "ai_assist", False)):
        return design

    rounds = design_regression_rounds(cfg, backend_module=backend)
    current = design
    if rounds <= 0:
        return current

    build_task_fn = getattr(backend, "_build_function_design_task", None) or build_function_design_task
    task_cfg_fn = getattr(backend, "_cfg_with_function_task_context", None) or cfg_with_function_task_context
    prefer_fn = getattr(backend, "_prefer_regression_design", None) or prefer_regression_design
    deterministic_baseline = None

    for round_idx in range(1, rounds + 1):
        # 检查停止信号，避免停止后继续触发新的 AI 回归调用
        if backend.stop_requested(cfg) or getattr(cfg, "_user_cancelled", False):
            backend.vlog(cfg, f"[停止] AI 回归补跑检测到停止信号，退出回归循环")
            break

        meta = current.ai_meta if isinstance(current.ai_meta, AIBuildMeta) else AIBuildMeta()
        if not meta.regression_needed:
            break

        reasons = ", ".join(meta.regression_reasons) or "unknown"
        func_name = str(task.get("func_name") or "")
        backend.gui_event(cfg, {
            "type": "ai_regression_start",
            "file": str(task.get("file") or ""),
            "func_name": func_name,
            "func_index": int(task.get("index") or 0),
            "func_pos": int(task.get("func_pos") or 0),
            "round": round_idx,
            "reasons": list(meta.regression_reasons),
        })
        backend.vlog(cfg, f"函数 {func_name or task.get('index')} 触发 AI 回归补跑，第 {round_idx} 轮：{reasons}")

        retry_cfg = make_regression_cfg(cfg, round_idx=round_idx, meta=meta, backend_module=backend)
        # Inject previous design as revision context for iterative improvement
        if isinstance(current, FunctionDesign):
            _prev = {
                "title": getattr(current, "title", ""),
                "description": "\n".join(getattr(current, "description_lines", ()) or ()),
                "logic_lines": list(getattr(current, "logic_lines", ()) or [])[:20],
                "io_elements": [{"name": getattr(e, "name", ""), "ident": getattr(e, "ident", "")} for e in (getattr(current, "io_elements", ()) or ())],
                "local_elements": [{"name": getattr(e, "name", ""), "ident": getattr(e, "ident", "")} for e in (getattr(current, "local_elements", ()) or ())],
                "quality_issues": list(getattr(meta, "quality_issues", ()) or []),
                "unresolved_symbols": list(getattr(meta, "unresolved_local_symbols", ()) or []) + list(getattr(meta, "unresolved_param_symbols", ()) or []) + list(getattr(meta, "unresolved_logic_symbols", ()) or []),
            }
            if retry_cfg.extra_params is None:
                retry_cfg.extra_params = {}
            retry_cfg.extra_params["revision_context"] = _prev
        has_hard_error = quality_gate.has_structural_logic_error(meta.quality_issues)
        if has_hard_error:
            if deterministic_baseline is None:
                try:
                    deterministic_baseline = _build_deterministic_baseline(
                        task, cfg, build_task_fn, task_cfg_fn, backend_module=backend,
                    )
                except Exception as exc:
                    backend.vlog(cfg, f"确定性逻辑基线构建失败：{exc}")
            retry, improved = _try_targeted_structural_repair(
                task, current, retry_cfg, backend_module=backend,
            )
            if improved:
                current = retry
            new_meta = current.ai_meta if isinstance(current.ai_meta, AIBuildMeta) else AIBuildMeta()
            _emit_quality_recovery_event(cfg, task, new_meta.quality_recovery[-1:])
            backend.gui_event(cfg, {
                "type": "ai_regression_end", "file": str(task.get("file") or ""),
                "func_name": func_name, "func_index": int(task.get("index") or 0),
                "func_pos": int(task.get("func_pos") or 0), "round": round_idx,
                "ok": bool(improved), "improved": bool(improved),
                "reasons": list(new_meta.regression_reasons),
            })
            # Structural repair gets the entire configured budget even when a
            # model response was rejected; the final fallback is deterministic.
            continue
        try:
            task_cfg = task_cfg_fn(retry_cfg, task)
            if build_task_fn is build_function_design_task:
                retry = build_task_fn(
                    task["func_data"],
                    task["module_req_prefix"],
                    int(task["index"]),
                    task_cfg,
                    backend_module=backend,
                )
            else:
                retry = build_task_fn(
                    task["func_data"],
                    task["module_req_prefix"],
                    int(task["index"]),
                    task_cfg,
                )
        except Exception as exc:
            backend.vlog(cfg, f"AI 回归补跑失败：{func_name or task.get('index')}；原因：{exc}")
            backend.gui_event(cfg, {
                "type": "ai_regression_end",
                "file": str(task.get("file") or ""),
                "func_name": func_name,
                "func_index": int(task.get("index") or 0),
                "func_pos": int(task.get("func_pos") or 0),
                "round": round_idx,
                "ok": False,
                "improved": False,
                "error": str(exc),
                "reasons": list(meta.regression_reasons),
            })
            break

        improved = prefer_fn(current, retry)
        if improved:
            current = retry
        new_meta = current.ai_meta if isinstance(current.ai_meta, AIBuildMeta) else AIBuildMeta()
        backend.gui_event(cfg, {
            "type": "ai_regression_end",
            "file": str(task.get("file") or ""),
            "func_name": func_name,
            "func_index": int(task.get("index") or 0),
            "func_pos": int(task.get("func_pos") or 0),
            "round": round_idx,
            "ok": True,
            "improved": bool(improved),
            "reasons": list(new_meta.regression_reasons),
        })
        if (not improved) or (not new_meta.regression_needed):
            break

    final_meta = current.ai_meta if isinstance(current.ai_meta, AIBuildMeta) else AIBuildMeta()
    policy = utils_module.cfg_get_str(
        cfg,
        "ai_quality_hard_fail_policy",
        str(getattr(cfg, "ai_quality_hard_fail_policy", "line_deterministic_fallback") or "line_deterministic_fallback"),
    )
    if quality_gate.has_structural_logic_error(final_meta.quality_issues) and policy == "line_deterministic_fallback":
        if deterministic_baseline is None:
            try:
                deterministic_baseline = _build_deterministic_baseline(
                    task, cfg, build_task_fn, task_cfg_fn, backend_module=backend,
                )
            except Exception as exc:
                backend.vlog(cfg, f"确定性逻辑基线构建失败：{exc}")
        current = _fallback_structural_logic_lines(current, deterministic_baseline)
        final_meta = current.ai_meta if isinstance(current.ai_meta, AIBuildMeta) else AIBuildMeta()
        _emit_quality_recovery_event(cfg, task, final_meta.quality_recovery[-1:])
    return current


def should_degrade_one_call_logic(
    body: str,
    unknowns: Sequence[dict],
    logic_map: dict[str, str],
    logic_text: str,
    cfg,
    *,
    backend_module=None,
) -> bool:
    backend = backend_module or legacy_backend()
    unknown_count = len(unknowns or [])
    if unknown_count <= 0:
        return False

    raw_lines = backend._join_c_line_continuations(body or "").splitlines()
    stmt_lines = [
        line.strip()
        for line in raw_lines
        if line.strip() and line.strip() not in ("{", "}", "};")
    ]
    stmt_count = len(stmt_lines)
    placeholder_count = logic_utils._count_logic_placeholder_lines(logic_text, backend_module=backend)
    logic_got = len(logic_map or {})

    large_unknown_min = max(12, int(utils_module.cfg_get_int(cfg, "one_call_logic_degrade_unknown_min", 24)))
    large_stmt_min = max(40, int(utils_module.cfg_get_int(cfg, "one_call_logic_degrade_stmt_min", 80)))

    is_large_func = unknown_count >= large_unknown_min or stmt_count >= large_stmt_min
    if not is_large_func:
        return False

    if logic_got <= 0:
        return True

    sparse_limit = max(1, unknown_count // 8)
    heavy_placeholder_limit = max(8, unknown_count // 2)
    return logic_got <= sparse_limit and placeholder_count >= heavy_placeholder_limit


def build_degraded_one_call_logic(
    body: str,
    local_vars: Sequence[dict],
    name_map: dict[str, str],
    cfg,
    *,
    backend_module=None,
):
    backend = backend_module or legacy_backend()
    fallback_cfg = clone_cfg(
        cfg,
        ai_assist=False,
        ai_one_call=False,
        ai_logic_policy="hybrid",
    )
    return backend.generate_logic_from_body(body, local_vars, fallback_cfg, name_map=name_map)


def prepare_design_context(func_data: dict, cfg, *, backend_module=None) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    comment_info = dict(func_data.get("comment_info", {}) or {})
    raw_comment_desc = utils_module._safe_strip((func_data.get("comment_info") or {}).get("desc"))
    func_info = dict(func_data.get("func_info", {}) or {})
    body = func_data.get("body", "") or ""
    file_context = dict(func_data.get("file_context", {}) or {})
    family_prefix = utils_module._safe_strip(file_context.get("family_prefix")) or backend._identifier_family_prefix(
        utils_module._safe_strip(func_info.get("func_name"))
    )
    module_key = utils_module._safe_strip(file_context.get("module_key"))
    owner_func = utils_module._safe_strip(func_info.get("func_name"))
    source_file = utils_module._safe_strip(file_context.get("source_file"))
    owner_ret_type = utils_module._safe_strip(func_info.get("ret_type"))

    local_vars = backend.parse_local_variables_from_body(body)
    params = backend.parse_params_from_prototype(func_info)
    local_vars = backend._filter_local_vars_against_params(
        local_vars,
        params,
        cfg=cfg,
        func_name=(func_info.get("func_name") or ""),
    )
    for v in local_vars:
        v["family_prefix"] = family_prefix
        v["module_key"] = module_key
        v["owner_func"] = owner_func
        v["source_file"] = source_file
        v["owner_ret_type"] = owner_ret_type
        v["scope"] = "local"
        v["direction"] = "local"

    input_desc = utils_module._safe_strip(comment_info.get("input_desc"))
    output_desc = utils_module._safe_strip(comment_info.get("output_desc"))
    in_map = backend.parse_param_desc(input_desc, strip_paren_content=True)
    out_map = backend.parse_param_desc(output_desc)
    param_ai_name_map = backend._seed_symbol_memory_into_scope(
        comment_info,
        func_info,
        local_vars,
        params,
        in_map,
        out_map,
    )
    scope_inference_log = backend.infer_scope_symbol_names(
        local_vars,
        params,
        body=body,
        func_info=func_info,
        comment_info=comment_info,
        in_map=in_map,
        out_map=out_map,
        cfg=cfg,
    )
    for name, inference in scope_inference_log.items():
        if inference.candidate_cn and name in in_map:
            param_ai_name_map[name] = inference.candidate_cn
    global_symbol_map, file_symbol_inference_log = backend._build_canonical_file_symbol_map(
        file_context,
        body,
        local_vars,
        params,
        cfg,
    )
    ctx = {
        "cfg": cfg,
        "comment_info": comment_info,
        "raw_comment_desc": raw_comment_desc,
        "func_info": func_info,
        "body": body,
        "file_context": file_context,
        "family_prefix": family_prefix,
        "module_key": module_key,
        "owner_func": owner_func,
        "source_file": source_file,
        "owner_ret_type": owner_ret_type,
        "local_vars": local_vars,
        "params": params,
        "in_map": in_map,
        "out_map": out_map,
        "param_ai_name_map": param_ai_name_map,
        "scope_inference_log": scope_inference_log,
        "global_symbol_map": global_symbol_map,
        "file_symbol_inference_log": file_symbol_inference_log,
    }
    ctx["lsp_fact_pack"] = lsp_fact_utils.build_function_fact_pack(func_data, cfg, backend_module=backend)
    ctx["logic_semantic_pack"] = build_logic_semantic_pack(ctx, backend_module=backend)
    initial_gaps = backend.detect_gaps(
        comment_info,
        local_vars,
        params,
        in_map,
        out_map,
        func_info,
        cfg,
        logic_semantic_pack=ctx.get("logic_semantic_pack"),
        body=body,
    )
    ctx["initial_gaps"] = initial_gaps
    return ctx


def _build_base_name_map(ctx: dict[str, Any], *, backend_module=None) -> dict[str, str]:
    backend = backend_module or legacy_backend()
    typedef_blocks = (ctx.get("file_context") or {}).get("typedefs") or []
    member_symbol_map = backend._extract_member_symbol_map_from_typedefs(typedef_blocks)
    header_member_symbol_map = dict((ctx.get("file_context") or {}).get("member_symbol_map") or {})
    scope_symbol_map = backend._build_local_param_symbol_map(
        ctx.get("local_vars") or [],
        ctx.get("params") or [],
        ctx.get("in_map") or {},
        ctx.get("out_map") or {},
        ctx.get("param_ai_name_map") or {},
    )
    name_map = dict(ctx.get("global_symbol_map") or {})
    name_map.update(member_symbol_map)
    name_map.update(header_member_symbol_map)
    name_map.update(scope_symbol_map)
    for item in (ctx.get("local_vars") or []):
        ident = utils_module._safe_strip((item or {}).get("name"))
        cn_name = utils_module._safe_strip((item or {}).get("cn_name"))
        if ident and cn_name and not backend._is_missing_gap_text(cn_name):
            name_map[ident] = cn_name
    for ident, cn_name in dict(ctx.get("var_cn_map") or {}).items():
        ident_s = utils_module._safe_strip(ident)
        cn_s = utils_module._safe_strip(cn_name)
        if ident_s and cn_s and not backend._is_missing_gap_text(cn_s):
            name_map[ident_s] = cn_s
    name_map.update(_build_control_expression_aliases(ctx, backend_module=backend))
    # AI-supplied field text participates in every later logic rendering step.
    # Drop structurally unsafe values before they can contaminate a bitwise
    # chain or a control condition; deterministic maps remain available.
    name_map = {
        key: value for key, value in name_map.items()
        if quality_gate.is_safe_ai_text(value)
    }
    # Disambiguate: 合并后的 name_map 中若有不同 key 映射到同一中文名，
    # 追加原始名末段以避免逻辑文本中出现重复/混淆。
    def _eligible_disambiguation_key(c_name: str) -> bool:
        key = utils_module._safe_strip(c_name)
        if not key or any(token in key for token in (".", "->", "[", "]")):
            return False
        if re.fullmatch(r"(?:bit|all|mem|word\d+)(?:_(?:u|i)\d+)?", key, flags=re.IGNORECASE):
            return False
        return True

    seen: dict[str, list[str]] = {}
    for c_name, cn_name in name_map.items():
        if not _eligible_disambiguation_key(c_name):
            continue
        seen.setdefault(cn_name, []).append(c_name)
    for cn_name, c_names in seen.items():
        if len(c_names) <= 1:
            continue
        tails = {utils_module._safe_strip(str(c_name).replace("->", ".").split(".")[-1]) for c_name in c_names}
        if len(tails) == 1:
            continue
        for c_name in c_names:
            if not _eligible_disambiguation_key(c_name):
                continue
            compact = re.sub(r"^(?:[glsvp]_)?", "", c_name)
            compact = re.sub(r"_(?:u|i)(?:8|16|32|64)\b", "", compact, flags=re.IGNORECASE)
            compact = compact.strip("_")
            if compact.lower() == c_name.lower():
                compact = c_name[-6:]
            name_map[c_name] = f"{cn_name}({compact})"
    return name_map


def build_design_name_map(ctx: dict[str, Any], *, backend_module=None) -> dict[str, str]:
    backend = backend_module or legacy_backend()
    name_map = _build_base_name_map(ctx, backend_module=backend)
    semantic_pack = dict(ctx.get("logic_semantic_pack") or {})
    for key, value in dict(semantic_pack.get("entity_aliases") or {}).items():
        key_s = utils_module._safe_strip(key)
        value_s = utils_module._safe_strip(value)
        if not key_s or not value_s:
            continue
        existing = utils_module._safe_strip(name_map.get(key_s))
        if existing and (not _is_low_specificity_design_label(existing)) and _is_low_specificity_design_label(value_s):
            continue
        if not quality_gate.is_safe_ai_text(value_s):
            continue
        name_map[key_s] = value_s
    return name_map


def _is_low_specificity_design_label(text: str) -> bool:
    compact = re.sub(r"\s+", "", utils_module._safe_strip(text))
    if not compact:
        return True
    compact = re.sub(r"[（(][^（）()]*[）)]", "", compact).strip()
    return compact in {
        "上一周期值",
        "缓存值",
        "缓存数组",
        "状态值",
        "状态快照",
        "标志值",
        "标志位",
        "请求值",
        "当前值",
        "临时值",
        "中间值",
        "结果值",
        "计数值",
        "循环索引",
        "索引",
        "下标",
        "指令",
        "有效",
    }


_CONTROL_SLOT_LABELS = {
    "CTRL_A1": "A1",
    "CTRL_A2": "A2",
    "CTRL_A3": "A3",
    "CTRL_A4": "A4",
    "0": "A1",
    "1": "A2",
    "2": "A3",
    "3": "A4",
}

_CONTROL_BASE_EXPR_LABELS = {
    "s_actDev_t": "作动器控制参数",
    "gc_act_p": "作动器控制参数",
}

_CONTROL_MEMBER_ARRAY_LABELS = {
    "gc_fpga_p->ramPos_f": "位置反馈",
    "gc_fpga_p->actLoadRatio_f": "载荷系数",
    "g_mmF2degHexRatio_f": "角度码换算系数",
}


def _build_control_expression_aliases(ctx: dict[str, Any], *, backend_module=None) -> dict[str, str]:
    backend = backend_module or legacy_backend()
    body = utils_module._safe_text(ctx.get("body"))
    if not body:
        return {}

    aliases: dict[str, str] = {}

    def _remember(expr: str, label: str) -> None:
        raw_expr = utils_module._safe_strip(expr)
        raw_label = utils_module._safe_strip(label)
        if (not raw_expr) or (not raw_label):
            return
        aliases.setdefault(raw_expr, raw_label)
        aliases.setdefault(raw_expr.replace("->", "."), raw_label)
        aliases.setdefault(re.sub(r"\s+", "", raw_expr.replace("->", ".")), raw_label)

    for base, base_label in _CONTROL_BASE_EXPR_LABELS.items():
        pattern = re.compile(rf"({re.escape(base)}\s*\[\s*([A-Za-z_]\w*|\d+)\s*\])")
        for match in pattern.finditer(body):
            expr = utils_module._safe_strip(match.group(1))
            slot_key = utils_module._safe_strip(match.group(2))
            slot_label = _CONTROL_SLOT_LABELS.get(slot_key)
            if slot_label:
                _remember(expr, f"{slot_label}{base_label}")

    for base, member_label in _CONTROL_MEMBER_ARRAY_LABELS.items():
        pattern = re.compile(rf"({re.escape(base)}\s*\[\s*([A-Za-z_]\w*|\d+)\s*\])")
        for match in pattern.finditer(body):
            expr = utils_module._safe_strip(match.group(1))
            slot_key = utils_module._safe_strip(match.group(2))
            slot_label = _CONTROL_SLOT_LABELS.get(slot_key)
            if slot_label:
                _remember(expr, f"{slot_label}{member_label}")

    return aliases


def _semantic_range_line(range_data: dict[str, Any], key: str) -> int:
    try:
        return int((range_data or {}).get(key) or 0)
    except Exception:
        return 0


def _semantic_raw_code_for_range(ctx: dict[str, Any], range_data: dict[str, Any], item: dict[str, Any]) -> str:
    for key in ("raw_code", "code", "expr", "condition"):
        value = utils_module._safe_strip(item.get(key))
        if value:
            return value
    evidence = item.get("evidence")
    if isinstance(evidence, dict):
        for key in ("raw_code", "code", "expr", "lhs", "rhs", "callee"):
            value = utils_module._safe_strip(evidence.get(key))
            if value:
                return value
    start_line = _semantic_range_line(range_data, "start_line")
    end_line = _semantic_range_line(range_data, "end_line") or start_line
    body_lines = utils_module._safe_text(ctx.get("body")).splitlines()
    if start_line > 0 and end_line >= start_line and end_line <= len(body_lines):
        return "\n".join(body_lines[start_line - 1:end_line]).strip()
    source_file = utils_module._safe_strip((ctx.get("file_context") or {}).get("source_file"))
    if source_file and os.path.exists(source_file) and start_line > 0 and end_line >= start_line:
        try:
            with open(source_file, encoding="utf-8", errors="ignore") as fh:
                source_lines = utils_module._safe_text(fh.read()).splitlines()
            if end_line <= len(source_lines):
                return "\n".join(source_lines[start_line - 1:end_line]).strip()
        except Exception:
            pass
    parts = [utils_module._safe_strip(item.get(key)) for key in ("lhs", "rhs", "callee", "subject", "object", "expr")]
    return " ".join(part for part in parts if part).strip()


def _semantic_source_anchor(ctx: dict[str, Any], item: dict[str, Any]) -> dict[str, Any]:
    range_data = dict((item or {}).get("range") or {})
    file_context = dict(ctx.get("file_context") or {})
    fact_pack = dict(ctx.get("lsp_fact_pack") or {})
    metadata = dict(fact_pack.get("metadata") or {})
    source_file = utils_module._safe_strip(file_context.get("source_file") or metadata.get("source_file"))
    start_line = _semantic_range_line(range_data, "start_line")
    end_line = _semantic_range_line(range_data, "end_line") or start_line
    return {
        "file": source_file,
        "start_line": start_line,
        "end_line": end_line,
        "raw_code": _semantic_raw_code_for_range(ctx, range_data, item),
    }


def _semantic_provenance(item: dict[str, Any], *, provider: str) -> dict[str, Any]:
    try:
        confidence = float((item or {}).get("confidence", 0.0) or 0.0)
    except Exception:
        confidence = 0.0
    return {
        "provider": provider or "structured",
        "source": utils_module._safe_strip((item or {}).get("source") or provider or "structured"),
        "confidence": confidence,
        "verified": bool((item or {}).get("verified")),
    }


_SEMANTIC_REF_RE = re.compile(r"[A-Za-z_]\w*(?:\s*(?:->|\.)\s*[A-Za-z_]\w*)*")


def _semantic_ref_candidates(item: dict[str, Any]) -> list[str]:
    chunks: list[str] = []
    for key in ("condition", "lhs", "rhs", "callee", "expr", "subject", "object", "label"):
        value = utils_module._safe_strip((item or {}).get(key))
        if value:
            chunks.append(value)
    evidence = (item or {}).get("evidence")
    if isinstance(evidence, dict):
        for key in ("lhs", "rhs", "callee", "expr", "subject", "object"):
            value = utils_module._safe_strip(evidence.get(key))
            if value:
                chunks.append(value)
    refs: list[str] = []
    keywords = getattr(legacy_backend(), "_C_KEYWORDS", set())
    for chunk in chunks:
        for match in _SEMANTIC_REF_RE.finditer(chunk):
            text = utils_module._safe_strip(match.group(0))
            if not text or text in keywords:
                continue
            refs.append(text)
            if "->" in text or "." in text:
                parts = [part for part in re.split(r"\s*(?:->|\.)\s*", text) if part]
                refs.extend(parts)
    return list(dict.fromkeys(refs))[:16]


def _semantic_name_refs(
    ctx: dict[str, Any],
    item: dict[str, Any],
    *,
    name_map: dict[str, str],
    entity_aliases: dict[str, str],
    backend_module=None,
) -> tuple[dict[str, Any], ...]:
    backend = backend_module or legacy_backend()
    resolver_ctx = dict(ctx or {})
    resolver_ctx["name_map"] = dict(name_map or {})
    resolver_ctx["entity_aliases"] = dict(entity_aliases or {})
    refs = []
    for raw in _semantic_ref_candidates(item):
        resolved = naming_utils.resolve_symbol_display(
            raw,
            ctx=resolver_ctx,
            file_context=ctx.get("file_context") or {},
            name_map=name_map,
            backend_module=backend,
        )
        if resolved.get("raw"):
            refs.append(resolved)
    return tuple(refs)


def _enrich_semantic_item_v2(
    ctx: dict[str, Any],
    item: dict[str, Any],
    *,
    name_map: dict[str, str],
    entity_aliases: dict[str, str],
    provider: str,
    backend_module=None,
) -> dict[str, Any]:
    enriched = dict(item or {})
    enriched.setdefault("source_anchor", _semantic_source_anchor(ctx, enriched))
    enriched.setdefault("provenance", _semantic_provenance(enriched, provider=provider))
    enriched.setdefault(
        "name_refs",
        _semantic_name_refs(ctx, enriched, name_map=name_map, entity_aliases=entity_aliases, backend_module=backend_module),
    )
    return enriched


def _semantic_quality_summary(pack: dict[str, Any]) -> dict[str, Any]:
    items: list[dict[str, Any]] = []
    for key in ("control_blocks", "state_updates", "call_roles", "return_actions", "flow_actions", "pattern_hits"):
        items.extend([dict(item) for item in (pack.get(key) or ()) if isinstance(item, dict)])
    missing_anchor = 0
    low_confidence_names = 0
    raw_macro_refs = 0
    all_refs: list[dict[str, Any]] = []
    for item in items:
        anchor = dict(item.get("source_anchor") or {})
        if not anchor.get("raw_code"):
            missing_anchor += 1
        for ref in (item.get("name_refs") or ()):
            if not isinstance(ref, dict):
                continue
            all_refs.append(ref)
            try:
                confidence = float(ref.get("confidence", 0.0) or 0.0)
            except Exception:
                confidence = 0.0
            if confidence < 0.6:
                low_confidence_names += 1
            raw = utils_module._safe_strip(ref.get("raw"))
            display = utils_module._safe_strip(ref.get("display"))
            if raw and raw == display and re.fullmatch(r"[A-Z][A-Z0-9_]*", raw):
                raw_macro_refs += 1
    resolver_stats = naming_utils.summarize_name_resolutions(all_refs)
    return {
        "item_count": len(items),
        "missing_source_anchor_count": missing_anchor,
        "low_confidence_name_count": low_confidence_names,
        "raw_macro_ref_count": raw_macro_refs,
        "resolver_stats": resolver_stats,
    }


def _count_effective_statements(body: str, *, backend_module=None) -> int:
    backend = backend_module or legacy_backend()
    return len(logic_utils._collect_statement_like_units(body, backend_module=backend))


def _strip_return_expr_outer_parens(expr: str) -> str:
    value = utils_module._safe_strip(expr)
    while value.startswith("(") and value.endswith(")"):
        depth = 0
        balanced = True
        for idx, ch in enumerate(value):
            if ch == "(":
                depth += 1
            elif ch == ")":
                depth -= 1
                if depth == 0 and idx != len(value) - 1:
                    balanced = False
                    break
            if depth < 0:
                balanced = False
                break
        if not balanced or depth != 0:
            break
        value = utils_module._safe_strip(value[1:-1])
    return value


def _collect_flow_actions(
    body: str,
    *,
    function_start_line: int = 0,
) -> list[dict[str, Any]]:
    actions: list[dict[str, Any]] = []
    for idx, raw in enumerate(parse_utils._join_c_line_continuations(body or "").splitlines(), start=1):
        code, _comments = parse_utils._split_code_and_comments_for_symbol(raw)
        line_text = utils_module._safe_strip(code)
        if not line_text:
            continue
        line_no = function_start_line + idx if function_start_line > 0 else idx
        for match in re.finditer(r"\b(?P<kind>break|continue)\s*;", line_text):
            kind = utils_module._safe_strip(match.group("kind"))
            if not kind:
                continue
            actions.append(
                {
                    "kind": kind,
                    "range": {
                        "start_line": line_no,
                        "end_line": line_no,
                        "start_col": int(match.start() + 1),
                        "end_col": int(match.end() + 1),
                    },
                    "source": "structured",
                    "confidence": 0.9,
                    "verified": True,
                }
            )
    return actions


def _clean_return_expr(expr: str) -> str:
    expr = utils_module._safe_strip(expr)
    expr = re.sub(r"/\*.*?\*/", "", expr, flags=re.S)
    expr = re.sub(r"//.*$", "", expr).strip()
    return _strip_return_expr_outer_parens(expr)


def _parse_return_exprs_from_body(body: str) -> tuple[str, ...]:
    backend = legacy_backend()
    text = backend._strip_c_comments_keep_layout(utils_module._safe_text(body))
    exprs: list[str] = []
    for match in re.finditer(r"\breturn\b(?P<expr>.*?);", text, flags=re.S):
        expr = _clean_return_expr(match.group("expr"))
        if expr:
            exprs.append(expr)
    return tuple(exprs)


def _parse_return_expr_from_body(body: str) -> str:
    exprs = _parse_return_exprs_from_body(body)
    return exprs[0] if exprs else ""


_KNOWN_STATEMENT_MACROS = (
    "SPI_FLASH_CS_LOW",
    "SPI_FLASH_CS_HIGH",
    "NOP",
)


def _collect_statement_macro_roles(body: str, *, function_start_line: int = 0) -> list[dict[str, Any]]:
    """Capture bare macro statements that represent observable embedded-C actions."""
    out: list[dict[str, Any]] = []
    seen: set[tuple[str, int]] = set()
    lines = legacy_backend()._join_c_line_continuations(body or "").splitlines()
    for idx, raw in enumerate(lines, start=1):
        code = re.sub(r"//.*", "", raw)
        code = re.sub(r"/\*.*?\*/", "", code)
        if not utils_module._safe_strip(code):
            continue
        line_no = function_start_line + idx if function_start_line > 0 else idx
        for macro in _KNOWN_STATEMENT_MACROS:
            if not re.search(rf"(?<![A-Za-z0-9_]){re.escape(macro)}(?![A-Za-z0-9_])", code):
                continue
            key = (macro, line_no)
            if key in seen:
                continue
            seen.add(key)
            out.append(
                {
                    "callee": macro,
                    "role": _classify_call_role(macro),
                    "range": {"start_line": line_no, "end_line": line_no, "start_col": 1, "end_col": 1},
                    "source": "structured_macro",
                    "confidence": 0.78,
                    "verified": True,
                }
            )
    return out


def _collect_memory_call_roles(body: str, *, function_start_line: int = 0) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    lines = legacy_backend()._join_c_line_continuations(body or "").splitlines()
    for idx, raw in enumerate(lines, start=1):
        code = re.sub(r"//.*", "", raw)
        code = re.sub(r"/\*.*?\*/", "", code).strip()
        if not code:
            continue
        match = re.match(r"^(?P<callee>memset|memcpy|memmove)\s*\((?P<args>.*)\)\s*;?\s*$", code)
        if not match:
            continue
        callee = utils_module._safe_strip(match.group("callee"))
        args = logic_utils._split_c_call_args(match.group("args"))
        line_no = function_start_line + idx if function_start_line > 0 else idx
        role = "内存设置" if callee == "memset" else "内存拷贝"
        out.append(
            {
                "callee": callee,
                "args": args,
                "role": role,
                "range": {"start_line": line_no, "end_line": line_no, "start_col": 1, "end_col": 1},
                "source": "structured_call",
                "confidence": 0.86,
                "verified": True,
            }
        )
    return out


def _collect_statement_comment_hints_by_line(body: str, *, function_start_line: int = 0) -> dict[int, list[dict[str, Any]]]:
    out: dict[int, list[dict[str, Any]]] = {}
    pending_comments: list[str] = []
    lines = legacy_backend()._join_c_line_continuations(body or "").splitlines()
    for idx, raw in enumerate(lines, start=1):
        code, comments = parse_utils._split_code_and_comments_for_symbol(raw)
        inline_comments = [c for c in comments if c and not parse_utils._is_non_semantic_comment(c)]
        core = code.replace("{", "").replace("}", "").replace(";", "").strip()
        if not core:
            pending_comments.extend(inline_comments)
            continue
        if logic_utils.is_declaration_line(code):
            pending_comments = []
            continue
        attached = inline_comments[:] if inline_comments else pending_comments[:]
        pending_comments = []
        hints = parse_utils.extract_statement_hints(code, attached)
        hint_dicts = [
            {
                "kind": utils_module._safe_strip(getattr(hint, "kind", "")),
                "text": utils_module._safe_strip(getattr(hint, "text", "")),
                "confidence": float(getattr(hint, "confidence", 0.0) or 0.0),
            }
            for hint in hints
            if utils_module._safe_strip(getattr(hint, "kind", ""))
            and utils_module._safe_strip(getattr(hint, "text", ""))
        ]
        if not hint_dicts:
            continue
        out[idx] = hint_dicts
        if function_start_line > 0:
            out[function_start_line + idx] = hint_dicts
    return out


def _repair_function_desc_by_domain(func_name: str, desc: str, *, current_desc: str = "") -> str:
    value = utils_module._safe_strip(desc)
    if not value:
        return ""
    ident = utils_module._safe_strip(func_name).lower()
    context = value + " " + utils_module._safe_strip(current_desc)
    if (
        re.search(r"write\s*(?:dis|disable)|writedis", ident)
        or "写禁止" in context
        or "写禁用" in context
    ):
        if re.search(r"(?:关闭|关).*写保护|写保护(?:关闭|失能)|写禁关", value):
            return "发送写禁止指令，禁止后续FLASH写入"
    if "spiflash" in ident and "datatrans" in ident:
        if "转换" in value:
            return "通过SPI接口完成FLASH数据交互传输"
    if "refuelstagepreset" in ident:
        return "根据加油模式和目标油箱发送开阀预位命令，并在阀位和泵低压检查后切换加油执行或故障结束状态"
    return value


def _collect_local_initializer_updates(
    body: str,
    local_vars: Sequence[dict[str, Any]],
    *,
    function_start_line: int = 0,
) -> list[dict[str, Any]]:
    local_names = {
        utils_module._safe_strip((item or {}).get("name"))
        for item in (local_vars or [])
        if utils_module._safe_strip((item or {}).get("name"))
    }
    if not local_names:
        return []
    type_prefix = r"(?:static\s+|const\s+|volatile\s+|register\s+)*"
    type_name = r"(?:struct\s+\w+|union\s+\w+|enum\s+\w+|[A-Za-z_]\w*)"
    type_suffix = r"(?:\s+(?:const|volatile)\b)*"
    pointer = r"(?:\s*(?:(?:const|volatile)\s*)?\*\s*(?:const\s+|volatile\s+)*)*"
    decl_init_re = re.compile(
        rf"^\s*(?P<type>{type_prefix}{type_name}{type_suffix}{pointer})\s+"
        rf"(?P<name>[A-Za-z_]\w*)(?:\s*\[[^\]]*\])?\s*=\s*(?P<rhs>[^;]+?)\s*;\s*$"
    )
    out: list[dict[str, Any]] = []
    lines = legacy_backend()._join_c_line_continuations(body or "").splitlines()
    for idx, raw in enumerate(lines, start=1):
        code = re.sub(r"//.*", "", raw)
        code = re.sub(r"/\*.*?\*/", "", code).strip()
        if not code:
            continue
        match = decl_init_re.match(code)
        if not match:
            continue
        name = utils_module._safe_strip(match.group("name"))
        rhs = utils_module._safe_strip(match.group("rhs"))
        if name not in local_names or not rhs:
            continue
        line_no = function_start_line + idx if function_start_line > 0 else idx
        out.append(
            {
                "kind": "local_init",
                "lhs": name,
                "rhs": rhs,
                "label": "",
                "range": {"start_line": line_no, "end_line": line_no, "start_col": 1, "end_col": 1},
                "source": "structured_decl",
                "confidence": 0.86,
                "verified": True,
            }
        )
    return out


def _collect_complete_assignment_rhs(
    body: str,
    *,
    function_start_line: int = 0,
) -> tuple[dict[tuple[str, int], str], dict[str, str]]:
    by_line: dict[tuple[str, int], str] = {}
    by_lhs_values: dict[str, list[str]] = {}
    pending: list[str] = []
    pending_start_idx = 0
    lines = parse_utils._join_c_line_continuations(body or "").splitlines()
    for idx, raw in enumerate(lines, start=1):
        code = re.sub(r"/\*.*?\*/", "", raw)
        code = re.sub(r"//.*", "", code).strip()
        if not code or code.startswith("#"):
            continue
        if not pending and (
            code in {"{", "}", "};"}
            or re.match(r"^(?:if|else\s+if|else|while|for|switch|case|default)\b", code)
        ):
            continue
        if not pending:
            pending_start_idx = idx
        pending.append(code)
        joined = " ".join(part for part in pending if part).strip()
        if not joined:
            pending.clear()
            continue
        if not joined.endswith(";"):
            continue
        pending.clear()
        if re.match(r"^(?:if|while|for|switch)\s*\(", joined):
            continue
        parts = logic_utils._split_plain_assignment(joined)
        if not parts:
            continue
        lhs, rhs = (utils_module._safe_strip(parts[0]), utils_module._safe_strip(parts[1]))
        if not lhs or not rhs:
            continue
        candidate_lines = {pending_start_idx}
        if function_start_line > 0:
            candidate_lines.add(function_start_line + pending_start_idx)
            candidate_lines.add(max(1, function_start_line + pending_start_idx - 1))
        for line_no in candidate_lines:
            by_line[(lhs, line_no)] = rhs
        by_lhs_values.setdefault(lhs, []).append(rhs)
    unique_by_lhs = {
        lhs: values[0]
        for lhs, values in by_lhs_values.items()
        if len(set(values)) == 1
    }
    return by_line, unique_by_lhs


def _root_identifier_from_expr(expr: str) -> str:
    value = _strip_return_expr_outer_parens(utils_module._safe_strip(expr))
    value = re.sub(r"^(?:const|volatile)\s+", "", value)
    value = value.lstrip("*& \t")
    value = _strip_return_expr_outer_parens(value)
    match = re.search(r"\b[A-Za-z_]\w*\b", value)
    return match.group(0) if match else ""


def _is_simple_return_io_expr(expr: str) -> bool:
    value = _strip_return_expr_outer_parens(utils_module._safe_strip(expr)).replace("->", ".")
    if not value:
        return False
    return bool(re.fullmatch(r"[A-Za-z_]\w*(?:\s*(?:\.\s*[A-Za-z_]\w*|\[[^\]]+\]))*", value))


def _strip_return_candidate_expr(expr: str) -> str:
    value = _strip_return_expr_outer_parens(utils_module._safe_strip(expr))
    while True:
        match = re.match(
            r"^\(\s*(?:const\s+|volatile\s+)?(?:struct\s+\w+|union\s+\w+|enum\s+\w+|[A-Za-z_]\w*(?:\s+[A-Za-z_]\w*)*)\s*(?:\*\s*)*\)\s*(.+)$",
            value,
        )
        if not match:
            break
        value = _strip_return_expr_outer_parens(match.group(1))
    return value


def _return_ternary_candidate_exprs(expr: str) -> tuple[str, ...]:
    value = _strip_return_expr_outer_parens(utils_module._safe_strip(expr))
    if "?" not in value or ":" not in value:
        return ()

    def walk(part: str) -> list[str]:
        part2 = _strip_return_expr_outer_parens(utils_module._safe_strip(part))
        split = logic_utils._split_top_level_ternary(part2)
        if split:
            _cond, true_expr, false_expr = split
            return walk(true_expr) + walk(false_expr)
        candidate = _strip_return_candidate_expr(part2)
        return [candidate] if _is_simple_return_io_expr(candidate) else []

    split = logic_utils._split_top_level_ternary(value)
    if not split:
        return ()
    _cond, true_expr, false_expr = split
    candidates = walk(true_expr) + walk(false_expr)
    return tuple(dict.fromkeys(candidate for candidate in candidates if candidate))


def _return_simple_candidate_exprs_from_body(body: str) -> tuple[str, ...]:
    candidates: list[str] = []
    for expr in _parse_return_exprs_from_body(body):
        ternary_candidates = _return_ternary_candidate_exprs(expr)
        if ternary_candidates:
            candidates.extend(ternary_candidates)
            continue
        candidate = _strip_return_candidate_expr(expr)
        if _is_simple_return_io_expr(candidate):
            candidates.append(candidate)
    return tuple(dict.fromkeys(candidate for candidate in candidates if candidate))


def _last_identifier_from_expr(expr: str) -> str:
    names = re.findall(r"\b[A-Za-z_]\w*\b", utils_module._safe_strip(expr))
    return names[-1] if names else ""


def _member_identifier_from_expr(expr: str) -> str:
    value = utils_module._safe_strip(expr).replace("->", ".")
    if "." not in value:
        return ""
    value = re.sub(r"\[[^\]]*\]", "", value)
    parts = [part for part in value.split(".") if part]
    if len(parts) < 2:
        return ""
    match = re.search(r"\b[A-Za-z_]\w*\b", parts[-1])
    return match.group(0) if match else ""


def _last_array_index_from_expr(expr: str) -> str:
    indexes = re.findall(r"\[\s*([^\]]+?)\s*\]", utils_module._safe_strip(expr))
    return utils_module._safe_strip(indexes[-1]) if indexes else ""


def _lookup_label_from_maps(keys: Sequence[str], maps: Sequence[dict[str, str]]) -> str:
    for key in keys:
        key_s = utils_module._safe_strip(key)
        if not key_s:
            continue
        for mapping in maps:
            text = utils_module._safe_strip((mapping or {}).get(key_s))
            if text:
                return text
    return ""


def _clean_member_label_for_display(text: str) -> str:
    value = utils_module._safe_strip(text)
    if not value:
        return ""
    value = re.sub(r"^bit\d+\s*[:：]\s*", "", value, flags=re.IGNORECASE).strip()
    bit_wrapped = re.fullmatch(r"bit\d+\(([^()]+)\)", value, flags=re.IGNORECASE)
    if bit_wrapped:
        return utils_module._safe_strip(bit_wrapped.group(1))
    if re.fullmatch(r"bit\d+", value, flags=re.IGNORECASE):
        return ""
    return value


def _combine_index_member_display_name(index_cn: str, member_cn: str) -> str:
    index = re.sub(r"\s+", "", utils_module._safe_strip(index_cn))
    member = re.sub(r"\s+", "", utils_module._safe_strip(member_cn))
    if not index:
        return member
    if not member:
        return index
    member = re.sub(r"^(?:高、低|高低)", "", member)
    if index.endswith("握手") and member.startswith("握手"):
        member = member[len("握手"):]
    if index.endswith("状态") and member.startswith("状态"):
        member = member[len("状态"):]
    return f"{index}{member}"


def _infer_scalar_c_type_from_ident(ident: str) -> str:
    value = utils_module._safe_strip(ident)
    root = _root_identifier_from_expr(value) or value
    candidates = []
    member = _member_identifier_from_expr(value)
    last = _last_identifier_from_expr(value)
    for item in (member, root, last, value):
        text = utils_module._safe_strip(item)
        if text and text not in candidates:
            candidates.append(text)
    suffix_match = None
    for candidate in candidates:
        suffix_match = re.search(r"_([uif])(?:int)?(8|16|32|64)\b", candidate, flags=re.IGNORECASE)
        if suffix_match:
            break
    if not suffix_match:
        return ""
    kind = suffix_match.group(1).lower()
    bits = suffix_match.group(2)
    if kind == "u":
        return f"Uint{bits}"
    if kind == "i":
        return f"Sint{bits}"
    if kind == "f":
        return "float" if bits == "32" else "double"
    return ""


def _normalize_member_decl_type(type_text: str) -> str:
    value = utils_module._safe_strip(type_text)
    if not value:
        return ""
    value = re.sub(r"\b(?:static|extern|register)\b", " ", value)
    value = re.sub(r"\s+", " ", value.replace(" *", "*").replace("* ", "*")).strip()
    return value


def _extract_member_type_map_from_typedefs(typedef_blocks: Sequence[str]) -> dict[str, str]:
    type_map: dict[str, str] = {}
    for block in typedef_blocks or ():
        text = utils_module._safe_text(block)
        if not text:
            continue
        type_names: list[str] = []
        open_match = re.search(r"\b(?:struct|union)\s+([A-Za-z_]\w*)", text)
        if open_match:
            type_names.append(open_match.group(1))
        close_match = re.search(r"}\s*([^;]+);", text, flags=re.S)
        if close_match:
            for part in close_match.group(1).split(","):
                part = re.sub(r"\[[^\]]*\]", "", part)
                names = re.findall(r"\b[A-Za-z_]\w*\b", part)
                if names and names[-1] not in type_names:
                    type_names.append(names[-1])
        for raw in text.splitlines():
            code = re.sub(r"/\*.*?\*/", "", raw)
            code = re.sub(r"//.*", "", code)
            core = utils_module._safe_strip(code)
            if not core or not core.endswith(";") or "(" in core or core.lstrip().startswith(("#", "typedef", "}")):
                continue
            stmt = core[:-1].strip()
            stmt = stmt.split("=", 1)[0].strip()
            stmt = stmt.split(":", 1)[0].strip()
            stmt = re.sub(r"\[[^\]]*\]", "", stmt)
            match = re.match(r"(?P<type>.+?)(?P<ptr>\s*\*+\s*)?(?P<name>[A-Za-z_]\w*)$", stmt)
            if not match:
                continue
            member_name = utils_module._safe_strip(match.group("name"))
            member_type = _normalize_member_decl_type((match.group("type") or "") + ("*" if match.group("ptr") else ""))
            if not member_name or not member_type:
                continue
            type_map.setdefault(member_name, member_type)
            for type_name in type_names:
                type_key = utils_module._safe_strip(type_name)
                if type_key:
                    type_map.setdefault(f"{type_key}.{member_name}", member_type)
    return type_map


def _infer_io_c_type(ctx: dict[str, Any], ident: str, *, backend_module=None) -> str:
    _ = backend_module or legacy_backend()
    value = utils_module._safe_strip(ident)
    suffix_type = _infer_scalar_c_type_from_ident(value)
    if suffix_type:
        return suffix_type
    file_context = dict(ctx.get("file_context") or {})
    typedef_blocks = list(file_context.get("typedefs") or []) + list(file_context.get("header_typedefs") or [])
    member_type_map = _extract_member_type_map_from_typedefs(typedef_blocks)
    if not member_type_map:
        return ""
    member = _member_identifier_from_expr(value)
    root = _root_identifier_from_expr(value)
    keys = [value, value.replace("->", "."), member, root, _last_identifier_from_expr(value)]
    for key in keys:
        type_text = utils_module._safe_strip(member_type_map.get(utils_module._safe_strip(key)))
        if type_text:
            return type_text
    return ""


def _lookup_io_display_name(ctx: dict[str, Any], ident: str, *, backend_module=None, fallback: str = "") -> str:
    backend = backend_module or legacy_backend()
    value = utils_module._safe_strip(ident)
    root = _root_identifier_from_expr(value)
    last = _last_identifier_from_expr(value)
    file_context = dict(ctx.get("file_context") or {})
    member_symbol_map = dict(file_context.get("member_symbol_map") or {})
    maps_for_member: list[dict[str, str]] = [
        member_symbol_map,
        dict(ctx.get("global_symbol_map") or {}),
        dict(file_context.get("symbol_map") or {}),
        dict(file_context.get("glossary") or {}),
    ]
    if last and re.search(r"(?:\.|->)", value):
        normalized = re.sub(r"\s+", "", value.replace("->", "."))
        chain_parts = [part for part in normalized.split(".") if part]
        has_transparent_container = any(
            logic_utils._is_transparent_union_container_member(part)
            for part in chain_parts[1:]
        )
        if len(chain_parts) >= 2 and has_transparent_container:
            root_part = chain_parts[0]
            filtered_parts = [
                part
                for part in chain_parts[1:]
                if not logic_utils._is_transparent_union_container_member(part)
            ]
            if filtered_parts:
                owner_keys: list[str] = []
                if len(filtered_parts) == 1:
                    owner_keys.extend([f"{root_part}.{filtered_parts[0]}", filtered_parts[0]])
                else:
                    owner_keys.extend(
                        [
                            f"{root_part}.{'.'.join(filtered_parts)}",
                            ".".join(filtered_parts),
                            f"{filtered_parts[0]}.{filtered_parts[-1]}",
                        ]
                    )
                owner_cn = _clean_member_label_for_display(
                    _lookup_label_from_maps(
                        owner_keys,
                        [
                            member_symbol_map,
                            dict(ctx.get("global_symbol_map") or {}),
                            dict(file_context.get("symbol_map") or {}),
                        ],
                    )
                )
                if owner_cn:
                    return backend._shorten_element_display_name(owner_cn, fallback=filtered_parts[-1])
        member_keys = [value, value.replace("->", "."), normalized]
        if root:
            member_keys.append(f"{root}.{last}")
        if not logic_utils._is_transparent_union_container_member(last):
            member_keys.append(last)
        member_cn = _lookup_label_from_maps(member_keys, [member_symbol_map])
        if member_cn:
            return backend._shorten_element_display_name(_clean_member_label_for_display(member_cn), fallback=last)
        member_ident = _member_identifier_from_expr(value)
        if member_ident:
            raw_member_cn = _lookup_label_from_maps([member_ident], maps_for_member)
            index_raw = _last_array_index_from_expr(value)
            index_cn = _lookup_label_from_maps([index_raw], maps_for_member)
            if raw_member_cn and index_cn:
                return backend._shorten_element_display_name(
                    _combine_index_member_display_name(index_cn, _clean_member_label_for_display(raw_member_cn)),
                    fallback=member_ident,
                )
            if raw_member_cn:
                return backend._shorten_element_display_name(_clean_member_label_for_display(raw_member_cn), fallback=member_ident)
    local_symbol_map = {
        utils_module._safe_strip((item or {}).get("name")): utils_module._safe_strip((item or {}).get("cn_name") or (item or {}).get("usage"))
        for item in (ctx.get("local_vars") or [])
        if utils_module._safe_strip((item or {}).get("name"))
        and utils_module._safe_strip((item or {}).get("cn_name") or (item or {}).get("usage"))
    }
    maps: list[dict[str, str]] = [
        dict(ctx.get("var_cn_map") or {}),
        local_symbol_map,
        dict(ctx.get("global_symbol_map") or {}),
        dict(file_context.get("symbol_map") or {}),
        dict(file_context.get("glossary") or {}),
        member_symbol_map,
    ]
    for key in (value, root, last):
        if not key:
            continue
        for mapping in maps:
            cn = utils_module._safe_strip(mapping.get(key))
            if cn:
                return backend._shorten_element_display_name(cn, fallback=key)
    guess_key = root or last or value
    guessed = utils_module._safe_strip(backend._guess_cn_from_ident(guess_key, glossary=getattr(backend, "DOMAIN_GLOSSARY", {})))
    return backend._shorten_element_display_name(guessed or fallback or guess_key, fallback=fallback or guess_key)


def _clean_return_candidate_display_label(text: str) -> str:
    value = utils_module._safe_strip(text)
    if not value:
        return ""
    value = re.sub(r"\s+", "", value)
    value = value.strip("：:，,；;。.-")
    return value


def _return_desc_label_for_candidate(return_desc: str, candidate: str) -> str:
    candidate_text = utils_module._safe_strip(candidate)
    if not candidate_text:
        return ""
    for raw_line in re.split(r"[\r\n；;]+", utils_module._safe_text(return_desc)):
        line = utils_module._safe_strip(raw_line)
        if not line or candidate_text not in line:
            continue
        for sep in ("==>", "=>", "----", "---", "--", "->", "：", ":"):
            if sep not in line:
                continue
            left, right = line.split(sep, 1)
            left = utils_module._safe_strip(left)
            right = utils_module._safe_strip(right)
            if candidate_text in right and left:
                return _clean_return_candidate_display_label(left)
            if candidate_text in left and right:
                return _clean_return_candidate_display_label(right.replace(candidate_text, ""))
        stripped = _clean_return_candidate_display_label(line.replace(candidate_text, ""))
        if stripped:
            return stripped
    return ""


def _return_candidate_label_needs_status(label: str, candidate: str) -> bool:
    value = utils_module._safe_strip(label)
    upper = utils_module._safe_strip(candidate).upper()
    if not value or not upper:
        return False
    if any(token in upper for token in ("ERR", "ERROR", "FAIL", "FAILED", "INVALID", "NG")):
        return not re.search(r"(?:未通过|错误|故障|失败|异常|无效)", value)
    if any(token in upper for token in ("OK", "PASS", "PASSED", "VALID", "SUCCESS")):
        return not re.search(r"(?:通过|正常|有效|成功|无故障)", value)
    return False


def _return_candidate_status_from_ident(candidate: str) -> str:
    upper = utils_module._safe_strip(candidate).upper()
    if not upper:
        return ""
    if any(token in upper for token in ("ERR", "ERROR", "FAIL", "FAILED", "INVALID", "NG")):
        return "无效" if "INVALID" in upper else "异常"
    if any(token in upper for token in ("OK", "PASS", "PASSED", "VALID", "SUCCESS")):
        return "有效"
    return ""


def _merge_return_candidate_status_label(base: str, status: str) -> str:
    base_text = _clean_return_candidate_display_label(base)
    status_text = _clean_return_candidate_display_label(status)
    if not base_text:
        return status_text
    if not status_text or status_text in base_text:
        return base_text
    limit = min(len(base_text), len(status_text))
    for size in range(limit, 0, -1):
        if base_text.endswith(status_text[:size]):
            return f"{base_text}{status_text[size:]}"
    return f"{base_text}{status_text}"


def _specific_return_candidate_macro_label(candidate: str, current_label: str, *, backend_module=None) -> str:
    backend = backend_module or legacy_backend()
    value = utils_module._safe_strip(candidate)
    status = _return_candidate_status_from_ident(value)
    current = _clean_return_candidate_display_label(current_label)
    if not value or not status or (current and status in current):
        return ""
    if not backend._is_macro_identifier(_last_identifier_from_expr(value) or value):
        return ""
    guessed = utils_module._safe_strip(
        backend._guess_cn_from_ident(_last_identifier_from_expr(value) or value, glossary=getattr(backend, "DOMAIN_GLOSSARY", {}))
    )
    guessed = _clean_return_candidate_display_label(guessed)
    if guessed and guessed != value and status in guessed:
        return guessed
    if current:
        return _merge_return_candidate_status_label(current, status)
    return status


def _lookup_return_candidate_display_name(
    ctx: dict[str, Any],
    candidate: str,
    *,
    backend_module=None,
    fallback: str = "",
) -> str:
    backend = backend_module or legacy_backend()
    value = utils_module._safe_strip(candidate)
    if not value:
        return utils_module._safe_strip(fallback)
    root = _root_identifier_from_expr(value)
    last = _last_identifier_from_expr(value)
    file_context = dict(ctx.get("file_context") or {})
    raw_maps: list[dict[str, str]] = [
        dict(ctx.get("global_symbol_map") or {}),
        dict(file_context.get("symbol_map") or {}),
        dict(file_context.get("glossary") or {}),
    ]
    comment_info = dict(ctx.get("comment_info") or {})
    desc_label = _return_desc_label_for_candidate(comment_info.get("return_desc", ""), value)
    raw_label = _lookup_label_from_maps([value, root, last], raw_maps)
    if raw_label and value not in raw_label:
        raw_clean = _clean_return_candidate_display_label(raw_label)
        if desc_label and _return_candidate_label_needs_status(raw_clean, value):
            return _merge_return_candidate_status_label(raw_clean, desc_label)
        specific_label = _specific_return_candidate_macro_label(value, raw_clean, backend_module=backend)
        if specific_label:
            return specific_label
        return raw_clean

    mapped = _lookup_io_display_name(ctx, value, backend_module=backend, fallback="")
    if mapped and mapped != value:
        if desc_label and _return_candidate_label_needs_status(mapped, value):
            return _merge_return_candidate_status_label(mapped, desc_label)
        specific_label = _specific_return_candidate_macro_label(value, mapped, backend_module=backend)
        if specific_label:
            return specific_label
        return mapped

    if desc_label:
        return desc_label
    return utils_module._safe_strip(fallback or value)


def _global_write_io_elements(ctx: dict[str, Any], seen_idents: set[str], *, backend_module=None) -> list[IOElement]:
    backend = backend_module or legacy_backend()
    fact_pack = dict(ctx.get("lsp_fact_pack") or {})
    writes = [dict(item) for item in (fact_pack.get("writes") or []) if isinstance(item, dict)]
    if not writes:
        return []
    local_names = {
        utils_module._safe_strip((item or {}).get("name"))
        for item in (ctx.get("local_vars") or [])
        if utils_module._safe_strip((item or {}).get("name"))
    }
    param_names = {
        utils_module._safe_strip((item or {}).get("name"))
        for item in (ctx.get("params") or [])
        if utils_module._safe_strip((item or {}).get("name"))
    }
    result: list[IOElement] = []
    for item in writes:
        lhs = utils_module._safe_strip(item.get("lhs"))
        if not lhs or lhs in seen_idents:
            continue
        root = _root_identifier_from_expr(lhs)
        if (not root) or root in local_names or root in param_names or root in getattr(backend, "_C_KEYWORDS", set()):
            continue
        display_name = _lookup_io_display_name(ctx, lhs, backend_module=backend, fallback="全局状态")
        result.append(
            IOElement(
                name=display_name or "全局状态",
                ident=lhs,
                c_type=_infer_io_c_type(ctx, lhs, backend_module=backend),
                direction="输出",
            )
        )
        seen_idents.add(lhs)
    return result


def _is_pointer_or_array_param(param_info: dict[str, Any]) -> bool:
    ptype = utils_module._safe_strip((param_info or {}).get("type"))
    name = utils_module._safe_strip((param_info or {}).get("name"))
    return bool("*" in ptype or "[" in ptype or name.startswith(("p_", "pp_", "vp_", "v_p_", "gp_", "lp_", "sp_", "cp_", "tp_")))


def _is_const_pointer_param(param_info: dict[str, Any]) -> bool:
    return bool("*" in utils_module._safe_strip((param_info or {}).get("type")) and re.search(r"\bconst\b", utils_module._safe_strip((param_info or {}).get("type")), re.I))


def _param_has_external_write(ctx: dict[str, Any], param_info: dict[str, Any]) -> bool:
    name = utils_module._safe_strip((param_info or {}).get("name"))
    if not name or not _is_pointer_or_array_param(param_info) or _is_const_pointer_param(param_info):
        return False
    fact_pack = dict(ctx.get("lsp_fact_pack") or {})
    writes = [dict(item) for item in (fact_pack.get("writes") or []) if isinstance(item, dict)]
    if not writes:
        return False
    ident = re.escape(name)
    patterns = (
        re.compile(rf"^\s*\*+\s*\(?\s*{ident}\s*\)?(?:\b|\s*(?:->|\.|\[))"),
        re.compile(rf"^\s*{ident}\s*(?:->|\[)"),
    )
    for item in writes:
        lhs = utils_module._safe_strip(item.get("lhs"))
        if not lhs:
            continue
        if any(pattern.search(lhs) for pattern in patterns):
            return True
    return False


def _function_param_direction(ctx: dict[str, Any], param_info: dict[str, Any]) -> str:
    name = utils_module._safe_strip((param_info or {}).get("name"))
    in_map = dict(ctx.get("in_map") or {})
    out_map = dict(ctx.get("out_map") or {})
    is_input = name in in_map
    is_output = name in out_map or _param_has_external_write(ctx, param_info)
    if is_input and is_output:
        return "输入/输出"
    if is_output:
        return "输出"
    return "输入"


def _effect_analysis_mode(cfg) -> str:
    value = utils_module.cfg_get_str(cfg, "effect_analysis_mode", str(getattr(cfg, "effect_analysis_mode", "one_hop") or "one_hop"))
    return value if value in {"off", "direct", "one_hop"} else "one_hop"


def _effect_index_for_context(ctx: dict[str, Any], cfg):
    source_file = utils_module._safe_strip(ctx.get("source_file") or (ctx.get("file_context") or {}).get("source_file"))
    project_root = utils_module._safe_strip(getattr(cfg, "project_root", ""))
    if not project_root:
        project_root = os.path.dirname(source_file) if source_file else ""
    if not project_root:
        return effects_utils.EffectIndex()
    try:
        project_root = os.path.abspath(project_root)
    except Exception:
        pass
    cached_root = utils_module._safe_strip(getattr(cfg, "_effect_index_root", ""))
    cached = getattr(cfg, "_effect_index", None)
    if cached is not None and cached_root == project_root:
        return cached
    index = effects_utils.build_effect_index(project_root, cfg)
    try:
        cfg._effect_index = index
        cfg._effect_index_root = project_root
    except Exception:
        pass
    return index


def build_function_effect_facts(ctx: dict[str, Any], func_data: dict[str, Any], cfg, name_map: dict[str, str]):
    """Build direct and one-hop effects without allowing AI to invent facts."""
    if _effect_analysis_mode(cfg) == "off":
        return (), (), ()
    direct, returns = effects_utils.extract_direct_effects(
        func_data,
        params=ctx.get("params") or (),
        local_vars=ctx.get("local_vars") or (),
        fact_pack=dict(ctx.get("lsp_fact_pack") or {}),
        name_map=name_map,
    )
    normalized_returns = []
    for item in returns:
        condition = utils_module._safe_strip(item.condition)
        if condition:
            rendered = logic_utils._render_supported_c_expr_cn(condition, name_map)
            condition = rendered or "条件分支"
        normalized_returns.append(replace(item, condition=condition))
    inherited: tuple = ()
    issues: tuple[dict[str, Any], ...] = ()
    if _effect_analysis_mode(cfg) == "one_hop":
        inherited, issues = effects_utils.resolve_one_hop_effects(
            dict(ctx.get("lsp_fact_pack") or {}),
            index=_effect_index_for_context(ctx, cfg),
            source_file=utils_module._safe_strip(ctx.get("source_file")),
            source_function=utils_module._safe_strip(ctx.get("owner_func")),
            name_map=name_map,
        )
    return tuple(direct) + tuple(inherited), tuple(normalized_returns), tuple(issues)


def _classify_state_update(lhs: str, rhs: str) -> str:
    registry_kind = semantic_registry.classify_state_update(lhs, rhs)
    if registry_kind:
        return registry_kind
    lhs_s = str(lhs or "").strip()
    rhs_s = str(rhs or "").strip()
    rhs_compact = re.sub(r"\s+", "", rhs_s)
    if rhs_compact in {"0", "0U", "0UL", "0.0F", "0.0f"}:
        return "reset_or_clear"
    lhs_lower = lhs_s.lower()
    rhs_lower = rhs_s.lower()
    if re.search(r"\bl_data\d{3}_", lhs_lower) and ".bit_" in lhs_lower:
        return "pack_buffer_fill"
    if (".filtout_" in lhs_lower or ".filtout_" in rhs_lower) and re.search(r"\bl_data\d{3}_", lhs_lower + " " + rhs_lower):
        return "result_surface_write"
    if any(token in lhs_s for token in ("TxPack", "PackDat", "maint422TxPack", "toFpga")):
        return "pack_output"
    if any(token in rhs_s for token in ("DataTrans", "PackUp", "*", "/", "+", "-")):
        return "feedback_compute"
    if "(" in rhs_s and ")" in rhs_s:
        return "control_compute"
    if "." in lhs_s or "->" in lhs_s:
        return "state_sync"
    return "control_compute"


def _classify_call_role(callee: str, definition_comment: str = "") -> str:
    return semantic_registry.classify_call_role(callee, definition_comment)


def _is_placeholder_definition_comment(text: str, *, backend_module=None) -> bool:
    backend = backend_module or legacy_backend()
    value = utils_module._safe_strip(text)
    if not value:
        return True
    lower = value.lower()
    return (
        value.startswith("→")
        or lower.startswith("provided by ")
        or lower.startswith('provided by "')
        or backend._is_noop_comment(value)
        or backend._looks_like_logic_noise_comment(value)
    )


def _infer_entity_class(name: str, decl_type: str = "", *, backend_module=None) -> tuple[str, str]:
    backend = backend_module or legacy_backend()
    ident = utils_module._safe_strip(name)
    ident_lower = ident.lower()
    decl_lower = utils_module._safe_strip(decl_type).lower()
    registry_meta = semantic_registry.classify_entity(ident, decl_lower)
    if registry_meta:
        return utils_module._safe_strip(registry_meta.get("class")), utils_module._safe_strip(registry_meta.get("label"))
    if (re.search(r"(?:pmfl|data)(\d{3})", ident_lower) or re.search(r"(?:pmfl|data)(\d{3})", decl_lower)) and any(
        token in (ident_lower + " " + decl_lower) for token in ("1553b", "pmfl", "revdef")
    ):
        match = re.search(r"(?:pmfl|data)(\d{3})", ident_lower) or re.search(r"(?:pmfl|data)(\d{3})", decl_lower)
        label = f"{match.group(1)}字打包缓存" if match else "打包缓存"
        return "pack_buffer", label
    if "mode" in ident_lower and "compat" in ident_lower:
        return "mode_word", "模式源字"
    if "compat" in ident_lower:
        if ("act" in ident_lower) and ("flt" in decl_lower):
            return "compat_word", "作动器故障兼容字"
        return "compat_word", "兼容字"
    if ident_lower.endswith("srcerr_u16"):
        return "error_flag", "源有效性错误标志"
    if ident_lower.endswith("modeerr_u16"):
        return "error_flag", "模式错误标志"
    if any(token in ident_lower for token in ("ratio", "scale")):
        return "convert_ratio", "换算系数"
    if "gain" in ident_lower:
        return "convert_ratio", "增益系数"
    if ident_lower.startswith("l_s_") or ident_lower.startswith("s_last") or any(token in ident_lower for token in ("last", "prev", "snapshot")):
        return "snapshot_value", "状态快照"
    return "", ""


def _build_entity_classes(fact_pack: dict[str, Any], *, backend_module=None) -> dict[str, dict[str, Any]]:
    backend = backend_module or legacy_backend()
    entity_classes: dict[str, dict[str, Any]] = {}

    def _remember(name: str, decl_type: str = "", *, source: str = "heuristic", confidence: float = 0.78, verified: bool = False) -> None:
        ident = utils_module._safe_strip(name)
        if not ident or ident in entity_classes:
            return
        entity_class, label = _infer_entity_class(ident, decl_type, backend_module=backend)
        if not entity_class:
            return
        entity_classes[ident] = {
            "class": entity_class,
            "label": label,
            "source": source,
            "confidence": confidence,
            "verified": verified,
        }

    for item in (fact_pack.get("locals") or []):
        if not isinstance(item, dict):
            continue
        _remember(
            utils_module._safe_strip(item.get("name")),
            utils_module._safe_strip(item.get("decl_type")),
            source=utils_module._safe_strip(item.get("source") or "hover"),
            confidence=float(item.get("confidence", 0.0) or 0.0),
            verified=bool(item.get("verified")),
        )
    for item in (fact_pack.get("writes") or []):
        if not isinstance(item, dict):
            continue
        lhs = utils_module._safe_strip(item.get("lhs"))
        subject = lhs
        if "." in subject:
            subject = subject.split(".", 1)[0]
        elif "->" in subject:
            subject = subject.split("->", 1)[0]
        if "[" in subject:
            subject = _root_identifier_from_expr(subject) or subject
        _remember(subject, "", source="references", confidence=float(item.get("confidence", 0.0) or 0.0), verified=bool(item.get("verified")))
    return entity_classes


def _build_pattern_hits(
    fact_pack: dict[str, Any],
    entity_classes: dict[str, dict[str, Any]],
    *,
    body: str = "",
    backend_module=None,
) -> list[dict[str, Any]]:
    backend = backend_module or legacy_backend()
    hits: list[dict[str, Any]] = []
    seen: set[tuple[str, str, int]] = set()

    def _normalize_subject(expr: str) -> str:
        value = utils_module._safe_strip(expr)
        if not value:
            return ""
        if "." in value:
            value = value.split(".", 1)[0]
        elif "->" in value:
            value = value.split("->", 1)[0]
        value = utils_module._safe_strip(value)
        match = re.search(r"([A-Za-z_]\w*)\s*(?:\[[^\]]*\])?$", value)
        return utils_module._safe_strip(match.group(1)) if match else value

    def _line_of(item: dict[str, Any]) -> int:
        rng = dict(item.get("range") or {})
        try:
            return int(rng.get("start_line") or 0)
        except Exception:
            return 0

    def _body_has_snapshot_compare(subject: str) -> bool:
        ident = utils_module._safe_strip(subject)
        if not ident or not body:
            return False
        compact = str(body)
        pattern = re.compile(rf"\bif\s*\([^)]*\b{re.escape(ident)}\b[^)]*(?:!=|==)[^)]*\)")
        if pattern.search(compact):
            return True
        compare_pattern = re.compile(rf"\b{re.escape(ident)}\b\s*(?:!=|==)")
        return bool(compare_pattern.search(compact))

    for item in (fact_pack.get("writes") or []):
        if not isinstance(item, dict):
            continue
        lhs = utils_module._safe_strip(item.get("lhs"))
        rhs = utils_module._safe_strip(item.get("rhs"))
        if not lhs:
            continue
        subject = _normalize_subject(lhs)
        entity_meta = dict(entity_classes.get(subject) or {})
        entity_class = utils_module._safe_strip(entity_meta.get("class"))
        label = utils_module._safe_strip(entity_meta.get("label"))
        pattern = ""
        category = ""
        object_text = ""
        if entity_class == "pack_buffer" and ".bit_" in lhs:
            pattern = "pack_buffer_fill"
            category = "pack_output"
            object_text = lhs
            label = label.replace("打包缓存", "故障输出数据") if label.endswith("打包缓存") else (label or "输出数据")
        elif entity_class == "compat_word" and ".bit_" in lhs:
            pattern = "compat_word_fill"
            category = "compat_word"
            object_text = lhs
            label = label or "兼容字"
        elif entity_class == "mode_word" and rhs:
            pattern = "mode_word_sync"
            category = "state_sync"
            object_text = rhs
            label = label or "模式源字"
        elif entity_class == "error_flag" and rhs:
            pattern = "error_flag_assign"
            category = "state_sync"
            object_text = rhs
            label = label or "错误标志"
        elif entity_class == "validity_flag" and rhs:
            pattern = "validity_flag_assign"
            category = "state_sync"
            object_text = rhs
            label = label or "有效性标志"
        elif entity_class == "counter_value" and rhs:
            pattern = "counter_update"
            category = "state_sync"
            object_text = rhs
            label = label or "计数值"
        elif entity_class == "convert_ratio":
            pattern = "convert_ratio"
            category = "control_compute"
            object_text = rhs
            label = label or "换算系数"
        elif entity_class == "snapshot_value" and rhs and _body_has_snapshot_compare(subject):
            pattern = "snapshot_compare"
            category = "state_sync"
            object_text = rhs
            label = label or "状态快照"
        elif (".filtout_" in lhs.lower() or ".filtout_" in rhs.lower()) and rhs:
            rhs_subject = _normalize_subject(rhs)
            rhs_meta = dict(entity_classes.get(rhs_subject) or {})
            if utils_module._safe_strip(rhs_meta.get("class")) == "pack_buffer":
                pattern = "result_surface_write"
                category = "pack_output"
                object_text = lhs
                rhs_label = utils_module._safe_strip(rhs_meta.get("label"))
                label = rhs_label.replace("打包缓存", "故障输出数据") if rhs_label.endswith("打包缓存") else (rhs_label or "输出数据")
        if not pattern:
            continue
        line_no = _line_of(item)
        key = (pattern, subject or lhs, line_no)
        if key in seen:
            continue
        seen.add(key)
        hits.append(
            {
                "pattern": pattern,
                "category": category,
                "subject": subject,
                "object": object_text,
                "label": label,
                "confidence": float(item.get("confidence", 0.0) or entity_meta.get("confidence", 0.0) or 0.0),
                "verified": bool(item.get("verified") or entity_meta.get("verified")),
                "range": dict(item.get("range") or {}),
                "evidence": {
                    "lhs": lhs,
                    "rhs": rhs,
                    "entity_class": entity_class,
                    "source": utils_module._safe_strip(item.get("source") or entity_meta.get("source") or "references"),
                },
            }
        )
    for item in (fact_pack.get("calls") or []):
        if not isinstance(item, dict):
            continue
        callee = utils_module._safe_strip(item.get("callee"))
        lower = callee.lower()
        if (not callee) or ("filt" not in lower and "filter" not in lower):
            continue
        line_no = _line_of(item)
        key = ("filter_output", callee, line_no)
        if key in seen:
            continue
        seen.add(key)
        hits.append(
            {
                "pattern": "filter_output",
                "category": "filter_output",
                "subject": callee,
                "object": "",
                "label": "数字滤波",
                "confidence": float(item.get("confidence", 0.0) or 0.0),
                "verified": bool(item.get("verified")),
                "range": dict(item.get("range") or {}),
                "evidence": {
                    "callee": callee,
                    "source": utils_module._safe_strip(item.get("source") or "callHierarchy"),
                },
            }
        )
    return hits


def build_logic_semantic_pack(ctx: dict[str, Any], *, backend_module=None) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    body = utils_module._safe_text(ctx.get("body"))
    fact_pack = dict(ctx.get("lsp_fact_pack") or {})
    base_name_map = _build_base_name_map(ctx, backend_module=backend)
    entity_aliases = dict(_build_control_expression_aliases(ctx, backend_module=backend))
    for key, value in dict(base_name_map or {}).items():
        if ("[" in str(key or "")) or ("->" in str(key or "")) or ("." in str(key or "")):
            entity_aliases.setdefault(str(key), str(value))
    def _clean_owner_type_cn(raw_type: str) -> str:
        """Strip C qualifiers/pointers from a type and translate the typedef name to CN."""
        text = utils_module._safe_strip(raw_type)
        if not text:
            return ""
        # 剥掉 C 类型修饰符：const / volatile / static / restrict / * / &
        text = re.sub(r"\b(?:const|volatile|static|restrict)\b", "", text, flags=re.IGNORECASE)
        text = text.replace("*", "").replace("&", "").strip()
        # 去掉多余空白
        text = re.sub(r"\s+", "", text)
        if not text:
            return ""
        # 对类型名查 symbol_dict / 猜测中文
        cn = utils_module._safe_strip(base_name_map.get(text)) or backend._guess_cn_from_ident(text)
        return cn if (cn and utils_module._safe_strip(cn) != text) else text

    for item in (fact_pack.get("members") or []):
        if not isinstance(item, dict):
            continue
        access_text = utils_module._safe_strip(item.get("access_text"))
        member = utils_module._safe_strip(item.get("member"))
        owner_type = utils_module._safe_strip(item.get("owner_type"))
        if not access_text:
            continue
        if access_text not in entity_aliases and member and owner_type:
            owner_cn = _clean_owner_type_cn(owner_type)
            member_cn = (
                utils_module._safe_strip(base_name_map.get(member))
                or backend._guess_cn_from_ident(member)
            )
            if owner_cn and member_cn:
                entity_aliases[access_text] = f"{owner_cn}.{member_cn}"
            elif member_cn:
                entity_aliases[access_text] = member_cn
            else:
                entity_aliases[access_text] = f"{owner_type}.{member}"
        elif access_text not in entity_aliases and member:
            entity_aliases[access_text] = (
                utils_module._safe_strip(base_name_map.get(member))
                or backend._guess_cn_from_ident(member)
            )
    entity_classes = _build_entity_classes(fact_pack, backend_module=backend)
    for key, meta in dict(entity_classes or {}).items():
        label = utils_module._safe_strip((meta or {}).get("label"))
        if key and label:
            entity_aliases.setdefault(key, label)
    control_blocks = []
    for item in (fact_pack.get("blocks") or []):
        if not isinstance(item, dict):
            continue
        control_blocks.append(
            {
                "id": utils_module._safe_strip(item.get("id")),
                "kind": utils_module._safe_strip(item.get("kind")),
                "parent": utils_module._safe_strip(item.get("parent")),
                "condition": utils_module._safe_strip(item.get("condition")),
                "range": dict(item.get("range") or {}),
                "metadata": dict(item.get("metadata") or {}),
                "source": utils_module._safe_strip(item.get("source")),
                "confidence": float(item.get("confidence", 0.0) or 0.0),
                "verified": bool(item.get("verified")),
            }
        )
    pattern_hits = _build_pattern_hits(fact_pack, entity_classes, body=body, backend_module=backend)
    pattern_by_line: dict[int, list[dict[str, Any]]] = {}
    for item in (pattern_hits or []):
        if not isinstance(item, dict):
            continue
        try:
            line_no = int(dict(item.get("range") or {}).get("start_line") or 0)
        except Exception:
            line_no = 0
        if line_no > 0:
            pattern_by_line.setdefault(line_no, []).append(item)
    function_range = dict(dict(fact_pack.get("function") or {}).get("range") or {})
    try:
        function_start_line = int(function_range.get("start_line") or 0)
    except Exception:
        function_start_line = 0
    statement_comment_hints_by_line = _collect_statement_comment_hints_by_line(
        body,
        function_start_line=function_start_line,
    )
    complete_rhs_by_line, complete_rhs_by_lhs = _collect_complete_assignment_rhs(
        body,
        function_start_line=function_start_line,
    )
    state_updates = []
    seen_state_updates: set[tuple[str, str, int]] = set()
    for item in (fact_pack.get("writes") or []):
        if not isinstance(item, dict):
            continue
        lhs = utils_module._safe_strip(item.get("lhs"))
        rhs = utils_module._safe_strip(item.get("rhs"))
        if not lhs:
            continue
        range_data = dict(item.get("range") or {})
        try:
            line_no = int(range_data.get("start_line") or 0)
        except Exception:
            line_no = 0
        complete_rhs = complete_rhs_by_line.get((lhs, line_no)) or complete_rhs_by_lhs.get(lhs, "")
        if complete_rhs and len(re.sub(r"\s+", "", complete_rhs)) > len(re.sub(r"\s+", "", rhs)):
            rhs = complete_rhs
        matched_patterns = list(pattern_by_line.get(line_no) or [])
        metadata = dict(item.get("metadata") or {})
        op = utils_module._safe_strip(metadata.get("op"))
        kind = _classify_state_update(lhs, rhs)
        if op and op != "=":
            kind = "compound_assign"
        label = ""
        if matched_patterns:
            preferred = matched_patterns[0]
            pattern = utils_module._safe_strip(preferred.get("pattern"))
            label = utils_module._safe_strip(preferred.get("label"))
            if pattern == "pack_buffer_fill":
                kind = "pack_buffer_fill"
            elif pattern == "result_surface_write":
                kind = "result_surface_write"
            elif pattern == "compat_word_fill":
                kind = "compat_word_fill"
            elif pattern == "mode_word_sync":
                kind = "mode_word_sync"
            elif pattern == "error_flag_assign":
                kind = "error_flag_assign"
            elif pattern == "validity_flag_assign":
                kind = "validity_flag_assign"
            elif pattern == "counter_update":
                kind = "counter_update"
            elif pattern == "snapshot_compare":
                kind = "snapshot_compare"
            elif pattern == "convert_ratio":
                kind = "control_compute"
        seen_state_updates.add((lhs, rhs, line_no))
        state_updates.append(
            {
                "kind": kind,
                "lhs": lhs,
                "rhs": rhs,
                "op": op,
                "label": label,
                "range": range_data,
                "source": utils_module._safe_strip(item.get("source") or "references"),
                "confidence": float(item.get("confidence", 0.0) or 0.0),
                "verified": bool(item.get("verified")),
            }
        )
    for item in _collect_local_initializer_updates(
        body,
        ctx.get("local_vars") or [],
        function_start_line=function_start_line,
    ):
        lhs = utils_module._safe_strip(item.get("lhs"))
        rhs = utils_module._safe_strip(item.get("rhs"))
        try:
            line_no = int(dict(item.get("range") or {}).get("start_line") or 0)
        except Exception:
            line_no = 0
        key = (lhs, rhs, line_no)
        if key in seen_state_updates:
            continue
        seen_state_updates.add(key)
        state_updates.append(item)
    return_actions = []
    for match in re.finditer(r"\breturn\b(?P<expr>.*?);", body, flags=re.S):
        expr = _strip_return_expr_outer_parens(utils_module._safe_strip(match.group("expr")))
        line_no = 0
        if function_start_line > 0:
            line_no = function_start_line + body[: match.start()].count("\n") + 1
        return_actions.append(
            {
                "expr": expr,
                "range": {"start_line": line_no, "end_line": line_no, "start_col": 1, "end_col": 1},
                "source": "structured",
                "confidence": 0.9,
                "verified": True,
            }
        )
    flow_actions = _collect_flow_actions(
        body,
        function_start_line=function_start_line,
    )
    call_roles = []
    seen_call_roles: set[tuple[str, int]] = set()
    callee_definition_comments = dict((ctx.get("file_context") or {}).get("func_comment_map") or {})
    for item in (fact_pack.get("calls") or []):
        if not isinstance(item, dict):
            continue
        callee = utils_module._safe_strip(item.get("callee"))
        if not callee:
            continue
        range_data = dict(item.get("range") or {})
        try:
            line_no = int(range_data.get("start_line") or 0)
        except Exception:
            line_no = 0
        seen_call_roles.add((callee, line_no))
        comment_hints = list(statement_comment_hints_by_line.get(line_no) or [])
        definition_comment = utils_module._safe_strip(item.get("definition_comment"))
        if _is_placeholder_definition_comment(definition_comment, backend_module=backend):
            definition_comment = utils_module._safe_strip(callee_definition_comments.get(callee))
        call_roles.append(
            {
                "callee": callee,
                "role": _classify_call_role(callee, definition_comment),
                "definition_comment": definition_comment,
                "range": range_data,
                "source": utils_module._safe_strip(item.get("source") or "callHierarchy"),
                "confidence": float(item.get("confidence", 0.0) or 0.0),
                "verified": bool(item.get("verified")),
                "comment_hints": comment_hints,
            }
        )
    for item in _collect_statement_macro_roles(body, function_start_line=function_start_line):
        callee = utils_module._safe_strip(item.get("callee"))
        range_data = dict(item.get("range") or {})
        try:
            line_no = int(range_data.get("start_line") or 0)
        except Exception:
            line_no = 0
        if (callee, line_no) in seen_call_roles:
            for existing in call_roles:
                existing_range = dict(existing.get("range") or {})
                try:
                    existing_line = int(existing_range.get("start_line") or 0)
                except Exception:
                    existing_line = 0
                if utils_module._safe_strip(existing.get("callee")) == callee and existing_line == line_no:
                    if item.get("args") and not existing.get("args"):
                        existing["args"] = list(item.get("args") or [])
                    break
            continue
        seen_call_roles.add((callee, line_no))
        if line_no in statement_comment_hints_by_line:
            item = dict(item)
            item["comment_hints"] = list(statement_comment_hints_by_line.get(line_no) or [])
        call_roles.append(item)
    for item in _collect_memory_call_roles(body, function_start_line=function_start_line):
        callee = utils_module._safe_strip(item.get("callee"))
        range_data = dict(item.get("range") or {})
        try:
            line_no = int(range_data.get("start_line") or 0)
        except Exception:
            line_no = 0
        if (callee, line_no) in seen_call_roles:
            for existing in call_roles:
                existing_range = dict(existing.get("range") or {})
                try:
                    existing_line = int(existing_range.get("start_line") or 0)
                except Exception:
                    existing_line = 0
                if utils_module._safe_strip(existing.get("callee")) == callee and existing_line == line_no:
                    if item.get("args") and not existing.get("args"):
                        existing["args"] = list(item.get("args") or [])
                    break
            continue
        seen_call_roles.add((callee, line_no))
        if line_no in statement_comment_hints_by_line:
            item = dict(item)
            item["comment_hints"] = list(statement_comment_hints_by_line.get(line_no) or [])
        call_roles.append(item)
    bad_static_lines: list[str] = []
    try:
        rule_cfg = clone_cfg(ctx.get("cfg")) if ctx.get("cfg") is not None else None
    except Exception:
        rule_cfg = None
    if rule_cfg is not None:
        try:
            rule_cfg.ai_assist = False
            rule_cfg.ai_one_call = False
            rule_cfg.ai_logic_policy = "hybrid"
            logic_text, _ = backend.generate_logic_from_body(body, ctx.get("local_vars") or [], rule_cfg, name_map=base_name_map)
            for line in str(logic_text or "").splitlines():
                text = utils_module._safe_strip(line)
                if any(token in text for token in ("对应项", "数据指针", "缓存值", "待人工修改", "存放l_")):
                    bad_static_lines.append(text)
                elif (text.startswith("调用函数") or (text.startswith("调用") and text.endswith("函数"))) and len(text) <= 14:
                    bad_static_lines.append(text)
                elif "调用函数围绕当前" in text:
                    bad_static_lines.append(text)
        except Exception:
            bad_static_lines = []
    provider = utils_module.cfg_get_str(ctx.get("cfg"), "semantic_provider", "structured")
    control_blocks = [
        _enrich_semantic_item_v2(
            ctx,
            item,
            name_map=base_name_map,
            entity_aliases=entity_aliases,
            provider=provider,
            backend_module=backend,
        )
        for item in control_blocks
    ]
    state_updates = [
        _enrich_semantic_item_v2(
            ctx,
            item,
            name_map=base_name_map,
            entity_aliases=entity_aliases,
            provider=provider,
            backend_module=backend,
        )
        for item in state_updates
    ]
    return_actions = [
        _enrich_semantic_item_v2(
            ctx,
            item,
            name_map=base_name_map,
            entity_aliases=entity_aliases,
            provider=provider,
            backend_module=backend,
        )
        for item in return_actions
    ]
    flow_actions = [
        _enrich_semantic_item_v2(
            ctx,
            item,
            name_map=base_name_map,
            entity_aliases=entity_aliases,
            provider=provider,
            backend_module=backend,
        )
        for item in flow_actions
    ]
    call_roles = [
        _enrich_semantic_item_v2(
            ctx,
            item,
            name_map=base_name_map,
            entity_aliases=entity_aliases,
            provider=provider,
            backend_module=backend,
        )
        for item in call_roles
    ]
    pattern_hits = [
        _enrich_semantic_item_v2(
            ctx,
            item,
            name_map=base_name_map,
            entity_aliases=entity_aliases,
            provider=provider,
            backend_module=backend,
        )
        for item in pattern_hits
    ]
    pack = {
        "semantic_pack_version": 2,
        "semantic_registry": semantic_registry.registry_snapshot(),
        "name_map": dict(base_name_map or {}),
        "entity_aliases": entity_aliases,
        "entity_classes": entity_classes,
        "pattern_hits": pattern_hits,
        "control_blocks": control_blocks,
        "state_updates": state_updates,
        "return_actions": return_actions,
        "flow_actions": flow_actions,
        "call_roles": call_roles,
        "bad_static_lines": list(dict.fromkeys([x for x in bad_static_lines if x]))[:12],
        "statement_count": _count_effective_statements(body, backend_module=backend),
        "member_access_count": len(fact_pack.get("members") or []),
        "key_call_count": len(call_roles),
        "provider": provider,
    }
    quality_summary = _semantic_quality_summary(pack)
    pack["quality_summary"] = quality_summary
    pack["resolver_stats"] = dict(quality_summary.get("resolver_stats") or {})
    return pack


def run_one_call_design_enrichment(
    ctx: dict[str, Any],
    func_data: dict,
    cfg,
    index: int,
    *,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    comment_info = ctx["comment_info"]
    func_info = ctx["func_info"]
    body = ctx["body"]
    file_context = ctx["file_context"]
    local_vars = ctx["local_vars"]
    params = ctx["params"]
    in_map = ctx["in_map"]
    out_map = ctx["out_map"]
    param_ai_name_map = ctx["param_ai_name_map"]
    global_symbol_map = ctx["global_symbol_map"]
    initial_gaps = ctx["initial_gaps"]

    state: dict[str, Any] = {
        "func_data": func_data,
        "one_call_bundle": None,
        "one_call_guard_fallback": False,
        "one_call_locals_expected": 0,
        "one_call_locals_got": 0,
        "one_call_params_expected": 0,
        "one_call_params_got": 0,
        "one_call_logic_expected": 0,
        "one_call_logic_got": 0,
        "one_call_logic_degraded": False,
    }
    if not (cfg.ai_assist and getattr(cfg, "ai_one_call", False)):
        return state

    glossary = file_context.get("glossary") or backend.build_project_glossary(
        file_symbols=file_context.get("file_symbol_map")
    )
    gaps = initial_gaps
    need_bundle = False
    unknowns = []

    def _fallback_from_one_call_guard(reason: str) -> None:
        nonlocal local_vars
        nonlocal comment_info
        nonlocal param_ai_name_map
        nonlocal unknowns
        nonlocal need_bundle

        # 用户已取消时不触发回退的多步 AI 调用
        if backend.stop_requested(cfg) or getattr(cfg, "_user_cancelled", False):
            state["one_call_bundle"] = None
            state["one_call_guard_fallback"] = True
            need_bundle = False
            unknowns = []
            return

        backend.vlog(
            cfg,
            f"函数 {func_info.get('func_name', '') or index} 的 one-call 回包命中质量闸门，"
            f"已自动回退多次 AI 调用：{reason}",
        )
        cfg_one_call_off = clone_cfg(cfg, ai_one_call=False)
        # 将取消标记传播到克隆的 cfg
        if getattr(cfg, "_user_cancelled", False):
            try:
                cfg_one_call_off._user_cancelled = True
            except Exception:
                pass
        func_data2 = dict(func_data)
        func_data2["comment_info"] = comment_info
        func_data2["func_info"] = func_info
        func_data2["body"] = body
        local_vars, comment_info, param_ai_name_map = backend.enrich_with_ai(func_data2, cfg_one_call_off)
        state["one_call_bundle"] = None
        state["one_call_logic_expected"] = 0
        state["one_call_locals_got"] = 0
        state["one_call_params_got"] = 0
        state["one_call_logic_got"] = 0
        state["one_call_logic_degraded"] = False
        state["one_call_guard_fallback"] = True
        need_bundle = False
        unknowns = []

    def _normalize_name_list(items) -> list[str]:
        out: list[str] = []
        for it in (items or []):
            if isinstance(it, str):
                name = it.strip()
            elif isinstance(it, dict):
                name = str(it.get("name", "")).strip()
            else:
                name = str(it).strip()
            if name:
                out.append(name)
        return out

    missing_locals = _normalize_name_list(gaps.get("need_local_usages") or [])
    locked_usage_refine = [
        utils_module._safe_strip(v.get("name"))
        for v in (local_vars or [])
        if utils_module._safe_strip(v.get("name")) and utils_module._safe_strip(v.get("cn_name")) and backend._needs_ai_local_usage_refine(v)
    ]
    missing_locals = list(dict.fromkeys(missing_locals + locked_usage_refine))
    missing_params = _normalize_name_list(gaps.get("need_param_names") or [])
    state["one_call_locals_expected"] = len(missing_locals)
    state["one_call_params_expected"] = len(missing_params)

    typedef_blocks = file_context.get("typedefs") or []
    member_symbol_map = backend._extract_member_symbol_map_from_typedefs(typedef_blocks)
    header_member_symbol_map = dict(file_context.get("member_symbol_map") or {})
    scope_symbol_map = backend._build_local_param_symbol_map(local_vars, params, in_map, out_map, param_ai_name_map)
    name_map = dict(global_symbol_map)
    name_map.update(member_symbol_map)
    name_map.update(header_member_symbol_map)
    name_map.update(scope_symbol_map)

    cfg_logic = clone_cfg(cfg, ai_logic_policy="ai_non_structured")
    logic, unknowns = backend.generate_logic_from_body(body, local_vars, cfg_logic, name_map=name_map)
    logic_semantic_pack = dict(ctx.get("logic_semantic_pack") or {})
    if (not unknowns) and (
        utils_module.cfg_get_int(cfg, "logic_use_lsp", 1)
        and (
            logic_semantic_pack.get("control_blocks")
            or logic_semantic_pack.get("state_updates")
            or logic_semantic_pack.get("call_roles")
            or logic_semantic_pack.get("pattern_hits")
        )
    ):
        logic, unknowns = backend.generate_logic_from_semantic_pack(
            logic_semantic_pack,
            cfg_logic,
            name_map=name_map,
        )
    if (not unknowns) and logic and bool(gaps.get("need_func_semantic_rewrite") or getattr(cfg, "ai_mode", 1) == 2):
        unknowns = backend.select_ai_logic_polish_unknowns(
            logic,
            max_items=utils_module.cfg_get_int(cfg, "ai_logic_polish_max_items", 12),
        )
    state["one_call_logic_expected"] = len(unknowns or [])

    if backend._should_disable_local_one_call(body, unknowns, cfg):
        # 用户已取消时不触发回退的多步 AI 调用
        if backend.stop_requested(cfg) or getattr(cfg, "_user_cancelled", False):
            state["one_call_logic_expected"] = 0
            unknowns = []
            state["one_call_bundle"] = None
        else:
            backend.vlog(
                cfg,
                f"函数 {func_info.get('func_name', '') or index} 上下文过大，"
                "已自动禁用 one-call 并回退多次 AI 调用",
            )
            cfg_one_call_off = clone_cfg(cfg, ai_one_call=False)
            # 将取消标记传播到克隆的 cfg
            if getattr(cfg, "_user_cancelled", False):
                try:
                    cfg_one_call_off._user_cancelled = True
                except Exception:
                    pass
            func_data2 = dict(func_data)
            func_data2["comment_info"] = comment_info
            func_data2["func_info"] = func_info
            func_data2["body"] = body
            local_vars, comment_info, param_ai_name_map = backend.enrich_with_ai(func_data2, cfg_one_call_off)
        state["one_call_logic_expected"] = 0
        unknowns = []
        state["one_call_bundle"] = None
    else:
        need_bundle = bool(
            gaps.get("need_func_desc")
            or gaps.get("need_func_cn_name")
            or gaps.get("need_func_cn_refine")
            or gaps.get("need_func_semantic_rewrite")
            or missing_locals
            or missing_params
            or unknowns
        )

        if need_bundle:
            locked_local_names = {
                utils_module._safe_strip(v.get("name"))
                for v in (local_vars or [])
                if utils_module._safe_strip(v.get("name")) in set(missing_locals or []) and utils_module._safe_strip(v.get("cn_name"))
            }
            state["one_call_bundle"] = backend.ai_suggest_bundle_one_call(
                func_info=func_info,
                body=body,
                comment_info=comment_info if cfg.ai_mode != 2 else {},
                local_vars=local_vars,
                params=params,
                missing_locals=missing_locals,
                missing_params=missing_params,
                unknowns=unknowns,
                cfg=cfg,
                file_context=file_context,
                evidence_pack=backend.build_llm_evidence_pack(func_data, cfg, task="logic"),
            )
            # 检查停止信号
            if backend.stop_requested(cfg):
                return state
            backend.ai_debug_log(cfg, "one_call_request_summary", {
                "func_name": func_info.get("func_name", ""),
                "missing_locals": len(missing_locals),
                "missing_params": len(missing_params),
                "unknowns": len(unknowns),
                "body_len": len(body or ""),
                "ai_mode": getattr(cfg, "ai_mode", None),
                "provider": getattr(cfg, "ai_provider", ""),
                "model": getattr(cfg, "ai_model", ""),
            })
            bundle_guard = (state["one_call_bundle"] or {}).get("_guard") or {}
            if isinstance(bundle_guard, dict) and bool(bundle_guard.get("force_fallback")):
                guard_reasons = bundle_guard.get("reasons") or ()
                _fallback_from_one_call_guard(",".join([str(x) for x in guard_reasons if str(x)]))
            elif isinstance(state["one_call_bundle"], dict):
                raw_locals_part = state["one_call_bundle"].get("locals") or {}
                if isinstance(raw_locals_part, dict):
                    effective_local_hits = 0
                    for key, value in raw_locals_part.items():
                        local_name = utils_module._safe_strip(key)
                        if not local_name:
                            continue
                        if local_name in locked_local_names:
                            if isinstance(value, dict) and utils_module._safe_strip(value.get("usage") or value.get("desc") or value.get("purpose")):
                                effective_local_hits += 1
                            continue
                        if isinstance(value, dict) and utils_module._safe_strip(value.get("cn_name") or value.get("name")):
                            effective_local_hits += 1
                    state["one_call_locals_got"] = effective_local_hits
                else:
                    state["one_call_locals_got"] = 0
                state["one_call_params_got"] = len(state["one_call_bundle"].get("params") or {}) if isinstance(state["one_call_bundle"].get("params"), dict) else 0
                state["one_call_logic_got"] = len(state["one_call_bundle"].get("logic") or {}) if isinstance(state["one_call_bundle"].get("logic"), dict) else 0
                func_part = state["one_call_bundle"].get("func") or {}
                logic_part = state["one_call_bundle"].get("logic") or {}
                has_func_payload = isinstance(func_part, dict) and bool(
                    utils_module._safe_strip(func_part.get("func_cn_name")) or utils_module._safe_strip(func_part.get("desc"))
                )
                has_logic_payload = isinstance(logic_part, dict) and bool(logic_part)
                if state["one_call_locals_expected"] > 0 and state["one_call_locals_got"] < state["one_call_locals_expected"] and (not has_func_payload) and (not has_logic_payload):
                    _fallback_from_one_call_guard("locals_partial_or_empty")
                elif (
                    state["one_call_locals_expected"] > 0
                    and state["one_call_locals_got"] <= 0
                    and (not has_logic_payload)
                    and (has_func_payload or state["one_call_params_got"] > 0)
                ):
                    _fallback_from_one_call_guard("locals_empty_params_only")

    if isinstance(state["one_call_bundle"], dict):
        state["one_call_locals_got"] = len(state["one_call_bundle"].get("locals") or {}) if isinstance(state["one_call_bundle"].get("locals"), dict) else 0
        state["one_call_params_got"] = len(state["one_call_bundle"].get("params") or {}) if isinstance(state["one_call_bundle"].get("params"), dict) else 0
        state["one_call_logic_got"] = len(state["one_call_bundle"].get("logic") or {}) if isinstance(state["one_call_bundle"].get("logic"), dict) else 0
        func_part = state["one_call_bundle"].get("func") or {}
        if isinstance(func_part, dict):
            func_cn_name = utils_module._safe_strip(func_part.get("func_cn_name"))
            desc = utils_module._safe_strip(func_part.get("desc"))
            one_call_func_conf = float(func_part.get("confidence", utils_module.cfg_get_float(cfg, "symbol_memory_one_call_conf", 0.82)) or 0.0)
            if func_cn_name and not text_utils._contains_cjk(func_cn_name):
                guessed = backend._guess_cn_from_ident(func_cn_name, glossary=glossary)
                if guessed:
                    func_cn_name = guessed
            if not func_cn_name:
                guessed = backend._guess_cn_from_ident(func_info.get("func_name", ""), glossary=glossary)
                if guessed:
                    func_cn_name = guessed
            if func_cn_name:
                current_title = utils_module._safe_strip(comment_info.get("func_cn_name"))
                retrieved_examples = backend.retrieve_function_title_context(
                    {
                        "comment_info": dict(comment_info or {}),
                        "func_info": dict(func_info or {}),
                        "file_context": dict(file_context or {}),
                        "body": body,
                    },
                    cfg,
                )
                if (
                    backend._is_missing_gap_text(current_title)
                    or backend._should_accept_refined_function_title(
                        current_title,
                        func_cn_name,
                        func_name=utils_module._safe_strip(func_info.get("func_name")),
                        comment_desc=utils_module._safe_strip(comment_info.get("desc")),
                        examples=retrieved_examples,
                    )
                ):
                    comment_info["func_cn_name"] = func_cn_name
                    backend._remember_ai_symbol(
                        str(func_info.get("func_name") or ""),
                        func_cn_name,
                        kind="functions",
                        confidence=one_call_func_conf,
                        cfg=cfg,
                        source="ai_one_call_func",
                    )
            if desc:
                comment_info["desc"] = _repair_function_desc_by_domain(
                    utils_module._safe_strip(func_info.get("func_name")),
                    desc,
                    current_desc=utils_module._safe_strip(comment_info.get("desc")),
                )

        locals_part = state["one_call_bundle"].get("locals") or {}
        if isinstance(locals_part, dict):
            for v in local_vars:
                name = v.get("name")
                if not name or name not in locals_part:
                    continue
                item = locals_part.get(name) or {}
                if not isinstance(item, dict):
                    continue
                locked_cn = utils_module._safe_strip(v.get("cn_name"))
                cn = utils_module._safe_strip(item.get("cn_name") or item.get("name"))
                usage = backend._sanitize_ai_usage_text(item.get("usage") or item.get("purpose") or item.get("desc"))
                if cn and backend._is_strict_symbol_candidate_rejected(cn, raw_ident=name):
                    cn = ""
                if cn and (not locked_cn):
                    v["cn_name"] = cn
                    backend._remember_ai_symbol(
                        str(name),
                        cn,
                        kind="symbols",
                        confidence=utils_module.cfg_get_float(cfg, "symbol_memory_one_call_conf", 0.82),
                        cfg=cfg,
                        source="ai_one_call_symbol",
                    )
                if usage and backend._should_replace_local_usage_with_ai(v.get("usage"), v.get("cn_name")):
                    v["usage"] = usage

        params_part = state["one_call_bundle"].get("params") or {}
        if isinstance(params_part, dict):
            for p in params:
                name = (p or {}).get("name") or ""
                if not name or name not in params_part:
                    continue
                item = params_part.get(name) or {}
                if isinstance(item, dict):
                    cn = utils_module._safe_strip(item.get("cn_name") or item.get("name"))
                else:
                    cn = utils_module._safe_strip(item)
                if cn and not backend._is_strict_symbol_candidate_rejected(cn, raw_ident=name):
                    param_ai_name_map[name] = cn
                    backend._remember_ai_symbol(
                        str(name),
                        cn,
                        kind="symbols",
                        confidence=utils_module.cfg_get_float(cfg, "symbol_memory_one_call_conf", 0.82),
                        cfg=cfg,
                        source="ai_one_call_symbol",
                    )

        logic_part = state["one_call_bundle"].get("logic") or {}
        logic_map: dict[str, str] = {}
        if unknowns:
            if isinstance(logic_part, dict):
                if isinstance(logic_part.get("steps"), list):
                    for it in (logic_part.get("steps") or []):
                        if not isinstance(it, dict):
                            continue
                        idx_v = it.get("idx")
                        txt = utils_module._safe_strip(it.get("action") or it.get("text") or it.get("desc"))
                        txt = logic_utils._sanitize_ai_logic_action(txt)
                        if idx_v is None or not txt:
                            continue
                        logic_map[str(idx_v)] = txt
                else:
                    for key, value in logic_part.items():
                        if isinstance(value, str) and value.strip():
                            logic_map[str(key)] = logic_utils._sanitize_ai_logic_action(utils_module._safe_strip(value))
                        elif isinstance(value, dict):
                            txt = utils_module._safe_strip(value.get("action") or value.get("text") or value.get("desc"))
                            if txt:
                                logic_map[str(key)] = logic_utils._sanitize_ai_logic_action(txt)
            elif isinstance(logic_part, list):
                for it in logic_part:
                    if not isinstance(it, dict):
                        continue
                    idx_v = it.get("idx")
                    txt = utils_module._safe_strip(it.get("action") or it.get("text") or it.get("desc"))
                    txt = logic_utils._sanitize_ai_logic_action(txt)
                    if idx_v is None or not txt:
                        continue
                    logic_map[str(idx_v)] = txt

        if unknowns and logic_map:
            allowed_logic_keys = [str(u.get("idx")) for u in (unknowns or []) if u.get("idx") is not None]
            logic_map, logic_guard = backend._sanitize_one_call_section_map(
                logic_map,
                allowed_keys=allowed_logic_keys,
            )
            if logic_guard.get("dropped_keys"):
                backend.vlog(
                    cfg,
                    f"函数 {func_info.get('func_name', '') or index} 的 one-call logic 已过滤无效 idx: "
                    f"{logic_guard.get('dropped_keys')}",
                )
                backend.ai_debug_log(cfg, "one_call_logic_strict_guard", {
                    "func_name": func_info.get("func_name", ""),
                    "allowed_keys": logic_guard.get("allowed_keys"),
                    "dropped_keys": logic_guard.get("dropped_keys"),
                })
            state["one_call_logic_got"] = len(logic_map or {})

        if logic_map and unknowns:
            lines = logic.splitlines()
            fallback_count = 0
            filled_count = 0
            for u in unknowns:
                idx2 = u.get("idx")
                if idx2 is None:
                    continue
                new_text = utils_module._safe_strip(logic_map.get(str(idx2)))
                if new_text and logic_utils._is_control_logic_line(new_text):
                    new_text = ""
                if not new_text:
                    if bool(u.get("polish_only")):
                        continue
                    guess = backend.heuristic_logic_line(
                        u.get("code") or u.get("code_cn") or "",
                        name_map=name_map,
                        literal=(backend._get_logic_comment_mode(cfg) == "off"),
                    )
                    if not guess:
                        guess = backend.fallback_logic_line(u.get("code") or "", name_map=name_map)
                    new_text = utils_module._safe_strip(guess)
                    if not new_text:
                        continue
                    fallback_count += 1
                else:
                    filled_count += 1
                lines[int(idx2)] = u.get("indent", "") + new_text
            logic = "\n".join(lines)
            backend.ai_debug_log(cfg, "one_call_logic_fill_stats", {
                "func_name": func_info.get("func_name", ""),
                "unknowns": len(unknowns),
                "ai_filled": filled_count,
                "fallback_filled": fallback_count,
            })

        if should_degrade_one_call_logic(body, unknowns, logic_map, logic, cfg, backend_module=backend):
            degraded_logic, degraded_unknowns = build_degraded_one_call_logic(
                body,
                local_vars,
                name_map,
                cfg,
                backend_module=backend,
            )
            degraded_placeholders = logic_utils._count_logic_placeholder_lines(degraded_logic, backend_module=backend)
            current_placeholders = logic_utils._count_logic_placeholder_lines(logic, backend_module=backend)
            if degraded_logic and degraded_placeholders < current_placeholders:
                logic = degraded_logic
                state["one_call_logic_got"] = max(state["one_call_logic_expected"], len(logic_map or {}))
                state["one_call_logic_degraded"] = True
                backend.vlog(
                    cfg,
                    f"函数 {func_info.get('func_name', '') or index} 的 one-call logic 返回偏空，"
                    f"已自动降级为规则分段生成（unknowns={len(unknowns or [])}, "
                    f"degraded_unknowns={len(degraded_unknowns or [])}, "
                    f"placeholders {current_placeholders}->{degraded_placeholders}）",
                )

    if need_bundle:
        state["func_data"] = dict(func_data)
        state["func_data"]["_one_call_logic"] = logic
        state["func_data"]["_one_call_unknowns_used"] = True
        state["func_data"]["_one_call_logic_degraded"] = bool(state["one_call_logic_degraded"])

    ctx["comment_info"] = comment_info
    ctx["func_info"] = func_info
    ctx["body"] = body
    ctx["file_context"] = file_context
    ctx["local_vars"] = local_vars
    ctx["params"] = params
    ctx["in_map"] = in_map
    ctx["out_map"] = out_map
    ctx["param_ai_name_map"] = param_ai_name_map
    return state


def run_multi_call_design_enrichment(
    ctx: dict[str, Any],
    func_data: dict,
    cfg,
    *,
    backend_module=None,
) -> None:
    backend = backend_module or legacy_backend()
    if not (cfg.ai_assist and (not getattr(cfg, "ai_one_call", False))):
        return
    # 用户已取消时跳过多步 AI 调用
    if backend.stop_requested(cfg) or getattr(cfg, "_user_cancelled", False):
        return

    comment_info = ctx["comment_info"]
    func_info = ctx["func_info"]
    body = ctx["body"]
    local_vars = ctx["local_vars"]
    params = ctx["params"]
    in_map = ctx["in_map"]
    out_map = ctx["out_map"]
    scope_inference_log = ctx["scope_inference_log"]

    func_data2 = dict(func_data)
    func_data2["comment_info"] = comment_info
    func_data2["func_info"] = func_info
    func_data2["body"] = body
    local_vars, comment_info, param_ai_name_map = backend.enrich_with_ai(func_data2, cfg)
    scope_inference_log.update(
        backend.infer_scope_symbol_names(
            local_vars,
            params,
            body=body,
            func_info=func_info,
            comment_info=comment_info,
            in_map=in_map,
            out_map=out_map,
            cfg=cfg,
        )
    )

    ctx["comment_info"] = comment_info
    ctx["func_info"] = func_info
    ctx["body"] = body
    ctx["local_vars"] = local_vars
    ctx["params"] = params
    ctx["in_map"] = in_map
    ctx["out_map"] = out_map
    ctx["param_ai_name_map"] = param_ai_name_map
    ctx["scope_inference_log"] = scope_inference_log


def build_design_logic_lines(
    ctx: dict[str, Any],
    func_data: dict,
    cfg,
    name_map: dict[str, str],
    *,
    backend_module=None,
    cached_logic_lines: Optional[dict[int, str]] = None,
    changed_statement_lines: Optional[set[int]] = None,
):
    """
    构建逻辑语句列表，支持语句级增量合并。

    Args:
        ctx: 上下文
        func_data: 函数数据
        cfg: 配置
        name_map: 名称映射
        backend_module: 后端模块
        cached_logic_lines: 缓存的逻辑语句 {源代码行号: 逻辑描述}
        changed_statement_lines: 变更的源代码行号集合
    """
    backend = backend_module or legacy_backend()
    if not cfg.include_logic:
        return None

    body = ctx["body"]
    local_vars = ctx["local_vars"]
    logic_semantic_pack = dict(ctx.get("logic_semantic_pack") or {})
    use_lsp_logic = bool(utils_module.cfg_get_int(cfg, "logic_use_lsp", 1))
    line_start = int(func_data.get("line_start", 0) or func_data.get("func_info", {}).get("line_start", 0) or 0)

    # 语句级增量：如果提供了缓存且变更行数较少，尝试合并
    use_statement_merge = (
        cached_logic_lines is not None
        and changed_statement_lines is not None
        and len(changed_statement_lines) < len(cached_logic_lines) * 0.5  # 变更少于50%才合并
    )

    def _apply_ai_logic_replacements(base_logic: str, items: Sequence[dict]) -> str:
        if not (cfg.ai_assist and items and base_logic):
            return base_logic
        if getattr(cfg, "ai_circuit_break", False) or getattr(cfg, "_skip_ai_current_func", False):
            return base_logic
        repl = backend.ai_refine_logic_unknowns(items, body, cfg)
        lines = base_logic.splitlines()
        repl = repl if isinstance(repl, dict) else {}

        def _has_bad_logic_residual_ident(text: str) -> bool:
            value = utils_module._safe_strip(text)
            if not value:
                return False
            if value.startswith("#") or any(ch in value for ch in "{}"):
                return True
            return bool(
                re.search(
                    r"\b(?:[lgsvcp]_|[A-Za-z_]\w*_(?:u|i|f)(?:8|16|32|64)?\b|[A-Z][A-Z0-9_]{2,})",
                    value,
                )
            )

        for item in items:
            idx2 = item.get("idx")
            if idx2 is None:
                continue
            try:
                line_idx = int(idx2)
            except Exception:
                continue
            if line_idx < 0 or line_idx >= len(lines):
                continue
            key = str(idx2)
            new_text = utils_module._safe_strip(repl.get(key))
            new_text = logic_utils._sanitize_control_logic_line(new_text, backend_module=backend) or new_text
            if new_text and logic_utils._is_control_logic_line(new_text):
                new_text = ""
            if new_text and _has_bad_logic_residual_ident(new_text):
                new_text = ""
            if new_text:
                lines[line_idx] = item.get("indent", "") + new_text
                continue
            if bool(item.get("polish_only")):
                continue
            fb = utils_module._safe_strip(item.get("fallback_text"))
            if fb:
                lines[line_idx] = item.get("indent", "") + fb
            else:
                lines[line_idx] = item.get("indent", "") + backend.fallback_logic_line(item.get("code") or "", name_map=name_map)
        return "\n".join(lines)

    if cfg.ai_assist and getattr(cfg, "ai_one_call", False) and func_data.get("_one_call_unknowns_used"):
        ai_logic = func_data.get("_one_call_logic") or ""
        # Over-compression guard: if AI compressed too much vs static, fall back
        ai_lines = [ln for ln in ai_logic.splitlines() if ln.strip()]
        body_line_count = len([ln for ln in body.splitlines() if ln.strip()])
        if ai_lines and body_line_count > 40:
            try:
                static_logic, static_unknowns = backend.generate_logic_from_body(
                    body, local_vars, cfg, name_map=name_map
                )
                static_lines = [ln for ln in static_logic.splitlines() if ln.strip()]
                # If AI compressed to <40% of static lines, fall back to static + AI refinement
                if static_lines and len(ai_lines) < len(static_lines) * 0.4:
                    logic = static_logic
                    unknowns = static_unknowns
                else:
                    logic = ai_logic
                    unknowns = []
            except Exception:
                logic = ai_logic
                unknowns = []
        else:
            logic = ai_logic
            unknowns = []
    else:
        logic = ""
        unknowns = []
        # Opt-in primary path: LogicStep IR → deterministic Chinese lines.
        if logic_step_ir_primary(cfg):
            try:
                from .logic_step_ir import build_logic_steps, render_logic_steps_to_lines

                steps = build_logic_steps(body, local_vars, cfg, name_map=name_map, backend_module=backend)
                step_lines = render_logic_steps_to_lines(steps, name_map=name_map, backend_module=backend)
                if step_lines:
                    logic = "\n".join(step_lines)
                    unknowns = []
            except Exception as exc:
                backend.vlog(cfg, f"[LogicStep] primary 渲染失败，回退规则链：{exc}")
                logic = ""
        if not logic:
            # 逻辑渲染始终走规则解析：逐语句完整（含 ++/-- / compound assign），
            # LSP/clangd 的结构体成员 owner 翻译通过 build_design_name_map 注入 name_map 生效，
            # 不再用 generate_logic_from_semantic_pack 整体覆盖（会漏 LSP 未归类的语句）。
            logic, unknowns = backend.generate_logic_from_body(body, local_vars, cfg, name_map=name_map)
            if cfg.ai_assist and unknowns:
                logic = _apply_ai_logic_replacements(logic, unknowns)
            if cfg.ai_assist and logic and not getattr(cfg, "ai_one_call", False):
                polish_unknowns = backend.select_ai_logic_polish_unknowns(
                    logic,
                    max_items=utils_module.cfg_get_int(cfg, "ai_logic_polish_max_items", 12),
                )
                if polish_unknowns:
                    logic = _apply_ai_logic_replacements(logic, polish_unknowns)
    if bool(getattr(cfg, "enhanced_single_func_pseudocode", False)):
        enhanced_logic = backend._build_enhanced_single_function_logic(body, local_vars, name_map=name_map)
        if enhanced_logic:
            logic = enhanced_logic

    if not logic:
        logic_lines = ()
    else:
        out_lines = []
        for line in logic.splitlines():
            text = line.rstrip()
            text = logic_utils._sanitize_control_logic_line(text, backend_module=backend) or text
            text = logic_utils._refresh_control_logic_line_idents(text, name_map, backend_module=backend)
            if text and (not text.endswith("；")) and (not logic_utils._is_control_logic_line(text)):
                text += "；"
            out_lines.append(text)
        logic_lines = tuple(out_lines)
    if logic_lines:
        repair_symbol_map = dict(getattr(backend, "SYMBOL_DICTIONARY_RUNTIME", {}) or {})
        repair_symbol_map.update(dict(name_map or {}))
        logic_lines = tuple(backend.repair_unresolved_logic_lines(logic_lines, repair_symbol_map))
        if sum(1 for line in logic_lines if "调用函数" in line) >= 3:
            logic_lines = backend.repair_generic_logic_calls(logic_lines, body=body, name_map=name_map)
        logic_lines = tuple(logic_utils._polish_logic_lines(logic_lines, backend_module=backend))
        logic_lines = tuple(logic_utils._validate_control_blocks(logic_lines, backend_module=backend))
        logic_lines = tuple(
            logic_utils._sanitize_control_logic_line(line, backend_module=backend) or line
            for line in logic_lines
        )
        logic_lines = tuple(
            line for line in logic_lines
            if utils_module._safe_strip(line)
            and (
                not utils_module._safe_strip(line).startswith("初始化")
                and not utils_module._safe_strip(line).startswith("定义")
            )
        )
        if len(logic_lines) <= 1:
            logic_lines = backend.expand_thin_logic(logic_lines, body, name_map=name_map, cfg=cfg)

    # 语句级增量合并：将新生成的逻辑语句与缓存的未变更语句合并
    if use_statement_merge and cached_logic_lines and logic_lines:
        from .incremental import merge_logic_lines_with_cache
        merged = merge_logic_lines_with_cache(
            logic_lines,
            cached_logic_lines,
            changed_statement_lines,
            line_start,
            body,
        )
        if merged:
            logic_lines = tuple(merged)

    revision_patch = ctx.get("_revision_patch")
    if revision_patch:
        logic_lines = revision_utils.apply_revision_to_logic_lines(logic_lines, revision_patch)

    return logic_lines


def build_design_io_elements(
    ctx: dict[str, Any],
    func_data: dict,
    cfg,
    *,
    backend_module=None,
    one_call_bundle=None,
):
    backend = backend_module or legacy_backend()
    comment_info = ctx["comment_info"]
    func_info = ctx["func_info"]
    body = ctx["body"]
    params = ctx["params"]
    local_vars = ctx["local_vars"]
    in_map = ctx["in_map"]
    out_map = ctx["out_map"]
    param_ai_name_map = ctx["param_ai_name_map"]

    ret_type = utils_module._safe_strip(func_info.get("ret_type"))
    ret_var_name = backend.parse_return_var_from_body(body)
    ret_expr = _parse_return_expr_from_body(body)
    body_return_candidates = _return_simple_candidate_exprs_from_body(body)
    if len(body_return_candidates) > 1:
        ret_var_name = None
    ret_var = None
    if ret_var_name:
        for item in local_vars:
            if item.get("name") == ret_var_name:
                ret_var = item
                break

    var_cn_map = ctx.get("var_cn_map") or {}
    io_elements: list[Any] = []
    for pinfo in params:
        name = pinfo["name"]
        ptype = pinfo["type"]
        # LLM batch variable name takes priority over comment-derived names
        llm_cn = var_cn_map.get(name, "")
        if llm_cn:
            desc_cn = llm_cn
        else:
            raw_desc_cn = in_map.get(name) or out_map.get(name) or param_ai_name_map.get(name) or name
            desc_cn = backend._shorten_element_display_name(
                backend.resolve_canonical_symbol_name(name, kind="symbols", comment_cn=raw_desc_cn, fallback=name),
                fallback=name,
            )
        direction = _function_param_direction(ctx, pinfo)
        io_elements.append(IOElement(name=desc_cn, ident=name, c_type=ptype, direction=direction))

    if not _is_void_return_type(ret_type):
        ret_c_type = _normalize_return_c_type(ret_type)
        return_candidates = body_return_candidates if len(body_return_candidates) > 1 else ()
        if not return_candidates and not ret_var_name:
            return_candidates = _return_ternary_candidate_exprs(ret_expr)
        ret_ident = ret_var_name or (ret_expr if _is_simple_return_io_expr(ret_expr) else "") or "return"
        ret_lookup_name = ret_var_name or (return_candidates[0] if len(return_candidates) == 1 else "") or _root_identifier_from_expr(ret_expr) or ret_ident
        ret_display_name = None
        # Human revision return_desc has highest priority for regenerated docs.
        revision_return_desc = utils_module._safe_strip((ctx.get("comment_info") or {}).get("return_desc"))
        if ctx.get("_revision_patch") and revision_return_desc and not backend._is_noop_comment(revision_return_desc):
            ret_display_name = backend._shorten_element_display_name(revision_return_desc, fallback="返回值")
        # LLM batch variable name takes priority for return variable
        llm_ret_cn = var_cn_map.get(ret_lookup_name, "") or var_cn_map.get(ret_ident, "")
        if not ret_display_name and llm_ret_cn:
            ret_display_name = llm_ret_cn
        elif not ret_display_name and ret_var:
            ret_display_name = (
                backend.resolve_canonical_symbol_name(
                    ret_var_name,
                    kind="symbols",
                    comment_cn=ret_var.get("cn_name") or ret_var.get("usage") or "",
                    fallback="返回值",
                )
                or ret_var.get("usage")
                or "返回值"
            )
        elif cfg.ai_assist and getattr(cfg, "ai_one_call", False) and isinstance(one_call_bundle, dict):
            ret_part = one_call_bundle.get("return") or {}
            if isinstance(ret_part, dict):
                cn = utils_module._safe_strip(ret_part.get("cn_name"))
                if cn:
                    ret_display_name = cn
        elif ret_lookup_name and cfg.ai_assist:
            fake_locals = [{
                "name": ret_lookup_name,
                "type": ret_c_type,
                "usage": "",
                "cn_name": "",
                "family_prefix": ctx.get("family_prefix", ""),
                "module_key": ctx.get("module_key", ""),
                "owner_func": ctx.get("owner_func", ""),
                "source_file": ctx.get("source_file", ""),
                "owner_ret_type": ctx.get("owner_ret_type", ""),
                "scope": "local",
                "direction": "local",
            }]
            sugg_map = backend.ai_suggest_for_locals_batch(
                [ret_lookup_name],
                fake_locals,
                body,
                cfg,
                func_cn_name=comment_info.get("func_cn_name", ""),
                func_desc=comment_info.get("desc", ""),
                glossary=(func_data.get("file_context", {}) or {}).get("glossary") or backend.DOMAIN_GLOSSARY,
            )
            item = sugg_map.get(ret_lookup_name) if isinstance(sugg_map, dict) else None
            if item:
                conf = float(item.get("confidence", 0) or 0)
                if cfg.force_ai or conf >= cfg.ai_conf_symbol:
                    ret_display_name = item.get("cn_name") or "返回值"
        if not ret_display_name:
            return_desc = utils_module._safe_strip(comment_info.get("return_desc"))
            if return_desc and not backend._is_noop_comment(return_desc):
                ret_display_name = backend._shorten_element_display_name(return_desc, fallback="返回值")
        if not ret_display_name:
            if ret_var_name or _is_simple_return_io_expr(ret_expr):
                ret_display_name = _lookup_io_display_name(ctx, ret_ident, backend_module=backend, fallback="返回值")
            else:
                ret_display_name = "返回值"
        if return_candidates:
            for index, candidate in enumerate(return_candidates, start=1):
                candidate_name = _lookup_return_candidate_display_name(ctx, candidate, backend_module=backend, fallback="")
                if not candidate_name or candidate_name == candidate:
                    candidate_name = f"{ret_display_name or '返回值'}候选{index}"
                io_elements.append(IOElement(name=candidate_name, ident=candidate, c_type=ret_c_type, direction="输出"))
        else:
            io_elements.append(IOElement(name=ret_display_name or "返回值", ident=ret_ident, c_type=ret_c_type, direction="输出"))

    return tuple(io_elements), (not io_elements), ret_var_name


def _is_void_return_type(ret_type: str) -> bool:
    value = _normalize_return_c_type(ret_type)
    if not value:
        return False
    value = re.sub(r"\b(?:const|volatile|register)\b", " ", value)
    value = re.sub(r"\s+", "", value).lower()
    return value == "void"


def _normalize_return_c_type(ret_type: str) -> str:
    value = utils_module._safe_strip(ret_type)
    if not value:
        return ""
    value = re.sub(r"\b(?:static|extern|inline|__inline|__inline__)\b", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def build_design_local_elements(
    ctx: dict[str, Any],
    cfg,
    *,
    backend_module=None,
    ret_var_name: str = "",
):
    backend = backend_module or legacy_backend()
    if not cfg.include_locals:
        return None

    comment_info = ctx["comment_info"]
    body = ctx["body"]
    local_vars = ctx["local_vars"]
    locals_for_table = [v for v in local_vars if v.get("name") != ret_var_name] if ret_var_name else local_vars
    if not locals_for_table:
        return ()

    local_comment_desc = utils_module._safe_strip((comment_info or {}).get("desc"))
    name_map = build_design_name_map(ctx, backend_module=backend)
    var_cn_map = ctx.get("var_cn_map") or {}
    var_usage_map = ctx.get("var_usage_map") or {}
    local_rows = []
    for item in locals_for_table:
        vname = utils_module._safe_strip((item or {}).get("name"))
        # LLM batch variable name takes priority
        llm_cn = var_cn_map.get(vname, "")
        if llm_cn:
            cn_name = llm_cn
            usage_text = utils_module._safe_strip(var_usage_map.get(vname) or item.get("usage") or "")
            usage_text = backend._select_local_usage_text(
                {**dict(item or {}), "cn_name": cn_name, "usage": usage_text},
                body=body,
                comment_desc=local_comment_desc,
                cfg=cfg,
            )
            usage_text = backend._normalize_explanatory_text_for_output(usage_text, name_map=name_map)
            local_rows.append(
                LocalDataElement(
                    name=cn_name,
                    ident=vname,
                    c_type=item.get("type") or "",
                    usage=usage_text,
                )
            )
            continue
        neighbor_symbols = [
            utils_module._safe_strip((x or {}).get("name"))
            for x in locals_for_table
            if utils_module._safe_strip((x or {}).get("name")) and utils_module._safe_strip((x or {}).get("name")) != vname
        ]
        backend._repair_local_cn_name_with_profile(
            item,
            body=body,
            neighbor_symbols=neighbor_symbols,
            comment_desc=local_comment_desc,
            cfg=cfg,
        )
        local_item = dict(item or {})
        local_item["neighbor_symbols"] = neighbor_symbols
        cn_name = backend._select_local_display_name(local_item)
        usage_text = backend._select_local_usage_text(
            local_item,
            body=body,
            comment_desc=local_comment_desc,
            cfg=cfg,
        )
        usage_text = backend._normalize_explanatory_text_for_output(usage_text, name_map=name_map)
        local_rows.append(
            LocalDataElement(
                name=cn_name,
                ident=item.get("name") or "",
                c_type=item.get("type") or "",
                usage=usage_text,
            )
        )
    return tuple(local_rows)


def repair_design_local_profiles(
    ctx: dict[str, Any],
    cfg,
    *,
    backend_module=None,
) -> None:
    backend = backend_module or legacy_backend()
    comment_info = ctx["comment_info"]
    body = ctx["body"]
    local_vars = ctx["local_vars"]

    neighbor_symbols = [
        str((x or {}).get("name") or "").strip()
        for x in (local_vars or [])
        if str((x or {}).get("name") or "").strip()
    ]
    for item in (local_vars or []):
        backend._repair_local_cn_name_with_profile(
            item,
            body=body,
            neighbor_symbols=[x for x in neighbor_symbols if x and x != utils_module._safe_strip((item or {}).get("name"))],
            comment_desc=utils_module._safe_strip((comment_info or {}).get("desc")),
            cfg=cfg,
        )


def build_design_text_sections(
    ctx: dict[str, Any],
    module_req_prefix: str,
    index: int,
    cfg,
    *,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    comment_info = ctx["comment_info"]
    func_info = ctx["func_info"]
    body = ctx["body"]

    registered_title = utils_module._safe_strip((ctx.get("file_context") or {}).get("function_title"))
    title_cn = registered_title or backend._normalize_function_cn_title(
        backend.get_function_chinese_name_rich(
            ctx,
            cfg=cfg,
        ) if getattr(cfg, "ai_assist", False) else backend.get_function_chinese_name(comment_info, func_info),
        func_name=func_info.get("func_name", ""),
        comment_desc=ctx.get("raw_comment_desc") or utils_module._safe_strip(comment_info.get("desc")),
    )
    # AI ident decompose fallback: only fire when the static ident decomposition
    # (token dictionary) produced nothing useful — empty or still has English
    # beyond known domain acronyms. Avoids re-decomposing cases where the C
    # comment already gave a clean Chinese title.
    if (not registered_title
            and bool(getattr(cfg, "ai_assist", False))
            and func_info.get("func_name")
            and title_cn != func_info.get("func_name")):
        _ident_guessed = naming_utils.guess_cn_from_ident(func_info.get("func_name", ""))
        # Only fire when ident decomposition COMPLETELY failed (returned empty).
        # If guess_cn_from_ident already produced partial Chinese (even with
        # acronyms like PBIT), let normalize_function_cn_title handle it.
        _ident_failed = not _ident_guessed
        if _ident_failed:
            from . import ai as ai_module

            def _call_llm_text(prompt):
                return backend.call_llm_text(prompt, cfg, max_tokens=32)
            ai_name = naming_utils._ai_decompose_ident(
                func_info.get("func_name", ""),
                comment_desc=ctx.get("raw_comment_desc") or utils_module._safe_strip(comment_info.get("desc")),
                backend_module=backend,
                call_llm=_call_llm_text,
            )
            if ai_name:
                title_cn = backend._normalize_function_cn_title(
                    ai_name,
                    func_name=func_info.get("func_name", ""),
                    comment_desc=ctx.get("raw_comment_desc") or utils_module._safe_strip(comment_info.get("desc")),
                )
    if bool(getattr(cfg, "enhanced_single_func_pseudocode", False)):
        enhanced_desc = backend._suggest_enhanced_single_function_desc(
            func_info,
            body,
            current_desc=(comment_info.get("desc") or ""),
        )
        if enhanced_desc:
            comment_info["desc"] = enhanced_desc
    if backend._looks_like_pseudo_function_desc(comment_info.get("desc") or ""):
        from . import ai as ai_utils

        fallback_desc = ai_utils._fallback_function_description(
            func_info,
            body,
            current_desc=comment_info.get("desc") or "",
        )
        if fallback_desc:
            comment_info["desc"] = fallback_desc
    comment_info["desc"] = _repair_function_desc_by_domain(
        utils_module._safe_strip(func_info.get("func_name")),
        utils_module._safe_strip(comment_info.get("desc")),
        current_desc=ctx.get("raw_comment_desc") or "",
    )
    return {
        "title_cn": title_cn,
        "req_id": f"{backend.normalize_req_prefix(module_req_prefix)}_{index:03d}",
        "prototype": func_info.get("prototype") or "",
        "description_lines": logic_utils._clean_description_lines(utils_module._safe_strip(comment_info.get("desc"))),
    }


def collect_design_quality_inputs(
    ctx: dict[str, Any],
    logic_lines,
    name_map: dict[str, str],
    *,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    local_vars = ctx["local_vars"]
    params = ctx["params"]
    in_map = ctx["in_map"]
    out_map = ctx["out_map"]
    param_ai_name_map = ctx["param_ai_name_map"]
    scope_inference_log = ctx["scope_inference_log"]
    file_symbol_inference_log = ctx["file_symbol_inference_log"]

    post_missing_params = []
    for item in params:
        pname = item.get("name") or ""
        final_param_name = utils_module._safe_strip(in_map.get(pname) or out_map.get(pname) or param_ai_name_map.get(pname) or "")
        if pname and backend._is_missing_gap_text(final_param_name):
            post_missing_params.append(pname)

    post_missing_locals = []
    for item in local_vars:
        vname = item.get("name") or ""
        if vname and backend._is_missing_gap_text(item.get("usage") or ""):
            post_missing_locals.append(vname)

    logic_placeholders = sum(1 for line in (logic_lines or ()) if "待人工修改" in str(line or ""))
    quality_report = backend._collect_function_quality_report(
        local_vars,
        params,
        in_map,
        out_map,
        param_ai_name_map,
        logic_lines,
        name_map,
        inferences=tuple(list(scope_inference_log.values()) + list(file_symbol_inference_log.values())),
    )
    return {
        "post_missing_params": tuple(post_missing_params),
        "post_missing_locals": tuple(post_missing_locals),
        "logic_placeholders": int(logic_placeholders),
        "quality_report": quality_report,
    }


def _clean_logic_audit_text(text: Any) -> str:
    value = utils_module._safe_strip(text)
    value = re.sub(r"[；;]+$", "", value).strip()
    return re.sub(r"\s+", " ", value)


def _logic_source_audit_from_context(ctx: dict[str, Any], logic_lines: Optional[Sequence[str]], *, backend_module=None) -> tuple[dict[str, Any], ...]:
    backend = backend_module or legacy_backend()
    pack = dict(ctx.get("logic_semantic_pack") or {})
    if not logic_lines:
        return ()

    statement_hints: list[str] = []
    callee_comment_actions: list[str] = []
    for item in (pack.get("call_roles") or ()):
        if not isinstance(item, dict):
            continue
        for hint in (item.get("comment_hints") or ()):
            text = utils_module._safe_strip((hint or {}).get("text") if isinstance(hint, dict) else getattr(hint, "text", ""))
            if text:
                statement_hints.append(text)
        definition_comment = utils_module._safe_strip(item.get("definition_comment"))
        if definition_comment:
            action = utils_module._safe_strip(
                logic_utils._definition_comment_action(definition_comment, backend_module=backend)
            )
            if action:
                callee_comment_actions.append(action)

    has_control = bool(pack.get("control_blocks"))
    has_state = bool(pack.get("state_updates"))
    has_returns = bool(pack.get("return_actions"))
    has_patterns = bool(pack.get("pattern_hits"))
    # The rendered lines do not retain the parser node directly.  Keep a
    # monotonic, kind-aware association to the enriched semantic facts so a
    # quality issue can tell the repair prompt which C statement it represents.
    anchor_items = {
        "control_block": list(pack.get("control_blocks") or ()),
        "return_action": list(pack.get("return_actions") or ()),
        "statement_comment": list(pack.get("call_roles") or ()),
        "callee_comment": list(pack.get("call_roles") or ()),
        "call_role": list(pack.get("call_roles") or ()),
        "state_update": list(pack.get("state_updates") or ()),
        "semantic_pattern": list(pack.get("pattern_hits") or ()),
    }
    anchor_cursors = {key: 0 for key in anchor_items}

    def _take_source_anchor(kind: str) -> dict[str, Any]:
        items = anchor_items.get(kind) or ()
        pos = anchor_cursors.get(kind, 0)
        if pos >= len(items):
            return {}
        anchor_cursors[kind] = pos + 1
        item = items[pos] if isinstance(items[pos], dict) else {}
        anchor = dict(item.get("source_anchor") or {})
        if anchor:
            anchor.setdefault("statement_kind", kind)
        return anchor
    out: list[dict[str, Any]] = []

    for idx, raw_line in enumerate(logic_lines or (), start=1):
        text = _clean_logic_audit_text(raw_line)
        if not text:
            continue
        source = "post_polish"
        refinements: list[str] = []

        if logic_utils._is_control_logic_line(text):
            source = "control_block" if has_control else "control_flow"
        elif has_returns and text.startswith("返回"):
            source = "return_action"
        elif any(hint and hint in text for hint in statement_hints):
            source = "statement_comment"
        elif any(action and action in text for action in callee_comment_actions):
            source = "callee_comment"
        elif has_patterns and re.match(r"^(?:组装|执行|标记|同步|记录|计算)", text):
            source = "semantic_pattern"
        elif has_state and re.match(r"^(?:设置|将|计算|标记|清零|置位|置|读取|记录|选取|由)", text):
            source = "state_update"
        elif re.match(r"^(?:调用|执行)", text):
            source = "call_role"

        if "为空" in text or "不为空" in text:
            refinements.append("null_condition_polish")
        if any(token in text for token in ("RIU受油", "阀位超时原因", "无检查原因", "预位失败原因")):
            refinements.append("macro_display_name")
        if "低压故障" in text and source != "callee_comment":
            refinements.append("callee_comment")
        if "当前系统时间" in text:
            refinements.append("call_result_polish")

        item = {
            "idx": idx,
            "text": text,
            "source": source,
            "refinements": tuple(dict.fromkeys(refinements)),
        }
        source_anchor = _take_source_anchor(source)
        if source_anchor:
            item["source_anchor"] = source_anchor
        out.append(item)
    return tuple(out)


def _quality_issue(code: str, message: str, *, severity: str = "warning", source: str = "", anchor: Optional[dict[str, Any]] = None) -> dict[str, Any]:
    item = {
        "code": utils_module._safe_strip(code),
        "severity": utils_module._safe_strip(severity) or "warning",
        "message": utils_module._safe_strip(message),
        "source": utils_module._safe_strip(source),
    }
    if anchor:
        item["source_anchor"] = dict(anchor)
    return item


def _logic_rendering_quality_issues(
    logic_lines: Optional[Sequence[str]],
    source_anchors: Optional[Sequence[dict[str, Any]]] = None,
) -> tuple[dict[str, Any], ...]:
    return quality_gate.inspect_logic_lines(logic_lines, source_anchors=source_anchors)


def _quality_issues_from_context(
    ctx: dict[str, Any],
    logic_lines: Optional[Sequence[str]],
    quality_report: dict[str, Any],
    logic_source_audit: Sequence[dict[str, Any]],
    *,
    backend_module=None,
) -> tuple[dict[str, Any], ...]:
    _ = backend_module or legacy_backend()
    issues: list[dict[str, Any]] = []
    # One-hop resolution failures are deliberately warnings: they are visible
    # to review but must not make the generator fabricate external effects.
    issues.extend(
        dict(item) for item in (ctx.get("effect_quality_issues") or ())
        if isinstance(item, dict)
    )
    pack = dict(ctx.get("logic_semantic_pack") or {})
    summary = dict(pack.get("quality_summary") or {})
    if int(summary.get("missing_source_anchor_count") or 0) > 0:
        issues.append(
            _quality_issue(
                "semantic_anchor_missing",
                f"{int(summary.get('missing_source_anchor_count') or 0)} 条语义事实缺少源码锚点",
                source="semantic_pack",
            )
        )
    if int(summary.get("raw_macro_ref_count") or 0) > 0:
        issues.append(
            _quality_issue(
                "raw_macro_ref",
                f"{int(summary.get('raw_macro_ref_count') or 0)} 个宏名仍未解析为文档显示名",
                source="naming_resolver",
            )
        )
    if int(summary.get("low_confidence_name_count") or 0) > 0:
        issues.append(
            _quality_issue(
                "low_confidence_name",
                f"{int(summary.get('low_confidence_name_count') or 0)} 个名称来自低置信度猜测",
                severity="info",
                source="naming_resolver",
            )
        )
    for key, code, label in (
        ("unresolved_locals", "locals_unresolved", "局部变量未解析"),
        ("unresolved_params", "params_unresolved", "参数未解析"),
        ("unresolved_logic_symbols", "logic_symbol_unresolved", "逻辑语句残留源码标识符"),
    ):
        values = tuple(quality_report.get(key) or ())
        if values:
            issues.append(_quality_issue(code, f"{label}：{', '.join(str(x) for x in values[:8])}", source="quality_report"))
    numeric_checks = (
        ("generic_logic_count", "generic_logic", "泛化逻辑句"),
        ("comment_leak_count", "comment_leak", "注释泄露风险"),
        ("term_drift_count", "term_drift", "术语漂移"),
        ("over_translation_count", "over_translation", "过度翻译"),
        ("bad_symbol_guess_count", "bad_symbol_guess", "疑似错误命名猜测"),
    )
    for key, code, label in numeric_checks:
        count = int(quality_report.get(key) or 0)
        if count > 0:
            issues.append(_quality_issue(code, f"{label}：{count}", source="quality_report"))
    issues.extend(_logic_rendering_quality_issues(logic_lines, logic_source_audit))
    internal_terms = ("logic_source_audit", "callee_comment", "state_update")
    for idx, line in enumerate(logic_lines or (), start=1):
        text = utils_module._safe_text(line)
        hit = next((term for term in internal_terms if term in text), "")
        if hit:
            issues.append(_quality_issue("internal_field_leak", f"逻辑第 {idx} 行包含内部字段：{hit}", severity="error", source="logic_lines"))
    if not logic_source_audit and logic_lines:
        issues.append(_quality_issue("logic_audit_empty", "逻辑存在但来源审计为空", severity="info", source="logic_source_audit"))
    deduped: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for item in issues:
        key = (utils_module._safe_strip(item.get("code")), utils_module._safe_strip(item.get("message")))
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return tuple(deduped[:32])


def _build_lsp_fact_snapshot_for_meta(
    ctx: dict[str, Any],
    quality_report: dict[str, Any],
    *,
    backend_module=None,
) -> dict[str, Any]:
    """Extract a compact LSP fact snapshot for regression-round quality feedback.

    Only includes facts relevant to unresolved symbols, so the regression prompt
    stays small while giving the AI precise type/member/call context.
    """
    fact_pack = dict(ctx.get("lsp_fact_pack") or {})
    if not fact_pack:
        return {}
    focus = set()
    for key in ("unresolved_locals", "unresolved_params", "unresolved_logic_symbols"):
        for name in (quality_report.get(key) or []):
            name = utils_module._safe_strip(name)
            if name:
                focus.add(name)
    if not focus:
        return {}

    # local/param types
    type_facts: dict[str, str] = {}
    for item in list((fact_pack.get("locals") or [])) + list((fact_pack.get("params") or [])):
        if not isinstance(item, dict):
            continue
        name = utils_module._safe_strip(item.get("name"))
        if name in focus:
            decl_type = utils_module._safe_strip(item.get("decl_type"))
            if decl_type:
                type_facts[name] = decl_type

    # member accesses for focus symbols
    member_facts: list[dict[str, str]] = []
    for item in (fact_pack.get("members") or []):
        if not isinstance(item, dict):
            continue
        base = utils_module._safe_strip(item.get("base"))
        if base in focus:
            member_facts.append({
                "symbol": base,
                "member": utils_module._safe_strip(item.get("member")),
                "owner_type": utils_module._safe_strip(item.get("owner_type")),
                "access_text": utils_module._safe_strip(item.get("access_text")),
            })

    # calls involving focus symbols (as arguments or callees)
    call_facts: list[dict[str, str]] = []
    for item in (fact_pack.get("calls") or []):
        if not isinstance(item, dict):
            continue
        callee = utils_module._safe_strip(item.get("callee"))
        call_text = utils_module._safe_strip(item.get("call_text") or item.get("signature") or "")
        if callee in focus or any(s in call_text for s in focus if len(s) > 2):
            call_facts.append({"callee": callee, "call_text": call_text[:120]})

    # control blocks (for logic context)
    block_facts: list[dict[str, str]] = []
    for item in (fact_pack.get("blocks") or [])[:8]:
        if not isinstance(item, dict):
            continue
        block_facts.append({
            "kind": utils_module._safe_strip(item.get("kind")),
            "condition": utils_module._safe_strip(item.get("condition"))[:120],
        })

    snapshot = {}
    if type_facts:
        snapshot["type_facts"] = type_facts
    if member_facts:
        snapshot["member_facts"] = member_facts[:12]
    if call_facts:
        snapshot["call_facts"] = call_facts[:8]
    if block_facts:
        snapshot["block_facts"] = block_facts
    return snapshot


def build_design_ai_meta(
    ctx: dict[str, Any],
    cfg,
    quality_inputs: dict[str, Any],
    *,
    backend_module=None,
    one_call_guard_fallback: bool = False,
    one_call_locals_expected: int = 0,
    one_call_locals_got: int = 0,
    one_call_params_expected: int = 0,
    one_call_params_got: int = 0,
    one_call_logic_expected: int = 0,
    one_call_logic_got: int = 0,
    one_call_logic_degraded: bool = False,
):
    backend = backend_module or legacy_backend()
    comment_info = ctx["comment_info"]
    logic_lines = quality_inputs.get("logic_lines")
    initial_gaps = ctx["initial_gaps"]
    post_missing_params = list(quality_inputs.get("post_missing_params") or ())
    post_missing_locals = list(quality_inputs.get("post_missing_locals") or ())
    logic_placeholders = int(quality_inputs.get("logic_placeholders") or 0)
    quality_report = quality_inputs.get("quality_report") or {}
    title_debug = dict(getattr(cfg, "_current_func_title_debug", {}) or {})
    logic_source_audit = _logic_source_audit_from_context(ctx, logic_lines, backend_module=backend)
    quality_issues = _quality_issues_from_context(
        ctx,
        logic_lines,
        quality_report,
        logic_source_audit,
        backend_module=backend,
    )

    ai_regression_reasons = []
    ai_failed = bool(getattr(cfg, "_current_func_ai_failed", False))
    if bool(getattr(cfg, "ai_assist", False)):
        if ai_failed:
            ai_regression_reasons.append("ai_failed")
        if initial_gaps.get("need_func_desc") and backend._is_missing_gap_text(comment_info.get("desc") or ""):
            ai_regression_reasons.append("func_desc_unfilled")
        if initial_gaps.get("need_func_cn_name") and backend._is_missing_gap_text(comment_info.get("func_cn_name") or ""):
            ai_regression_reasons.append("func_name_unfilled")
        if initial_gaps.get("need_param_names") and post_missing_params:
            ai_regression_reasons.append("params_unfilled")
        if initial_gaps.get("need_local_usages") and post_missing_locals:
            ai_regression_reasons.append("locals_unfilled")
        if logic_placeholders > 0:
            ai_regression_reasons.append("logic_placeholder")
        for issue in quality_issues:
            code = utils_module._safe_strip(issue.get("code"))
            if code in quality_gate.STRUCTURAL_LOGIC_CODES:
                ai_regression_reasons.append(code)
        if quality_report["unresolved_locals"]:
            pass  # 不触发回归：符号不在术语表中，重试无法翻译
        if quality_report["unresolved_params"]:
            pass  # 不触发回归：同上
        if quality_report["unresolved_logic_symbols"]:
            pass  # 不触发回归：同上
        generic_logic_limit = max(2, len(logic_lines or ()) // 4) if logic_lines else 2
        if int(quality_report["generic_logic_count"] or 0) >= generic_logic_limit:
            ai_regression_reasons.append("logic_generic")
        if int(quality_report["comment_leak_count"] or 0) > 0:
            ai_regression_reasons.append("comment_leak")
        if int(quality_report["term_drift_count"] or 0) > 0:
            ai_regression_reasons.append("term_drift")
        if int(quality_report["over_translation_count"] or 0) > 0:
            ai_regression_reasons.append("over_translation")
        if int(quality_report["bad_symbol_guess_count"] or 0) > 0:
            ai_regression_reasons.append("bad_symbol_guess")
        if quality_report.get("thin_logic"):
            pass  # 不触发回归：简单函数逻辑行少是正常的
        if quality_report.get("generic_call_count", 0) >= 3:
            ai_regression_reasons.append("generic_calls_excess")
        if bool(getattr(cfg, "ai_one_call", False)) and (not one_call_guard_fallback):
            if one_call_locals_expected > 0 and one_call_locals_got <= 0:
                ai_regression_reasons.append("one_call_locals_empty")
            if one_call_params_expected > 0 and one_call_params_got <= 0:
                ai_regression_reasons.append("one_call_params_empty")
            if one_call_logic_expected > 0 and one_call_logic_got <= 0 and (not one_call_logic_degraded):
                ai_regression_reasons.append("one_call_logic_empty")

    # LSP 事实快照：供回归轮 compose_quality_feedback_text 使用
    lsp_fact_snapshot = _build_lsp_fact_snapshot_for_meta(ctx, quality_report, backend_module=backend)
    return AIBuildMeta(
        ai_enabled=bool(getattr(cfg, "ai_assist", False)),
        ai_failed=ai_failed,
        regression_needed=bool(ai_regression_reasons),
        regression_round=max(0, int(getattr(cfg, "_ai_regression_round", 0) or 0)),
        regression_reasons=tuple(dict.fromkeys(ai_regression_reasons)),
        logic_placeholders=int(logic_placeholders),
        unresolved_local_symbols=tuple(quality_report["unresolved_locals"]),
        unresolved_param_symbols=tuple(quality_report["unresolved_params"]),
        unresolved_logic_symbols=tuple(quality_report["unresolved_logic_symbols"]),
        generic_logic_count=int(quality_report["generic_logic_count"] or 0),
        comment_leak_count=int(quality_report["comment_leak_count"] or 0),
        term_drift_count=int(quality_report["term_drift_count"] or 0),
        over_translation_count=int(quality_report["over_translation_count"] or 0),
        bad_symbol_guess_count=int(quality_report["bad_symbol_guess_count"] or 0),
        raw_func_title=utils_module._safe_strip(title_debug.get("raw_func_cn_name")),
        pre_rerank_func_title=utils_module._safe_strip(title_debug.get("pre_rerank_func_cn_name")),
        title_candidates=tuple(str(x) for x in (title_debug.get("candidates") or ()) if utils_module._safe_strip(x)),
        title_pattern=utils_module._safe_strip(title_debug.get("pattern")),
        title_rerank_changed=bool(title_debug.get("rerank_changed")),
        title_fallback_used=bool(title_debug.get("fallback_used")),
        title_model_confidence=float(title_debug.get("model_confidence", 0.0) or 0.0),
        logic_source_audit=logic_source_audit,
        quality_issues=quality_issues,
        lsp_fact_snapshot=lsp_fact_snapshot,
    )


def assemble_function_design(
    text_sections: dict[str, Any],
    io_elements,
    io_none: bool,
    local_elements,
    logic_lines,
    ai_meta,
    name_map: dict[str, str],
    effects=(),
    return_effects=(),
    *,
    backend_module=None,
):
    backend = backend_module or legacy_backend()
    deterministic_logic = tuple(logic_lines or ())
    design = FunctionDesign(
        title=text_sections["title_cn"],
        req_id=text_sections["req_id"],
        prototype=text_sections["prototype"],
        description_lines=text_sections["description_lines"],
        io_elements=tuple(io_elements),
        io_none=io_none,
        local_elements=local_elements,
        logic_lines=logic_lines,
        ai_meta=ai_meta,
        effects=tuple(effects or ()),
        return_effects=tuple(return_effects or ()),
    )
    design = backend._normalize_function_design_texts(design, name_map=name_map)
    # Text normalization may consult persisted symbol memory.  Never allow a
    # stale AI name from that layer to turn an otherwise valid deterministic
    # logic sentence into malformed text.
    meta = design.ai_meta if isinstance(design.ai_meta, AIBuildMeta) else AIBuildMeta()
    post_normalize_issues = quality_gate.inspect_logic_lines(
        design.logic_lines, source_anchors=meta.logic_source_audit,
    )
    if post_normalize_issues and deterministic_logic:
        repaired_lines = list(design.logic_lines or ())
        restored: list[int] = []
        omitted: list[int] = []
        for issue in post_normalize_issues:
            try:
                idx = int(issue.get("logic_line")) - 1
            except (TypeError, ValueError):
                continue
            if idx < 0 or idx >= len(repaired_lines):
                continue
            fallback = deterministic_logic[idx] if idx < len(deterministic_logic) else ""
            if fallback and not quality_gate.inspect_logic_lines((fallback,)):
                repaired_lines[idx] = fallback
                restored.append(idx + 1)
            else:
                repaired_lines[idx] = ""
                omitted.append(idx + 1)
        if restored or omitted:
            recovery = ({
                "action": "post_normalize_deterministic_fallback",
                "lines": tuple(restored),
                "omitted_lines": tuple(omitted),
            },)
            design = replace(
                design,
                logic_lines=tuple(repaired_lines),
                ai_meta=_meta_with_structural_quality(meta, tuple(repaired_lines), recovery),
            )
    # Symbol coverage check: scan for identifiers missing from name_map
    unresolved = logic_utils._collect_unresolved_logic_symbols(
        design.logic_lines, name_map=name_map, backend_module=backend
    )
    if unresolved:
        enriched_map = dict(name_map or {})
        for sym in unresolved:
            try:
                resolved = backend.resolve_canonical_symbol_name(
                    sym, kind="symbols", comment_cn="", fallback=""
                )
                if resolved and resolved != sym and text_utils._contains_cjk(resolved):
                    enriched_map[sym] = resolved
            except Exception:
                pass
        if len(enriched_map) > len(name_map or {}):
            design = backend._normalize_function_design_texts(design, name_map=enriched_map)
    return design


def collect_design_components(
    ctx: dict[str, Any],
    func_data: dict,
    module_req_prefix: str,
    index: int,
    cfg,
    *,
    backend_module=None,
    one_call_bundle=None,
    one_call_guard_fallback: bool = False,
    one_call_locals_expected: int = 0,
    one_call_locals_got: int = 0,
    one_call_params_expected: int = 0,
    one_call_params_got: int = 0,
    one_call_logic_expected: int = 0,
    one_call_logic_got: int = 0,
    one_call_logic_degraded: bool = False,
    cached_logic_lines: Optional[dict[int, str]] = None,
    changed_statement_lines: Optional[set[int]] = None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()

    # ---- Batch LLM variable naming (before repair/display) ----
    if (
        getattr(cfg, "ai_assist", False)
        and (not backend.stop_requested(cfg))
        and (not getattr(cfg, "_user_cancelled", False))
    ):
        try:
            from . import naming_context as nc

            func_info = ctx.get("func_info") or {}
            func_name = utils_module._safe_strip(func_info.get("func_name"))
            source_file = utils_module._safe_strip((ctx.get("file_context") or {}).get("source_file"))
            body = ctx.get("body") or ""

            # Get cached body summary (already computed by rich naming)
            body_summary = nc.get_cached_summary(func_name, source_file, body) if func_name else ""

            # Get function Chinese name (already computed by text_sections)
            # We'll compute it here too since text_sections hasn't been built yet
            func_cn_name = utils_module._safe_strip(
                (ctx.get("comment_info") or {}).get("func_cn_name")
            )

            var_cn_map = naming_utils.get_variable_chinese_names_batch(
                ctx,
                func_cn_name=func_cn_name,
                body_summary=body_summary,
                cfg=cfg,
                backend_module=backend,
            )
            if var_cn_map:
                # Batch naming is an AI ingress point.  Validate each field
                # before it reaches local metadata or the shared name map;
                # one malformed candidate must not contaminate many logic
                # lines during later symbol replacement.
                safe_var_cn_map: dict[str, str] = {}
                for raw_name, payload in var_cn_map.items():
                    ident = utils_module._safe_strip(raw_name)
                    candidate = utils_module._safe_strip(
                        payload.get("cn_name") if isinstance(payload, dict) else payload
                    )
                    if not ident or not candidate:
                        continue
                    if not quality_gate.is_safe_ai_text(candidate):
                        continue
                    safe_var_cn_map[ident] = candidate
                ctx["var_cn_map"] = safe_var_cn_map
                ctx["var_usage_map"] = {
                    k: naming_utils.sanitize_ai_usage_text(v.get("usage"))
                    for k, v in var_cn_map.items()
                    if isinstance(v, dict) and k and naming_utils.sanitize_ai_usage_text(v.get("usage"))
                }
        except Exception:
            pass

    revision_patch = ctx.get("_revision_patch")
    if revision_patch:
        revision_utils.apply_revision_to_context(ctx, revision_patch)

    repair_design_local_profiles(ctx, cfg, backend_module=backend)

    text_sections = build_design_text_sections(
        ctx,
        module_req_prefix,
        index,
        cfg,
        backend_module=backend,
    )
    io_elements, io_none, ret_var_name = build_design_io_elements(
        ctx,
        func_data,
        cfg,
        backend_module=backend,
        one_call_bundle=one_call_bundle,
    )
    local_elements = build_design_local_elements(
        ctx,
        cfg,
        backend_module=backend,
        ret_var_name=ret_var_name,
    )
    name_map = build_design_name_map(ctx, backend_module=backend)
    # Merge LLM batch variable names into name_map for logic line rendering
    var_cn_map = ctx.get("var_cn_map") or {}
    if var_cn_map:
        for k, v in var_cn_map.items():
            if k and v and k not in name_map:
                name_map[k] = v
    # Enrich local_vars cn_name with var_cn_map so generate_logic_from_body sees them
    if var_cn_map:
        for item in (ctx.get("local_vars") or []):
            vname = item.get("name", "")
            if vname in var_cn_map and not item.get("cn_name"):
                item["cn_name"] = var_cn_map[vname]
    effect_facts, return_effects, effect_issues = build_function_effect_facts(ctx, func_data, cfg, name_map)
    ctx["effect_quality_issues"] = tuple(effect_issues)
    logic_lines = build_design_logic_lines(
        ctx,
        func_data,
        cfg,
        name_map,
        backend_module=backend,
        cached_logic_lines=cached_logic_lines,
        changed_statement_lines=changed_statement_lines,
    )
    quality_inputs = collect_design_quality_inputs(
        ctx,
        logic_lines,
        name_map,
        backend_module=backend,
    )
    quality_inputs["logic_lines"] = logic_lines
    ai_meta = build_design_ai_meta(
        ctx,
        cfg,
        quality_inputs,
        backend_module=backend,
        one_call_guard_fallback=one_call_guard_fallback,
        one_call_locals_expected=one_call_locals_expected,
        one_call_locals_got=one_call_locals_got,
        one_call_params_expected=one_call_params_expected,
        one_call_params_got=one_call_params_got,
        one_call_logic_expected=one_call_logic_expected,
        one_call_logic_got=one_call_logic_got,
        one_call_logic_degraded=one_call_logic_degraded,
    )
    return {
        "text_sections": text_sections,
        "io_elements": io_elements,
        "io_none": io_none,
        "local_elements": local_elements,
        "logic_lines": logic_lines,
        "ai_meta": ai_meta,
        "name_map": name_map,
        "effects": effect_facts,
        "return_effects": return_effects,
    }


def build_design_output(
    ctx: dict[str, Any],
    func_data: dict,
    module_req_prefix: str,
    index: int,
    cfg,
    *,
    backend_module=None,
    one_call_bundle=None,
    one_call_guard_fallback: bool = False,
    one_call_locals_expected: int = 0,
    one_call_locals_got: int = 0,
    one_call_params_expected: int = 0,
    one_call_params_got: int = 0,
    one_call_logic_expected: int = 0,
    one_call_logic_got: int = 0,
    one_call_logic_degraded: bool = False,
    cached_logic_lines: Optional[dict[int, str]] = None,
    changed_statement_lines: Optional[set[int]] = None,
):
    backend = backend_module or legacy_backend()
    components = collect_design_components(
        ctx,
        func_data,
        module_req_prefix,
        index,
        cfg,
        backend_module=backend,
        one_call_bundle=one_call_bundle,
        one_call_guard_fallback=one_call_guard_fallback,
        one_call_locals_expected=one_call_locals_expected,
        one_call_locals_got=one_call_locals_got,
        one_call_params_expected=one_call_params_expected,
        one_call_params_got=one_call_params_got,
        one_call_logic_expected=one_call_logic_expected,
        one_call_logic_got=one_call_logic_got,
        one_call_logic_degraded=one_call_logic_degraded,
        cached_logic_lines=cached_logic_lines,
        changed_statement_lines=changed_statement_lines,
    )
    design = assemble_function_design(
        components["text_sections"],
        components["io_elements"],
        components["io_none"],
        components["local_elements"],
        components["logic_lines"],
        components["ai_meta"],
        components["name_map"],
        components["effects"],
        components["return_effects"],
        backend_module=backend,
    )
    return revision_utils.apply_revision_to_design(design, ctx.get("_revision_patch"))


def build_function_design_impl(
    func_data: dict,
    module_req_prefix: str,
    index: int,
    cfg,
    *,
    backend_module=None,
    cached_logic_lines: Optional[dict[int, str]] = None,
    changed_statement_lines: Optional[set[int]] = None,
):
    backend = backend_module or legacy_backend()
    from . import render as render_module
    setattr(cfg, "_current_func_title_debug", {})

    # Design cache: reuse rule-based parsing across AI mode switches
    _dkey_file = str((func_data.get("file_context") or {}).get("source_file") or "")
    _dkey_name = str((func_data.get("func_info") or {}).get("func_name") or "")
    _dkey_body = str(func_data.get("body") or "")
    _design_key = render_module._design_cache_key(_dkey_file, _dkey_name, _dkey_body)
    cached_ctx = render_module.get_design_cache(_design_key)
    if cached_ctx is not None:
        ctx = dict(cached_ctx)
        ctx["cfg"] = cfg
    else:
        ctx = prepare_design_context(func_data, cfg, backend_module=backend)
        render_module.put_design_cache(_design_key, ctx)
    # ── P0#3 Evidence 旁路采集（shadow mode，不影响生成）──
    try:
        _evidence_shadow_collect(func_data, ctx, cfg)
    except Exception:
        pass
    revision_profile = revision_utils.load_revision_profile(cfg)
    revision_patch = revision_utils.find_function_patch(
        revision_profile,
        (ctx.get("file_context") or {}).get("source_file") or ctx.get("source_file") or "",
        (ctx.get("func_info") or {}).get("func_name") or "",
    )
    if revision_patch:
        revision_utils.apply_revision_to_context(ctx, revision_patch)
    one_call_state = run_one_call_design_enrichment(
        ctx,
        func_data,
        cfg,
        index,
        backend_module=backend,
    )
    if revision_patch:
        revision_utils.apply_revision_to_context(ctx, revision_patch)
    # 检查停止信号（包括用户取消后的 _user_cancelled 标记）
    if backend.stop_requested(cfg) or getattr(cfg, "_user_cancelled", False):
        return build_design_output(
            ctx,
            func_data,
            module_req_prefix,
            index,
            cfg,
            backend_module=backend,
            one_call_bundle=None,
            one_call_guard_fallback=True,
            one_call_locals_expected=0,
            one_call_locals_got=0,
            one_call_params_expected=0,
            one_call_params_got=0,
            one_call_logic_expected=0,
            one_call_logic_got=0,
            one_call_logic_degraded=False,
            cached_logic_lines=cached_logic_lines,
            changed_statement_lines=changed_statement_lines,
        )
    func_data = one_call_state["func_data"]
    # 用户取消时跳过多步 AI 调用，直接输出已有结果
    if backend.stop_requested(cfg) or getattr(cfg, "_user_cancelled", False):
        return build_design_output(
            ctx,
            func_data,
            module_req_prefix,
            index,
            cfg,
            backend_module=backend,
            one_call_bundle=one_call_state.get("one_call_bundle"),
            one_call_guard_fallback=one_call_state.get("one_call_guard_fallback", False),
            one_call_locals_expected=one_call_state.get("one_call_locals_expected", 0),
            one_call_locals_got=one_call_state.get("one_call_locals_got", 0),
            one_call_params_expected=one_call_state.get("one_call_params_expected", 0),
            one_call_params_got=one_call_state.get("one_call_params_got", 0),
            one_call_logic_expected=one_call_state.get("one_call_logic_expected", 0),
            one_call_logic_got=one_call_state.get("one_call_logic_got", 0),
            one_call_logic_degraded=one_call_state.get("one_call_logic_degraded", False),
            cached_logic_lines=cached_logic_lines,
            changed_statement_lines=changed_statement_lines,
        )
    run_multi_call_design_enrichment(
        ctx,
        func_data,
        cfg,
        backend_module=backend,
    )
    if revision_patch:
        revision_utils.apply_revision_to_context(ctx, revision_patch)
    return build_design_output(
        ctx,
        func_data,
        module_req_prefix,
        index,
        cfg,
        backend_module=backend,
        one_call_bundle=one_call_state["one_call_bundle"],
        one_call_guard_fallback=one_call_state["one_call_guard_fallback"],
        one_call_locals_expected=one_call_state["one_call_locals_expected"],
        one_call_locals_got=one_call_state["one_call_locals_got"],
        one_call_params_expected=one_call_state["one_call_params_expected"],
        one_call_params_got=one_call_state["one_call_params_got"],
        one_call_logic_expected=one_call_state["one_call_logic_expected"],
        one_call_logic_got=one_call_state["one_call_logic_got"],
        one_call_logic_degraded=one_call_state["one_call_logic_degraded"],
        cached_logic_lines=cached_logic_lines,
        changed_statement_lines=changed_statement_lines,
    )


def build_function_design(func_data: dict, module_req_prefix: str, index: int, cfg):
    backend = legacy_backend()
    impl = getattr(backend, "_legacy_build_function_design_impl", None) or backend.build_function_design
    return impl(func_data, module_req_prefix, index, cfg)


def _continue_on_function_error(cfg, *, backend_module=None) -> bool:
    backend = backend_module or legacy_backend()
    try:
        return bool(int(utils_module.cfg_get_str(cfg, "continue_on_function_error", "1") or "1"))
    except Exception:
        extra = dict(getattr(cfg, "extra_params", {}) or {})
        return str(extra.get("continue_on_function_error", "1")).strip() != "0"


def _build_failed_function_design(task: dict[str, Any], error: Exception, cfg, *, backend_module=None):
    backend = backend_module or legacy_backend()
    func_data = dict(task.get("func_data") or {})
    func_info = dict(func_data.get("func_info") or {})
    func_name = utils_module._safe_strip(task.get("func_name")) or utils_module._safe_strip(func_info.get("func_name")) or f"函数{int(task.get('index') or 0)}"
    prototype = utils_module._safe_strip(func_info.get("prototype")) or f"void {func_name}(void)"
    ai_meta = AIBuildMeta(
        ai_enabled=bool(getattr(cfg, "ai_assist", False)),
        ai_failed=True,
        regression_needed=False,
        regression_reasons=("build_error",),
    )

    # 发送函数失败事件到 GUI
    try:
        from .models import FunctionFailureRecord
        failure = FunctionFailureRecord.from_exception(task, error)
        backend.gui_event(cfg, {
            "type": "func_failure",
            "func_name": failure.func_name,
            "file": failure.file_path,
            "func_index": task.get("index"),
            "func_pos": task.get("func_pos"),
            "error_type": failure.error_type,
            "error_message": failure.error_message,
            "task": failure.to_dict(),
        })
    except Exception:
        pass

    return FunctionDesign(
        title=func_name,
        req_id=f"{task.get('module_req_prefix')}_{int(task.get('index') or 0):03d}",
        prototype=prototype,
        description_lines=(f"函数生成失败：{error}",),
        io_elements=(),
        io_none=True,
        local_elements=(),
        logic_lines=(),
        ai_meta=ai_meta,
    )


def run_function_design_task(task: dict, cfg, *, backend_module=None):
    backend = backend_module or legacy_backend()
    task_cfg = cfg_with_function_task_context(cfg, task)
    build_task_fn = getattr(backend, "_build_function_design_task", None) or build_function_design_task
    if build_task_fn is build_function_design_task:
        design = build_task_fn(
            task["func_data"],
            task["module_req_prefix"],
            int(task["index"]),
            task_cfg,
            backend_module=backend,
        )
    else:
        design = build_task_fn(
            task["func_data"],
            task["module_req_prefix"],
            int(task["index"]),
            task_cfg,
        )
    regress_fn = getattr(backend, "_maybe_regress_function_design", None) or maybe_regress_function_design
    return regress_fn(task, design, cfg)


def iter_function_build_results(
    tasks: Sequence[dict],
    cfg,
    *,
    on_submit: Optional[Any] = None,
    backend_module=None,
):
    backend = backend_module or legacy_backend()
    if not tasks:
        return

    def _submit_cb(task: dict) -> None:
        if callable(on_submit):
            try:
                on_submit(task)
            except Exception:
                pass

    continue_on_error = _continue_on_function_error(cfg, backend_module=backend)

    if (not should_parallel_build_design(cfg)) or len(tasks) <= 1:
        for task in tasks:
            if backend.stop_requested(cfg):
                backend.vlog(cfg, "[停止] 检测到停止信号，退出任务循环")
                break
            _submit_cb(task)
            try:
                design = run_function_design_task(task, cfg, backend_module=backend)
            except Exception as exc:
                if not continue_on_error:
                    raise backend.FunctionBuildTaskError(task, exc) from exc
                design = _build_failed_function_design(task, exc, cfg, backend_module=backend)
                yield FunctionBuildResult(task=dict(task), design=design, error=exc)
                continue
            # 任务完成后再次检查停止信号
            if backend.stop_requested(cfg):
                backend.vlog(cfg, "[停止] 任务完成后检测到停止信号，退出任务循环")
                break
            yield FunctionBuildResult(task=dict(task), design=design, error=None)
        return

    max_workers = min(len(tasks), max(1, int(getattr(cfg, "ai_workers", 1) or 1)))
    future_map: dict[int, concurrent.futures.Future] = {}
    next_submit = 0
    next_yield = 0

    executor = concurrent.futures.ThreadPoolExecutor(max_workers=max_workers)
    stop_now = False
    try:
        while next_yield < len(tasks):
            if backend.stop_requested(cfg):
                stop_now = True
                for f in future_map.values():
                    if not f.done():
                        f.cancel()
                break

            while next_submit < len(tasks) and len(future_map) < max_workers and (not backend.stop_requested(cfg)):
                task = tasks[next_submit]
                _submit_cb(task)
                future_map[next_submit] = executor.submit(
                    run_function_design_task,
                    task,
                    cfg,
                    backend_module=backend,
                )
                next_submit += 1

            future = future_map.get(next_yield)
            if future is None:
                break

            # 检查停止信号，如果停止则取消未完成的任务，但仍回收当前任务的结果
            if backend.stop_requested(cfg):
                stop_now = True
                for f in future_map.values():
                    if f is not future and not f.done():
                        f.cancel()
                try:
                    design = future.result(timeout=2.0)
                except Exception:
                    pass
                else:
                    task = tasks[next_yield]
                    yield FunctionBuildResult(task=dict(task), design=design, error=None)
                    next_yield += 1
                break

            task = tasks[next_yield]
            try:
                design = _future_result_until_stop(future, cfg, backend)
            except TimeoutError:
                stop_now = True
                for other in future_map.values():
                    if not other.done():
                        other.cancel()
                break
            except Exception as exc:
                if not continue_on_error:
                    for other in future_map.values():
                        if other is not future:
                            other.cancel()
                    raise backend.FunctionBuildTaskError(task, exc) from exc
                design = _build_failed_function_design(task, exc, cfg, backend_module=backend)
                result = FunctionBuildResult(task=dict(task), design=design, error=exc)
            else:
                result = FunctionBuildResult(task=dict(task), design=design, error=None)
            finally:
                future_map.pop(next_yield, None)

            yield result
            next_yield += 1
    finally:
        for future in future_map.values():
            if not future.done():
                future.cancel()
        executor.shutdown(wait=(not stop_now))


def _iter_function_design_results(
    tasks: Sequence[dict],
    cfg,
    *,
    on_submit: Optional[Any] = None,
    backend_module=None,
):
    for result in iter_function_build_results(tasks, cfg, on_submit=on_submit, backend_module=backend_module):
        yield result.task, result.design, result.error


def build_project_layer_sets(
    app_modules: Sequence[dict],
    mid_modules: Sequence[dict],
    drv_modules: Sequence[dict],
) -> list[tuple[str, str, Sequence[dict]]]:
    return [
        ("", "应用层", app_modules),
        ("", "中间层", mid_modules),
        ("", "驱动层", drv_modules),
    ]


def should_include_function(func_data: dict, *, only_with_comment: bool) -> bool:
    if not only_with_comment:
        return True
    comment_info = (func_data or {}).get("comment_info") or {}
    if bool(
        comment_info.get("func_cn_name")
        or comment_info.get("func_name")
        or comment_info.get("desc")
        or comment_info.get("input_desc")
        or comment_info.get("output_desc")
        or comment_info.get("other_desc")
        or comment_info.get("return_desc")
    ):
        return True
    # 无注释块时，若从头文件/符号表已补全出中文名（func_cn_map），同样纳入生成，
    # 避免因注释缺失遗漏应生成 CSU 的函数。
    file_context = (func_data or {}).get("file_context") or {}
    func_cn_map = file_context.get("func_cn_map") or {}
    func_name = ((func_data or {}).get("func_info") or {}).get("func_name", "")
    return bool(func_cn_map.get(func_name))


def resolve_source_module_names(
    source: str,
    *,
    func_cn_name: str = "",
    func_ident: str = "",
    backend_module=None,
) -> tuple[str, str]:
    """返回 (原始文件名, 紧凑中文模块名)。

    优先级：
    1. 函数中文名（<=12字，AI给出的紧凑名）
    2. 函数C标识符拆词（guess_cn_from_ident）
    3. 文件头部推导（_derive_module_display_name）
    """
    backend = backend_module or legacy_backend()
    module_name_raw = os.path.splitext(os.path.basename(source))[0]
    candidate = utils_module._safe_strip(func_cn_name)
    if candidate and len(candidate) <= 12 and text_utils._contains_cjk(candidate):
        # 函数中文名已由 get_function_chinese_name 做过 normalize，直接用
        return module_name_raw, candidate
    # 函数中文名太长或为空 —— 尝试从 C 标识符拆词
    ident_hint = utils_module._safe_strip(func_ident) or module_name_raw
    guessed = backend._guess_cn_from_ident(ident_hint, glossary=backend.DOMAIN_GLOSSARY)
    if guessed and len(guessed) <= 12 and text_utils._contains_cjk(guessed):
        guessed = backend._normalize_function_cn_title(
            guessed, func_name=module_name_raw, comment_desc=guessed
        )
        return module_name_raw, guessed
    module_code = ""
    try:
        module_code = backend.load_c_file(source)
    except Exception:
        module_code = ""
    raw = backend._derive_module_display_name(source, module_code) if module_code else (
        backend._guess_cn_from_ident(module_name_raw, glossary=backend.DOMAIN_GLOSSARY) or module_name_raw
    )
    if text_utils._contains_cjk(raw):
        module_name = backend._normalize_function_cn_title(
            raw, func_name=module_name_raw, comment_desc=raw
        )
    else:
        module_name = raw
    return module_name_raw, module_name


def count_single_file_progress_total(
    func_list: Sequence[dict],
    *,
    resume_func_pos: int,
    only_with_comment: bool,
) -> int:
    total = 0
    for idx, func_data in enumerate(func_list or (), start=1):
        if idx < resume_func_pos:
            continue
        if not should_include_function(func_data, only_with_comment=only_with_comment):
            continue
        total += 1
    return int(total)


def build_single_file_table_payload(
    func_list: Sequence[dict],
    *,
    resume_func_pos: int,
    module_id: str,
    module_name: str,
    only_with_comment: bool,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    entries: list[dict] = []
    func_rows: list[dict] = []
    for idx, func_data in enumerate(func_list or (), start=1):
        if idx < resume_func_pos:
            continue
        if not should_include_function(func_data, only_with_comment=only_with_comment):
            continue
        csu_name = _registered_function_title(func_data, backend_module=backend)
        entries.append(
            {
                "csu_name": csu_name,
                "csu_id": f"{backend.normalize_req_prefix(module_id)}_{len(entries) + 1:03d}",
            }
        )
        func_rows.append(
            {
                "index": len(func_rows) + 1,
                "name": csu_name,
                "prototype": backend._format_func_prototype(func_data.get("func_info") or {}),
            }
        )
    return {
        "entries": entries,
        "unit_func_table": {
            "unit_name": module_name,
            "func_rows": func_rows,
        },
    }


def build_single_file_tasks(
    func_list: Sequence[dict],
    *,
    resume_func_pos: int,
    func_index: int,
    module_id: str,
    source: str,
    only_with_comment: bool,
) -> dict[str, Any]:
    tasks: list[dict] = []
    marker_updates: dict[str, Any] = {}
    next_func_index = int(func_index)
    for idx, func_data in enumerate(func_list or (), start=1):
        if idx < resume_func_pos:
            continue
        if not should_include_function(func_data, only_with_comment=only_with_comment):
            marker_updates["func_pos"] = idx + 1
            continue
        func_name = ((func_data.get("func_info") or {}).get("func_name") or "")
        tasks.append(
            {
                "func_data": func_data,
                "module_req_prefix": module_id,
                "index": next_func_index,
                "func_pos": idx,
                "func_name": func_name,
                "file": source,
                "module_id": module_id,
            }
        )
        next_func_index += 1
    return {"tasks": tasks, "resume_marker_updates": marker_updates}


def build_single_export_task(
    target: dict[str, Any],
    *,
    target_pos: int,
    func_name: str,
    module_id: str,
    source: str,
) -> dict[str, Any]:
    return {
        "func_data": target,
        "module_req_prefix": module_id,
        "index": 1,
        "func_pos": int(target_pos or 1),
        "func_name": func_name,
        "file": source,
        "module_id": module_id,
    }


def collect_project_c_files_by_layer(project_dir: str, cfg, *, backend_module=None) -> tuple[str, list[str], list[str], list[str]]:
    backend = backend_module or legacy_backend()
    if not project_dir or not os.path.isdir(project_dir):
        return "", [], [], []
    if backend.stop_requested(cfg):
        return "", [], [], []

    src_dir = None
    try:
        for name in os.listdir(project_dir):
            candidate = os.path.join(project_dir, name)
            if os.path.isdir(candidate) and name.lower() == "src":
                src_dir = candidate
                break
    except Exception:
        src_dir = None
    if not src_dir:
        for dirpath, dirnames, _ in scanner_utils.walk_filtered(project_dir, exclude_dirs=cfg.exclude_dirs):
            if backend.stop_requested(cfg):
                return "", [], [], []
            for dirname in dirnames:
                if dirname.lower() == "src":
                    src_dir = os.path.join(dirpath, dirname)
                    break
            if src_dir:
                break
    if not src_dir:
        return "", [], [], []

    app_files: list[str] = []
    mid_files: list[str] = []
    drv_files: list[str] = []
    mid_keys = [str(item).strip().lower() for item in (getattr(cfg, "mid_dir_keywords", None) or ()) if str(item).strip()]
    drv_keys = [str(item).strip().lower() for item in (getattr(cfg, "drv_dir_keywords", None) or ()) if str(item).strip()]
    for dirpath, _, files in scanner_utils.walk_filtered(src_dir, exclude_dirs=cfg.exclude_dirs):
        if backend.stop_requested(cfg):
            break
        rel = os.path.relpath(dirpath, src_dir)
        rel_lower = rel.lower()
        for filename in files:
            if backend.stop_requested(cfg):
                break
            if not filename.lower().endswith(".c"):
                continue
            full_path = os.path.join(dirpath, filename)
            if any(key in rel_lower for key in mid_keys):
                mid_files.append(full_path)
            elif any(key in rel_lower for key in drv_keys):
                drv_files.append(full_path)
            else:
                app_files.append(full_path)

    app_files.sort()
    mid_files.sort()
    drv_files.sort()
    return src_dir, app_files, mid_files, drv_files


def apply_project_file_order_override(
    app_files: list[str],
    mid_files: list[str],
    drv_files: list[str],
    order: Optional[dict[str, list[Any]]],
) -> tuple[list[str], list[str], list[str]]:
    if not order:
        return app_files, mid_files, drv_files
    explicit_only = bool(order.get("_explicit")) if isinstance(order, dict) else False

    def _flatten_entries(xs: Sequence[Any]) -> list[str]:
        out: list[str] = []
        for item in xs or []:
            if isinstance(item, str):
                out.append(item)
                continue
            if isinstance(item, dict):
                files = item.get("files") or item.get("paths") or item.get("c_files") or []
                if isinstance(files, (list, tuple)):
                    for path in files:
                        if isinstance(path, str) and path.strip():
                            out.append(path.strip())
        return out

    def norm_list(xs: Sequence[Any]) -> list[str]:
        out: list[str] = []
        seen: set[str] = set()
        for path in _flatten_entries(xs):
            abs_path = os.path.abspath(str(path))
            if abs_path in seen:
                continue
            seen.add(abs_path)
            out.append(abs_path)
        return out

    all_files = {os.path.abspath(path) for path in (app_files + mid_files + drv_files)}
    ordered_app = [path for path in norm_list(order.get("app") or []) if path in all_files]
    ordered_mid = [path for path in norm_list(order.get("mid") or []) if path in all_files]
    ordered_drv = [path for path in norm_list(order.get("drv") or []) if path in all_files]
    assigned = set(ordered_app) | set(ordered_mid) | set(ordered_drv)

    def append_missing_keep_layer(ordered: list[str], originals: list[str]) -> list[str]:
        seen = {os.path.abspath(path) for path in ordered}
        for path in originals:
            abs_path = os.path.abspath(path)
            if abs_path in assigned:
                continue
            if abs_path not in seen:
                ordered.append(abs_path)
                seen.add(abs_path)
        return ordered

    if explicit_only:
        return ordered_app, ordered_mid, ordered_drv
    return (
        append_missing_keep_layer(ordered_app, app_files),
        append_missing_keep_layer(ordered_mid, mid_files),
        append_missing_keep_layer(ordered_drv, drv_files),
    )


def build_project_modules_from_order(
    *,
    app_files: list[str],
    mid_files: list[str],
    drv_files: list[str],
    order: Optional[dict[str, list[Any]]],
) -> tuple[list[dict], list[dict], list[dict]]:
    all_files = {os.path.abspath(path) for path in (app_files + mid_files + drv_files)}

    def norm_files(xs: Sequence[Any]) -> list[str]:
        out: list[str] = []
        seen: set[str] = set()
        for path in xs or []:
            if not isinstance(path, str):
                continue
            abs_path = os.path.abspath(path)
            if abs_path in seen:
                continue
            seen.add(abs_path)
            if abs_path in all_files:
                out.append(abs_path)
        return out

    def parse_entries(entries: Sequence[Any]) -> list[dict]:
        modules: list[dict] = []
        for item in entries or []:
            if isinstance(item, str):
                abs_path = os.path.abspath(item)
                if abs_path in all_files:
                    modules.append({"name": None, "files": [abs_path]})
                continue
            if isinstance(item, dict):
                name = (item.get("module") or item.get("name") or item.get("module_name") or "").strip() or None
                files = item.get("files") or item.get("paths") or item.get("c_files") or []
                files_norm = norm_files(files if isinstance(files, (list, tuple)) else [])
                if files_norm:
                    modules.append({"name": name, "files": files_norm})
        return modules

    if not order:
        def as_single(files: list[str]) -> list[dict]:
            return [{"name": None, "files": [os.path.abspath(path)]} for path in files]

        return as_single(app_files), as_single(mid_files), as_single(drv_files)

    explicit_only = bool(order.get("_explicit")) if isinstance(order, dict) else False
    ordered_app = parse_entries(order.get("app") or [])
    ordered_mid = parse_entries(order.get("mid") or [])
    ordered_drv = parse_entries(order.get("drv") or [])

    assigned: set[str] = set()
    for module in (ordered_app + ordered_mid + ordered_drv):
        for path in module.get("files") or []:
            assigned.add(os.path.abspath(str(path)))

    def append_missing_keep_layer(modules: list[dict], originals: list[str]) -> list[dict]:
        for path in originals or []:
            abs_path = os.path.abspath(path)
            if abs_path in assigned:
                continue
            modules.append({"name": None, "files": [abs_path]})
            assigned.add(abs_path)
        return modules

    if explicit_only:
        return ordered_app, ordered_mid, ordered_drv
    return (
        append_missing_keep_layer(ordered_app, app_files),
        append_missing_keep_layer(ordered_mid, mid_files),
        append_missing_keep_layer(ordered_drv, drv_files),
    )


def plan_project_source_layout(
    root_dir: str,
    cfg,
    *,
    order_override: Optional[dict[str, list[Any]]] = None,
    backend_module=None,
) -> dict[str, Any]:
    src_dir, app_files, mid_files, drv_files = collect_project_c_files_by_layer(
        root_dir,
        cfg,
        backend_module=backend_module,
    )
    app_modules = [{"name": None, "files": [os.path.abspath(path)]} for path in (app_files or [])]
    mid_modules = [{"name": None, "files": [os.path.abspath(path)]} for path in (mid_files or [])]
    drv_modules = [{"name": None, "files": [os.path.abspath(path)]} for path in (drv_files or [])]
    if order_override:
        app_modules, mid_modules, drv_modules = build_project_modules_from_order(
            app_files=app_files,
            mid_files=mid_files,
            drv_files=drv_files,
            order=order_override,
        )
        app_files = [path for module in app_modules for path in (module.get("files") or [])]
        mid_files = [path for module in mid_modules for path in (module.get("files") or [])]
        drv_files = [path for module in drv_modules for path in (module.get("files") or [])]
    return {
        "src_dir": src_dir,
        "app_files": list(app_files or []),
        "mid_files": list(mid_files or []),
        "drv_files": list(drv_files or []),
        "app_modules": list(app_modules or []),
        "mid_modules": list(mid_modules or []),
        "drv_modules": list(drv_modules or []),
        "total_files": len(app_files or []) + len(mid_files or []) + len(drv_files or []),
        "total_modules": len(app_modules or []) + len(mid_modules or []) + len(drv_modules or []),
    }


def initialize_project_run_state(
    output: str,
    cfg,
    *,
    continuing: bool,
    resume_state: Optional[dict],
    backend_module=None,
    runtime_module=None,
    render_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    if continuing:
        if runtime_module is None:
            from . import runtime as runtime_module
        resume_info = runtime_module.normalize_project_resume_state(resume_state)
        if not os.path.exists(output):
            raise backend.RenderError("未找到上次生成的文档，无法继续。")
        return {
            "doc": backend.Document(output),
            "placeholder": None,
            "body_start_idx": 0,
            "resume_info": resume_info,
            "start_layer_idx": resume_info.layer_index,
            "start_file_idx": resume_info.file_index,
            "start_func_pos": resume_info.func_pos,
            "start_func_index": resume_info.func_index,
            "module_counter": resume_info.module_counter,
            "layer_heading_started": resume_info.layer_started,
            "module_started": resume_info.module_started,
            "resume_module_id": resume_info.module_id,
        }

    if render_module is None:
        from . import render as render_module
    doc_state = render_module.init_generation_document(
        cfg,
        main_heading="CSCI详细设计",
        heading_level=1,
        backend_module=backend,
    )
    return {
        "doc": doc_state["doc"],
        "placeholder": doc_state.get("placeholder"),
        "body_start_idx": int(doc_state.get("body_start_idx") or 0),
        "resume_info": None,
        "start_layer_idx": 0,
        "start_file_idx": 0,
        "start_func_pos": 1,
        "start_func_index": 1,
        "module_counter": 1,
        "layer_heading_started": False,
        "module_started": False,
        "resume_module_id": None,
    }


def _flatten_project_module_files(layer_modules: Sequence[Sequence[dict]]) -> list[str]:
    ordered_files: list[str] = []
    for modules in layer_modules:
        for module in modules or ():
            for path in (module.get("files") or []):
                text = str(path).strip()
                if text:
                    ordered_files.append(os.path.abspath(text))
    return ordered_files


def preprocess_project_files(
    file_list: Sequence[str],
    *,
    project_root: str,
    cfg,
    prefilter: bool,
    backend_module=None,
) -> dict[str, dict]:
    backend = backend_module or legacy_backend()
    results: dict[str, dict] = {}
    if not file_list:
        return results

    worker_count = int(getattr(cfg, "preprocess_workers", 0) or 0)
    if worker_count <= 0:
        worker_count = min(4, (os.cpu_count() or 1))
    worker_count = max(1, worker_count)

    def _job(path: str) -> tuple[str, list[dict], Optional[str], Optional[Exception]]:
        try:
            func_list, reason = _prepare_func_list_for_c_file(
                path,
                project_root=project_root,
                cfg=cfg,
                prefilter=prefilter,
                backend_module=backend,
            )
            return path, func_list, reason, None
        except Exception as exc:
            return path, [], None, exc

    if worker_count <= 1 or len(file_list) <= 1:
        for path in file_list:
            c_path, func_list, reason, err = _job(path)
            results[c_path] = {"func_list": func_list, "skip_reason": reason, "error": err}
        return results

    executor = concurrent.futures.ThreadPoolExecutor(max_workers=worker_count)
    stop_now = False
    try:
        futures = {executor.submit(_job, path): path for path in file_list}
        pending = set(futures)
        while pending:
            if backend.stop_requested(cfg):
                stop_now = True
                break
            done, pending = concurrent.futures.wait(
                pending,
                timeout=0.2,
                return_when=concurrent.futures.FIRST_COMPLETED,
            )
            for future in done:
                if backend.stop_requested(cfg):
                    stop_now = True
                    break
                c_path, func_list, reason, err = future.result()
                results[c_path] = {"func_list": func_list, "skip_reason": reason, "error": err}
            if stop_now:
                break
    finally:
        if 'futures' in locals():
            for future in futures:
                if not future.done():
                    future.cancel()
        executor.shutdown(wait=(not stop_now))
    return results


def prepare_project_progress(
    *,
    cfg,
    continuing: bool,
    preprocessed: dict[str, dict],
    app_modules: Sequence[dict],
    mid_modules: Sequence[dict],
    drv_modules: Sequence[dict],
    layer_sets: Sequence[tuple[str, str, Sequence[dict]]],
    start_layer_idx: int,
    start_file_idx: int,
    start_func_pos: int,
    module_started: bool,
    root_dir: str,
    prefilter: bool,
    backend_module=None,
) -> Optional[dict[str, Any]]:
    backend = backend_module or legacy_backend()
    want_progress = callable(getattr(cfg, "gui_event", None))
    if not want_progress:
        return None

    progress_pre = preprocessed
    if continuing:
        ordered_files = _flatten_project_module_files((app_modules, mid_modules, drv_modules))
        progress_pre = preprocess_project_files(
            ordered_files,
            project_root=root_dir,
            cfg=cfg,
            prefilter=prefilter,
            backend_module=backend,
        )

    total_work = count_project_progress_total(
        layer_sets,
        start_layer_idx=start_layer_idx,
        start_file_idx=start_file_idx,
        start_func_pos=start_func_pos,
        continuing=continuing,
        module_started=module_started,
        progress_pre=progress_pre,
        root_dir=root_dir,
        cfg=cfg,
        prefilter=prefilter,
        backend_module=backend,
    )
    return {
        "payload": {
            "type": "progress_init",
            "mode": "project",
            "total": int(total_work),
            "unit": "函数",
        },
        "progress_pre": progress_pre,
    }


def prepare_project_layer_iteration(
    doc,
    *,
    cfg,
    layer_idx: int,
    layer_name: str,
    files: Sequence[dict],
    continuing: bool,
    start_layer_idx: int,
    layer_heading_started: bool,
    resume_marker: dict[str, Any],
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    marker = dict(resume_marker or {})
    marker["layer_index"] = layer_idx
    if backend.stop_requested(cfg):
        cfg.resume_state = dict(marker)
        backend.vlog(cfg, "收到停止请求，结束工程生成。")
        return {
            "stopped": True,
            "resume_marker": marker,
            "layer_heading_started": False,
            "empty_layer": False,
        }

    current_layer_heading = bool(layer_heading_started) if (continuing and layer_idx == start_layer_idx) else False
    if not current_layer_heading:
        doc.add_paragraph(layer_name, style=backend.pick_heading_style(doc, 2))
        current_layer_heading = True
    marker["layer_started"] = current_layer_heading
    if not files:
        doc.add_paragraph("无对应C文件。", style="Normal")
    return {
        "stopped": False,
        "resume_marker": marker,
        "layer_heading_started": current_layer_heading,
        "empty_layer": not bool(files),
    }


def prepare_project_module_section(
    doc,
    func_list_all: Sequence[dict],
    *,
    reuse_module: bool,
    module_display: str,
    module_id: str,
    collect_unit_func_lists: bool,
    backend_module=None,
    render_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    if reuse_module:
        return {
            "module_table": None,
            "unit_func_table": None,
            "module_started": False,
        }

    table_payload = build_project_module_table_payload(
        func_list_all,
        module_id=module_id,
        module_display=module_display,
        include_unit_func_table=collect_unit_func_lists,
        backend_module=backend,
    )
    if render_module is None:
        from . import render as render_module
    module_table = render_module.add_module_section(
        doc,
        module_display,
        module_id,
        table_payload["entries"],
        backend_module=backend,
    )
    return {
        "module_table": module_table,
        "unit_func_table": table_payload.get("unit_func_table") if collect_unit_func_lists else None,
        "module_started": True,
    }


def prepare_project_file_iteration(
    mod: Any,
    *,
    cfg,
    file_idx: int,
    resume_marker: dict[str, Any],
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    marker = dict(resume_marker or {})
    marker["file_index"] = int(file_idx)
    if backend.stop_requested(cfg):
        cfg.resume_state = dict(marker)
        backend.vlog(cfg, "收到停止请求，终止后续文件处理。")
        return {
            "stopped": True,
            "resume_marker": marker,
            "c_paths": [],
        }

    module = mod if isinstance(mod, dict) else {"name": None, "files": []}
    c_paths = [
        os.path.abspath(str(path).strip())
        for path in (module.get("files") or [])
        if path is not None and str(path).strip()
    ]
    return {
        "stopped": False,
        "resume_marker": marker,
        "c_paths": c_paths,
    }


def log_project_module_plan(
    *,
    cfg,
    layer_name: str,
    file_idx: int,
    module_name: str,
    module_id: str,
    c_paths: Sequence[str],
    root_dir: str,
    backend_module=None,
) -> None:
    backend = backend_module or legacy_backend()
    backend.vlog(
        cfg,
        f"[{layer_name}] 模块{int(file_idx) + 1:03d}: {module_name} <- {len(c_paths)} 文件 -> 模块ID {module_id}",
    )
    try:
        rels = [os.path.relpath(path, root_dir) for path in c_paths]
        if len(rels) <= 6:
            for rel_path in rels:
                backend.vlog(cfg, f"  - {rel_path}")
    except Exception:
        pass


def finalize_project_module_iteration(
    *,
    cfg,
    processed_tasks: int,
    module_tasks: Sequence[dict],
    resume_marker: dict[str, Any],
    file_idx: int,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    marker = dict(resume_marker or {})
    if backend.stop_requested(cfg) and processed_tasks < len(module_tasks):
        next_task = module_tasks[processed_tasks]
        marker = record_project_task_resume_progress(
            marker,
            next_task,
            advance=False,
        )
        cfg.resume_state = dict(marker)
        backend.vlog(cfg, "收到停止请求，提前结束当前模块。")
        return {
            "stopped": True,
            "resume_marker": marker,
            "start_func_pos": 1,
            "start_func_index": 1,
            "resume_module_id": None,
            "module_started": False,
        }

    marker = advance_project_resume_after_module(
        marker,
        file_idx=file_idx,
    )
    return {
        "stopped": False,
        "resume_marker": marker,
        "start_func_pos": 1,
        "start_func_index": 1,
        "resume_module_id": None,
        "module_started": False,
    }


def execute_project_module_tasks(
    doc,
    module_tasks: Sequence[dict],
    *,
    cfg,
    root_dir: str,
    module_table: Any,
    collect_unit_rows: bool,
    unit_rows: list[dict],
    unit_index: int,
    resume_marker: dict[str, Any],
    designs: Optional[list] = None,
    backend_module=None,
    render_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    marker = dict(resume_marker or {})
    processed_tasks = 0
    wrote_any = False
    next_unit_index = int(unit_index)
    next_func_index = int(marker.get("func_index", 1) or 1)
    if render_module is None:
        from . import render as render_module

    def _on_submit(task: dict) -> None:
        backend.gui_event(cfg, build_project_func_start_event(task, root_dir=root_dir))

    _parallel = should_parallel_build_design(cfg) and len(module_tasks) > 1

    try:
        if _parallel:
            _design_iter = _iter_function_design_results(
                module_tasks, cfg, on_submit=_on_submit, backend_module=backend,
            )
        else:
            _design_iter = None

        for task in module_tasks:
            if backend.stop_requested(cfg):
                backend.vlog(cfg, "[停止] 检测到停止信号，退出任务循环")
                break

            # Compute render cache key before building design
            cfg._current_render_func_data = task.get("func_data") or {}
            _rkey = render_module._render_cache_key(
                task.get("source_file", ""),
                task.get("func_name", ""),
                (task.get("func_data") or {}).get("body", ""),
                _registered_function_title(task.get("func_data") or {}, backend_module=backend),
                ai_mode=int(getattr(cfg, "ai_mode", 0) or 0),
            )
            _body_start = len(list(doc.element.body))
            _has_revision = bool(
                getattr(cfg, "extra_params", None)
                and (
                    (cfg.extra_params or {}).get("revision_profile")
                    or (cfg.extra_params or {}).get("revision_profile_json")
                )
            )

            design = None
            err = None

            if not _has_revision and render_module.try_replay_rendered(doc, _rkey):
                # Cache hit: skip design building entirely (saves AI calls)
                pass
            elif _parallel:
                # Parallel path: get design from generator
                try:
                    task_r, design, err = next(_design_iter)
                except StopIteration:
                    break
                if design is not None:
                    render_module.render_function_design(doc, design, cfg)
                    render_module.capture_rendered_elements(doc, _body_start, _rkey)
            else:
                # Sequential path: build design on demand
                _on_submit(task)
                try:
                    design = run_function_design_task(task, cfg, backend_module=backend)
                except Exception as exc:
                    err = exc
                    if not _continue_on_function_error(cfg, backend_module=backend):
                        raise backend.FunctionBuildTaskError(task, exc) from exc
                    design = _build_failed_function_design(task, exc, cfg, backend_module=backend)
                if design is not None:
                    render_module.render_function_design(doc, design, cfg)
                    render_module.capture_rendered_elements(doc, _body_start, _rkey)

            try:
                if design is not None:
                    _collect_review_function(cfg, design, task)
                    _collect_design_workspace_pair(cfg, design, task)
                    _doc_path = getattr(cfg, "_output_doc_path", "")
                    _csu_id = f"{task.get('module_id', '')}_{int(task.get('index', 0) or 0):03d}"
                    if _doc_path and _csu_id and design is not None:
                        save_design_snapshot(_doc_path, _csu_id, design)
            finally:
                try:
                    cfg._current_render_func_data = None
                except Exception:
                    pass
            # 收集 design 用于术语一致性检查
            if designs is not None and design is not None:
                designs.append({
                    "func_name": task.get("func_name", ""),
                    "title": getattr(design, "title", ""),
                    "source_file": task.get("source_file", ""),
                    "io_elements": [{"ident": e.ident, "name": e.name} for e in (design.io_elements or [])],
                    "local_elements": [{"ident": e.ident, "name": e.name} for e in (design.local_elements or [])],
                    "logic_lines": list(design.logic_lines or []),
                })
            backend.gui_event(
                cfg,
                build_project_func_end_event(
                    task,
                    root_dir=root_dir,
                    design=design,
                    error=err,
                    ok=(err is None),
                ),
            )
            update_project_module_table_title(module_table, task, design)
            if collect_unit_rows:
                unit_rows.append(
                    build_project_unit_row(
                        task,
                        design,
                        unit_index=next_unit_index,
                        backend_module=backend,
                    )
                )
                next_unit_index += 1
            if backend.should_log_step(cfg, int(task["index"])):
                backend.vlog(cfg, f"    已生成函数 {int(task['index'])}: {task['func_name']}")
            wrote_any = True
            processed_tasks += 1
            next_func_index = int(task["index"]) + 1
            marker = record_project_task_resume_progress(
                marker,
                task,
                advance=True,
            )
    except backend.FunctionBuildTaskError as exc:
        task = exc.task
        marker = record_project_task_resume_progress(
            marker,
            task,
            advance=False,
        )
        cfg.resume_state = dict(marker)
        backend.gui_event(
            cfg,
            build_project_func_end_event(
                task,
                root_dir=root_dir,
                error=exc.cause,
                ok=False,
            ),
        )
        raise exc.cause from exc

    return {
        "processed_tasks": processed_tasks,
        "wrote_any": wrote_any,
        "resume_marker": marker,
        "unit_index": next_unit_index,
        "func_index": next_func_index,
    }


def execute_single_file_tasks(
    doc,
    tasks: Sequence[dict],
    *,
    cfg,
    source: str,
    module_table: Any,
    resume_marker: dict[str, Any],
    backend_module=None,
    render_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    marker = dict(resume_marker or {})
    processed_tasks = 0
    wrote_any = False
    if render_module is None:
        from . import render as render_module

    def _on_submit(task: dict) -> None:
        backend.gui_event(cfg, build_single_file_func_start_event(task, mode="single", source=source))

    _parallel = should_parallel_build_design(cfg) and len(tasks) > 1

    try:
        if _parallel:
            _design_iter = _iter_function_design_results(
                tasks, cfg, on_submit=_on_submit, backend_module=backend,
            )
        else:
            _design_iter = None

        for task in tasks:
            if backend.stop_requested(cfg):
                backend.vlog(cfg, "[停止] 检测到停止信号，退出任务循环")
                break

            cfg._current_render_func_data = task.get("func_data") or {}
            _rkey = render_module._render_cache_key(
                task.get("source_file", ""),
                task.get("func_name", ""),
                (task.get("func_data") or {}).get("body", ""),
                _registered_function_title(task.get("func_data") or {}, backend_module=backend),
                ai_mode=int(getattr(cfg, "ai_mode", 0) or 0),
            )
            _body_start = len(list(doc.element.body))
            _has_revision = bool(
                getattr(cfg, "extra_params", None)
                and (
                    (cfg.extra_params or {}).get("revision_profile")
                    or (cfg.extra_params or {}).get("revision_profile_json")
                )
            )

            design = None
            err = None

            if not _has_revision and render_module.try_replay_rendered(doc, _rkey):
                pass
            elif _parallel:
                try:
                    task_r, design, err = next(_design_iter)
                except StopIteration:
                    break
                if design is not None:
                    render_module.render_function_design(doc, design, cfg)
                    render_module.capture_rendered_elements(doc, _body_start, _rkey)
            else:
                _on_submit(task)
                try:
                    design = run_function_design_task(task, cfg, backend_module=backend)
                except Exception as exc:
                    err = exc
                    if not _continue_on_function_error(cfg, backend_module=backend):
                        raise backend.FunctionBuildTaskError(task, exc) from exc
                    design = _build_failed_function_design(task, exc, cfg, backend_module=backend)
                if design is not None:
                    render_module.render_function_design(doc, design, cfg)
                    render_module.capture_rendered_elements(doc, _body_start, _rkey)

            try:
                if design is not None:
                    _collect_review_function(cfg, design, task)
                    _collect_design_workspace_pair(cfg, design, task)
                    _doc_path = getattr(cfg, "_output_doc_path", "")
                    _csu_id = f"{task.get('module_id', '')}_{int(task.get('index', 0) or 0):03d}"
                    if _doc_path and _csu_id and design is not None:
                        save_design_snapshot(_doc_path, _csu_id, design)
            finally:
                try:
                    cfg._current_render_func_data = None
                except Exception:
                    pass
            backend.gui_event(
                cfg,
                build_single_file_func_end_event(
                    task,
                    mode="single",
                    source=source,
                    design=design,
                    error=err,
                    ok=(err is None),
                ),
            )
            update_project_module_table_title(module_table, task, design)
            if backend.should_log_step(cfg, int(task["index"])):
                backend.vlog(cfg, f"已生成函数 {int(task['index'])}：{task['func_name']}")
            wrote_any = True
            processed_tasks += 1
            marker = record_single_file_resume_progress(
                marker,
                task,
                advance=True,
            )
    except backend.FunctionBuildTaskError as exc:
        task = exc.task
        marker = record_single_file_resume_progress(
            marker,
            task,
            advance=False,
        )
        cfg.resume_state = dict(marker)
        backend.gui_event(
            cfg,
            build_single_file_func_end_event(
                task,
                mode="single",
                source=source,
                error=exc.cause,
                ok=False,
            ),
        )
        raise exc.cause from exc

    return {
        "processed_tasks": processed_tasks,
        "wrote_any": wrote_any,
        "resume_marker": marker,
    }


def execute_single_export_task(
    doc,
    task: dict[str, Any],
    *,
    cfg,
    source: str,
    module_table: Any,
    backend_module=None,
    render_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    if render_module is None:
        from . import render as render_module

    backend.gui_event(cfg, build_single_file_func_start_event(task, mode="single_export", source=source))
    try:
        design = run_single_export_design(
            task,
            cfg,
            func_name=str(task.get("func_name") or ""),
            backend_module=backend,
        )
        try:
            cfg._current_render_func_data = task.get("func_data") or {}
            # docx 渲染缓存：未变函数直接复用缓存的 XML 元素
            _rkey = render_module._render_cache_key(
                task.get("source_file", ""),
                task.get("func_name", ""),
                (task.get("func_data") or {}).get("body", ""),
                _registered_function_title(task.get("func_data") or {}, backend_module=backend),
                ai_mode=int(getattr(cfg, "ai_mode", 0) or 0),
            )
            _body_start = len(list(doc.element.body))
            # Revision profiles can change rendered text without source-body changes.
            _has_revision = bool(
                getattr(cfg, "extra_params", None)
                and (
                    (cfg.extra_params or {}).get("revision_profile")
                    or (cfg.extra_params or {}).get("revision_profile_json")
                )
            )
            if _has_revision or not render_module.try_replay_rendered(doc, _rkey):
                if _has_revision:
                    # Drop stale cache entry so later capture stores revised XML.
                    try:
                        render_module._RENDER_CACHE.pop(_rkey, None)
                    except Exception:
                        pass
                render_module.render_function_design(doc, design, cfg)
                render_module.capture_rendered_elements(doc, _body_start, _rkey)
            _collect_review_function(cfg, design, task)
            _collect_design_workspace_pair(cfg, design, task)
            _doc_path = getattr(cfg, "_output_doc_path", "")
            _csu_id = f"{task.get('module_id', '')}_{int(task.get('index', 0) or 0):03d}"
            if _doc_path and _csu_id and design is not None:
                save_design_snapshot(_doc_path, _csu_id, design)
        finally:
            try:
                cfg._current_render_func_data = None
            except Exception:
                pass
    except Exception as exc:
        backend.gui_event(
            cfg,
            build_single_file_func_end_event(
                task,
                mode="single_export",
                source=source,
                error=exc,
                ok=False,
            ),
        )
        raise

    backend.gui_event(
        cfg,
        build_single_file_func_end_event(
            task,
            mode="single_export",
            source=source,
            design=design,
            ok=True,
        ),
    )
    updated = update_project_module_table_title(module_table, task, design)
    return {
        "design": design,
        "updated_module_table": updated,
        "wrote_any": True,
    }


def finalize_single_file_task_iteration(
    *,
    cfg,
    processed_tasks: int,
    tasks: Sequence[dict],
    resume_marker: dict[str, Any],
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    marker = dict(resume_marker or {})
    if backend.stop_requested(cfg) and processed_tasks < len(tasks):
        next_task = tasks[processed_tasks]
        marker = record_single_file_resume_progress(
            marker,
            next_task,
            advance=False,
        )
        cfg.resume_state = dict(marker)
        backend.vlog(cfg, "收到停止请求，结束单文件生成。")
        return {
            "stopped": True,
            "resume_marker": marker,
        }
    return {
        "stopped": False,
        "resume_marker": marker,
    }


def get_ordered_project_c_files(project_dir: str, cfg, *, backend_module=None) -> list[str]:
    if not project_dir or not os.path.isdir(project_dir):
        return []

    _, app_files, mid_files, drv_files = collect_project_c_files_by_layer(
        project_dir,
        cfg,
        backend_module=backend_module,
    )
    if not (app_files or mid_files or drv_files):
        return []

    app_files, mid_files, drv_files = apply_project_file_order_override(
        app_files,
        mid_files,
        drv_files,
        getattr(cfg, "project_file_order", None),
    )
    return list(app_files or []) + list(mid_files or []) + list(drv_files or [])


def compute_project_module_id(project_dir: str, c_file: str, cfg, *, backend_module=None) -> Optional[str]:
    backend = backend_module or legacy_backend()
    if not project_dir or not c_file or not os.path.isdir(project_dir):
        return None

    ordered = get_ordered_project_c_files(project_dir, cfg, backend_module=backend)
    if not ordered:
        return None

    target = os.path.abspath(c_file)
    for module_counter, path in enumerate(ordered, start=1):
        if os.path.abspath(path) == target:
            return f"{backend.normalize_req_prefix(cfg.req_id_prefix)}_{module_counter:03d}"
    return None


def _extend_project_unit_rows(
    unit_rows: list[dict],
    func_list: Sequence[dict],
    *,
    module_id: str,
    c_path: str,
    unit_index: int,
    backend_module=None,
) -> int:
    backend = backend_module or legacy_backend()
    next_unit_index = int(unit_index)
    for fd in func_list or ():
        func_info = fd.get("func_info") or {}
        csu_name = _registered_function_title(fd, backend_module=backend)
        prototype = backend._format_func_prototype(func_info or fd)
        req_id = f"{module_id}_{next_unit_index:03d}"
        unit_rows.append(
            {
                "index": next_unit_index,
                "name": csu_name,
                "prototype": prototype,
                "req_id": req_id,
                "location": os.path.basename(c_path),
                "status": "新研",
                "purpose": csu_name,
            }
        )
        next_unit_index += 1
    return next_unit_index


def collect_project_module_entries_and_units(
    project_dir: str,
    cfg,
    prefilter: bool,
    *,
    backend_module=None,
) -> tuple[list[tuple[str, int]], list[dict]]:
    backend = backend_module or legacy_backend()
    ordered = get_ordered_project_c_files(project_dir, cfg, backend_module=backend)
    module_entries: list[tuple[str, int]] = []
    unit_rows: list[dict] = []
    if not ordered:
        return module_entries, unit_rows

    unit_index = 1
    for module_counter, c_path in enumerate(ordered, start=1):
        func_list, _ = _prepare_func_list_for_c_file(
            c_path,
            project_root=project_dir,
            cfg=cfg,
            prefilter=prefilter,
            need_symbol_map=False,
            backend_module=backend,
        )
        module_id = f"{backend.normalize_req_prefix(cfg.req_id_prefix)}_{module_counter:03d}"
        csu_count = len(func_list or [])
        if csu_count > 0:
            module_entries.append((module_id, csu_count))
        unit_index = _extend_project_unit_rows(
            unit_rows,
            func_list or [],
            module_id=module_id,
            c_path=c_path,
            unit_index=unit_index,
            backend_module=backend,
        )

    return module_entries, unit_rows


def collect_project_module_tables_data(
    project_dir: str,
    cfg,
    prefilter: bool,
    *,
    backend_module=None,
) -> tuple[list[dict], list[dict]]:
    backend = backend_module or legacy_backend()
    ordered = get_ordered_project_c_files(project_dir, cfg, backend_module=backend)
    modules: list[dict] = []
    unit_rows: list[dict] = []
    if not ordered:
        return modules, unit_rows

    unit_index = 1
    for module_counter, c_path in enumerate(ordered, start=1):
        func_list, _ = _prepare_func_list_for_c_file(
            c_path,
            project_root=project_dir,
            cfg=cfg,
            prefilter=prefilter,
            need_symbol_map=False,
            backend_module=backend,
        )
        module_id = f"{backend.normalize_req_prefix(cfg.req_id_prefix)}_{module_counter:03d}"
        module_name_raw = os.path.splitext(os.path.basename(c_path))[0]
        module_display = backend._guess_cn_from_ident(module_name_raw, glossary=backend.DOMAIN_GLOSSARY) or module_name_raw

        csu_entries: list[dict] = []
        for fd in func_list or ():
            csu_name = _registered_function_title(fd, backend_module=backend)
            csu_entries.append({"csu_name": csu_name})

        modules.append(
            {
                "module_id": module_id,
                "module_name": module_display,
                "csu_entries": csu_entries,
            }
        )
        unit_index = _extend_project_unit_rows(
            unit_rows,
            func_list or [],
            module_id=module_id,
            c_path=c_path,
            unit_index=unit_index,
            backend_module=backend,
        )

    return modules, unit_rows


def resolve_project_resume_checkpoint(
    layer_sets: Sequence[tuple[str, str, Sequence[dict]]],
    resume_info: ProjectResumeState,
    *,
    root_dir: str,
    cfg,
    prefilter: bool,
    backend_module=None,
) -> dict[str, int]:
    backend = backend_module or legacy_backend()
    start_layer_idx = resume_info.layer_index
    start_file_idx = resume_info.file_index
    start_func_pos = resume_info.func_pos
    resume_file = resume_info.file
    resume_func_name = resume_info.func_name
    resume_module_files = list(resume_info.module_files)
    if not (resume_file or resume_module_files):
        return {
            "layer_index": start_layer_idx,
            "file_index": start_file_idx,
            "func_pos": start_func_pos,
        }

    files_hint = {os.path.abspath(p) for p in resume_module_files if str(p).strip()}
    file_hint = os.path.abspath(resume_file) if resume_file else ""
    match: Optional[tuple[int, int, list[str]]] = None
    for layer_idx, (_, _layer_name, modules) in enumerate(layer_sets):
        for file_idx, mod in enumerate(modules or ()):
            c_paths = _normalize_project_module_paths(mod)
            if file_hint and file_hint in c_paths:
                match = (layer_idx, file_idx, c_paths)
                break
            if files_hint and any(path in files_hint for path in c_paths):
                match = (layer_idx, file_idx, c_paths)
                break
        if match is not None:
            break
    if match is None:
        return {
            "layer_index": start_layer_idx,
            "file_index": start_file_idx,
            "func_pos": start_func_pos,
        }

    start_layer_idx, start_file_idx, c_paths = match
    if resume_func_name:
        func_list_all: list[dict] = []
        for c_path in c_paths:
            try:
                func_list, _ = _prepare_func_list_for_c_file(
                    c_path,
                    project_root=root_dir,
                    cfg=cfg,
                    prefilter=prefilter,
                    backend_module=backend,
                )
            except Exception:
                func_list = []
            for fd in func_list or []:
                fd2 = dict(fd)
                fc = dict(fd.get("file_context") or {})
                fc["source_file"] = c_path
                fd2["file_context"] = fc
                func_list_all.append(fd2)
        for pos, fd in enumerate(func_list_all, start=1):
            fi = fd.get("func_info") or {}
            fname = str(fi.get("func_name") or "")
            fsrc = str((fd.get("file_context") or {}).get("source_file") or "")
            if fname == resume_func_name and (not resume_file or os.path.abspath(fsrc) == os.path.abspath(resume_file)):
                start_func_pos = max(1, int(pos))
                break

    return {
        "layer_index": start_layer_idx,
        "file_index": start_file_idx,
        "func_pos": start_func_pos,
    }


def build_project_resume_marker(
    *,
    root_dir: str,
    output: str,
    start_layer_idx: int,
    start_file_idx: int,
    start_func_pos: int,
    start_func_index: int,
    module_counter: int,
    layer_heading_started: bool,
    module_started: bool,
    resume_module_id: Optional[str],
    resume_info: Optional[ProjectResumeState] = None,
) -> dict[str, Any]:
    info = resume_info or ProjectResumeState()
    return {
        "mode": "project",
        "project_dir": root_dir,
        "output": output,
        "layer_index": start_layer_idx,
        "file_index": start_file_idx,
        "func_pos": start_func_pos,
        "func_index": start_func_index,
        "module_counter": module_counter,
        "layer_started": layer_heading_started,
        "module_started": module_started,
        "module_id": resume_module_id,
        "module_files": list(info.module_files),
        "file": info.file,
        "func_name": info.func_name,
    }


def plan_project_module_run(
    *,
    continuing: bool,
    layer_idx: int,
    start_layer_idx: int,
    file_idx: int,
    start_file_idx: int,
    module_started: bool,
    resume_module_id: Optional[str],
    module_counter: int,
    req_id_prefix: str,
    start_func_index: int,
    start_func_pos: int,
    c_paths: Sequence[str],
    resume_marker: dict[str, Any],
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    reuse_module = bool(
        continuing and layer_idx == start_layer_idx and file_idx == start_file_idx and module_started
    )
    if reuse_module and resume_module_id:
        module_id = resume_module_id
        next_module_counter = module_counter
    else:
        module_id = f"{backend.normalize_req_prefix(req_id_prefix)}_{module_counter:03d}"
        next_module_counter = module_counter + 1
    func_index = start_func_index if reuse_module else 1
    func_start = start_func_pos if reuse_module else 1
    marker = dict(resume_marker or {})
    marker["module_counter"] = next_module_counter
    marker["module_started"] = reuse_module
    marker["module_id"] = module_id
    marker["module_files"] = list(c_paths)
    marker["func_pos"] = func_start
    marker["func_index"] = func_index
    return {
        "reuse_module": reuse_module,
        "module_id": module_id,
        "module_counter": next_module_counter,
        "func_index": func_index,
        "func_start_pos": func_start,
        "resume_marker": marker,
    }


def count_project_progress_total(
    layer_sets: Sequence[tuple[str, str, Sequence[dict]]],
    *,
    start_layer_idx: int,
    start_file_idx: int,
    start_func_pos: int,
    continuing: bool,
    module_started: bool,
    progress_pre: dict[str, dict],
    root_dir: str,
    cfg,
    prefilter: bool,
    backend_module=None,
) -> int:
    backend = backend_module or legacy_backend()

    def _count_module_funcs(c_paths: Sequence[str]) -> int:
        total = 0
        for c_path in c_paths:
            pre = (progress_pre or {}).get(c_path)
            if pre is not None:
                if pre.get("error"):
                    continue
                func_list = pre.get("func_list") or []
                if func_list:
                    total += len(func_list)
                    continue
            try:
                func_list, _ = _prepare_func_list_for_c_file(
                    c_path,
                    project_root=root_dir,
                    cfg=cfg,
                    prefilter=prefilter,
                    backend_module=backend,
                )
                total += len(func_list or [])
            except Exception:
                continue
        return int(total)

    total_work = 0
    for layer_idx, (_, _layer_name, modules) in enumerate(layer_sets):
        if layer_idx < start_layer_idx:
            continue
        file_start = start_file_idx if (continuing and layer_idx == start_layer_idx) else 0
        for file_idx in range(file_start, len(modules)):
            mod = modules[file_idx] if isinstance(modules[file_idx], dict) else {"name": None, "files": []}
            c_paths = _normalize_project_module_paths(mod)
            if not c_paths:
                continue
            func_count = _count_module_funcs(c_paths)
            if func_count <= 0:
                continue
            reuse_module = (
                continuing and layer_idx == start_layer_idx and file_idx == start_file_idx and module_started
            )
            func_start = start_func_pos if reuse_module else 1
            func_start = max(1, int(func_start))
            total_work += max(0, int(func_count) - (func_start - 1))
    return int(total_work)


def record_project_task_resume_progress(
    resume_marker: dict[str, Any],
    task: dict[str, Any],
    *,
    advance: bool,
) -> dict[str, Any]:
    marker = dict(resume_marker or {})
    marker["file"] = str(task.get("file") or "")
    marker["func_name"] = str(task.get("func_name") or "")
    marker["func_pos"] = int(task.get("func_pos") or 1) + (1 if advance else 0)
    marker["func_index"] = int(task.get("index") or 1) + (1 if advance else 0)
    return marker


def advance_project_resume_after_module(
    resume_marker: dict[str, Any],
    *,
    file_idx: int,
) -> dict[str, Any]:
    marker = dict(resume_marker or {})
    marker["file_index"] = int(file_idx) + 1
    marker["func_pos"] = 1
    marker["func_index"] = 1
    marker["module_started"] = False
    marker["module_id"] = None
    return marker


def advance_project_resume_after_layer(resume_marker: dict[str, Any]) -> dict[str, Any]:
    marker = dict(resume_marker or {})
    marker["layer_started"] = False
    return marker


def build_single_file_func_start_event(
    task: dict[str, Any],
    *,
    mode: str,
    source: str,
) -> dict[str, Any]:
    return {
        "type": "func_start",
        "mode": mode,
        "source": source,
        "file": str(task.get("file") or source),
        "module_id": str(task.get("module_id") or ""),
        "func_index": int(task.get("index") or 0),
        "func_pos": int(task.get("func_pos") or 0),
        "func_name": str(task.get("func_name") or ""),
    }


def build_single_file_func_end_event(
    task: dict[str, Any],
    *,
    mode: str,
    source: str,
    design: Any = None,
    error: Any = "",
    ok: Optional[bool] = None,
) -> dict[str, Any]:
    resolved_ok = bool(ok) if ok is not None else (not error)
    payload = {
        "type": "func_end",
        "mode": mode,
        "file": str(task.get("file") or source),
        "module_id": str(task.get("module_id") or ""),
        "func_index": int(task.get("index") or 0),
        "func_pos": int(task.get("func_pos") or 0),
        "func_name": str(task.get("func_name") or ""),
        "ok": resolved_ok,
        "error": "" if resolved_ok else str(error or ""),
    }
    if mode == "single_export":
        payload["source"] = source
    if design is not None:
        payload.update(
            {
                "io_ok": bool(getattr(design, "io_none", False) or getattr(design, "io_elements", ()) or ()),
                "locals_ok": True,
                "locals_empty": bool(getattr(design, "local_elements", ()) == ()),
                "logic_ok": True,
                "logic_empty": bool(getattr(design, "logic_lines", ()) == ()),
            }
        )
    return payload


def record_single_file_resume_progress(
    resume_marker: dict[str, Any],
    task: dict[str, Any],
    *,
    advance: bool,
) -> dict[str, Any]:
    marker = dict(resume_marker or {})
    marker["func_pos"] = int(task.get("func_pos") or 1) + (1 if advance else 0)
    marker["func_index"] = int(task.get("index") or 1) + (1 if advance else 0)
    return marker


def run_single_export_design(task: dict[str, Any], cfg, *, func_name: str, backend_module=None):
    backend = backend_module or legacy_backend()
    try:
        cfg._in_func_context = True
        cfg._current_func_ai_failed = False
        cfg._skip_ai_current_func = False
    except Exception:
        pass
    try:
        design = backend._run_function_design_task(task, cfg)
        original_desc = utils_module._safe_strip(((task.get("func_data") or {}).get("comment_info") or {}).get("desc"))
        forced_title = backend._normalize_function_cn_title(
            design.title,
            func_name=func_name,
            comment_desc=original_desc,
        )
        if forced_title and forced_title != design.title:
            design = backend.replace(design, title=forced_title)
        revision_patch = revision_utils.find_function_patch(
            revision_utils.load_revision_profile(cfg),
            str(task.get("file") or ""),
            func_name,
        )
        if revision_patch:
            design = revision_utils.apply_revision_to_design(design, revision_patch)
        return design
    finally:
        try:
            cfg._in_func_context = False
        except Exception:
            pass


def build_project_func_start_event(task: dict[str, Any], *, root_dir: str) -> dict[str, Any]:
    return {
        "type": "func_start",
        "mode": "project",
        "project_dir": root_dir,
        "file": str(task.get("file") or ""),
        "layer": str(task.get("layer") or ""),
        "module_id": str(task.get("module_id") or ""),
        "func_index": int(task.get("index") or 0),
        "func_pos": int(task.get("func_pos") or 0),
        "func_name": str(task.get("func_name") or ""),
    }


def build_project_func_end_event(
    task: dict[str, Any],
    *,
    root_dir: str,
    design: Any = None,
    error: Any = "",
    ok: Optional[bool] = None,
) -> dict[str, Any]:
    resolved_ok = bool(ok) if ok is not None else (not error)
    payload = {
        "type": "func_end",
        "mode": "project",
        "project_dir": root_dir,
        "file": str(task.get("file") or ""),
        "layer": str(task.get("layer") or ""),
        "module_id": str(task.get("module_id") or ""),
        "func_index": int(task.get("index") or 0),
        "func_pos": int(task.get("func_pos") or 0),
        "func_name": str(task.get("func_name") or ""),
        "ok": resolved_ok,
        "error": "" if resolved_ok else str(error or ""),
    }
    if design is not None:
        payload.update(
            {
                "io_ok": bool(getattr(design, "io_none", False) or getattr(design, "io_elements", ()) or ()),
                "locals_ok": True,
                "locals_empty": bool(getattr(design, "local_elements", ()) == ()),
                "logic_ok": True,
                "logic_empty": bool(getattr(design, "logic_lines", ()) == ()),
            }
        )
    return payload


def update_project_module_table_title(module_table: Any, task: dict[str, Any], design: Any) -> bool:
    if module_table is None or design is None:
        return False
    try:
        module_table.rows[int(task.get("index") or 0)].cells[2].text = getattr(design, "title", "")
    except Exception:
        return False
    return True


def finalize_project_generation(
    output: str,
    root_dir: str,
    cfg,
    *,
    stopped: bool,
    resume_marker: dict[str, Any],
    collect_unit_rows: bool,
    unit_rows: Sequence[dict],
    collect_unit_func_lists: bool,
    unit_func_tables: Sequence[dict],
    designs: Optional[Sequence[dict]] = None,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    result = {
        "stopped": bool(stopped),
        "resume_state": None,
        "unit_output": "",
        "func_list_output": "",
    }
    if stopped:
        backend.vlog(cfg, f"文档生成已停止，已输出部分内容：{output}")
        if getattr(cfg, "resume_state", None) is None:
            cfg.resume_state = dict(resume_marker)
        result["resume_state"] = getattr(cfg, "resume_state", None)
        return result

    backend.vlog(cfg, f"文档生成完成：{output}")
    cfg.resume_state = None
    result["resume_state"] = None

    if collect_unit_rows:
        try:
            unit_output = backend.derive_software_unit_output_path(output)
            backend.build_software_unit_table_doc(unit_rows, unit_output)
            backend.vlog(cfg, f"软件单元清单已生成：{unit_output}")
            result["unit_output"] = unit_output
        except Exception as exc:
            backend.write_error_log(
                "software_unit_table_failed",
                {
                    "project_dir": root_dir,
                    "output": output,
                    "error": repr(exc),
                },
            )
            backend.vlog(cfg, f"软件单元清单生成失败：{exc}")

    if collect_unit_func_lists:
        try:
            func_list_output = backend.derive_unit_function_list_output_path(output)
            backend.build_unit_function_list_doc(unit_func_tables, func_list_output)
            backend.vlog(cfg, f"单元函数列表已生成：{func_list_output}")
            result["func_list_output"] = func_list_output
        except Exception as exc:
            backend.write_error_log(
                "unit_function_list_failed",
                {
                    "project_dir": root_dir,
                    "output": output,
                    "error": repr(exc),
                },
            )
            backend.vlog(cfg, f"单元函数列表生成失败：{exc}")
    else:
        backend.vlog(cfg, "继续模式：跳过软件单元清单生成（避免不完整）。")

    # 术语一致性检查
    if designs and not stopped:
        try:
            from .term_checker import collect_term_mappings, check_consistency
            import json

            symbol_dict_path = backend.app_root() / "symbol_dictionary.json"
            symbol_dict = {}
            if symbol_dict_path.exists():
                symbol_dict = json.loads(symbol_dict_path.read_text(encoding="utf-8"))

            term_map = collect_term_mappings(list(designs), symbol_dict=symbol_dict)
            report = check_consistency(term_map, symbol_dict=symbol_dict)

            result["consistency_score"] = report.score
            result["consistency_inconsistencies"] = len(report.inconsistencies)
            result["consistency_dict_conflicts"] = len(report.symbol_dict_conflicts)

            if report.score < 80:
                backend.vlog(cfg, f"术语一致性警告：评分 {report.score:.1f}/100，发现 {len(report.inconsistencies)} 处不一致")
            else:
                backend.vlog(cfg, f"术语一致性检查通过：评分 {report.score:.1f}/100")

            report_path = ""
            try:
                from .term_checker import report_to_dict, write_consistency_report

                report_path = os.path.splitext(str(output or "out.docx"))[0] + ".consistency_report.json"
                write_consistency_report(report, report_path)
                result["consistency_report_path"] = report_path
                backend.vlog(cfg, f"术语一致性报告已写入：{report_path}")
            except Exception as write_exc:
                backend.vlog(cfg, f"术语一致性报告写入失败：{write_exc}")

            event_payload = {
                "type": "consistency_report",
                "score": report.score,
                "total_symbols": report.total_symbols,
                "inconsistencies": len(report.inconsistencies),
                "dict_conflicts": len(report.symbol_dict_conflicts),
            }
            if report_path:
                event_payload["report_path"] = report_path
            try:
                from .term_checker import report_to_dict

                # Optional full payload for agents that can consume it; GUI may ignore.
                event_payload["report"] = report_to_dict(report)
            except Exception:
                pass
            backend.gui_event(cfg, event_payload)
        except Exception as exc:
            backend.vlog(cfg, f"术语一致性检查失败：{exc}")

    return result


def save_project_stop_placeholder_doc(
    output: str,
    cfg,
    *,
    stage: str,
    message: str,
    backend_module=None,
    render_module=None,
) -> str:
    backend = backend_module or legacy_backend()
    if render_module is None:
        from . import render as render_module
    try:
        output = backend.normalize_docx_output_path(output, ensure_parent_dir=True)
    except Exception as exc:
        raise backend.RenderError(f"输出路径无效：{exc}") from exc
    doc_state = render_module.init_generation_document(
        cfg,
        main_heading="CSCI详细设计",
        heading_level=1,
        backend_module=backend,
    )
    doc = doc_state["doc"]
    doc.add_paragraph("生成状态", style=backend.pick_heading_style(doc, 2))
    doc.add_paragraph(message or "用户已请求停止，当前阶段尚未生成函数内容。", style="Normal")
    doc.add_paragraph(f"停止阶段：{stage}", style="Normal")
    try:
        backend.safe_save_docx(doc, output)
    except Exception as exc:
        raise backend.RenderError(f"保存 Word 失败：{exc}") from exc
    backend.vlog(cfg, f"文档生成已停止，已输出停止说明：{output}")
    return output


def finalize_single_file_generation(
    output: str,
    source: str,
    cfg,
    *,
    stopped: bool,
    resume_marker: dict[str, Any],
    unit_func_tables: Sequence[dict],
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    result = {
        "stopped": bool(stopped),
        "resume_state": None,
        "func_list_output": "",
    }
    if stopped:
        backend.vlog(cfg, f"文档生成已停止，已输出部分内容：{output}")
        if getattr(cfg, "resume_state", None) is None:
            cfg.resume_state = dict(resume_marker)
        result["resume_state"] = getattr(cfg, "resume_state", None)
        return result

    backend.vlog(cfg, f"文档生成完成：{output}")
    cfg.resume_state = None
    result["resume_state"] = None
    if unit_func_tables:
        try:
            func_list_output = backend.derive_unit_function_list_output_path(output)
            backend.build_unit_function_list_doc(unit_func_tables, func_list_output)
            backend.vlog(cfg, f"单元函数列表已生成：{func_list_output}")
            result["func_list_output"] = func_list_output
        except Exception as exc:
            backend.write_error_log(
                "unit_function_list_failed",
                {
                    "source": source,
                    "output": output,
                    "error": repr(exc),
                },
            )
            backend.vlog(cfg, f"单元函数列表生成失败：{exc}")
    return result


def run_single_file_generation(
    source: str,
    output: str,
    cfg,
    *,
    resume_state: Optional[dict] = None,
    backend_module=None,
    render_module=None,
    runtime_module=None,
):
    backend = backend_module or legacy_backend()
    if render_module is None:
        from . import render as render_module
    if runtime_module is None:
        from . import runtime as runtime_module

    backend.vlog(cfg, f"开始处理单文件：{source}")
    if _EVIDENCE_ENABLED:
        clear_recorded_evidence()
        clear_untranslated_idents()
    project_root = backend._guess_project_root_for_source(source)
    # 加载 docx 渲染缓存
    try:
        if project_root:
            render_module.load_render_cache(os.path.join(project_root, ".autodoc", "render_cache.json"))
    except Exception:
        pass
    runtime_ctx = runtime_module.ensure_project_runtime(
        cfg,
        project_root=project_root,
        resume_state=resume_state,
    )
    generation = runtime_ctx.generation
    continuing = generation.continuing
    try:
        output = backend.normalize_docx_output_path(output, ensure_parent_dir=(not continuing))
        _reset_review_collection(cfg)
        _reset_design_workspace_pairs(cfg)
    except Exception as exc:
        raise backend.RenderError(f"输出路径无效：{exc}") from exc
    try:
        from . import graph_visuals

        graph_visuals.configure_graph_output(cfg, output)
    except Exception as exc:
        backend.vlog(cfg, f"[Graph] 输出目录初始化失败，已关闭图谱输出：{exc}")

    func_list, _ = _prepare_func_list_for_c_file(
        source,
        project_root=project_root,
        cfg=cfg,
        prefilter=False,
        backend_module=backend,
    )
    for fd in func_list or ():
        file_context = fd.get("file_context") if isinstance(fd, dict) else None
        if not isinstance(file_context, dict):
            file_context = {}
            if isinstance(fd, dict):
                fd["file_context"] = file_context
        file_context["source_file"] = source
    apply_project_function_title_registry(
        func_list,
        project_root or os.path.dirname(os.path.abspath(source)),
        cfg,
        backend_module=backend,
    )
    if len(func_list or []) > 1:
        backend._warmup_symbol_memory_once(func_list, cfg, scope_label=f"single_file:{os.path.basename(source)}")
    backend.vlog(cfg, f"扫描到 {len(func_list)} 个函数（含注释）。")
    if not func_list:
        raise backend.NoDataError(f"未解析到任何函数：{source}")

    cfg.resume_state = None
    cfg._output_doc_path = output
    if continuing:
        if not os.path.exists(output):
            raise backend.RenderError("未找到已生成的文档，无法继续。")
        doc = backend.Document(output)
        placeholder = None
        body_start_idx = 0
        resume_info = runtime_module.normalize_single_file_resume_state(generation.resume_state)
        resume_func_pos = resume_info.func_pos
        func_index = resume_info.func_index
    else:
        doc_state = render_module.init_generation_document(cfg, backend_module=backend)
        doc = doc_state["doc"]
        placeholder = doc_state["placeholder"]
        body_start_idx = int(doc_state["body_start_idx"])
        resume_func_pos = 1
        func_index = 1

    try:
        total_work = count_single_file_progress_total(
            func_list,
            resume_func_pos=resume_func_pos,
            only_with_comment=generation.only_with_comment,
        )
        backend.gui_event(cfg, {"type": "progress_init", "mode": "single", "total": int(total_work), "unit": "函数"})
    except Exception:
        pass

    wrote_any = False
    stopped = False
    resume_marker = {
        "mode": "single_file",
        "source": source,
        "output": output,
        "func_pos": resume_func_pos,
        "func_index": func_index,
    }
    module_id = backend.normalize_req_prefix(cfg.req_id_prefix)
    _, module_name = resolve_source_module_names(
        source,
        backend_module=backend,
    )

    module_table = None
    unit_func_tables: list[dict] = []
    if not continuing:
        table_payload = build_single_file_table_payload(
            func_list,
            resume_func_pos=resume_func_pos,
            module_id=module_id,
            module_name=module_name,
            only_with_comment=generation.only_with_comment,
            backend_module=backend,
        )
        module_table = render_module.add_module_section(
            doc,
            module_name,
            module_id,
            table_payload["entries"],
            backend_module=backend,
        )
        unit_func_tables.append(table_payload["unit_func_table"])

    task_payload = build_single_file_tasks(
        func_list,
        resume_func_pos=resume_func_pos,
        func_index=func_index,
        module_id=module_id,
        source=source,
        only_with_comment=generation.only_with_comment,
    )
    tasks = task_payload["tasks"]
    resume_marker.update(task_payload["resume_marker_updates"])
    execution = execute_single_file_tasks(
        doc,
        tasks,
        cfg=cfg,
        source=source,
        module_table=module_table,
        resume_marker=resume_marker,
        backend_module=backend,
        render_module=render_module,
    )
    processed_tasks = int(execution["processed_tasks"])
    wrote_any = wrote_any or bool(execution["wrote_any"])
    resume_marker = dict(execution["resume_marker"])

    single_finalize = finalize_single_file_task_iteration(
        cfg=cfg,
        processed_tasks=processed_tasks,
        tasks=tasks,
        resume_marker=resume_marker,
        backend_module=backend,
    )
    resume_marker = dict(single_finalize["resume_marker"])
    stopped = bool(single_finalize["stopped"])

    if (not wrote_any) and (not stopped):
        raise backend.NoDataError("没有符合条件的函数可生成。")

    if not continuing:
        backend.relocate_generated_blocks(doc, body_start_idx, placeholder)
    try:
        backend.safe_save_docx(doc, output)
    except Exception as exc:
        raise backend.RenderError(f"保存 Word 失败：{exc}") from exc
    _write_review_workspace_if_enabled(cfg, output, project_root=_review_project_root(source=source), merge_existing=continuing)
    _write_design_workspace_if_enabled(cfg, output, project_root=_review_project_root(source=source), merge_existing=continuing)

    finalize_single_file_generation(
        output,
        source,
        cfg,
        stopped=stopped,
        resume_marker=resume_marker,
        unit_func_tables=unit_func_tables,
        backend_module=backend,
    )
    try:
        from . import graph_visuals

        html_path = graph_visuals.write_html_report(cfg, title=f"AutoDocGen 调用图谱 - {os.path.basename(source)}")
        if html_path:
            backend.vlog(cfg, f"调用图谱 HTML 已生成：{html_path}")
    except Exception as exc:
        backend.vlog(cfg, f"[Graph] HTML 图谱生成失败：{exc}")
    if generation.open_after_done and os.name == "nt":
        try:
            os.startfile(output)
        except Exception:
            pass
    backend.finalize_project_symbol_memory(cfg)
    # 保存 docx 渲染缓存
    try:
        if project_root:
            render_module.save_render_cache(os.path.join(project_root, ".autodoc", "render_cache.json"))
    except Exception:
        pass
    maybe_write_evidence_report(output, cfg, backend_module=backend)
    if _EVIDENCE_ENABLED and logic_step_ir_enabled(cfg):
        try:
            suggested = auto_suggest_symbol_translations(cfg, project_root=project_root)
            if suggested:
                backend.vlog(cfg, f"[LogicStep] AI 自动建议 {len(suggested)} 个符号翻译，已写入符号记忆库")
        except Exception as exc:
            backend.vlog(cfg, f"[LogicStep] 自动翻译建议失败：{exc}")
    return output


def run_single_export_generation(
    source: str,
    func_name: str,
    output: str,
    cfg,
    *,
    project_root: Optional[str] = None,
    backend_module=None,
    render_module=None,
    runtime_module=None,
) -> str:
    backend = backend_module or legacy_backend()
    if render_module is None:
        from . import render as render_module
    if runtime_module is None:
        from . import runtime as runtime_module

    if not source:
        raise ValueError("source 不能为空")
    if not func_name:
        raise ValueError("func_name 不能为空")
    if _EVIDENCE_ENABLED:
        clear_recorded_evidence()

    cfg = clone_cfg(cfg, enhanced_single_func_pseudocode=True)
    backend.vlog(cfg, f"导出单函数：{func_name} <- {source}")
    try:
        output = backend.normalize_docx_output_path(output, ensure_parent_dir=True)
        _reset_review_collection(cfg)
        _reset_design_workspace_pairs(cfg)
    except Exception as exc:
        raise backend.RenderError(f"输出路径无效：{exc}") from exc
    try:
        from . import graph_visuals

        graph_visuals.configure_graph_output(cfg, output)
    except Exception as exc:
        backend.vlog(cfg, f"[Graph] 输出目录初始化失败，已关闭图谱输出：{exc}")

    project_root = project_root or backend._guess_project_root_for_source(source)
    runtime_ctx = runtime_module.ensure_project_runtime(cfg, project_root=project_root)
    generation = runtime_ctx.generation
    prepare_cfg = cfg
    try:
        from . import ai as ai_utils
        if ai_utils.ai_context_scope(cfg) == "target_only":
            prepare_cfg = clone_cfg(cfg, ai_assist=False)
    except Exception:
        prepare_cfg = cfg
    func_list, _ = _prepare_func_list_for_c_file(
        source,
        project_root=project_root,
        cfg=prepare_cfg,
        prefilter=False,
        backend_module=backend,
    )
    if not func_list:
        raise backend.NoDataError(f"未解析到任何函数：{source}")

    target = None
    target_pos = 1
    for idx, fd in enumerate(func_list, start=1):
        fi = fd.get("func_info") or {}
        if (fi.get("func_name") or "").strip() == func_name.strip():
            target = fd
            target_pos = idx
            break
    if target is None:
        candidates = [(fd.get("func_info") or {}).get("func_name") or "" for fd in func_list]
        candidates = [c for c in candidates if c]
        hint = ("；可选函数：\n- " + "\n- ".join(candidates[:50])) if candidates else ""
        raise backend.NoDataError(f"未在文件中找到函数：{func_name}{hint}")

    try:
        backend.gui_event(cfg, {"type": "progress_init", "mode": "single_export", "total": 1, "unit": "函数"})
    except Exception:
        pass

    doc_state = render_module.init_generation_document(cfg, backend_module=backend)
    doc = doc_state["doc"]
    placeholder = doc_state["placeholder"]
    body_start_idx = int(doc_state["body_start_idx"])

    module_id = backend.normalize_req_prefix(cfg.req_id_prefix)
    ci = target.get("comment_info") or {}
    fi = target.get("func_info") or {}
    csu_name = backend.get_function_chinese_name_rich(
        target,
        cfg=cfg,
    ) if getattr(cfg, "ai_assist", False) else backend.get_function_chinese_name(ci, fi)
    func_ident = utils_module._safe_strip(fi.get("func_name", ""))
    _, module_name = resolve_source_module_names(
        source,
        func_cn_name=csu_name,
        func_ident=func_ident,
        backend_module=backend,
    )
    csu_id = f"{backend.normalize_req_prefix(module_id)}_{1:03d}"
    module_table = render_module.add_module_section(
        doc,
        module_name,
        module_id,
        [{"csu_name": csu_name, "csu_id": csu_id}],
        backend_module=backend,
    )

    task = build_single_export_task(
        target,
        target_pos=int(target_pos or 1),
        func_name=func_name,
        module_id=module_id,
        source=source,
    )
    execute_single_export_task(
        doc,
        task,
        cfg=cfg,
        source=source,
        module_table=module_table,
        backend_module=backend,
        render_module=render_module,
    )
    backend.relocate_generated_blocks(doc, body_start_idx, placeholder)
    try:
        backend.safe_save_docx(doc, output)
    except Exception as exc:
        raise backend.RenderError(f"保存 Word 失败：{exc}") from exc
    _write_review_workspace_if_enabled(cfg, output, project_root=_review_project_root(source=source))
    _write_design_workspace_if_enabled(cfg, output, project_root=_review_project_root(source=source))

    backend.vlog(cfg, f"单函数导出完成：{output}")
    try:
        from . import graph_visuals

        html_path = graph_visuals.write_html_report(cfg, title=f"AutoDocGen 调用图谱 - {func_name}")
        if html_path:
            backend.vlog(cfg, f"调用图谱 HTML 已生成：{html_path}")
    except Exception as exc:
        backend.vlog(cfg, f"[Graph] HTML 图谱生成失败：{exc}")
    if generation.open_after_done and os.name == "nt":
        try:
            os.startfile(output)
        except Exception:
            pass
    backend.finalize_project_symbol_memory(cfg)
    maybe_write_evidence_report(output, cfg, backend_module=backend)
    if _EVIDENCE_ENABLED and logic_step_ir_enabled(cfg):
        try:
            suggested = auto_suggest_symbol_translations(cfg, project_root=project_root)
            if suggested:
                backend.vlog(cfg, f"[LogicStep] AI 自动建议 {len(suggested)} 个符号翻译，已写入符号记忆库")
        except Exception as exc:
            backend.vlog(cfg, f"[LogicStep] 自动翻译建议失败：{exc}")
    return output


def run_project_generation(
    root_dir: str,
    output: str,
    cfg,
    *,
    resume_state: Optional[dict] = None,
    incremental: bool = False,
    backend_module=None,
    runtime_module=None,
) -> None:
    backend = backend_module or legacy_backend()
    if runtime_module is None:
        from . import runtime as runtime_module

    backend.vlog(cfg, f"开始处理工程目录：{root_dir}")
    if _EVIDENCE_ENABLED:
        clear_recorded_evidence()
    # 加载 docx 渲染缓存（增量复跑时跳过未变函数的渲染）
    try:
        from . import render as _render_mod
        _render_mod.load_render_cache(os.path.join(root_dir, ".autodoc", "render_cache.json"))
    except Exception:
        pass
    if not os.path.isdir(root_dir):
        raise backend.SourceReadError(f"目录不存在：{root_dir}")
    runtime_ctx = runtime_module.ensure_project_runtime(
        cfg,
        project_root=root_dir,
        resume_state=resume_state,
    )
    generation = runtime_ctx.generation
    continuing = generation.continuing
    try:
        output = backend.normalize_docx_output_path(output, ensure_parent_dir=(not continuing))
        _reset_review_collection(cfg)
        _reset_design_workspace_pairs(cfg)
    except Exception as exc:
        raise backend.RenderError(f"输出路径无效：{exc}") from exc
    try:
        from . import graph_visuals

        graph_visuals.configure_graph_output(cfg, output)
    except Exception as exc:
        backend.vlog(cfg, f"[Graph] 输出目录初始化失败，已关闭图谱输出：{exc}")
    codegraph_mode = "auto"
    try:
        from . import codegraph_adapter

        codegraph_mode = codegraph_adapter.graph_mode_from_cfg(cfg)
        cg_status = codegraph_adapter.prepare_project_index(root_dir, cfg)
        if cg_status.enabled:
            backend.vlog(cfg, f"[CodeGraph] 索引可用：{cg_status.index_path}")
        else:
            backend.vlog(cfg, f"[CodeGraph] 未启用：{cg_status.message}")
    except Exception as exc:
        backend.vlog(cfg, f"[CodeGraph] 准备失败：{exc}")
        if codegraph_mode == "force":
            raise
    try:
        prebuilt = _prebuild_project_symbols_into_runtime(
            root_dir,
            cfg=cfg,
            backend_module=backend,
        )
        backend.vlog(cfg, f"[prebuild] loaded {sum(len(v) for v in prebuilt.values())} symbols from project headers")
    except Exception as exc:
        backend.vlog(cfg, f"[prebuild] skipped ({exc})")
    if backend.stop_requested(cfg):
        backend.vlog(cfg, "收到停止请求，结束工程预构建。")
        cfg.resume_state = {"stage": "project_prebuild"}
        if not continuing:
            save_project_stop_placeholder_doc(
                output,
                cfg,
                stage="project_prebuild",
                message="已在工程预构建阶段停止，尚未进入工程扫描和函数生成。",
                backend_module=backend,
            )
        return output

    source_layout = plan_project_source_layout(
        root_dir,
        cfg,
        order_override=generation.project_file_order,
        backend_module=backend,
    )
    if backend.stop_requested(cfg):
        backend.vlog(cfg, "收到停止请求，结束工程扫描。")
        cfg.resume_state = {"stage": "project_scan"}
        if not continuing:
            save_project_stop_placeholder_doc(
                output,
                cfg,
                stage="project_scan",
                message="已在工程扫描阶段停止，尚未进入函数文档生成。",
                backend_module=backend,
            )
        return output
    src_dir = str(source_layout.get("src_dir") or "")
    app_files = list(source_layout.get("app_files") or [])
    mid_files = list(source_layout.get("mid_files") or [])
    drv_files = list(source_layout.get("drv_files") or [])
    app_modules = list(source_layout.get("app_modules") or [])
    mid_modules = list(source_layout.get("mid_modules") or [])
    drv_modules = list(source_layout.get("drv_modules") or [])
    if not src_dir:
        raise backend.NoDataError("未找到 SRC 目录（不区分大小写）。")
    total_files = int(source_layout.get("total_files", 0) or 0)
    total_modules = int(source_layout.get("total_modules", 0) or 0)
    backend.vlog(
        cfg,
        f"工程扫描完成：共 {total_files} 个 C 文件 / {total_modules} 个模块；应用层 {len(app_files)}，中间层 {len(mid_files)}，驱动层 {len(drv_files)}。",
    )
    if total_files == 0:
        raise backend.NoDataError("SRC 目录下未找到任何 .c 文件。")

    cfg.resume_state = None
    cfg._output_doc_path = output
    prefilter = generation.prefilter_project_files
    preprocessed: dict[str, dict] = {}
    func_entries_for_graph: list[dict] = []
    ordered_files = app_files + mid_files + drv_files
    # A resumed project must use the same project-wide title choices as a new
    # run; otherwise a later colliding title can silently change in the DOCX.
    preprocessed = preprocess_project_files(
        ordered_files,
        project_root=root_dir,
        cfg=cfg,
        prefilter=prefilter,
        backend_module=backend,
    )
    apply_preprocessed_function_title_registry(
        preprocessed,
        root_dir,
        cfg,
        backend_module=backend,
    )
    func_entries = backend._flatten_preprocessed_func_entries(preprocessed)
    if not continuing:
        func_entries_for_graph = list(func_entries or [])
        try:
            term_table = backend.build_project_term_table(
                root_dir,
                cfg,
                prebuilt=prebuilt if isinstance(prebuilt, dict) else None,
                func_entries=func_entries,
                save=True,
            )
            total_terms = sum(len(term_table.get(section) or {}) for section in ("functions", "symbols", "members", "macros"))
            backend.vlog(cfg, f"[term_table] loaded {total_terms} project terms")
        except Exception as exc:
            backend.vlog(cfg, f"[term_table] skipped ({exc})")
            backend._warmup_symbol_memory_once(
                func_entries,
                cfg,
                scope_label=f"project:{os.path.basename(os.path.abspath(root_dir))}",
            )
    collect_unit_rows = not continuing
    collect_unit_func_lists = not continuing
    unit_rows: list[dict] = []
    unit_func_tables: list[dict] = []
    unit_index = 1
    designs: list[dict] = []  # 收集所有函数设计用于术语一致性检查

    # 增量生成：加载上次状态
    incremental_state = None
    skipped_count = 0
    if incremental and not continuing:
        try:
            from .incremental import load_incremental_state
            incremental_state = load_incremental_state(root_dir)
            backend.vlog(cfg, f"[增量] 加载上次生成状态")
        except Exception as exc:
            backend.vlog(cfg, f"[增量] 状态加载失败: {exc}")

    run_state = initialize_project_run_state(
        output,
        cfg,
        continuing=continuing,
        resume_state=generation.resume_state,
        backend_module=backend,
        runtime_module=runtime_module,
    )
    doc = run_state["doc"]
    if not continuing and func_entries_for_graph:
        try:
            from . import render as render_module

            render_module.render_project_graph_overview(
                doc,
                func_entries_for_graph,
                cfg,
                root_dir=root_dir,
                backend_module=backend,
            )
        except Exception as exc:
            backend.vlog(cfg, f"[Graph] 项目调用关系总览生成失败：{exc}")
    if not continuing and func_entries_for_graph:
        try:
            from . import callgraph as cgmod
            from . import render as render_module

            callees_map = cgmod.build_project_callees_map(func_entries_for_graph)
            if callees_map:
                title_map: dict[str, str] = {}
                for fd in func_entries_for_graph:
                    fi = (fd or {}).get("func_info") or {}
                    fn = str(fi.get("func_name") or "").strip()
                    fc = (fd or {}).get("file_context") or {}
                    title = str(fc.get("function_title") or "").strip()
                    if fn and title:
                        title_map[fn] = title
                entries = cgmod.find_entry_functions(callees_map)
                if not entries:
                    entries = list(callees_map.keys())[:1]
                all_tree_rows: list[tuple[str, str, str, str]] = []
                for entry_fn in entries:
                    tree_rows = cgmod.flatten_call_tree(
                        callees_map,
                        entry_fn,
                        max_depth=3,
                        name_map=title_map,
                    )
                    all_tree_rows.extend(tree_rows)
                if all_tree_rows:
                    render_module.render_static_call_relation_table(
                        doc,
                        all_tree_rows,
                        entry_label="主函数",
                        backend_module=backend,
                    )
        except Exception as exc:
            backend.vlog(cfg, f"[Graph] 软件单元静态关系表生成失败：{exc}")
    resume_info = run_state["resume_info"]
    start_layer_idx = int(run_state["start_layer_idx"])
    start_file_idx = int(run_state["start_file_idx"])
    start_func_pos = int(run_state["start_func_pos"])
    start_func_index = int(run_state["start_func_index"])
    module_counter = int(run_state["module_counter"])
    layer_heading_started = bool(run_state["layer_heading_started"])
    module_started = bool(run_state["module_started"])
    resume_module_id = run_state["resume_module_id"]

    layer_sets = build_project_layer_sets(app_modules, mid_modules, drv_modules)
    if continuing:
        checkpoint = resolve_project_resume_checkpoint(
            layer_sets,
            resume_info,
            root_dir=root_dir,
            cfg=cfg,
            prefilter=prefilter,
            backend_module=backend,
        )
        start_layer_idx = int(checkpoint.get("layer_index", start_layer_idx))
        start_file_idx = int(checkpoint.get("file_index", start_file_idx))
        start_func_pos = int(checkpoint.get("func_pos", start_func_pos))

    stopped = False
    wrote_any = False
    resume_marker = build_project_resume_marker(
        root_dir=root_dir,
        output=output,
        start_layer_idx=start_layer_idx,
        start_file_idx=start_file_idx,
        start_func_pos=start_func_pos,
        start_func_index=start_func_index,
        module_counter=module_counter,
        layer_heading_started=layer_heading_started,
        module_started=module_started,
        resume_module_id=resume_module_id,
        resume_info=resume_info if continuing else None,
    )

    try:
        progress_state = prepare_project_progress(
            cfg=cfg,
            continuing=continuing,
            preprocessed=preprocessed,
            app_modules=app_modules,
            mid_modules=mid_modules,
            drv_modules=drv_modules,
            layer_sets=layer_sets,
            start_layer_idx=start_layer_idx,
            start_file_idx=start_file_idx,
            start_func_pos=start_func_pos,
            module_started=module_started,
            root_dir=root_dir,
            prefilter=prefilter,
            backend_module=backend,
        )
        if progress_state is not None:
            backend.gui_event(cfg, progress_state["payload"])
    except Exception:
        pass

    for layer_idx, (_, layer_name, files) in enumerate(layer_sets):
        if layer_idx < start_layer_idx:
            continue
        layer_state = prepare_project_layer_iteration(
            doc,
            cfg=cfg,
            layer_idx=layer_idx,
            layer_name=layer_name,
            files=files,
            continuing=continuing,
            start_layer_idx=start_layer_idx,
            layer_heading_started=layer_heading_started,
            resume_marker=resume_marker,
            backend_module=backend,
        )
        resume_marker = dict(layer_state["resume_marker"])
        layer_heading_started = bool(layer_state["layer_heading_started"])
        if layer_state["stopped"]:
            stopped = True
            break
        if layer_state["empty_layer"]:
            start_file_idx = 0
            layer_heading_started = False
            resume_marker = advance_project_resume_after_layer(resume_marker)
            continue

        file_start = start_file_idx if (continuing and layer_idx == start_layer_idx) else 0
        for file_idx in range(file_start, len(files)):
            mod = files[file_idx]
            file_state = prepare_project_file_iteration(
                mod,
                cfg=cfg,
                file_idx=file_idx,
                resume_marker=resume_marker,
                backend_module=backend,
            )
            resume_marker = dict(file_state["resume_marker"])
            c_paths = list(file_state["c_paths"])
            if file_state["stopped"]:
                stopped = True
                break
            if not c_paths:
                continue

            module_run = plan_project_module_run(
                continuing=continuing,
                layer_idx=layer_idx,
                start_layer_idx=start_layer_idx,
                file_idx=file_idx,
                start_file_idx=start_file_idx,
                module_started=module_started,
                resume_module_id=resume_module_id,
                module_counter=module_counter,
                req_id_prefix=cfg.req_id_prefix,
                start_func_index=start_func_index,
                start_func_pos=start_func_pos,
                c_paths=c_paths,
                resume_marker=resume_marker,
                backend_module=backend,
            )
            reuse_module = bool(module_run["reuse_module"])
            module_id = str(module_run["module_id"] or "")
            module_counter = int(module_run["module_counter"])
            func_index = int(module_run["func_index"])
            func_start_pos = int(module_run["func_start_pos"])
            resume_marker = dict(module_run["resume_marker"])

            func_list_all = collect_project_module_functions(
                c_paths,
                preprocessed=preprocessed,
                root_dir=root_dir,
                cfg=cfg,
                prefilter=prefilter,
                backend_module=backend,
            )
            if not func_list_all:
                backend.vlog(cfg, f"未解析到任何函数，跳过模块：{os.path.basename(c_paths[0])}")
                resume_marker = advance_project_resume_after_module(
                    resume_marker,
                    file_idx=file_idx,
                )
                start_func_pos = 1
                start_func_index = 1
                resume_module_id = None
                module_started = False
                continue

            module_name, module_display = resolve_project_module_names(
                mod,
                c_paths,
                backend_module=backend,
            )
            log_project_module_plan(
                cfg=cfg,
                layer_name=layer_name,
                file_idx=file_idx,
                module_name=module_name,
                module_id=module_id,
                c_paths=c_paths,
                root_dir=root_dir,
                backend_module=backend,
            )
            section_state = prepare_project_module_section(
                doc,
                func_list_all,
                reuse_module=reuse_module,
                module_display=module_display,
                module_id=module_id,
                collect_unit_func_lists=collect_unit_func_lists,
                backend_module=backend,
            )
            module_table = section_state["module_table"]
            if section_state["unit_func_table"] is not None:
                unit_func_tables.append(section_state["unit_func_table"])
            if section_state["module_started"]:
                module_started = True
                resume_marker["module_started"] = True

            module_tasks = build_project_module_tasks(
                func_list_all,
                func_start_pos=func_start_pos,
                func_index=func_index,
                layer_name=layer_name,
                module_id=module_id,
            )

            # 增量生成：过滤未变更的任务
            if incremental_state is not None:
                try:
                    from .incremental import (
                        filter_tasks_for_incremental,
                        build_consistency_constraints,
                        detect_dependent_changes,
                    )

                    # 检测依赖变更
                    changed_func_names = []
                    for task in module_tasks:
                        body = task.get("body", "")
                        func_name = task.get("func_name", "")
                        file_path = task.get("source_file", "")
                        func_key = f"{file_path}::{func_name}"
                        prev_fp = incremental_state.function_fingerprints.get(func_key)
                        if prev_fp is None:
                            changed_func_names.append(func_name)

                    # 检测调用变更函数的依赖
                    dependent_keys = detect_dependent_changes(
                        changed_func_names,
                        func_list_all,
                        incremental_state,
                    )

                    tasks_to_gen, tasks_to_skip = filter_tasks_for_incremental(
                        module_tasks,
                        incremental_state,
                        force_all=False,
                    )

                    # 将依赖变更的函数也加入生成列表
                    for task in tasks_to_skip:
                        func_name = task.get("func_name", "")
                        file_path = task.get("source_file", "")
                        func_key = f"{file_path}::{func_name}"
                        if func_key in dependent_keys:
                            tasks_to_gen.append(task)
                            tasks_to_skip = [t for t in tasks_to_skip if f"{t.get('source_file', '')}::{t.get('func_name', '')}" != func_key]

                    if tasks_to_skip:
                        skipped_count += len(tasks_to_skip)
                        backend.vlog(cfg, f"[增量] 跳过 {len(tasks_to_skip)} 个未变更函数")
                        # 对于跳过的任务，直接使用缓存的设计
                        for skip_task in tasks_to_skip:
                            cached = skip_task.get("_cached_design")
                            if cached:
                                designs.append(cached)

                    # 构建一致性约束
                    if tasks_to_gen:
                        symbols_in_tasks = set()
                        for task in tasks_to_gen:
                            for elem in (task.get("local_vars") or []):
                                symbols_in_tasks.add(elem.get("name", ""))
                            for elem in (task.get("params") or []):
                                symbols_in_tasks.add(elem.get("name", ""))
                        consistency_constraints = build_consistency_constraints(
                            incremental_state,
                            list(symbols_in_tasks),
                        )
                        if consistency_constraints:
                            backend.vlog(cfg, f"[增量] 应用 {len(consistency_constraints)} 个一致性约束")
                            # 将约束注入到 cfg 的 extra_params
                            if cfg.extra_params is None:
                                cfg.extra_params = {}
                            cfg.extra_params["_consistency_constraints"] = consistency_constraints

                        # 语句级增量：注入缓存的逻辑语句和变更行号
                        from .incremental import (
                            compute_function_fingerprint,
                            detect_statement_changes,
                        )
                        for task in tasks_to_gen:
                            func_name = task.get("func_name", "")
                            file_path = task.get("source_file", "")
                            func_key = f"{file_path}::{func_name}"
                            prev_fp = incremental_state.function_fingerprints.get(func_key)
                            if prev_fp and prev_fp.cached_logic_lines:
                                # 计算当前指纹
                                body = task.get("body", "")
                                line_start = int(task.get("line_start", 0) or 0)
                                line_end = int(task.get("line_end", 0) or 0)
                                current_fp = compute_function_fingerprint(
                                    func_name=func_name,
                                    file_path=file_path,
                                    line_start=line_start,
                                    line_end=line_end,
                                    body=body,
                                    signature=task.get("prototype", ""),
                                )
                                # 检测语句级变更
                                changed_lines, unchanged_lines = detect_statement_changes(
                                    current_fp, prev_fp
                                )
                                if changed_lines and len(changed_lines) < len(prev_fp.cached_logic_lines):
                                    # 注入缓存到任务
                                    task["_cached_logic_lines"] = prev_fp.cached_logic_lines
                                    task["_changed_statement_lines"] = changed_lines
                                    backend.vlog(cfg, f"[增量] 语句级合并: {func_name} 变更 {len(changed_lines)}/{len(prev_fp.cached_logic_lines)} 行")

                    module_tasks = tasks_to_gen
                except Exception as exc:
                    backend.vlog(cfg, f"[增量] 过滤失败: {exc}")

            execution = execute_project_module_tasks(
                doc,
                module_tasks,
                cfg=cfg,
                root_dir=root_dir,
                module_table=module_table,
                collect_unit_rows=collect_unit_rows,
                unit_rows=unit_rows,
                unit_index=unit_index,
                resume_marker=resume_marker,
                designs=designs,
                backend_module=backend,
            )
            processed_tasks = int(execution["processed_tasks"])
            wrote_any = wrote_any or bool(execution["wrote_any"])
            resume_marker = dict(execution["resume_marker"])
            unit_index = int(execution["unit_index"])
            func_index = int(execution["func_index"])

            # Flush review bundle after each module so review HTML is available
            # even if generation is interrupted partway.
            _write_review_workspace_if_enabled(
                cfg, output, project_root=root_dir, merge_existing=True,
            )

            module_finalize = finalize_project_module_iteration(
                cfg=cfg,
                processed_tasks=processed_tasks,
                module_tasks=module_tasks,
                resume_marker=resume_marker,
                file_idx=file_idx,
                backend_module=backend,
            )
            resume_marker = dict(module_finalize["resume_marker"])
            stopped = bool(module_finalize["stopped"])
            if stopped:
                break
            start_func_pos = int(module_finalize["start_func_pos"])
            start_func_index = int(module_finalize["start_func_index"])
            resume_module_id = module_finalize["resume_module_id"]
            module_started = bool(module_finalize["module_started"])
        if stopped:
            break
        start_layer_idx = layer_idx + 1
        start_file_idx = 0
        layer_heading_started = False
        resume_marker = advance_project_resume_after_layer(resume_marker)

    if (not wrote_any) and (not stopped):
        raise backend.NoDataError("工程内没有可生成的函数内容。")

    try:
        _placeholder = run_state.get("placeholder")
        _body_start_idx = int(run_state.get("body_start_idx") or 0)
        if _placeholder is not None:
            backend.relocate_generated_blocks(doc, _body_start_idx, _placeholder)
        backend.safe_save_docx(doc, output)
    except Exception as exc:
        raise backend.RenderError(f"保存 Word 失败：{exc}") from exc
    _write_review_workspace_if_enabled(cfg, output, project_root=root_dir, merge_existing=continuing)
    _write_design_workspace_if_enabled(cfg, output, project_root=root_dir, merge_existing=continuing)
    finalize_project_generation(
        output,
        root_dir,
        cfg,
        stopped=stopped,
        resume_marker=resume_marker,
        collect_unit_rows=collect_unit_rows,
        unit_rows=unit_rows,
        collect_unit_func_lists=collect_unit_func_lists,
        unit_func_tables=unit_func_tables,
        designs=designs,
        backend_module=backend,
    )
    try:
        from . import graph_visuals

        html_path = graph_visuals.write_html_report(cfg, title=f"AutoDocGen 调用图谱 - {os.path.basename(os.path.abspath(root_dir))}")
        if html_path:
            backend.vlog(cfg, f"调用图谱 HTML 已生成：{html_path}")
    except Exception as exc:
        backend.vlog(cfg, f"[Graph] HTML 图谱生成失败：{exc}")

    # 增量生成：保存状态
    if incremental_state is not None:
        try:
            from .incremental import (
                update_state_with_new_designs,
                save_incremental_state,
            )
            # 提取一致性上下文（从已有设计中）
            incremental_state.extract_consistency_from_designs()
            # 更新状态（包含语句级缓存）
            update_state_with_new_designs(incremental_state, designs, func_data_list=func_list_all)
            save_incremental_state(root_dir, incremental_state)
            backend.vlog(cfg, f"[增量] 状态已保存，跳过 {skipped_count} 个函数")
        except Exception as exc:
            backend.vlog(cfg, f"[增量] 状态保存失败: {exc}")

    # 保存 docx 渲染缓存（下次复跑可直接插入 XML 元素，跳过渲染）
    try:
        from . import render as _render_mod
        _render_mod.save_render_cache(os.path.join(root_dir, ".autodoc", "render_cache.json"))
    except Exception:
        pass

    if generation.open_after_done and os.name == "nt":
        try:
            os.startfile(output)
        except Exception:
            pass
    backend.finalize_project_symbol_memory(cfg)
    maybe_write_evidence_report(output, cfg, backend_module=backend)
    if _EVIDENCE_ENABLED and logic_step_ir_enabled(cfg):
        try:
            suggested = auto_suggest_symbol_translations(cfg, project_root=root_dir)
            if suggested:
                backend.vlog(cfg, f"[LogicStep] AI 自动建议 {len(suggested)} 个符号翻译，已写入符号记忆库")
        except Exception as exc:
            backend.vlog(cfg, f"[LogicStep] 自动翻译建议失败：{exc}")


def collect_project_module_functions(
    c_paths: Sequence[str],
    *,
    preprocessed: Optional[dict[str, dict]],
    root_dir: str,
    cfg,
    prefilter: bool,
    backend_module=None,
) -> list[dict]:
    backend = backend_module or legacy_backend()
    func_list_all: list[dict] = []
    for c_path in c_paths:
        func_list = None
        skip_reason = None
        pre = (preprocessed or {}).get(c_path)
        if pre is not None:
            if pre.get("error"):
                backend.vlog(cfg, f"预解析失败，跳过：{c_path}；原因：{pre['error']}")
                continue
            func_list = pre.get("func_list")
            skip_reason = pre.get("skip_reason")
        if func_list is None:
            try:
                func_list, skip_reason = _prepare_func_list_for_c_file(
                    c_path,
                    project_root=root_dir,
                    cfg=cfg,
                    prefilter=prefilter,
                    backend_module=backend,
                )
            except backend.SourceReadError as exc:
                backend.vlog(cfg, f"读取失败，跳过：{c_path}；原因：{exc}")
                continue
        if not func_list:
            if skip_reason == "no_comment":
                backend.vlog(cfg, f"未发现注释，跳过：{c_path}")
            elif skip_reason == "no_func":
                backend.vlog(cfg, f"未发现函数，跳过：{c_path}")
            else:
                backend.vlog(cfg, f"未解析到函数，跳过：{c_path}")
            continue
        for fd in func_list:
            fd2 = dict(fd)
            file_context = dict(fd.get("file_context") or {})
            file_context["source_file"] = c_path
            fd2["file_context"] = file_context
            func_list_all.append(fd2)
    return func_list_all


def resolve_project_module_names(
    mod: dict,
    c_paths: Sequence[str],
    *,
    backend_module=None,
) -> tuple[str, str]:
    backend = backend_module or legacy_backend()
    module_name = utils_module._safe_strip((mod or {}).get("name"))
    if not module_name:
        module_code = ""
        try:
            module_code = backend.load_c_file(c_paths[0])
        except Exception:
            module_code = ""
        if module_code:
            module_name = backend._derive_module_display_name(c_paths[0], module_code)
        else:
            module_name = os.path.splitext(os.path.basename(c_paths[0]))[0]
    display_raw = backend._guess_cn_from_ident(module_name, glossary=backend.DOMAIN_GLOSSARY) or module_name
    if text_utils._contains_cjk(display_raw):
        module_display = backend._normalize_function_cn_title(
            display_raw, func_name=module_name, comment_desc=display_raw,
        )
    else:
        module_display = display_raw
    return module_name, module_display


def build_project_module_table_payload(
    func_list_all: Sequence[dict],
    *,
    module_id: str,
    module_display: str,
    include_unit_func_table: bool,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    entries: list[dict] = []
    func_rows: list[dict] = []
    for index, fd in enumerate(func_list_all, start=1):
        func_info = fd.get("func_info") or {}
        csu_name = _registered_function_title(fd, backend_module=backend)
        csu_id = f"{backend.normalize_req_prefix(module_id)}_{index:03d}"
        entries.append({"csu_name": csu_name, "csu_id": csu_id})
        if include_unit_func_table:
            func_rows.append({
                "index": index,
                "name": csu_name,
                "prototype": backend._format_func_prototype(func_info),
            })
    payload: dict[str, Any] = {"entries": entries}
    if include_unit_func_table:
        payload["unit_func_table"] = {
            "unit_name": module_display,
            "func_rows": func_rows,
        }
    return payload


def build_project_module_tasks(
    func_list_all: Sequence[dict],
    *,
    func_start_pos: int,
    func_index: int,
    layer_name: str,
    module_id: str,
) -> list[dict]:
    tasks: list[dict] = []
    next_module_func_index = int(func_index)
    for func_pos, func_data in enumerate(func_list_all, start=1):
        if func_pos < func_start_pos:
            continue
        func_name = ((func_data.get("func_info") or {}).get("func_name") or "")
        src_file = ((func_data.get("file_context") or {}).get("source_file") or "")
        func_info = dict(func_data.get("func_info") or {})
        # 把 filter_tasks_for_incremental 需要的字段直接挂在 task 上,
        # 否则 filter 拿到的 task 没有 body / prototype / line_start / line_end / source_file,
        # 导致 current_fp.body_hash 永远基于空 body,与 prev_fp 永远不等,
        # 增量模式跳过未变更函数的逻辑形同失效。
        # 这里 source_file 一律存空串,与 filter 的 func_key = "::func_name" 格式保持一致
        # (pipeline 之前用空串是因为 task["source_file"] 字段根本不存在)。
        tasks.append({
            "func_data": func_data,
            "module_req_prefix": module_id,
            "index": next_module_func_index,
            "func_pos": func_pos,
            "func_name": func_name,
            "file": str(src_file or ""),
            "source_file": "",  # 与 filter_tasks_for_incremental 的 func_key 格式一致
            "body": str(func_data.get("body", "") or ""),
            "prototype": str(func_info.get("prototype", "") or ""),
            "line_start": int(func_info.get("start", 0) or 0),
            "line_end": int(func_info.get("end", 0) or 0),
            "layer": layer_name,
            "module_id": module_id,
        })
        next_module_func_index += 1
    return tasks


def build_project_unit_row(
    task: dict[str, Any],
    design: Any,
    *,
    unit_index: int,
    backend_module=None,
) -> dict[str, Any]:
    backend = backend_module or legacy_backend()
    func_info = (task.get("func_data") or {}).get("func_info") or {}
    prototype = backend._format_func_prototype(func_info or task.get("func_data") or {})
    return {
        "index": unit_index,
        "name": getattr(design, "title", ""),
        "prototype": prototype,
        "req_id": getattr(design, "req_id", ""),
        "location": os.path.basename(str(task.get("file") or "")),
        "status": "新研",
        "purpose": getattr(design, "title", ""),
    }


def build_design_model(func_data: dict, module_req_prefix: str, index: int, cfg) -> DesignModel:
    backend = legacy_backend()
    design = build_function_design(func_data, module_req_prefix, index, cfg)
    func_info = dict((func_data or {}).get("func_info") or {})
    file_context = dict((func_data or {}).get("file_context") or {})
    return DesignModel(
        func_name=utils_module._safe_strip(func_info.get("func_name")),
        func_cn_name=utils_module._safe_strip(getattr(design, "title", "")),
        desc="\n".join(getattr(design, "description_lines", ()) or ()),
        params=[{"name": e.name, "ident": e.ident, "c_type": e.c_type, "direction": e.direction} for e in (getattr(design, "io_elements", ()) or ())],
        locals=[{"name": e.name, "ident": e.ident, "c_type": e.c_type, "usage": e.usage} for e in (getattr(design, "local_elements", ()) or ())],
        logic_steps=list(getattr(design, "logic_lines", ()) or ()),
        file_context=file_context,
    )


def generate_design_doc_from_file(source: str, output: str, cfg, resume_state: Optional[dict] = None):
    backend = legacy_backend()
    impl = getattr(backend, "_legacy_generate_design_doc_from_file_impl", None) or backend.generate_design_doc_from_file
    return impl(source, output, cfg, resume_state=resume_state)


def generate_design_doc_for_single_function(source: str, func_name: str, output: str, cfg, *, project_root: Optional[str] = None):
    backend = legacy_backend()
    impl = getattr(backend, "_legacy_generate_design_doc_for_single_function_impl", None) or backend.generate_design_doc_for_single_function
    return impl(source, func_name, output, cfg, project_root=project_root)


def _design_snapshot_path(doc_path: str) -> str:
    return doc_path + ".autodesign.json"


def save_design_snapshot(doc_path: str, csu_id: str, design: Any) -> None:
    path = _design_snapshot_path(doc_path)
    try:
        data = {}
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
        meta = getattr(design, "ai_meta", None)
        quality_issues = []
        unresolved = []
        if meta is not None:
            quality_issues = list(getattr(meta, "quality_issues", ()) or ())
            unresolved = list(getattr(meta, "unresolved_local_symbols", ()) or ()) + list(getattr(meta, "unresolved_param_symbols", ()) or ()) + list(getattr(meta, "unresolved_logic_symbols", ()) or ())
        data[csu_id] = {
            "title": getattr(design, "title", ""),
            "description": "\n".join(getattr(design, "description_lines", ()) or ()),
            "logic_lines": list(getattr(design, "logic_lines", ()) or ()),
            "io_elements": [{"name": getattr(e, "name", ""), "ident": getattr(e, "ident", "")} for e in (getattr(design, "io_elements", ()) or ())],
            "local_elements": [{"name": getattr(e, "name", ""), "ident": getattr(e, "ident", "")} for e in (getattr(design, "local_elements", ()) or ())],
            "quality_issues": quality_issues,
            "unresolved_symbols": unresolved,
        }
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def load_design_snapshot(doc_path: str, csu_id: str) -> Optional[dict]:
    path = _design_snapshot_path(doc_path)
    try:
        if not os.path.exists(path):
            return None
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get(csu_id)
    except Exception:
        return None


def _build_csu_body_elements(
    doc,
    target: dict,
    target_pos: int,
    csu_id: str,
    cfg,
    *,
    backend,
    render_module,
    include_heading: bool = False,
) -> tuple[list, Any]:
    """Build design, render into *target doc* (not a temp doc), extract body elements.

    Rendering into the target doc ensures all styles (609_4, Table Grid, etc.)
    are already present, so inserted elements inherit correct formatting.
    Returns (body_elements_after_heading, design).
    """
    module_id = backend.normalize_req_prefix(cfg.req_id_prefix)
    # Disable call graph rendering for in-place regen (no project context).
    regen_cfg = clone_cfg(cfg)
    try:
        regen_cfg._autodoc_graph_configured = False
    except Exception:
        pass
    # Reset per-function AI state to prevent cross-function leakage.
    for attr in ("_skip_ai_current_func", "_current_func_ai_failed",
                 "_ai_quality_feedback", "_current_render_func_data"):
        try:
            setattr(regen_cfg, attr, False if "skip" in attr or "failed" in attr else None)
        except Exception:
            pass
    try:
        regen_cfg._current_render_func_data = target
    except Exception:
        pass
    design = build_function_design_impl(
        target, module_id, target_pos, regen_cfg, backend_module=backend
    )
    design = replace(design, req_id=csu_id)

    # Insert a temporary marker, render after it, then extract only the new
    # elements between the marker and the trailing sectPr. python-docx inserts
    # paragraphs before sectPr, so len(body)-based slicing is not reliable.
    body = doc._body._element
    marker = doc.add_paragraph("__AUTODOC_REGEN_MARKER__")
    marker_elem = marker._element
    render_module.render_function_design(doc, design, regen_cfg)
    new_elems = []
    body_after_render = list(body)
    try:
        marker_idx = body_after_render.index(marker_elem)
    except ValueError:
        marker_idx = -1
    if marker_idx < 0:
        raise RuntimeError("failed to locate temporary CSU render marker")
    end_idx = len(body_after_render)
    if marker_idx >= 0:
        for idx in range(marker_idx + 1, len(body_after_render)):
            tag = getattr(body_after_render[idx], "tag", "")
            if tag.endswith("}sectPr"):
                end_idx = idx
                break
    for elem in body_after_render[marker_idx + 1:end_idx]:
        tag = getattr(elem, "tag", "")
        if tag.endswith("}p") or tag.endswith("}tbl"):
            new_elems.append(elem)
    # Detach from doc (remove from body) so replace_csu_in_doc can re-insert.
    for elem in new_elems:
        body.remove(elem)
    try:
        body.remove(marker_elem)
    except Exception:
        pass
    # First element is the Heading 4. Replacements keep the old heading in
    # place; insertions need the generated heading too.
    if (not include_heading) and new_elems and getattr(new_elems[0], "tag", "").endswith("}p"):
        new_elems = new_elems[1:]
    try:
        regen_cfg._current_render_func_data = None
    except Exception:
        pass
    return new_elems, design


def _update_module_table_csu_name(
    doc,
    csu_id: str,
    new_title: str,
    *,
    backend,
    render_module,
) -> None:
    """Update the CSU name in the module function table, robustly."""
    m = re.match(r"^(.+?)_(\d+)$", csu_id)
    if not m:
        return
    mod_id = m.group(1)
    csu_row = int(m.group(2))
    for table in doc.tables:
        if not render_module.is_csu_table(table, backend_module=backend):
            continue
        try:
            # Check column 1 (CSC标识) for module id — handle merged cells.
            cell_text = ""
            if len(table.rows) > 1:
                try:
                    cell_text = table.cell(1, 1).text or ""
                except Exception:
                    cell_text = ""
            if cell_text.strip() != mod_id:
                continue
            # Row index = csu_row (1-based: row 0 = header, row 1 = first CSU).
            if len(table.rows) > csu_row:
                target_cell = table.cell(csu_row, 2)
                # Clear existing paragraphs, set new text.
                target_cell.text = new_title
        except Exception:
            continue


def _write_back_incremental_cache(
    project_root: str,
    source: str,
    func_name: str,
    design,
) -> None:
    """Update incremental_state.json so next project-gen skips this function."""
    try:
        from .incremental import load_incremental_state, save_incremental_state, update_state_with_new_designs
        state = load_incremental_state(project_root)
        func_key = f"::{func_name}"  # matches existing key format (empty file path)
        design_dict = {
            "func_name": func_name,
            "title": getattr(design, "title", ""),
            "source_file": "",
            "io_elements": [
                {"ident": e.ident, "name": e.name}
                for e in (getattr(design, "io_elements", ()) or ())
            ],
            "local_elements": [
                {"ident": e.ident, "name": e.name}
                for e in (getattr(design, "local_elements", ()) or ())
            ],
            "logic_lines": list(getattr(design, "logic_lines", ()) or ()),
        }
        state.generated_designs[func_key] = design_dict
        save_incremental_state(project_root, state)
    except Exception:
        pass


def regenerate_csu_in_doc(
    doc_path: str,
    source: str,
    func_name: str,
    csu_id: str,
    cfg,
    *,
    project_root: Optional[str] = None,
    doc: Optional[Any] = None,
    save: bool = True,
    prepared_func_cache: Optional[dict[tuple[str, str], tuple[list[dict[str, Any]], Any]]] = None,
) -> dict[str, Any]:
    """Regenerate a single CSU and replace it in-place in an existing document.

    Renders into the *target* document (not a temp doc) so styles are inherited.
    Disables call graph rendering (no project context for in-place regen).
    Resets per-function AI state to prevent cross-function leakage.
    Writes back incremental cache so next project-gen skips this function.

    Args:
        doc_path: Path to the .docx to modify.
        source: Path to the .c source file.
        func_name: Function name to regenerate.
        csu_id: CSU identifier (e.g. D/R_SDD01_009_007) to locate in doc.
        cfg: GenConfig.
        project_root: Project root (for incremental cache and header resolution).
        doc: Pre-opened Document object. If None, opens and saves doc_path.
             When provided, caller is responsible for saving (batch mode).
        save: If True (default), save the document after replacement.
              Set False in batch mode.

    Returns ``{"found": bool, "replaced": int, "old_title": str, "new_title": str, "saved": bool}``.
    """
    backend = legacy_backend()
    from . import render as render_module
    from docx import Document

    func_name = (func_name or "").strip()
    csu_id = (csu_id or "").strip()
    if not func_name or not csu_id:
        raise ValueError("func_name 和 csu_id 不能为空")

    project_root = project_root or backend._guess_project_root_for_source(source)

    # 1. Parse source and locate the target function.
    cache_key = (os.path.abspath(source), os.path.abspath(project_root or ""))
    if prepared_func_cache is not None and cache_key in prepared_func_cache:
        func_list, _prepare_meta = prepared_func_cache[cache_key]
    else:
        func_list, _prepare_meta = _prepare_func_list_for_c_file(
            source,
            project_root=project_root,
            cfg=cfg,
            prefilter=False,
            backend_module=backend,
        )
        if prepared_func_cache is not None:
            prepared_func_cache[cache_key] = (func_list, _prepare_meta)
    if not func_list:
        raise backend.NoDataError(f"未解析到任何函数：{source}")

    target = None
    target_pos = 1
    for idx, fd in enumerate(func_list, start=1):
        fi = fd.get("func_info") or {}
        if (fi.get("func_name") or "").strip() == func_name:
            target = fd
            target_pos = idx
            break
    if target is None:
        raise backend.NoDataError(f"未在文件中找到函数：{func_name}")

    # 2. Open target doc (or use caller-provided doc for batch mode).
    owns_doc = doc is None
    if owns_doc:
        doc = Document(doc_path)

    # 2b. Load previous design snapshot for iterative improvement
    revision_ctx = load_design_snapshot(doc_path, csu_id)
    if revision_ctx is not None:
        if cfg.extra_params is None:
            cfg.extra_params = {}
        cfg.extra_params["revision_context"] = revision_ctx

    # 3. Render into the TARGET doc (inherits all styles: 609_4, Table Grid, etc.)
    new_elements, design = _build_csu_body_elements(
        doc, target, target_pos, csu_id, cfg,
        backend=backend, render_module=render_module,
    )

    # 4. Replace in-place.
    new_heading_text = f"{design.title}（{csu_id}）"
    result = render_module.replace_csu_in_doc(
        doc,
        csu_id=csu_id,
        new_heading_text=new_heading_text,
        new_body_elements=new_elements,
        backend_module=backend,
    )

    if not result.get("found"):
        if owns_doc:
            pass  # No changes, don't save.
        return {**result, "saved": False}

    # 5. Update module table if title changed.
    _update_module_table_csu_name(doc, csu_id, design.title,
                                  backend=backend, render_module=render_module)

    # 6. Save (unless in batch mode where caller saves).
    if owns_doc and save:
        backend.safe_save_docx(doc, doc_path)
        result["saved"] = True
    elif not owns_doc:
        result["saved"] = False  # Caller will save.
    else:
        result["saved"] = True  # owns_doc but save=False (unusual).

    # 6b. Save updated design snapshot
    save_design_snapshot(doc_path, csu_id, design)

    # 7. Write back incremental cache.
    _write_back_incremental_cache(project_root, source, func_name, design)

    return result


def insert_csu_after_in_doc(
    doc_path: str,
    source: str,
    func_name: str,
    csu_id: str,
    after_csu_id: str,
    cfg,
    *,
    project_root: Optional[str] = None,
    doc: Optional[Any] = None,
    save: bool = True,
    prepared_func_cache: Optional[dict[tuple[str, str], tuple[list[dict[str, Any]], Any]]] = None,
) -> dict[str, Any]:
    """Regenerate a function as a complete CSU and insert it after another CSU."""
    backend = legacy_backend()
    from . import render as render_module
    from docx import Document

    func_name = (func_name or "").strip()
    csu_id = (csu_id or "").strip()
    after_csu_id = (after_csu_id or "").strip()
    if not func_name or not csu_id or not after_csu_id:
        raise ValueError("func_name、csu_id 和 after_csu_id 不能为空")

    project_root = project_root or backend._guess_project_root_for_source(source)

    cache_key = (os.path.abspath(source), os.path.abspath(project_root or ""))
    if prepared_func_cache is not None and cache_key in prepared_func_cache:
        func_list, _prepare_meta = prepared_func_cache[cache_key]
    else:
        func_list, _prepare_meta = _prepare_func_list_for_c_file(
            source,
            project_root=project_root,
            cfg=cfg,
            prefilter=False,
            backend_module=backend,
        )
        if prepared_func_cache is not None:
            prepared_func_cache[cache_key] = (func_list, _prepare_meta)
    if not func_list:
        raise backend.NoDataError(f"未解析到任何函数：{source}")

    target = None
    target_pos = 1
    for idx, fd in enumerate(func_list, start=1):
        fi = fd.get("func_info") or {}
        if (fi.get("func_name") or "").strip() == func_name:
            target = fd
            target_pos = idx
            break
    if target is None:
        raise backend.NoDataError(f"未在文件中找到函数：{func_name}")

    owns_doc = doc is None
    if owns_doc:
        doc = Document(doc_path)

    new_elements, design = _build_csu_body_elements(
        doc,
        target,
        target_pos,
        csu_id,
        cfg,
        backend=backend,
        render_module=render_module,
        include_heading=True,
    )

    result = render_module.insert_csu_after_in_doc(
        doc,
        after_csu_id=after_csu_id,
        new_elements=new_elements,
        backend_module=backend,
    )
    result["new_title"] = getattr(design, "title", "")
    result["csu_id"] = csu_id

    if not result.get("found"):
        return {**result, "saved": False}

    module_id_match = re.match(r"^(.+?)_\d+$", csu_id)
    if module_id_match:
        result["module_table"] = render_module.sync_module_function_table_for_module(
            doc,
            module_id_match.group(1),
            backend_module=backend,
        )

    if owns_doc and save:
        backend.safe_save_docx(doc, doc_path)
        result["saved"] = True
    elif not owns_doc:
        result["saved"] = False
    else:
        result["saved"] = True

    _write_back_incremental_cache(project_root, source, func_name, design)
    return result


def generate_design_doc_for_project(root_dir: str, output: str, cfg, resume_state: Optional[dict] = None):
    backend = legacy_backend()
    impl = getattr(backend, "_legacy_generate_design_doc_for_project_impl", None) or backend.generate_design_doc_for_project
    return impl(root_dir, output, cfg, resume_state=resume_state)


def _normalize_project_module_paths(mod: Any) -> list[str]:
    if not isinstance(mod, dict):
        return []
    return [
        os.path.abspath(str(path))
        for path in (mod.get("files") or [])
        if str(path).strip()
    ]


def __getattr__(name: str) -> Any:
    return getattr(legacy_backend(), name)


__all__ = [
    "DesignModel",
    "FunctionBuildResult",
    "_iter_function_design_results",
    "build_design_model",
    "build_design_name_map",
    "build_design_io_elements",
    "build_design_ai_meta",
    "apply_project_file_order_override",
    "build_project_modules_from_order",
    "build_project_layer_sets",
    "build_project_resume_marker",
    "execute_single_export_task",
    "execute_single_file_tasks",
    "execute_project_module_tasks",
    "finalize_project_module_iteration",
    "initialize_project_run_state",
    "log_project_module_plan",
    "prepare_project_file_iteration",
    "prepare_project_layer_iteration",
    "plan_project_module_run",
    "plan_project_source_layout",
    "preprocess_project_files",
    "prepare_project_module_section",
    "prepare_project_progress",
    "build_project_func_end_event",
    "build_project_func_start_event",
    "compute_project_module_id",
    "collect_project_module_entries_and_units",
    "collect_project_module_tables_data",
    "collect_project_c_files_by_layer",
    "build_project_module_table_payload",
    "build_project_module_tasks",
    "build_project_unit_row",
    "build_single_file_func_end_event",
    "build_single_file_func_start_event",
    "build_single_file_table_payload",
    "build_single_file_tasks",
    "collect_design_components",
    "collect_project_module_functions",
    "get_ordered_project_c_files",
    "assemble_function_design",
    "count_single_file_progress_total",
    "finalize_single_file_task_iteration",
    "finalize_project_generation",
    "finalize_single_file_generation",
    "record_project_task_resume_progress",
    "record_single_file_resume_progress",
    "advance_project_resume_after_module",
    "advance_project_resume_after_layer",
    "count_project_progress_total",
    "resolve_project_resume_checkpoint",
    "resolve_project_module_names",
    "resolve_source_module_names",
    "repair_design_local_profiles",
    "should_include_function",
    "build_design_text_sections",
    "update_project_module_table_title",
    "collect_design_quality_inputs",
    "build_design_local_elements",
    "build_design_output",
    "build_design_logic_lines",
    "build_function_design_impl",
    "build_function_design",
    "generate_design_doc_for_project",
    "generate_design_doc_for_single_function",
    "generate_design_doc_from_file",
    "iter_function_build_results",
    "prepare_design_context",
    "run_project_generation",
    "run_single_export_generation",
    "run_single_file_generation",
    "run_multi_call_design_enrichment",
    "run_one_call_design_enrichment",
    "run_function_design_task",
    "regenerate_csu_in_doc",
    "insert_csu_after_in_doc",
]
