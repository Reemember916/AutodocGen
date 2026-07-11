#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 PROJECT-2007 维护422通信协议 Word 文档。
依赖: python-docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "/Users/ree/Downloads/PROJECT-2007-0613/output/doc/PROJECT-2007_维护422通信协议.docx"

# ---------- 样式辅助 ----------

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '808080')
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_table(doc, headers, rows, col_widths_cm=None, header_bg='1F4E79', header_fg='FFFFFF'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    # 表头
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_bg(hdr[i], header_bg)
        set_cell_borders(hdr[i])
    # 数据行
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            set_cell_borders(cells[c_idx])
            # 隔行底色
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], 'F2F2F2')
    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 0:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor.from_string('1F4E79')
    elif level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string('1F4E79')
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string('2E74B5')
    return h

def add_para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string('808080')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

# ---------- 文档构建 ----------

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# 默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 标题页 =====
add_heading(doc, 'PROJECT-2007 空中加受油控制器', level=0)
add_heading(doc, '维护422通信协议说明书', level=0)
add_para(doc, '', size=6)
p = doc.add_paragraph()
run = p.add_run('用于编写维护422上位机软件')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string('808080')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
run = p.add_run('文档版本: V1.0    日期: 2026-06-30')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor.from_string('808080')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ===== 1. 物理层与链路层 =====
add_heading(doc, '1. 物理层与链路层', level=1)
add_table(doc,
    ['项', '值'],
    [
        ['接口', 'RS422 全双工异步'],
        ['SCI端口', 'SCIA'],
        ['波特率', '115200 bit/s'],
        ['数据格式', '8数据位、无校验、1停止位'],
        ['字节序', '多字节字段小端(低字节在前)'],
        ['帧间隔', '维护态: 10ms/包, 5包/轮; 下载态: 2ms/16字节'],
    ],
    col_widths_cm=[4, 12])

doc.add_paragraph()

# ===== 2. 通用帧结构 =====
add_heading(doc, '2. 通用帧结构', level=1)
add_para(doc, '所有维护帧固定16字节, 结构统一:')
add_table(doc,
    ['字节偏移', '长度', '内容'],
    [
        ['0', '1', '帧头1'],
        ['1', '1', '帧头2'],
        ['2', '1', '包号(发送) / 帧计数(接收)'],
        ['3', '1', '帧计数(发送) / 维护指令码(接收)'],
        ['4~13', '10', '载荷(按包号/指令码不同定义)'],
        ['14', '1', '载荷(部分包号使用)'],
        ['15', '1', '校验和'],
    ],
    col_widths_cm=[3, 2, 11])
add_para(doc, '校验算法: checksum = (~(sum(byte[0..14])) + 1) & 0xFF', bold=True)

doc.add_paragraph()

# ===== 3. 上行: 控制器->上位机 =====
add_heading(doc, '3. 上行: 控制器 -> 上位机 (发送方向)', level=1)
add_para(doc, '帧头固定 0x55 0xAA, 帧长16字节。字节2为包号, 决定载荷语义。每轮5包顺序发送: ID0->ID1->ID2->ID3->ID4。')
add_table(doc,
    ['字节', '内容'],
    [
        ['0', '0x55'],
        ['1', '0xAA'],
        ['2', '包号(0~4)'],
        ['3', '帧计数(0~255循环)'],
        ['4~14', '载荷(按包号定义)'],
        ['15', '校验和'],
    ],
    col_widths_cm=[3, 13])

# --- 包0 ---
add_heading(doc, '3.1 包0 基础状态', level=2)
add_table(doc,
    ['字节', '位段', '内容'],
    [
        ['4', '[7:0]', '系统时间字节0(最低)'],
        ['5', '[7:0]', '系统时间字节1'],
        ['6', '[7:0]', '系统时间字节2'],
        ['7', '[7:0]', '系统时间字节3(最高), 单位ms'],
        ['8', '[3:0]', '当前系统状态(0=初始,1=工作,2=安全,3=维护,4=掉电)'],
        ['8', '[7:4]', '上一拍系统状态'],
        ['9', '[3:0]', '当前工作模式'],
        ['9', '[7:4]', '上一拍工作模式'],
        ['10', '[3:0]', '运行期控制权(0=备,1=主)'],
        ['10', '[7:4]', '静态主备身份'],
        ['11', 'bit0', '控制输出有效'],
        ['11', 'bit1', '本端CHV资格有效'],
        ['11', 'bit3', '本端通道号(0=通道1,1=通道2)'],
        ['11', '[7:4]', '空中加油模式'],
        ['12', '[3:0]', '控制功能码'],
        ['12', '[7:4]', '上一拍控制功能码'],
        ['13', '[7:0]', 'CHV回采信号快照'],
        ['14', '[7:0]', '空中加油结束状态'],
    ],
    col_widths_cm=[2, 2, 12])

# --- 包1 ---
add_heading(doc, '3.2 包1 主备轮值诊断', level=2)
add_table(doc,
    ['字节', '内容'],
    [
        ['4', '本端记录的冷启动默认主通道ID'],
        ['5', '对端基础帧上报的冷启动默认主通道ID'],
        ['6', '本次启动仲裁得到的主通道ID'],
        ['7', '通道类型判别结果码'],
        ['8~14', '保留(0)'],
    ],
    col_widths_cm=[3, 13])

# --- 包2 ---
add_heading(doc, '3.3 包2 输出授权诊断', level=2)
add_table(doc,
    ['字节', '位段', '内容'],
    [
        ['4', '[3:0]', '运行期控制权'],
        ['4', '[7:4]', '静态主备身份'],
        ['5', '[3:0]', '本通道ID'],
        ['5', '[7:4]', '启动判型结果码'],
        ['6', '[7:0]', '本端CHV资格'],
        ['7', '[7:0]', '控制输出状态'],
        ['8', '[7:0]', 'CHV输入快照低字节'],
        ['9', '[7:0]', 'CHV输入快照高字节'],
        ['10', '[7:0]', '对端在线状态'],
        ['11', '[7:0]', '对端控制权观测'],
        ['12', '[3:0]', '系统状态'],
        ['12', '[7:4]', '工作模式'],
        ['13', '[3:0]', '控制功能'],
        ['13', '[7:4]', '维护功能'],
        ['14', 'bit0', '当前为主控'],
        ['14', 'bit1', '本端CHV资格有效'],
        ['14', 'bit2', '本端myCHV有效'],
        ['14', 'bit3', '控制输出有效'],
        ['14', 'bit4', '对端在线'],
        ['14', 'bit5', '观测到对端持有控制权'],
    ],
    col_widths_cm=[2, 2, 12])

# --- 包3 ---
add_heading(doc, '3.4 包3 BIT/故障摘要', level=2)
add_table(doc,
    ['字节', '内容'],
    [
        ['4', '上电BIT关键故障(0=正常,1=故障)'],
        ['5', '周期BIT故障等级(0~7)'],
        ['6', '维护BIT故障等级(0~7)'],
        ['7', '控制故障位图: bit0=通信故障, bit1=测量故障, bit2=不平衡故障, bit3=综合故障'],
        ['8', '控制故障原因码'],
        ['9', 'IFBIT低32位签名 字节0'],
        ['10', 'IFBIT低32位签名 字节1'],
        ['11', 'IFBIT低32位签名 字节2'],
        ['12', 'IFBIT低32位签名 字节3'],
        ['13', 'MBIT低16位签名 字节0'],
        ['14', 'MBIT低16位签名 字节1'],
    ],
    col_widths_cm=[3, 13])

# --- 包4 ---
add_heading(doc, '3.5 包4 通信来源/余度来源', level=2)
add_table(doc,
    ['字节', '位段', '内容'],
    [
        ['4', '[1:0]', 'RIU来源'],
        ['4', '[3:2]', 'CCDL来源'],
        ['4', '[5:4]', 'KZZZ来源'],
        ['5', '[7:0]', 'RIU心跳数据状态'],
        ['6', '[7:0]', '左吊舱KZZZ数据状态'],
        ['7', '[7:0]', '右吊舱KZZZ数据状态'],
        ['8', '[7:0]', 'CCDL系统状态数据状态'],
        ['9', '[7:0]', 'RIU心跳数据快照'],
        ['10', '[7:0]', '左吊舱KZZZ数据快照'],
        ['11', '[7:0]', '右吊舱KZZZ数据快照'],
        ['12', '[7:0]', 'CCDL系统状态数据快照'],
        ['13', '[3:0]', 'CCDL对端主备身份'],
        ['13', '[7:4]', 'CCDL对端默认主通道ID'],
        ['14', 'bit0', 'RIU心跳有效'],
        ['14', 'bit1', '左KZZZ有效'],
        ['14', 'bit2', '右KZZZ有效'],
        ['14', 'bit3', 'CCDL系统状态有效'],
        ['14', 'bit4', 'RIU来源无效标志'],
        ['14', 'bit5', 'CCDL来源无效标志'],
        ['14', 'bit6', 'KZZZ来源无效标志'],
    ],
    col_widths_cm=[2, 2, 12])
add_note(doc, '数据状态码: 0=错误(ERR), 1=本地RIU1/左吊舱, 2=本地RIU2/右吊舱, 3=对端CCDL镜像')

doc.add_page_break()

# ===== 4. 下行: 上位机->控制器 =====
add_heading(doc, '4. 下行: 上位机 -> 控制器 (接收方向)', level=1)
add_para(doc, '帧头固定 0xEB 0x90, 帧长16字节。字节3为维护指令码, 决定载荷语义。')
add_table(doc,
    ['字节', '内容'],
    [
        ['0', '0xEB'],
        ['1', '0x90'],
        ['2', '帧计数(上位机维护, 用于丢帧检测)'],
        ['3', '维护指令码(见下表)'],
        ['4~12', '载荷(按指令码定义)'],
        ['13', '维护状态码'],
        ['15', '校验和'],
    ],
    col_widths_cm=[3, 13])

# --- 维护指令码 ---
add_heading(doc, '4.1 维护指令码 (字节3)', level=2)
add_table(doc,
    ['码值', '宏定义', '含义', '载荷定义'],
    [
        ['0x00', 'MAINT_CODE_CMD_INVALID', '无效', '-'],
        ['0x11', 'MAINT_CODE_COMM_INFO', '通信状态信息请求', '见4.2'],
        ['0x66', 'MAINT_CODE_GROUND_CON', '地面控制指令', '见4.3'],
        ['0x99', 'MAINT_CODE_MAINT_FUNC', '地面维护功能', '见4.4'],
    ],
    col_widths_cm=[2, 5, 4, 5])

# --- 维护状态码 ---
add_heading(doc, '4.2 维护状态码 (字节13)', level=2)
add_table(doc,
    ['码值', '含义'],
    [
        ['0x00', '无效'],
        ['0x11', '通信状态信息'],
        ['0x22', '地面维护状态'],
    ],
    col_widths_cm=[3, 13])

# --- 指令0x11 ---
add_heading(doc, '4.3 指令0x11 通信状态信息 载荷', level=2)
add_table(doc,
    ['字节', '内容'],
    [
        ['4', '低4位=429通信信息ID(0~7), 高4位=流量测量盒422信息ID(0~1)'],
        ['5', '429接收标号'],
        ['6~7', '保留'],
        ['8', '读取地址字节0(最低字节)'],
        ['9', '读取地址字节1'],
        ['10', '读取地址字节2(最高字节)'],
        ['11~12', '保留'],
    ],
    col_widths_cm=[3, 13])
add_note(doc, '429信息ID: 0/1=RMC1/2, 2/3=RIU1/2, 4/5=JYB1/2, 6=DMP, 7=KZZZ')

# --- 指令0x66 ---
add_heading(doc, '4.4 指令0x66 地面控制 载荷', level=2)
add_table(doc,
    ['字节', '内容'],
    [
        ['4', '指令更新标志(0x34=有效, 0x00=无效)'],
        ['5', '维护指令ID(0/1=加油泵1/2, 2=排气阀, 3=电磁阀, 4=控制装置)'],
        ['6~7', '指令数据(当前版本仅透传泵/阀索引, KZZZ已清退)'],
    ],
    col_widths_cm=[3, 13])

# --- 指令0x99 ---
add_heading(doc, '4.5 指令0x99 地面维护功能 载荷', level=2)
add_table(doc,
    ['字节', '内容'],
    [
        ['4', '维护功能码(见4.6)'],
        ['5~7', '保留'],
        ['8~11', '下载起始地址(小端4字节, 仅DATA_DOWNLOAD时有效)'],
        ['12~14', '下载地址长度(小端3字节, 仅DATA_DOWNLOAD时有效)'],
    ],
    col_widths_cm=[3, 13])
add_note(doc, '注: 源码中下载长度按4字节解析, 建议上位机按4字节小端填写字节12~15, 字节15为校验和不变。')

# --- 维护功能码 ---
add_heading(doc, '4.6 维护功能码 (字节4, 配合指令0x99)', level=2)
add_table(doc,
    ['码值', '宏定义', '含义', '触发动作'],
    [
        ['0x00', 'GROUND_MAINT_FUNC_INVALID', '无效', '-'],
        ['0x01', 'GROUND_MAINT_FUNC_SOFT_CRC', '软件CRC计算', '计算应用区CRC16并上报'],
        ['0x02', 'GROUND_MAINT_FUNC_DATA_DOWNLOAD', '数据下载', '按起始地址+长度发送FLASH裸数据'],
        ['0x03', 'GROUND_MAINT_FUNC_DATA_ERASE', '信息擦除', '擦除FLASH存储区'],
        ['0x04', 'GROUND_MAINT_FUNC_PID_PARA_ADJUST', '控制参数修改', '当前版本不支持'],
    ],
    col_widths_cm=[2, 6, 3, 5])

# --- 执行结果 ---
add_heading(doc, '4.7 维护功能执行结果', level=2)
add_table(doc,
    ['码值', '含义'],
    [
        ['0x00', '未执行'],
        ['0x01', '执行成功'],
        ['0x02', '功能暂不支持'],
        ['0x03', '参数无效'],
    ],
    col_widths_cm=[3, 13])

doc.add_page_break()

# ===== 5. 下载模式 =====
add_heading(doc, '5. 下载模式 (特殊发送)', level=1)
add_para(doc, '收到指令0x99+功能码0x02后, 控制器进入下载模式, 不再发送常规16字节帧, 改为按2ms间隔连续发送FLASH裸数据:')
add_table(doc,
    ['项', '值'],
    [
        ['帧格式', '无帧头无校验, 纯裸数据'],
        ['每帧长度', '16字节'],
        ['发送间隔', '2ms'],
        ['地址折算', '窗口内取模 + 应用区基址'],
        ['长度截断', '超窗时截断到窗口末端'],
        ['提前终止', '中途28V掉电则提前终止'],
    ],
    col_widths_cm=[4, 12])

doc.add_paragraph()

# ===== 6. 时序要点 =====
add_heading(doc, '6. 时序要点', level=1)
add_table(doc,
    ['项', '说明'],
    [
        ['维护态进入', '地面条件 + 维护开关有效 + 连续1s接收维护指令(指令0x11且状态码0x22)'],
        ['维护态退出', '维护开关无效即退出'],
        ['下载门禁', '仅地面维护状态可执行下载'],
        ['发送使能', '仅在控制输出有效(CON_OUT_STATE_VALID)时发送, 否则静默'],
        ['轮询周期', '5包/轮, 每10ms发1包, 单包更新率50ms'],
    ],
    col_widths_cm=[4, 12])

# ===== 保存 =====
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print("SAVED:", OUT_PATH)
print("SIZE:", os.path.getsize(OUT_PATH), "bytes")
