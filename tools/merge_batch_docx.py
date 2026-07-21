#!/usr/bin/env python3
"""Merge multiple batch-generated docx files into one complete document.

Reads all `batch_*.docx` files in an input directory (excluding `*_unit_func_list.docx`),
preserves styles/bold/tables, and promotes:
  - Heading 3 (CSC title) → Heading 1 (top-level module)
  - Heading 4 (CSU title) → Heading 2 (function)
Page breaks are inserted before each CSC title (except the first).

Usage:
    python3 tools/merge_batch_docx.py [INPUT_DIR] [-o OUTPUT] [--dry-run]
    python3 tools/merge_batch_docx.py --help

Default INPUT_DIR is tmp/project_batch, OUTPUT is INPUT_DIR/../<dirname>_merged.docx.
"""
import argparse
import os
import re
import sys
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


CSC_HEADING_RE = re.compile(r"（D/R_")
AI_HEADER_RE = re.compile(r"【AI 辅助已启用】")


def is_csc_heading(p):
    return p.style.name == "Heading 3" and bool(CSC_HEADING_RE.search(p.text))


def is_csu_heading(p):
    return p.style.name == "Heading 4"


def is_ai_header(p):
    return bool(AI_HEADER_RE.search(p.text))


def append_to_target(target_doc, element, *, style_name=None, page_break=False):
    """Append a deep copy of `element` to target body, before sectPr.
    If style_name is given and element is a paragraph, set its pStyle val to that style_id.
    If page_break is True, add <w:pageBreakBefore/> to the paragraph's pPr.
    """
    sectPr = target_doc.element.body.find(qn("w:sectPr"))
    new_elem = deepcopy(element)
    target_doc.element.body.insert(list(target_doc.element.body).index(sectPr), new_elem)
    if style_name and new_elem.tag == qn("w:p"):
        pPr = new_elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            new_elem.insert(0, pPr)
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            pStyle = OxmlElement("w:pStyle")
            pPr.insert(0, pStyle)
        try:
            style_id = target_doc.styles[style_name].style_id
            pStyle.set(qn("w:val"), style_id)
        except KeyError:
            pass
    if page_break and new_elem.tag == qn("w:p"):
        pPr = new_elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            new_elem.insert(0, pPr)
        if pPr.find(qn("w:pageBreakBefore")) is None:
            pPr.append(OxmlElement("w:pageBreakBefore"))


def collect_batch_files(input_dir):
    """Return sorted list of main docx files (excluding *_unit_func_list.docx)."""
    files = []
    for fname in sorted(os.listdir(input_dir)):
        if not fname.endswith(".docx"):
            continue
        if "_unit_func_list" in fname:
            continue
        if fname.startswith("~$"):  # skip Word lock files
            continue
        files.append(os.path.join(input_dir, fname))
    return files


def merge(input_dir, output_path, dry_run=False):
    files = collect_batch_files(input_dir)
    print(f"Found {len(files)} batch docx files in {input_dir}")
    if not files:
        print("ERROR: no batch docx files found", file=sys.stderr)
        return 1

    if dry_run:
        print("DRY RUN — would merge:")
        for f in files:
            print(f"  {os.path.basename(f)}")
        return 0

    # Clone first file to inherit its style definitions
    target = Document(files[0])
    body = target.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        print(f"ERROR: first file {files[0]} has no sectPr", file=sys.stderr)
        return 1
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    ai_header_inserted = False
    csc_count = 0
    for fp in files:
        print(f"\n=== merging: {os.path.basename(fp)} ===", flush=True)
        src = Document(fp)
        src_paras = src.paragraphs
        p_idx = 0
        for child in list(src.element.body):
            tag = child.tag.split("}")[-1]
            if tag == "sectPr":
                continue
            if tag == "p":
                p = src_paras[p_idx] if p_idx < len(src_paras) else None
                p_idx += 1
                if p is None:
                    continue
                if is_ai_header(p):
                    if not ai_header_inserted:
                        append_to_target(target, child)
                        ai_header_inserted = True
                    continue
                if is_csc_heading(p):
                    append_to_target(target, child, style_name="Heading 1",
                                     page_break=(csc_count > 0))
                    csc_count += 1
                    continue
                if is_csu_heading(p):
                    append_to_target(target, child, style_name="Heading 2")
                    continue
                append_to_target(target, child)
                continue
            if tag == "tbl":
                append_to_target(target, child)
                continue
            # Other body elements (e.g., sectPr) are skipped

    os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
    target.save(output_path)

    print(f"\n=== DONE ===")
    print(f"output: {output_path}")
    print(f"size: {os.path.getsize(output_path) // 1024} KB")
    print(f"CSC count: {csc_count}")

    # Verify
    d = Document(output_path)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    tables = len(d.tables)
    h1 = sum(1 for p in d.paragraphs if p.style.name == "Heading 1")
    h2 = sum(1 for p in d.paragraphs if p.style.name == "Heading 2")
    print(f"verification: paragraphs={len(paras)} tables={tables} "
          f"Heading1={h1} Heading2={h2}")
    return 0


def main():
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
    default_input = os.path.join(repo_root, "tmp", "project_batch")
    default_output = os.path.join(repo_root, "tmp", "PROJECT_merged.docx")

    ap = argparse.ArgumentParser(
        description="Merge batch docx files into a complete document with Heading 1/2 structure.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    ap.add_argument("input_dir", nargs="?", default=default_input,
                    help=f"directory containing batch_*.docx files (default: {default_input})")
    ap.add_argument("-o", "--output", default=default_output,
                    help=f"output path (default: {default_output})")
    ap.add_argument("--dry-run", action="store_true",
                    help="list files that would be merged without writing output")
    args = ap.parse_args()

    if not os.path.isdir(args.input_dir):
        print(f"ERROR: input dir not found: {args.input_dir}", file=sys.stderr)
        return 1

    return merge(args.input_dir, args.output, dry_run=args.dry_run)


if __name__ == "__main__":
    sys.exit(main())
