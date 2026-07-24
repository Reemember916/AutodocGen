#!/usr/bin/env python3
"""Plan or safely apply AutoDocGen document updates from code diffs.

First version is intentionally conservative:
- modified functions with a unique CSU match can be applied automatically;
- reviewed replacement decisions can be applied explicitly;
- new/deleted/header/ambiguous items are otherwise reported only.
"""

from __future__ import annotations

import argparse
import difflib
import inspect
import json
import os
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Sequence


CSU_RE = re.compile(r"(.+?)[（(](D/R_[A-Za-z0-9_]+)[）)]")
CSU_ID_SUFFIX_RE = re.compile(r"_\d{3}_\d{3}$")
FUNC_RE = re.compile(r"(?:interrupt\s+)?[A-Za-z_][\w\s\*\(\)]*?\s+([A-Za-z_]\w*)\s*\(")
INCLUDE_RE = re.compile(r'^\s*#\s*include\s*[<"]([^">]+)[">]', re.M)
C_FUNC_DEF_RE = re.compile(
    r"(?m)^\s*(?:static\s+|extern\s+|inline\s+|interrupt\s+|const\s+|volatile\s+)*"
    r"[A-Za-z_][\w\s\*\(\)]*?\s+([A-Za-z_]\w*)\s*\([^;{}]*\)\s*\{"
)
DOC_CODE_ALIGNMENT_SCHEMA = 1


def _open_docx_safe(path: str):
    """Open a docx file with validation and wrapped Document() call.

    Returns the Document instance on success.
    Raises ValueError if the file is corrupted or unreadable.
    """
    _validate_docx(path)
    from docx import Document
    try:
        return Document(path)
    except Exception as e:
        raise ValueError(f"无法打开 docx 文件（可能已损坏）：{path}\n  错误：{e}")


def _validate_docx(path: str) -> None:
    """Check that *path* is a valid ZIP archive before opening it with python-docx.

    python-docx wraps lxml which is a C extension — a corrupted .docx can
    trigger a segfault inside lxml that kills the entire process immediately.
    This early check catches the most common corruption so the caller can
    fail gracefully instead of crashing.
    """
    import zipfile
    if not os.path.isfile(path):
        raise FileNotFoundError(f"docx 文件不存在：{path}")
    if not path.lower().endswith(".docx"):
        raise ValueError(f"文件不是 .docx 格式：{path}")
    try:
        with zipfile.ZipFile(path) as zf:
            bad = zf.testzip()
            if bad is not None:
                raise ValueError(f"docx 文件已损坏（ZIP 校验失败）：{bad}")
    except zipfile.BadZipFile:
        raise ValueError(f"docx 文件不是有效的 ZIP 格式（可能已损坏）：{path}")


@dataclass
class PlannedItem:
    action: str
    status: str
    rel_path: str
    func_name: str = ""
    csu_id: str = ""
    reason: str = ""
    change: dict[str, Any] = field(default_factory=dict)
    result: dict[str, Any] = field(default_factory=dict)
    alignment: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "action": self.action,
            "status": self.status,
            "rel_path": self.rel_path,
            "func_name": self.func_name,
            "csu_id": self.csu_id,
            "reason": self.reason,
            "change": self.change,
            "result": self.result,
            "alignment": self.alignment,
        }


def _abs(path: str) -> str:
    return os.path.abspath(os.path.expanduser(path or ""))


def _default_docdiff_root() -> str:
    # 环境变量优先，便于跨机器部署
    env_root = os.environ.get("AUTODOCGEN_DOCDIFF_ROOT", "").strip()
    if env_root and os.path.isdir(env_root) and os.path.isfile(os.path.join(env_root, "cli.py")):
        return env_root
    candidates: list[str] = []
    frozen_root = getattr(sys, "_MEIPASS", "")
    if frozen_root:
        candidates.append(os.path.join(frozen_root, "DocDiff-main"))
    # tools/ 的上一级是项目根 AutoDocGen/（不是 parents[2] 的上上级）
    here = Path(__file__).resolve()
    candidates.append(str(here.parents[1] / "DocDiff-main"))
    # 兼容旧布局：仓库旁的 DocDiff-main
    if len(here.parents) > 2:
        candidates.append(str(here.parents[2] / "DocDiff-main"))
    # cwd 相对路径兜底
    candidates.append(os.path.abspath(os.path.join(os.getcwd(), "DocDiff-main")))
    for candidate in candidates:
        if candidate and os.path.isdir(candidate) and os.path.isfile(os.path.join(candidate, "cli.py")):
            return candidate
    return ""


def _default_sidecar(out_path: str, suffix: str) -> str:
    base, _ = os.path.splitext(_abs(out_path))
    return base + suffix


def _run_docdiff(
    *,
    docdiff_root: str,
    old_code: str,
    new_code: str,
    change_docx: str,
    change_json: str,
) -> None:
    cli = os.path.join(docdiff_root, "cli.py")
    if not os.path.isfile(cli):
        raise FileNotFoundError(f"DocDiff cli.py not found: {cli}")
    cmd = [
        sys.executable,
        cli,
        "--mode",
        "code",
        "--old",
        old_code,
        "--new",
        new_code,
        "--out",
        change_docx,
        "--json-out",
        change_json,
    ]
    subprocess.run(cmd, cwd=docdiff_root, check=True)


def _load_changes(path: str) -> list[dict[str, Any]]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, list):
        raise ValueError("code diff JSON must be a list")
    return [x for x in data if isinstance(x, dict)]


def _load_review_decisions(path: str) -> list[dict[str, Any]]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, dict):
        decisions = data.get("decisions", [])
    else:
        decisions = data
    if not isinstance(decisions, list):
        raise ValueError("review decisions JSON must contain a decisions list")
    return [
        x for x in decisions
        if isinstance(x, dict) and str(x.get("decision") or "").strip()
    ]


def _load_alignment_decisions(path: str) -> list[dict[str, Any]]:
    if not path:
        return []
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, dict):
        decisions = data.get("alignment_decisions", [])
    else:
        decisions = []
    if not isinstance(decisions, list):
        raise ValueError("alignment decisions JSON must contain an alignment_decisions list")
    return [
        x for x in decisions
        if isinstance(x, dict)
        and (
            str(x.get("manual_function") or "").strip()
            or str(x.get("manual_rel_path") or "").strip()
            or str(x.get("notes") or "").strip()
        )
    ]


def _paragraph_style_name(paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def _paragraph_style_names(paragraph) -> list[str]:
    names: list[str] = []
    try:
        style = paragraph.style
    except Exception:
        return names
    seen: set[int] = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        try:
            name = (style.name or "").strip()
        except Exception:
            name = ""
        if name:
            names.append(name)
        try:
            style = style.base_style
        except Exception:
            break
    return names


def _style_name_matches_heading_level(name: str, level: int) -> bool:
    normalized = (name or "").strip()
    if not normalized:
        return False
    simple = re.sub(r"[\s_]+", "", normalized)
    level_str = str(level)
    return (
        f"Heading {level}" in normalized
        or f"Heading{level_str}" in simple
        or f"标题 {level}" in normalized
        or f"标题{level_str}" in simple
        or normalized == f"609_{level}"
    )


def _looks_like_csu_heading(paragraph) -> bool:
    text = (paragraph.text or "").strip()
    match = CSU_RE.search(text)
    if not match:
        return False
    if not CSU_ID_SUFFIX_RE.search(match.group(2) or ""):
        return False
    names = _paragraph_style_names(paragraph)
    if any(_style_name_matches_heading_level(name, 4) for name in names):
        return True
    name = (names[0] if names else _paragraph_style_name(paragraph)).lower()
    simple = re.sub(r"[\s_]+", "", name)
    return "4" in simple and ("heading" in simple or "标题" in simple)


def _extract_function_after_heading(paragraphs: list[Any], index: int) -> tuple[str, str]:
    limit = min(index + 12, len(paragraphs))
    for pos in range(index + 1, limit):
        text = (paragraphs[pos].text or "").strip()
        if not text:
            continue
        if "函数原型" in text:
            for proto_pos in range(pos + 1, min(pos + 5, len(paragraphs))):
                proto = (paragraphs[proto_pos].text or "").strip()
                if not proto:
                    continue
                match = FUNC_RE.search(proto)
                if match:
                    return match.group(1), proto
            return "", ""
        match = FUNC_RE.search(text)
        if match and (";" in text or ")" in text):
            return match.group(1), text
    return "", ""


def _doc_csu_entries(doc_path: str) -> list[dict[str, str]]:
    doc = _open_docx_safe(doc_path)
    paragraphs = list(doc.paragraphs)
    entries: list[dict[str, str]] = []
    for pos, paragraph in enumerate(paragraphs):
        text = (paragraph.text or "").strip()
        if not _looks_like_csu_heading(paragraph):
            continue
        match = CSU_RE.search(text)
        if not match:
            continue
        title = match.group(1).strip()
        csu_id = match.group(2).strip()
        func_name, prototype = _extract_function_after_heading(paragraphs, pos)
        if not func_name:
            # Fallback: some generated docs keep the C name in the title.
            title_match = re.search(r"\b([A-Za-z_]\w*)\b", title)
            func_name = title_match.group(1) if title_match else ""
        entries.append({
            "title": title,
            "csu_id": csu_id,
            "heading": text,
            "func_name": func_name,
            "prototype": prototype,
        })
    return entries


def build_csu_index(doc_path: str) -> dict[str, list[dict[str, str]]]:
    index: dict[str, list[dict[str, str]]] = {}
    for entry in _doc_csu_entries(doc_path):
        func_name = entry.get("func_name") or ""
        if func_name:
            index.setdefault(func_name, []).append({
                "title": entry.get("title", ""),
                "csu_id": entry.get("csu_id", ""),
                "heading": entry.get("heading", ""),
                "prototype": entry.get("prototype", ""),
            })
    return index


def _normalize_signature_text(text: str) -> str:
    text = re.sub(r"/\*.*?\*/", "", text or "", flags=re.S)
    text = re.sub(r"//.*", "", text)
    text = text.replace("{", "").replace(";", "")
    return re.sub(r"\s+", "", text)


def _scan_code_functions(code_root: str) -> dict[str, list[dict[str, str]]]:
    root = Path(code_root)
    index: dict[str, list[dict[str, str]]] = {}
    if not root.is_dir():
        return index
    for path in sorted(root.rglob("*.c")):
        rel_path = _normalized_rel(os.path.relpath(path, root))
        text = _read_text_if_possible(str(path))
        for match in C_FUNC_DEF_RE.finditer(text):
            name = match.group(1)
            if name in {"if", "for", "while", "switch", "return", "sizeof"}:
                continue
            signature = match.group(0).rstrip("{").strip()
            line_no = text[:match.start()].count("\n") + 1
            index.setdefault(name, []).append({
                "func_name": name,
                "rel_path": rel_path,
                "signature": signature,
                "line": str(line_no),
            })
    return index


def _alignment_for_doc_entry(entry: dict[str, str], code_index: dict[str, list[dict[str, str]]]) -> dict[str, Any]:
    func_name = (entry.get("func_name") or "").strip()
    csu_id = (entry.get("csu_id") or "").strip()
    base = {
        "schema_version": DOC_CODE_ALIGNMENT_SCHEMA,
        "csu_id": csu_id,
        "doc_title": entry.get("title", ""),
        "doc_heading": entry.get("heading", ""),
        "doc_func_name": func_name,
        "doc_prototype": entry.get("prototype", ""),
        "matched_function": "",
        "rel_path": "",
        "signature": "",
        "confidence": 0.0,
        "status": "unmatched",
        "evidence": {},
    }
    if not func_name:
        base["status"] = "no_doc_function"
        base["evidence"] = {"doc_function_name": "missing"}
        return base

    candidates = code_index.get(func_name, [])
    if len(candidates) != 1:
        base["status"] = "ambiguous" if candidates else "unmatched"
        base["evidence"] = {
            "function_name": "exact" if candidates else "missing_in_code",
            "candidate_count": len(candidates),
            "candidates": [
                {"rel_path": candidate.get("rel_path", ""), "line": candidate.get("line", "")}
                for candidate in candidates[:20]
            ],
        }
        return base

    candidate = candidates[0]
    doc_sig = _normalize_signature_text(entry.get("prototype", ""))
    code_sig = _normalize_signature_text(candidate.get("signature", ""))
    signature_match = bool(doc_sig and code_sig and doc_sig == code_sig)
    confidence = 0.98 if signature_match else 0.92
    base.update({
        "matched_function": func_name,
        "rel_path": candidate.get("rel_path", ""),
        "signature": candidate.get("signature", ""),
        "confidence": confidence,
        "status": "matched_high",
        "evidence": {
            "function_name": "exact_unique",
            "candidate_count": 1,
            "prototype": "exact" if signature_match else ("missing" if not doc_sig else "name_only"),
        },
    })
    return base


def build_doc_code_alignment_index(doc_path: str, code_root: str) -> dict[str, dict[str, Any]]:
    code_index = _scan_code_functions(code_root)
    alignment: dict[str, dict[str, Any]] = {}
    for entry in _doc_csu_entries(doc_path):
        csu_id = (entry.get("csu_id") or "").strip()
        if not csu_id:
            continue
        alignment[csu_id] = _alignment_for_doc_entry(entry, code_index)
    return alignment


def _find_code_function(code_root: str, rel_path: str, func_name: str) -> dict[str, str]:
    rel_path = _normalized_rel(rel_path)
    func_name = (func_name or "").strip()
    if not rel_path or not func_name:
        return {}
    source = os.path.join(code_root, rel_path)
    text = _read_text_if_possible(source)
    for match in C_FUNC_DEF_RE.finditer(text):
        if match.group(1) != func_name:
            continue
        return {
            "func_name": func_name,
            "rel_path": rel_path,
            "signature": match.group(0).rstrip("{").strip(),
            "line": str(text[:match.start()].count("\n") + 1),
        }
    return {}


def _append_csu_index_entry(
    csu_index: dict[str, list[dict[str, str]]],
    func_name: str,
    entry: dict[str, str],
) -> None:
    func_name = (func_name or "").strip()
    csu_id = (entry.get("csu_id") or "").strip()
    if not func_name or not csu_id:
        return
    bucket = csu_index.setdefault(func_name, [])
    if not any((x.get("csu_id") or "") == csu_id for x in bucket):
        bucket.append(entry)


def apply_alignment_decisions(
    alignment_index: dict[str, dict[str, Any]],
    csu_index: dict[str, list[dict[str, str]]],
    alignment_decisions: list[dict[str, Any]],
    *,
    code_roots: Sequence[str],
) -> None:
    for decision in alignment_decisions:
        csu_id = str(decision.get("csu_id") or "").strip()
        manual_function = str(decision.get("manual_function") or "").strip()
        manual_rel_path = _normalized_rel(str(decision.get("manual_rel_path") or ""))
        notes = str(decision.get("notes") or "").strip()
        if not csu_id or not manual_function or not manual_rel_path:
            continue

        current = dict(alignment_index.get(csu_id) or {})
        if not current:
            current = {
                "schema_version": DOC_CODE_ALIGNMENT_SCHEMA,
                "csu_id": csu_id,
                "doc_title": str(decision.get("doc_title") or ""),
                "doc_heading": "",
                "doc_func_name": str(decision.get("doc_func_name") or ""),
                "doc_prototype": "",
            }
        previous_status = str(current.get("status") or "")
        code_func: dict[str, str] = {}
        for code_root in code_roots:
            code_func = _find_code_function(code_root, manual_rel_path, manual_function)
            if code_func:
                break
        current.update({
            "schema_version": DOC_CODE_ALIGNMENT_SCHEMA,
            "matched_function": manual_function,
            "rel_path": manual_rel_path,
            "signature": code_func.get("signature", ""),
            "confidence": 1.0,
            "status": "manual_matched",
            "evidence": {
                "source": "alignment_decisions",
                "manual_function": manual_function,
                "manual_rel_path": manual_rel_path,
                "previous_status": previous_status,
                "notes": notes,
            },
        })
        alignment_index[csu_id] = current

        _append_csu_index_entry(
            csu_index,
            manual_function,
            {
                "title": str(current.get("doc_title") or ""),
                "csu_id": csu_id,
                "heading": str(current.get("doc_heading") or ""),
                "prototype": str(current.get("doc_prototype") or ""),
                "alignment_status": "manual_matched",
            },
        )


def attach_alignment_to_items(items: list[PlannedItem], alignment_index: dict[str, dict[str, Any]]) -> None:
    for item in items:
        csu_id = (item.csu_id or "").strip()
        if csu_id and csu_id in alignment_index:
            item.alignment = alignment_index[csu_id]


def collect_csu_ids(doc_path: str) -> list[str]:
    doc = _open_docx_safe(doc_path)
    ids: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        match = CSU_RE.search(text)
        if match:
            ids.append(match.group(2).strip())
    return ids


def allocate_next_csu_id(doc_path: str, after_csu_id: str) -> str:
    """Allocate the next CSU ID in the same module as ``after_csu_id``."""
    after_csu_id = (after_csu_id or "").strip()
    match = re.match(r"^(.+?)_(\d+)$", after_csu_id)
    if not match:
        return ""
    prefix, suffix = match.group(1), match.group(2)
    width = len(suffix)
    max_index = 0
    pattern = re.compile(rf"^{re.escape(prefix)}_(\d+)$")
    for csu_id in collect_csu_ids(doc_path):
        id_match = pattern.match(csu_id)
        if not id_match:
            continue
        try:
            max_index = max(max_index, int(id_match.group(1)))
        except ValueError:
            continue
    if max_index <= 0:
        try:
            max_index = int(suffix)
        except ValueError:
            return ""
    return f"{prefix}_{max_index + 1:0{width}d}"


def _change_func_name(change: dict[str, Any], *, prefer: str = "any") -> str:
    keys = ["function_name", "new_function_name", "old_function_name"]
    if prefer == "new":
        keys = ["new_function_name", "function_name", "old_function_name"]
    elif prefer == "old":
        keys = ["old_function_name", "function_name", "new_function_name"]
    for key in keys:
        value = str(change.get(key) or "").strip()
        if value:
            return value
    return ""


def _normalized_function_text(change: dict[str, Any], *, old_name: str, new_name: str, side: str) -> str:
    text = str(change.get("old_text" if side == "old" else "new_text") or change.get("text") or "")
    text = re.sub(r"/\*.*?\*/", "", text, flags=re.S)
    text = re.sub(r"//.*", "", text)
    for name in {old_name, new_name}:
        if name:
            text = re.sub(rf"\b{re.escape(name)}\b", "__FUNC__", text)
    return re.sub(r"\s+", "", text)


def _rename_similarity(deleted_change: dict[str, Any], new_change: dict[str, Any]) -> float:
    old_name = _change_func_name(deleted_change, prefer="old")
    new_name = _change_func_name(new_change, prefer="new")
    old_text = _normalized_function_text(deleted_change, old_name=old_name, new_name=new_name, side="old")
    new_text = _normalized_function_text(new_change, old_name=old_name, new_name=new_name, side="new")
    if len(old_text) < 20 or len(new_text) < 20:
        return 0.0
    return difflib.SequenceMatcher(None, old_text, new_text).ratio()


def detect_renamed_functions(changes: list[dict[str, Any]], *, threshold: float = 0.78) -> tuple[list[dict[str, Any]], set[int]]:
    deleted = [
        (idx, change) for idx, change in enumerate(changes)
        if str(change.get("change_kind") or "") == "deleted_function"
        and str(change.get("language") or "") == "c"
    ]
    added = [
        (idx, change) for idx, change in enumerate(changes)
        if str(change.get("change_kind") or "") == "new_function"
        and str(change.get("language") or "") == "c"
    ]
    candidates: list[tuple[float, int, int, dict[str, Any], dict[str, Any]]] = []
    for deleted_idx, deleted_change in deleted:
        for added_idx, added_change in added:
            if str(deleted_change.get("key") or "") != str(added_change.get("key") or ""):
                continue
            score = _rename_similarity(deleted_change, added_change)
            if score >= threshold:
                candidates.append((score, deleted_idx, added_idx, deleted_change, added_change))

    renamed: list[dict[str, Any]] = []
    used: set[int] = set()
    for score, deleted_idx, added_idx, deleted_change, added_change in sorted(candidates, reverse=True):
        if deleted_idx in used or added_idx in used:
            continue
        old_name = _change_func_name(deleted_change, prefer="old")
        new_name = _change_func_name(added_change, prefer="new")
        combined = dict(added_change)
        combined.update({
            "change_kind": "renamed_function",
            "function_name": new_name,
            "old_function_name": old_name,
            "new_function_name": new_name,
            "old_signature": deleted_change.get("old_signature") or deleted_change.get("seg") or "",
            "new_signature": added_change.get("new_signature") or added_change.get("seg") or "",
            "old_text": deleted_change.get("old_text") or "",
            "new_text": added_change.get("new_text") or "",
            "rename_similarity": round(score, 4),
        })
        renamed.append(combined)
        used.update({deleted_idx, added_idx})
    return renamed, used


def _read_text_if_possible(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def _normalized_rel(path: str) -> str:
    return str(path or "").replace("\\", "/").strip().lstrip("./")


def _include_matches_header(include_name: str, header_rel_path: str) -> bool:
    include_name = _normalized_rel(include_name)
    header_rel_path = _normalized_rel(header_rel_path)
    header_base = os.path.basename(header_rel_path)
    return (
        include_name == header_rel_path
        or include_name.endswith("/" + header_rel_path)
        or include_name == header_base
        or include_name.endswith("/" + header_base)
    )


def _resolve_include_to_project_header(include_name: str, header_paths: list[str]) -> str:
    include_name = _normalized_rel(include_name)
    matches = [path for path in header_paths if _include_matches_header(include_name, path)]
    if not matches:
        return ""
    matches.sort(key=lambda value: (len(value), value))
    return matches[0]


def _scan_project_includes(root: Path) -> tuple[dict[str, list[str]], dict[str, str]]:
    files: dict[str, str] = {}
    for pattern in ("*.h", "*.hpp", "*.c"):
        for path in root.rglob(pattern):
            rel_path = _normalized_rel(os.path.relpath(path, root))
            files[rel_path] = _read_text_if_possible(str(path))

    header_paths = sorted([path for path in files if path.lower().endswith((".h", ".hpp"))])
    includes_by_file: dict[str, list[str]] = {}
    for rel_path, text in files.items():
        resolved: list[str] = []
        for match in INCLUDE_RE.finditer(text):
            header_rel = _resolve_include_to_project_header(match.group(1), header_paths)
            if header_rel and header_rel not in resolved:
                resolved.append(header_rel)
        includes_by_file[rel_path] = resolved
    return includes_by_file, files


def _impacted_headers_for(header_rel: str, includes_by_file: dict[str, list[str]]) -> set[str]:
    impacted = {header_rel}
    changed = True
    while changed:
        changed = False
        for rel_path, includes in includes_by_file.items():
            if not rel_path.lower().endswith((".h", ".hpp")):
                continue
            if rel_path in impacted:
                continue
            if any(include in impacted for include in includes):
                impacted.add(rel_path)
                changed = True
    return impacted


_HEADER_SYMBOL_RE = re.compile(r"\b[A-Za-z_]\w*\b")
_HEADER_SYMBOL_IGNORED = {
    "const", "extern", "static", "struct", "typedef", "union", "enum",
    "unsigned", "signed", "volatile", "void", "char", "short", "int",
    "long", "float", "double", "return", "include", "define", "ifdef",
    "ifndef", "endif", "pragma", "true", "false", "null",
    # Common project/toolchain type aliases are frequently present only in
    # truncated diff context and must never fan out header impact records.
    "uint8", "uint16", "uint32", "uint64",
    "int8", "int16", "int32", "int64",
    "uint8_t", "uint16_t", "uint32_t", "uint64_t",
    "int8_t", "int16_t", "int32_t", "int64_t",
}


def _changed_header_symbols(change: dict[str, Any]) -> set[str]:
    """Extract identifiers added or removed by a header diff.

    Header inclusion alone is too broad in embedded projects where nearly every
    C file includes a shared ``Global.h``.  These symbols provide a conservative
    second signal: a function is listed only when its body actually references
    an identifier changed by the header diff.
    """
    old_symbols = set(_HEADER_SYMBOL_RE.findall(str(change.get("old_text") or "")))
    new_symbols = set(_HEADER_SYMBOL_RE.findall(str(change.get("new_text") or "")))
    candidates = old_symbols.symmetric_difference(new_symbols)
    return {
        symbol
        for symbol in candidates
        if len(symbol) >= 3
        and symbol.lower() not in _HEADER_SYMBOL_IGNORED
        and not re.fullmatch(r"(?:u?int|uint|float|double|bool)\d*(?:_t)?", symbol, re.IGNORECASE)
        and not symbol.isdigit()
    }


def _extract_c_function_bodies(source_text: str) -> list[tuple[str, str]]:
    """Return (function_name, body) pairs using balanced-brace scanning."""
    text = source_text or ""
    out: list[tuple[str, str]] = []
    for match in C_FUNC_DEF_RE.finditer(text):
        name = str(match.group(1) or "").strip()
        brace = text.find("{", match.start(), match.end() + 1)
        if not name or brace < 0:
            continue
        depth = 0
        end = brace
        for pos in range(brace, len(text)):
            char = text[pos]
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    end = pos + 1
                    break
        if depth == 0:
            out.append((name, text[match.start():end]))
    return out


def _extract_c_function_names(source_text: str) -> list[str]:
    names: list[str] = []
    for match in C_FUNC_DEF_RE.finditer(source_text or ""):
        name = match.group(1)
        if name in {"if", "for", "while", "switch", "return", "sizeof"}:
            continue
        if name not in names:
            names.append(name)
    return names


def _file_level_c_function_names(change: dict[str, Any], *, new_code: str) -> list[str]:
    kind = str(change.get("change_kind") or "")
    rel_path = _normalized_rel(str(change.get("key") or ""))
    text = ""
    if kind == "new_file":
        text = str(change.get("new_text") or "")
        if (not text) and rel_path:
            text = _read_text_if_possible(os.path.join(new_code, rel_path))
    elif kind == "deleted_file":
        text = str(change.get("old_text") or "")
    return _extract_c_function_names(text)


def _planned_deleted_function_item(
    change: dict[str, Any],
    *,
    rel_path: str,
    func_name: str,
    csu_index: dict[str, list[dict[str, str]]],
) -> PlannedItem:
    matches = csu_index.get(func_name, [])
    csu_id = matches[0]["csu_id"] if len(matches) == 1 else ""
    reason = "deleted function removal needs confirmation"
    if len(matches) > 1:
        reason = "deleted function has multiple CSU matches"
    elif not matches:
        reason = "deleted function has no CSU match"
    return PlannedItem("deleted_function", "review", rel_path, func_name, csu_id, reason, change)


def find_header_impacted_items(
    header_change: dict[str, Any],
    *,
    new_code: str,
    csu_index: dict[str, list[dict[str, str]]],
    skip_functions: set[tuple[str, str]],
) -> list[PlannedItem]:
    header_rel = _normalized_rel(str(header_change.get("key") or ""))
    if not header_rel:
        return []

    impacted: list[PlannedItem] = []
    root = Path(new_code)
    if not root.is_dir():
        return impacted

    includes_by_file, file_texts = _scan_project_includes(root)
    impacted_headers = _impacted_headers_for(header_rel, includes_by_file)
    changed_symbols = _changed_header_symbols(header_change)
    if not changed_symbols:
        return impacted
    symbol_pattern = re.compile(
        r"\b(?:" + "|".join(re.escape(symbol) for symbol in sorted(changed_symbols)) + r")\b"
    )

    for rel_path in sorted(path for path in file_texts if path.lower().endswith(".c")):
        text = file_texts.get(rel_path, "")
        includes = includes_by_file.get(rel_path) or []
        matched_headers = sorted(set(includes).intersection(impacted_headers))
        if not matched_headers:
            continue
        for func_name, func_body in _extract_c_function_bodies(text):
            if (rel_path, func_name) in skip_functions:
                continue
            referenced_symbols = sorted(set(symbol_pattern.findall(func_body)))
            if not referenced_symbols:
                continue
            matches = csu_index.get(func_name, [])
            csu_id = matches[0]["csu_id"] if len(matches) == 1 else ""
            via_text = matched_headers[0]
            symbol_preview = ", ".join(referenced_symbols[:5])
            reason = f"function references changed header symbols ({symbol_preview}): {header_rel}"
            if via_text == header_rel:
                reason = f"function directly includes and references changed symbols ({symbol_preview}): {header_rel}"
            status = "review"
            if len(matches) > 1:
                status = "manual"
                reason = f"header impact has multiple CSU matches: {header_rel}"
            elif not matches:
                status = "manual"
                reason = f"header impact has no CSU match: {header_rel}"
            impact_change = dict(header_change)
            impact_change.update({
                "change_kind": "header_impacted_function",
                "function_name": func_name,
                "source_file": rel_path,
                "header_file": header_rel,
                "matched_header_file": via_text,
                "impacted_headers": sorted(impacted_headers),
                "changed_symbols": sorted(changed_symbols),
                "referenced_symbols": referenced_symbols,
            })
            impacted.append(PlannedItem("header_impacted_function", status, rel_path, func_name, csu_id, reason, impact_change))
    return impacted


def dedupe_header_impacted_items(items: list[PlannedItem]) -> list[PlannedItem]:
    merged: list[PlannedItem] = []
    header_items: dict[tuple[str, str, str], PlannedItem] = {}
    for item in items:
        if item.action != "header_impacted_function":
            merged.append(item)
            continue

        key = (item.rel_path, item.func_name, item.csu_id)
        existing = header_items.get(key)
        if existing is None:
            change = dict(item.change or {})
            header_file = str(change.get("header_file") or "").strip()
            if header_file:
                change["impacted_by_headers"] = [header_file]
            item.change = change
            header_items[key] = item
            merged.append(item)
            continue

        existing_change = existing.change or {}
        incoming_change = item.change or {}
        headers = list(existing_change.get("impacted_by_headers") or [])
        header_file = str(incoming_change.get("header_file") or "").strip()
        if header_file and header_file not in headers:
            headers.append(header_file)
        existing_change["impacted_by_headers"] = headers
        combined_impacted = set(existing_change.get("impacted_headers") or [])
        combined_impacted.update(incoming_change.get("impacted_headers") or [])
        existing_change["impacted_headers"] = sorted(combined_impacted)
        existing.change = existing_change

        if item.status == "manual":
            existing.status = "manual"
        count = len(headers) or 1
        if existing.status == "manual":
            if not existing.csu_id:
                existing.reason = f"header impact has no CSU match: {count} changed headers"
            else:
                existing.reason = f"header impact has multiple CSU matches: {count} changed headers"
        else:
            existing.reason = f"function impacted by {count} changed headers"

    return merged


def classify_changes(
    changes: list[dict[str, Any]],
    *,
    new_code: str,
    csu_index: dict[str, list[dict[str, str]]],
) -> list[PlannedItem]:
    items: list[PlannedItem] = []
    renamed_changes, renamed_used = detect_renamed_functions(changes)
    directly_changed_functions: set[tuple[str, str]] = set()
    for change in changes:
        kind = str(change.get("change_kind") or "")
        if kind in {"modified_function", "new_function", "deleted_function"}:
            name = _change_func_name(change)
            rel_path = _normalized_rel(str(change.get("key") or ""))
            if rel_path and name:
                directly_changed_functions.add((rel_path, name))
        elif kind in {"new_file", "deleted_file"}:
            rel_path = _normalized_rel(str(change.get("key") or ""))
            if rel_path.lower().endswith(".c"):
                for name in _file_level_c_function_names(change, new_code=new_code):
                    directly_changed_functions.add((rel_path, name))
    for change in renamed_changes:
        rel_path = str(change.get("key") or "")
        old_func_name = _change_func_name(change, prefer="old")
        new_func_name = _change_func_name(change, prefer="new")
        if rel_path and new_func_name:
            directly_changed_functions.add((_normalized_rel(rel_path), new_func_name))
        if rel_path and old_func_name:
            directly_changed_functions.add((_normalized_rel(rel_path), old_func_name))
        matches = csu_index.get(old_func_name, [])
        csu_id = matches[0]["csu_id"] if len(matches) == 1 else ""
        reason = "possible function rename; review replace_csu target"
        if len(matches) > 1:
            reason = "possible function rename has multiple old CSU matches"
        elif not matches:
            reason = "possible function rename has no old CSU match"
        items.append(PlannedItem("renamed_function", "review", rel_path, new_func_name, csu_id, reason, change))

    for index, change in enumerate(changes):
        if index in renamed_used:
            continue
        kind = str(change.get("change_kind") or "")
        rel_path = str(change.get("key") or "")
        func_name = str(
            change.get("function_name")
            or change.get("new_function_name")
            or change.get("old_function_name")
            or ""
        ).strip()
        language = str(change.get("language") or "")

        if kind == "modified_function" and language == "c":
            new_source = os.path.join(new_code, rel_path)
            matches = csu_index.get(func_name, [])
            if not os.path.isfile(new_source):
                items.append(PlannedItem("modified_function", "manual", rel_path, func_name, reason="new source file missing", change=change))
            elif len(matches) == 1:
                items.append(PlannedItem("modified_function", "safe", rel_path, func_name, matches[0]["csu_id"], "unique CSU match", change))
            elif len(matches) > 1:
                items.append(PlannedItem("modified_function", "manual", rel_path, func_name, reason="multiple CSU matches", change=change))
            else:
                items.append(PlannedItem("modified_function", "manual", rel_path, func_name, reason="no CSU match", change=change))
            continue

        if kind == "new_function":
            matches = csu_index.get(func_name, [])
            if len(matches) == 1:
                reason = "new function has manual CSU alignment; review replace_csu target"
                items.append(PlannedItem("new_function", "review", rel_path, func_name, matches[0]["csu_id"], reason, change))
            elif len(matches) > 1:
                items.append(PlannedItem("new_function", "manual", rel_path, func_name, reason="new function has multiple CSU matches", change=change))
            else:
                items.append(PlannedItem("new_function", "review", rel_path, func_name, reason="new function insertion needs section position and CSU ID", change=change))
            continue

        if kind == "deleted_function":
            items.append(_planned_deleted_function_item(
                change,
                rel_path=rel_path,
                func_name=func_name,
                csu_index=csu_index,
            ))
            continue

        if kind in {"new_file", "deleted_file"} and _normalized_rel(rel_path).lower().endswith(".c"):
            file_funcs = _file_level_c_function_names(change, new_code=new_code)
            if not file_funcs:
                items.append(PlannedItem(kind, "manual", rel_path, reason=f"{kind} has no parseable C functions", change=change))
                continue
            for file_func_name in file_funcs:
                expanded_change = dict(change)
                if kind == "new_file":
                    expanded_change.update({
                        "change_kind": "new_function",
                        "function_name": file_func_name,
                        "new_function_name": file_func_name,
                    })
                    matches = csu_index.get(file_func_name, [])
                    if len(matches) == 1:
                        reason = "new function has manual CSU alignment; review replace_csu target"
                        items.append(PlannedItem("new_function", "review", rel_path, file_func_name, matches[0]["csu_id"], reason, change=expanded_change))
                    elif len(matches) > 1:
                        items.append(PlannedItem("new_function", "manual", rel_path, file_func_name, reason="new function has multiple CSU matches", change=expanded_change))
                    else:
                        items.append(PlannedItem("new_function", "review", rel_path, file_func_name, reason="new function insertion needs section position and CSU ID", change=expanded_change))
                else:
                    expanded_change.update({
                        "change_kind": "deleted_function",
                        "function_name": file_func_name,
                        "old_function_name": file_func_name,
                    })
                    items.append(_planned_deleted_function_item(
                        expanded_change,
                        rel_path=rel_path,
                        func_name=file_func_name,
                        csu_index=csu_index,
                    ))
            continue

        if kind == "header_changed":
            items.append(PlannedItem("header_changed", "manual", rel_path, reason="header changes may affect many functions", change=change))
            items.extend(find_header_impacted_items(
                change,
                new_code=new_code,
                csu_index=csu_index,
                skip_functions=directly_changed_functions,
            ))
            continue

        items.append(PlannedItem(kind or "unknown", "manual", rel_path, func_name, reason="unsupported change kind", change=change))
    return dedupe_header_impacted_items(items)


def _make_cfg(
    *,
    ai_assist: bool,
    template_path: str,
    verbose: bool = False,
    stop_event: Any = None,
    settings: Any = None,
    gui_log: Any = None,
):
    root = Path(__file__).resolve().parents[1]
    if str(root) not in sys.path:
        sys.path.insert(0, str(root))
    from autodoc.backend import GenConfig

    cfg = GenConfig(
        verbose=bool(verbose),
        ai_assist=bool(ai_assist),
        template_path=template_path or "",
        stop_event=stop_event,
    )
    if settings is not None:
        cfg.ai_mode = int(getattr(settings, "ai_mode", 1) or 1)
        cfg.ai_provider = str(getattr(settings, "ai_provider", "local") or "local")
        cfg.ai_api_base = str(getattr(settings, "ai_api_base", "") or "")
        cfg.ai_api_key = str(getattr(settings, "ai_api_key", "") or "")
        cfg.ai_model = str(getattr(settings, "ai_model", "") or "")
        cfg.ai_read_timeout = float(getattr(settings, "ai_read_timeout", 40) or 40)
        cfg.ai_workers = max(1, int(getattr(settings, "ai_workers", 1) or 1))
        cfg.ai_max_tokens = int(getattr(settings, "ai_max_tokens", 16384) or 16384)
        cfg.ai_num_ctx = int(getattr(settings, "ai_num_ctx", 0) or 0)
        cfg.proxy = str(getattr(settings, "proxy", "") or "") if bool(getattr(settings, "use_proxy", False)) else ""
        cfg.no_proxy = bool(getattr(settings, "no_proxy", False))
    if gui_log is not None:
        cfg.gui_log = gui_log
    return cfg


def _copy_old_doc_to_out(*, old_doc: str, out_doc: str, label: str) -> str:
    old_abs = _abs(old_doc)
    out_abs = _abs(out_doc)
    if old_abs == out_abs:
        raise ValueError(f"{label} requires --out to differ from --old-doc")
    os.makedirs(os.path.dirname(out_abs) or ".", exist_ok=True)
    shutil.copy2(old_abs, out_abs)
    return out_abs


def _supports_kwarg(func: Any, name: str) -> bool:
    try:
        sig = inspect.signature(func)
    except Exception:
        return False
    return name in sig.parameters or any(
        param.kind == inspect.Parameter.VAR_KEYWORD
        for param in sig.parameters.values()
    )


def apply_safe_items(
    items: list[PlannedItem],
    *,
    old_doc: str,
    out_doc: str,
    new_code: str,
    ai_assist: bool,
    template_path: str,
    verbose: bool = False,
    copy_doc: bool = True,
    stop_event: Any = None,
    settings: Any = None,
    gui_log: Any = None,
) -> None:
    out_abs = _copy_old_doc_to_out(old_doc=old_doc, out_doc=out_doc, label="apply-safe") if copy_doc else _abs(out_doc)

    root = Path(__file__).resolve().parents[1]
    if str(root) not in sys.path:
        sys.path.insert(0, str(root))
    from autodoc.pipeline import regenerate_csu_in_doc
    from autodoc._legacy_support import legacy_backend

    cfg = _make_cfg(ai_assist=ai_assist, template_path=template_path, verbose=verbose,
                    stop_event=stop_event, settings=settings, gui_log=gui_log)
    backend = legacy_backend()
    doc = _open_docx_safe(out_abs)
    prepared_func_cache: dict[tuple[str, str], tuple[list[dict[str, Any]], Any]] = {}
    changed = False
    for item in items:
        if item.status != "safe":
            continue
        if stop_event is not None and stop_event.is_set():
            break
        source = os.path.join(new_code, item.rel_path)
        try:
            kwargs = {
                "project_root": new_code,
                "doc": doc,
                "save": False,
            }
            if _supports_kwarg(regenerate_csu_in_doc, "prepared_func_cache"):
                kwargs["prepared_func_cache"] = prepared_func_cache
            result = regenerate_csu_in_doc(
                out_abs,
                source,
                item.func_name,
                item.csu_id,
                cfg,
                **kwargs,
            )
            item.result = result
            if result.get("found"):
                item.status = "applied"
                changed = True
            else:
                item.status = "failed"
                item.reason = "CSU not found during apply"
        except Exception as exc:
            item.status = "failed"
            item.reason = str(exc)
    if changed:
        backend.safe_save_docx(doc, out_abs)


def _decision_matches_item(item: PlannedItem, decision: dict[str, Any]) -> bool:
    rel_path = str(decision.get("rel_path") or "").strip()
    func_name = str(decision.get("func_name") or "").strip()
    action = str(decision.get("action") or "").strip()
    csu_id = str(decision.get("csu_id") or "").strip()
    if not csu_id and str(decision.get("decision") or "").strip() != "insert_after_csu":
        csu_id = str(decision.get("target_csu_id") or "").strip()
    return (
        (not rel_path or item.rel_path == rel_path)
        and (not func_name or item.func_name == func_name)
        and (not action or item.action == action)
        and (not csu_id or item.csu_id == csu_id)
    )


def _item_for_decision(items: list[PlannedItem], decision: dict[str, Any]) -> PlannedItem | None:
    raw_index = decision.get("item_index")
    try:
        index = int(raw_index)
    except Exception:
        index = -1
    if 0 <= index < len(items) and _decision_matches_item(items[index], decision):
        return items[index]

    matches = [
        item for item in items
        if _decision_matches_item(item, decision)
    ]
    return matches[0] if len(matches) == 1 else None


def apply_review_decisions(
    items: list[PlannedItem],
    *,
    out_doc: str,
    new_code: str,
    review_decisions: list[dict[str, Any]],
    ai_assist: bool,
    template_path: str,
    renumber_module_csu: bool = False,
    verbose: bool = False,
    stop_event: Any = None,
    settings: Any = None,
    gui_log: Any = None,
) -> None:
    root = Path(__file__).resolve().parents[1]
    if str(root) not in sys.path:
        sys.path.insert(0, str(root))
    from autodoc.pipeline import insert_csu_after_in_doc, regenerate_csu_in_doc
    from autodoc import render as render_module
    from autodoc._legacy_support import legacy_backend

    cfg = _make_cfg(ai_assist=ai_assist, template_path=template_path, verbose=verbose,
                    stop_event=stop_event, settings=settings, gui_log=gui_log)
    backend = legacy_backend()
    prepared_func_cache: dict[tuple[str, str], tuple[list[dict[str, Any]], Any]] = {}
    for decision in review_decisions:
        if stop_event is not None and stop_event.is_set():
            break
        item = _item_for_decision(items, decision)
        if item is None:
            continue

        action = str(decision.get("decision") or "").strip()
        item.result = {
            **(item.result or {}),
            "review_decision": decision,
        }

        if action == "skip":
            if item.status not in {"applied", "applied_review"}:
                item.status = "skipped_review"
                item.reason = "review decision: skip"
            continue

        if action == "manual":
            if item.status not in {"applied", "applied_review"}:
                item.status = "manual_review"
                item.reason = "review decision: manual"
            continue

        if action == "delete_csu":
            target_csu_id = str(decision.get("target_csu_id") or item.csu_id or "").strip()
            if item.status == "applied":
                item.result["review_result"] = {
                    "skipped": True,
                    "reason": "item was already applied as safe",
                }
                continue
            if not target_csu_id:
                item.status = "failed_review"
                item.reason = "delete_csu requires target_csu_id"
                continue
            try:
                doc = _open_docx_safe(out_doc)
                module_match = re.match(r"^(.+?)_\d+$", target_csu_id)
                module_id = module_match.group(1) if module_match else ""
                result = render_module.delete_csu_in_doc(
                    doc,
                    target_csu_id,
                    backend_module=backend,
                )
                item.result["review_result"] = result
                item.csu_id = target_csu_id
                if result.get("found"):
                    if module_id:
                        if renumber_module_csu:
                            result["renumber_module_csu"] = render_module.renumber_module_csu_ids(
                                doc,
                                module_id,
                                backend_module=backend,
                            )
                        else:
                            result["module_table"] = render_module.sync_module_function_table_for_module(
                                doc,
                                module_id,
                                backend_module=backend,
                            )
                    backend.safe_save_docx(doc, out_doc)
                    item.status = "applied_review"
                    item.reason = "review decision: delete_csu"
                else:
                    item.status = "failed_review"
                    item.reason = "target CSU not found during review delete"
            except Exception as exc:
                item.status = "failed_review"
                item.reason = str(exc)
            continue

        if action == "insert_after_csu":
            target_csu_id = str(decision.get("target_csu_id") or item.csu_id or "").strip()
            after_csu_id = str(decision.get("insert_after_csu_id") or "").strip()
            func_name = str(decision.get("func_name") or item.func_name or "").strip()
            rel_path = str(decision.get("rel_path") or item.rel_path or "").strip()
            source = os.path.join(new_code, rel_path)
            auto_allocated = False

            if item.status == "applied":
                item.result["review_result"] = {
                    "skipped": True,
                    "reason": "item was already applied as safe",
                }
                continue
            if not after_csu_id:
                item.status = "failed_review"
                item.reason = "insert_after_csu requires insert_after_csu_id"
                continue
            if not target_csu_id:
                target_csu_id = allocate_next_csu_id(out_doc, after_csu_id)
                auto_allocated = bool(target_csu_id)
            if not target_csu_id:
                item.status = "failed_review"
                item.reason = "insert_after_csu could not allocate target_csu_id"
                continue
            if not func_name:
                item.status = "failed_review"
                item.reason = "insert_after_csu requires func_name"
                continue
            if not rel_path or not os.path.isfile(source):
                item.status = "failed_review"
                item.reason = "insert_after_csu source file missing"
                continue

            try:
                kwargs = {"project_root": new_code}
                if _supports_kwarg(insert_csu_after_in_doc, "prepared_func_cache"):
                    kwargs["prepared_func_cache"] = prepared_func_cache
                result = insert_csu_after_in_doc(
                    out_doc,
                    source,
                    func_name,
                    target_csu_id,
                    after_csu_id,
                    cfg,
                    **kwargs,
                )
                if auto_allocated:
                    result["auto_allocated_csu_id"] = target_csu_id
                if result.get("found") and renumber_module_csu:
                    module_match = re.match(r"^(.+?)_\d+$", target_csu_id)
                    module_id = module_match.group(1) if module_match else ""
                    if module_id:
                        doc = _open_docx_safe(out_doc)
                        renumber_result = render_module.renumber_module_csu_ids(
                            doc,
                            module_id,
                            backend_module=backend,
                        )
                        result["renumber_module_csu"] = renumber_result
                        for entry in renumber_result.get("mapping") or []:
                            if entry.get("old_csu_id") == target_csu_id:
                                target_csu_id = entry.get("new_csu_id") or target_csu_id
                                break
                        backend.safe_save_docx(doc, out_doc)
                item.result["review_result"] = result
                item.csu_id = target_csu_id
                if result.get("found"):
                    item.status = "applied_review"
                    item.reason = "review decision: insert_after_csu"
                else:
                    item.status = "failed_review"
                    item.reason = "anchor CSU not found during review insert"
            except Exception as exc:
                item.status = "failed_review"
                item.reason = str(exc)
            continue

        if action != "replace_csu":
            if item.status not in {"applied", "applied_review"}:
                item.status = "failed_review"
                item.reason = f"unsupported review decision: {action}"
            continue

        target_csu_id = str(decision.get("target_csu_id") or item.csu_id or "").strip()
        func_name = str(decision.get("func_name") or item.func_name or "").strip()
        rel_path = str(decision.get("rel_path") or item.rel_path or "").strip()
        source = os.path.join(new_code, rel_path)

        if item.status == "applied" and target_csu_id == item.csu_id and func_name == item.func_name:
            item.result["review_result"] = {"skipped": True, "reason": "already applied as safe"}
            continue

        if not target_csu_id:
            item.status = "failed_review"
            item.reason = "replace_csu requires target_csu_id"
            continue
        if not func_name:
            item.status = "failed_review"
            item.reason = "replace_csu requires func_name"
            continue
        if not rel_path or not os.path.isfile(source):
            item.status = "failed_review"
            item.reason = "replace_csu source file missing"
            continue

        try:
            kwargs = {"project_root": new_code}
            if _supports_kwarg(regenerate_csu_in_doc, "prepared_func_cache"):
                kwargs["prepared_func_cache"] = prepared_func_cache
            result = regenerate_csu_in_doc(
                out_doc,
                source,
                func_name,
                target_csu_id,
                cfg,
                **kwargs,
            )
            item.result["review_result"] = result
            item.csu_id = target_csu_id
            if result.get("found"):
                item.status = "applied_review"
                item.reason = "review decision: replace_csu"
            else:
                item.status = "failed_review"
                item.reason = "target CSU not found during review apply"
        except Exception as exc:
            item.status = "failed_review"
            item.reason = str(exc)


def write_reports(
    *,
    items: list[PlannedItem],
    plan_path: str,
    report_path: str,
    metadata: dict[str, Any],
    alignment_index: dict[str, dict[str, Any]] | None = None,
) -> None:
    alignment_items = sorted(
        (alignment_index or {}).values(),
        key=lambda item: str(item.get("csu_id") or ""),
    )
    payload = {
        "schema_version": 1,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "metadata": metadata,
        "summary": {
            status: sum(1 for item in items if item.status == status)
            for status in sorted({item.status for item in items})
        },
        "items": [item.to_dict() for item in items],
        "alignment_items": alignment_items,
    }
    os.makedirs(os.path.dirname(_abs(plan_path)) or ".", exist_ok=True)
    with open(plan_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    lines = [
        "# AutoDocGen Document Update Report",
        "",
        f"- generated_at: {payload['generated_at']}",
        f"- mode: {metadata.get('mode', '')}",
        f"- old_code: {metadata.get('old_code', '')}",
        f"- new_code: {metadata.get('new_code', '')}",
        f"- old_doc: {metadata.get('old_doc', '')}",
        f"- out_doc: {metadata.get('out_doc', '')}",
        "",
        "## Summary",
        "",
    ]
    for key, value in payload["summary"].items():
        lines.append(f"- {key}: {value}")
    lines.extend(["", "## Items", ""])
    for item in items:
        label = item.func_name or item.rel_path
        lines.append(f"- [{item.status}] {item.action}: {label}")
        if item.csu_id:
            lines.append(f"  - csu_id: {item.csu_id}")
        if item.alignment:
            alignment = item.alignment
            lines.append(
                "  - alignment: "
                f"{alignment.get('status', '')} "
                f"{alignment.get('matched_function', '')} "
                f"{alignment.get('rel_path', '')} "
                f"confidence={alignment.get('confidence', '')}"
            )
        if item.reason:
            lines.append(f"  - reason: {item.reason}")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")


def write_review_html(*, plan_path: str, review_html_path: str) -> None:
    script_dir = Path(__file__).resolve().parent
    if str(script_dir) not in sys.path:
        sys.path.insert(0, str(script_dir))
    from render_update_review_html import render_review_html

    with open(plan_path, "r", encoding="utf-8") as f:
        plan = json.load(f)
    render_review_html(plan, review_html_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--old-code", required=True)
    parser.add_argument("--new-code", required=True)
    parser.add_argument("--old-doc", required=True)
    parser.add_argument("--out", required=True)
    parser.add_argument("--docdiff-root", default=_default_docdiff_root())
    parser.add_argument("--mode", choices=["plan-only", "apply-safe", "apply-review"], default="plan-only")
    parser.add_argument("--change-docx", default="")
    parser.add_argument("--change-json", default="")
    parser.add_argument("--plan-out", default="")
    parser.add_argument("--report-out", default="")
    parser.add_argument("--review-html", default="")
    parser.add_argument("--review-decisions", default="")
    parser.add_argument("--alignment-decisions", default="")
    parser.add_argument(
        "--renumber-module-csu",
        action="store_true",
        help="apply-review 时按模块内 H4 顺序重排 CSU 编号；默认关闭以保留既有编号",
    )
    parser.add_argument("--no-review-html", action="store_true")
    parser.add_argument("--template", default="")
    parser.add_argument("--ai-assist", action="store_true")
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="输出 AutoDocGen 详细解析日志；默认仅输出增量工具摘要",
    )
    parser.add_argument(
        "--skip-safe-apply",
        action="store_true",
        help="仅 apply-review 使用：跳过 safe CSU 自动替换，只应用 review-decisions",
    )
    return parser


def main() -> int:
    args = build_parser().parse_args()
    old_code = _abs(args.old_code)
    new_code = _abs(args.new_code)
    old_doc = _abs(args.old_doc)
    out_doc = _abs(args.out)
    docdiff_root = _abs(args.docdiff_root)

    change_docx = _abs(args.change_docx or _default_sidecar(out_doc, ".code_change.docx"))
    change_json = _abs(args.change_json or _default_sidecar(out_doc, ".code_changes.json"))
    plan_out = _abs(args.plan_out or _default_sidecar(out_doc, ".update_plan.json"))
    report_out = _abs(args.report_out or _default_sidecar(out_doc, ".update_report.md"))
    review_html = _abs(args.review_html or _default_sidecar(out_doc, ".update_review.html"))
    review_decisions_path = _abs(args.review_decisions) if args.review_decisions else ""
    alignment_decisions_path = _abs(args.alignment_decisions) if args.alignment_decisions else ""

    _run_docdiff(
        docdiff_root=docdiff_root,
        old_code=old_code,
        new_code=new_code,
        change_docx=change_docx,
        change_json=change_json,
    )
    changes = _load_changes(change_json)
    csu_index = build_csu_index(old_doc)
    alignment_index = build_doc_code_alignment_index(old_doc, old_code)
    alignment_decisions: list[dict[str, Any]] = []
    for path in [review_decisions_path, alignment_decisions_path]:
        if path:
            alignment_decisions.extend(_load_alignment_decisions(path))
    apply_alignment_decisions(
        alignment_index,
        csu_index,
        alignment_decisions,
        code_roots=[old_code, new_code],
    )
    items = classify_changes(changes, new_code=new_code, csu_index=csu_index)
    attach_alignment_to_items(items, alignment_index)

    if args.mode == "apply-safe":
        apply_safe_items(
            items,
            old_doc=old_doc,
            out_doc=out_doc,
            new_code=new_code,
            ai_assist=bool(args.ai_assist),
            template_path=args.template,
            verbose=bool(args.verbose),
        )
    elif args.mode == "apply-review":
        if not review_decisions_path:
            raise ValueError("--mode apply-review requires --review-decisions")
        if args.skip_safe_apply:
            _copy_old_doc_to_out(old_doc=old_doc, out_doc=out_doc, label="apply-review")
        else:
            apply_safe_items(
                items,
                old_doc=old_doc,
                out_doc=out_doc,
                new_code=new_code,
                ai_assist=bool(args.ai_assist),
                template_path=args.template,
                verbose=bool(args.verbose),
            )
        decisions = _load_review_decisions(review_decisions_path)
        apply_review_decisions(
            items,
            out_doc=out_doc,
            new_code=new_code,
            review_decisions=decisions,
            ai_assist=bool(args.ai_assist),
            template_path=args.template,
            renumber_module_csu=bool(args.renumber_module_csu),
            verbose=bool(args.verbose),
        )

    metadata = {
        "mode": args.mode,
        "old_code": old_code,
        "new_code": new_code,
        "old_doc": old_doc,
        "out_doc": out_doc,
        "docdiff_root": docdiff_root,
        "change_docx": change_docx,
        "change_json": change_json,
        "review_decisions": review_decisions_path,
        "alignment_decisions": alignment_decisions_path,
        "renumber_module_csu": bool(args.renumber_module_csu),
        "alignment": {
            "schema_version": DOC_CODE_ALIGNMENT_SCHEMA,
            "code_root": old_code,
            "total_csu": len(alignment_index),
            "matched_high": sum(1 for item in alignment_index.values() if item.get("status") == "matched_high"),
            "manual_matched": sum(1 for item in alignment_index.values() if item.get("status") == "manual_matched"),
            "ambiguous": sum(1 for item in alignment_index.values() if item.get("status") == "ambiguous"),
            "unmatched": sum(1 for item in alignment_index.values() if item.get("status") in {"unmatched", "no_doc_function"}),
        },
    }
    write_reports(
        items=items,
        plan_path=plan_out,
        report_path=report_out,
        metadata=metadata,
        alignment_index=alignment_index,
    )
    if not args.no_review_html:
        write_review_html(plan_path=plan_out, review_html_path=review_html)
    print(f"plan: {plan_out}")
    print(f"report: {report_out}")
    if not args.no_review_html:
        print(f"review_html: {review_html}")
    print(f"code_change_docx: {change_docx}")
    print(f"code_change_json: {change_json}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
